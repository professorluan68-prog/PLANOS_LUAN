import json
import os
import sqlite3
import logging
import re
import hashlib
import tempfile
import unicodedata
import uuid
from io import BytesIO
from pathlib import Path
from contextlib import contextmanager
from datetime import datetime

from config import DB_PATH, REGISTRO_PROXIMA_GERACAO_PATH, HISTORICO_DOCX_DIR, PLANOS_FEITOS_DIR
from core.lib.classificador import normalizar_texto

logger = logging.getLogger(__name__)


HISTORICO_PLANOS_LIMITE_PADRAO = 50
HISTORICO_PLANOS_LIMITE_MAXIMO = 500

_MESES_HISTORICO = {
    "JANEIRO": "01",
    "FEVEREIRO": "02",
    "MARCO": "03",
    "MARÇO": "03",
    "ABRIL": "04",
    "MAIO": "05",
    "JUNHO": "06",
    "JULHO": "07",
    "AGOSTO": "08",
    "SETEMBRO": "09",
    "OUTUBRO": "10",
    "NOVEMBRO": "11",
    "DEZEMBRO": "12",
}


class SafeConnectionWrapper:
    """
    Wrapper para sqlite3.Connection que garante o fechamento automático da
    conexão ao sair do bloco 'with'.
    """
    def __init__(self, conn):
        self.conn = conn

    def __getattr__(self, name):
        return getattr(self.conn, name)

    def __enter__(self):
        self.conn.__enter__()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        try:
            self.conn.__exit__(exc_type, exc_val, exc_tb)
        finally:
            self.conn.close()


def get_connection(db_path=None):
    """
    Retorna uma nova conexão SQLite configurada para concorrência (WAL).
    Cada worker/thread deve abrir sua própria conexão via esta função.
    """
    db_path = DB_PATH if db_path is None else db_path
    conn = sqlite3.connect(db_path, timeout=30, check_same_thread=False)
    # Pragmas aplicadas por conexão para garantir comportamento consistente
    conn.execute("PRAGMA journal_mode=WAL;")
    conn.execute("PRAGMA synchronous=NORMAL;")
    conn.execute("PRAGMA foreign_keys=ON;")
    conn.execute("PRAGMA busy_timeout=10000;")
    return SafeConnectionWrapper(conn)


@contextmanager
def connection_scope(db_path=None):
    """
    Context manager para uso de conexão com commit/rollback automático.
    Uso recomendado: with connection_scope() as conn: ...
    """
    conn = get_connection(db_path)
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()

def _normalizar_campo(valor):
    return str(valor or "").strip()


def _normalizar_campo_chave(valor):
    texto = _normalizar_campo(valor).replace("_", " ")
    if not texto:
        return ""
    texto = normalizar_texto(texto)
    texto = re.sub(r"[^a-z0-9]+", " ", texto, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", texto).strip().upper()


def _normalizar_turma_historico_chave(valor: str) -> str:
    chave = _normalizar_campo_chave(valor)
    chave = re.sub(r"\b([1-9])\s*[OA]\s+(ANO|SERIE|TERMO)\b", r"\1 \2", chave)
    return re.sub(r"\s+", " ", chave).strip()


def _normalizar_bimestre_historico_chave(valor: str) -> str:
    chave = _normalizar_campo_chave(valor)
    chave = re.sub(r"\b([1-4])\s*[OA]?\s*BIMESTRE\b", r"\1 BIMESTRE", chave)
    return re.sub(r"\s+", " ", chave).strip()


def _mes_geracao_historico(data_geracao: str = "") -> str:
    texto = str(data_geracao or "").strip()
    if re.match(r"^\d{4}-\d{2}", texto):
        return texto[:7]
    return ""


def _normalizar_mes_plano_historico(valor: str = "", data_geracao: str = "") -> str:
    texto = _normalizar_campo(valor)
    if not texto:
        return ""
    if re.match(r"^\d{4}-\d{2}$", texto):
        return texto

    ano = _mes_geracao_historico(data_geracao)[:4]
    if not ano:
        ano_match = re.search(r"\b(20\d{2})\b", texto)
        ano = ano_match.group(1) if ano_match else ""

    texto_chave = _normalizar_campo_chave(texto)
    mes_num = ""
    match_num = re.search(r"(?<!\d)(0?[1-9]|1[0-2])(?!\d)", texto_chave)
    if match_num:
        mes_num = f"{int(match_num.group(1)):02d}"
    else:
        for nome_mes, numero_mes in _MESES_HISTORICO.items():
            if nome_mes in texto_chave:
                mes_num = numero_mes
                break

    if not (ano and mes_num):
        return ""
    return f"{ano}-{mes_num}"


def _extrair_contexto_docx_historico(
    arquivo_docx_bytes: bytes | None,
    data_geracao: str = "",
) -> dict[str, str]:
    if not arquivo_docx_bytes:
        return {"bimestre": "", "mes_plano": ""}

    try:
        from docx import Document

        documento = Document(BytesIO(arquivo_docx_bytes))
    except Exception:
        return {"bimestre": "", "mes_plano": ""}

    bimestre = ""
    mes_plano = ""

    for tabela in documento.tables:
        linhas = tabela.rows
        for indice, linha in enumerate(linhas[:-1]):
            rotulos = [_normalizar_campo_chave(celula.text) for celula in linha.cells]
            valores = [_normalizar_campo(celula.text) for celula in linhas[indice + 1].cells]
            for posicao, rotulo in enumerate(rotulos):
                valor = valores[posicao] if posicao < len(valores) else ""
                if not mes_plano and rotulo == "MES":
                    mes_plano = _normalizar_mes_plano_historico(valor, data_geracao)
                if not bimestre and rotulo == "BIMESTRE":
                    bimestre = valor
            if bimestre and mes_plano:
                return {"bimestre": bimestre, "mes_plano": mes_plano}

    textos = []
    for paragrafo in documento.paragraphs:
        texto = _normalizar_campo(paragrafo.text)
        if texto:
            textos.append(texto)
    texto_completo = "\n".join(textos)
    if not bimestre:
        match_bimestre = re.search(
            r"\b([1-4]\s*[ºoO]?\s*Bimestre)\b",
            texto_completo,
            flags=re.I,
        )
        if match_bimestre:
            bimestre = _normalizar_campo(match_bimestre.group(1))
    if not mes_plano:
        match_mes = re.search(
            r"\b(Janeiro|Fevereiro|Mar[cç]o|Abril|Maio|Junho|Julho|Agosto|Setembro|Outubro|Novembro|Dezembro)\b",
            texto_completo,
            flags=re.I,
        )
        if match_mes:
            mes_plano = _normalizar_mes_plano_historico(match_mes.group(1), data_geracao)

    return {"bimestre": bimestre, "mes_plano": mes_plano}


def _inferir_origem_historico(arquivo_path: str = "") -> str:
    caminho = Path(str(arquivo_path or "").strip())
    if not str(caminho):
        return ""
    if not caminho.is_absolute():
        return "historico_docx"

    try:
        caminho_resolvido = caminho.resolve(strict=False)
        caminho_resolvido.relative_to(Path(PLANOS_FEITOS_DIR).resolve(strict=False))
        return "planos_feitos"
    except ValueError:
        return "externo"
    except OSError:
        return "externo"


def _metadados_historico(
    professor_nome: str = "",
    disciplina: str = "",
    turma: str = "",
    bimestre: str = "",
    mes_plano: str = "",
    data_geracao: str = "",
    arquivo_path: str = "",
    arquivo_docx_bytes: bytes | None = None,
    extrair_resumo_aulas: bool = False,
) -> dict[str, object]:
    arquivo_tamanho = None
    arquivo_hash = ""
    ultima_aula = None
    total_aulas = None
    arquivo_bytes = arquivo_docx_bytes
    if arquivo_docx_bytes is not None:
        arquivo_tamanho = len(arquivo_docx_bytes)
        arquivo_hash = hashlib.sha256(arquivo_docx_bytes).hexdigest()
        if extrair_resumo_aulas:
            ultima_aula, total_aulas = _extrair_resumo_aulas_historico(
                arquivo_docx_bytes,
                bimestre,
            )
    elif arquivo_path:
        caminho = _resolver_caminho_arquivo_historico(arquivo_path)
        try:
            arquivo_tamanho = caminho.stat().st_size if caminho.exists() else None
        except OSError:
            arquivo_tamanho = None
        if extrair_resumo_aulas and arquivo_tamanho:
            try:
                arquivo_bytes = caminho.read_bytes()
                ultima_aula, total_aulas = _extrair_resumo_aulas_historico(arquivo_bytes, bimestre)
            except OSError:
                ultima_aula, total_aulas = None, None

    bimestre_final = _normalizar_campo(bimestre)
    mes_plano_final = _normalizar_mes_plano_historico(mes_plano, data_geracao)
    if arquivo_bytes and (not bimestre_final or not mes_plano_final):
        contexto_docx = _extrair_contexto_docx_historico(arquivo_bytes, data_geracao)
        if not bimestre_final:
            bimestre_final = _normalizar_campo(contexto_docx.get("bimestre"))
        if not mes_plano_final:
            mes_plano_final = _normalizar_mes_plano_historico(
                contexto_docx.get("mes_plano", ""),
                data_geracao,
            )

    return {
        "bimestre": bimestre_final,
        "professor_chave": _normalizar_campo_chave(professor_nome),
        "disciplina_chave": _normalizar_campo_chave(disciplina),
        "turma_chave": _normalizar_turma_historico_chave(turma),
        "bimestre_chave": _normalizar_bimestre_historico_chave(bimestre_final),
        "mes_geracao": _mes_geracao_historico(data_geracao),
        "mes_plano": mes_plano_final,
        "arquivo_hash": arquivo_hash,
        "arquivo_tamanho": arquivo_tamanho,
        "origem": _inferir_origem_historico(arquivo_path),
        "ultima_aula": ultima_aula,
        "total_aulas": total_aulas,
    }


def _extrair_resumo_aulas_historico(
    arquivo_docx_bytes: bytes,
    bimestre: str = "",
) -> tuple[int | None, int | None]:
    if not arquivo_docx_bytes:
        return None, None
    try:
        from core.gestao_aulas import detectar_resumo_aulas_de_docx_bytes

        resumo = detectar_resumo_aulas_de_docx_bytes(arquivo_docx_bytes, bimestre)
        ultima_aula = int(resumo.get("ultima_aula") or 0)
        total_aulas = int(resumo.get("total_aulas") or 0)
        return ultima_aula, total_aulas
    except Exception as exc:
        logger.debug("Não foi possível extrair resumo de aulas do histórico: %s", exc)
        return 0, 0


def _normalizar_chave_vinculo(disciplina: str, turma: str, componente_curricular: str = "") -> tuple[str, str, str]:
    disciplina_chave = _normalizar_campo_chave(disciplina)
    turma_chave = _normalizar_campo_chave(turma)
    componente_base = componente_curricular if _normalizar_campo(componente_curricular) else disciplina
    componente_chave = _normalizar_campo_chave(componente_base)
    return disciplina_chave, turma_chave, componente_chave


def _buscar_vinculo_existente_equivalente(
    cursor,
    professor_id: int,
    disciplina: str,
    turma: str,
    componente_curricular: str = "",
    ignorar_id: int | None = None,
) -> int | None:
    disciplina_chave, turma_chave, componente_chave = _normalizar_chave_vinculo(
        disciplina,
        turma,
        componente_curricular,
    )
    cursor.execute(
        """
        SELECT id, disciplina, turma, COALESCE(componente_curricular, '')
        FROM professor_turmas
        WHERE professor_id = ?
        ORDER BY id
        """,
        (professor_id,),
    )
    for row in cursor.fetchall():
        vinculo_id = int(row[0])
        if ignorar_id is not None and vinculo_id == ignorar_id:
            continue
        atual_disciplina, atual_turma, atual_componente = _normalizar_chave_vinculo(
            row[1],
            row[2],
            row[3],
        )
        if (
            atual_disciplina == disciplina_chave
            and atual_turma == turma_chave
            and atual_componente == componente_chave
        ):
            return vinculo_id
    return None


def _resolver_caminho_arquivo_historico(arquivo_path: str) -> Path:
    caminho = Path(str(arquivo_path or "").strip())
    if not str(caminho):
        return Path()
    return caminho if caminho.is_absolute() else Path(HISTORICO_DOCX_DIR) / caminho


def _caminho_historico_gerenciado(arquivo_path: str) -> Path | None:
    """Resolve somente arquivos relativos mantidos pelo diretório de histórico."""
    caminho = Path(str(arquivo_path or "").strip())
    if not str(caminho) or caminho.is_absolute():
        return None
    raiz = Path(HISTORICO_DOCX_DIR).resolve(strict=False)
    candidato = (raiz / caminho).resolve(strict=False)
    try:
        candidato.relative_to(raiz)
    except ValueError:
        return None
    return candidato


def _gravar_arquivo_historico_atomico(destino: Path, conteudo: bytes) -> None:
    destino.parent.mkdir(parents=True, exist_ok=True)
    temporario: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="wb",
            dir=destino.parent,
            prefix=f".{destino.stem}.",
            suffix=".tmp",
            delete=False,
        ) as stream:
            temporario = Path(stream.name)
            stream.write(conteudo)
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporario, destino)
        if destino.stat().st_size != len(conteudo):
            destino.unlink(missing_ok=True)
            raise OSError(f"Tamanho divergente após gravar {destino.name}")
    finally:
        if temporario is not None and temporario.exists():
            temporario.unlink(missing_ok=True)


def _remover_arquivo_historico_gerenciado(arquivo_path: str) -> bool:
    caminho = _caminho_historico_gerenciado(arquivo_path)
    if caminho is None:
        return False
    try:
        caminho.unlink(missing_ok=True)
    except OSError as exc:
        logger.error("Erro ao remover arquivo físico do histórico %s: %s", caminho, exc)
        return False
    return True


def _chave_caminho_historico(arquivo_path: str) -> str:
    caminho = _resolver_caminho_arquivo_historico(arquivo_path)
    if not str(caminho):
        return ""
    try:
        return str(caminho.resolve()).upper()
    except OSError:
        return str(caminho).upper()


def _normalizar_nome_pasta_historico(valor: str) -> str:
    texto = str(valor or "").replace("_", " ").strip()
    return re.sub(r"\s+", " ", normalizar_texto(texto)).strip().upper()


def _slug_nome_arquivo_historico(texto: str) -> str:
    texto = str(texto or "").replace("º", "o").replace("°", "o").replace("ª", "a")
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    return re.sub(r"[^A-Za-z0-9_-]+", "_", texto).strip("_")


def _inferir_turma_por_nome_arquivo(arquivo_nome: str, disciplina: str) -> str:
    stem = Path(str(arquivo_nome or "")).stem
    if stem.lower().startswith("plano_"):
        stem = stem[6:]
    if stem.lower().endswith("_in"):
        stem = stem[:-3]

    disciplina_slug = _slug_nome_arquivo_historico(disciplina)
    sufixo_disciplina = f"_{disciplina_slug.lower()}" if disciplina_slug else ""
    if sufixo_disciplina and stem.lower().endswith(sufixo_disciplina):
        stem = stem[: -len(sufixo_disciplina)]

    turma = re.sub(r"\s+", " ", stem.replace("_", " ")).strip()
    return turma.upper()


def _resolver_nome_professor_por_pasta(nome_pasta: str, nomes_professores: list[str]) -> str:
    chave_pasta = _normalizar_nome_pasta_historico(nome_pasta)
    for nome in nomes_professores:
        if _normalizar_nome_pasta_historico(nome) == chave_pasta:
            return str(nome or "").strip()
    return re.sub(r"\s+", " ", str(nome_pasta or "").replace("_", " ")).strip().upper()


def _buscar_id_historico_por_caminho(cursor, arquivo_path: str) -> int | None:
    alvo = _chave_caminho_historico(arquivo_path)
    if not alvo:
        return None

    cursor.execute(
        """
        SELECT id, arquivo_path
        FROM historico_planos
        WHERE COALESCE(TRIM(arquivo_path), '') <> ''
        """
    )
    for row in cursor.fetchall():
        if _chave_caminho_historico(row[1]) == alvo:
            return int(row[0])
    return None


def _existe_historico_mesmo_contexto_arquivo(
    cursor,
    professor_nome: str,
    disciplina: str,
    turma: str,
    arquivo_nome: str,
) -> bool:
    return _buscar_id_historico_mesmo_contexto_arquivo(
        cursor,
        professor_nome,
        disciplina,
        turma,
        arquivo_nome,
    ) is not None


def _buscar_id_historico_mesmo_contexto_arquivo(
    cursor,
    professor_nome: str,
    disciplina: str,
    turma: str,
    arquivo_nome: str,
) -> int | None:
    cursor.execute(
        """
        SELECT id
        FROM historico_planos
        WHERE UPPER(TRIM(professor_nome)) = UPPER(TRIM(?))
          AND UPPER(TRIM(disciplina)) = UPPER(TRIM(?))
          AND UPPER(TRIM(turma)) = UPPER(TRIM(?))
          AND UPPER(TRIM(arquivo_nome)) = UPPER(TRIM(?))
        LIMIT 1
        """,
        (professor_nome, disciplina, turma, arquivo_nome),
    )
    row = cursor.fetchone()
    return int(row[0]) if row else None


def sincronizar_historico_planos_com_planos_feitos() -> int:
    pasta_planos = Path(PLANOS_FEITOS_DIR)
    if not pasta_planos.exists():
        return 0

    # Fase 1: I/O de disco
    arquivos_docx = [
        caminho
        for caminho in pasta_planos.rglob("*.docx")
        if caminho.is_file() and not caminho.name.startswith("~$")
    ]
    if not arquivos_docx:
        return 0

    # Fase 2: Carregar dados em memória
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT nome FROM professores ORDER BY nome")
        nomes_professores = [row[0] for row in cursor.fetchall()]

        cursor.execute("SELECT arquivo_path FROM historico_planos WHERE arquivo_path IS NOT NULL")
        paths_banco = cursor.fetchall()

    paths_existentes = set()
    for row in paths_banco:
        chave = _chave_caminho_historico(row[0])
        if chave:
            paths_existentes.add(chave)

    # Fase 3: Processamento em memória
    dados_para_inserir = []
    for caminho in arquivos_docx:
        caminho_str = str(caminho.resolve())
        chave = _chave_caminho_historico(caminho_str)
        if chave in paths_existentes:
            continue

        try:
            relativo = caminho.relative_to(pasta_planos)
        except ValueError:
            continue
        partes = relativo.parts
        if len(partes) < 3:
            continue

        professor_nome = _resolver_nome_professor_por_pasta(partes[0], nomes_professores)
        disciplina = str(partes[1] or "").replace("_", " ").strip()
        arquivo_nome = caminho.name
        turma = _inferir_turma_por_nome_arquivo(arquivo_nome, disciplina)
        if not professor_nome or not disciplina or not turma:
            continue

        data_geracao = datetime.fromtimestamp(caminho.stat().st_mtime).strftime(
            "%Y-%m-%d %H:%M:%S"
        )
        metadados = _metadados_historico(
            professor_nome=professor_nome,
            disciplina=disciplina,
            turma=turma,
            bimestre="",
            data_geracao=data_geracao,
            arquivo_path=caminho_str,
            extrair_resumo_aulas=True,
        )
        dados_para_inserir.append((
            professor_nome,
            disciplina,
            turma,
            metadados["bimestre"],
            data_geracao,
            arquivo_nome,
            caminho_str,
            metadados["professor_chave"],
            metadados["disciplina_chave"],
            metadados["turma_chave"],
            metadados["bimestre_chave"],
            metadados["mes_geracao"],
            metadados["mes_plano"],
            metadados["arquivo_hash"],
            metadados["arquivo_tamanho"],
            metadados["origem"],
            metadados["ultima_aula"],
            metadados["total_aulas"],
        ))

    if not dados_para_inserir:
        return 0

    # Fase 4: Transação de Inserção curta
    inseridos = 0
    with get_connection() as conn:
        cursor = conn.cursor()
        for dados in dados_para_inserir:
            (
                prof,
                disc,
                turma,
                bim,
                dt,
                arq,
                path,
                professor_chave,
                disciplina_chave,
                turma_chave,
                bimestre_chave,
                mes_geracao,
                mes_plano,
                arquivo_hash,
                arquivo_tamanho,
                origem,
                ultima_aula,
                total_aulas,
            ) = dados
            registro_existente_id = _buscar_id_historico_mesmo_contexto_arquivo(
                cursor,
                prof,
                disc,
                turma,
                arq,
            )
            if registro_existente_id is not None:
                cursor.execute(
                    """
                    UPDATE historico_planos
                    SET bimestre = COALESCE(NULLIF(bimestre, ''), ?),
                        data_geracao = ?,
                        arquivo_path = ?,
                        professor_chave = ?,
                        disciplina_chave = ?,
                        turma_chave = ?,
                        bimestre_chave = COALESCE(NULLIF(bimestre_chave, ''), ?),
                        mes_geracao = ?,
                        mes_plano = COALESCE(NULLIF(mes_plano, ''), ?),
                        arquivo_hash = COALESCE(NULLIF(arquivo_hash, ''), ?),
                        arquivo_tamanho = ?,
                        origem = ?,
                        ultima_aula = COALESCE(ultima_aula, ?),
                        total_aulas = COALESCE(total_aulas, ?)
                    WHERE id = ?
                    """,
                    (
                        bim,
                        dt,
                        path,
                        professor_chave,
                        disciplina_chave,
                        turma_chave,
                        bimestre_chave,
                        mes_geracao,
                        mes_plano,
                        arquivo_hash,
                        arquivo_tamanho,
                        origem,
                        ultima_aula,
                        total_aulas,
                        registro_existente_id,
                    ),
                )
                inseridos += 1
                continue

            cursor.execute(
                """
                INSERT INTO historico_planos
                    (
                        professor_nome,
                        disciplina,
                        turma,
                        bimestre,
                        data_geracao,
                        arquivo_nome,
                        arquivo_path,
                        professor_chave,
                        disciplina_chave,
                        turma_chave,
                        bimestre_chave,
                        mes_geracao,
                        mes_plano,
                        arquivo_hash,
                        arquivo_tamanho,
                        origem,
                        ultima_aula,
                        total_aulas
                    )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                dados,
            )
            inseridos += 1

        conn.commit()
        return inseridos


def _obter_ou_criar_professor(cursor, nome: str) -> int:
    nome = _normalizar_campo(nome).upper()
    cursor.execute("INSERT OR IGNORE INTO professores (nome) VALUES (?)", (nome,))
    cursor.execute("SELECT id FROM professores WHERE nome = ?", (nome,))
    row = cursor.fetchone()
    if not row:
        raise ValueError("Nao foi possivel localizar ou criar o professor.")
    return int(row[0])


def _remover_professor_sem_turmas(cursor, professor_id: int) -> None:
    cursor.execute("SELECT COUNT(*) FROM professor_turmas WHERE professor_id = ?", (professor_id,))
    total = int(cursor.fetchone()[0] or 0)
    cursor.execute(
        """
        SELECT COUNT(*)
        FROM professor_dados
        WHERE professor_id = ?
          AND (
              COALESCE(TRIM(cpf), '') <> ''
              OR COALESCE(TRIM(email), '') <> ''
              OR COALESCE(TRIM(valor_mensal), '') <> ''
              OR COALESCE(TRIM(telefone), '') <> ''
              OR COALESCE(TRIM(observacoes), '') <> ''
          )
        """,
        (professor_id,),
    )
    tem_dados_administrativos = int(cursor.fetchone()[0] or 0) > 0
    if total == 0 and not tem_dados_administrativos:
        cursor.execute("DELETE FROM professores WHERE id = ?", (professor_id,))


def _criar_indices_banco(cursor) -> None:
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_historico_planos_data_id
        ON historico_planos (data_geracao DESC, id DESC)
        """
    )
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_historico_planos_contexto_data
        ON historico_planos (professor_nome, disciplina, turma, data_geracao DESC, id DESC)
        """
    )
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_historico_planos_contexto_bimestre_data
        ON historico_planos (professor_nome, disciplina, turma, bimestre, data_geracao DESC, id DESC)
        """
    )
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_historico_planos_chaves_data
        ON historico_planos
        (professor_chave, disciplina_chave, turma_chave, bimestre_chave, data_geracao DESC, id DESC)
        """
    )
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_historico_planos_contexto_chaves_data
        ON historico_planos
        (professor_chave, disciplina_chave, turma_chave, data_geracao DESC, id DESC)
        """
    )
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_historico_planos_prof_data
        ON historico_planos (professor_chave, data_geracao DESC, id DESC)
        """
    )
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_historico_planos_prof_mes_data
        ON historico_planos (professor_chave, mes_geracao, data_geracao DESC, id DESC)
        """
    )
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_historico_planos_mes_data
        ON historico_planos (mes_geracao DESC, data_geracao DESC, id DESC)
        """
    )
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_historico_planos_mes_plano_data
        ON historico_planos (mes_plano DESC, data_geracao DESC, id DESC)
        """
    )
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_historico_planos_hash
        ON historico_planos (arquivo_hash)
        """
    )
    cursor.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_professor_turmas_prof_id
        ON professor_turmas (professor_id)
        """
    )


def _limpar_historico_planos_incompletos(cursor) -> None:
    cursor.execute(
        """
        DELETE FROM historico_planos
        WHERE arquivo_path IS NULL
           OR COALESCE(TRIM(arquivo_path), '') = ''
           OR COALESCE(TRIM(arquivo_nome), '') = ''
        """
    )


def _normalizar_limite_historico(limite) -> int:
    try:
        valor = int(limite)
    except (TypeError, ValueError):
        valor = HISTORICO_PLANOS_LIMITE_PADRAO
    return max(1, min(valor, HISTORICO_PLANOS_LIMITE_MAXIMO))


def _atualizar_metadados_historico(cursor) -> None:
    cursor.execute(
        """
        SELECT
            id,
            professor_nome,
            disciplina,
            turma,
            COALESCE(bimestre, ''),
            COALESCE(data_geracao, ''),
            COALESCE(arquivo_path, ''),
            COALESCE(arquivo_hash, '')
        FROM historico_planos
        WHERE COALESCE(TRIM(professor_chave), '') = ''
           OR COALESCE(TRIM(disciplina_chave), '') = ''
           OR COALESCE(TRIM(turma_chave), '') = ''
           OR COALESCE(TRIM(bimestre_chave), '') = ''
           OR COALESCE(TRIM(mes_geracao), '') = ''
           OR COALESCE(TRIM(mes_plano), '') = ''
           OR COALESCE(TRIM(bimestre), '') = ''
           OR origem IS NULL
           OR arquivo_tamanho IS NULL
           OR ultima_aula IS NULL
           OR total_aulas IS NULL
        """
    )
    registros = cursor.fetchall()
    for row in registros:
        metadados = _metadados_historico(
            professor_nome=row[1],
            disciplina=row[2],
            turma=row[3],
            bimestre=row[4],
            data_geracao=row[5],
            arquivo_path=row[6],
            extrair_resumo_aulas=True,
        )
        arquivo_hash = row[7] or metadados["arquivo_hash"]
        cursor.execute(
            """
            UPDATE historico_planos
            SET professor_chave = ?,
                disciplina_chave = ?,
                turma_chave = ?,
                bimestre_chave = ?,
                mes_geracao = ?,
                mes_plano = COALESCE(NULLIF(mes_plano, ''), ?),
                bimestre = COALESCE(NULLIF(bimestre, ''), ?),
                origem = ?,
                arquivo_tamanho = COALESCE(arquivo_tamanho, ?),
                arquivo_hash = COALESCE(NULLIF(arquivo_hash, ''), ?),
                ultima_aula = COALESCE(ultima_aula, ?),
                total_aulas = COALESCE(total_aulas, ?)
            WHERE id = ?
            """,
            (
                metadados["professor_chave"],
                metadados["disciplina_chave"],
                metadados["turma_chave"],
                metadados["bimestre_chave"],
                metadados["mes_geracao"],
                metadados["mes_plano"],
                metadados["bimestre"],
                metadados["origem"],
                metadados["arquivo_tamanho"],
                arquivo_hash,
                metadados["ultima_aula"],
                metadados["total_aulas"],
                int(row[0]),
            ),
        )


# ---------- Sistema de migrações versionadas ----------
MIGRACOES = [
    # Versão 1
    "ALTER TABLE professor_turmas ADD COLUMN arquivo_modelo TEXT",
    # Versão 2
    "ALTER TABLE professor_turmas ADD COLUMN template_id TEXT",
    # Versão 3
    "ALTER TABLE professor_turmas ADD COLUMN componente_curricular TEXT",
    # Versão 4
    "ALTER TABLE historico_planos ADD COLUMN bimestre TEXT",
    # Versão 5
    "ALTER TABLE historico_planos ADD COLUMN professor_chave TEXT",
    # Versão 6
    "ALTER TABLE historico_planos ADD COLUMN disciplina_chave TEXT",
    # Versão 7
    "ALTER TABLE historico_planos ADD COLUMN turma_chave TEXT",
    # Versão 8
    "ALTER TABLE historico_planos ADD COLUMN bimestre_chave TEXT",
    # Versão 9
    "ALTER TABLE historico_planos ADD COLUMN mes_geracao TEXT",
    # Versão 10
    "ALTER TABLE historico_planos ADD COLUMN arquivo_hash TEXT",
    # Versão 11
    "ALTER TABLE historico_planos ADD COLUMN arquivo_tamanho INTEGER",
    # Versão 12
    "ALTER TABLE historico_planos ADD COLUMN origem TEXT",
    # Versão 13
    "ALTER TABLE historico_planos ADD COLUMN ultima_aula INTEGER",
    # Versão 14
    "ALTER TABLE historico_planos ADD COLUMN total_aulas INTEGER",
    # Versão 15
    "ALTER TABLE historico_planos ADD COLUMN mes_plano TEXT",
    # Versão 16
    """
    CREATE TABLE IF NOT EXISTS professor_dados (
        professor_id INTEGER PRIMARY KEY,
        cpf TEXT,
        email TEXT,
        valor_mensal TEXT,
        telefone TEXT,
        observacoes TEXT,
        atualizado_em TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(professor_id) REFERENCES professores(id) ON DELETE CASCADE
    )
    """,
]


def _aplicar_migracoes(cursor):
    """Aplica migrações pendentes de schema com controle de versão.

    Política fail-closed: a versão só é registrada após execução bem-sucedida
    do SQL da migração. Erros reais abortam o processo sem avançar a versão,
    impedindo que o banco fique em estado parcialmente migrado.
    """
    cursor.execute(
        "CREATE TABLE IF NOT EXISTS schema_version (versao INTEGER PRIMARY KEY)"
    )
    cursor.execute("SELECT COALESCE(MAX(versao), 0) FROM schema_version")
    versao_atual = cursor.fetchone()[0]

    for i, sql in enumerate(MIGRACOES[versao_atual:], start=versao_atual + 1):
        try:
            cursor.execute(sql)
        except Exception as e:
            if "duplicate column name" in str(e).lower():
                # Coluna já existe: migração é idempotente — avançar versão
                logger.warning(
                    "Migração %d: coluna já existia (idempotente). Avançando versão.", i
                )
            else:
                # Erro real: NÃO registrar versão — manter banco em estado consistente
                logger.error("Migração %d falhou: %s", i, e)
                raise RuntimeError(f"Falha na migração {i}: {e}") from e
        # Registrar versão somente após execução (ou idempotência) confirmada
        cursor.execute("INSERT OR IGNORE INTO schema_version VALUES (?)", (i,))


def _migrar_blob_para_path(conn) -> None:
    """
    Verifica se a tabela historico_planos tem a coluna antiga 'arquivo_docx' (BLOB).
    Se sim, migra os blobs para arquivos físicos em HISTORICO_DOCX_DIR,
    salva a referência do caminho relativo em uma nova tabela e substitui a antiga.

    Estratégia atômica (auditoria P0):
    - Todos os arquivos físicos são gravados ANTES do DROP TABLE;
    - Cada arquivo é verificado por tamanho após a escrita;
    - Qualquer falha aborta TODA a migração, preservando a tabela original;
    - O DROP TABLE só ocorre após todos os registros confirmados.
    """
    cursor = conn.cursor()
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='historico_planos'")
    if not cursor.fetchone():
        return

    cursor.execute("PRAGMA table_info(historico_planos)")
    colunas = [row[1] for row in cursor.fetchall()]
    if "arquivo_docx" not in colunas:
        return

    logger.info("Iniciando migração de BLOBs do SQLite para arquivos físicos...")
    cursor.execute(
        "SELECT id, professor_nome, disciplina, turma, bimestre, data_geracao, arquivo_nome, arquivo_docx "
        "FROM historico_planos"
    )
    registros = cursor.fetchall()

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS historico_planos_nova (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            professor_nome TEXT,
            disciplina TEXT,
            turma TEXT,
            bimestre TEXT,
            data_geracao TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            arquivo_nome TEXT,
            arquivo_path TEXT
        )
        """
    )

    os.makedirs(HISTORICO_DOCX_DIR, exist_ok=True)
    from core.lib.classificador import normalizar_texto

    # --- Fase 1: gravar todos os arquivos físicos e validar ---
    # Só avançamos para o DROP TABLE se TODOS os arquivos forem confirmados.
    arquivos_gravados = []  # lista de (filepath, unique_name, tamanho_esperado)
    registros_nova = []    # dados para inserção na nova tabela

    for reg in registros:
        r_id, prof, disc, turma, bim, data_gen, arq_nome, blob = reg
        if not arq_nome:
            arq_nome = f"plano_{r_id}.docx"

        stem = Path(arq_nome).stem
        ext = Path(arq_nome).suffix or ".docx"

        prof_clean = normalizar_texto(prof or "").replace(" ", "_")
        disc_clean = normalizar_texto(disc or "").replace(" ", "_")
        turma_clean = normalizar_texto(turma or "").replace(" ", "_")

        unique_name = f"{prof_clean}_{disc_clean}_{turma_clean}_{r_id}_{stem}{ext}"
        filepath = Path(HISTORICO_DOCX_DIR) / unique_name

        if blob:
            tamanho_blob = len(blob)
            try:
                filepath.write_bytes(blob)
                # Verificar tamanho após escrita
                tamanho_gravado = filepath.stat().st_size
                if tamanho_gravado != tamanho_blob:
                    raise RuntimeError(
                        f"Tamanho divergente após gravação: esperado {tamanho_blob}B, "
                        f"gravado {tamanho_gravado}B em {filepath}"
                    )
            except Exception as e:
                # Abortar: limpar arquivos já gravados nesta migração
                logger.error(
                    "Erro ao gravar arquivo físico na migração (id=%s): %s. "
                    "Abortando — tabela original preservada.",
                    r_id, e,
                )
                for fp, _, _ in arquivos_gravados:
                    try:
                        fp.unlink(missing_ok=True)
                    except Exception:
                        pass
                raise RuntimeError(
                    f"Migração abortada: erro ao gravar arquivo {filepath}"
                ) from e
            arquivos_gravados.append((filepath, unique_name, tamanho_blob))
        else:
            # Blob nulo: sem arquivo físico
            arquivos_gravados.append((None, unique_name, 0))

        registros_nova.append(
            (r_id, prof, disc, turma, bim, data_gen, arq_nome, unique_name)
        )

    logger.info(
        "Fase 1 concluída: %d arquivos físicos gravados e verificados.",
        sum(1 for fp, _, _ in arquivos_gravados if fp is not None),
    )

    # --- Fase 2: popular nova tabela e substituir a original ---
    for reg_nova in registros_nova:
        cursor.execute(
            """
            INSERT INTO historico_planos_nova 
                (id, professor_nome, disciplina, turma, bimestre, data_geracao, arquivo_nome, arquivo_path)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            reg_nova,
        )

    # DROP só ocorre após todos os arquivos e registros confirmados
    cursor.execute("DROP TABLE historico_planos")
    cursor.execute("ALTER TABLE historico_planos_nova RENAME TO historico_planos")
    logger.info(
        "Migração de BLOBs concluída com sucesso! %d registros migrados.",
        len(registros_nova),
    )


def init_db():
    with get_connection() as conn:
        _migrar_blob_para_path(conn)
        cursor = conn.cursor()

        cursor.execute(
            """
            CREATE TABLE IF NOT EXISTS professores (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nome TEXT UNIQUE NOT NULL
            )
            """
        )

        cursor.execute(
            """
            CREATE TABLE IF NOT EXISTS professor_turmas (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                professor_id INTEGER,
                disciplina TEXT,
                turma TEXT,
                dia_semana TEXT,
                horario TEXT,
                aulas_semana TEXT,
                arquivo_modelo TEXT,
                template_id TEXT,
                componente_curricular TEXT,
                FOREIGN KEY(professor_id) REFERENCES professores(id) ON DELETE CASCADE
            )
            """
        )

        cursor.execute(
            """
            CREATE TABLE IF NOT EXISTS professor_dados (
                professor_id INTEGER PRIMARY KEY,
                cpf TEXT,
                email TEXT,
                valor_mensal TEXT,
                telefone TEXT,
                observacoes TEXT,
                atualizado_em TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(professor_id) REFERENCES professores(id) ON DELETE CASCADE
            )
            """
        )
        
        

        cursor.execute(
            """
            CREATE TABLE IF NOT EXISTS historico_planos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                professor_nome TEXT,
                disciplina TEXT,
                turma TEXT,
                bimestre TEXT,
                data_geracao TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                arquivo_nome TEXT,
                arquivo_path TEXT,
                professor_chave TEXT,
                disciplina_chave TEXT,
                turma_chave TEXT,
                bimestre_chave TEXT,
                mes_geracao TEXT,
                mes_plano TEXT,
                arquivo_hash TEXT,
                arquivo_tamanho INTEGER,
                origem TEXT,
                ultima_aula INTEGER,
                total_aulas INTEGER
            )
            """
        )


        cursor.execute(
            """
            CREATE TABLE IF NOT EXISTS configuracoes (
                chave TEXT PRIMARY KEY,
                valor TEXT
            )
            """
        )
        _aplicar_migracoes(cursor)
        _atualizar_metadados_historico(cursor)
        _criar_indices_banco(cursor)
        _limpar_historico_planos_incompletos(cursor)
        conn.commit()


def migrar_json_para_sqlite():
    json_path = Path(__file__).resolve().parent.parent / "professores.json"
    if not json_path.exists():
        return

    try:
        with open(json_path, "r", encoding="utf-8") as f:
            dados = json.load(f)

        with get_connection() as conn:
            cursor = conn.cursor()
            for professor, info in dados.items():
                cursor.execute("INSERT OR IGNORE INTO professores (nome) VALUES (?)", (professor,))
                cursor.execute("SELECT id FROM professores WHERE nome = ?", (professor,))
                prof_id = cursor.fetchone()[0]

                for d in info.get("disciplinas", []):
                    existente_id = _buscar_vinculo_existente_equivalente(
                        cursor,
                        prof_id,
                        d.get("disciplina"),
                        d.get("turma"),
                        d.get("componente_curricular", ""),
                    )
                    if not existente_id:
                        cursor.execute(
                            """
                            INSERT INTO professor_turmas
                            (professor_id, disciplina, turma, dia_semana, horario, aulas_semana, arquivo_modelo, template_id, componente_curricular)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                            """,
                            (
                                prof_id,
                                d.get("disciplina"),
                                d.get("turma"),
                                d.get("dia_semana"),
                                d.get("horario"),
                                d.get("aulas_semana"),
                                d.get("arquivo_modelo") or d.get("arquivo") or "",
                                d.get("template_id") or "",
                                d.get("componente_curricular") or "",
                            ),
                        )
            conn.commit()

        os.rename(json_path, json_path.with_suffix(".json.backup"))
    except Exception as e:
        print(f"Erro na migracao do JSON: {e}")


def obter_professor_id_por_nome(nome: str) -> int | None:
    """Retorna o ID do professor pelo nome cadastrado (sem criar se não existir)."""
    if not nome or not isinstance(nome, str):
        return None
    try:
        with get_connection() as conn:
            cursor = conn.cursor()
            nome_norm = _normalizar_campo(nome).upper()
            cursor.execute("SELECT id FROM professores WHERE nome = ?", (nome_norm,))
            row = cursor.fetchone()
            if row:
                return int(row[0])
    except Exception:
        pass
    return None


def _professor_dados_vazio(nome: str = "") -> dict[str, str]:
    return {
        "professor": _normalizar_campo(nome).upper(),
        "cpf": "",
        "email": "",
        "valor_mensal": "",
        "telefone": "",
        "observacoes": "",
        "atualizado_em": "",
    }


def obter_dados_administrativos_professor(nome: str) -> dict[str, str]:
    nome = _normalizar_campo(nome).upper()
    if not nome:
        return _professor_dados_vazio()

    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT
                p.nome,
                COALESCE(d.cpf, ''),
                COALESCE(d.email, ''),
                COALESCE(d.valor_mensal, ''),
                COALESCE(d.telefone, ''),
                COALESCE(d.observacoes, ''),
                COALESCE(d.atualizado_em, '')
            FROM professores p
            LEFT JOIN professor_dados d ON d.professor_id = p.id
            WHERE p.nome = ?
            """,
            (nome,),
        )
        row = cursor.fetchone()

    if not row:
        return _professor_dados_vazio(nome)
    return {
        "professor": row[0] or nome,
        "cpf": row[1] or "",
        "email": row[2] or "",
        "valor_mensal": row[3] or "",
        "telefone": row[4] or "",
        "observacoes": row[5] or "",
        "atualizado_em": row[6] or "",
    }


def salvar_dados_administrativos_professor(
    nome: str,
    *,
    cpf: str = "",
    email: str = "",
    valor_mensal: str = "",
    telefone: str = "",
    observacoes: str = "",
) -> dict[str, str]:
    nome = _normalizar_campo(nome).upper()
    if not nome:
        raise ValueError("Nome do professor é obrigatório.")

    dados = {
        "cpf": _normalizar_campo(cpf),
        "email": _normalizar_campo(email),
        "valor_mensal": _normalizar_campo(valor_mensal),
        "telefone": _normalizar_campo(telefone),
        "observacoes": _normalizar_campo(observacoes),
    }
    with connection_scope() as conn:
        cursor = conn.cursor()
        professor_id = _obter_ou_criar_professor(cursor, nome)
        cursor.execute(
            """
            INSERT INTO professor_dados
                (professor_id, cpf, email, valor_mensal, telefone, observacoes, atualizado_em)
            VALUES (?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
            ON CONFLICT(professor_id) DO UPDATE SET
                cpf = excluded.cpf,
                email = excluded.email,
                valor_mensal = excluded.valor_mensal,
                telefone = excluded.telefone,
                observacoes = excluded.observacoes,
                atualizado_em = CURRENT_TIMESTAMP
            """,
            (
                professor_id,
                dados["cpf"],
                dados["email"],
                dados["valor_mensal"],
                dados["telefone"],
                dados["observacoes"],
            ),
        )

    return obter_dados_administrativos_professor(nome)


def obter_professores_db():
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT p.nome, t.disciplina, t.turma, t.dia_semana, t.horario, t.aulas_semana, t.arquivo_modelo, t.template_id, t.componente_curricular
            FROM professores p
            LEFT JOIN professor_turmas t ON p.id = t.professor_id
            ORDER BY
                p.nome,
                CASE
                    WHEN COALESCE(t.dia_semana, '') <> ''
                     AND COALESCE(t.horario, '') <> ''
                     AND COALESCE(t.aulas_semana, '') <> '' THEN 0
                    ELSE 1
                END,
                t.disciplina,
                t.turma
                , COALESCE(t.componente_curricular, '')
            """
        )

        resultado = {}
        for row in cursor.fetchall():
            nome = row[0]
            if nome not in resultado:
                resultado[nome] = {"disciplinas": []}
            if row[1] and row[2]:
                resultado[nome]["disciplinas"].append(
                    {
                        "disciplina": row[1],
                        "turma": row[2],
                        "dia_semana": row[3] or "",
                        "horario": row[4] or "",
                        "aulas_semana": row[5] or "",
                        "arquivo": row[6] or "",
                        "arquivo_modelo": row[6] or "",
                        "template_id": row[7] or "",
                        "componente_curricular": row[8] or "",
                        "origem": "banco",
                    }
                )
        return resultado


def salvar_professor_turma(
    nome,
    disciplina,
    turma,
    dia_semana,
    horario,
    aulas_semana,
    arquivo_modelo="",
    componente_curricular="",
    template_id="",
):
    with get_connection() as conn:
        cursor = conn.cursor()
        prof_id = _obter_ou_criar_professor(cursor, nome)
        disciplina = _normalizar_campo(disciplina)
        turma = _normalizar_campo(turma)
        dia_semana = _normalizar_campo(dia_semana)
        horario = _normalizar_campo(horario)
        aulas_semana = _normalizar_campo(aulas_semana)
        arquivo_modelo = _normalizar_campo(arquivo_modelo)
        template_id = _normalizar_campo(template_id)
        componente_curricular = _normalizar_campo(componente_curricular)
        existente_id = _buscar_vinculo_existente_equivalente(
            cursor,
            prof_id,
            disciplina,
            turma,
            componente_curricular,
        )

        if existente_id:
            cursor.execute(
                """
                UPDATE professor_turmas
                SET dia_semana = ?, horario = ?, aulas_semana = ?, arquivo_modelo = ?, template_id = ?, componente_curricular = ?
                WHERE id = ?
                """,
                (dia_semana, horario, aulas_semana, arquivo_modelo, template_id, componente_curricular, existente_id),
            )
        else:
            cursor.execute(
                """
                INSERT INTO professor_turmas
                (professor_id, disciplina, turma, dia_semana, horario, aulas_semana, arquivo_modelo, template_id, componente_curricular)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (prof_id, disciplina, turma, dia_semana, horario, aulas_semana, arquivo_modelo, template_id, componente_curricular),
            )
        conn.commit()


def listar_vinculos_professores():
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT
                t.id,
                p.id,
                p.nome,
                t.disciplina,
                t.turma,
                t.dia_semana,
                t.horario,
                t.aulas_semana,
                t.arquivo_modelo,
                t.template_id,
                t.componente_curricular
            FROM professor_turmas t
            JOIN professores p ON p.id = t.professor_id
            ORDER BY p.nome, t.disciplina, t.turma, t.id
                     , COALESCE(t.componente_curricular, '')
            """
        )
        return [
            {
                "id": row[0],
                "professor_id": row[1],
                "professor": row[2] or "",
                "disciplina": row[3] or "",
                "turma": row[4] or "",
                "dia_semana": row[5] or "",
                "horario": row[6] or "",
                "aulas_semana": row[7] or "",
                "arquivo": row[8] or "",
                "arquivo_modelo": row[8] or "",
                "template_id": row[9] or "",
                "componente_curricular": row[10] or "",
                "origem": "banco",
            }
            for row in cursor.fetchall()
        ]


def obter_vinculo_professor(vinculo_id):
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT
                t.id,
                p.id,
                p.nome,
                t.disciplina,
                t.turma,
                t.dia_semana,
                t.horario,
                t.aulas_semana,
                t.arquivo_modelo,
                t.template_id,
                t.componente_curricular
            FROM professor_turmas t
            JOIN professores p ON p.id = t.professor_id
            WHERE t.id = ?
            """,
            (vinculo_id,),
        )
        row = cursor.fetchone()
        if not row:
            return None
        return {
            "id": row[0],
            "professor_id": row[1],
            "professor": row[2] or "",
            "disciplina": row[3] or "",
            "turma": row[4] or "",
            "dia_semana": row[5] or "",
            "horario": row[6] or "",
            "aulas_semana": row[7] or "",
            "arquivo": row[8] or "",
            "arquivo_modelo": row[8] or "",
            "template_id": row[9] or "",
            "componente_curricular": row[10] or "",
            "origem": "banco",
        }


def atualizar_vinculo_professor(
    vinculo_id,
    nome,
    disciplina,
    turma,
    dia_semana,
    horario,
    aulas_semana,
    arquivo_modelo="",
    componente_curricular="",
    template_id="",
):
    nome = _normalizar_campo(nome).upper()
    disciplina = _normalizar_campo(disciplina)
    turma = _normalizar_campo(turma)
    if not nome or not disciplina or not turma:
        raise ValueError("Professor, disciplina e turma sao obrigatorios.")

    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT professor_id FROM professor_turmas WHERE id = ?", (vinculo_id,))
        row = cursor.fetchone()
        if not row:
            raise ValueError("Cadastro nao encontrado.")
        professor_antigo_id = int(row[0])
        professor_id = _obter_ou_criar_professor(cursor, nome)
        vinculo_equivalente_id = _buscar_vinculo_existente_equivalente(
            cursor,
            professor_id,
            disciplina,
            turma,
            componente_curricular,
            ignorar_id=vinculo_id,
        )
        if vinculo_equivalente_id is not None:
            cursor.execute(
                """
                UPDATE professor_turmas
                SET dia_semana = ?,
                    horario = ?,
                    aulas_semana = ?,
                    arquivo_modelo = ?,
                    template_id = ?,
                    componente_curricular = ?
                WHERE id = ?
                """,
                (
                    _normalizar_campo(dia_semana),
                    _normalizar_campo(horario),
                    _normalizar_campo(aulas_semana),
                    _normalizar_campo(arquivo_modelo),
                    _normalizar_campo(template_id),
                    _normalizar_campo(componente_curricular),
                    vinculo_equivalente_id,
                ),
            )
            cursor.execute("DELETE FROM professor_turmas WHERE id = ?", (vinculo_id,))
            if professor_antigo_id != professor_id:
                _remover_professor_sem_turmas(cursor, professor_antigo_id)
            conn.commit()
            return obter_vinculo_professor(vinculo_equivalente_id)
        cursor.execute(
            """
            UPDATE professor_turmas
            SET professor_id = ?,
                disciplina = ?,
                turma = ?,
                dia_semana = ?,
                horario = ?,
                aulas_semana = ?,
                arquivo_modelo = ?,
                template_id = ?,
                componente_curricular = ?
            WHERE id = ?
            """,
            (
                professor_id,
                disciplina,
                turma,
                _normalizar_campo(dia_semana),
                _normalizar_campo(horario),
                _normalizar_campo(aulas_semana),
                _normalizar_campo(arquivo_modelo),
                _normalizar_campo(template_id),
                _normalizar_campo(componente_curricular),
                vinculo_id,
            ),
        )
        if professor_antigo_id != professor_id:
            _remover_professor_sem_turmas(cursor, professor_antigo_id)
        conn.commit()
    return obter_vinculo_professor(vinculo_id)


def excluir_vinculo_professor(vinculo_id) -> bool:
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT professor_id FROM professor_turmas WHERE id = ?", (vinculo_id,))
        row = cursor.fetchone()
        if not row:
            return False
        professor_id = int(row[0])
        cursor.execute("DELETE FROM professor_turmas WHERE id = ?", (vinculo_id,))
        _remover_professor_sem_turmas(cursor, professor_id)
        conn.commit()
    return True


def duplicar_vinculo_professor(
    vinculo_id,
    nome=None,
    disciplina=None,
    turma=None,
    dia_semana=None,
    horario=None,
    aulas_semana=None,
    arquivo_modelo=None,
    componente_curricular=None,
    template_id=None,
) -> int:
    original = obter_vinculo_professor(vinculo_id)
    if not original:
        raise ValueError("Cadastro original nao encontrado.")

    with get_connection() as conn:
        cursor = conn.cursor()
        professor_id = _obter_ou_criar_professor(cursor, nome or original["professor"])
        disciplina_final = _normalizar_campo(disciplina if disciplina is not None else original["disciplina"])
        turma_final = _normalizar_campo(turma if turma is not None else original["turma"])
        componente_final = _normalizar_campo(
            componente_curricular
            if componente_curricular is not None
            else original["componente_curricular"]
        )
        vinculo_existente_id = _buscar_vinculo_existente_equivalente(
            cursor,
            professor_id,
            disciplina_final,
            turma_final,
            componente_final,
        )
        if vinculo_existente_id is not None:
            return vinculo_existente_id
        cursor.execute(
            """
            INSERT INTO professor_turmas
            (professor_id, disciplina, turma, dia_semana, horario, aulas_semana, arquivo_modelo, template_id, componente_curricular)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                professor_id,
                disciplina_final,
                turma_final,
                _normalizar_campo(dia_semana if dia_semana is not None else original["dia_semana"]),
                _normalizar_campo(horario if horario is not None else original["horario"]),
                _normalizar_campo(aulas_semana if aulas_semana is not None else original["aulas_semana"]),
                _normalizar_campo(arquivo_modelo if arquivo_modelo is not None else original["arquivo_modelo"]),
                _normalizar_campo(template_id if template_id is not None else original["template_id"]),
                componente_final,
            ),
        )
        novo_id = int(cursor.lastrowid)
        conn.commit()
    return novo_id


def salvar_historico_plano(
    professor_nome,
    disciplina,
    turma,
    arquivo_nome,
    arquivo_docx_bytes,
    limite_retencao: int = 5,
    bimestre: str = "",
    mes_plano: str = "",
):
    professor_nome = _normalizar_campo(professor_nome)
    disciplina = _normalizar_campo(disciplina)
    turma = _normalizar_campo(turma)
    bimestre = _normalizar_campo(bimestre)
    arquivo_nome = _normalizar_campo(arquivo_nome)
    arquivo_docx_bytes = bytes(arquivo_docx_bytes or b"")

    prof_clean = normalizar_texto(professor_nome).replace(" ", "_")
    disc_clean = normalizar_texto(disciplina).replace(" ", "_")
    turma_clean = normalizar_texto(turma).replace(" ", "_")
    ts = uuid.uuid4().hex[:8]
    nome_fisico = Path(arquivo_nome or "plano.docx").name
    nome_fisico = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", nome_fisico).strip(" .")
    nome_fisico = nome_fisico or "plano.docx"
    unique_filename = f"{prof_clean}_{disc_clean}_{turma_clean}_{ts}_{nome_fisico}"
    filepath = Path(HISTORICO_DOCX_DIR) / unique_filename
    data_geracao = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    metadados = _metadados_historico(
        professor_nome=professor_nome,
        disciplina=disciplina,
        turma=turma,
        bimestre=bimestre,
        mes_plano=mes_plano,
        data_geracao=data_geracao,
        arquivo_path=unique_filename,
        arquivo_docx_bytes=arquivo_docx_bytes,
        extrair_resumo_aulas=True,
    )
    bimestre = str(metadados["bimestre"] or bimestre)

    try:
        _gravar_arquivo_historico_atomico(filepath, arquivo_docx_bytes)
    except OSError as exc:
        logger.error("Erro ao salvar arquivo físico de histórico: %s", exc)
        return

    caminhos_remover: list[str] = []
    deletados = 0
    try:
        with connection_scope() as conn:
            cursor = conn.cursor()
            cursor.execute(
                """
                INSERT INTO historico_planos
                    (
                        professor_nome,
                        disciplina,
                        turma,
                        bimestre,
                        data_geracao,
                        arquivo_nome,
                        arquivo_path,
                        professor_chave,
                        disciplina_chave,
                        turma_chave,
                        bimestre_chave,
                        mes_geracao,
                        mes_plano,
                        arquivo_hash,
                        arquivo_tamanho,
                        origem,
                        ultima_aula,
                        total_aulas
                    )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    professor_nome,
                    disciplina,
                    turma,
                    bimestre,
                    data_geracao,
                    arquivo_nome,
                    unique_filename,
                    metadados["professor_chave"],
                    metadados["disciplina_chave"],
                    metadados["turma_chave"],
                    metadados["bimestre_chave"],
                    metadados["mes_geracao"],
                    metadados["mes_plano"],
                    metadados["arquivo_hash"],
                    metadados["arquivo_tamanho"],
                    metadados["origem"],
                    metadados["ultima_aula"],
                    metadados["total_aulas"],
                ),
            )

            if limite_retencao > 0:
                _atualizar_metadados_historico(cursor)
                cursor.execute(
                    """
                    SELECT id, arquivo_path FROM historico_planos
                    WHERE professor_chave = ?
                      AND disciplina_chave = ?
                      AND turma_chave = ?
                      AND bimestre_chave = ?
                    ORDER BY data_geracao DESC, id DESC
                    LIMIT -1 OFFSET ?
                    """,
                    (
                        metadados["professor_chave"],
                        metadados["disciplina_chave"],
                        metadados["turma_chave"],
                        metadados["bimestre_chave"],
                        limite_retencao,
                    ),
                )
                registros_remover = cursor.fetchall()
                caminhos_remover = [row[1] for row in registros_remover if row[1]]
                if registros_remover:
                    cursor.executemany(
                        "DELETE FROM historico_planos WHERE id = ?",
                        [(row[0],) for row in registros_remover],
                    )
                    deletados = len(registros_remover)
    except Exception:
        _remover_arquivo_historico_gerenciado(unique_filename)
        raise

    for arquivo_path in caminhos_remover:
        if not _remover_arquivo_historico_gerenciado(arquivo_path):
            logger.warning(
                "Arquivo antigo não removido pela retenção por estar fora do histórico "
                "gerenciado ou indisponível: %s",
                arquivo_path,
            )

    if deletados > 0:
        logger.info(
            "Política de retenção aplicada: %d planos antigos removidos para %s - %s - %s - %s",
            deletados,
            professor_nome,
            disciplina,
            turma,
            bimestre,
        )



def listar_historico_planos(limite=HISTORICO_PLANOS_LIMITE_PADRAO):
    limite = _normalizar_limite_historico(limite)
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT id, professor_nome, disciplina, turma, data_geracao, arquivo_nome
            FROM historico_planos
            ORDER BY data_geracao DESC, id DESC
            LIMIT ?
            """,
            (limite,),
        )
        return cursor.fetchall()


def listar_ultimos_planos_por_contexto(bimestre: str = "") -> list[dict]:
    bimestre = _normalizar_campo(bimestre)
    bimestre_chave = _normalizar_bimestre_historico_chave(bimestre)

    with get_connection() as conn:
        cursor = conn.cursor()
        if bimestre_chave:
            cursor.execute(
                """
                SELECT
                    h.id,
                    h.professor_nome,
                    h.disciplina,
                    h.turma,
                    COALESCE(h.bimestre, ''),
                    h.data_geracao,
                    h.arquivo_nome
                FROM historico_planos h
                JOIN (
                    SELECT
                        professor_chave,
                        disciplina_chave,
                        turma_chave,
                        bimestre_chave,
                        MAX(id) AS ultimo_id
                    FROM historico_planos
                    WHERE bimestre_chave = ?
                    GROUP BY
                        professor_chave,
                        disciplina_chave,
                        turma_chave,
                        bimestre_chave
                ) ultimos
                    ON h.id = ultimos.ultimo_id
                ORDER BY
                    h.professor_chave,
                    h.disciplina_chave,
                    h.turma_chave,
                    h.data_geracao DESC,
                    h.id DESC
                """,
                (bimestre_chave,),
            )
        else:
            cursor.execute(
                """
                SELECT
                    h.id,
                    h.professor_nome,
                    h.disciplina,
                    h.turma,
                    COALESCE(h.bimestre, ''),
                    h.data_geracao,
                    h.arquivo_nome
                FROM historico_planos h
                JOIN (
                    SELECT
                        professor_chave,
                        disciplina_chave,
                        turma_chave,
                        bimestre_chave,
                        MAX(id) AS ultimo_id
                    FROM historico_planos
                    GROUP BY
                        professor_chave,
                        disciplina_chave,
                        turma_chave,
                        bimestre_chave
                ) ultimos
                    ON h.id = ultimos.ultimo_id
                ORDER BY
                    h.professor_chave,
                    h.disciplina_chave,
                    h.turma_chave,
                    h.bimestre_chave,
                    h.data_geracao DESC,
                    h.id DESC
                """
            )

        return [
            {
                "id": int(row[0]),
                "professor_nome": row[1] or "",
                "disciplina": row[2] or "",
                "turma": row[3] or "",
                "bimestre": row[4] or "",
                "data_geracao": row[5] or "",
                "arquivo_nome": row[6] or "",
            }
            for row in cursor.fetchall()
        ]


def obter_meses_historico_planos() -> list[str]:
    """Retorna anos-meses (YYYY-MM) do plano, com fallback para a geração."""
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT DISTINCT
                COALESCE(NULLIF(mes_plano, ''), NULLIF(mes_geracao, ''), strftime('%Y-%m', data_geracao)) as mes
            FROM historico_planos
            WHERE data_geracao IS NOT NULL
            ORDER BY mes DESC
            """
        )
        return [row[0] for row in cursor.fetchall() if row[0]]


def obter_bimestres_historico_planos() -> list[str]:
    """Retorna bimestres preenchidos no histórico, preservando o texto salvo."""
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT bimestre
            FROM (
                SELECT
                    COALESCE(NULLIF(TRIM(bimestre), ''), '') AS bimestre,
                    COALESCE(NULLIF(TRIM(bimestre_chave), ''), '') AS bimestre_chave,
                    MAX(data_geracao) AS ultima_data,
                    MAX(id) AS ultimo_id
                FROM historico_planos
                WHERE COALESCE(NULLIF(TRIM(bimestre), ''), '') <> ''
                GROUP BY COALESCE(NULLIF(TRIM(bimestre_chave), ''), '')
            )
            WHERE bimestre <> ''
            ORDER BY bimestre_chave, ultima_data DESC, ultimo_id DESC
            """
        )
        return [row[0] for row in cursor.fetchall() if row[0]]


def buscar_historico_planos(professor_nome: str, mes: str = "") -> list[dict]:
    """Busca os planos gerados por um professor, filtrando opcionalmente por mês (YYYY-MM)."""
    professor_chave = _normalizar_campo_chave(professor_nome)
    mes = _normalizar_campo(mes)

    with get_connection() as conn:
        cursor = conn.cursor()
        
        query = """
            SELECT id, professor_nome, disciplina, turma, bimestre, data_geracao, arquivo_nome, arquivo_path, mes_plano
            FROM historico_planos
            WHERE professor_chave = ?
        """
        params = [professor_chave]
        
        if mes:
            query += " AND COALESCE(NULLIF(mes_plano, ''), NULLIF(mes_geracao, ''), strftime('%Y-%m', data_geracao)) = ?"
            params.append(mes)
            
        query += " ORDER BY data_geracao DESC, id DESC"
        
        cursor.execute(query, params)
        return [
            {
                "id": int(row[0]),
                "professor_nome": row[1] or "",
                "disciplina": row[2] or "",
                "turma": row[3] or "",
                "bimestre": row[4] or "",
                "data_geracao": row[5] or "",
                "arquivo_nome": row[6] or "",
                "arquivo_path": row[7] or "",
                "mes_plano": row[8] or "",
            }
            for row in cursor.fetchall()
        ]


def buscar_historico_planos_avancado(
    professor_nome: str = "",
    mes: str = "",
    disciplina: str = "",
    turma: str = "",
    bimestre: str = "",
    termo_busca: str = "",
    somente_disponiveis: bool = False,
    limite: int = HISTORICO_PLANOS_LIMITE_MAXIMO,
) -> list[dict]:
    """Busca histórico com filtros combináveis e metadados para a tela.

    A consulta usa as chaves normalizadas da Fase 1 quando possível. O filtro
    textual fica como busca auxiliar em campos curtos do histórico.
    """
    limite = _normalizar_limite_historico(limite)
    filtros = []
    params: list[object] = []

    professor_chave = _normalizar_campo_chave(professor_nome)
    if professor_chave:
        filtros.append("professor_chave = ?")
        params.append(professor_chave)

    mes = _normalizar_campo(mes)
    if mes:
        filtros.append(
            "COALESCE(NULLIF(mes_plano, ''), NULLIF(mes_geracao, ''), strftime('%Y-%m', data_geracao)) = ?"
        )
        params.append(mes)

    disciplina_chave = _normalizar_campo_chave(disciplina)
    if disciplina_chave:
        filtros.append("disciplina_chave = ?")
        params.append(disciplina_chave)

    turma_chave = _normalizar_turma_historico_chave(turma)
    if turma_chave:
        filtros.append("turma_chave = ?")
        params.append(turma_chave)

    bimestre_chave = _normalizar_bimestre_historico_chave(bimestre)
    if bimestre_chave:
        filtros.append("(bimestre_chave = ? OR COALESCE(TRIM(bimestre_chave), '') = '')")
        params.append(bimestre_chave)

    termo = _normalizar_campo(termo_busca)
    if termo:
        termo_like = f"%{termo}%"
        termo_chave_like = f"%{_normalizar_campo_chave(termo)}%"
        filtros.append(
            """
            (
                arquivo_nome LIKE ? COLLATE NOCASE
                OR professor_nome LIKE ? COLLATE NOCASE
                OR disciplina LIKE ? COLLATE NOCASE
                OR turma LIKE ? COLLATE NOCASE
                OR professor_chave LIKE ?
                OR disciplina_chave LIKE ?
                OR turma_chave LIKE ?
            )
            """
        )
        params.extend(
            [
                termo_like,
                termo_like,
                termo_like,
                termo_like,
                termo_chave_like,
                termo_chave_like,
                termo_chave_like,
            ]
        )

    where_sql = f"WHERE {' AND '.join(filtros)}" if filtros else ""
    params.append(limite)

    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            f"""
            SELECT
                id,
                professor_nome,
                disciplina,
                turma,
                bimestre,
                data_geracao,
                arquivo_nome,
                arquivo_path,
                origem,
                arquivo_tamanho,
                ultima_aula,
                total_aulas,
                mes_plano
            FROM historico_planos
            {where_sql}
            ORDER BY data_geracao DESC, id DESC
            LIMIT ?
            """,
            params,
        )
        rows = cursor.fetchall()

    resultados = []
    for row in rows:
        arquivo_path = row[7] or ""
        arquivo_disponivel = False
        if arquivo_path:
            try:
                arquivo_disponivel = _resolver_caminho_arquivo_historico(arquivo_path).exists()
            except OSError:
                arquivo_disponivel = False
        if somente_disponiveis and not arquivo_disponivel:
            continue
        resultados.append(
            {
                "id": int(row[0]),
                "professor_nome": row[1] or "",
                "disciplina": row[2] or "",
                "turma": row[3] or "",
                "bimestre": row[4] or "",
                "data_geracao": row[5] or "",
                "arquivo_nome": row[6] or "",
                "arquivo_path": arquivo_path,
                "origem": row[8] or "",
                "arquivo_tamanho": int(row[9]) if row[9] is not None else None,
                "ultima_aula": int(row[10]) if row[10] is not None else None,
                "total_aulas": int(row[11]) if row[11] is not None else None,
                "mes_plano": row[12] or "",
                "arquivo_disponivel": arquivo_disponivel,
            }
        )
    return resultados



def obter_arquivo_historico(plano_id):
    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT arquivo_nome, arquivo_path FROM historico_planos WHERE id = ?", (plano_id,))
        row = cursor.fetchone()
        if not row:
            return None
        nome, path_rel = row
        bytes_content = b""
        if path_rel:
            p_abs = _resolver_caminho_arquivo_historico(path_rel)
            if p_abs.exists():
                try:
                    bytes_content = p_abs.read_bytes()
                except Exception as e:
                    logger.error(f"Erro ao ler arquivo fisico do historico: {e}")
        return nome, bytes_content



def obter_ultimo_historico_por_contexto(
    professor_nome: str,
    disciplina: str,
    turma: str,
    bimestre: str = "",
) -> dict | None:
    """Retorna o plano recente do contexto, inclusive registros antigos sem bimestre."""
    professor_nome = _normalizar_campo(professor_nome)
    disciplina = _normalizar_campo(disciplina)
    turma = _normalizar_campo(turma)
    bimestre = _normalizar_campo(bimestre)

    if not professor_nome or not disciplina or not turma:
        return None

    chave_professor = _normalizar_campo_chave(professor_nome)
    chave_disciplina = _normalizar_campo_chave(disciplina)
    chave_turma_esperada = _normalizar_turma_historico_chave(turma)
    chave_bimestre_esperada = _normalizar_bimestre_historico_chave(bimestre)

    with get_connection() as conn:
        cursor = conn.cursor()
        if chave_bimestre_esperada:
            cursor.execute(
                """
                SELECT id, bimestre, data_geracao, arquivo_nome, ultima_aula, total_aulas
                FROM historico_planos
                WHERE professor_chave = ?
                  AND disciplina_chave = ?
                  AND turma_chave = ?
                  AND (bimestre_chave = ? OR COALESCE(TRIM(bimestre_chave), '') = '')
                ORDER BY
                    CASE WHEN bimestre_chave = ? THEN 0 ELSE 1 END,
                    data_geracao DESC,
                    id DESC
                LIMIT 1
                """,
                (
                    chave_professor,
                    chave_disciplina,
                    chave_turma_esperada,
                    chave_bimestre_esperada,
                    chave_bimestre_esperada,
                ),
            )
        else:
            cursor.execute(
                """
                SELECT id, bimestre, data_geracao, arquivo_nome, ultima_aula, total_aulas
                FROM historico_planos
                WHERE professor_chave = ?
                  AND disciplina_chave = ?
                  AND turma_chave = ?
                ORDER BY data_geracao DESC, id DESC
                LIMIT 1
                """,
                (chave_professor, chave_disciplina, chave_turma_esperada),
            )
        registro = cursor.fetchone()

    if not registro:
        return None

    return {
        "id": int(registro[0]),
        "bimestre": registro[1] or "",
        "data_geracao": registro[2] or "",
        "arquivo_nome": registro[3] or "",
        "ultima_aula": int(registro[4]) if registro[4] is not None else None,
        "total_aulas": int(registro[5]) if registro[5] is not None else None,
    }



def obter_ultimo_plano_docx(professor_nome: str, disciplina: str, turma: str) -> bytes | None:
    professor_nome = _normalizar_campo(professor_nome)
    disciplina = _normalizar_campo(disciplina)
    turma = _normalizar_campo(turma)
    professor_chave = _normalizar_campo_chave(professor_nome)
    disciplina_chave = _normalizar_campo_chave(disciplina)
    turma_chave = _normalizar_turma_historico_chave(turma)

    with get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT arquivo_path FROM historico_planos
            WHERE professor_chave = ?
              AND disciplina_chave = ?
              AND turma_chave = ?
            ORDER BY data_geracao DESC, id DESC
            LIMIT 1
            """,
            (professor_chave, disciplina_chave, turma_chave),
        )
        row = cursor.fetchone()
        if row and row[0]:
            p_abs = _resolver_caminho_arquivo_historico(row[0])
            if p_abs.exists():
                try:
                    return p_abs.read_bytes()
                except Exception as e:
                    logger.error(f"Erro ao ler ultimo arquivo docx: {e}")
        return None



def obter_ultima_aula_gerada_sistema(professor: str, disciplina: str, turma: str, bimestre: str = "") -> int:
    from core.gestao_aulas import obter_ultima_aula_gerada_sistema_impl
    return obter_ultima_aula_gerada_sistema_impl(professor, disciplina, turma, bimestre)


def verificar_plano_gerado_por_outro_professor(
    professor_nome: str,
    disciplina: str,
    turma: str,
    bimestre: str = "",
) -> list[dict]:
    """
    Retorna lista de registros do historico onde a mesma disciplina e turma
    foram geradas para um professor diferente do atual.
    """
    professor_nome = _normalizar_campo(professor_nome)
    disciplina = _normalizar_campo(disciplina)
    turma = _normalizar_campo(turma)
    bimestre = _normalizar_campo(bimestre)
    professor_chave = _normalizar_campo_chave(professor_nome)
    disciplina_chave = _normalizar_campo_chave(disciplina)
    turma_chave = _normalizar_turma_historico_chave(turma)
    bimestre_chave = _normalizar_bimestre_historico_chave(bimestre)

    if not professor_nome or not disciplina or not turma:
        return []

    with get_connection() as conn:
        cursor = conn.cursor()
        if bimestre:
            cursor.execute(
                """
                SELECT DISTINCT professor_nome, data_geracao, arquivo_nome, bimestre
                FROM historico_planos
                WHERE professor_chave <> ?
                  AND disciplina_chave = ?
                  AND turma_chave = ?
                  AND bimestre_chave = ?
                ORDER BY data_geracao DESC
                """,
                (professor_chave, disciplina_chave, turma_chave, bimestre_chave),
            )
        else:
            cursor.execute(
                """
                SELECT DISTINCT professor_nome, data_geracao, arquivo_nome, bimestre
                FROM historico_planos
                WHERE professor_chave <> ?
                  AND disciplina_chave = ?
                  AND turma_chave = ?
                ORDER BY data_geracao DESC
                """,
                (professor_chave, disciplina_chave, turma_chave),
            )
        return [
            {
                "professor_nome": row[0],
                "data_geracao": row[1],
                "arquivo_nome": row[2],
                "bimestre": row[3],
            }
            for row in cursor.fetchall()
        ]

