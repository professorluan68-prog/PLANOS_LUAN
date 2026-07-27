import re
import unicodedata
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from typing import Dict, Iterable, Tuple

from config import PDF_AULAS_DIR, REFERENCIAS_METODOLOGICAS_DIR

REFERENCIA_LEITURA_REDACAO = "🧠🔥 GUIA METODOLÓGICO ESTRUTURADO - LEITURA E REDAÇÃO.md"

# Títulos de etapa que NUNCA podem aparecer em planos de História
TITULOS_PROIBIDOS_HISTORIA: tuple[str, ...] = (
    "pause e responda",
    "pause e responda:",
)

# Mapa global de títulos proibidos por perfil (usado em ia.py e revisao_final.py)
TITULOS_PROIBIDOS_POR_PERFIL: dict[str, tuple[str, ...]] = {
    "historia": TITULOS_PROIBIDOS_HISTORIA,
    # Adicionar outros perfis aqui conforme necessário
}


def get_titulos_proibidos(perfil: str) -> tuple[str, ...]:
    """Retorna os títulos de etapa proibidos para um dado perfil."""
    import unicodedata
    perfil_norm = unicodedata.normalize("NFKD", str(perfil or "").lower().strip())
    perfil_norm = "".join(ch for ch in perfil_norm if not unicodedata.combining(ch))
    return TITULOS_PROIBIDOS_POR_PERFIL.get(perfil_norm, ())


REGRAS_ESTRUTURAIS_HISTORIA = """\
REGRAS ESTRUTURAIS OBRIGATÓRIAS — HISTÓRIA (aplicar sempre):

REGRA 1 [PROIBIÇÃO ABSOLUTA]: NUNCA gere uma etapa com título "Pause e responda".
  - Esta regra não tem exceção para História.
  - Se o PDF contiver a seção "PAUSE E RESPONDA", ignore-a completamente.
  - Não mencione, não incorpore, não parafraseie seu conteúdo.
  - Se você gerar uma etapa com esse título, o plano será REJEITADO automaticamente.

REGRA 2 [FUSÃO OBRIGATÓRIA]: NUNCA use "Para começar" e "Relembre" na mesma aula.
  - Se o PDF tiver ambos, sintetize os dois em um único bloco "Para começar".

REGRA 3 [FUSÃO DE CONSECUTIVOS]: Se houver múltiplos "Foco no conteúdo" CONSECUTIVOS
  (sem "Na prática" entre eles), reúna-os em UM ÚNICO bloco "Foco no conteúdo" mais conciso.

REGRA 4 [SEPARAÇÃO OBRIGATÓRIA]: Se houver "Foco no conteúdo" ANTES de uma "Na prática"
  E outros "Foco no conteúdo" DEPOIS dessa mesma "Na prática", gere DOIS blocos separados:
  um antes e um depois da "Na prática". NÃO funda os dois.

REGRA 5 [NUMERAÇÃO DE ATIVIDADES]: Se houver múltiplas atividades em "Na prática"
  CONSECUTIVAS (sem "Foco no conteúdo" entre elas), descreva-as em UM ÚNICO bloco
  "Na prática", numerando cada atividade (Atividade 1, Atividade 2...).

REGRA 6 [ENCERRAMENTO OBRIGATÓRIO]: O "Encerramento" DEVE sempre aparecer ao final,
  incluindo obrigatoriamente a técnica "COM SUAS PALAVRAS" e as perguntas finais do PDF.

REGRA 7 [LIMITE DE TAMANHO — CRÍTICO]: Cada etapa da metodologia de História deve ter
  NO MÁXIMO 350 caracteres. Conte os caracteres antes de finalizar cada etapa.
  Se ultrapassar 350 caracteres, corte na última frase completa antes do limite.
  Uma etapa com mais de 350 caracteres será truncada automaticamente pelo sistema.
"""

REGRAS_TECNICAS_HISTORIA = """\
REGRAS DE POSICIONAMENTO DE TÉCNICAS LEMOV — HISTÓRIA:
- "VIREM E CONVERSEM": usar APENAS em "Para começar" ou momentos de discussão inicial.
- "HORA DA LEITURA": usar APENAS quando há leitura de texto/fonte no material.
- "TODO MUNDO ESCREVE": usar APENAS para registro individual de atividade.
- "COM SUAS PALAVRAS": usar APENAS no "Encerramento" ou síntese final.
- "DE OLHO NO MODELO": usar APENAS antes de atividade de produção com modelo.
NUNCA usar "COM SUAS PALAVRAS" em "Para começar".
NUNCA usar "VIREM E CONVERSEM" no "Encerramento".
NUNCA usar "PAUSE E RESPONDA" em nenhuma etapa de História.
"""


def get_regras_estruturais_historia() -> str:
    """Retorna as regras estruturais + técnicas LEMOV para injeção no prompt de História."""
    return REGRAS_ESTRUTURAIS_HISTORIA + "\n" + REGRAS_TECNICAS_HISTORIA


PASTAS_BUSCA = (
    REFERENCIAS_METODOLOGICAS_DIR,
    PDF_AULAS_DIR,
)


def _caminho_em_pasta_oficial(caminho: Path) -> bool:
    """Impede que referencias externas ou legadas entrem no prompt da IA."""
    try:
        caminho_resolvido = caminho.resolve()
    except OSError:
        return False

    for pasta in PASTAS_BUSCA:
        try:
            caminho_resolvido.relative_to(pasta.resolve())
            return True
        except ValueError:
            continue
        except OSError:
            continue
    return False


def resolver_caminho_referencia(arquivo: str | Path) -> Path | None:
    arq_path = Path(arquivo)
    if arq_path.is_absolute():
        if arq_path.is_file() and _caminho_em_pasta_oficial(arq_path):
            return arq_path
        return None
    for pasta in PASTAS_BUSCA:
        caminho = pasta / arquivo
        if caminho.is_file():
            return caminho
    return None

LIMITE_REFERENCIA_CHARS = 6200
LIMITE_INTERDISCIPLINAR_CHARS = 1800


@dataclass(frozen=True)
class DiagnosticoReferenciaMetodologica:
    texto: str
    arquivos_solicitados: tuple[str, ...]
    arquivos_encontrados: tuple[str, ...]
    arquivos_ausentes: tuple[str, ...]
    aviso: str = ""

REFERENCIA_INTERDISCIPLINAR = "ADAPTAÇÃO METODOLÓGICA INTERDISCIPLINAR.md"

REFERENCIA_PORTUGUES_GERAL = "ANÁLISE METODOLÓGICA COMPLETA - LÍNGUA PORTUGUESA.md"
REFERENCIA_PORTUGUES_FUNDAMENTAL = "ANÁLISE METODOLÓGICA - LÍNGUA PORTUGUESA ENSINO FUNDAMENTAL.md"
REFERENCIA_PORTUGUES_MEDIO = "ANÁLISE METODOLÓGICA - LÍNGUA PORTUGUESA ENSINO MÉDIO.md"

REFERENCIA_PV_GERAL = "ANÁLISE METODOLÓGICA COMPLETA - PROJETO DE VIDA.md"
REFERENCIA_PV_FUNDAMENTAL = "ANÁLISE METODOLÓGICA - PROJETO DE VIDA ENSINO FUNDAMENTAL.md"
REFERENCIA_PV_REFINADA_FUNDAMENTAL = "ANÁLISE METODOLÓGICA REFINADA - PROJETO DE VIDA ENSINO FUNDAMENTAL.md"
REFERENCIA_PV_FUNDAMENTAL_ANOS_FINAIS = "ANÁLISE METODOLÓGICA - PROJETO DE VIDA - ENSINO FUNDAMENTAL ANOS FINAIS.md"
REFERENCIA_PV_7_ANO = "ANÁLISE METODOLÓGICA - PROJETO DE VIDA 7º ANO.md"

MAPA_REFERENCIAS_NOVAS = {
    "arte": "melhorias_metodologicas_arte.md",
    "biologia": "melhorias_metodologicas_biologia.md",
    "ciencias": "melhorias_metodologicas_ciencias.md",
    "ciencia": "melhorias_metodologicas_ciencias.md",
    "educacao financeira": "melhorias_metodologicas_educacao_financeira.md",
    "geografia": "melhorias_metodologicas_geografia.md",
    "historia": "melhorias_metodologicas_historia.md",
    "ingles": "melhorias_metodologicas_lingua_inglesa.md",
    "english": "melhorias_metodologicas_lingua_inglesa.md",
    "lideranca e oratoria": "melhorias_metodologicas_lideranca_e_oratoria.md",
    "lingua portuguesa": "melhorias_metodologicas_lingua_portuguesa.md",
    "portugues": "melhorias_metodologicas_lingua_portuguesa.md",
    "matematica": "melhorias_metodologicas_matematica.md",
    "orientacao de estudos": "melhorias_metodologicas_orientacao_estudos.md",
    "projeto de vida": "melhorias_metodologicas_projeto_de_vida.md",
    "quimica": "melhorias_metodologicas_quimica.md",
    "redacao e leitura": "melhorias_metodologicas_redacao_leitura.md",
    "leitura e redacao": "melhorias_metodologicas_redacao_leitura.md",
    "tecnologia e inovacao": "melhorias_metodologicas_tecnologia_e_inovacao.md",
}

MAPA_REFERENCIAS = {
    "matematica": (
        "ANÁLISE METODOLÓGICA - MATEMÁTICA.md",
        "analise_metodologica_matematica_ensino_medio_seduc_sp.md",
    ),
    "lingua portuguesa": REFERENCIA_PORTUGUES_GERAL,
    "portugues": REFERENCIA_PORTUGUES_GERAL,
    "redacao e leitura": REFERENCIA_LEITURA_REDACAO,
    "leitura e redacao": REFERENCIA_LEITURA_REDACAO,
    "ciencias": (
        "ANALISE_METODOLOGICA_CIENCIAS_EF_ANOS_FINAIS_3B.md",
        "CIÊNCIAS-6ANO_metodologias_ciencias_6ano_versao_final_completa_ajustada.docx",
        "ANÁLISE METODOLÓGICA - CIÊNCIAS 7º ANO.md",
    ),
    "ciencia": (
        "ANALISE_METODOLOGICA_CIENCIAS_EF_ANOS_FINAIS_3B.md",
        "CIÊNCIAS-6ANO_metodologias_ciencias_6ano_versao_final_completa_ajustada.docx",
        "ANÁLISE METODOLÓGICA - CIÊNCIAS 7º ANO.md",
    ),
    "geografia": "GEOGRAFIA-1EM_metodologia.docx",
    "arte": "ANÁLISE METODOLÓGICA - ARTE - ENSINO FUNDAMENTAL ANOS FINAIS.md",
    "artes": "ANÁLISE METODOLÓGICA - ARTE - ENSINO FUNDAMENTAL ANOS FINAIS.md",
    "historia": "ANÁLISE METODOLÓGICA - HISTÓRIA ENSINO FUNDAMENTAL.md",
    "projeto de vida": (
        REFERENCIA_PV_GERAL,
        REFERENCIA_PV_REFINADA_FUNDAMENTAL,
        REFERENCIA_PV_FUNDAMENTAL,
        REFERENCIA_PV_FUNDAMENTAL_ANOS_FINAIS,
    ),
    "ingles": "ANÁLISE METODOLÓGICA - INGLÊS ENSINO FUNDAMENTAL.md",
    "english": "ANÁLISE METODOLÓGICA - INGLÊS ENSINO FUNDAMENTAL.md",
    "orientacao de estudos": "ANÁLISE METODOLÓGICA PROFUNDA - ORIENTAÇÃO DE ESTUDOS.md",
    "educacao financeira": "EDUCAÇÃO FINANCEIRA-7ANO_METODOLOGIA.docx",
    "lideranca e oratoria_cdp": "LIDERANCA_E_ORATORIA_CDP_metodologia.docx",
    "cdp": (
        "metodologiacdp.docx",
        "HABILIDADES POR DISCIPLINA - EDUCAÇÃO DE JOVENS E ADULTOS (EJA).md",
        "HABILIDADES POR DISCIPLINA - EDUCAÇÃO DE JOVENS E ADULTOS (EJA).mdparte2.md",
    ),
    "eja": (
        "metodologiacdp.docx",
        "HABILIDADES POR DISCIPLINA - EDUCAÇÃO DE JOVENS E ADULTOS (EJA).md",
        "HABILIDADES POR DISCIPLINA - EDUCAÇÃO DE JOVENS E ADULTOS (EJA).mdparte2.md",
    ),
}


def normalizar_disciplina(texto: str = "") -> str:
    texto = str(texto or "").strip().lower()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    texto = texto.replace("º", "").replace("ª", "").replace("°", "")
    return re.sub(r"\s+", " ", texto)


def _eh_turma_fundamental(turma: str = "") -> bool:
    turma_norm = normalizar_disciplina(turma)
    return bool(re.search(r"\b(?:6|7|8|9)\s*(?:ano|a|b|c|d|e)?\b", turma_norm))


def _ano_turma(turma: str = "") -> int:
    turma_norm = normalizar_disciplina(turma)
    match = re.search(r"\b([1-9])\s*(?:ano|em|a|b|c|d|e)?\b", turma_norm)
    return int(match.group(1)) if match else 0


def _eh_portugues(disciplina: str = "") -> bool:
    disciplina_norm = normalizar_disciplina(disciplina)
    return "portugues" in disciplina_norm or "lingua portuguesa" in disciplina_norm


def _eh_projeto_vida(disciplina: str = "") -> bool:
    return "projeto de vida" in normalizar_disciplina(disciplina)


def _combinar_arquivos(*grupos: Iterable[str]) -> Tuple[str, ...]:
    vistos = set()
    combinados = []
    for grupo in grupos:
        for arquivo in grupo or ():
            if not arquivo or arquivo in vistos:
                continue
            vistos.add(arquivo)
            combinados.append(arquivo)
    return tuple(combinados)


def _arquivos_novos_para_disciplina(disciplina: str = "") -> Tuple[str, ...]:
    disciplina_norm = normalizar_disciplina(disciplina)
    for chave, arquivo in MAPA_REFERENCIAS_NOVAS.items():
        if chave in disciplina_norm:
            return (arquivo,)
    return ()


def _arquivos_para_disciplina(disciplina: str = "") -> Tuple[str, ...]:
    disciplina_norm = normalizar_disciplina(disciplina)
    for chave, arquivos in MAPA_REFERENCIAS.items():
        if chave in disciplina_norm:
            if isinstance(arquivos, str):
                return (arquivos,)
            return tuple(arquivos)
    return ()


def _limpar_markdown(texto: str) -> str:
    texto = re.sub(r"```.*?```", " ", texto or "", flags=re.DOTALL)
    texto = re.sub(r"#{1,6}\s*", "", texto)
    texto = re.sub(r"\*\*(.*?)\*\*", r"\1", texto)
    texto = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def _priorizar_analise_nova_markdown(texto: str) -> str:
    bruto = str(texto or "")
    titulo = ""
    match_titulo = re.search(r"(?m)^#\s+(.+)$", bruto)
    if match_titulo:
        titulo = match_titulo.group(1).strip()

    secoes_prioritarias = []
    for numero in ("4", "5", "6", "7", "8", "9", "10"):
        match = re.search(
            rf"(?ims)^##\s*{numero}\.\s+.*?(?=^##\s+\d+\.|\Z)",
            bruto,
        )
        if match:
            secoes_prioritarias.append(match.group(0).strip())

    secoes_contexto = []
    for numero in ("2", "3"):
        match = re.search(
            rf"(?ims)^##\s*{numero}\.\s+.*?(?=^##\s+\d+\.|\Z)",
            bruto,
        )
        if match:
            secoes_contexto.append(match.group(0).strip())

    partes = []
    if titulo:
        partes.append(f"# {titulo}")
    partes.extend(secoes_prioritarias)
    partes.extend(secoes_contexto)

    if not partes:
        return _limpar_markdown(bruto)
    return _limpar_markdown("\n\n".join(partes))


def _limpar_interdisciplinar(texto: str) -> str:
    texto = _limpar_markdown(texto)
    texto = re.sub(
        r"(?is)1\.\s*Aplicabilidade dos Padrões por Disciplina.*?(?=2\.\s*Adaptações Específicas por Disciplina)",
        "",
        texto,
    )
    texto = re.sub(r"(?is)3\.\s*Riscos de Confusão no Código Python.*", "", texto)
    texto = re.sub(r"(?im)^#+\s*", "", texto)
    return texto.strip()


def _reforcar_regras_do_sistema(texto: str) -> str:
    reforco = (
        "REGRAS FIXAS DO SISTEMA:\\n"
        "- Use esta biblioteca apenas como referência de estilo e qualidade, sem copiar trechos prontos.\\n"
        "- Priorize textos completos; não use reticências para encurtar frases em desenvolvimento, acompanhamento ou acessibilidade.\\n"
        "- Se precisar reduzir, reescreva a frase de forma mais curta e completa.\\n"
        "- Não invente técnicas pedagógicas; só cite técnicas quando estiverem explicitamente presentes nos slides.\\n"
        "- Em modalidade EJA, não cite nomes de técnicas pedagógicas; descreva as ações diretamente.\\n"
        "- Para HISTÓRIA: 'Pause e responda' é SEMPRE PROIBIDO, sem exceção.\\n"
        "- Respeite a ordem real dos slides enviados.\\n"
        "- Mantenha metodologia fluida, objetiva e adequada ao conteúdo da aula.\\n\\n"
    )
    return reforco + texto


def _carregar_referencia_interdisciplinar() -> str:
    caminho = resolver_caminho_referencia(REFERENCIA_INTERDISCIPLINAR)
    if not caminho:
        return ""

    texto = _limpar_interdisciplinar(caminho.read_text(encoding="utf-8", errors="ignore"))
    if not texto:
        return ""

    texto = (
        "REFERÊNCIA INTERDISCIPLINAR COMPLEMENTAR:\n"
        "- Use apenas para variar verbos, progressão pedagógica e linguagem de mediação.\n"
        "- Não use esta referência para criar tempos fixos, etapas inexistentes ou técnicas não presentes nos slides.\n"
        "- A referência específica da disciplina e a ordem real dos slides têm prioridade.\n\n"
        + texto
    )
    if len(texto) <= LIMITE_INTERDISCIPLINAR_CHARS:
        return texto
    return texto[:LIMITE_INTERDISCIPLINAR_CHARS].rsplit("\n", 1)[0].strip()


def _ler_docx(caminho: Path) -> str:
    try:
        import docx
        doc = docx.Document(caminho)
        texto_partes = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                texto_partes.append(txt)
        for table in doc.tables:
            for row in table.rows:
                celulas = []
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt and (not celulas or celulas[-1] != txt):
                        celulas.append(txt)
                linha_txt = " | ".join(celulas)
                if linha_txt:
                    texto_partes.append(linha_txt)
        return "\n".join(texto_partes)
    except Exception as e:
        return f"Referência {caminho.name} (erro ao ler: {e})"


def _ler_arquivos_referencia(arquivos: Iterable[str]) -> str:
    partes = []
    for arquivo in arquivos:
        caminho = resolver_caminho_referencia(arquivo)
        if not caminho:
            continue
        suffix = caminho.suffix.lower()
        if suffix == ".md":
            texto = caminho.read_text(encoding="utf-8", errors="ignore")
            if _caminho_em_pasta_oficial(caminho) and caminho.is_relative_to(
                REFERENCIAS_METODOLOGICAS_DIR.resolve()
            ):
                texto = _priorizar_analise_nova_markdown(texto)
            else:
                texto = _limpar_markdown(texto)
        elif suffix == ".docx":
            texto = _ler_docx(caminho)
        else:
            try:
                texto = caminho.read_text(encoding="utf-8", errors="ignore")
            except Exception:
                texto = caminho.name
        if texto:
            partes.append(texto)
    return "\n\n".join(partes)


def _buscar_metodologia_automatica(disciplina: str, turma: str) -> tuple[str, ...]:
    try:
        from core.helpers import normalizar_para_pasta, resolver_raiz_disciplina_pdfs
        import re
        
        disc_norm = normalizar_para_pasta(disciplina)
        pasta_disc = resolver_raiz_disciplina_pdfs(
            PDF_AULAS_DIR,
            disciplina,
            modalidade_eja="EJA" in disc_norm,
        )
        
        if not pasta_disc.exists():
            # Fallback flexível de nome de disciplina
            for d in Path(PDF_AULAS_DIR).iterdir():
                if d.is_dir() and normalizar_para_pasta(d.name) == disc_norm:
                    pasta_disc = d
                    break
                    
        if not pasta_disc.exists():
            return ()
            
        turma_norm = normalizar_para_pasta(turma)
        ano_str = ""
        m = re.search(r"(\d+)_?ANO", turma_norm)
        if m:
            ano_str = f"{m.group(1)}_ANO"
            
        encontrados = []
        for ext in ("*.docx", "*.md"):
            for arquivo in pasta_disc.rglob(f"METODOLOGIA{ext}"):
                caminho_str = str(arquivo).upper()
                if ano_str and ano_str in caminho_str:
                    encontrados.append(str(arquivo))
                elif not ano_str:
                    encontrados.append(str(arquivo))
                    
        return tuple(encontrados)
    except Exception as e:
        import logging
        logging.getLogger("PLANOS_LUAN").error(f"Erro na busca automatica de metodologia: {e}")
        return ()

@lru_cache(maxsize=32)
def diagnosticar_referencia_metodologica(
    disciplina: str = "", turma: str = ""
) -> DiagnosticoReferenciaMetodologica:
    arquivos_automaticos = _buscar_metodologia_automatica(disciplina, turma)
    arquivos_novos = _combinar_arquivos(arquivos_automaticos, _arquivos_novos_para_disciplina(disciplina))
    if _eh_portugues(disciplina) and _eh_turma_fundamental(turma):
        arquivos_padrao = (REFERENCIA_PORTUGUES_FUNDAMENTAL, REFERENCIA_PORTUGUES_GERAL)
    elif _eh_portugues(disciplina):
        arquivos_padrao = (REFERENCIA_PORTUGUES_MEDIO, REFERENCIA_PORTUGUES_GERAL)
    elif _eh_projeto_vida(disciplina) and _ano_turma(turma) == 7:
        arquivos_padrao = (
            REFERENCIA_PV_GERAL,
            REFERENCIA_PV_REFINADA_FUNDAMENTAL,
            REFERENCIA_PV_7_ANO,
            REFERENCIA_PV_FUNDAMENTAL,
            REFERENCIA_PV_FUNDAMENTAL_ANOS_FINAIS,
        )
    elif _eh_projeto_vida(disciplina):
        arquivos_padrao = (
            REFERENCIA_PV_GERAL,
            REFERENCIA_PV_REFINADA_FUNDAMENTAL,
            REFERENCIA_PV_FUNDAMENTAL,
            REFERENCIA_PV_FUNDAMENTAL_ANOS_FINAIS,
        )
    else:
        arquivos_padrao = _arquivos_para_disciplina(disciplina)
    
    disc_norm = normalizar_disciplina(disciplina)
    if "redacao" in disc_norm or "leitura" in disc_norm:
        arquivos = _combinar_arquivos(arquivos_padrao, arquivos_novos)
    else:
        arquivos = arquivos_novos if arquivos_novos else arquivos_padrao
        
    if not arquivos:
        return DiagnosticoReferenciaMetodologica("", (), (), ())

    caminhos_encontrados = []
    arquivos_ausentes = []
    for arquivo in arquivos:
        caminho = resolver_caminho_referencia(arquivo)
        if caminho:
            caminhos_encontrados.append(caminho)
        else:
            arquivos_ausentes.append(str(arquivo))

    texto = _ler_arquivos_referencia(caminhos_encontrados)
    if not texto:
        aviso = (
            "Referência metodológica oficial não encontrada para "
            f"{disciplina or 'a disciplina selecionada'}"
            + (f" ({turma})" if turma else "")
            + ". O plano foi gerado sem essa referência complementar."
        )
        return DiagnosticoReferenciaMetodologica(
            "",
            tuple(str(arquivo) for arquivo in arquivos),
            tuple(str(caminho) for caminho in caminhos_encontrados),
            tuple(arquivos_ausentes),
            aviso,
        )

    interdisciplinar = _carregar_referencia_interdisciplinar()
    if interdisciplinar:
        limite_texto_principal = max(2600, LIMITE_REFERENCIA_CHARS - len(interdisciplinar) - 400)
        if len(texto) > limite_texto_principal:
            texto = texto[:limite_texto_principal].rsplit("\n", 1)[0].strip()
        texto = texto + "\n\n" + interdisciplinar

    texto = _reforcar_regras_do_sistema(texto)
    if len(texto) <= LIMITE_REFERENCIA_CHARS:
        texto_final = texto
    else:
        texto_final = texto[:LIMITE_REFERENCIA_CHARS].rsplit("\n", 1)[0].strip()

    return DiagnosticoReferenciaMetodologica(
        texto_final,
        tuple(str(arquivo) for arquivo in arquivos),
        tuple(str(caminho) for caminho in caminhos_encontrados),
        tuple(arquivos_ausentes),
    )


@lru_cache(maxsize=32)
def carregar_referencia_metodologica(disciplina: str = "", turma: str = "") -> str:
    """Mantem a API legada que fornece somente o texto da referencia."""
    return diagnosticar_referencia_metodologica(disciplina, turma).texto


def listar_referencias_disponiveis() -> Dict[str, str]:
    disponiveis = {}
    chaves = set(MAPA_REFERENCIAS) | set(MAPA_REFERENCIAS_NOVAS)
    for disciplina in sorted(chaves):
        lista = _combinar_arquivos(
            _arquivos_novos_para_disciplina(disciplina),
            _arquivos_para_disciplina(disciplina),
        )
        caminhos = []
        for arquivo in lista:
            caminho = resolver_caminho_referencia(arquivo)
            if caminho:
                caminhos.append(str(caminho))
        if caminhos:
            disponiveis[disciplina] = " | ".join(caminhos)
    return disponiveis
