from __future__ import annotations

from io import BytesIO
import re

from docx import Document

from core.lote import (
    acessibilidade_cdp_contextual,
    acompanhamento_cdp_contextual,
    formatar_material_cdp_contextual,
    limpar_tema_cdp_contextual,
    metodologia_cdp_contextual,
    normalizar_texto_lote,
    perfil_disciplina,
)
from docx_generator.preencher import (
    _eh_tabela_aulas,
    _indices_linha_aula,
    _preencher_celula_lista,
    _preencher_celula_metodologia,
    _preencher_celula_tema_material,
)


DISCIPLINAS_CDP = {
    "História": ["história", "historia", "tempo histórico", "fonte histórica", "colonização", "império", "revolução"],
    "Geografia": ["geografia", "espaço agrário", "cidade", "território", "paisagem", "mapa", "cartografia", "onu"],
    "Sociologia": ["sociologia", "relações de classe", "desigualdades", "relações étnico", "relações sociais de gênero"],
    "Liderança e Oratória": ["liderança e oratória", "oratoria", "persuasão", "falácias", "discurso", "storytelling"],
    "Língua Portuguesa": ["língua portuguesa", "português", "gênero textual", "produção textual", "interpretação"],
    "Matemática": ["matemática", "equação", "função", "fração", "plano cartesiano", "porcentagem"],
    "Ciências": ["ciências", "sistema digestório", "célula", "vírus", "ecossistema", "genética"],
    "Arte": ["arte", "obra", "artista", "pintura", "linguagem artística"],
}

TERMOS_PROIBIDOS_CDP = [
    "virem e conversem",
    "todo mundo escreve",
    "com suas palavras",
    "hora da leitura",
    "um passo de cada vez",
    "de olho no modelo",
    "pausa produtiva",
    "youtube",
    "internet",
    "celular",
    "aplicativo",
    "plataforma",
    "projetor",
    "datashow",
    "assistir ao vídeo",
    "assistir ao video",
]


def _texto_docx_completo(doc: Document) -> str:
    partes: list[str] = []
    for paragrafo in doc.paragraphs:
        if paragrafo.text:
            partes.append(paragrafo.text)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if celula.text:
                    partes.append(celula.text)
    return "\n".join(partes)


def _detectar_disciplina_cdp(texto: str) -> str:
    base = normalizar_texto_lote(texto)
    pontuacoes: dict[str, int] = {}
    for disciplina, termos in DISCIPLINAS_CDP.items():
        score = 0
        for termo in termos:
            termo_norm = normalizar_texto_lote(termo)
            if not termo_norm:
                continue
            if re.search(rf"(?<!\w){re.escape(termo_norm)}(?!\w)", base):
                score += 10
            elif termo_norm in base:
                score += 3
        pontuacoes[disciplina] = score
    melhor, score = max(pontuacoes.items(), key=lambda item: item[1])
    return melhor if score > 0 else "Geral"


def _extrair_tema_material(texto: str) -> str:
    bruto = re.sub(r"\s+", " ", str(texto or "")).strip()
    if not bruto:
        return "Conteúdo da aula"

    linhas = [linha.strip(" -:–—") for linha in str(texto or "").splitlines() if linha.strip()]
    candidatas: list[str] = []
    for linha in linhas:
        linha_limpa = re.sub(r"^\s*TEMA\s*:\s*", "", linha, flags=re.I).strip(" -:–—")
        linha_limpa = re.sub(r"^\s*AULA\s*\d+\s*[-:–—]?\s*", "", linha_limpa, flags=re.I).strip(" -:–—")
        linha_limpa = re.sub(
            r"^\s*(MATEMÁTICA|HISTÓRIA|GEOGRAFIA|SOCIOLOGIA|LIDERANÇA E ORATÓRIA|LÍNGUA PORTUGUESA|CIÊNCIAS|ARTE)\s*[-:–—]?\s*",
            "",
            linha_limpa,
            flags=re.I,
        ).strip(" -:–—")
        if linha_limpa and linha_limpa.upper() not in {"TEMA"}:
            candidatas.append(linha_limpa)

    if not candidatas:
        return bruto.split(" - ", 1)[-1].strip() if " - " in bruto else bruto
    return max(candidatas, key=len)


def _limpar_resultado_cdp(texto: str) -> str:
    from core.qualidade_metodologica import sanitizar_texto_cdp_estrito
    return sanitizar_texto_cdp_estrito(texto)


def _metodologia_generica_cdp(disciplina: str, tema: str, aprendizagem: str) -> list[str]:
    perfil = perfil_disciplina(disciplina)
    tema_frase = limpar_tema_cdp_contextual(tema, disciplina)

    if perfil == "historia":
        texto = (
            f"O professor inicia a aula situando {tema_frase} no tempo e no espaço, registrando na lousa os acontecimentos, sujeitos históricos e conceitos centrais. "
            "Em seguida, realiza leitura mediada do material impresso e orienta a análise de textos, imagens, charges ou fontes históricas, quando aparecerem na aula. "
            "Os estudantes registram as ideias principais no caderno, respondem às questões propostas e participam da correção coletiva, retomando causas, consequências e relações entre passado e presente."
        )
    elif perfil == "geografia":
        texto = (
            f"O professor inicia a aula relacionando {tema_frase} a situações concretas do espaço vivido pelos estudantes. "
            "Na lousa, organiza palavras-chave e conduz a leitura de mapas, gráficos, tabelas, imagens ou textos do material impresso, destacando título, legenda, fonte, localização e informações comparáveis. "
            "Os estudantes registram observações no caderno, respondem às atividades e participam da correção coletiva, justificando as respostas com base no material analisado."
        )
    elif perfil == "sociologia":
        texto = (
            f"O professor inicia a aula apresentando {tema_frase} por meio da pergunta, charge, texto ou situação presente no material impresso, sem exigir exposição pessoal dos estudantes. "
            "Depois, conduz debate guiado para diferenciar opinião, exemplo cotidiano e conceito sociológico, registrando na lousa as ideias centrais. "
            "Os estudantes fazem síntese no caderno, respondem às questões e participam de correção coletiva com retomada do vocabulário sociológico essencial."
        )
    elif perfil == "lideranca_oratoria":
        texto = (
            f"O professor inicia a aula apresentando {tema_frase} e realizando leitura mediada do texto, tirinha, exemplo de discurso ou situação comunicativa do material impresso. "
            "Na lousa, organiza um roteiro de análise com ideia central, público, argumento, estratégia utilizada, efeito produzido e responsabilidade discursiva. "
            "Os estudantes registram respostas curtas no caderno e socializam oralmente de forma mediada, com fechamento voltado à comunicação clara, ética e responsável."
        )
    else:
        texto = " ".join(metodologia_cdp_contextual(perfil, "", tema_frase, aprendizagem, 0))

    return [_limpar_resultado_cdp(texto)]


def _lista_generica_cdp(disciplina: str, tema: str, tipo: str) -> list[str]:
    perfil = perfil_disciplina(disciplina)
    if tipo == "acompanhamento":
        if perfil == "historia":
            return [
                "☑ Verificar se o estudante identifica contexto, sujeitos históricos, causas e consequências do tema estudado.",
                "☑ Observar se utiliza textos, imagens ou fontes do material impresso para justificar as respostas.",
                "☑ Acompanhar os registros no caderno e retomar dúvidas durante a correção coletiva.",
            ]
        if perfil == "geografia":
            return [
                "☑ Verificar se o estudante lê informações de mapas, gráficos, tabelas, imagens ou textos do material impresso.",
                "☑ Observar se compara dados, lugares ou fenômenos e justifica respostas com base no conteúdo trabalhado.",
                "☑ Conferir se relaciona o tema a território, paisagem, cidade, campo, população ou ambiente.",
            ]
        if perfil == "sociologia":
            return [
                "☑ Verificar se o estudante diferencia exemplo cotidiano, opinião e conceito sociológico.",
                "☑ Observar se relaciona o material analisado aos temas sociais estudados.",
                "☑ Conferir se os registros apresentam justificativas coerentes e vocabulário adequado.",
            ]
        if perfil == "lideranca_oratoria":
            return [
                "☑ Verificar se o estudante identifica intenção comunicativa, argumento e estratégia discursiva.",
                "☑ Observar clareza, respeito e responsabilidade nas respostas orais e escritas.",
                "☑ Conferir se aplica o roteiro de análise do discurso na atividade proposta.",
            ]
    else:
        if perfil == "historia":
            return [
                "☑ Organizar o conteúdo em linha do tempo, lista de causas e consequências ou esquema simples na lousa.",
                "☑ Explicar o vocabulário histórico antes da atividade escrita.",
                "☑ Permitir resposta oral mediada antes do registro escrito quando necessário.",
            ]
        if perfil == "geografia":
            return [
                "☑ Realizar leitura guiada de mapas, gráficos, tabelas ou imagens, destacando título, legenda, fonte e comparação principal.",
                "☑ Reproduzir na lousa os dados essenciais para apoiar a leitura do material impresso.",
                "☑ Permitir registros por tópicos, setas ou frases curtas, mantendo a interpretação geográfica central.",
            ]
        if perfil == "sociologia":
            return [
                "☑ Fazer leitura pausada de charges, textos ou situações, evitando exposição pessoal dos estudantes.",
                "☑ Registrar na lousa conceitos-chave para consulta durante a atividade.",
                "☑ Possibilitar respostas por tópicos ou frases curtas antes da versão final no caderno.",
            ]
        if perfil == "lideranca_oratoria":
            return [
                "☑ Disponibilizar roteiro simples na lousa para análise do discurso.",
                "☑ Retomar vocabulário como persuasão, argumento, credibilidade, falácia e responsabilidade.",
                "☑ Permitir socialização oral breve e mediada, sem exposição constrangedora.",
            ]
    return acompanhamento_cdp_contextual(perfil, tema)[:3] if tipo == "acompanhamento" else acessibilidade_cdp_contextual(perfil, tema)[:3]


def reescrever_docx_cdp_ensino_medio(docx_bytes: bytes) -> tuple[bytes, dict[str, object]]:
    doc = Document(BytesIO(docx_bytes))
    texto_doc = _texto_docx_completo(doc)
    disciplina_doc = _detectar_disciplina_cdp(texto_doc)
    linhas_reescritas = 0
    temas: list[str] = []
    disciplinas: list[str] = []

    for tabela in doc.tables:
        if not _eh_tabela_aulas(tabela):
            continue
        cabecalho = tabela.rows[0] if tabela.rows else None
        for linha in tabela.rows[1:]:
            indices = _indices_linha_aula(linha, cabecalho)
            if not indices:
                continue

            texto_material = linha.cells[indices["material"]].text
            texto_aprendizagem = linha.cells[indices["aprendizagem"]].text
            disciplina = _detectar_disciplina_cdp(f"{texto_doc}\n{texto_material}\n{texto_aprendizagem}") or disciplina_doc
            if disciplina == "Geral":
                disciplina = disciplina_doc or "Geral"
            tema = limpar_tema_cdp_contextual(_extrair_tema_material(texto_material), disciplina)
            if not tema:
                continue

            metodologia = _metodologia_generica_cdp(disciplina, tema, texto_aprendizagem)
            acompanhamento = _lista_generica_cdp(disciplina, tema, "acompanhamento")
            acessibilidade = _lista_generica_cdp(disciplina, tema, "acessibilidade")

            _preencher_celula_tema_material(linha.cells[indices["material"]], formatar_material_cdp_contextual(tema, disciplina))
            _preencher_celula_metodologia(linha.cells[indices["desenvolvimento"]], metodologia)
            _preencher_celula_lista(linha.cells[indices["acompanhamento"]], acompanhamento[:3])
            _preencher_celula_lista(linha.cells[indices["acessibilidade"]], acessibilidade[:3])

            linhas_reescritas += 1
            temas.append(tema)
            disciplinas.append(disciplina)

    saida = BytesIO()
    doc.save(saida)
    saida.seek(0)
    return saida.getvalue(), {"linhas_reescritas": linhas_reescritas, "temas": temas, "disciplinas": disciplinas}


# Compatibilidade com o nome antigo. Agora ele não força Matemática.
def reescrever_docx_cdp_contextual_matematica(docx_bytes: bytes) -> tuple[bytes, dict[str, object]]:
    return reescrever_docx_cdp_ensino_medio(docx_bytes)
