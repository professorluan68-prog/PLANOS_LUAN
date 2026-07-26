from __future__ import annotations

import re
import shutil
import unicodedata
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document

from config import PASTA_PLANOS_PROFESSORES, PLANOS_FEITOS_DIR

MESES_PT = {
    "JANEIRO": 1,
    "FEVEREIRO": 2,
    "MARCO": 3,
    "MARÇO": 3,
    "ABRIL": 4,
    "MAIO": 5,
    "JUNHO": 6,
    "JULHO": 7,
    "AGOSTO": 8,
    "SETEMBRO": 9,
    "OUTUBRO": 10,
    "NOVEMBRO": 11,
    "DEZEMBRO": 12,
}
DIAS_PT = ["Segunda", "Terca", "Quarta", "Quinta", "Sexta", "Sabado", "Domingo"]


def _norm(texto: str) -> str:
    return re.sub(r"\s+", " ", str(texto or "")).strip()


def _norm_chave_mescla(texto: str) -> str:
    valor = unicodedata.normalize("NFKD", _norm(texto).upper())
    valor = "".join(ch for ch in valor if not unicodedata.combining(ch))
    valor = re.sub(r"[^A-Z0-9]+", " ", valor)
    return " ".join(valor.split())


def _safe_filename_part(texto: str) -> str:
    texto = _norm(texto)
    texto = re.sub(r'[\\/:*?"<>|]+', "-", texto)
    texto = re.sub(r"\s*-\s*", " - ", texto)
    return texto.strip(" .-") or "NAO INFORMADO"


def nome_padronizado_plano(disciplina: str, turma: str, mes: str = "") -> str:
    from datetime import date
    MESES_ABREV = {
        1: "JANEIRO", 2: "FEVEREIRO", 3: "MARCO", 4: "ABRIL",
        5: "MAIO", 6: "JUNHO", 7: "JULHO", 8: "AGOSTO",
        9: "SETEMBRO", 10: "OUTUBRO", 11: "NOVEMBRO", 12: "DEZEMBRO",
    }
    if mes:
        mes_upper = mes.strip().upper()
    else:
        mes_upper = MESES_ABREV.get(date.today().month, "MES")
    return f"PLANO_{mes_upper} - {_safe_filename_part(disciplina)} - {_safe_filename_part(turma)}.docx"



def _mes_para_nome_plano(origem_path: Path | None, professor: str, mes: str = "") -> str:
    mes_informado = _norm(mes).upper()
    if mes_informado:
        return mes_informado
    if origem_path and origem_path.exists():
        try:
            info = extrair_info_plano(origem_path, professor)
            mes_origem = _norm(str(info.get("mes") or "")).upper()
            if mes_origem:
                return mes_origem
        except Exception:
            pass
    return ""


def _texto_tabela(tabela) -> str:
    return " ".join(celula.text.upper() for linha in tabela.rows for celula in linha.cells)


def _eh_cabecalho_plano(tabela) -> bool:
    texto = _texto_tabela(tabela)
    return len(tabela.rows) >= 4 and "PLANO DE AULAS" in texto and "PROFESSOR" in texto


def _eh_tabela_aulas(tabela) -> bool:
    if not tabela.rows:
        return False
    texto = " ".join(celula.text.upper() for celula in tabela.rows[0].cells)
    return "AULA" in texto and "APRENDIZAGEM" in texto and "DESENVOLVIMENTO" in texto


def _cabecalho_info(tabela, professor_pasta: str) -> dict[str, str]:
    info = {
        "professor": professor_pasta,
        "disciplina": "",
        "turma": "",
        "mes": "",
        "aulas_semana": "",
        "componente_curricular": "",
        "semana": "",
    }
    if not _eh_cabecalho_plano(tabela):
        return info
    dados = tabela.rows[2].cells
    if len(dados) >= 9:
        info["professor"] = _norm(dados[2].text) or professor_pasta
        info["disciplina"] = _norm(dados[3].text)
        info["componente_curricular"] = _norm(dados[3].text)
        info["turma"] = _norm(dados[6].text)
        info["mes"] = _norm(dados[7].text).upper()
    semana = tabela.rows[3].cells
    if len(semana) >= 4:
        info["semana"] = _norm(semana[1].text)
        info["aulas_semana"] = _norm(semana[3].text)
    return info


def _mes_numero(mes: str) -> int:
    return MESES_PT.get(_norm(mes).upper(), date.today().month)


def _data_por_ddmm(texto: str, mes: str = ""):
    match = re.search(r"\b(\d{1,2})/(\d{1,2})\b", str(texto or ""))
    if not match:
        return None
    dia = int(match.group(1))
    mes_num = int(match.group(2) or _mes_numero(mes))
    try:
        return date(date.today().year, mes_num, dia)
    except ValueError:
        return None


def _partes_data_horario(texto: str) -> dict[str, object] | None:
    linhas = [_norm(linha) for linha in str(texto or "").splitlines() if _norm(linha)]
    if not linhas:
        return None
    primeira_linha = linhas[0]
    data_aula = _data_por_ddmm(primeira_linha)
    if not data_aula:
        return None
    partes_restantes = []
    match_primeira = re.match(r"^\s*\d{1,2}/\d{1,2}\s*(.*)$", primeira_linha)
    if match_primeira:
        restante = _norm(match_primeira.group(1))
        if restante:
            partes_restantes.append(restante)
    for linha in linhas[1:]:
        partes_restantes.append(linha)
    texto_restante = _norm(" ".join(parte for parte in partes_restantes if parte))
    aula = ""
    horario = ""
    if texto_restante:
        if re.search(r"\d{1,2}h", texto_restante, flags=re.I):
            horario = texto_restante
        else:
            aula = texto_restante
    return {"data": data_aula, "aula": aula, "horario": horario}


def _extrair_datas_horarios(tabela) -> list[dict[str, object]]:
    itens = []
    if not _eh_tabela_aulas(tabela):
        return itens
    for linha in tabela.rows[1:]:
        item = _partes_data_horario(linha.cells[0].text)
        if item:
            itens.append(item)
    return itens


def extrair_datas_horarios_de_bytes(docx_bytes: bytes) -> list[dict[str, object]]:
    if not docx_bytes:
        return []
    doc = Document(BytesIO(docx_bytes))
    itens: list[dict[str, object]] = []
    for tabela in doc.tables:
        itens.extend(_extrair_datas_horarios(tabela))
    return itens


def _texto_horario_grade(item: dict[str, object]) -> str:
    horario = _norm(str(item.get("horario") or ""))
    aula = _norm(str(item.get("aula") or ""))
    if horario and aula:
        return f"{horario} ({aula})"
    return horario or aula


def _resumir_grade_semanal(datas_horarios: list[dict[str, object]]) -> tuple[str, str]:
    itens_ordenados = sorted(
        datas_horarios,
        key=lambda item: (
            item["data"],
            _texto_horario_grade(item),
        ),
    )
    grade = []
    vistos = set()
    for item in itens_ordenados:
        data_aula = item.get("data")
        if not hasattr(data_aula, "weekday"):
            continue
        horario_grade = _texto_horario_grade(item)
        chave = (data_aula.weekday(), horario_grade)
        if chave in vistos:
            continue
        vistos.add(chave)
        grade.append(
            {
                "dia_semana": DIAS_PT[data_aula.weekday()],
                "horario": horario_grade,
                "indice_dia": data_aula.weekday(),
            }
        )

    grade.sort(key=lambda item: (item["indice_dia"], item["horario"]))
    dias = " - ".join(item["dia_semana"] for item in grade if item["dia_semana"])
    horarios = ", ".join(item["horario"] for item in grade if item["horario"])
    return dias, horarios


def extrair_info_plano(caminho: Path, professor_pasta: str) -> dict[str, object]:
    doc = Document(caminho)
    info = {
        "professor": professor_pasta,
        "disciplina": "",
        "turma": "",
        "aulas_semana": "",
        "componente_curricular": "",
        "dia_semana": "",
        "horario": "",
        "datas_horarios": [],
        "arquivo": str(caminho),
    }

    ultimo_cabecalho = None
    datas_horarios: list[dict[str, object]] = []
    for tabela in doc.tables:
        if _eh_cabecalho_plano(tabela):
            ultimo_cabecalho = _cabecalho_info(tabela, professor_pasta)
            if not info["disciplina"]:
                info.update({k: ultimo_cabecalho.get(k, "") for k in ["professor", "disciplina", "turma", "aulas_semana", "componente_curricular"]})
        elif _eh_tabela_aulas(tabela):
            datas_horarios.extend(_extrair_datas_horarios(tabela))

    if datas_horarios:
        info["datas_horarios"] = datas_horarios
        info["dia_semana"], info["horario"] = _resumir_grade_semanal(datas_horarios)
    if ultimo_cabecalho and not info["aulas_semana"]:
        info["aulas_semana"] = ultimo_cabecalho.get("aulas_semana", "")
    return info


def _registrar_info_professor(
    resultado: dict[str, dict[str, list[dict[str, object]]]],
    caminho: Path,
    professor_pasta: str,
) -> None:
    try:
        info = extrair_info_plano(caminho, professor_pasta)
    except Exception:
        return

    professor = str(info.get("professor") or professor_pasta)
    disciplina = str(info.get("disciplina") or "").strip()
    turma = str(info.get("turma") or "").strip()
    if not disciplina or not turma:
        return

    resultado.setdefault(professor, {"disciplinas": []})["disciplinas"].append(
        {
            "disciplina": disciplina,
            "turma": turma,
            "dia_semana": str(info.get("dia_semana") or ""),
            "horario": str(info.get("horario") or ""),
            "aulas_semana": str(info.get("aulas_semana") or ""),
            "componente_curricular": str(info.get("componente_curricular") or ""),
            "datas_horarios": info.get("datas_horarios") or [],
            "arquivo": str(info.get("arquivo") or caminho),
        }
    )


def carregar_professores_dos_planos(base_dir: Path = PASTA_PLANOS_PROFESSORES) -> dict[str, dict[str, list[dict[str, object]]]]:
    resultado: dict[str, dict[str, list[dict[str, object]]]] = {}

    if base_dir.exists():
        for pasta_professor in sorted(p for p in base_dir.iterdir() if p.is_dir() and not p.name.startswith("_")):
            for caminho in sorted(pasta_professor.glob("*.docx")):
                _registrar_info_professor(resultado, caminho, pasta_professor.name)

    # Fallback importante: quando a pasta principal dos modelos não existe,
    # usamos os planos já gerados para recuperar grade semanal e contexto.
    planos_feitos_dir = Path(PLANOS_FEITOS_DIR)
    if planos_feitos_dir.exists():
        for caminho in sorted(planos_feitos_dir.rglob("*.docx")):
            if caminho.name.startswith("~$"):
                continue
            try:
                relativo = caminho.relative_to(planos_feitos_dir)
            except ValueError:
                continue
            if len(relativo.parts) < 3:
                continue
            professor_pasta = relativo.parts[0]
            _registrar_info_professor(resultado, caminho, professor_pasta)

    return resultado


def _arquivos_modelo_vinculados() -> list[Path]:
    try:
        from core.database import listar_vinculos_professores
    except Exception:
        return []
    caminhos = []
    for vinculo in listar_vinculos_professores():
        arquivo = str(vinculo.get("arquivo_modelo") or vinculo.get("arquivo") or "").strip()
        if arquivo:
            caminhos.append(Path(arquivo))
    return caminhos


def diagnosticar_modelos_professores(base_dir: Path = PASTA_PLANOS_PROFESSORES) -> dict[str, object]:
    diagnostico: dict[str, object] = {
        "base_dir": str(base_dir),
        "total_docx": 0,
        "lidos_ok": 0,
        "professores_pasta": 0,
        "erros_leitura": [],
        "sem_disciplina_turma": [],
        "sem_datas_horarios": [],
        "duplicidades": [],
    }
    arquivos_para_ler = sorted(dict.fromkeys(_arquivos_modelo_vinculados()))
    diagnostico["base_dir"] = "Modelos vinculados no banco"

    chaves: dict[tuple[str, str, str], list[str]] = {}
    diagnostico["professores_pasta"] = len({caminho.parent for caminho in arquivos_para_ler})

    for caminho in arquivos_para_ler:
        pasta_professor = caminho.parent
        diagnostico["total_docx"] = int(diagnostico["total_docx"]) + 1
        if not caminho.exists():
            diagnostico["erros_leitura"].append(
                {
                    "professor": pasta_professor.name,
                    "arquivo": str(caminho),
                    "erro": "Arquivo vinculado nao encontrado.",
                }
            )
            continue
        try:
            info = extrair_info_plano(caminho, pasta_professor.name)
        except Exception as exc:
            diagnostico["erros_leitura"].append(
                {
                    "professor": pasta_professor.name,
                    "arquivo": caminho.name,
                    "erro": str(exc),
                }
            )
            continue

        diagnostico["lidos_ok"] = int(diagnostico["lidos_ok"]) + 1
        professor = str(info.get("professor") or pasta_professor.name).strip()
        disciplina = str(info.get("disciplina") or "").strip()
        turma = str(info.get("turma") or "").strip()
        registro = {
            "professor": professor or pasta_professor.name,
            "disciplina": disciplina,
            "turma": turma,
            "arquivo": str(caminho),
        }

        if not disciplina or not turma:
            diagnostico["sem_disciplina_turma"].append(registro)
            continue

        if not info.get("datas_horarios"):
            diagnostico["sem_datas_horarios"].append(registro)

        chaves.setdefault((professor, disciplina, turma), []).append(str(caminho))

    diagnostico["duplicidades"] = [
        {
            "professor": professor,
            "disciplina": disciplina,
            "turma": turma,
            "arquivos": arquivos,
        }
        for (professor, disciplina, turma), arquivos in sorted(chaves.items())
        if len(arquivos) > 1
    ]
    return diagnostico


def criar_ou_atualizar_modelo_professor(
    professor: str,
    disciplina: str,
    turma: str,
    origem: str = "",
    aulas_semana: str = "",
    componente_curricular: str = "",
    mes: str = "",
    base_dir: Path = PASTA_PLANOS_PROFESSORES,
) -> str:
    pasta_professor = base_dir / _safe_filename_part(professor)
    pasta_professor.mkdir(parents=True, exist_ok=True)

    origem_path = Path(origem) if origem else None
    mes_arquivo = _mes_para_nome_plano(origem_path, professor, mes)
    destino = pasta_professor / nome_padronizado_plano(disciplina, turma, mes_arquivo)

    if origem_path and origem_path.exists() and origem_path.parent == pasta_professor and origem_path.resolve() != destino.resolve() and not destino.exists():
        origem_path.rename(destino)
    elif origem_path and origem_path.exists() and origem_path.resolve() != destino.resolve():
        shutil.copy2(origem_path, destino)
    elif not destino.exists() and origem_path and origem_path.exists():
        shutil.copy2(origem_path, destino)
    elif not destino.exists():
        template = Path(__file__).resolve().parent.parent / "templates" / "MODELOEGLE.docx"
        if template.exists():
            shutil.copy2(template, destino)
        else:
            raise FileNotFoundError("Nao foi encontrado um DOCX base para criar o modelo do professor.")

    doc = Document(destino)
    for tabela in doc.tables:
        if not _eh_cabecalho_plano(tabela):
            continue
        dados = tabela.rows[2].cells
        if len(dados) >= 9:
            dados[2].text = professor
            dados[3].text = componente_curricular or disciplina
            dados[6].text = turma
            if mes_arquivo:
                dados[7].text = mes_arquivo
        semana = tabela.rows[3].cells
        if aulas_semana and len(semana) >= 4:
            semana[3].text = str(aulas_semana)
    doc.save(destino)
    return str(destino)


def atualizar_cabecalho_modelo_professor(
    caminho: str,
    professor: str,
    disciplina: str,
    turma: str,
    aulas_semana: str = "",
    componente_curricular: str = "",
) -> str:
    destino = Path(caminho)
    if not destino.exists():
        raise FileNotFoundError("Modelo DOCX nao encontrado.")

    doc = Document(destino)
    atualizou = False
    for tabela in doc.tables:
        if not _eh_cabecalho_plano(tabela):
            continue
        dados = tabela.rows[2].cells
        if len(dados) >= 9:
            dados[2].text = professor
            dados[3].text = componente_curricular or disciplina
            dados[6].text = turma
            atualizou = True
        semana = tabela.rows[3].cells
        if aulas_semana and len(semana) >= 4:
            semana[3].text = str(aulas_semana)
    if not atualizou:
        raise ValueError("O DOCX nao possui cabecalho de plano reconhecido.")
    doc.save(destino)
    return str(destino)


def mesclar_professores(base: dict, planos: dict) -> dict:
    mesclado = {prof: {"disciplinas": list(dados.get("disciplinas", []))} for prof, dados in (base or {}).items()}
    for prof, dados in (planos or {}).items():
        mesclado.setdefault(prof, {"disciplinas": []})
        for item in dados.get("disciplinas", []):
            chave = (
                _norm_chave_mescla(item.get("disciplina", "")),
                _norm_chave_mescla(item.get("turma", "")),
            )
            existente = next(
                (
                    d
                    for d in mesclado[prof]["disciplinas"]
                    if (
                        _norm_chave_mescla(d.get("disciplina", "")),
                        _norm_chave_mescla(d.get("turma", "")),
                    ) == chave
                ),
                None,
            )
            if existente:
                for campo in (
                    "arquivo",
                    "datas_horarios",
                    "componente_curricular",
                    "dia_semana",
                    "horario",
                    "aulas_semana",
                ):
                    if not existente.get(campo) and item.get(campo):
                        existente[campo] = item.get(campo)
                continue
            mesclado[prof]["disciplinas"].append(item)
    return mesclado
