import re
import json
from io import BytesIO
from pathlib import Path
from docx import Document

from config import REGISTRO_PROXIMA_GERACAO_PATH


_PADRAO_AULA = re.compile(r"\bAULA\s*(\d+)\b", re.IGNORECASE)
_PADRAO_BIMESTRE = re.compile(r"\b([1-4])\s*(?:[º°]|[oa])?\s*BIMESTRE\b", re.IGNORECASE)


def _numero_bimestre(texto: str) -> int:
    correspondencia = _PADRAO_BIMESTRE.search(str(texto or ""))
    return int(correspondencia.group(1)) if correspondencia else 0


def _bimestres_do_cabecalho(documento: Document) -> set[int]:
    """Localiza o valor do bimestre no cabeçalho, não em notas das aulas."""
    bimestres = set()
    for tabela in documento.tables:
        for indice_linha, linha in enumerate(tabela.rows[:-1]):
            for indice_coluna, celula in enumerate(linha.cells):
                if str(celula.text or "").strip().upper() != "BIMESTRE":
                    continue
                for linha_dados in tabela.rows[indice_linha + 1:]:
                    numero = _numero_bimestre(linha_dados.cells[indice_coluna].text)
                    if numero:
                        bimestres.add(numero)
                        break
    return bimestres


def detectar_ultima_aula_de_docx_bytes(docx_bytes: bytes, bimestre: str = "") -> int:
    """
    Analisa os bytes de um arquivo .docx para extrair o número máximo de aula gerado.
    """
    if not docx_bytes:
        return 0

    try:
        documento = Document(BytesIO(docx_bytes))
        bimestre_esperado = _numero_bimestre(bimestre)
        bimestres_cabecalho = _bimestres_do_cabecalho(documento)
        if bimestre_esperado and bimestres_cabecalho and bimestre_esperado not in bimestres_cabecalho:
            return 0

        aulas_detectadas = []
        textos = [paragrafo.text for paragrafo in documento.paragraphs]
        celulas_vistas = set()
        for tabela in documento.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    chave_celula = celula._tc
                    if chave_celula not in celulas_vistas:
                        celulas_vistas.add(chave_celula)
                        textos.append(celula.text)

        for texto in textos:
            aulas_detectadas.extend(
                int(correspondencia.group(1))
                for correspondencia in _PADRAO_AULA.finditer(texto or "")
            )

        if aulas_detectadas:
            return max(aulas_detectadas)
    except Exception:
        pass

    return 0

def obter_aula_parada_do_json(professor: str, disciplina: str, turma: str, bimestre: str = "") -> int:
    """
    Tenta obter o número da aula de parada a partir do arquivo JSON de mapeamento.
    """
    prof_upper = str(professor or "").strip().upper()
    disc_upper = str(disciplina or "").strip().upper()
    turma_upper = str(turma or "").strip().upper()

    try:
        json_path = Path(REGISTRO_PROXIMA_GERACAO_PATH)
        if json_path.exists():
            with open(json_path, "r", encoding="utf-8") as fj:
                dados = json.load(fj)
                for item in dados:
                    if (str(item.get("professor") or "").strip().upper() == prof_upper and
                        str(item.get("disciplina") or "").strip().upper() == disc_upper and
                        str(item.get("turma") or "").strip().upper() == turma_upper):
                        if bimestre and item.get("bimestre"):
                            if str(item.get("bimestre")).strip().lower() != bimestre.strip().lower():
                                continue
                        return int(item.get("aula_parada") or 0)
    except Exception:
        pass

    return 0


def obter_referencia_ultima_aula_historico(
    professor: str,
    disciplina: str,
    turma: str,
    bimestre: str = "",
) -> dict | None:
    """Consulta o último plano salvo e identifica a última aula registrada nele.

    Esta função é apenas informativa para a tela de geração. Ela não altera a
    aula inicial nem a seleção automática dos PDFs.
    """
    from core.database import (
        obter_arquivo_historico,
        obter_ultimo_historico_por_contexto,
    )

    historico = obter_ultimo_historico_por_contexto(
        professor,
        disciplina,
        turma,
        bimestre,
    )
    if not historico:
        return None

    _, docx_bytes = obter_arquivo_historico(historico["id"])
    ultima_aula = detectar_ultima_aula_de_docx_bytes(docx_bytes, bimestre)

    return {
        **historico,
        "ultima_aula": ultima_aula,
    }


def obter_ultima_aula_gerada_sistema_impl(professor: str, disciplina: str, turma: str, bimestre: str = "") -> int:
    """
    Regra atual do projeto: novos planos sempre começam pela Aula 1.

    O histórico continua salvo para consulta e download, mas não deve mais
    interferir na aula inicial sugerida para novas gerações.
    """
    return 0
