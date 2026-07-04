from __future__ import annotations

import re
import logging
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


TECNICAS_PEDAGOGICAS = [
    "VIREM E CONVERSEM",
    "HORA DA LEITURA",
    "TODO MUNDO ESCREVE",
    "UM PASSO DE CADA VEZ",
    "COM SUAS PALAVRAS",
    "DE OLHO NO MODELO",
]

MAPA_SECOES = {
    "para comecar": "Para começar",
    "relembre": "Relembre",
    "foco no conteudo": "Foco no conteúdo",
    "na pratica": "Na prática",
    "encerramento": "Encerramento",
    "pause e responda": "IGNORAR",
}

TERMOS_COMPARACAO = (
    "comparacao",
    "comparar",
    "versus",
    "diferenca",
    "diferenças",
    "quadro comparativo",
)

PADROES_DESCARTE_TITULO = (
    "produzido pela",
    "disponivel em",
    "disponível em",
    "acesso em",
    "fonte:",
    "veja no livro",
    "link para video",
    "link para vídeo",
    "getty images",
    "seduc-sp",
    "continue",
    "continua",
    "veja no livro",
    "gabarito",
    "expectativas de respostas",
    "dinamica de conducao",
    "conceito-base",
    "caderno de exercicios",
    "referencias",
    "para professores",
    "habilidade:",
    "clique para adicionar uma legenda",
    "fonte do video",
    "fonte do vídeo",
)

COR_TITULO = RGBColor(0x1F, 0x49, 0x7D)
COR_SUBTITULO = RGBColor(0x2E, 0x74, 0xB5)
COR_ROTULO = RGBColor(0x00, 0x47, 0x70)
COR_OBSERVACAO = RGBColor(0x7F, 0x60, 0x00)
FONTE_PADRAO = "Arial"
logger = logging.getLogger(__name__)


@dataclass
class PaginaExtraida:
    numero: int
    texto: str
    linhas: list[str]
    tem_imagem: bool
    tem_tabela: bool
    tecnicas: list[str] = field(default_factory=list)
    secao_detectada: str | None = None
    eh_correcao: bool = False
    numero_atividade_detectada: int | None = None


@dataclass
class BlocoSecao:
    secao: str
    paginas: list[PaginaExtraida] = field(default_factory=list)
    numero_atividade: int | None = None


@dataclass
class AulaExtraida:
    caminho_pdf: Path
    titulo: str
    blocos: list[BlocoSecao]
    esboco: list[str]
    ancoras: list[str]
    observacoes: list[str]
    docx_auxiliar: Path | None = None


def _normalizar_espacos(texto: str) -> str:
    return re.sub(r"\s+", " ", str(texto or "")).strip()


def _sem_acento(texto: str) -> str:
    texto = unicodedata.normalize("NFD", str(texto or ""))
    texto = "".join(ch for ch in texto if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", texto).strip().lower()


def _linhas_texto(texto: str) -> list[str]:
    return [
        _normalizar_espacos(linha)
        for linha in str(texto or "").splitlines()
        if _normalizar_espacos(linha)
    ]


def _extrair_numero_aula(nome_arquivo: str) -> int:
    match = re.search(r"AULA[_\s-]*(\d{1,3})", str(nome_arquivo or ""), flags=re.I)
    if match:
        return int(match.group(1))
    match = re.search(r"(\d{1,3})", str(nome_arquivo or ""))
    return int(match.group(1)) if match else 0


def _eh_linha_generica_titulo(linha: str) -> bool:
    texto = _normalizar_espacos(linha)
    texto_norm = _sem_acento(texto)
    if not texto:
        return True
    if texto_norm in {"historia", "geografia", "ciencias", "arte"}:
        return True
    if texto_norm in {_sem_acento(secao) for secao in MAPA_SECOES.values() if secao != "IGNORAR"}:
        return True
    if texto.upper() in TECNICAS_PEDAGOGICAS:
        return True
    if any(padrao in texto_norm for padrao in PADROES_DESCARTE_TITULO):
        return True
    if re.search(r"\b\d+\s*min", texto_norm):
        return True
    if re.search(r"\b3o bimestre\b|\b3º bimestre\b|\banos finais\b|\baula\s*\d+\b", texto_norm):
        return True
    if re.fullmatch(r"slides?\s+\d+(\s*[ae]\s*\d+)?", texto_norm):
        return True
    if re.fullmatch(r"slide\s+\d+", texto_norm):
        return True
    if re.search(r"https?://|www\.", texto, flags=re.I):
        return True
    if re.fullmatch(r"[\W_]+", texto):
        return True
    return False


def _eh_pagina_correcao(linhas: list[str]) -> bool:
    if not linhas:
        return False
    primeiras = " ".join(linhas[:6])
    return bool(re.search(r"\b(correc[aã]o|gabarito|expectativas de respostas)\b", _sem_acento(primeiras)))


def _extrair_numero_atividade(linhas: list[str]) -> int | None:
    trecho = " ".join(linhas[:6])
    match = re.search(r"\batividade\s*(\d{1,2})\b", trecho, flags=re.I)
    return int(match.group(1)) if match else None


def _pontuar_linha_titulo(linha: str) -> float:
    texto = _normalizar_espacos(linha)
    texto_norm = _sem_acento(texto)
    if _eh_linha_generica_titulo(texto):
        return -100.0

    palavras = texto.split()
    score = 0.0
    if 1 <= len(palavras) <= 9:
        score += 4.0
    elif len(palavras) <= 14:
        score += 2.0
    else:
        score -= 1.5

    if texto[:1].islower():
        score -= 2.6

    if not texto.endswith("."):
        score += 1.0
    if ":" in texto and len(palavras) <= 8:
        score += 0.7
    if "?" in texto and len(palavras) <= 10:
        score += 0.8
    if re.search(r"\bmapa mental\b", texto_norm):
        score += 2.5
    if re.search(r"\bmapa\b", texto_norm):
        score += 1.2
    if re.match(r"^mapa\b", texto_norm) and "mapa mental" not in texto_norm:
        score -= 2.0
    if re.match(r"^(link para video|link para vídeo|fonte)\b", texto_norm):
        score -= 3.0
    if re.search(r"\bcompar", texto_norm):
        score += 1.4
    if re.search(r"\bmulheres?\b|\batenas\b|\besparta\b|\bp[oó]lis\b|\broma\b", texto_norm):
        score += 1.2
    if re.match(r"^(atividade|correcao|correção)\b", texto_norm):
        score -= 2.0
    return score


def classificar_secao(texto: str) -> str | None:
    for linha in _linhas_texto(texto)[:3]:
        linha_norm = _sem_acento(linha)
        for chave, secao in MAPA_SECOES.items():
            if re.match(rf"^{re.escape(chave)}\b", linha_norm):
                return secao
    return None


def extrair_paginas_pdf(caminho_pdf: str | Path) -> list[PaginaExtraida]:
    caminho = Path(caminho_pdf)
    paginas: list[PaginaExtraida] = []
    with pdfplumber.open(caminho) as pdf:
        for indice, page in enumerate(pdf.pages, start=1):
            texto = page.extract_text() or ""
            linhas = _linhas_texto(texto)
            texto_upper = texto.upper()
            tecnicas = [tecnica for tecnica in TECNICAS_PEDAGOGICAS if tecnica in texto_upper]
            pagina = PaginaExtraida(
                numero=indice,
                texto=texto,
                linhas=linhas,
                tem_imagem=bool(page.images),
                tem_tabela=bool(page.extract_tables() or []),
                tecnicas=tecnicas,
                secao_detectada=classificar_secao(texto),
                eh_correcao=_eh_pagina_correcao(linhas),
                numero_atividade_detectada=_extrair_numero_atividade(linhas),
            )
            paginas.append(pagina)
    return paginas


def agrupar_paginas_por_secao(paginas: list[PaginaExtraida]) -> list[BlocoSecao]:
    blocos: list[BlocoSecao] = []
    bloco_atual: BlocoSecao | None = None
    contador_atividade = 0

    for pagina in paginas:
        secao = pagina.secao_detectada
        if secao == "IGNORAR":
            continue

        if blocos and blocos[-1].secao == "Encerramento" and secao not in {None, "Encerramento"}:
            continue
        if secao is None:
            if bloco_atual is not None and bloco_atual.secao != "Encerramento":
                bloco_atual.paginas.append(pagina)
            continue

        if bloco_atual is not None and secao == bloco_atual.secao and secao in {
            "Para começar",
            "Relembre",
            "Foco no conteúdo",
            "Encerramento",
        }:
            bloco_atual.paginas.append(pagina)
            continue

        if secao == "Na prática":
            if bloco_atual is not None and bloco_atual.secao == "Na prática":
                if pagina.eh_correcao:
                    bloco_atual.paginas.append(pagina)
                    continue
                if (
                    pagina.numero_atividade_detectada is not None
                    and pagina.numero_atividade_detectada == bloco_atual.numero_atividade
                ):
                    bloco_atual.paginas.append(pagina)
                    continue

            numero_atividade = pagina.numero_atividade_detectada
            if numero_atividade is None:
                contador_atividade += 1
                numero_atividade = contador_atividade
            else:
                contador_atividade = max(contador_atividade, numero_atividade)

            bloco_atual = BlocoSecao(secao=secao, paginas=[pagina], numero_atividade=numero_atividade)
            blocos.append(bloco_atual)
            continue

        bloco_atual = BlocoSecao(secao=secao, paginas=[pagina])
        blocos.append(bloco_atual)

    return blocos


def _linhas_relevantes_para_titulo(pagina: PaginaExtraida) -> list[str]:
    linhas: list[str] = []
    for linha in pagina.linhas[:12]:
        texto = _normalizar_espacos(linha)
        if _eh_linha_generica_titulo(texto):
            continue
        if texto.upper() in TECNICAS_PEDAGOGICAS:
            continue
        if re.match(r"^\d+[.)]?\s*$", texto):
            continue
        linhas.append(texto)
    return linhas


def _extrair_pergunta_inicial(pagina: PaginaExtraida) -> str | None:
    linhas = _linhas_relevantes_para_titulo(pagina)
    if not linhas:
        return None

    def limpar_pergunta(texto: str) -> str:
        pergunta = _normalizar_espacos(texto)
        pergunta = re.sub(r"\?\s*\d+[.)].*$", "?", pergunta)
        if pergunta.count("?") > 1:
            pergunta = pergunta.split("?")[0] + "?"
        return pergunta

    for indice, linha in enumerate(linhas):
        if not re.match(r"^\d+[.)]?\s*", linha):
            continue
        linha_limpa = re.sub(r"^\d+[.)]?\s*", "", linha).strip()
        if "?" in linha_limpa:
            return limpar_pergunta(linha_limpa)

        acumulado = [linha_limpa]
        for proxima in linhas[indice + 1 : indice + 4]:
            proxima_norm = _sem_acento(proxima)
            if proxima_norm.startswith("recomendamos"):
                continue
            if proxima_norm.startswith(("disponivel em", "acesso em")):
                continue
            acumulado.append(proxima)
            if "?" in proxima:
                break
        pergunta = limpar_pergunta(" ".join(acumulado))
        if "?" in pergunta:
            return pergunta

    for indice, linha in enumerate(linhas):
        if "?" not in linha or len(linha.split()) > 16:
            continue

        acumulado = [linha]
        for anterior in reversed(linhas[max(0, indice - 2) : indice]):
            anterior_norm = _sem_acento(anterior)
            if re.match(r"^\d+[.)]?\s*", anterior):
                break
            if anterior_norm.startswith(("recomendamos", "disponivel em", "acesso em")):
                continue
            if len(anterior.split()) <= 8:
                acumulado.insert(0, anterior)
                continue
            break

        return limpar_pergunta(" ".join(acumulado))

    return None


def _extrair_titulo_curto_inicial(pagina: PaginaExtraida) -> str | None:
    linhas = _linhas_relevantes_para_titulo(pagina)
    if not linhas:
        return None

    primeira = linhas[0]
    if 1 <= len(primeira.split()) <= 12:
        if len(linhas) > 1:
            segunda = linhas[1]
            ultima_palavra = _sem_acento(primeira).split()[-1]
            if (
                len(segunda.split()) <= 4
                and (segunda[:1].islower() or ultima_palavra in {"a", "o", "as", "os", "de", "da", "do", "das", "dos", "e", "em", "para"})
            ):
                return _normalizar_espacos(f"{primeira} {segunda}")
        return primeira

    return None


def detectar_elemento_principal(pagina: PaginaExtraida) -> str:
    texto_norm = _sem_acento(pagina.texto)
    if "youtube.com" in texto_norm or "youtu.be" in texto_norm or "link para video" in texto_norm:
        return "VÍDEO"
    if "mapa mental" in texto_norm:
        return "MAPA MENTAL"
    if re.search(r"^\s*\d+\s*$", pagina.texto, flags=re.M) and len(re.findall(r"^\s*\d+\s*$", pagina.texto, flags=re.M)) >= 3:
        return "Lista numerada"
    if any(termo in texto_norm for termo in ("quadro comparativo", "tabela comparativa", "vamos comparar")):
        return "QUADRO DE COMPARAÇÃO"
    if "mapa" in texto_norm and pagina.tem_imagem:
        return "MAPA"
    if pagina.tem_imagem and any(
        termo in texto_norm
        for termo in (
            "ilustracao",
            "imagem",
            "ruinas",
            "detalhe de um vaso",
            "estatua",
            "urna eletr",
            "captura de tela",
            "tear vertical",
        )
    ):
        return "IMAGEM"
    if "tabela" in texto_norm or "quadro" in texto_norm:
        return "QUADRO/TABELA"
    if pagina.tem_imagem:
        return "IMAGEM"
    return "TEXTO"


def extrair_titulo_pagina(pagina: PaginaExtraida) -> str:
    if pagina.secao_detectada in {"Para começar", "Relembre"}:
        pergunta = _extrair_pergunta_inicial(pagina)
        if pergunta:
            return pergunta

    texto_norm = _sem_acento(pagina.texto)

    if "tabela comparativa" in texto_norm:
        return "Tabela comparativa entre sociedade e política na Grécia Antiga"

    if pagina.secao_detectada == "Encerramento" and "COM SUAS PALAVRAS" in pagina.texto.upper():
        if len(pagina.linhas) <= 3:
            return "Com suas palavras"
        return "Perguntas finais sobre o tema"

    titulo_curto = _extrair_titulo_curto_inicial(pagina)
    if titulo_curto:
        return titulo_curto

    candidatos: list[tuple[float, str]] = []
    for linha in pagina.linhas[:12]:
        score = _pontuar_linha_titulo(linha)
        if score > 0:
            candidatos.append((score, linha))

    if candidatos:
        candidatos.sort(key=lambda item: (-item[0], len(item[1])))
        return candidatos[0][1][:100]

    if "mapa mental" in texto_norm:
        return "Mapa mental"
    if any(termo in texto_norm for termo in TERMOS_COMPARACAO):
        return "Comparação"
    if pagina.tem_imagem:
        return "Conteúdo visual da página"
    return "Sem título"


def extrair_descricao_atividade(paginas: list[PaginaExtraida]) -> str:
    padrao = re.compile(
        r"(Leia|Observe|Analise|Analise a|Responda|Com base|A partir|Associe|Explique|Compare)[^.?!:]{10,120}[.?!]?",
        flags=re.I,
    )
    for pagina in paginas:
        texto_corrente = _normalizar_espacos(pagina.texto.replace("\n", " "))
        match = padrao.search(texto_corrente)
        if match:
            return _normalizar_espacos(match.group(0)).strip(" .")
        for linha in pagina.linhas:
            if _eh_linha_generica_titulo(linha):
                continue
            match = padrao.search(linha)
            if match:
                return _normalizar_espacos(match.group(0)).strip(" .")
    for pagina in paginas:
        for linha in pagina.linhas:
            if len(linha.split()) >= 4 and not _eh_linha_generica_titulo(linha):
                return linha[:120].strip(" .")
    return "Atividade de fixação"


def _bloco_tem_correcao(bloco: BlocoSecao) -> bool:
    return any(
        re.search(r"\b(correcao|gabarito|expectativas de respostas)\b", _sem_acento(pagina.texto))
        for pagina in bloco.paginas
    )


def _formatar_tecnicas(tecnicas: list[str]) -> str:
    itens = []
    vistos = set()
    for tecnica in tecnicas:
        chave = _sem_acento(tecnica)
        if chave in vistos:
            continue
        vistos.add(chave)
        itens.append(f'"{tecnica}"')
    return " E ".join(itens)


def gerar_esboco(blocos: list[BlocoSecao]) -> list[str]:
    linhas: list[str] = []
    for bloco in blocos:
        tecnicas_bloco = []
        for pagina in bloco.paginas:
            tecnicas_bloco.extend(pagina.tecnicas)
        tecnicas_txt = _formatar_tecnicas(tecnicas_bloco)

        if bloco.secao in {"Para começar", "Relembre"}:
            pagina = bloco.paginas[0]
            elemento = detectar_elemento_principal(pagina)
            titulo = extrair_titulo_pagina(pagina)
            linha = f"{bloco.secao.upper()}: {elemento} / {titulo}"
            if tecnicas_txt:
                linha += f" – {tecnicas_txt}"
            linhas.append(linha)
            continue

        if bloco.secao == "Foco no conteúdo":
            linhas.append(f"FOCO NO CONTEÚDO: ({len(bloco.paginas)} PÁGINAS)")
            for indice, pagina in enumerate(bloco.paginas, start=1):
                elemento = detectar_elemento_principal(pagina)
                titulo = extrair_titulo_pagina(pagina)
                linhas.append(f"  {indice}ª PÁGINA: {elemento} – {titulo}")
            continue

        if bloco.secao == "Na prática":
            descricao = extrair_descricao_atividade(bloco.paginas)
            linha = f"NA PRÁTICA: ATIVIDADE {bloco.numero_atividade}: {descricao}"
            if tecnicas_txt:
                linha += f" – {tecnicas_txt}"
            if _bloco_tem_correcao(bloco):
                linha += " + correção"
            linhas.append(linha)
            continue

        if bloco.secao == "Encerramento":
            titulo = extrair_titulo_pagina(bloco.paginas[0])
            linha = f"ENCERRAMENTO: {titulo}"
            if tecnicas_txt:
                linha += f" – {tecnicas_txt}"
            linhas.append(linha)

    return linhas


def _frases_para_ancoras(texto: str) -> list[str]:
    partes = re.split(r"[/:;]| - |\u2013", _normalizar_espacos(texto))
    frases = []
    for parte in partes:
        parte = _normalizar_espacos(parte).strip(" .")
        if not parte:
            continue
        if _eh_linha_generica_titulo(parte):
            continue
        if len(parte) < 4:
            continue
        if parte.lower().startswith("conteúdo visual"):
            continue
        if parte.lower().startswith(("slide", "slides", "referenc", "referên", "para professores")):
            continue
        frases.append(parte[:90])
    return frases


def coletar_ancoras(blocos: list[BlocoSecao]) -> list[str]:
    ancoras: list[str] = []
    vistos = set()

    def adicionar(texto: str) -> None:
        chave = _sem_acento(texto)
        if not texto or chave in vistos:
            return
        vistos.add(chave)
        ancoras.append(texto)

    for bloco in blocos:
        if bloco.secao == "Na prática":
            adicionar(f"ATIVIDADE {bloco.numero_atividade}")
        for pagina in bloco.paginas:
            if pagina.eh_correcao:
                continue
            elemento = detectar_elemento_principal(pagina)
            if elemento in {"MAPA MENTAL", "QUADRO DE COMPARAÇÃO", "Lista numerada", "VÍDEO"}:
                adicionar(elemento)
            for tecnica in pagina.tecnicas:
                adicionar(tecnica)
            if bloco.secao != "Na prática":
                titulo = extrair_titulo_pagina(pagina)
                for frase in _frases_para_ancoras(titulo):
                    adicionar(frase)
        if bloco.secao == "Na prática":
            descricao = extrair_descricao_atividade(bloco.paginas)
            for frase in _frases_para_ancoras(descricao):
                adicionar(frase)
    return ancoras


def extrair_titulo_aula(caminho_pdf: str | Path, paginas: list[PaginaExtraida]) -> str:
    nome_base = Path(caminho_pdf).stem.replace("_", " ").strip()
    nome_limpo = re.sub(r"\s+", " ", nome_base)
    match_arquivo = re.match(r"AULA\s*[- ]*(\d{1,3})\s*-\s*(.+)", nome_limpo, flags=re.I)
    if match_arquivo:
        return f"AULA {int(match_arquivo.group(1))} - {match_arquivo.group(2).strip()}"
    if paginas:
        linhas = paginas[0].linhas[:6]
        juntas = " ".join(linhas)
        match = re.search(r"(Aula\s*\d+\s+.+)", juntas, flags=re.I)
        if match:
            return _normalizar_espacos(match.group(1))
    return nome_limpo


def converter_pdf_para_docx_auxiliar(caminho_pdf: str | Path, pasta_destino: str | Path) -> Path | None:
    try:
        from pdf2docx import Converter
    except Exception as exc:
        logger.warning("pdf2docx indisponível para %s: %s", caminho_pdf, exc)
        return None

    caminho = Path(caminho_pdf)
    destino_dir = Path(pasta_destino)
    destino_dir.mkdir(parents=True, exist_ok=True)
    destino = destino_dir / f"{caminho.stem}__extracao.docx"

    try:
        if destino.exists() and destino.stat().st_mtime >= caminho.stat().st_mtime:
            return destino
    except OSError:
        pass

    logger.info("Convertendo PDF para DOCX auxiliar: %s", caminho.name)
    logging.getLogger("pdf2docx").setLevel(logging.WARNING)
    try:
        cv = Converter(str(caminho))
        cv.convert(str(destino), start=0, end=None)
        cv.close()
        return destino
    except Exception as exc:
        logger.warning("Falha ao converter PDF para DOCX auxiliar %s: %s", caminho, exc)
        return None


def processar_pdf_palavras_chave(caminho_pdf: str | Path, pasta_docx_auxiliares: str | Path | None = None) -> AulaExtraida:
    paginas = extrair_paginas_pdf(caminho_pdf)
    blocos = agrupar_paginas_por_secao(paginas)
    esboco = gerar_esboco(blocos)
    observacoes: list[str] = []
    if any(extrair_titulo_pagina(pagina) == "Conteúdo visual da página" for bloco in blocos for pagina in bloco.paginas):
        observacoes.append(
            "Algumas páginas têm pouco texto extraível e foram resumidas como conteúdo visual da página."
        )
    docx_auxiliar = None
    if pasta_docx_auxiliares:
        docx_auxiliar = converter_pdf_para_docx_auxiliar(caminho_pdf, pasta_docx_auxiliares)
        if docx_auxiliar is not None:
            observacoes.append(
                "Foi gerado um DOCX auxiliar via pdf2docx para apoiar leitura, revisão e próximos refinamentos."
            )
    return AulaExtraida(
        caminho_pdf=Path(caminho_pdf),
        titulo=extrair_titulo_aula(caminho_pdf, paginas),
        blocos=blocos,
        esboco=esboco,
        ancoras=coletar_ancoras(blocos),
        observacoes=observacoes,
        docx_auxiliar=docx_auxiliar,
    )


def processar_pasta_pdfs(caminho_pasta: str | Path, pasta_docx_auxiliares: str | Path | None = None) -> list[AulaExtraida]:
    pasta = Path(caminho_pasta)
    arquivos = sorted(
        pasta.glob("*.pdf"),
        key=lambda caminho: (_extrair_numero_aula(caminho.name), caminho.name.lower()),
    )
    return [processar_pdf_palavras_chave(caminho, pasta_docx_auxiliares=pasta_docx_auxiliares) for caminho in arquivos]


def _configurar_documento(doc: Document) -> None:
    secao = doc.sections[0]
    secao.top_margin = Inches(1)
    secao.bottom_margin = Inches(1)
    secao.left_margin = Inches(1)
    secao.right_margin = Inches(1)


def _configurar_paragrafo(paragrafo, *, alinhamento=None, antes: float = 0, depois: float = 0, linha: float = 1.15) -> None:
    if alinhamento is not None:
        paragrafo.alignment = alinhamento
    formato = paragrafo.paragraph_format
    formato.space_before = Pt(antes)
    formato.space_after = Pt(depois)
    formato.line_spacing = linha


def _adicionar_run(paragrafo, texto: str, *, negrito: bool = False, tamanho: float = 10.5, cor: RGBColor | None = None) -> None:
    run = paragrafo.add_run(texto)
    run.bold = negrito
    run.font.name = FONTE_PADRAO
    run.font.size = Pt(tamanho)
    if cor is not None:
        run.font.color.rgb = cor


def gerar_docx_palavras_chave(
    aulas: list[AulaExtraida],
    caminho_saida: str | Path,
    *,
    titulo_documento: str,
    subtitulo: str = "",
) -> Path:
    doc = Document()
    _configurar_documento(doc)

    p = doc.add_paragraph()
    _configurar_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.LEFT, depois=3)
    _adicionar_run(p, titulo_documento, negrito=True, tamanho=18, cor=COR_TITULO)

    if subtitulo:
        p = doc.add_paragraph()
        _configurar_paragrafo(p, depois=6)
        _adicionar_run(p, subtitulo, tamanho=11.5, cor=COR_SUBTITULO)

    p = doc.add_paragraph()
    _configurar_paragrafo(p, depois=8)
    _adicionar_run(
        p,
        (
            "Documento de teste com o esboço automático de palavras-chave, etapas e técnicas "
            "detectadas diretamente nos PDFs pedagógicos. Quando uma página tem pouco texto extraível, "
            "ela aparece como conteúdo visual da página para revisão posterior."
        ),
        tamanho=10.5,
    )

    for aula in aulas:
        p = doc.add_paragraph()
        _configurar_paragrafo(p, antes=10, depois=2)
        _adicionar_run(p, aula.titulo, negrito=True, tamanho=14, cor=COR_ROTULO)

        p = doc.add_paragraph()
        _configurar_paragrafo(p, depois=2)
        _adicionar_run(p, "Arquivo fonte: ", negrito=True, tamanho=10.5)
        _adicionar_run(p, aula.caminho_pdf.name, tamanho=10.5)

        if aula.docx_auxiliar is not None:
            p = doc.add_paragraph()
            _configurar_paragrafo(p, depois=2)
            _adicionar_run(p, "DOCX auxiliar: ", negrito=True, tamanho=10.5)
            _adicionar_run(p, aula.docx_auxiliar.name, tamanho=10.5)

        if aula.ancoras:
            p = doc.add_paragraph()
            _configurar_paragrafo(p, depois=2)
            _adicionar_run(p, "Âncoras detectadas: ", negrito=True, tamanho=10.5)
            _adicionar_run(p, "; ".join(aula.ancoras), tamanho=10.5)

        p = doc.add_paragraph()
        _configurar_paragrafo(p, antes=4, depois=2)
        _adicionar_run(p, "Esboço extraído", negrito=True, tamanho=11.5, cor=COR_TITULO)

        for linha in aula.esboco:
            p = doc.add_paragraph()
            _configurar_paragrafo(p, depois=0.5)
            _adicionar_run(p, linha, tamanho=10.5)

        for observacao in aula.observacoes:
            p = doc.add_paragraph()
            _configurar_paragrafo(p, antes=2, depois=0)
            _adicionar_run(p, f"Observação: {observacao}", tamanho=10, cor=COR_OBSERVACAO)

    saida = Path(caminho_saida)
    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida)
    return saida


def extrair_palavras_chave_docx(caminho_docx: str | Path) -> list[str]:
    """
    Abre um arquivo Word (.docx) e extrai todas as palavras-chave destacadas em amarelo
    (highlight yellow) nos parágrafos e nas tabelas, mantendo a sequência de aparição.
    """
    import docx
    from docx.enum.text import WD_COLOR_INDEX
    
    caminho = Path(caminho_docx)
    if not caminho.exists():
        return []
        
    doc = docx.Document(caminho)
    palavras_chave = []
    vistas = set()
    
    def adicionar(texto: str):
        texto_limpo = re.sub(r"\s+", " ", str(texto or "")).strip()
        if not texto_limpo or len(texto_limpo) < 2:
            return
        chave = texto_limpo.lower()
        if chave not in vistas:
            vistas.add(chave)
            palavras_chave.append(texto_limpo)

    # 1. Varre os parágrafos do documento principal
    for p in doc.paragraphs:
        for run in p.runs:
            destacado = False
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                destacado = True
            else:
                highlights = run._r.xpath('w:rPr/w:highlight')
                if highlights:
                    val = highlights[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if val == 'yellow':
                        destacado = True
            if destacado:
                adicionar(run.text)
                
    # 2. Varre as tabelas (células, linhas e parágrafos dentro de tabelas)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        destacado = False
                        if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                            destacado = True
                        else:
                            highlights = run._r.xpath('w:rPr/w:highlight')
                            if highlights:
                                val = highlights[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                if val == 'yellow':
                                    destacado = True
                        if destacado:
                            adicionar(run.text)
                            
    return palavras_chave

