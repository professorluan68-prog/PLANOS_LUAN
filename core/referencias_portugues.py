"""Referencias prontas de Língua Portuguesa a partir de DOCX na pasta dos PDFs."""

from __future__ import annotations

import re
import unicodedata
from functools import lru_cache
from pathlib import Path
from typing import Any

from core.referencias_base import (
    finalizar_aula,
    normalizar_busca as _normalizar_busca,
    normalizar_espacos as _normalizar_espacos,
    normalizar_numero_aula,
    paragrafos_docx as _paragrafos_docx,
    parte_titulo as _parte_titulo,
    pontuar_titulo as _pontuar_titulo,
    selecionar_referencia as _selecionar_referencia,
    tokens_titulo as _tokens_titulo,
)












def _normalizar_numero_aula(valor: Any) -> int:
    return normalizar_numero_aula(valor, max_digitos=3)




def _documento_docx(caminho_docx: str):
    try:
        from docx import Document
    except Exception:
        return None

    try:
        return Document(caminho_docx)
    except Exception:
        return None


def _finalizar_aula(aula: dict[str, Any] | None, aulas: dict[int, dict[str, Any]]) -> None:
    finalizar_aula(aula, aulas, max_digitos_numero=3)


@lru_cache(maxsize=16)
def _carregar_referencias_docx(caminho_docx: str) -> dict[int, dict[str, Any]]:
    doc = _documento_docx(caminho_docx)
    if doc is None:
        return {}

    aulas_tabelas = _carregar_referencias_docx_tabelas(doc)
    if aulas_tabelas:
        return aulas_tabelas

    paragrafos = [
        _normalizar_espacos(paragrafo.text)
        for paragrafo in doc.paragraphs
        if _normalizar_espacos(paragrafo.text)
    ]
    aulas: dict[int, dict[str, Any]] = {}
    aula_atual: dict[str, Any] | None = None
    secao = ""

    for texto in paragrafos:
        match_aula = re.match(r"^AULA\s+(\d{1,2})\s*[-–—]\s*(.+)$", texto, flags=re.I)
        if match_aula:
            _finalizar_aula(aula_atual, aulas)
            aula_atual = {
                "numero": int(match_aula.group(1)),
                "titulo": _normalizar_espacos(match_aula.group(2)),
                "metodologia": [],
                "acompanhamento": [],
                "acessibilidade": [],
            }
            secao = ""
            continue

        if not aula_atual:
            continue

        texto_norm = _normalizar_busca(texto)
        if texto_norm == "metodologia":
            secao = "metodologia"
            continue
        if texto_norm == "acompanhamento da aprendizagem":
            secao = "acompanhamento"
            continue
        if texto_norm == "acessibilidade":
            secao = "acessibilidade"
            continue

        if secao == "metodologia":
            match_etapa = re.match(r"^([^:]{2,80}):\s*(.+)$", texto)
            if match_etapa:
                aula_atual["metodologia"].append(
                    {
                        "titulo": _normalizar_espacos(match_etapa.group(1)),
                        "texto": _normalizar_espacos(match_etapa.group(2)),
                    }
                )
            elif aula_atual["metodologia"]:
                aula_atual["metodologia"][-1]["texto"] = _normalizar_espacos(
                    f"{aula_atual['metodologia'][-1]['texto']} {texto}"
                )
        elif secao in {"acompanhamento", "acessibilidade"}:
            item = texto if texto.startswith("☑") else f"☑ {texto.lstrip('☑ ').strip()}"
            aula_atual[secao].append(_normalizar_espacos(item))

    _finalizar_aula(aula_atual, aulas)
    return aulas


def _separar_itens_checklist(texto: str) -> list[str]:
    bruto = str(texto or "").replace("\r", "\n")
    itens = []
    partes = re.split(r"(?:^|\n)\s*☑\s*", bruto)
    for parte in partes:
        parte = _normalizar_espacos(parte)
        if not parte:
            continue
        itens.append(f"☑ {parte.lstrip('☑ ').strip()}")
    if itens:
        return itens[:3]

    linhas = [linha.strip("•- \t") for linha in bruto.splitlines() if linha.strip()]
    return [f"☑ {_normalizar_espacos(linha)}" for linha in linhas[:3]]


def _extrair_metodologia_texto(texto: str) -> list[dict[str, str]]:
    bruto = str(texto or "").replace("\r", "\n")
    padrao = re.compile(
        r"(?i)(Para começar|Para comecar|Foco no conteúdo|Foco no conteudo|Pause e responda|Na prática|Na pratica|Socialização|Socializacao|Encerramento)\s*:\s*"
    )
    ocorrencias = list(padrao.finditer(bruto))
    if not ocorrencias:
        return []

    blocos = []
    for indice, match in enumerate(ocorrencias):
        inicio_texto = match.end()
        fim_texto = ocorrencias[indice + 1].start() if indice + 1 < len(ocorrencias) else len(bruto)
        titulo = _normalizar_espacos(match.group(1))
        texto_bloco = _normalizar_espacos(bruto[inicio_texto:fim_texto])
        if texto_bloco:
            blocos.append({"titulo": titulo, "texto": texto_bloco})
    return blocos


def _carregar_referencias_docx_tabelas(doc) -> dict[int, dict[str, Any]]:
    aulas: dict[int, dict[str, Any]] = {}
    for table in doc.tables:
        if not table.rows:
            continue

        cabecalho = [_normalizar_busca(cell.text) for cell in table.rows[0].cells]
        if not cabecalho:
            continue

        try:
            idx_titulo = next(i for i, valor in enumerate(cabecalho) if "numero e titulo do material digital" in valor)
            idx_aprendizagem = next(
                i for i, valor in enumerate(cabecalho) if "aprendizagem essencial" in valor
            )
            idx_desenvolvimento = next(i for i, valor in enumerate(cabecalho) if valor == "desenvolvimento")
            idx_acompanhamento = next(i for i, valor in enumerate(cabecalho) if "acompanhamento da aprendizagem" in valor)
            idx_acessibilidade = next(i for i, valor in enumerate(cabecalho) if valor == "acessibilidade")
        except StopIteration:
            continue

        for row in table.rows[1:]:
            celulas = row.cells
            if len(celulas) <= max(
                idx_titulo,
                idx_aprendizagem,
                idx_desenvolvimento,
                idx_acompanhamento,
                idx_acessibilidade,
            ):
                continue

            titulo_material = _normalizar_espacos(celulas[idx_titulo].text)
            match_aula = re.match(r"^AULA\s+(\d{1,2})\s*[-–—]\s*(.+)$", titulo_material, flags=re.I)
            if not match_aula:
                continue

            numero = int(match_aula.group(1))
            titulo = _normalizar_espacos(match_aula.group(2))
            aprendizagem = _normalizar_espacos(celulas[idx_aprendizagem].text)
            metodologia = _extrair_metodologia_texto(celulas[idx_desenvolvimento].text)
            acompanhamento = _separar_itens_checklist(celulas[idx_acompanhamento].text)
            acessibilidade = _separar_itens_checklist(celulas[idx_acessibilidade].text)

            if metodologia and len(acompanhamento) >= 3 and len(acessibilidade) >= 3:
                aulas[numero] = {
                    "numero": numero,
                    "titulo": titulo,
                    "habilidade": aprendizagem,
                    "metodologia": metodologia,
                    "acompanhamento": acompanhamento[:3],
                    "acessibilidade": acessibilidade[:3],
                }
    return aulas


def _score_docx_referencia(caminho: Path) -> tuple[int, int, float, str]:
    nome = _normalizar_busca(caminho.name)
    prioridade_pasta = 1 if "relatorios conferencia planos" in _normalizar_busca(str(caminho.parent)) else 0
    prioridade_nome = 0
    if "revisado" in nome:
        prioridade_nome = 3
    elif "atualizado" in nome or "corrigido" in nome:
        prioridade_nome = 2
    elif "backup" in nome:
        prioridade_nome = -2
    try:
        modificado = caminho.stat().st_mtime
    except OSError:
        modificado = 0.0
    return prioridade_pasta, prioridade_nome, modificado, caminho.name.lower()


def localizar_docx_referencia_portugues(caminho_pdf: str | Path) -> Path | None:
    caminho = Path(caminho_pdf)
    if not caminho_pdf or not caminho.parent.exists():
        return None

    pastas_busca = [
        caminho.parent / "RELATORIOS_CONFERENCIA_PLANOS",
        caminho.parent,
    ]

    candidatos: list[Path] = []
    padroes = [
        "Plano_*Lingua_Portuguesa*_REVISADO.docx",
        "Plano_*Língua_Portuguesa*_REVISADO.docx",
        "*REVISADO*.docx",
        "Plano_*Lingua_Portuguesa*.docx",
        "Plano_*Língua_Portuguesa*.docx",
        "*.docx",
    ]
    for pasta in pastas_busca:
        if not pasta.exists():
            continue
        for padrao in padroes:
            candidatos.extend(pasta.glob(padrao))

    candidatos_unicos = {candidato.resolve(): candidato for candidato in candidatos}.values()
    candidatos_validos = [
        candidato
        for candidato in candidatos_unicos
        if not candidato.name.startswith("~$")
    ]
    if not candidatos_validos:
        return None
    return max(candidatos_validos, key=_score_docx_referencia)




def referencia_portugues_por_pdf(
    caminho_pdf: str | Path, numero_aula: Any, tema: str = ""
) -> dict[str, Any] | None:
    docx = localizar_docx_referencia_portugues(caminho_pdf)
    if not docx:
        return None
    numero = _normalizar_numero_aula(numero_aula)
    if not numero:
        numero = _normalizar_numero_aula(Path(caminho_pdf).stem)
    if not numero:
        return None
    referencias = _carregar_referencias_docx(str(docx))
    referencia = _selecionar_referencia(referencias, numero, tema)
    if not referencia:
        return None
    return {
        "numero": str(referencia.get("numero") or numero),
        "titulo": referencia.get("titulo", ""),
        "habilidade": referencia.get("habilidade", ""),
        "metodologia": list(referencia.get("metodologia") or []),
        "acompanhamento": list(referencia.get("acompanhamento") or [])[:3],
        "acessibilidade": list(referencia.get("acessibilidade") or [])[:3],
        "fonte": str(docx),
        "referencia_pedagogica_aplicada": True,
    }
