#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerador de DOCX pedagogicos para Lingua Portuguesa - Ensino Medio - 3 Bimestre
Le os JSONs de cada serie e gera um arquivo .docx formatado com alta qualidade pedagogica
Reescreve as metodologias diretamente a partir do texto_fonte (texto real do PDF).
"""

import json
import glob
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paleta de cores
COR_AZUL_PRINCIPAL = RGBColor(0x1F, 0x49, 0x7D)
COR_AZUL_SUBTITULO = RGBColor(0x2E, 0x74, 0xB5)
COR_CINZA_INTRO    = RGBColor(0x40, 0x40, 0x40)
COR_PRETO          = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------------------
# Utilitarios de formatacao
# ---------------------------------------------------------------------------

def set_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spc = OxmlElement('w:spacing')
    spc.set(qn('w:before'), str(before * 20))
    spc.set(qn('w:after'), str(after * 20))
    if line:
        spc.set(qn('w:line'), str(line * 20))
        spc.set(qn('w:lineRule'), 'exact')
    pPr.append(spc)


def p_titulo_principal(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=6)
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = COR_AZUL_PRINCIPAL


def p_subtitulo(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=10)
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = COR_AZUL_SUBTITULO


def p_intro(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=0, after=14)
    r = p.add_run(texto)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = COR_CINZA_INTRO


def p_titulo_aula(doc, texto):
    p = doc.add_paragraph()
    set_spacing(p, before=14, after=3)
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = COR_AZUL_PRINCIPAL


def p_secao(doc, titulo):
    p = doc.add_paragraph()
    set_spacing(p, before=5, after=2)
    r = p.add_run(titulo)
    r.bold = True
    r.underline = True
    r.font.size = Pt(11)
    r.font.color.rgb = COR_AZUL_PRINCIPAL


def p_etapa(doc, nome, texto):
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(nome + ': ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = COR_AZUL_SUBTITULO
    r2 = p.add_run(texto)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = COR_PRETO


def p_item(doc, texto):
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=1)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto)
    r.font.size = Pt(10.5)
    r.font.color.rgb = COR_PRETO


def p_divisoria(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'BFBFBF')
    pBdr.append(bot)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Extracao e reescrita pedagogica
# ---------------------------------------------------------------------------

TECNICAS_LEMOV = [
    'VIREM E CONVERSEM',
    'TODO MUNDO ESCREVE',
    'HORA DA LEITURA',
    'COM SUAS PALAVRAS',
    'UM PASSO DE CADA VEZ',
    'CHAMADA FRIA',
    'LEITURA EM VOZ ALTA',
]


def detectar_tecnicas(txt):
    encontradas = []
    txt_up = txt.upper()
    for t in TECNICAS_LEMOV:
        if t in txt_up:
            encontradas.append(t)
    return encontradas


def formatar_tecnica(texto, tecnicas):
    """Garante que tecnicas Lemov aparecam em MAIUSCULAS entre aspas."""
    for t in tecnicas:
        # Forma minuscula / mista
        padrao = re.compile(re.escape(t), re.IGNORECASE)
        if padrao.search(texto):
            texto = padrao.sub(f'"{t}"', texto, count=1)
    return texto


def get_titulo_arquivo(dados):
    """Extrai titulo limpo a partir do caminho do arquivo fonte."""
    arq = dados.get('arquivo_fonte', '') or dados.get('arquivo_fonte_extracao', '')
    if arq:
        base = os.path.basename(arq)
        base = os.path.splitext(base)[0]
        match = re.match(r'AULA_\d+\s*-\s*(.+)', base)
        if match:
            titulo = match.group(1).strip()
            # Normalizar espacos multiplos
            titulo = re.sub(r'\s+', ' ', titulo)
            return titulo
    return dados.get('tema', 'Aula')


def construir_metodologia_a_partir_fonte(dados):
    """
    Constroi metodologia pedagogica de alta qualidade a partir do texto_fonte.
    Usa a estrutura do JSON como guia e enriquece com o conteudo real do PDF.
    Retorna lista de dict {titulo, texto}
    """
    texto_fonte = dados.get('texto_fonte', '')
    metodologia_json = dados.get('metodologia', [])
    tema = dados.get('tema', '')
    numero = dados.get('numero_aula', '?')

    if not metodologia_json:
        return []

    tecnicas = detectar_tecnicas(texto_fonte)

    # Construir mapa de momentos do PDF
    tem_virem_conversem   = 'VIREM E CONVERSEM' in texto_fonte.upper()
    tem_todo_mundo        = 'TODO MUNDO ESCREVE' in texto_fonte.upper()
    tem_hora_leitura      = 'HORA DA LEITURA' in texto_fonte.upper()
    tem_com_suas_palavras = 'COM SUAS PALAVRAS' in texto_fonte.upper()
    tem_pause_responda    = 'PAUSE E RESPONDA' in texto_fonte.upper()
    tem_producao          = bool(re.search(r'produz|produção|Escrev[ae]|rascunho|revisão|revisao|elabore', texto_fonte, re.IGNORECASE))
    tem_encenacao         = bool(re.search(r'encenar|ensaiar|esquete|apresentar.*cena|roleplay', texto_fonte, re.IGNORECASE))
    tem_comparacao        = bool(re.search(r'compare|compar|diferença|semelhan', texto_fonte, re.IGNORECASE))
    tem_debate            = bool(re.search(r'debate|debater|argumento|contra-argumento', texto_fonte, re.IGNORECASE))

    # Contar atividades no texto_fonte
    atividades = re.findall(r'Atividade\s+\d+', texto_fonte)
    n_atividades = len(set(atividades))

    etapas_resultado = []

    for etapa in metodologia_json:
        titulo_orig = etapa.get('titulo', '').strip()
        texto_orig  = etapa.get('texto', '').strip()

        if not titulo_orig or not texto_orig:
            continue

        # Limpar textos gericos que nao descrevem a aula real
        texto_novo = texto_orig

        # Remover referencias genericas ao "schema" equivocado
        texto_novo = texto_novo.replace('esquema do material conceitual', 'esquema conceitual')
        texto_novo = texto_novo.replace('esquema do material geográfico', 'esquema conceitual')
        texto_novo = texto_novo.replace('pessoa biografada', tema)

        # Remover duplicacao "por meio de leitura orientada, destacando..."
        texto_novo = re.sub(
            r'Promover leitura orientada do material por meio de leitura orientada,\s*destacando informações essenciais e pontos de atenção\.',
            '',
            texto_novo
        )

        # Enriquecer Para começar: mencionar tecnica Lemov se detectada
        if titulo_orig in ('Para começar', 'Para comecar'):
            if tem_virem_conversem and '"VIREM E CONVERSEM"' not in texto_novo.upper():
                texto_novo = texto_novo.rstrip('.')
                texto_novo += (' A aula se abre com a estrategia "VIREM E CONVERSEM",'
                               ' mobilizando conhecimentos previos e ampliando a participacao da turma.')

        # Enriquecer Na pratica: mencionar tecnica se detectada
        if 'Na prática' in titulo_orig or 'Na pratica' in titulo_orig:
            if tem_hora_leitura and '"HORA DA LEITURA"' not in texto_novo.upper():
                texto_novo = texto_novo.rstrip('.')
                texto_novo += ' O momento inclui "HORA DA LEITURA" com texto do material.'
            if tem_todo_mundo and '"TODO MUNDO ESCREVE"' not in texto_novo.upper():
                texto_novo = texto_novo.rstrip('.')
                texto_novo += ' Os estudantes realizam o exercicio "TODO MUNDO ESCREVE".'
            if tem_com_suas_palavras and '"COM SUAS PALAVRAS"' not in texto_novo.upper():
                texto_novo = texto_novo.rstrip('.')
                texto_novo += ' A resposta e elaborada com "COM SUAS PALAVRAS".'

        # Encerramento: mencionar virem e conversem final se houver
        if titulo_orig == 'Encerramento':
            if tem_virem_conversem and '"VIREM E CONVERSEM"' not in texto_novo.upper():
                # Pode ser que virem e conversem aparece apenas no fim
                pass  # Nao adicionar redundantemente

        # Formatar tecnicas existentes
        texto_novo = formatar_tecnica(texto_novo, tecnicas)

        # Limpar espacos duplos
        texto_novo = re.sub(r'\s+', ' ', texto_novo).strip()

        # Remover ponto duplo ao final
        texto_novo = re.sub(r'\.+$', '.', texto_novo)
        if texto_novo and texto_novo[-1] not in '.!?':
            texto_novo += '.'

        etapas_resultado.append({'titulo': titulo_orig, 'texto': texto_novo})

    return etapas_resultado


# ---------------------------------------------------------------------------
# Carregamento
# ---------------------------------------------------------------------------

def carregar_aulas(pasta):
    arquivos = sorted(glob.glob(os.path.join(pasta, '*.json')))
    aulas = []
    for arq in arquivos:
        try:
            with open(arq, encoding='utf-8-sig') as f:
                dados = json.load(f)
            aulas.append(dados)
        except Exception as e:
            print(f"  [AVISO] Erro ao ler {arq}: {e}")
    aulas.sort(key=lambda x: int(str(x.get('numero_aula', 0))))
    return aulas


# ---------------------------------------------------------------------------
# Gerador
# ---------------------------------------------------------------------------

def configurar_doc(doc):
    s = doc.sections[0]
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    estilo = doc.styles['Normal']
    estilo.font.name = 'Calibri'
    estilo.font.size = Pt(10.5)


def gerar_docx_serie(pasta, saida, titulo, subtitulo, intro):
    print(f"\n{'='*60}")
    print(f"Gerando: {saida}")

    aulas = carregar_aulas(pasta)
    print(f"Aulas carregadas: {len(aulas)}")

    doc = Document()
    configurar_doc(doc)

    p_titulo_principal(doc, titulo)
    p_subtitulo(doc, subtitulo)
    p_intro(doc, intro)

    for i, dados in enumerate(aulas):
        numero = dados.get('numero_aula', '?')
        titulo_aula = get_titulo_arquivo(dados)
        metodologia  = dados.get('metodologia', [])
        acompanhamento = dados.get('acompanhamento', [])
        acessibilidade = dados.get('acessibilidade', [])

        print(f"  Aula {numero}: {titulo_aula[:55]}...")

        if i > 0:
            p_divisoria(doc)

        p_titulo_aula(doc, f"AULA {numero} - {titulo_aula}")

        # Metodologia
        p_secao(doc, "Metodologia")
        etapas = construir_metodologia_a_partir_fonte(dados)
        if not etapas:
            etapas = metodologia
        for etapa in etapas:
            t  = etapa.get('titulo', '').strip()
            tx = etapa.get('texto', '').strip()
            if t and tx:
                p_etapa(doc, t, tx)

        # Acompanhamento
        p_secao(doc, "Acompanhamento da aprendizagem")
        for item in acompanhamento:
            if item and item.strip():
                p_item(doc, item.strip())

        # Acessibilidade
        p_secao(doc, "Acessibilidade")
        for item in acessibilidade:
            if item and item.strip():
                p_item(doc, item.strip())

    doc.save(saida)
    print(f"[OK] Salvo: {saida}")
    return len(aulas)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

BASE = r"D:\PDF novos\LINGUA_PORTUGUESA\EM\3_BIMESTRE"

SERIES = [
    {
        "pasta":     os.path.join(BASE, "1_ANO"),
        "saida":     os.path.join(BASE, "1_ANO", "Metodologias_Lingua_Portuguesa_1_Ano_3_Bimestre.docx"),
        "titulo":    "Metodologias - Lingua Portuguesa",
        "subtitulo": "1 Ano - 3 Bimestre",
        "intro": (
            "Este material reune a metodologia, o acompanhamento da aprendizagem e a "
            "acessibilidade das aulas de Lingua Portuguesa do 1 Ano do Ensino Medio para o "
            "3 Bimestre. Cada registro esta organizado por aula e serve como referencia "
            "pedagogica para o sistema PLANOS_LUAN na geracao dos planos de aula em Word."
        ),
    },
    {
        "pasta":     os.path.join(BASE, "2_ANO"),
        "saida":     os.path.join(BASE, "2_ANO", "Metodologias_Lingua_Portuguesa_2_Ano_3_Bimestre.docx"),
        "titulo":    "Metodologias - Lingua Portuguesa",
        "subtitulo": "2 Ano - 3 Bimestre",
        "intro": (
            "Este material reune a metodologia, o acompanhamento da aprendizagem e a "
            "acessibilidade das aulas de Lingua Portuguesa do 2 Ano do Ensino Medio para o "
            "3 Bimestre. Cada registro esta organizado por aula e serve como referencia "
            "pedagogica para o sistema PLANOS_LUAN na geracao dos planos de aula em Word."
        ),
    },
    {
        "pasta":     os.path.join(BASE, "3_ANO"),
        "saida":     os.path.join(BASE, "3_ANO", "Metodologias_Lingua_Portuguesa_3_Ano_3_Bimestre.docx"),
        "titulo":    "Metodologias - Lingua Portuguesa",
        "subtitulo": "3 Ano - 3 Bimestre",
        "intro": (
            "Este material reune a metodologia, o acompanhamento da aprendizagem e a "
            "acessibilidade das aulas de Lingua Portuguesa do 3 Ano do Ensino Medio para o "
            "3 Bimestre. Cada registro esta organizado por aula e serve como referencia "
            "pedagogica para o sistema PLANOS_LUAN na geracao dos planos de aula em Word."
        ),
    },
]


def main():
    total = 0
    for s in SERIES:
        n = gerar_docx_serie(**s)
        total += n

    print(f"\n{'='*60}")
    print(f"CONCLUIDO! Total de aulas: {total}")
    print("Arquivos:")
    for s in SERIES:
        ok = "[OK]" if os.path.exists(s["saida"]) else "[FALTA]"
        print(f"  {ok} {s['saida']}")


if __name__ == "__main__":
    main()
