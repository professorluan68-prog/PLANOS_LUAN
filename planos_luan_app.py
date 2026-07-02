import streamlit as st
import tempfile
import re
import os
import json
import base64
import html
import math
import traceback
import zipfile
import unicodedata
import hashlib
from datetime import date, timedelta, datetime
from io import BytesIO
from pathlib import Path
import time
import threading
import signal

# Monitorador para fechar o servidor automaticamente quando o navegador for fechado
def _monitorar_sessoes_ativas():
    time.sleep(10)  # Período de tolerância inicial
    has_connected = False
    consecutive_zero_sessions = 0
    while True:
        try:
            from streamlit.runtime import get_instance
            runtime = get_instance()
            if runtime:
                active_sessions = runtime._session_mgr.list_active_sessions()
                count = len(active_sessions)
                if count > 0:
                    has_connected = True
                    consecutive_zero_sessions = 0
                else:
                    if has_connected:
                        consecutive_zero_sessions += 1
            else:
                consecutive_zero_sessions = 0
        except Exception:
            consecutive_zero_sessions = 0
        
        if has_connected and consecutive_zero_sessions >= 5:
            # Encerra o processo do Streamlit
            os._exit(0)
        time.sleep(1)

_monitor_thread = threading.Thread(target=_monitorar_sessoes_ativas, daemon=True)
_monitor_thread.start()
# ── Modernização da tela inicial ──────────────────────────────────────────────
from ui.tela_inicial_moderna import (
    HERO_CSS,
    HERO_HTML,
    STATS_HTML,
    SECTION_HEADER_HTML,
    OPTION_MENU_STYLES,
)
from ui.ui_components import render_sidebar
from core.constantes import (
    HORARIOS_AULA,
    HORARIOS_SIMPLES,
    HORARIOS_DUPLAS,
    TURNOS_HORARIOS,
    MESES,
    DIAS_SEMANA_CADASTRO,
    AULAS_SEMANA_OPCOES,
    EXTENSAO_MES_ANTECIPACOES,
    EXTENSAO_MES_OPCOES,
    EXTENSAO_MES_VALORES,
)

from ui.shared import (
    _rotulo_horario,
    _rotulo_data_aula_com_dia,
    _serializar_horarios_padronizados,
    _tipo_horario,
    nome_arquivo_plano,
    _normalizar_texto_simples,
    _normalizar_label_aula,
    _slug_key,
    _chave_cadastro,
    _eh_cadastro_cdp_eja,
    _arquivo_existe,
    _ler_bytes_arquivo_cache,
    _carregar_professores_dos_planos_cache,
    _diagnosticar_modelos_professores_cache,
    carregar_css,
    carregar_chaves_locais,
    _sincronizar_divisao_pdf_padrao,
    _proxima_data_pelo_dia,
    _sugerir_horario_e_tipo,
    _normalizar_horario_cadastro,
    _horarios_extraidos_texto,
    _dia_semana_numero,
    _partes_dia_config,
    _partes_horario_config,
    _sugerir_horario_cadastrado,
    _indice_horario,
    _horarios_padronizados_de_texto,
    _defaults_grade_horarios,
    _asset_data_uri,
    _selecionar_turma,
    _selecionar_mes,
    _selecionar_aulas_semana,
    _datas_horarios_do_mes,
    _datas_do_mes_por_dia,
    _padroes_horario_config,
    _mes_numero_app,
    DIAS_SEMANA_COMPLETOS,
)
from ui.cadastro import _renderizar_cadastro_professor
from ui.diagnostico import _renderizar_diagnostico_modelos
from ui.reescrita_cdp import _renderizar_reescrita_cdp_em

# ── Banco de Dados e Cadastro ──────────────────────────────────────────
from core.database import (
    listar_vinculos_professores,
    atualizar_vinculo_professor,
    salvar_professor_turma,
    duplicar_vinculo_professor,
    excluir_vinculo_professor,
    init_db,
    obter_professores_db,
    salvar_historico_plano,
    migrar_json_para_sqlite,
    verificar_plano_gerado_por_outro_professor,
)
from core.disciplinas import (
    BIMESTRES,
    TURMAS_CDP_MULTISSERIADA,
    eh_cdp,
    eh_cdp_contextual,
    eh_cdp_fundamental,
    eh_cdp_multisseriada,
    nomes_disciplinas,
    obter_config,
)
from core.cdp import SEQUENCIA_PADRAO_CDP_MULTISSERIADA
from core.cdp_em_docx import reescrever_docx_cdp_contextual_matematica
from core.calendario import (
    datas_do_periodo as _datas_do_periodo,
    datas_feriado_padrao as _datas_feriado_padrao,
    datas_sem_aula_padrao as _datas_sem_aula_padrao,
    datas_por_dia_ate_limite as _datas_por_dia_ate_limite,
    fim_periodo_mes_com_extensao as _fim_periodo_mes_com_extensao,
    inicio_periodo_mes_com_antecipacao as _inicio_periodo_mes_com_antecipacao,
    filtrar_datas_sem_aula as _filtrar_datas_sem_aula,
    rotulo_data_sem_aula as _rotulo_data_sem_aula,
)
from core.lote import processar_varios_pdfs
from core.validador_plano import validar_aulas_geradas
from config import MODELO_OPENAI_PADRAO, MODELO_GEMINI_PADRAO, PASTA_PLANOS_PROFESSORES, PLANOS_FINALIZADOS_DIR, TEMPLATES_DOCX_DIR, PASTA_BACKUP, inicializar_pastas, BASE_DIR
from docx_generator.preencher import preencher_documento
from docx_generator.preencher_cdp import preencher_documento_cdp, prever_aulas_cdp
from core.helpers import (
    LocalFileWrapper,
    filtrar_pdfs_para_aulas,
    arquivos_na_ordem_de_envio,
    horario_para_plano,
    listar_falhas_ia,
    montar_relatorio_geracao,
    normalizar_para_pasta,
    numeros_pdfs_faltantes,
    ordenar_pdfs_por_numero,
    ordenar_pdfs_por_sequencia,
    resolver_pasta_pdfs,
    resumir_falhas_ia,
    texto_lista as _texto_lista,
    numero_aula_pdf,
)
from core.turmas import turmas_espelho_mesma_serie
from core.professores_planos import (
    atualizar_cabecalho_modelo_professor,
    carregar_professores_dos_planos,
    criar_ou_atualizar_modelo_professor,
    diagnosticar_modelos_professores,
    extrair_datas_horarios_de_bytes,
    mesclar_professores,
)
from core.modelos_docx import (
    caminho_template_central,
    resolver_template_id_geracao,
    template_id_por_contexto,
)
from core.ae_priorizado import (
    aplicar_ae_priorizado_nas_aulas,
    contexto_ae_priorizado_disponivel,
    disciplina_ae_priorizado_disponivel,
    sequencia_aulas_ae_priorizado,
)

APP_ICON_PNG = BASE_DIR / "assets" / "planos_luan_icon.png"

st.set_page_config(
    page_title="PLANOS_LUAN",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)

CAMPOS_TELA = {
    "modelo_file",
    "pdfs_aulas_files",
    "modo_upload_pdf",
    "salvar_historico_geracao",
    "novo_modelo_file",
    "escolha_template",
    "escolha_template_manual",
    "modo_tela",
    "modo_ia",
    "professor",
    "professor_select",
    "aula_prof_select",
    "last_aula_prof",
    "disciplina_opcao",
    "disciplina_cdp_opcao",
    "disciplina_outra",
    "turma",
    "turma_select",
    "turma_cdp",
    "cdp_aula_inicial",
    "gerar_turma_espelho",
    "turma_espelho",
    "turma_espelho_select",
    "bimestre",
    "mes",
    "mes_select",
    "aulas_previstas_manual",
    "aulas_previstas_manual_select",
    "extensao_mes",
    "datas_sem_aula",
    "datas_sem_aula_assinatura",
    "escola",
    "componente_curricular",
    "last_componente_curricular",
    "professor_cadastro_select",
    "cadastro_busca",
    "cadastro_filtro_disciplina",
    "cadastro_filtro_origem",
    "cadastro_filtro_professor",
    "cadastro_filtro_sem_modelo",
    "cadastro_filtro_turma",
    "cadastro_selecionado",
}

PREFIXOS_TELA = (
    "pdfs_aulas_files_auto_",
    "data_aula_",
    "horario_aula_",
    "tipo_horario_aula_",
    "dividir_pdf_aula_",
    "data_turma2_aula_",
    "horario_turma2_aula_",
    "tipo_horario_turma2_aula_",
    "dividir_pdf_turma2_aula_",
    "cadastro_grade_",
    "ajuste_grade_",
)

PREFIXOS_REVISAO = ("tema_", "apr_", "acomp_", "acess_", "met_")
MODO_UPLOAD_PDF_PADRAO = "Todos de uma vez"


def _limpar_revisao_aulas() -> None:
    for chave in list(st.session_state.keys()):
        if any(str(chave).startswith(prefixo) for prefixo in PREFIXOS_REVISAO):
            del st.session_state[chave]


def limpar_dados_tela() -> None:
    _limpar_revisao_aulas()
    for chave in list(st.session_state.keys()):
        if chave in CAMPOS_TELA or any(str(chave).startswith(prefixo) for prefixo in PREFIXOS_TELA):
            del st.session_state[chave]


def _limpar_erro_processamento() -> None:
    st.session_state.pop("erro_processamento", None)
    st.session_state.pop("erro_processamento_detalhe", None)


def _assinatura_pdfs_automaticos(arquivos) -> str:
    """Identifica a lista atual de PDFs para evitar selecao antiga do Streamlit."""
    partes = []
    for arquivo in ordenar_pdfs_por_numero(arquivos or []):
        caminho = Path(arquivo)
        try:
            stat = caminho.stat()
            partes.append(f"{caminho.name}|{stat.st_size}|{stat.st_mtime_ns}|{numero_aula_pdf(caminho) or ''}")
        except OSError:
            partes.append(f"{caminho.name}|0|0|{numero_aula_pdf(caminho) or ''}")
    base = "\n".join(partes)
    return hashlib.md5(base.encode("utf-8")).hexdigest()[:12] if base else "sem_pdfs"

def _salvar_planos_na_pasta_finalizados(planos_gerados, disciplina: str) -> list[str]:
    caminhos_salvos = []
    try:
        PLANOS_FINALIZADOS_DIR.mkdir(parents=True, exist_ok=True)
        for plano in planos_gerados or []:
            nome_arq = nome_arquivo_plano(plano["turma"], disciplina, ia_usada=plano.get("ia_usada", False))
            caminho_completo = PLANOS_FINALIZADOS_DIR / nome_arq
            with open(caminho_completo, "wb") as f:
                f.write(plano["docx_bytes"].getvalue())
            caminhos_salvos.append(str(caminho_completo))
    except Exception as e:
        st.warning(f"Não foi possível salvar os arquivos localmente em {PLANOS_FINALIZADOS_DIR}: {e}")
    return caminhos_salvos


def _salvar_planos_gerados_se_configurado(
    planos_gerados,
    professor: str,
    disciplina: str,
    bimestre: str,
) -> bool:
    if not st.session_state.get("salvar_historico_geracao", False):
        return False

    for plano in planos_gerados or []:
        salvar_historico_plano(
            professor,
            disciplina,
            plano["turma"],
            nome_arquivo_plano(plano["turma"], disciplina),
            plano["docx_bytes"].getvalue(),
            bimestre=bimestre,
        )
    return True


def _registrar_mensagem_memoria_plano(salvou_historico: bool) -> None:
    if salvou_historico:
        st.session_state["mensagem_historico_planos_tipo"] = "success"
        st.session_state["mensagem_historico_planos"] = (
            "Plano salvo no histórico. Os próximos planos continuarão começando pela Aula 1."
        )
    else:
        st.session_state["mensagem_historico_planos_tipo"] = "info"
        st.session_state["mensagem_historico_planos"] = (
            "Plano gerado sem salvar no histórico."
        )


def _texto_lista_conferencia(itens) -> str:
    linhas = []
    for item in itens or []:
        texto = str(item or "").strip()
        if texto:
            linhas.append(texto)
    return "\n".join(linhas)


def _texto_metodologia_conferencia(aula: dict) -> str:
    return _texto_metodologia_app(aula)


def _linhas_relatorio_tecnico_conferencia(aula: dict) -> list[str]:
    return [
        "Relatório Técnico da Geração",
        f"Provedor da IA: {aula.get('ia_provedor') or 'Sem IA'}",
        f"Cache Reutilizado: {'Sim' if aula.get('cache_reutilizado') else 'Não'}",
        f"Versão do Gerador: {aula.get('versao_gerador') or '1.2.9'}",
        f"Origem da Metodologia: {aula.get('origem_metodologia') or 'Desconhecida'}",
        f"Score de Confiança: {aula.get('confidence_score', 100)}%",
    ]


def _texto_diagnostico_conferencia(aula: dict) -> str:
    diag = aula.get("diagnostico_geracao") or {}
    if not diag:
        return "Sem diagnóstico técnico detalhado."

    secoes = [
        ("1. Rascunho Local Heurístico", diag.get("metodologia_local") or []),
        ("2. Resposta IA Crua", diag.get("metodologia_ia_crua") or []),
        ("3. Higienização/Polimento", diag.get("metodologia_higienizada") or []),
        ("4. Metodologia Final", diag.get("metodologia_final") or []),
    ]
    partes = []
    for titulo, valor in secoes:
        partes.append(titulo)
        if isinstance(valor, str):
            partes.append(valor.strip() or "Nenhum conteúdo registrado.")
        elif valor:
            partes.append(_texto_metodologia_app({"metodologia": valor}))
        else:
            partes.append("Nenhum conteúdo registrado.")
        partes.append("")
    return "\n".join(partes).strip()


def _montar_texto_conferencia_aula(aula: dict, numero_aula: int, frases_redundantes=None) -> str:
    frases_redundantes = [str(frase).strip() for frase in (frases_redundantes or []) if str(frase).strip()]
    avisos_val = [str(aviso).strip() for aviso in (aula.get("avisos_validacao") or []) if str(aviso).strip()]
    score = aula.get("confidence_score")
    linhas = [
        f"Aula {numero_aula} - {aula.get('tema', '')}",
        "",
    ]

    if score is not None and score < 70:
        linhas.extend(
            [
                f"Baixo Score de Confiança ({score}%): Este plano de aula pode necessitar de ajustes manuais significativos.",
                "",
            ]
        )

    if avisos_val:
        linhas.append("Alertas de Qualidade Pedagógica:")
        linhas.extend(f"- {aviso}" for aviso in avisos_val)
        linhas.append("")

    if frases_redundantes:
        linhas.append("Aviso de Redundância (frases repetidas em mais de 2 aulas do lote):")
        linhas.extend(f'- "{frase}"' for frase in frases_redundantes)
        linhas.append("")

    linhas.extend(
        [
            "Tema",
            str(aula.get("tema", "") or ""),
            "",
            "Aprendizagem",
            str(aula.get("aprendizagem", "") or ""),
            "",
            "Acompanhamento",
            _texto_lista_conferencia(aula.get("acompanhamento") or []),
            "",
            "Acessibilidade",
            _texto_lista_conferencia(aula.get("acessibilidade") or []),
            "",
            "Metodologia",
            _texto_metodologia_conferencia(aula),
            "",
        ]
    )
    linhas.extend(_linhas_relatorio_tecnico_conferencia(aula))
    linhas.extend(["", "Transformação da Metodologia (Pipeline)", _texto_diagnostico_conferencia(aula), ""])
    return "\n".join(linhas).strip() + "\n"


def _resolver_pasta_base_conferencia(pasta_pdfs_auto: str = "", pdfs_selecionados=None) -> Path:
    if pasta_pdfs_auto:
        pasta = Path(pasta_pdfs_auto)
        if pasta.exists():
            return pasta

    for arquivo in pdfs_selecionados or []:
        caminho = getattr(arquivo, "path", None)
        if caminho:
            caminho = Path(caminho)
            if caminho.exists():
                return caminho.parent

    return BASE_DIR


def _salvar_relatorios_conferencia(
    *,
    turmas_processadas,
    duplicadas_por_aula,
    professor: str,
    disciplina: str,
    turma: str,
    mes: str,
    bimestre: str,
    modo_ia: str,
    modo_upload_pdf: str,
    pasta_pdfs_auto: str = "",
    pdfs_selecionados=None,
):
    token = st.session_state.get("revisao_token", 0)
    resumo_chave = {
        "token": token,
        "professor": professor,
        "disciplina": disciplina,
        "turma": turma,
        "mes": mes,
        "bimestre": bimestre,
        "modo_ia": modo_ia,
        "modo_upload_pdf": modo_upload_pdf,
        "aulas": [
            [aula.get("tema", ""), aula.get("confidence_score"), aula.get("avisos_validacao") or []]
            for bloco in (turmas_processadas or [])
            for aula in (bloco.get("aulas") or [])
        ],
    }
    chave = hashlib.md5(json.dumps(resumo_chave, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest()
    if st.session_state.get("relatorio_conferencia_chave") == chave:
        paths_salvos = st.session_state.get("relatorio_conferencia_paths") or []
        if paths_salvos and all(Path(caminho).exists() for caminho in paths_salvos):
            return paths_salvos

    pasta_base = _resolver_pasta_base_conferencia(pasta_pdfs_auto, pdfs_selecionados)
    pasta_relatorios = pasta_base / "RELATORIOS_CONFERENCIA_PLANOS"
    carimbo = datetime.now().strftime("%Y%m%d_%H%M%S")
    nome_execucao = "_".join(
        parte
        for parte in [
            carimbo,
            normalizar_para_pasta(disciplina) or "DISCIPLINA",
            normalizar_para_pasta(turma) or "TURMA",
            normalizar_para_pasta(modo_ia) or "MODO",
        ]
        if parte
    )
    pasta_execucao = pasta_relatorios / nome_execucao
    pasta_execucao.mkdir(parents=True, exist_ok=True)

    cabecalho = [
        "RELATÓRIO DE CONFERÊNCIA DO PLANO",
        f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}",
        f"Professor: {professor}",
        f"Disciplina: {disciplina}",
        f"Turma selecionada: {turma}",
        f"Mês: {mes}",
        f"Bimestre: {bimestre}",
        f"Modo IA: {modo_ia}",
        f"Modo de envio dos PDFs: {modo_upload_pdf}",
        f"Pasta dos relatórios: {pasta_execucao}",
        "Observação: esta pasta é apenas para conferência e pode ser apagada depois sem afetar o sistema.",
        "",
    ]

    relatorio_lote = list(cabecalho)
    arquivos_salvos = []
    for t_idx, bloco in enumerate(turmas_processadas or []):
        turma_bloco = str(bloco.get("turma") or "").strip() or turma
        relatorio_lote.extend([f"TURMA: {turma_bloco}", ""])
        for a_idx, aula in enumerate(bloco.get("aulas") or [], start=1):
            texto_aula = _montar_texto_conferencia_aula(
                aula,
                a_idx,
                duplicadas_por_aula.get((t_idx, a_idx - 1), []),
            )
            relatorio_lote.append(texto_aula)
            nome_aula = f"aula_{a_idx:02d}_{normalizar_para_pasta(aula.get('tema') or 'sem_tema')[:60]}.txt"
            caminho_aula = pasta_execucao / nome_aula
            caminho_aula.write_text("\n".join(cabecalho) + f"TURMA: {turma_bloco}\n\n" + texto_aula, encoding="utf-8")
            arquivos_salvos.append(str(caminho_aula))

    caminho_lote = pasta_execucao / "relatorio_conferencia_lote.md"
    caminho_lote.write_text("\n".join(relatorio_lote), encoding="utf-8")
    paths = [str(caminho_lote)] + arquivos_salvos
    st.session_state["relatorio_conferencia_chave"] = chave
    st.session_state["relatorio_conferencia_paths"] = paths
    return paths


def _registrar_erro_processamento(exc: Exception) -> None:
    mensagem = str(exc).strip() or exc.__class__.__name__
    st.session_state["erro_processamento"] = (
        "Nao foi possivel concluir o processamento das aulas. "
        f"Motivo: {mensagem}"
    )
    st.session_state["erro_processamento_detalhe"] = traceback.format_exc()


def _asset_data_uri(nome_arquivo: str, mime_type: str = "image/svg+xml") -> str:
    caminho = BASE_DIR / "assets" / nome_arquivo
    if not caminho.exists():
        return ""
    dados = base64.b64encode(caminho.read_bytes()).decode("ascii")
    return f"data:{mime_type};base64,{dados}"


TURMAS_PADRAO = ["(selecione a turma)"]
TURMAS_PADRAO += [f"{ano}º ANO {letra}" for ano in range(1, 10) for letra in ["A", "B", "C", "D", "E", "F"]]
TURMAS_PADRAO += [f"{ano}º ANO" for ano in range(1, 10)]
TURMAS_PADRAO += [
    "8º e 9º ano",
    "1º Termo",
    "2º Termo",
    "3º Termo",
    "MULTISSERIADO 1º, 2º e 3º ano",
    "MULTISSERIADO 4º e 5º ano",
]
TURMAS_PADRAO += [turma for turma in TURMAS_CDP_MULTISSERIADA if turma not in TURMAS_PADRAO]
TURMAS_PADRAO += ["Outra (digitar)"]

def _selecionar_turma(label: str, key_select: str, key_texto: str, placeholder: str = "Ex.: 7º ANO A") -> str:
    valor_atual = str(st.session_state.get(key_texto, "") or "").strip()
    opcoes = list(TURMAS_PADRAO)
    if valor_atual and valor_atual not in opcoes:
        opcoes.insert(-1, valor_atual)
    indice = opcoes.index(valor_atual) if valor_atual in opcoes else 0
    escolha = st.selectbox(label, opcoes, index=indice, key=key_select)
    if escolha == "Outra (digitar)":
        return st.text_input("Digite a turma", key=key_texto, placeholder=placeholder, autocomplete="off").strip()
    if escolha == "(selecione a turma)":
        st.session_state[key_texto] = ""
        return ""
    st.session_state[key_texto] = escolha
    return escolha


def _selecionar_turma_espelho(turma_principal: str, turmas_cadastradas: list[str]) -> str:
    opcoes = turmas_espelho_mesma_serie(turma_principal, list(dict.fromkeys(turmas_cadastradas or [])))
    if not turma_principal:
        st.warning("Selecione a turma principal antes de gerar para a 2ª turma.")
        st.session_state["turma_espelho"] = ""
        return ""

    if not opcoes:
        st.warning("Não encontrei outra turma cadastrada da mesma série para este professor e disciplina.")
        st.session_state["turma_espelho"] = ""
        return ""

    if len(opcoes) == 1:
        turma_unica = opcoes[0]
        st.session_state["turma_espelho"] = turma_unica
        st.info(f"2ª turma selecionada automaticamente: {turma_unica}.")
        return turma_unica

    valor_atual = str(st.session_state.get("turma_espelho", "") or "").strip()
    if valor_atual not in opcoes:
        valor_atual = opcoes[0]
        st.session_state["turma_espelho"] = valor_atual
        if "turma_espelho_select" in st.session_state:
            del st.session_state["turma_espelho_select"]

    indice = opcoes.index(valor_atual)
    escolha = st.selectbox("2ª Série/Turma", opcoes, index=indice, key="turma_espelho_select")
    st.session_state["turma_espelho"] = escolha
    st.caption("Opções limitadas às turmas cadastradas da mesma série.")
    return escolha


def _selecionar_mes() -> str:
    valor_atual = str(st.session_state.get("mes", "") or "").strip().upper()
    mes_padrao = MESES[date.today().month - 1]
    indice = MESES.index(valor_atual) if valor_atual in MESES else MESES.index(mes_padrao)
    mes_escolhido = st.selectbox("Mês", MESES, index=indice, key="mes_select")
    st.session_state["mes"] = mes_escolhido
    return mes_escolhido

def _selecionar_aulas_semana(label: str, key_select: str, key_texto: str) -> str:
    valor_atual = str(st.session_state.get(key_texto, "") or "").strip()
    opcoes = list(AULAS_SEMANA_OPCOES)

    if valor_atual and valor_atual not in opcoes:
        opcoes.append(valor_atual)

    valor_widget = str(st.session_state.get(key_select, "") or "").strip()
    if valor_widget and valor_widget not in opcoes:
        opcoes.append(valor_widget)

    if key_select not in st.session_state:
        st.session_state[key_select] = valor_atual if valor_atual in opcoes else "(selecione)"

    escolha = st.selectbox(label, opcoes, key=key_select)
    valor = "" if escolha == "(selecione)" else escolha
    st.session_state[key_texto] = valor
    return valor

# ── Banco de Dados e Cadastro ──────────────────────────────────────────
inicializar_pastas()
carregar_chaves_locais(BASE_DIR)
init_db()
migrar_json_para_sqlite()
PROFESSORES_DB = obter_professores_db()

PROFESSORES = {}
for prof, dados_prof in PROFESSORES_DB.items():
    disciplinas_unicas = []
    for d in dados_prof.get("disciplinas", []):
        nome_disc = d.get("disciplina")
        if nome_disc and nome_disc not in disciplinas_unicas:
            disciplinas_unicas.append(nome_disc)
    PROFESSORES[prof] = disciplinas_unicas

_NOMES_PROFESSORES = ["(selecione o professor)"] + sorted(PROFESSORES.keys()) + ["Outro (digitar)"]

def _slug_key(texto: str) -> str:
    return re.sub(r"[^A-Za-z0-9_]+", "_", str(texto or "")).strip("_") or "item"

def _chave_cadastro(
    professor: str,
    disciplina: str,
    turma: str,
    componente_curricular: str = "",
) -> tuple[str, str, str, str]:
    def norm(valor: str) -> str:
        valor = unicodedata.normalize("NFKD", str(valor or ""))
        valor = "".join(ch for ch in valor if not unicodedata.combining(ch))
        return re.sub(r"\s+", " ", valor).strip().upper()
    return norm(professor), norm(disciplina), norm(turma), norm(componente_curricular)

def _eh_cadastro_cdp_eja(disciplina: str, componente_curricular: str = "") -> bool:
    base = f"{disciplina} {componente_curricular}".upper()
    return eh_cdp(disciplina) or eh_cdp_contextual(disciplina) or "CDP" in base or "EJA" in base

def _arquivo_existe(caminho: str) -> bool:
    try:
        return bool(caminho and Path(caminho).exists())
    except OSError:
        return False

def _abrir_cadastro_com_filtros(professor: str, disciplina: str, turma: str) -> None:
    st.session_state["modo_tela"] = "Cadastro"
    st.session_state["cadastro_filtro_professor"] = professor
    st.session_state["cadastro_filtro_disciplina"] = disciplina
    st.session_state["cadastro_busca"] = turma

def _falhas_ia(aulas, exigir_ia: bool = True) -> list[str]:
    falhas = []
    if not exigir_ia:
        return falhas
    for idx, aula in enumerate(aulas or [], start=1):
        if aula.get("ia_usada"):
            continue
        erro = str(aula.get("ia_erro") or "").strip()
        tema = str(aula.get("tema") or f"Aula {idx}").strip()
        if erro:
            falhas.append(f"Aula {idx} ({tema}): {erro}")
        else:
            falhas.append(f"Aula {idx} ({tema}): a IA não retornou desenvolvimento completo.")
    return falhas

def _falhas_ia_atualizadas(aulas, exigir_ia: bool = True) -> list[str]:
    return listar_falhas_ia(aulas, exigir_ia=exigir_ia)


_falhas_ia = _falhas_ia_atualizadas


def _extrair_primeiro_texto_metodologia(aula) -> str:
    metodologia = aula.get("metodologia") or []
    if not metodologia:
        return ""

    primeiro_bloco = metodologia[0]
    return primeiro_bloco.get("texto", "") if isinstance(primeiro_bloco, dict) else str(primeiro_bloco)

def _salvar_pdf_temporario(pdf_file) -> str:
    """
    Guarda o PDF temporariamente, mas mantém o nome original para que 
    a Inteligência Artificial o consiga ler e usar como contexto!
    """
    # 1. Tenta descobrir o nome original do ficheiro que enviaste
    nome_original = getattr(pdf_file, "name", "aula.pdf")
    
    # 2. Limpa caracteres estranhos do nome por segurança
    nome_seguro = re.sub(r'[^A-Za-z0-9_.-]', '_', nome_original)
    
    # 3. Cria um nome final misturando um código curto com o nome original
    # Exemplo: xyZ1_Aula05.pdf (Assim a IA consegue ler o 'Aula05')
    pasta_temp = tempfile.gettempdir()
    codigo_aleatorio = base64.b64encode(os.urandom(3)).decode('utf-8').replace('/', '_').replace('+', '-')
    caminho_completo = os.path.join(pasta_temp, f"{codigo_aleatorio}_{nome_seguro}")
    
    try:
        pdf_file.seek(0)
    except Exception:
        pass
        
    # 4. Escreve o ficheiro no computador com o nome correto
    with open(caminho_completo, "wb") as f:
        f.write(pdf_file.read())
        
    return caminho_completo


def _preparar_pdf_para_processamento(pdf_file) -> tuple[str, bool]:
    """Retorna o caminho do PDF e se ele deve ser apagado ao final.

    No modo automatico, o arquivo ja existe em D:\\PDF novos. Usar esse caminho
    real preserva os JSONs, DOCXs de referencia e hashes da pasta original.
    No upload manual, criamos uma copia temporaria como antes.
    """
    caminho_local = getattr(pdf_file, "path", None)
    if caminho_local:
        caminho = Path(caminho_local)
        if caminho.exists():
            return str(caminho), False

    return _salvar_pdf_temporario(pdf_file), True


def _proxima_data_pelo_dia(dia_nome: str, data_referencia: date) -> date:
    dias = {"segunda": 0, "terça": 1, "quarta": 2, "quinta": 3, "sexta": 4, "sábado": 5, "domingo": 6}
    dia_alvo = dias.get(dia_nome.lower().strip(), 0)
    dias_para_frente = (dia_alvo - data_referencia.weekday() + 7) % 7
    if dias_para_frente == 0: dias_para_frente = 7
    return data_referencia + timedelta(days=dias_para_frente)

def _sugerir_horario_e_tipo(horario_str: str) -> tuple:
    horario_str = (horario_str or "").strip().lower()
    for h, label in HORARIOS_AULA:
        if h.lower() in horario_str:
            return h, label
    return HORARIOS_AULA[0]

def _normalizar_horario_cadastro(trecho: str) -> str:
    texto = (trecho or "").strip().lower()
    match = re.search(r"\b(\d{1,2})\s*(?:h|:)?\s*(\d{2})?\b", texto)
    if not match:
        return ""
    hora = int(match.group(1))
    minuto = match.group(2)
    if minuto:
        return f"{hora:02d}h{minuto}"
    return f"{hora:02d}h"

def _horarios_extraidos_texto(texto: str) -> list[str]:
    horarios = []
    for hora, minuto in re.findall(r"\b0?(\d{1,2})\s*(?:h|:)\s*(\d{2})?\b", str(texto or ""), flags=re.I):
        valor = f"{int(hora):02d}h{minuto or ''}"
        horarios.append(valor)
    return horarios

def _normalizar_label_aula(texto: str) -> str:
    texto = (texto or "").lower()
    texto = texto.replace("º", "ª").replace("°", "ª")
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()

def _normalizar_texto_simples(texto: str) -> str:
    texto = unicodedata.normalize("NFD", texto or "")
    texto = "".join(ch for ch in texto if unicodedata.category(ch) != "Mn")
    return texto.upper()

def _disciplina_suporta_modalidade_eja(disciplina: str) -> bool:
    texto = _normalizar_texto_simples(disciplina)
    return "BIOLOGIA" in texto or "INGLES" in texto

def _horarios_padronizados_de_texto(texto: str, contexto: str = "") -> list[tuple[str, str]]:
    horarios = []
    vistos = set()
    partes = _partes_horario_config(texto)
    if not partes and texto:
        partes = [str(texto)]
    for parte in partes:
        sugestao = _sugerir_horario_cadastrado(parte, contexto)
        if sugestao and sugestao not in vistos:
            horarios.append(sugestao)
            vistos.add(sugestao)
    return horarios

def _prefixo_turno(contexto: str) -> str:
    texto = _normalizar_texto_simples(contexto)
    if "NOITE" in texto or "NOTURNO" in texto:
        return "19"
    if "TARDE" in texto:
        return "13"
    return "07"

def _partes_horario_config(texto: str) -> list[str]:
    texto = str(texto or "").strip()
    if not texto:
        return []
    partes = [parte.strip() for parte in re.split(r"[;\n]+", texto) if parte.strip()]
    if len(partes) == 1:
        partes = [parte.strip() for parte in re.split(r",\s*(?=\d{1,2}h)", texto, flags=re.I) if parte.strip()]
    return partes


def _resumo_grade_cadastrada(config: dict | None) -> str:
    if not config:
        return ""
    dias = _partes_dia_config(str((config or {}).get("dia_semana") or ""))
    horarios = _partes_horario_config(str((config or {}).get("horario") or ""))
    aulas_semana = str((config or {}).get("aulas_semana") or "").strip()

    grade = []
    total = max(len(dias), len(horarios))
    for idx in range(total):
        dia = dias[idx] if idx < len(dias) else ""
        horario = horarios[idx] if idx < len(horarios) else ""
        if dia and horario:
            grade.append(f"{dia}: {horario}")
        elif dia:
            grade.append(dia)
        elif horario:
            grade.append(horario)

    prefixo = f"{aulas_semana} aula(s) na semana" if aulas_semana else "Grade cadastrada"
    if grade:
        return f"{prefixo} • " + " • ".join(grade)
    return prefixo if aulas_semana else ""

def _partes_dia_config(texto: str) -> list[str]:
    original = str(texto or "").strip()
    if not original:
        return []

    normalizado = _normalizar_texto_simples(original).replace("-", " ")
    padrao = re.compile(
        r"\b(SEGUNDA(?: FEIRA)?|TERCA(?: FEIRA)?|QUARTA(?: FEIRA)?|QUINTA(?: FEIRA)?|"
        r"SEXTA(?: FEIRA)?|SABADO|DOMINGO)\b"
    )
    dias = []
    vistos = set()
    for encontrado in padrao.finditer(normalizado):
        dia_num = _dia_semana_numero(encontrado.group(1))
        if dia_num is not None and dia_num not in vistos:
            vistos.add(dia_num)
            dias.append(DIAS_SEMANA_CADASTRO[dia_num])
    if dias:
        return dias

    return [parte.strip() for parte in re.split(r"[;,\n]+|\s+-\s+", original) if parte.strip()]

def _numeros_aulas_de_texto(texto: str) -> list[int]:
    base = re.sub(r"\b\d{1,2}h\d*\b", " ", str(texto or "").lower())
    numeros = []
    for numero in range(1, 7):
        if re.search(rf"\b{numero}\s*(?:ª|º|a|o)?\b", base):
            numeros.append(numero)
    return numeros

def _turno_por_horario_inicio(horario: str, contexto: str = "") -> str:
    horario_norm = _normalizar_horario_cadastro(horario)
    if horario_norm.startswith(("13", "14", "15", "16", "17")):
        return "Tarde"
    if horario_norm.startswith(("19", "20", "21", "22")):
        return "Noite"
    prefixo = _prefixo_turno(contexto)
    if prefixo == "13":
        return "Tarde"
    if prefixo == "19":
        return "Noite"
    return "Manhã"

def _formatar_label_aulas(numeros: list[int]) -> str:
    partes = [f"{numero}ª" for numero in sorted(set(numeros))]
    if not partes:
        return ""
    if len(partes) == 1:
        return f"{partes[0]} aula"
    if len(partes) == 2:
        return f"{partes[0]} e {partes[1]} aula"
    return f"{', '.join(partes[:-1])} e {partes[-1]} aula"

def _montar_horario_flexivel(turno: str, aulas: list[int | str]):
    slots = TURNOS_HORARIOS.get(turno) or TURNOS_HORARIOS["Manhã"]
    max_aulas = len(slots) - 1
    numeros = []
    for aula in aulas or []:
        match = re.search(r"\d+", str(aula))
        if match:
            numero = int(match.group(0))
            if 1 <= numero <= max_aulas:
                numeros.append(numero)
    numeros = sorted(set(numeros))
    if not numeros:
        return None

    primeira = numeros[0]
    ultima = numeros[-1]
    consecutivas = numeros == list(range(primeira, ultima + 1))
    inicio = slots[primeira - 1]
    fim = slots[ultima if consecutivas else ultima - 1]
    label = _formatar_label_aulas(numeros)
    return (inicio, label) if len(numeros) == 1 else (f"{inicio} - {fim}", label)

def _turno_e_aulas_de_horario(horario, contexto: str = "") -> tuple[str, list[str]]:
    if not horario:
        return ("Manhã", [])
    texto = _rotulo_horario(horario)
    horarios_texto = _horarios_extraidos_texto(texto)
    turno = _turno_por_horario_inicio(horarios_texto[0], contexto) if horarios_texto else _turno_por_horario_inicio("", contexto)
    numeros = _numeros_aulas_de_texto(texto)
    return turno, [f"{numero}ª" for numero in numeros]

def _horario_flexivel_por_texto(trecho: str, contexto: str = ""):
    numeros = _numeros_aulas_de_texto(trecho)
    if not numeros:
        return None
    horarios_texto = _horarios_extraidos_texto(trecho)
    turno = _turno_por_horario_inicio(horarios_texto[0], contexto) if horarios_texto else _turno_por_horario_inicio("", contexto)
    return _montar_horario_flexivel(turno, numeros)

def _indice_horario(horario) -> int:
    if horario in HORARIOS_AULA:
        return HORARIOS_AULA.index(horario)
    horarios_texto = _horarios_extraidos_texto(_rotulo_horario(horario))
    if not horarios_texto:
        return len(HORARIOS_AULA) + 99
    inicio = horarios_texto[0].lower()
    todos = [hora for slots in TURNOS_HORARIOS.values() for hora in slots[:-1]]
    for idx, hora in enumerate(todos):
        if hora.lower() == inicio:
            return idx
    return len(HORARIOS_AULA) + 50

def _defaults_grade_horarios(dia_texto: str = "", horario_texto: str = "", contexto: str = "") -> dict[str, dict[str, object]]:
    defaults: dict[str, dict[str, object]] = {}
    dias = _partes_dia_config(dia_texto)
    horarios = _horarios_padronizados_de_texto(horario_texto, contexto)
    for idx, dia in enumerate(dias):
        dia_num = _dia_semana_numero(dia)
        if dia_num is None or dia_num >= len(DIAS_SEMANA_CADASTRO):
            continue
        horario = horarios[idx] if idx < len(horarios) else (horarios[0] if horarios else None)
        turno, aulas = _turno_e_aulas_de_horario(horario, contexto)
        defaults[DIAS_SEMANA_CADASTRO[dia_num]] = {"turno": turno, "aulas": aulas}
    return defaults

def _renderizar_grade_horarios(prefixo: str, dia_texto: str = "", horario_texto: str = "", contexto: str = "") -> tuple[str, str, int]:
    st.markdown("**Grade semanal de horários**")
    st.caption("Selecione o turno e as aulas de cada dia. Deixe vazio o dia em que não há aula.")
    defaults = _defaults_grade_horarios(dia_texto, horario_texto, contexto)
    selecionados = []
    turnos = list(TURNOS_HORARIOS.keys())

    for indice, dia in enumerate(DIAS_SEMANA_CADASTRO):
        default_dia = defaults.get(dia, {})
        turno_default = str(default_dia.get("turno") or "Manhã")
        aulas_opcoes_default = [f"{numero}ª" for numero in range(1, len(TURNOS_HORARIOS.get(turno_default, TURNOS_HORARIOS["Manhã"])))]
        aulas_default = [aula for aula in default_dia.get("aulas", []) if aula in aulas_opcoes_default]

        col_dia, col_turno, col_aulas, col_previa = st.columns([1.1, 1.1, 2.2, 2.1])
        with col_dia:
            st.markdown(f"**{dia}**")
        with col_turno:
            turno = st.selectbox(
                "Turno",
                turnos,
                index=turnos.index(turno_default) if turno_default in turnos else 0,
                key=f"{prefixo}_turno_{indice}",
                label_visibility="collapsed",
            )
        aulas_opcoes = [f"{numero}ª" for numero in range(1, len(TURNOS_HORARIOS[turno]))]
        aulas_default = [aula for aula in aulas_default if aula in aulas_opcoes]
        with col_aulas:
            aulas = st.multiselect(
                "Aulas",
                aulas_opcoes,
                default=aulas_default,
                key=f"{prefixo}_aulas_{indice}",
                label_visibility="collapsed",
            )
        horario = _montar_horario_flexivel(turno, aulas)
def _prefixo_turno(contexto: str) -> str:
    texto = _normalizar_texto_simples(contexto)
    if "NOITE" in texto or "NOTURNO" in texto:
        return "19"
    if "TARDE" in texto:
        return "13"
    return "07"

def _partes_horario_config(texto: str) -> list[str]:
    texto = str(texto or "").strip()
    if not texto:
        return []
    partes = [parte.strip() for parte in re.split(r"[;\n]+", texto) if parte.strip()]
    if len(partes) == 1:
        partes = [parte.strip() for parte in re.split(r",\s*(?=\d{1,2}h)", texto, flags=re.I) if parte.strip()]
    return partes


def _resumo_grade_cadastrada(config: dict | None) -> str:
    if not config:
        return ""
    dias = _partes_dia_config(str((config or {}).get("dia_semana") or ""))
    horarios = _partes_horario_config(str((config or {}).get("horario") or ""))
    aulas_semana = str((config or {}).get("aulas_semana") or "").strip()

    grade = []
    total = max(len(dias), len(horarios))
    for idx in range(total):
        dia = dias[idx] if idx < len(dias) else ""
        horario = horarios[idx] if idx < len(horarios) else ""
        if dia and horario:
            grade.append(f"{dia}: {horario}")
        elif dia:
            grade.append(dia)
        elif horario:
            grade.append(horario)

    prefixo = f"{aulas_semana} aula(s) na semana" if aulas_semana else "Grade cadastrada"
    if grade:
        return f"{prefixo} • " + " • ".join(grade)
    return prefixo if aulas_semana else ""

def _partes_dia_config(texto: str) -> list[str]:
    original = str(texto or "").strip()
    if not original:
        return []

    normalizado = _normalizar_texto_simples(original).replace("-", " ")
    padrao = re.compile(
        r"\b(SEGUNDA(?: FEIRA)?|TERCA(?: FEIRA)?|QUARTA(?: FEIRA)?|QUINTA(?: FEIRA)?|"
        r"SEXTA(?: FEIRA)?|SABADO|DOMINGO)\b"
    )
    dias = []
    vistos = set()
    for encontrado in padrao.finditer(normalizado):
        dia_num = _dia_semana_numero(encontrado.group(1))
        if dia_num is not None and dia_num not in vistos:
            vistos.add(dia_num)
            dias.append(DIAS_SEMANA_CADASTRO[dia_num])
    if dias:
        return dias

    return [parte.strip() for parte in re.split(r"[;,\n]+|\s+-\s+", original) if parte.strip()]

def _numeros_aulas_de_texto(texto: str) -> list[int]:
    base = re.sub(r"\b\d{1,2}h\d*\b", " ", str(texto or "").lower())
    numeros = []
    for numero in range(1, 7):
        if re.search(rf"\b{numero}\s*(?:ª|º|a|o)?\b", base):
            numeros.append(numero)
    return numeros

def _turno_por_horario_inicio(horario: str, contexto: str = "") -> str:
    horario_norm = _normalizar_horario_cadastro(horario)
    if horario_norm.startswith(("13", "14", "15", "16", "17")):
        return "Tarde"
    if horario_norm.startswith(("19", "20", "21", "22")):
        return "Noite"
    prefixo = _prefixo_turno(contexto)
    if prefixo == "13":
        return "Tarde"
    if prefixo == "19":
        return "Noite"
    return "Manhã"

def _formatar_label_aulas(numeros: list[int]) -> str:
    partes = [f"{numero}ª" for numero in sorted(set(numeros))]
    if not partes:
        return ""
    if len(partes) == 1:
        return f"{partes[0]} aula"
    if len(partes) == 2:
        return f"{partes[0]} e {partes[1]} aula"
    return f"{', '.join(partes[:-1])} e {partes[-1]} aula"

def _montar_horario_flexivel(turno: str, aulas: list[int | str]):
    slots = TURNOS_HORARIOS.get(turno) or TURNOS_HORARIOS["Manhã"]
    max_aulas = len(slots) - 1
    numeros = []
    for aula in aulas or []:
        match = re.search(r"\d+", str(aula))
        if match:
            numero = int(match.group(0))
            if 1 <= numero <= max_aulas:
                numeros.append(numero)
    numeros = sorted(set(numeros))
    if not numeros:
        return None

    primeira = numeros[0]
    ultima = numeros[-1]
    consecutivas = numeros == list(range(primeira, ultima + 1))
    inicio = slots[primeira - 1]
    fim = slots[ultima if consecutivas else ultima - 1]
    label = _formatar_label_aulas(numeros)
    return (inicio, label) if len(numeros) == 1 else (f"{inicio} - {fim}", label)

def _turno_e_aulas_de_horario(horario, contexto: str = "") -> tuple[str, list[str]]:
    if not horario:
        return ("Manhã", [])
    texto = _rotulo_horario(horario)
    horarios_texto = _horarios_extraidos_texto(texto)
    turno = _turno_por_horario_inicio(horarios_texto[0], contexto) if horarios_texto else _turno_por_horario_inicio("", contexto)
    numeros = _numeros_aulas_de_texto(texto)
    return turno, [f"{numero}ª" for numero in numeros]

def _horario_flexivel_por_texto(trecho: str, contexto: str = ""):
    numeros = _numeros_aulas_de_texto(trecho)
    if not numeros:
        return None
    horarios_texto = _horarios_extraidos_texto(trecho)
    turno = _turno_por_horario_inicio(horarios_texto[0], contexto) if horarios_texto else _turno_por_horario_inicio("", contexto)
    return _montar_horario_flexivel(turno, numeros)

def _indice_horario(horario) -> int:
    if horario in HORARIOS_AULA:
        return HORARIOS_AULA.index(horario)
    horarios_texto = _horarios_extraidos_texto(_rotulo_horario(horario))
    if not horarios_texto:
        return len(HORARIOS_AULA) + 99
    inicio = horarios_texto[0].lower()
    todos = [hora for slots in TURNOS_HORARIOS.values() for hora in slots[:-1]]
    for idx, hora in enumerate(todos):
        if hora.lower() == inicio:
            return idx
    return len(HORARIOS_AULA) + 50

def _defaults_grade_horarios(dia_texto: str = "", horario_texto: str = "", contexto: str = "") -> dict[str, dict[str, object]]:
    defaults: dict[str, dict[str, object]] = {}
    dias = _partes_dia_config(dia_texto)
    horarios = _horarios_padronizados_de_texto(horario_texto, contexto)
    for idx, dia in enumerate(dias):
        dia_num = _dia_semana_numero(dia)
        if dia_num is None or dia_num >= len(DIAS_SEMANA_CADASTRO):
            continue
        horario = horarios[idx] if idx < len(horarios) else (horarios[0] if horarios else None)
        turno, aulas = _turno_e_aulas_de_horario(horario, contexto)
        defaults[DIAS_SEMANA_CADASTRO[dia_num]] = {"turno": turno, "aulas": aulas}
    return defaults

def _renderizar_grade_horarios(prefixo: str, dia_texto: str = "", horario_texto: str = "", contexto: str = "") -> tuple[str, str, int]:
    st.markdown("**Grade semanal de horários**")
    st.caption("Selecione o turno e as aulas de cada dia. Deixe vazio o dia em que não há aula.")
    defaults = _defaults_grade_horarios(dia_texto, horario_texto, contexto)
    selecionados = []
    turnos = list(TURNOS_HORARIOS.keys())

    for indice, dia in enumerate(DIAS_SEMANA_CADASTRO):
        default_dia = defaults.get(dia, {})
        turno_default = str(default_dia.get("turno") or "Manhã")
        aulas_opcoes_default = [f"{numero}ª" for numero in range(1, len(TURNOS_HORARIOS.get(turno_default, TURNOS_HORARIOS["Manhã"])))]
        aulas_default = [aula for aula in default_dia.get("aulas", []) if aula in aulas_opcoes_default]

        col_dia, col_turno, col_aulas, col_previa = st.columns([1.1, 1.1, 2.2, 2.1])
        with col_dia:
            st.markdown(f"**{dia}**")
        with col_turno:
            turno = st.selectbox(
                "Turno",
                turnos,
                index=turnos.index(turno_default) if turno_default in turnos else 0,
                key=f"{prefixo}_turno_{indice}",
                label_visibility="collapsed",
            )
        aulas_opcoes = [f"{numero}ª" for numero in range(1, len(TURNOS_HORARIOS[turno]))]
        aulas_default = [aula for aula in aulas_default if aula in aulas_opcoes]
        with col_aulas:
            aulas = st.multiselect(
                "Aulas",
                aulas_opcoes,
                default=aulas_default,
                key=f"{prefixo}_aulas_{indice}",
                label_visibility="collapsed",
            )
        horario = _montar_horario_flexivel(turno, aulas)
        with col_previa:
            st.caption(_rotulo_horario(horario) if horario else "Sem aula neste dia")
        if horario:
            selecionados.append({"dia": dia, "horario": horario, "aulas": aulas})

    dia_serializado = " - ".join(item["dia"] for item in selecionados)
    horario_serializado = _serializar_horarios_padronizados([item["horario"] for item in selecionados])
    total_aulas = sum(len(item["aulas"]) for item in selecionados)
    if total_aulas:
        st.caption(f"Total selecionado na semana: {total_aulas} aula(s).")
    return dia_serializado, horario_serializado, total_aulas

def _mes_numero_app(mes: str) -> int:
    meses = {
        "JANEIRO": 1, "FEVEREIRO": 2, "MARCO": 3, "ABRIL": 4, "MAIO": 5, "JUNHO": 6,
        "JULHO": 7, "AGOSTO": 8, "SETEMBRO": 9, "OUTUBRO": 10, "NOVEMBRO": 11, "DEZEMBRO": 12,
    }
    return meses.get(_normalizar_texto_simples(mes), date.today().month)

def _dia_semana_numero(texto: str):
    dias = {
        "SEGUNDA": 0, "SEGUNDA FEIRA": 0, "TERCA": 1, "TERCA FEIRA": 1,
        "QUARTA": 2, "QUARTA FEIRA": 2, "QUINTA": 3, "QUINTA FEIRA": 3,
        "SEXTA": 4, "SEXTA FEIRA": 4, "SABADO": 5, "DOMINGO": 6,
    }
    return dias.get(_normalizar_texto_simples(texto).replace("-", " "))

def _datas_do_mes_por_dia(
    mes: str,
    dia_semana: int,
    ano: int | None = None,
    extensao: int = 0,
    antecipacao: int = 0,
) -> list[date]:
    ano = ano or date.today().year
    mes_num = _mes_numero_app(mes)
    inicio = _inicio_periodo_mes_com_antecipacao(ano, mes_num, antecipacao)
    fim = _fim_periodo_mes_com_extensao(ano, mes_num, extensao)
    return _datas_por_dia_ate_limite(inicio, fim, dia_semana)

def _padroes_horario_config(config: dict, turma: str = "") -> list[dict]:
    padroes = []
    vistos = set()
    datas_horarios = list((config or {}).get("datas_horarios") or [])
    dias_semana = [
        _dia_semana_numero(parte.strip())
        for parte in _partes_dia_config(str((config or {}).get("dia_semana") or ""))
        if parte.strip()
    ]
    dias_semana = [dia for dia in dias_semana if dia is not None]
    horarios_cadastro = _horarios_padronizados_de_texto(str((config or {}).get("horario") or ""), turma)

    if horarios_cadastro and dias_semana:
        for idx, dia in enumerate(dias_semana):
            sugestao = horarios_cadastro[idx] if idx < len(horarios_cadastro) else horarios_cadastro[0]
            chave = (dia, sugestao)
            if chave in vistos:
                continue
            vistos.add(chave)
            padroes.append({"dia": dia, "horario": sugestao})
        return sorted(padroes, key=lambda item: (item["dia"], _indice_horario(item["horario"])))

    if horarios_cadastro and datas_horarios:
        dias_datas = []
        for item in datas_horarios:
            data_aula = item.get("data")
            if hasattr(data_aula, "weekday") and data_aula.weekday() not in dias_datas:
                dias_datas.append(data_aula.weekday())
        for idx, dia in enumerate(dias_datas):
            sugestao = horarios_cadastro[idx] if idx < len(horarios_cadastro) else horarios_cadastro[0]
            chave = (dia, sugestao)
            if chave in vistos:
                continue
            vistos.add(chave)
            padroes.append({"dia": dia, "horario": sugestao})
        if padroes:
            return sorted(padroes, key=lambda item: (item["dia"], _indice_horario(item["horario"])))

    for item in datas_horarios:
        data_aula = item.get("data")
        if not hasattr(data_aula, "weekday"):
            continue
        trecho_horario = " ".join(str(item.get(chave) or "").strip() for chave in ("horario", "aula")).strip()
        sugestao = _sugerir_horario_cadastrado(trecho_horario, turma) or HORARIOS_AULA[0]
        chave = (data_aula.weekday(), sugestao)
        if chave in vistos:
            continue
        vistos.add(chave)
        padroes.append({"dia": data_aula.weekday(), "horario": sugestao})

    if padroes:
        return sorted(padroes, key=lambda item: (item["dia"], _indice_horario(item["horario"])))

    partes_horario = _partes_horario_config(str((config or {}).get("horario") or ""))
    sugestoes_h = [_sugerir_horario_cadastrado(parte, turma) or HORARIOS_AULA[0] for parte in partes_horario]
    if dias_semana and not sugestoes_h:
        sugestoes_h = [HORARIOS_AULA[0]]

    for idx, dia in enumerate(dias_semana):
        sugestao = sugestoes_h[idx] if idx < len(sugestoes_h) else sugestoes_h[0]
        padroes.append({"dia": dia, "horario": sugestao})
    return padroes

def _datas_horarios_do_mes(
    config: dict,
    mes: str,
    turma: str = "",
    extensao: int = 0,
    antecipacao: int = 0,
) -> list[dict]:
    if not config or not mes:
        return []
    if config.get("repetir_modelo_semanal"):
        base = list(config.get("datas_horarios") or [])
        if not base:
            return []
        primeira_data = next((item.get("data") for item in base if hasattr(item.get("data"), "weekday")), None)
        if not primeira_data:
            return []
        inicio_semana_base = primeira_data - timedelta(days=primeira_data.weekday())

        # Construir molde a partir dos dias-da-semana únicos presentes em TODOS os registros
        # (não apenas da 1ª semana, que pode ter feriados omitindo dias)
        molde_por_dia: dict[int, dict] = {}  # dia_semana -> item de referência
        for item in base:
            data_aula = item.get("data")
            if not hasattr(data_aula, "weekday"):
                continue
            dia = data_aula.weekday()
            if dia not in molde_por_dia:
                # Guardar o deslocamento em dias desde o início da semana e o horário
                molde_por_dia[dia] = {
                    "offset_dias": dia,  # offset a partir de segunda (0=seg, 3=qui, etc.)
                    "horario": item.get("horario") or "",
                    "aula": item.get("aula") or "",
                }
        if not molde_por_dia:
            return []
        molde = sorted(molde_por_dia.values(), key=lambda m: m["offset_dias"])

        ano = date.today().year
        mes_num = _mes_numero_app(mes)
        inicio_periodo = _inicio_periodo_mes_com_antecipacao(ano, mes_num, antecipacao)
        fim_periodo = _fim_periodo_mes_com_extensao(ano, mes_num, extensao)

        # Encontrar a segunda-feira da semana em que o período começa
        inicio_bloco = inicio_periodo - timedelta(days=inicio_periodo.weekday())

        itens = []
        while inicio_bloco <= fim_periodo:
            for entrada in molde:
                nova_data = inicio_bloco + timedelta(days=entrada["offset_dias"])
                if nova_data < inicio_periodo or nova_data > fim_periodo:
                    continue
                itens.append({
                    "data": nova_data,
                    "horario": entrada["horario"],
                    "aula": entrada["aula"],
                })
            inicio_bloco += timedelta(days=7)
        return sorted(itens, key=lambda item: (item["data"], _indice_horario(item["horario"])))

    itens = []
    for padrao in _padroes_horario_config(config, turma):
        for data_aula in _datas_do_mes_por_dia(
            mes,
            padrao["dia"],
            extensao=extensao,
            antecipacao=antecipacao,
        ):
            itens.append({"data": data_aula, "horario": padrao["horario"]})
    return sorted(itens, key=lambda item: (item["data"], _indice_horario(item["horario"])))

def _sincronizar_datas_horarios_mes(
    config: dict,
    mes: str,
    professor: str,
    disciplina: str,
    turma: str,
    extensao: int = 0,
    antecipacao: int = 0,
    datas_sem_aula: list[date] | set[date] | None = None,
) -> list[dict]:
    itens = _filtrar_datas_sem_aula(
        _datas_horarios_do_mes(config, mes, turma, extensao=extensao, antecipacao=antecipacao),
        datas_sem_aula,
    )
    if not itens:
        for idx in range(40):
            for prefixo in ("data_aula_", "horario_aula_", "tipo_horario_aula_"):
                st.session_state.pop(f"{prefixo}{idx}", None)
        return []

    def _serializar_horario(item: dict) -> str:
        horario = item.get("horario")
        aula = item.get("aula")
        if isinstance(horario, (tuple, list)):
            return ":".join(str(parte) for parte in horario)
        return str(aula or horario or "")

    agenda = "|".join(f"{item['data'].isoformat()}:{_serializar_horario(item)}" for item in itens)
    cadastro = f"{config.get('dia_semana', '')}|{config.get('horario', '')}|{config.get('aulas_semana', '')}"
    datas_bloqueadas = ",".join(sorted(dt.isoformat() for dt in set(datas_sem_aula or [])))
    assinatura = f"{professor}|{disciplina}|{turma}|{mes}|{extensao}|{antecipacao}|{cadastro}|{agenda}|{datas_bloqueadas}"
    if st.session_state.get("agenda_mes_assinatura") == assinatura:
        return itens

    st.session_state["agenda_mes_assinatura"] = assinatura
    for idx, item in enumerate(itens):
        horario = item.get("horario")
        st.session_state[f"data_aula_{idx}"] = item["data"]
        if isinstance(horario, tuple):
            st.session_state[f"horario_aula_{idx}"] = horario
            st.session_state[f"tipo_horario_aula_{idx}"] = _tipo_horario(horario)

    for idx in range(len(itens), 40):
        for prefixo in ("data_aula_", "horario_aula_", "tipo_horario_aula_"):
            st.session_state.pop(f"{prefixo}{idx}", None)
    return itens


def _padroes_horario_config(config: dict, turma: str = "") -> list[dict]:
    padroes = []
    vistos = set()
    datas_horarios = list((config or {}).get("datas_horarios") or [])
    dias_semana = [
        _dia_semana_numero(parte.strip())
        for parte in _partes_dia_config(str((config or {}).get("dia_semana") or ""))
        if parte.strip()
    ]
    dias_semana = [dia for dia in dias_semana if dia is not None]
    horarios_cadastro = _horarios_padronizados_de_texto(str((config or {}).get("horario") or ""), turma)

    if horarios_cadastro and dias_semana:
        for idx, dia in enumerate(dias_semana):
            sugestao = horarios_cadastro[idx] if idx < len(horarios_cadastro) else horarios_cadastro[0]
            chave = (dia, sugestao)
            if chave in vistos:
                continue
            vistos.add(chave)
            padroes.append({"dia": dia, "horario": sugestao})
        return sorted(padroes, key=lambda item: (item["dia"], _indice_horario(item["horario"])))

    if horarios_cadastro and datas_horarios:
        dias_datas = []
        for item in datas_horarios:
            data_aula = item.get("data")
            if hasattr(data_aula, "weekday") and data_aula.weekday() not in dias_datas:
                dias_datas.append(data_aula.weekday())
        for idx, dia in enumerate(dias_datas):
            sugestao = horarios_cadastro[idx] if idx < len(horarios_cadastro) else horarios_cadastro[0]
            chave = (dia, sugestao)
            if chave in vistos:
                continue
            vistos.add(chave)
            padroes.append({"dia": dia, "horario": sugestao})
        if padroes:
            return sorted(padroes, key=lambda item: (item["dia"], _indice_horario(item["horario"])))

    for item in datas_horarios:
        data_aula = item.get("data")
        if not hasattr(data_aula, "weekday"):
            continue
        trecho_horario = " ".join(str(item.get(chave) or "").strip() for chave in ("horario", "aula")).strip()
        sugestao = _sugerir_horario_cadastrado(trecho_horario, turma) or HORARIOS_AULA[0]
        chave = (data_aula.weekday(), sugestao)
        if chave in vistos:
            continue
        vistos.add(chave)
        padroes.append({"dia": data_aula.weekday(), "horario": sugestao})

    if padroes:
        return sorted(padroes, key=lambda item: (item["dia"], _indice_horario(item["horario"])))

    partes_horario = _partes_horario_config(str((config or {}).get("horario") or ""))
    sugestoes_h = [_sugerir_horario_cadastrado(parte, turma) or HORARIOS_AULA[0] for parte in partes_horario]
    if dias_semana and not sugestoes_h:
        sugestoes_h = [HORARIOS_AULA[0]]

    for idx, dia in enumerate(dias_semana):
        sugestao = sugestoes_h[idx] if idx < len(sugestoes_h) else sugestoes_h[0]
        padroes.append({"dia": dia, "horario": sugestao})
    return padroes


def _sincronizar_datas_horarios_mes_turma2(
    config: dict,
    mes: str,
    professor: str,
    disciplina: str,
    turma: str,
    extensao: int = 0,
    antecipacao: int = 0,
    datas_sem_aula: list[date] | set[date] | None = None,
) -> list[dict]:
    """Versão da sincronização de datas/horários dedicada à 2ª turma (chaves prefixadas com 'turma2_')."""
    itens = _filtrar_datas_sem_aula(
        _datas_horarios_do_mes(config, mes, turma, extensao=extensao, antecipacao=antecipacao),
        datas_sem_aula,
    )
    if not itens:
        for idx in range(40):
            for prefixo in ("turma2_data_aula_", "turma2_horario_aula_", "turma2_tipo_horario_aula_"):
                st.session_state.pop(f"{prefixo}{idx}", None)
        return []

    def _serializar_horario_t2(item: dict) -> str:
        horario = item.get("horario")
        aula = item.get("aula")
        if isinstance(horario, (tuple, list)):
            return ":".join(str(parte) for parte in horario)
        return str(aula or horario or "")

    agenda = "|".join(f"{item['data'].isoformat()}:{_serializar_horario_t2(item)}" for item in itens)
    cadastro = f"{config.get('dia_semana', '')}|{config.get('horario', '')}|{config.get('aulas_semana', '')}"
    datas_bloqueadas = ",".join(sorted(dt.isoformat() for dt in set(datas_sem_aula or [])))
    assinatura = f"turma2|{professor}|{disciplina}|{turma}|{mes}|{extensao}|{antecipacao}|{cadastro}|{agenda}|{datas_bloqueadas}"
    if st.session_state.get("agenda_mes_assinatura_turma2") == assinatura:
        return itens

    st.session_state["agenda_mes_assinatura_turma2"] = assinatura
    for idx, item in enumerate(itens):
        horario = item.get("horario")
        st.session_state[f"turma2_data_aula_{idx}"] = item["data"]
        if isinstance(horario, tuple):
            st.session_state[f"turma2_horario_aula_{idx}"] = horario
            st.session_state[f"turma2_tipo_horario_aula_{idx}"] = _tipo_horario(horario)

    for idx in range(len(itens), 40):
        for prefixo in ("turma2_data_aula_", "turma2_horario_aula_", "turma2_tipo_horario_aula_"):
            st.session_state.pop(f"{prefixo}{idx}", None)
    return itens

def _config_agenda_a_partir_do_modelo(modelo_bytes: bytes | None) -> dict:
    if not modelo_bytes:
        return {}
    datas_horarios = extrair_datas_horarios_de_bytes(modelo_bytes)
    if not datas_horarios:
        return {}
    dias = []
    horarios = []
    for item in datas_horarios:
        data_aula = item.get("data")
        if hasattr(data_aula, "weekday"):
            dia_nome = DIAS_SEMANA_CADASTRO[data_aula.weekday()].upper()
            if dia_nome not in dias:
                dias.append(dia_nome)
        trecho_horario = " ".join(str(item.get(chave) or "").strip() for chave in ("horario", "aula")).strip()
        if trecho_horario and trecho_horario not in horarios:
            horarios.append(trecho_horario)
    return {
        "datas_horarios": datas_horarios,
        "dia_semana": " - ".join(dias),
        "horario": ", ".join(horarios),
        "repetir_modelo_semanal": True,
    }

def _sugerir_horario_cadastrado(trecho: str, contexto: str = ""):
    trecho_norm = _normalizar_label_aula(trecho)
    horarios_no_texto = _horarios_extraidos_texto(trecho)
    horario_flexivel = _horario_flexivel_por_texto(trecho, contexto)
    if horario_flexivel:
        return horario_flexivel

    if len(horarios_no_texto) >= 2:
        inicio, fim = horarios_no_texto[:2]
        for horario in HORARIOS_DUPLAS:
            if _horarios_extraidos_texto(horario[0])[:2] == [inicio, fim]:
                return horario

    for horario in HORARIOS_DUPLAS + HORARIOS_SIMPLES:
        label_norm = _normalizar_label_aula(horario[1])
        horario_norm = _normalizar_label_aula(HORARIOS_LABELS.get(horario, ""))
        if label_norm and label_norm in trecho_norm:
            return horario
        if horario_norm and horario_norm in trecho_norm:
            return horario

    horario_norm = _normalizar_horario_cadastro(trecho)
    if horario_norm:
        horario_sem_zero = horario_norm[1:] if horario_norm.startswith("0") else horario_norm
        for horario in HORARIOS_SIMPLES + HORARIOS_DUPLAS:
            base = horario[0].lower()
            if base == horario_norm or base == horario_sem_zero:
                return horario
            if base.startswith(f"{horario_sem_zero} ") or base.startswith(f"{horario_norm} "):
                return horario

    if not trecho_norm:
        return None
    candidatos = [h for h in HORARIOS_DUPLAS + HORARIOS_SIMPLES if _normalizar_label_aula(h[1]) in trecho_norm]
    if candidatos:
        prefixo = _prefixo_turno(contexto)
        for candidato in candidatos:
            if candidato[0].startswith(prefixo):
                return candidato
        return candidatos[0]
    return None

def _inicio_semana(dt: date) -> date:
    return dt - timedelta(days=dt.weekday())

def _tamanho_bloco_primeira_semana(datas: list[date]) -> int:
    """Retorna o número de aulas por semana (padrão semanal).

    Usa os dias-da-semana únicos presentes em TODAS as datas para determinar
    o tamanho do bloco semanal. Isso evita que feriados na 1ª semana reduzam
    o bloco incorretamente (ex: quinta feriada faz o sistema achar que a semana
    tem só terças).
    """
    if not datas:
        return 0
    # Coletar dias da semana únicos em todo o conjunto (não só na 1ª semana)
    dias_unicos = set()
    for dt in datas:
        dias_unicos.add(dt.weekday())
    bloco = len(dias_unicos)
    return max(1, bloco)

def _eh_data_antecipacao(data_aula: date, mes: str, antecipacao_mes: int) -> bool:
    if not mes or antecipacao_mes <= 0:
        return False
    ano = date.today().year
    mes_num = _mes_numero_app(mes)
    inicio_mes_oficial = date(ano, mes_num, 1)
    return data_aula < inicio_mes_oficial

def validar_entrada(
    modelo_bytes, disciplina: str, disciplina_config, aulas_envio,
    professor: str, turma: str, bimestre: str, mes: str,
    aulas_previstas_manual: str, pdfs_enviados: int = 0, pdfs_necessarios: int = 0,
    deixar_antecipacao_vazia: bool = False, antecipacao_mes: int = 0,
) -> str:
    disciplina_norm = re.sub(r"\s+", " ", str(disciplina or "")).strip().lower()
    orientacao_estudos = "orienta" in disciplina_norm and "estudo" in disciplina_norm
    if not modelo_bytes:
        return "Selecione ou envie o modelo DOCX."
    if not disciplina.strip():
        return "Selecione ou informe a disciplina."
    campos_obrigatorios = []
    if not (professor or "").strip(): campos_obrigatorios.append("Nome do professor")
    if not (turma or "").strip(): campos_obrigatorios.append("Turma")
    if not (bimestre or "").strip(): campos_obrigatorios.append("Bimestre")
    if not (mes or "").strip(): campos_obrigatorios.append("Mês")
    if not (aulas_previstas_manual or "").strip(): campos_obrigatorios.append("Aulas na semana")
    if campos_obrigatorios:
        return "Preencha os campos obrigatórios antes de gerar o plano: " + ", ".join(campos_obrigatorios) + "."
    if disciplina_config.exige_pdf and not aulas_envio:
        return "Envie os PDFs das aulas para gerar o plano."
    
    aulas_obrigatorias = [a for a in aulas_envio if not _eh_data_antecipacao(a["data"], mes, antecipacao_mes)] if deixar_antecipacao_vazia else aulas_envio
    pdf_unico_orientacao = bool(orientacao_estudos and pdfs_enviados == 1 and pdfs_necessarios >= 1)
    if disciplina_config.exige_pdf and pdfs_necessarios and pdfs_enviados != pdfs_necessarios and not pdf_unico_orientacao:
        return f"Quantidade de PDFs incorreta: foram adicionados {pdfs_enviados}, mas o plano possui {pdfs_necessarios} linha(s)."
    if disciplina_config.exige_pdf and any(not aula["pdf"] for aula in aulas_obrigatorias):
        return "Preencha data, horário e PDF em todas as aulas cadastradas do mês oficial."
    return ""

def validar_aulas_secundarias(gerar_turma_espelho: bool, turma_espelho: str, aulas_envio_espelho, exige_pdf: bool) -> str:
    if not gerar_turma_espelho: return ""
    if not (turma_espelho or "").strip(): return "Preencha a 2ª série/turma antes de gerar os planos em conjunto."
    if exige_pdf and any(not aula["pdf"] for aula in aulas_envio_espelho): return "Preencha data, horário e PDF da 2ª turma."
    return ""

def _grupos_pdf_por_aula(aulas_envio: list[dict]) -> list[dict]:
    grupos = []
    idx = 0
    while idx < len(aulas_envio):
        aula = aulas_envio[idx]
        dividir = bool(aula.get("dividir_pdf"))
        if dividir and idx + 1 < len(aulas_envio):
            grupos.append({"indices": [idx, idx + 1], "dividir": True})
            idx += 2
            continue
        grupos.append({"indices": [idx], "dividir": False})
        idx += 1
    return grupos

def _aplicar_pdfs_a_grupos(aulas_envio: list[dict], pdfs_aulas_files, replicar_pdf_unico: bool = False) -> tuple[list[dict], int]:
    grupos = _grupos_pdf_por_aula(aulas_envio)
    for grupo_idx, grupo in enumerate(grupos):
        if replicar_pdf_unico and len(pdfs_aulas_files or []) == 1:
            pdf = pdfs_aulas_files[0]
        else:
            pdf = pdfs_aulas_files[grupo_idx] if grupo_idx < len(pdfs_aulas_files) else None
        for indice in grupo["indices"]:
            aulas_envio[indice]["pdf"] = pdf
            aulas_envio[indice]["grupo_pdf"] = grupo_idx
            aulas_envio[indice]["dividir_pdf"] = grupo["dividir"]
    return aulas_envio, len(grupos)


def _divisao_pdf_padrao(idx: int, total_aulas: int) -> bool:
    return bool(idx % 2 == 0 and idx < total_aulas - 1)


def _sincronizar_divisao_pdf_padrao(
    num_rows: int,
    dividir_metodologia: bool,
    key_prefix: str = "",
    contexto: str = "",
) -> None:
    assinatura_chave = f"{key_prefix}dividir_metodologia_assinatura"
    assinatura_atual = f"v3|{bool(dividir_metodologia)}|{int(num_rows or 0)}|{contexto}"
    assinatura_anterior = st.session_state.get(assinatura_chave)
    acabou_de_ativar = bool(dividir_metodologia) and assinatura_anterior != assinatura_atual

    st.session_state[assinatura_chave] = assinatura_atual
    if not dividir_metodologia:
        return

    for idx in range(int(num_rows or 0)):
        chave = f"{key_prefix}dividir_pdf_aula_{idx}"
        if acabou_de_ativar or chave not in st.session_state:
            st.session_state[chave] = _divisao_pdf_padrao(idx, int(num_rows or 0))


def _estimar_pdfs_por_estado(num_rows: int, dividir_metodologia: bool, key_prefix: str = "") -> int:
    num_rows = int(num_rows or 0)
    if num_rows <= 0:
        return 0
    if not dividir_metodologia:
        return num_rows

    aulas_simuladas = []
    for idx in range(num_rows):
        chave = f"{key_prefix}dividir_pdf_aula_{idx}"
        dividir_pdf = st.session_state.get(chave, _divisao_pdf_padrao(idx, num_rows))
        aulas_simuladas.append({"dividir_pdf": bool(dividir_pdf)})
    return len(_grupos_pdf_por_aula(aulas_simuladas))

def _status_visual_aula(idx: int, num_rows: int, bloqueado: bool, continuidade_anterior: bool, dividir_pdf_ativo: bool) -> tuple:
    badges = []
    if continuidade_anterior:
        badges.extend(['<span class="lesson-badge lesson-badge--info">2o momento</span>', '<span class="lesson-badge lesson-badge--soft">PDF compartilhado</span>'])
        return "lesson-card lesson-card--continuation", "Continuacao da aula anterior", "Esta linha recebe a segunda parte da metodologia e usa o mesmo material.", badges
    if dividir_pdf_ativo:
        badges.extend(['<span class="lesson-badge lesson-badge--success">PDF em 2 aulas</span>', '<span class="lesson-badge lesson-badge--soft">1o momento</span>'])
        return "lesson-card lesson-card--paired", "Material compartilhado com a proxima aula", "Reaproveita o PDF na proxima linha.", badges
    if bloqueado:
        badges.append('<span class="lesson-badge lesson-badge--neutral">Repeticao semanal</span>')
        return "lesson-card lesson-card--locked", "Aula preenchida pela repeticao automatica", "Protegida para manter a sequencia.", badges
    if idx == num_rows - 1:
        badges.append('<span class="lesson-badge lesson-badge--neutral">Ultima aula</span>')
        return "lesson-card", "Configuracao individual", "Ajuste a data e o horario normalmente.", badges
    return "lesson-card", "Configuracao individual", "Defina a data, o horario e se o PDF continua.", badges

def _coletar_aulas_envio(
    num_rows: int,
    pdfs_aulas_files,
    dividir_metodologia: bool,
    auto_repetir_semana: bool,
    replicar_pdf_unico: bool = False,
    key_prefix: str = "",
    titulo_secao: str = "",
    modo_upload_individual: bool = False,
    preservar_datas_sincronizadas: bool = False,
    sequencia_pdf_esperada: list[int] | None = None,
    deixar_antecipacao_vazia: bool = False,
    mes: str = "",
    antecipacao_mes: int = 0,
):
    aulas_envio = []
    datas_cache = []
    horarios_cache = []

    for idx in range(num_rows):
        chave_data = f"{key_prefix}data_aula_{idx}"
        chave_horario = f"{key_prefix}horario_aula_{idx}"
        data_fallback = st.session_state.get(f"{key_prefix}data_aula_{idx}", date.today())
        horario_fallback = st.session_state.get(f"{key_prefix}horario_aula_{idx}", HORARIOS_AULA[0])
        if chave_data not in st.session_state: st.session_state[chave_data] = data_fallback
        if chave_horario not in st.session_state: st.session_state[chave_horario] = horario_fallback
        datas_cache.append(st.session_state[chave_data])
        horarios_cache.append(st.session_state[chave_horario])

    bloco_semana = 0 if preservar_datas_sincronizadas else _tamanho_bloco_primeira_semana(datas_cache)
    if (
        auto_repetir_semana
        and not preservar_datas_sincronizadas
        and bloco_semana > 0
        and num_rows > bloco_semana
    ):
        # Construir padrão semanal a partir dos dias-da-semana únicos em TODAS as datas
        # (não apenas da 1ª semana, que pode estar incompleta por feriados)
        dias_vistos: dict[int, tuple] = {}  # dia_semana -> horário representativo
        for dt, hr in zip(datas_cache, horarios_cache):
            dia = dt.weekday()
            if dia not in dias_vistos:
                dias_vistos[dia] = hr
        padrao = sorted(dias_vistos.items())  # [(dia_semana, horario), ...]

        # Para cada aula, determinar: qual dia da semana e qual semana
        for idx in range(num_rows):
            pos_no_padrao = idx % len(padrao)
            semana_num = idx // len(padrao)
            dia_semana_alvo, horario_padrao = padrao[pos_no_padrao]

            # Encontrar a data real: semana_num * 7 dias a partir da 1ª ocorrência desse dia
            # Usar a primeira data do padrão como âncora para a segunda-feira da semana 0
            data_ancora = next(
                (dt for dt in datas_cache if dt.weekday() == dia_semana_alvo),
                datas_cache[0]
            )
            segunda_ancora = data_ancora - timedelta(days=data_ancora.weekday())
            segunda_alvo = segunda_ancora + timedelta(weeks=semana_num)
            nova_data = segunda_alvo + timedelta(days=dia_semana_alvo)

            st.session_state[f"{key_prefix}data_aula_{idx}"] = nova_data
            st.session_state[f"{key_prefix}horario_aula_{idx}"] = horario_padrao
            st.session_state[f"{key_prefix}tipo_horario_aula_{idx}"] = _tipo_horario(horario_padrao)

    if titulo_secao: st.markdown(f"**{titulo_secao}**")

    for idx in range(num_rows):
        chave_data = f"{key_prefix}data_aula_{idx}"
        chave_horario = f"{key_prefix}horario_aula_{idx}"
        chave_tipo = f"{key_prefix}tipo_horario_aula_{idx}"
        chave_dividir = f"{key_prefix}dividir_pdf_aula_{idx}"
        data_fallback = st.session_state.get(chave_data, date.today())
        horario_fallback = st.session_state.get(chave_horario, HORARIOS_AULA[0])
        if chave_data not in st.session_state: st.session_state[chave_data] = data_fallback
        if chave_horario not in st.session_state: st.session_state[chave_horario] = horario_fallback
        
        horario_padrao_item = st.session_state.get(chave_horario, horario_fallback)
        bloqueado = (not preservar_datas_sincronizadas) and auto_repetir_semana and idx >= bloco_semana
        continuidade_anterior = bool(dividir_metodologia and idx > 0 and st.session_state.get(f"{key_prefix}dividir_pdf_aula_{idx - 1}", False))
        dividir_pdf_ativo = bool(dividir_metodologia and st.session_state.get(chave_dividir, False))
        card_class, status_titulo, status_texto, badges = _status_visual_aula(idx, num_rows, bloqueado, continuidade_anterior, dividir_pdf_ativo)
        badges_html = "".join([f'<span class="lesson-badge lesson-badge--index">Aula {idx + 1}</span>'] + badges)
        numero_pdf_esperado = None
        if not dividir_metodologia and sequencia_pdf_esperada and idx < len(sequencia_pdf_esperada):
            numero_pdf_esperado = sequencia_pdf_esperada[idx]

        with st.container():
            st.markdown(
                f"""
                <div class="{card_class}">
                    <div class="lesson-header">
                        <div>
                            <div class="lesson-title">{status_titulo}</div>
                            <div class="lesson-subtitle">{status_texto}</div>
                        </div>
                        <div class="lesson-badges">{badges_html}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True
            )
            if numero_pdf_esperado:
                st.caption(f"PDF esperado neste bloco: AULA {int(numero_pdf_esperado)}")
            col_data, col_horario = st.columns([1, 1])
            with col_data:
                data_label = _rotulo_data_aula_com_dia(st.session_state.get(chave_data, data_fallback))
                st.markdown(
                    f'<div class="lesson-field-label">Data da aula</div><div class="lesson-field-help">{data_label}</div>',
                    unsafe_allow_html=True,
                )
                data_aula = st.date_input(
                    "Data da aula",
                    format="DD/MM/YYYY",
                    key=chave_data,
                    disabled=bloqueado,
                    label_visibility="collapsed",
                )
                st.caption(f"Dia da semana: {DIAS_SEMANA_COMPLETOS[data_aula.weekday()]}")
            with col_horario:
                tipo_padrao = _tipo_horario(horario_padrao_item)
                if chave_tipo not in st.session_state or st.session_state[chave_tipo] not in ["Simples", "Dupla"]:
                    st.session_state[chave_tipo] = tipo_padrao
                tipo_horario = st.radio("Tipo de horário", ["Simples", "Dupla"], horizontal=True, key=chave_tipo, disabled=bloqueado)
                
                opcoes_horario = list(HORARIOS_SIMPLES if tipo_horario == "Simples" else HORARIOS_DUPLAS)
                horario_atual = st.session_state.get(chave_horario)
                if horario_atual not in opcoes_horario and isinstance(horario_atual, tuple):
                    opcoes_horario.insert(0, horario_atual)
                if st.session_state.get(chave_horario) not in opcoes_horario:
                    st.session_state[chave_horario] = opcoes_horario[0]
                horario_aula = st.selectbox("Horário", opcoes_horario, format_func=_rotulo_horario, key=chave_horario, disabled=bloqueado)

        dividir_pdf = False
        if dividir_metodologia:
            sugestao_dividir = _divisao_pdf_padrao(idx, num_rows)
            if chave_dividir not in st.session_state: st.session_state[chave_dividir] = sugestao_dividir
            if continuidade_anterior: st.session_state[chave_dividir] = False
            # Corrigido: bloqueado não impede o checkbox quando dividir_metodologia está ativo
            dividir_pdf = st.checkbox("Usar o mesmo PDF na próxima", key=chave_dividir, disabled=(bloqueado and not dividir_metodologia) or idx == num_rows - 1 or continuidade_anterior)
            if continuidade_anterior: st.caption("Esta já é continuação da anterior.")

        # Upload individual por aula
        pdf_individual = None
        eh_antecipacao_vazia = deixar_antecipacao_vazia and _eh_data_antecipacao(data_aula, mes, antecipacao_mes)
        if modo_upload_individual:
            if eh_antecipacao_vazia:
                st.caption("ℹ️ Aula na semana extra (bloco configurado para vir vazio).")
            elif continuidade_anterior:
                st.caption("📎 PDF compartilhado com a aula anterior.")
            else:
                chave_pdf_ind = f"{key_prefix}pdf_individual_aula_{idx}"
                uploaded = st.file_uploader(f"PDF da Aula {idx + 1}", type=["pdf"], key=chave_pdf_ind, label_visibility="collapsed")
                pdf_individual = uploaded
                if uploaded:
                    st.caption(f"✅ {uploaded.name}")
                else:
                    st.caption("⬆️ Adicione o PDF desta aula")

        aulas_envio.append({"data": data_aula, "horario": horario_aula, "pdf": pdf_individual, "dividir_pdf": dividir_pdf})

    if modo_upload_individual:
        # Propagar PDF para aulas de continuação (mesmo PDF da aula anterior)
        for i in range(1, len(aulas_envio)):
            if deixar_antecipacao_vazia and _eh_data_antecipacao(aulas_envio[i]["data"], mes, antecipacao_mes):
                continue
            if aulas_envio[i - 1].get("dividir_pdf") and aulas_envio[i].get("pdf") is None:
                aulas_envio[i]["pdf"] = aulas_envio[i - 1]["pdf"]
                aulas_envio[i]["grupo_pdf"] = aulas_envio[i - 1].get("grupo_pdf")
                aulas_envio[i]["dividir_pdf"] = False
    else:
        if deixar_antecipacao_vazia and antecipacao_mes > 0:
            aulas_antecipacao = [a for a in aulas_envio if _eh_data_antecipacao(a["data"], mes, antecipacao_mes)]
            aulas_oficiais = [a for a in aulas_envio if not _eh_data_antecipacao(a["data"], mes, antecipacao_mes)]
            _aplicar_pdfs_a_grupos(aulas_oficiais, pdfs_aulas_files, replicar_pdf_unico=replicar_pdf_unico)
            for a in aulas_antecipacao:
                a["pdf"] = None
                a["dividir_pdf"] = False
        else:
            aulas_envio, _ = _aplicar_pdfs_a_grupos(aulas_envio, pdfs_aulas_files, replicar_pdf_unico=replicar_pdf_unico)
    return aulas_envio

def _texto_metodologia_app(aula: dict) -> str:
    metodologia = aula.get("metodologia") or []
    blocos = []
    for item in metodologia:
        if isinstance(item, dict):
            titulo = item.get("titulo", "").strip()
            texto = item.get("texto", "").strip()
            if titulo: blocos.append(f"{titulo}: {texto}")
            else: blocos.append(texto)
        else: blocos.append(str(item))
    return "\n\n".join(blocos)

_TITULOS_METODOLOGIA_APP = {
    "para comecar": "Para comecar", "para começar": "Para comecar", "contextualizacao": "Contextualizacao",
    "leitura analitica": "Leitura analitica", "exploracao": "Exploracao", "formalizacao": "Formalizacao",
    "na pratica": "Na pratica", "sistematizacao": "Sistematizacao", "encerramento": "Encerramento",
}

def _normalizar_titulo_metodologia_app(texto: str) -> str:
    texto = (texto or "").strip().lower()
    mapa = str.maketrans("áàâãéêíóôõúç", "aaaaeeiooouc")
    return re.sub(r"\s+", " ", texto.translate(mapa)).strip()

def _metodologia_app_para_blocos(texto: str):
    linhas = [linha.rstrip() for linha in str(texto or "").splitlines()]
    blocos = []
    atual = None
    for linha in linhas:
        limpa = linha.strip()
        if not limpa: continue
        match = re.match(r"^([^:]{2,80}):\s*(.*)$", limpa)
        titulo_chave = _normalizar_titulo_metodologia_app(match.group(1)) if match else ""
        if match and titulo_chave in _TITULOS_METODOLOGIA_APP:
            if atual:
                atual["texto"] = " ".join(atual["texto"]).strip()
                blocos.append(atual)
            atual = {"titulo": _TITULOS_METODOLOGIA_APP[titulo_chave], "texto": [match.group(2).strip()] if match.group(2).strip() else []}
            continue
        if atual: atual["texto"].append(limpa)
        else: blocos.append(limpa)
    if atual:
        atual["texto"] = " ".join(atual["texto"]).strip()
        blocos.append(atual)
    return blocos or [str(texto or "").strip()]

def _extrair_aulas_dos_pdfs(
    aulas_envio, disciplina: str, turma_atual: str, bimestre: str, modo_ia: str,
    modelo_openai: str, modelo_gemini: str, dividir_metodologia: bool,
    modalidade_eja: bool = False, usar_ae_priorizado: bool = False,
    progress_callback=None, professor: str = "",
    deixar_antecipacao_vazia: bool = False, mes: str = "", antecipacao_mes: int = 0,
):
    temp_paths = []
    caminhos_para_apagar = []
    try:
        dados_aulas = []
        avisos_ia = []
        
        eh_ant_vazia = deixar_antecipacao_vazia and antecipacao_mes > 0
        if eh_ant_vazia:
            aulas_antecipacao = [a for a in aulas_envio if _eh_data_antecipacao(a["data"], mes, antecipacao_mes)]
            aulas_processamento = [a for a in aulas_envio if not _eh_data_antecipacao(a["data"], mes, antecipacao_mes)]
        else:
            aulas_antecipacao = []
            aulas_processamento = aulas_envio

        grupos = _grupos_pdf_por_aula(aulas_processamento) if dividir_metodologia else [{"indices": [idx], "dividir": False} for idx in range(len(aulas_processamento))]
        dividir_por_pdf = []
        for grupo in grupos:
            aula_envio = aulas_processamento[grupo["indices"][0]]
            caminho_pdf, apagar_ao_final = _preparar_pdf_para_processamento(aula_envio["pdf"])
            temp_paths.append(caminho_pdf)
            if apagar_ao_final:
                caminhos_para_apagar.append(caminho_pdf)
            dividir_por_pdf.append(bool(grupo["dividir"]))
            
        for aula_envio in aulas_processamento:
            dados_aulas.append({"data": aula_envio["data"].strftime("%d/%m"), "horario": horario_para_plano(aula_envio["horario"])})

        aulas = []
        avisos_ae = []
        avisos_repeticao = []
        if aulas_processamento:
            aulas = processar_varios_pdfs(
                temp_paths, disciplina=disciplina, turma=turma_atual, bimestre=bimestre, usar_ia=modo_ia != "Sem IA",
                provedor_ia=modo_ia.lower(), modelo_ia=(modelo_openai if modo_ia == "OpenAI" else modelo_gemini) if modo_ia != "Sem IA" else "",
                dividir_metodologia=dividir_metodologia, dividir_por_pdf=dividir_por_pdf, modalidade_eja=modalidade_eja,
                progress_callback=progress_callback, professor=professor,
            )
            if not aulas: raise RuntimeError("Nenhuma aula foi extraída dos PDFs oficiais.")
            
            if modo_ia != "Sem IA":
                falhas_ia = _falhas_ia(aulas, exigir_ia=not eh_cdp_contextual(disciplina))
                if falhas_ia and len(falhas_ia) == len(aulas or []):
                    raise RuntimeError("Falha de IA detectada em todas as aulas oficiais:\n" + "\n".join(falhas_ia))
                aviso_ia = resumir_falhas_ia(falhas_ia)
                if aviso_ia:
                    avisos_ia.append(aviso_ia)
            
            if usar_ae_priorizado:
                aulas, avisos_ae = aplicar_ae_priorizado_nas_aulas(
                    aulas,
                    disciplina=disciplina,
                    turma=turma_atual,
                    bimestre=bimestre,
                    caminho_planilha=str(st.session_state.get("caminho_ae_priorizado") or "").strip(),
                )

            for aula, dados in zip(aulas, dados_aulas): aula.update(dados)
            cdp_contextual = eh_cdp_contextual(disciplina)
            problemas_plano = validar_aulas_geradas(aulas, permitir_temas_repetidos=cdp_contextual, permitir_metodologia_simples=cdp_contextual or dividir_metodologia)
            
            for problema in problemas_plano:
                if "repetido de aula anterior" in str(problema).lower(): avisos_repeticao.append(problema)
                else: raise ValueError("Problemas encontrados:\n" + "\n".join(problemas_plano))

        aulas_vazias = []
        for aula_envio in aulas_antecipacao:
            aulas_vazias.append({
                "tema": "",
                "conteudo": "",
                "aprendizagem": "",
                "metodologia": [],
                "acompanhamento": [],
                "acessibilidade": [],
                "ia_usada": False,
                "data": aula_envio["data"].strftime("%d/%m"),
                "horario": horario_para_plano(aula_envio["horario"]),
            })
            
        aulas_completas = aulas_vazias + aulas

        for aula in aulas_completas:
            metodologia = aula.get("metodologia", [])
            for i, item in enumerate(metodologia):
                if isinstance(item, dict) and "texto" in item:
                    item["texto"] = re.sub(r'\s+', ' ', re.sub(r'\(\s*\)', '', re.sub(r'(?i)\s*(?:\(|-)?\s*\d+\s*min(?:uto)?s?(?:\))?', '', item["texto"]))).strip()
                elif isinstance(item, str):
                    metodologia[i] = re.sub(r'\s+', ' ', re.sub(r'\(\s*\)', '', re.sub(r'(?i)\s*(?:\(|-)?\s*\d+\s*min(?:uto)?s?(?:\))?', '', item))).strip()
                    
        return {
            "aulas": aulas_completas,
            "avisos_repeticao": avisos_repeticao,
            "avisos_ae": avisos_ae,
            "avisos_ia": avisos_ia
        }
    finally:
        for caminho_temp in caminhos_para_apagar:
            if caminho_temp:
                try: os.unlink(caminho_temp)
                except OSError: pass

def _gerar_docx_final(modelo_bytes: bytes, aulas, escola: str, professor: str, disciplina: str, componente_curricular: str, turma_atual: str, mes: str, bimestre: str, semana: str, observacao: str, aulas_previstas_manual: str):
    docx_bytes = preencher_documento(
        BytesIO(modelo_bytes), aulas, escola=escola, professor=professor, disciplina=componente_curricular or disciplina,
        turma=turma_atual, mes=mes, bimestre=bimestre, semana=semana, observacao=observacao, aulas_previstas_manual=aulas_previstas_manual,
    )
    relatorio = montar_relatorio_geracao(aulas, disciplina, turma_atual, bimestre, mes)
    return {"turma": turma_atual, "aulas": aulas, "docx_bytes": docx_bytes, "relatorio": relatorio, "ia_usada": any(aula.get("ia_usada") for aula in aulas)}

def _gerar_docx_cdp_final(modelo_bytes: bytes, escola: str, professor: str, disciplina: str, turma_atual: str, mes: str, bimestre: str, semana: str, observacao: str, aulas_previstas_manual: str, cdp_aula_inicial: int, turma_cdp: str = "", modo_ia: str = "Sem IA", modelo_openai: str = "", modelo_gemini: str = "", datas_horarios: list[dict] | None = None):
    docx_bytes = preencher_documento_cdp(
        BytesIO(modelo_bytes), escola=escola, professor=professor, turma=turma_atual, mes=mes, bimestre=bimestre,
        aula_inicial=int(cdp_aula_inicial or 1), fundamental=eh_cdp_fundamental(disciplina), multisseriada=eh_cdp_multisseriada(disciplina),
        serie_cdp=turma_cdp or "", usar_ia=modo_ia != "Sem IA", provedor_ia=modo_ia.lower() if modo_ia != "Sem IA" else "",
        modelo_ia=(modelo_openai if modo_ia == "OpenAI" else modelo_gemini) if modo_ia != "Sem IA" else "",
        datas_horarios=datas_horarios, semana=semana, observacao=observacao, aulas_previstas_manual=aulas_previstas_manual,
    )
    tipo = "CDP - Ciclo I" if eh_cdp_fundamental(disciplina) else "CDP/EJA Multisseriada"
    relatorio = f"Plano gerado em modo {tipo}.\nProfessor: {professor}\nDisciplina: {disciplina}\nTurma: {turma_atual}\nBimestre: {bimestre}\nMês: {mes}\nAula inicial CDP: {int(cdp_aula_inicial or 1)}\n"
    if turma_cdp: relatorio += f"Turma multisseriada: {turma_cdp}\n"
    if modo_ia != "Sem IA": relatorio += f"IA: {modo_ia}\n"
    return {"turma": turma_atual, "aulas": [], "docx_bytes": docx_bytes, "relatorio": relatorio, "ia_usada": modo_ia != "Sem IA"}

def _montar_zip_planos(planos: list[dict], disciplina: str) -> bytes:
    saida = BytesIO()
    with zipfile.ZipFile(saida, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for plano in planos:
            nome_docx = nome_arquivo_plano(plano["turma"], disciplina, ia_usada=plano["ia_usada"])
            zf.writestr(nome_docx, plano["docx_bytes"].getvalue())
            zf.writestr(nome_docx.replace(".docx", "_relatorio.txt"), plano["relatorio"].encode("utf-8"))
    saida.seek(0)
    return saida.read()

st.markdown(HERO_CSS, unsafe_allow_html=True)

st.markdown(HERO_HTML, unsafe_allow_html=True)
st.markdown(STATS_HTML, unsafe_allow_html=True)
render_sidebar()

col_limpar, _ = st.columns([1, 5])
with col_limpar: st.button("Limpar dados da tela", type="secondary", on_click=limpar_dados_tela)

st.markdown(SECTION_HEADER_HTML, unsafe_allow_html=True)

from streamlit_option_menu import option_menu

modos_disponiveis = ["Planos gerais", "CDP - Ciclo I", "Reescrita CDP", "Cadastro", "Diagnóstico"]
if st.session_state.get("modo_tela") == "Geração em Lote":
    st.session_state["modo_tela"] = "Planos gerais"

# Sincroniza o modo_tela default a partir do session_state se existir
default_modo = st.session_state.get("modo_tela", "Planos gerais")
if default_modo not in modos_disponiveis:
    default_modo = "Planos gerais"
idx_default = modos_disponiveis.index(default_modo)

modo_tela = option_menu(
    menu_title=None,
    options=modos_disponiveis,
    icons=["file-earmark-text", "file-earmark-spreadsheet", "pencil-square", "person-badge", "tools"],
    menu_icon="cast",
    default_index=idx_default,
    orientation="horizontal",
    styles=OPTION_MENU_STYLES,
)
st.session_state["modo_tela"] = modo_tela

modo_cdp_dedicado = modo_tela == "CDP - Ciclo I"
modo_reescrita_cdp_em = modo_tela == "Reescrita CDP"
modo_cadastro_professor = modo_tela == "Cadastro"
modo_diagnostico_modelos = modo_tela == "Diagnóstico"
if modo_cadastro_professor: _renderizar_cadastro_professor(PROFESSORES_DB); st.stop()
if modo_diagnostico_modelos: _renderizar_diagnostico_modelos(); st.stop()
if modo_reescrita_cdp_em: _renderizar_reescrita_cdp_em(); st.stop()

TEMPLATES_DIR = TEMPLATES_DOCX_DIR
TEMPLATES_DIR.mkdir(exist_ok=True)
templates_disponiveis = [f.name for f in TEMPLATES_DIR.glob("*.docx")]

OPCAO_MODELO_AUTOMATICO = "Automático pelo professor"
modelo_bytes = None
modelo_automatico_arquivo = ""
modelo_automatico_template_id = ""
escolha_template = "MODELOCDP.docx" if modo_cdp_dedicado else OPCAO_MODELO_AUTOMATICO
pdfs_aulas_files = []

st.markdown('<div class="section-title">🧠 Configuração de Inteligência</div>', unsafe_allow_html=True)
st.markdown('<div class="section-subtitle">Defina se o processamento será manual ou apoiado por IA.</div>', unsafe_allow_html=True)
modo_ia = st.radio("Motor de processamento", ["Sem IA", "OpenAI", "Gemini"], index=0, horizontal=True, key="modo_ia")
modelo_openai = os.environ.get("OPENAI_MODEL", MODELO_OPENAI_PADRAO) if modo_ia == "OpenAI" else ""
modelo_gemini = os.environ.get("GEMINI_MODEL", MODELO_GEMINI_PADRAO) if modo_ia == "Gemini" else ""

st.markdown('<div class="section-title">📝 Dados do Cabeçalho</div>', unsafe_allow_html=True)
st.markdown('<div class="section-subtitle">Preencha professor, disciplina, turma, período e dados que irão para o documento final.</div>', unsafe_allow_html=True)
col_prof, col_disciplina = st.columns([1, 1])
with col_prof:
    professor_selecionado = st.selectbox("Professor", _NOMES_PROFESSORES, key="professor_select")
    if professor_selecionado == "Outro (digitar)":
        professor = st.text_input("Nome do professor", key="professor").strip()
        if professor:
            from rapidfuzz import process, fuzz
            nomes_existentes = [p for p in PROFESSORES.keys() if p]
            if nomes_existentes:
                match = process.extractOne(professor, nomes_existentes, scorer=fuzz.WRatio)
                if match:
                    sugerido, score, _ = match
                    if 72 <= score < 100:
                        st.info(f"💡 **Dica de Digitação:** O nome digitado é semelhante ao professor cadastrado **'{sugerido}'** ({int(score)}% de similaridade). Se for ele, selecione-o no campo de seleção acima para evitar duplicidade.")
    else:
        professor = professor_selecionado if professor_selecionado != "(selecione o professor)" else ""

with col_disciplina:
    dados_prof = PROFESSORES_DB.get(professor, {})
    disciplinas_cadastradas = [] if modo_cdp_dedicado else dados_prof.get("disciplinas", [])
    disciplinas_gerais = [d for d in nomes_disciplinas() if not eh_cdp(d)]
    
    if disciplinas_cadastradas:
        disciplinas_unicas_prof = list(dict.fromkeys(d["disciplina"] for d in disciplinas_cadastradas))
        disc_selecionada = st.selectbox("Disciplina", ["(escolha a disciplina)"] + disciplinas_unicas_prof + ["Outra..."], key="disc_prof_select")
        disciplina = st.selectbox("Disciplina (Geral)", disciplinas_gerais, key="disciplina_opcao") if disc_selecionada == "Outra..." else (disc_selecionada if disc_selecionada != "(escolha a disciplina)" else "")
    else:
        disciplina = "" if modo_cdp_dedicado else st.selectbox("Disciplina", disciplinas_gerais, key="disciplina_opcao")

if disciplina == "Outra":
    disciplina = st.text_input("Informe a disciplina", key="disciplina_outra").strip()
    if disciplina:
        from rapidfuzz import process, fuzz
        match = process.extractOne(disciplina, disciplinas_gerais, scorer=fuzz.WRatio)
        if match:
            sugerido, score, _ = match
            if 70 <= score < 100:
                st.info(f"💡 **Dica de Digitação:** A disciplina digitada é semelhante a **'{sugerido}'** ({int(score)}% de similaridade). Se for ela, você pode usar o nome oficial para garantir que o sistema aplique as regras pedagógicas corretas.")
if modo_cdp_dedicado: disciplina = st.selectbox("Tipo de plano CDP", ["CDP- Multisseriada", "CDP - Ciclo I"], key="disciplina_cdp_opcao")

disciplina_config = obter_config(disciplina)
disciplina_norm = re.sub(r"\s+", " ", str(disciplina or "")).strip().lower()
orientacao_estudos = "orienta" in disciplina_norm and "estudo" in disciplina_norm
disciplina_cdp = eh_cdp(disciplina)

# Verificar disponibilidade das planilhas CDP
from core.cdp_legacy import PLANILHA_CDP, PLANILHA_CDP_MULTISSERIADA
if eh_cdp(disciplina) and not PLANILHA_CDP.exists():
    st.warning(
        f"⚠️ Planilha CDP não encontrada em: `{PLANILHA_CDP}`. "
        "O plano será gerado sem habilidades específicas. "
        "Verifique se o arquivo PLANILHACDP.xlsx está na pasta correta."
    )

if (disciplina_cdp or eh_cdp_fundamental(disciplina) or modo_cdp_dedicado) and escolha_template != "Upload de novo modelo...":
    modelo_cdp = TEMPLATES_DIR / "MODELOCDP.docx"
    if modelo_cdp.exists(): modelo_bytes = modelo_cdp.read_bytes()

config_turma_selecionada = None
col_turma, col_bimestre, col_mes, col_previstas = st.columns([2, 2, 2, 1])
with col_turma:
    turmas_cadastradas = [d["turma"] for d in dados_prof.get("disciplinas", []) if d["disciplina"] == disciplina]
    if turmas_cadastradas:
        turma_selecionada = st.selectbox("Série/Turma", ["(escolha a turma)"] + list(dict.fromkeys(turmas_cadastradas)) + ["Outra..."], key="turma_prof_select")
        if turma_selecionada == "Outra...": turma = _selecionar_turma("Série/Turma (Outra)", "turma_select", "turma")
        elif turma_selecionada == "(escolha a turma)": turma = ""
        else:
            turma = turma_selecionada
            st.session_state["turma"] = turma
            config_selecionada = next((d for d in dados_prof.get("disciplinas", []) if d["disciplina"] == disciplina and d["turma"] == turma), None)
            if config_selecionada:
                config_turma_selecionada = config_selecionada
                modelo_automatico_arquivo = str(config_selecionada.get("arquivo") or "")
                modelo_automatico_template_id = str(config_selecionada.get("template_id") or "")
                
                selecao_vaga_id = f"{professor}-{disciplina}-{turma}"
                if st.session_state.get("last_aula_prof") != selecao_vaga_id:
                    st.session_state["last_aula_prof"] = selecao_vaga_id
                    val_aulas = str(config_selecionada.get("aulas_semana") or "")
                    st.session_state["aulas_previstas_manual"] = val_aulas

                    if "aulas_previstas_manual_select" in st.session_state:
                        del st.session_state["aulas_previstas_manual_select"]
                        
                    datas_horarios = list(config_selecionada.get("datas_horarios") or [])
                    if datas_horarios:
                        for i, item in enumerate(datas_horarios):
                            st.session_state[f"data_aula_{i}"] = item.get("data", date.today())
                            sug = _sugerir_horario_cadastrado(" ".join(str(item.get(k) or "") for k in ("horario", "aula")), turma)
                            if sug: st.session_state[f"horario_aula_{i}"] = sug; st.session_state[f"tipo_horario_aula_{i}"] = _tipo_horario(sug)
    else:
        turma = _selecionar_turma("Série/Turma", "turma_select", "turma")

with col_bimestre: bimestre = st.selectbox("Bimestre", BIMESTRES, key="bimestre")
with col_mes: mes = _selecionar_mes()
with col_previstas: aulas_previstas_manual = _selecionar_aulas_semana("Aulas", "aulas_previstas_manual_select", "aulas_previstas_manual")

resumo_grade_cadastrada = _resumo_grade_cadastrada(config_turma_selecionada)
if resumo_grade_cadastrada:
    st.info(f"Horário cadastrado: {resumo_grade_cadastrada}", icon="🕒")

# ── Alerta: plano já gerado para outro professor ──────────────────────
if professor and disciplina and turma:
    outros = verificar_plano_gerado_por_outro_professor(
        professor,
        disciplina,
        turma,
        bimestre=bimestre,
    )
    if outros:
        nomes_outros = list(dict.fromkeys(r["professor_nome"] for r in outros))
        nomes_formatados = ", ".join(f"**{nome}**" for nome in nomes_outros[:3])
        data_recente = outros[0]["data_geracao"]
        try:
            if " " in str(data_recente):
                data_so = str(data_recente).split(" ")[0]
                partes = data_so.split("-")
                data_amigavel = f"{partes[2]}/{partes[1]}/{partes[0]}"
            else:
                data_amigavel = str(data_recente)
        except Exception:
            data_amigavel = str(data_recente)
            
        st.warning(
            f"⚠️ **ATENÇÃO:** O plano de **{disciplina} — {turma}** já foi gerado anteriormente neste **{bimestre}** para {nomes_formatados} (última geração em {data_amigavel}). Certifique-se de que é isso mesmo que deseja antes de prosseguir.",
            icon="⚠️"
        )


extensao_mes_rotulo = st.selectbox("Extensão após o mês", EXTENSAO_MES_OPCOES, index=0, key="extensao_mes")
extensao_mes = EXTENSAO_MES_VALORES.get(extensao_mes_rotulo, 0)
antecipacao_mes = EXTENSAO_MES_ANTECIPACOES.get(extensao_mes_rotulo, 0)

datas_horarios_mes, datas_sem_aula = [], []
config_agenda_mes = config_turma_selecionada
if modo_cdp_dedicado and modelo_bytes:
    config_agenda_mes = {**(config_turma_selecionada or {}), **_config_agenda_a_partir_do_modelo(modelo_bytes)}

if config_agenda_mes and mes and (not disciplina_cdp or modo_cdp_dedicado):
    datas_horarios_mes_base = _datas_horarios_do_mes(
        config_agenda_mes,
        mes,
        turma,
        extensao=extensao_mes,
        antecipacao=antecipacao_mes,
    )
    inicio_periodo = _inicio_periodo_mes_com_antecipacao(date.today().year, _mes_numero_app(mes), antecipacao_mes)
    fim_periodo = _fim_periodo_mes_com_extensao(date.today().year, _mes_numero_app(mes), extensao_mes)
    datas_opcoes_sem_aula = _datas_do_periodo(inicio_periodo, fim_periodo)
    
    assinatura_datas = f"{professor}|{disciplina}|{turma}|{mes}|{extensao_mes}|{antecipacao_mes}"
    if st.session_state.get("datas_sem_aula_assinatura") != assinatura_datas:
        st.session_state["datas_sem_aula_assinatura"] = assinatura_datas
        st.session_state["datas_sem_aula"] = _datas_feriado_padrao(datas_opcoes_sem_aula) or _datas_sem_aula_padrao(datas_horarios_mes_base)

    if datas_opcoes_sem_aula:
        datas_sem_aula = st.multiselect("Dias sem aula", options=datas_opcoes_sem_aula, format_func=_rotulo_data_sem_aula, key="datas_sem_aula")
    
    deixar_antecipacao_vazia = False
    if antecipacao_mes > 0:
        deixar_antecipacao_vazia = st.checkbox(
            "Deixar bloco da semana extra vazio",
            key="deixar_antecipacao_vazia",
            value=bool(st.session_state.get("deixar_antecipacao_vazia", False)),
            help="Se marcado, as aulas da semana extra anterior ao início do mês virão completamente em branco no plano, e a correspondência com os PDFs começará a partir da primeira data oficial do mês."
        )
    
    datas_horarios_mes = _sincronizar_datas_horarios_mes(
        config_agenda_mes,
        mes,
        professor,
        disciplina,
        turma,
        extensao=extensao_mes,
        antecipacao=antecipacao_mes,
        datas_sem_aula=datas_sem_aula,
    )

if professor and disciplina and turma and not disciplina_cdp and escolha_template == OPCAO_MODELO_AUTOMATICO:
    template_id_central = resolver_template_id_geracao(
        template_id=modelo_automatico_template_id or "",
        disciplina=disciplina,
        componente_curricular=str((config_turma_selecionada or {}).get("componente_curricular") or disciplina),
        escola=st.session_state.get("escola", ""),
        arquivo_modelo=modelo_automatico_arquivo,
    )
    caminho_template = caminho_template_central(template_id_central)
    if caminho_template.exists(): modelo_bytes = caminho_template.read_bytes()

if bool(professor and disciplina and turma and not disciplina_cdp and not modelo_bytes):
    st.markdown('<div class="section-title">Modelo DOCX</div>', unsafe_allow_html=True)
    escolha_template = st.selectbox("Modelo DOCX Base", templates_disponiveis + ["Upload de novo modelo..."], key="escolha_template_manual")
    if escolha_template == "Upload de novo modelo...":
        modelo_file = st.file_uploader("Novo Modelo", type=["docx"], key="novo_modelo_file")
        if modelo_file:
            modelo_bytes = modelo_file.getvalue()
            destino_modelo = TEMPLATES_DIR / modelo_file.name
            modelo_existe = destino_modelo.exists()
            confirmar_modelo = True
            if modelo_existe:
                confirmar_modelo = st.checkbox(
                    f"Confirmo que desejo substituir o modelo existente '{modelo_file.name}'.",
                    key="confirmar_substituir_modelo_manual",
                )
                st.warning("Ja existe um modelo com esse nome. Marque a confirmacao para substituir.")
            if st.button("Salvar para futuro", disabled=modelo_existe and not confirmar_modelo):
                destino_modelo.write_bytes(modelo_bytes)
                st.rerun()
    else:
        if (TEMPLATES_DIR / escolha_template).exists(): modelo_bytes = (TEMPLATES_DIR / escolha_template).read_bytes()

if professor and disciplina and turma:
    st.button("Editar cadastro", type="secondary", on_click=_abrir_cadastro_com_filtros, args=(professor, disciplina, turma))

assinatura_comp = f"{professor}|{disciplina}|{turma}|{(config_turma_selecionada or {}).get('componente_curricular', '')}"
if st.session_state.get("last_componente_curricular") != assinatura_comp:
    st.session_state["last_componente_curricular"] = assinatura_comp
    st.session_state["componente_curricular"] = str((config_turma_selecionada or {}).get("componente_curricular") or disciplina)

col_escola, col_comp = st.columns([1, 1])
with col_escola: escola = st.selectbox("Escola", ["EE PROFª. EGLE LUPORINI COSTA", "PADRE GERALDO LOURENÇO"], key="escola")
with col_comp: componente_curricular = st.text_input("Componente curricular", key="componente_curricular")

modalidade_eja = False
if not disciplina_cdp and _disciplina_suporta_modalidade_eja(disciplina):
    modalidade_eja = st.selectbox("Modalidade", ["Regular", "EJA"], key="modalidade_eja") == "EJA"


def _resolver_caminho_ae_priorizado(disciplina: str, turma: str, bimestre: str, professor: str = "") -> str:
    try:
        pasta = resolver_pasta_pdfs(r"D:\PDF novos", disciplina, turma, bimestre, professor=professor)
    except Exception:
        return ""

    if not getattr(pasta, "exists", lambda: False)():
        return ""

    candidatos = []
    padroes = ["GUIA*.xlsx", "*GUIA*.xlsx", "planilha.xlsx", "*.xlsx"]
    for padrao in padroes:
        for arquivo in sorted(pasta.glob(padrao)):
            nome = str(getattr(arquivo, "name", "") or "")
            if nome.startswith("~$"):
                continue
            if arquivo not in candidatos:
                candidatos.append(arquivo)

    return str(candidatos[0]) if candidatos else ""


caminho_ae_priorizado = _resolver_caminho_ae_priorizado(disciplina, turma, bimestre, professor)
st.session_state["caminho_ae_priorizado"] = caminho_ae_priorizado
usar_ae_priorizado = False
contexto_ae_ok = False
if disciplina_ae_priorizado_disponivel(disciplina) or caminho_ae_priorizado:
    contexto_ae_ok = contexto_ae_priorizado_disponivel(
        disciplina,
        turma,
        bimestre,
        caminho_planilha=caminho_ae_priorizado,
    )
    st.checkbox(
        "Usar AE no lugar da habilidade",
        value=bool(st.session_state.get("usar_ae_priorizado", False)),
        key="usar_ae_priorizado",
        disabled=not contexto_ae_ok,
        help="Quando ativado, o sistema troca a coluna de aprendizagem pelo AE correspondente do guia priorizado, quando houver base cadastrada para a disciplina, turma e bimestre.",
    )
    usar_ae_priorizado = bool(contexto_ae_ok and st.session_state.get("usar_ae_priorizado", False))
    if contexto_ae_ok:
        if caminho_ae_priorizado:
            st.caption("Guia priorizado encontrado para este contexto. Se alguma aula nao estiver na planilha, o sistema mantém a habilidade normal.")
        else:
            st.caption("Base AE encontrada para este contexto. Se alguma aula não estiver no mapa, o sistema mantém a habilidade normal.")
    else:
        st.caption("Esta opção fica disponível quando existe base AE ou guia priorizado para a disciplina, série e bimestre selecionados.")
else:
    st.session_state["usar_ae_priorizado"] = False

def _resumo_tela(valor: str, fallback: str = "Não definido") -> str:
    return str(valor).strip() if str(valor or "").strip() else fallback


def _rotulo_sequencia_pdfs_esperada(numeros: list[int]) -> str:
    return " | ".join(
        f"{indice + 1}. AULA {int(numero)}"
        for indice, numero in enumerate(numeros or [])
        if str(numero).strip()
    )


def _limitar_sequencia_ae(numeros: list[int], limite: int | None = None) -> list[int]:
    if limite is None:
        return list(numeros or [])
    try:
        limite_int = int(limite)
    except (TypeError, ValueError):
        limite_int = 0
    if limite_int <= 0:
        return list(numeros or [])
    return list(numeros or [])[:limite_int]


def _nome_pdf_para_tela(arquivo) -> str:
    return html.escape(str(getattr(arquivo, "name", None) or Path(str(arquivo)).name))


def _render_painel_pdfs(
    *,
    modo: str,
    necessarios: int,
    carregados: int,
    total_aulas: int = 0,
    dividir_metodologia: bool = False,
    encontrados: int = 0,
    pasta: str = "",
    selecionados=None,
    faltantes_ae=None,
) -> None:
    selecionados = list(selecionados or [])
    faltantes_ae = list(faltantes_ae or [])
    necessarios = max(0, int(necessarios or 0))
    carregados = max(0, int(carregados or 0))
    total_aulas = max(0, int(total_aulas or 0))
    encontrados = max(0, int(encontrados or 0))
    modo_texto = str(modo or "-").strip()
    if modo_texto != "Automatico":
        pasta = ""
        encontrados = 0
        faltantes_ae = []
    faltam = max(necessarios - carregados, 0)
    excedentes = max(carregados - necessarios, 0)
    progresso = 0 if necessarios <= 0 else min(100, int(round((carregados / necessarios) * 100)))

    if necessarios <= 0:
        status_texto = "Aguardando modelo"
        status_classe = "neutral"
        orientacao = "Selecione professor, turma e modelo para o sistema calcular quantos PDFs serao usados."
    elif faltam > 0:
        status_texto = f"Faltam {faltam}"
        status_classe = "warning"
        orientacao = f"Adicione mais {faltam} PDF(s) para completar o plano."
    elif excedentes > 0:
        status_texto = f"{excedentes} a mais"
        status_classe = "warning"
        orientacao = "Revise a selecao: ha mais PDFs do que a organizacao atual exige."
    else:
        status_texto = "Completo"
        status_classe = "success"
        orientacao = "Tudo certo: a quantidade de PDFs bate com a organizacao escolhida."

    criterio_pdfs = "1 PDF para cada par de aulas marcado" if dividir_metodologia else "1 PDF por aula"
    aulas_rotulo = total_aulas or necessarios

    with st.container():
        st.markdown("### Painel dos PDFs")
        st.caption(orientacao)

        if status_classe == "success":
            st.success(status_texto)
        elif status_classe == "warning":
            st.warning(status_texto)
        else:
            st.info(status_texto)

        col1, col2, col3, col4, col5 = st.columns(5)
        col1.metric("Modo", modo_texto)
        col2.metric("Aulas previstas", aulas_rotulo)
        col3.metric("PDFs necessarios", necessarios)
        col4.metric("Selecionados", carregados)
        col5.metric("Encontrados", encontrados)

        st.progress(progresso)
        st.caption(f"{carregados}/{necessarios or 0} PDF(s) prontos para processamento | {criterio_pdfs}")

        if pasta:
            st.caption(f"Pasta automatica: {pasta}")

        if faltantes_ae:
            faltantes_txt = ", ".join(f"AULA {int(numero)}" for numero in faltantes_ae)
            st.warning(f"PDFs AE nao encontrados: {faltantes_txt}")

        st.markdown("**Ordem que sera processada**")
        if selecionados:
            for indice, item in enumerate(selecionados, start=1):
                st.write(f"{indice}. {_nome_pdf_para_tela(item)}")
        else:
            st.caption("Nenhum PDF selecionado ainda.")


sequencia_ae_contexto = []
if usar_ae_priorizado and contexto_ae_ok:
    sequencia_ae_contexto = sequencia_aulas_ae_priorizado(
        disciplina,
        turma,
        bimestre,
        caminho_planilha=caminho_ae_priorizado,
    )
    if sequencia_ae_contexto:
        st.info(
            "Modo AE ativo neste contexto. Ordem base do guia priorizado: "
            f"{_rotulo_sequencia_pdfs_esperada(sequencia_ae_contexto)}."
        )
        st.caption("Mais abaixo, o envio dos PDFs do mês usará essa mesma ordem.")
    else:
        st.warning("Modo AE ativo, mas não encontrei a sequência do guia para este contexto.")


def _render_previa_aulas_cdp(preview: list[dict]):
    if not preview:
        st.warning("Não consegui localizar as aulas do CDP no modelo atual.")
        return

    st.markdown('<div class="section-subtitle">Prévia das aulas que serão puxadas da planilha</div>', unsafe_allow_html=True)
    st.info(
        "No CDP, a aula inicial é aplicada dentro de cada disciplina que aparece no modelo. A prévia abaixo mostra qual disciplina e qual aula da planilha entrarão em cada bloco.",
        icon="ℹ️",
    )

    cards = []
    for item in preview[:6]:
        aula_planilha = f"Aula {item.get('aula_planilha')}" if str(item.get("aula_planilha") or "").strip() else "Aula sem número"
        titulo = str(item.get("titulo") or "").strip() or "Sem título identificado"
        disciplina = str(item.get("disciplina") or "").strip() or "Disciplina não identificada"
        planilha = str(item.get("componente_planilha") or disciplina).strip()
        cards.append(
            (
                f'<div class="cdp-preview-card">'
                f'<div class="cdp-preview-card__top">'
                f'<span class="cdp-preview-card__ordem">Bloco {item.get("ordem")}</span>'
                f'<span class="cdp-preview-card__aula">{aula_planilha}</span>'
                f'</div>'
                f'<div class="cdp-preview-card__disciplina">{disciplina}</div>'
                f'<div class="cdp-preview-card__planilha">Planilha: {planilha}</div>'
                f'<div class="cdp-preview-card__titulo">{titulo}</div>'
                f'</div>'
            )
        )

    st.markdown(f'<div class="cdp-preview-grid">{"".join(cards)}</div>', unsafe_allow_html=True)
    if len(preview) > 6:
        st.caption(f"Mostrando os 6 primeiros blocos do modelo. Total identificado: {len(preview)}.")

semana = ""

# Pre-populate observation for August with the current fixed calendar text
default_agosto = (
    "Replanejamento – 22 e 23.07;\n"
    "Período do plano – 24.07 até 31.08;\n"
    "Reunião de pais/responsáveis – 04.08;\n"
    "Feriado Municipal – 06.08;"
)
default_agosto_legado = "06/08 - Aniversário da cidade\n07/08 - Ponto facultativo"
observacoes_automaticas_agosto = {default_agosto, default_agosto_legado}

if "observacao" not in st.session_state or not st.session_state["observacao"]:
    if mes.strip().upper() == "AGOSTO":
        st.session_state["observacao"] = default_agosto
elif (
    mes.strip().upper() == "AGOSTO"
    and str(st.session_state.get("observacao", "") or "").strip() == default_agosto_legado
):
    st.session_state["observacao"] = default_agosto

if "last_mes_for_obs" not in st.session_state:
    st.session_state["last_mes_for_obs"] = mes

if st.session_state["last_mes_for_obs"] != mes:
    current_obs = st.session_state.get("observacao", "")
    if mes.strip().upper() == "AGOSTO":
        if not current_obs or current_obs.strip() in {"", default_agosto_legado}:
            st.session_state["observacao"] = default_agosto
    else:
        if current_obs.strip() in observacoes_automaticas_agosto:
            st.session_state["observacao"] = ""
    st.session_state["last_mes_for_obs"] = mes

observacao = st.text_area("Observação", key="observacao")
gerar_turma_espelho = st.checkbox("Gerar para 2ª turma", value=False, key="gerar_turma_espelho")
turma_espelho = _selecionar_turma_espelho(turma, turmas_cadastradas) if gerar_turma_espelho else ""
aulas_envio_espelho = []
# Buscar configuração cadastrada da turma espelho para usar os horários corretos
config_turma_espelho = None
if gerar_turma_espelho and turma_espelho:
    config_turma_espelho = next(
        (d for d in dados_prof.get("disciplinas", [])
         if d["disciplina"] == disciplina and d["turma"] == turma_espelho),
        None,
    )

st.markdown('<div class="section-title">📚 Gestão das Aulas</div>', unsafe_allow_html=True)
st.markdown('<div class="section-subtitle">Confira os PDFs necessários, a ordem de processamento e o que ainda falta antes de gerar o plano.</div>', unsafe_allow_html=True)
if disciplina_cdp:
    if eh_cdp_multisseriada(disciplina):
        col1, col2 = st.columns([2, 1])
        with col1: turma_cdp = st.selectbox("Turma filtro", TURMAS_CDP_MULTISSERIADA, key="turma_cdp")
        with col2: cdp_aula_inicial = st.number_input("Aula inicial", min_value=1, value=1, key="cdp_aula_inicial")
    else: cdp_aula_inicial = st.number_input("Aula inicial", min_value=1, value=1, key="cdp_aula_inicial")
    if modelo_bytes:
        try:
            previa_cdp = prever_aulas_cdp(
                BytesIO(modelo_bytes),
                aula_inicial=int(cdp_aula_inicial or 1),
                fundamental=eh_cdp_fundamental(disciplina),
                multisseriada=eh_cdp_multisseriada(disciplina),
                serie_cdp=turma_cdp if eh_cdp_multisseriada(disciplina) else "",
                bimestre=bimestre,
            )
            _render_previa_aulas_cdp(previa_cdp)
        except Exception:
            st.warning("Não consegui montar a prévia das aulas do CDP com o modelo atual.")
else:
    linhas_modelo = len(datas_horarios_mes or []) or len((config_turma_selecionada or {}).get("datas_horarios") or [])
    sequencia_pdf_esperada_ae = []
    contexto_divisao_pdf = "|".join(str(valor or "") for valor in [professor, disciplina, turma, mes, bimestre])

    if bool(len(datas_horarios_mes or [])):
        st.session_state["auto_repetir_semana"] = False
    elif "auto_repetir_semana" not in st.session_state:
        st.session_state["auto_repetir_semana"] = True
    auto_repetir_semana = st.checkbox("Repetir semana", key="auto_repetir_semana", disabled=bool(len(datas_horarios_mes or [])))
    dividir_metodologia = st.checkbox("Dividir metodologia em dois dias", value=False, key="dividir_metodologia")
    _sincronizar_divisao_pdf_padrao(linhas_modelo, dividir_metodologia, contexto=contexto_divisao_pdf)

    opcoes_modo_upload = ["Automatico", "Todos de uma vez", "Um por aula"]
    if st.session_state.get("modo_upload_pdf") not in opcoes_modo_upload:
        st.session_state["modo_upload_pdf"] = MODO_UPLOAD_PDF_PADRAO
    modo_upload_pdf = st.radio(
        "Modo de envio dos PDFs",
        opcoes_modo_upload,
        horizontal=True,
        key="modo_upload_pdf",
        help=(
            "Automatico: busca os PDFs em D:\\PDF novos e aplica a ordem do sistema.\n"
            "Todos de uma vez: envie os arquivos manualmente em lote.\n"
            "Um por aula: envie o PDF diretamente em cada card."
        ),
    )
    modo_upload_individual = modo_upload_pdf == "Um por aula"
    modo_upload_automatico = modo_upload_pdf == "Automatico"

    pdfs_aulas_files = []
    qtd_aulas = 0
    pdfs_auto_total = 0
    try:
        pasta_pdfs_auto = str(resolver_pasta_pdfs(r"D:\PDF novos", disciplina, turma, bimestre, professor=professor))
    except Exception:
        pasta_pdfs_auto = ""
    faltantes_ae_auto = []
    pdfs_selecionados_tela = []

    deixar_ant_vazia = bool(st.session_state.get("deixar_antecipacao_vazia", False))
    if deixar_ant_vazia and datas_horarios_mes:
        aulas_oficiais_modelo = [d for d in datas_horarios_mes if not _eh_data_antecipacao(d["data"], mes, antecipacao_mes)]
        linhas_modelo_pdf = len(aulas_oficiais_modelo)
    else:
        linhas_modelo_pdf = linhas_modelo

    if not modo_upload_individual:
        # Calcular PDFs necessários estimados para o rótulo do uploader
        est_necessarios = 0
        if linhas_modelo_pdf > 0:
            est_necessarios = _estimar_pdfs_por_estado(linhas_modelo_pdf, dividir_metodologia)

        if usar_ae_priorizado and sequencia_ae_contexto:
            sequencia_pdf_esperada_ae = _limitar_sequencia_ae(
                sequencia_ae_contexto,
                est_necessarios or linhas_modelo_pdf,
            )

        label_uploader = "Envio Manual de PDFs"
        if est_necessarios > 0:
            label_uploader = f"Envio Manual (Insira exatamente {est_necessarios} PDF(s) para as {linhas_modelo_pdf} aulas do mês)"

        if sequencia_pdf_esperada_ae:
            st.info(
                "Modo AE ativo. Sequencia esperada dos PDFs neste contexto: "
                f"{_rotulo_sequencia_pdfs_esperada(sequencia_pdf_esperada_ae)}."
            )
            st.caption("Envie os arquivos nessa ordem do guia priorizado.")

        # Busca automatica de PDFs locais ou envio manual, conforme modo escolhido.
        if modo_upload_automatico:
            base_pdfs_dir = r"D:\PDF novos"
            pasta_pdfs = resolver_pasta_pdfs(base_pdfs_dir, disciplina, turma, bimestre, professor=professor)
            pasta_pdfs_auto = str(pasta_pdfs)

            pdf_files_disponiveis = []
            if pasta_pdfs.exists():
                pdfs_encontrados = filtrar_pdfs_para_aulas(pasta_pdfs.glob("*.pdf"))
                pdfs_auto_total = len(pdfs_encontrados)
                pdfs_com_numero = [pdf for pdf in pdfs_encontrados if numero_aula_pdf(pdf) is not None]
                pdfs_para_ordenar = pdfs_com_numero or pdfs_encontrados
                pdf_files_disponiveis = ordenar_pdfs_por_numero(pdfs_para_ordenar)

            if pdf_files_disponiveis:
                default_selection = []
                from core.database import obter_ultima_aula_gerada_sistema
                ultima_aula = obter_ultima_aula_gerada_sistema(professor, disciplina, turma, bimestre)
                
                # A continuidade automática foi desativada; a seleção padrão volta à Aula 1.
                pdf_files_filtrados = []
                for p in pdf_files_disponiveis:
                    num_aula = numero_aula_pdf(p)
                    if num_aula is not None and num_aula > ultima_aula:
                        pdf_files_filtrados.append(p)
                
                # Fallback se a lista filtrada ficar vazia
                if not pdf_files_filtrados:
                    pdf_files_filtrados = pdf_files_disponiveis

                if est_necessarios > 0:
                    default_selection = pdf_files_filtrados[:est_necessarios]

                faltantes_ae_auto = numeros_pdfs_faltantes(pdf_files_disponiveis, sequencia_pdf_esperada_ae)
                assinatura_pdfs_auto = _assinatura_pdfs_automaticos(pdf_files_disponiveis)

                chave_contexto_auto = "_".join(
                    [
                        "pdfs_aulas_files_auto_v3",
                        normalizar_para_pasta(disciplina),
                        normalizar_para_pasta(professor),
                        normalizar_para_pasta(turma),
                        normalizar_para_pasta(bimestre),
                        "-".join(str(numero) for numero in sequencia_pdf_esperada_ae) or "sem_ae",
                        f"ultima_{ultima_aula}",
                        str(est_necessarios or 0),
                        f"pdfs_{assinatura_pdfs_auto}",
                    ]
                )
                if est_necessarios > 0:
                    st.markdown(f"<div style='background-color: #ffe6e6; border: 2px solid #ff4b4b; padding: 15px; border-radius: 8px; text-align: center; margin-bottom: 15px; color: #2b0f14 !important;'><h3 style='color: #c81e2d !important; margin: 0;'>🚨 QUANTIDADE NECESSÁRIA: {est_necessarios} PDFs 🚨</h3><p style='color: #2b0f14 !important; margin-top: 5px; font-weight: 700;'>O sistema precisa de exatamente {est_necessarios} PDFs para montar o plano deste mês.</p></div>", unsafe_allow_html=True)

                selecionados = st.multiselect(
                    "PDFs automaticos na ordem de processamento",
                    options=pdf_files_disponiveis,
                    format_func=lambda p: p.name,
                    default=default_selection,
                    key=chave_contexto_auto,
                    help="A ordem abaixo ja e a ordem que o sistema vai usar. No modo AE, ela segue a sequencia do guia priorizado.",
                )
                selecionados = ordenar_pdfs_por_numero(selecionados)
                pdfs_selecionados_tela = list(selecionados)
                pdfs_aulas_files = [LocalFileWrapper(p) for p in selecionados]
        else:
            if est_necessarios > 0:
                st.markdown(
                    f"<div style='background-color: #ffe6e6; border: 2px solid #ff4b4b; padding: 15px; border-radius: 8px; text-align: center; margin-bottom: 15px; color: #2b0f14 !important;'><h3 style='color: #c81e2d !important; margin: 0;'>🚨 QUANTIDADE NECESSÁRIA: {est_necessarios} PDFs 🚨</h3><p style='color: #2b0f14 !important; margin-top: 5px; font-weight: 700;'>O sistema precisa de exatamente {est_necessarios} PDFs para montar o plano deste mês.</p></div>",
                    unsafe_allow_html=True,
                )
            pdfs_aulas_files = st.file_uploader(
                label_uploader,
                type=["pdf"],
                accept_multiple_files=True,
                key="pdfs_aulas_files",
                help="Envie todos os PDFs do plano em lote, na ordem em que devem ser processados.",
            ) or []
            qtd_aulas = len(pdfs_aulas_files)
            pdfs_selecionados_tela = list(pdfs_aulas_files)

    deixar_ant_vazia = bool(st.session_state.get("deixar_antecipacao_vazia", False))

    num_rows = linhas_modelo or int(qtd_aulas) * (2 if dividir_metodologia else 1)
    _sincronizar_divisao_pdf_padrao(num_rows, dividir_metodologia, contexto=contexto_divisao_pdf)
    aulas_envio = _coletar_aulas_envio(
        num_rows,
        pdfs_aulas_files,
        dividir_metodologia,
        auto_repetir_semana,
        replicar_pdf_unico=bool(orientacao_estudos and qtd_aulas == 1),
        modo_upload_individual=modo_upload_individual,
        preservar_datas_sincronizadas=bool(datas_horarios_mes),
        sequencia_pdf_esperada=sequencia_pdf_esperada_ae,
        deixar_antecipacao_vazia=deixar_ant_vazia,
        mes=mes,
        antecipacao_mes=antecipacao_mes,
    )

    aulas_para_pdf = [a for a in aulas_envio if not _eh_data_antecipacao(a["data"], mes, antecipacao_mes)] if deixar_ant_vazia else aulas_envio

    if modo_upload_individual:
        grupos_individuais = _grupos_pdf_por_aula(aulas_para_pdf) if dividir_metodologia else [{"indices": [idx]} for idx in range(len(aulas_para_pdf))]
        pdfs_individuais = [
            aulas_para_pdf[grupo["indices"][0]].get("pdf")
            for grupo in grupos_individuais
            if aulas_para_pdf[grupo["indices"][0]].get("pdf") is not None
        ]
        pdfs_necessarios = len(grupos_individuais) if dividir_metodologia else len(aulas_para_pdf)
        pdfs_prontos = len(pdfs_individuais)
        _render_painel_pdfs(
            modo=modo_upload_pdf,
            necessarios=pdfs_necessarios,
            carregados=pdfs_prontos,
            total_aulas=len(aulas_para_pdf),
            dividir_metodologia=dividir_metodologia,
            selecionados=pdfs_individuais,
        )
    else:
        pdfs_necessarios = len(_grupos_pdf_por_aula(aulas_para_pdf)) if dividir_metodologia else len(aulas_para_pdf)

        if linhas_modelo_pdf > 0:
            pdf_unico_orientacao = bool(orientacao_estudos and qtd_aulas == 1 and pdfs_necessarios >= 1)
            if qtd_aulas == pdfs_necessarios or pdf_unico_orientacao:
                st.success(f"Quantidade de PDFs correta: {qtd_aulas}/{pdfs_necessarios} PDF(s) carregado(s).")
            elif qtd_aulas > 0:
                st.warning(f"Quantidade de PDFs incorreta: foram adicionados {qtd_aulas}, mas o plano requer exatamente {pdfs_necessarios} PDF(s).")
            else:
                st.info(f"Aguardando o envio de {pdfs_necessarios} PDF(s) para {linhas_modelo_pdf} aula(s).")

    if gerar_turma_espelho:
        # Determine num_rows_espelho based on 2nd class config if month is selected
        num_rows_espelho = num_rows
        datas_horarios_mes_espelho = None
        
        if config_turma_espelho and mes:
            datas_horarios_mes_espelho = _sincronizar_datas_horarios_mes_turma2(
                config_turma_espelho, mes, professor, disciplina, turma_espelho,
                extensao=extensao_mes, antecipacao=antecipacao_mes, datas_sem_aula=datas_sem_aula,
            )
            linhas_modelo_espelho = len(datas_horarios_mes_espelho)
            if linhas_modelo_espelho > 0:
                num_rows_espelho = linhas_modelo_espelho

        contexto_divisao_pdf_espelho = f"{contexto_divisao_pdf}|{turma_espelho}"
        _sincronizar_divisao_pdf_padrao(num_rows_espelho, dividir_metodologia, key_prefix="turma2_", contexto=contexto_divisao_pdf_espelho)
        
        aulas_envio_espelho = _coletar_aulas_envio(
            num_rows_espelho,
            pdfs_aulas_files,
            dividir_metodologia,
            auto_repetir_semana,
            replicar_pdf_unico=bool(orientacao_estudos and qtd_aulas == 1),
            key_prefix="turma2_",
            titulo_secao="2ª turma",
            modo_upload_individual=modo_upload_individual,
            preservar_datas_sincronizadas=bool(datas_horarios_mes_espelho),
            sequencia_pdf_esperada=sequencia_pdf_esperada_ae,
        )

st.markdown('<div class="section-title">🚀 Passo 1: Extração e Processamento</div>', unsafe_allow_html=True)
st.markdown('<div class="section-subtitle">Confira a organização das aulas e inicie o processamento para transformar os PDFs em blocos prontos para revisão.</div>', unsafe_allow_html=True)
st.markdown(
    """
    <div class="process-panel">
        <div class="panel-title">Tudo pronto para processar</div>
        <div class="panel-text">Revise rapidamente as datas, horários e PDFs vinculados. Quando estiver ok, o sistema prepara as aulas para você revisar antes de gerar o documento final.</div>
        <div class="panel-pills">
            <span class="panel-pill">Fluxo guiado</span>
            <span class="panel-pill">Menos retrabalho</span>
            <span class="panel-pill">Revisão antes do DOCX</span>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)
erro_processamento = str(st.session_state.get("erro_processamento") or "").strip()
if erro_processamento:
    st.error(erro_processamento)
    erro_processamento_detalhe = str(st.session_state.get("erro_processamento_detalhe") or "").strip()
    if erro_processamento_detalhe:
        with st.expander("Ver detalhe tecnico"):
            st.code(erro_processamento_detalhe)
geracao_em_andamento = bool(st.session_state.get("geracao_em_andamento", False))
if disciplina_cdp:
    st.checkbox(
        "Salvar este plano no histórico",
        key="salvar_historico_geracao",
        value=bool(st.session_state.get("salvar_historico_geracao", False)),
        help="Marque apenas quando o plano estiver realmente ok. Isso guarda o DOCX para consulta futura, sem alterar a aula inicial dos próximos planos.",
    )
if st.button("PROCESSAR AULAS" if not disciplina_cdp else "GERAR PLANO", disabled=geracao_em_andamento, type="primary"):
    _limpar_erro_processamento()
    st.session_state["geracao_em_andamento"] = True
    pdfs_enviados_val = sum(1 for a in aulas_envio if a.get("pdf") is not None) if (not disciplina_cdp and st.session_state.get("modo_upload_pdf") == "Um por aula") else len(pdfs_aulas_files or [])
    deixar_ant_vazia = bool(st.session_state.get("deixar_antecipacao_vazia", False))
    erro = validar_entrada(
        modelo_bytes, disciplina, disciplina_config, aulas_envio, professor, turma,
        bimestre, mes, aulas_previstas_manual, pdfs_enviados_val, pdfs_necessarios,
        deixar_antecipacao_vazia=deixar_ant_vazia,
        antecipacao_mes=antecipacao_mes,
    )
    if not erro:
        erro = validar_aulas_secundarias(
            gerar_turma_espelho,
            turma_espelho,
            aulas_envio_espelho if not disciplina_cdp else [],
            bool(disciplina_config.exige_pdf and not disciplina_cdp),
        )
    if erro:
        st.error(erro); st.session_state["geracao_em_andamento"] = False
    elif disciplina_cdp:
        planos_gerados = []
        with st.status("📚 Gerando...", expanded=True) as status:
            for t in [turma] + ([turma_espelho] if gerar_turma_espelho else []):
                planos_gerados.append(_gerar_docx_cdp_final(modelo_bytes, escola, professor, disciplina, t, mes, bimestre, semana, observacao, aulas_previstas_manual, cdp_aula_inicial, turma_cdp, modo_ia, modelo_openai, modelo_gemini, datas_horarios_mes))
            st.session_state["planos_gerados"] = planos_gerados
            salvou_historico = _salvar_planos_gerados_se_configurado(
                planos_gerados,
                professor,
                disciplina,
                bimestre,
            )
            _registrar_mensagem_memoria_plano(salvou_historico)
            status.update(label="✅ Concluído", state="complete", expanded=False)
        st.session_state["geracao_em_andamento"] = False; st.rerun()
    else:
        turmas_processadas, avisos = [], []
        blocos_processamento = [(turma, aulas_envio)] + ([(turma_espelho, aulas_envio_espelho)] if gerar_turma_espelho else [])
        total_pdfs_processamento = sum(
            len(_grupos_pdf_por_aula(aulas_bloco) if dividir_metodologia else aulas_bloco)
            for _, aulas_bloco in blocos_processamento
        )
        progresso_estado = {"atual": 0}
        with st.status("⏳ Extraindo...", expanded=True) as status:
            progress_bar = st.progress(0, text="Preparando os PDFs para extração...")
            try:
                for t, a in blocos_processamento:
                    def _callback_pdf(indice_pdf, total_pdf_turma, caminho_pdf, turma_atual=t):
                        progresso_estado["atual"] += 1
                        total_base = max(1, total_pdfs_processamento)
                        pct = min(100, int(round((progresso_estado["atual"] / total_base) * 100)))
                        nome_pdf = Path(str(caminho_pdf)).name
                        progress_bar.progress(
                            pct,
                            text=f"Processando {progresso_estado['atual']}/{total_base} PDF(s) • {turma_atual} • {nome_pdf}",
                        )
                        st.write(f"✓ Aula {indice_pdf + 1}: {nome_pdf} processada para {turma_atual}")

                    res = _extrair_aulas_dos_pdfs(
                        a,
                        disciplina,
                        t,
                        bimestre,
                        modo_ia,
                        modelo_openai,
                        modelo_gemini,
                        dividir_metodologia,
                        modalidade_eja,
                        usar_ae_priorizado=usar_ae_priorizado,
                        progress_callback=_callback_pdf,
                        professor=professor,
                        deixar_antecipacao_vazia=deixar_ant_vazia,
                        mes=mes,
                        antecipacao_mes=antecipacao_mes,
                    )
                    turmas_processadas.append({"turma": t, "aulas": res["aulas"]})
                    avisos_turma = []
                    if res.get("avisos_repeticao"):
                        avisos_turma.extend(res["avisos_repeticao"])
                    if res.get("avisos_ae"):
                        avisos_turma.extend(res["avisos_ae"])
                    if res.get("avisos_ia"):
                        avisos_turma.extend(res["avisos_ia"])
                    if avisos_turma:
                        avisos.append({"turma": t, "avisos": avisos_turma})
                progress_bar.progress(100, text="Extração concluída. Preparando a revisão...")
                status.update(label="✅ Extraído!", state="complete", expanded=False)
                st.session_state["turmas_processadas"] = turmas_processadas
                st.session_state["avisos_processamento"] = avisos
                st.session_state["revisao_token"] = st.session_state.get("revisao_token", 0) + 1
            except Exception as e:
                _registrar_erro_processamento(e)
            st.session_state["geracao_em_andamento"] = False
            st.rerun()

if st.session_state.get("turmas_processadas"):
    avisos_processamento = st.session_state.get("avisos_processamento") or []
    for bloco in avisos_processamento:
        turma_aviso = str(bloco.get("turma") or "").strip()
        avisos_turma = [str(aviso).strip() for aviso in bloco.get("avisos", []) if str(aviso).strip()]
        if avisos_turma:
            prefixo = f"{turma_aviso}: " if turma_aviso else ""
            st.warning(prefixo + " | ".join(avisos_turma))
    st.markdown('<div class="section-title">✏️ Passo 2: Revisão</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-subtitle">Ajuste tema, aprendizagem, metodologia, acompanhamento e acessibilidade antes de montar o arquivo final.</div>', unsafe_allow_html=True)
    total_turmas_revisao = len(st.session_state["turmas_processadas"])
    total_aulas_revisao = sum(len(td.get("aulas", [])) for td in st.session_state["turmas_processadas"])
    st.markdown(
        f"""
        <div class="review-shell">
            <div class="panel-title">Revisão pedagógica centralizada</div>
            <div class="panel-text">Você está revisando <strong>{total_aulas_revisao}</strong> aula(s) distribuídas em <strong>{total_turmas_revisao}</strong> turma(s). Abra apenas os blocos que quiser ajustar.</div>
            <div class="panel-pills">
                <span class="panel-pill">Tema</span>
                <span class="panel-pill">Metodologia</span>
                <span class="panel-pill">Acompanhamento</span>
                <span class="panel-pill">Acessibilidade</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    turmas_revisadas = []
    rev_tok = st.session_state.get("revisao_token", 0)

    # Detecção de Frases Repetidas (Item 12/13)
    from collections import defaultdict
    import re
    from core.normalizacao import normalizar as normalizar_texto_aux

    contagem_sentencas = defaultdict(list)
    for t_idx_dup, td_dup in enumerate(st.session_state.get("turmas_processadas", [])):
        for a_idx_dup, aula_dup in enumerate(td_dup.get("aulas", [])):
            metodologia_dup = aula_dup.get("metodologia") or []
            textos_etapas = []
            for item in metodologia_dup:
                if isinstance(item, dict):
                    textos_etapas.append(item.get("texto", ""))
                else:
                    textos_etapas.append(str(item))
            texto_completo = " ".join(textos_etapas)
            # Divide por sentenças usando pontuação simples
            sentencas = re.split(r'[.!?\n]', texto_completo)
            vistas_nesta_aula = set()
            for s in sentencas:
                s_limpa = re.sub(r'\s+', ' ', s).strip()
                palavras = s_limpa.split()
                if len(palavras) > 8:
                    s_norm = normalizar_texto_aux(s_limpa)
                    if s_norm not in vistas_nesta_aula:
                        vistas_nesta_aula.add(s_norm)
                        contagem_sentencas[s_norm].append((t_idx_dup, a_idx_dup, s_limpa))

    duplicadas_por_aula = defaultdict(list)
    for frase_norm, ocorrencias in contagem_sentencas.items():
        if len(ocorrencias) > 2:
            # Esta frase está repetida em mais de 2 aulas
            for t_i, a_i, original_txt in ocorrencias:
                duplicadas_por_aula[(t_i, a_i)].append(original_txt)

    try:
        caminhos_relatorio = _salvar_relatorios_conferencia(
            turmas_processadas=st.session_state.get("turmas_processadas", []),
            duplicadas_por_aula=duplicadas_por_aula,
            professor=professor,
            disciplina=disciplina,
            turma=turma,
            mes=mes,
            bimestre=bimestre,
            modo_ia=modo_ia,
            modo_upload_pdf=modo_upload_pdf,
            pasta_pdfs_auto=pasta_pdfs_auto,
            pdfs_selecionados=pdfs_selecionados_tela,
        )
        if caminhos_relatorio:
            pasta_relatorio = Path(caminhos_relatorio[0]).parent
            st.info(
                "Relatórios de conferência salvos em: "
                f"{pasta_relatorio}. Esta pasta é só para análise e pode ser apagada depois sem afetar o sistema."
            )
    except Exception as err:
        st.warning(f"Não consegui salvar os relatórios de conferência automaticamente: {err}")

    for t_idx, td in enumerate(st.session_state["turmas_processadas"]):
        total_aulas_turma = len(td.get("aulas", []))
        st.markdown(f'<div class="review-class-title">{td["turma"]}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="review-class-meta">{total_aulas_turma} aula(s) prontas para conferência nesta turma.</div>', unsafe_allow_html=True)
        aulas_edit = []
        for a_idx, aula in enumerate(td["aulas"]):
            score = aula.get("confidence_score", 100)
            if score >= 80:
                status_emoji = "🟢"
            elif score >= 60:
                status_emoji = "🟡"
            else:
                status_emoji = "🔴"
            with st.expander(f"{status_emoji} Aula {a_idx+1} - {aula.get('tema','')}", expanded=False):
                # Alertas de Qualidade e Redundância (Item 13)
                if score < 60:
                    st.error(f"🔴 **Qualidade Crítica ({score}%)**: Este plano possui baixíssima aderência ao PDF ou problemas pedagógicos graves.")
                elif score < 80:
                    st.warning(f"🟡 **Qualidade Aceitável ({score}%)**: O plano possui ressalvas ou desvios menores em relação ao PDF.")
                else:
                    st.success(f"🟢 **Alta Qualidade ({score}%)**: Plano totalmente aderente e validado.")
                
                avisos_val = aula.get("avisos_validacao") or []
                if avisos_val:
                    st.warning("**Alertas de Qualidade Pedagógica:**\n" + "\n".join([f"- {aviso}" for aviso in avisos_val]))
                
                frases_dupl = duplicadas_por_aula.get((t_idx, a_idx))
                if frases_dupl:
                    st.warning("**Aviso de Redundância (frases repetidas em mais de 2 aulas do lote):**\n" + "\n".join([f"- \"{frase}\"" for frase in frases_dupl]))

                col1, col2 = st.columns(2)
                with col1:
                    t = st.text_input("Tema", value=aula.get("tema",""), key=f"tema_{rev_tok}_{t_idx}_{a_idx}")
                    a = st.text_area("Aprendizagem", value=aula.get("aprendizagem",""), key=f"apr_{rev_tok}_{t_idx}_{a_idx}")
                with col2:
                    acomp = st.text_area("Acompanhamento", value="\n".join(aula.get("acompanhamento",[])), key=f"acomp_{rev_tok}_{t_idx}_{a_idx}")
                    aces = st.text_area("Acessibilidade", value="\n".join(aula.get("acessibilidade",[])), key=f"acess_{rev_tok}_{t_idx}_{a_idx}")
                m = st.text_area("Metodologia", value=_texto_metodologia_app(aula), height=150, key=f"met_{rev_tok}_{t_idx}_{a_idx}")
                
                # Relatório Técnico (Item 12/14)
                if st.checkbox("🛠️ Exibir Relatório Técnico da Geração", value=False, key=f"tech_rep_{rev_tok}_{t_idx}_{a_idx}"):
                    st.markdown(
                        f"""
                        | Parâmetro | Valor |
                        |---|---|
                        | **Provedor da IA** | {aula.get("ia_provedor") or "Sem IA"} |
                        | **Cache Reutilizado** | {"Sim" if aula.get("cache_reutilizado") else "Não"} |
                        | **Versão do Gerador** | {aula.get("versao_gerador") or "1.2.9"} |
                        | **Origem da Metodologia** | {aula.get("origem_metodologia") or "Desconhecida"} |
                        | **Score de Confiança** | {aula.get('confidence_score', 100)}% |
                        """
                    )
                    
                    diag = aula.get("diagnostico_geracao") or {}
                    if diag:
                        st.markdown("#### Transformação da Metodologia (Pipeline)")
                        tabs = st.tabs(["1. Rascunho Local Heurístico", "2. Resposta IA Crua", "3. Higienização/Polimento", "4. Metodologia Final"])
                        with tabs[0]:
                            met_local = diag.get("metodologia_local") or []
                            if met_local:
                                st.write(_texto_metodologia_app({"metodologia": met_local}))
                            else:
                                st.info("Nenhuma etapa heurística local gerada.")
                        with tabs[1]:
                            met_ia = diag.get("metodologia_ia_crua") or []
                            if isinstance(met_ia, str):
                                st.text(met_ia)
                            elif met_ia:
                                st.write(_texto_metodologia_app({"metodologia": met_ia}))
                            else:
                                st.info("Sem resposta direta de IA (gerado localmente ou cached).")
                        with tabs[2]:
                            met_hig = diag.get("metodologia_higienizada") or []
                            if met_hig:
                                st.write(_texto_metodologia_app({"metodologia": met_hig}))
                            else:
                                st.info("Nenhum estágio higienizado intermediário.")
                        with tabs[3]:
                            met_fin = diag.get("metodologia_final") or []
                            if met_fin:
                                st.write(_texto_metodologia_app({"metodologia": met_fin}))
                            else:
                                st.info("Nenhuma metodologia final.")
                
                ae = aula.copy()
                ae.update({"tema": t, "aprendizagem": a, "acompanhamento": [x for x in acomp.split("\n") if x], "acessibilidade": [x for x in aces.split("\n") if x], "metodologia": _metodologia_app_para_blocos(m)})
                aulas_edit.append(ae)
        turmas_revisadas.append({"turma": td["turma"], "aulas": aulas_edit})

    # Botão para salvar alterações de volta nos arquivos de referência DOCX
    referencias_para_atualizar = {}
    for tr in turmas_revisadas:
        for aula in tr["aulas"]:
            ref_path = aula.get("fonte_referencia_metodologia")
            if ref_path and os.path.exists(ref_path):
                referencias_para_atualizar.setdefault(ref_path, []).append(aula)
                
    if referencias_para_atualizar:
        st.markdown('<div class="section-card"></div><div class="section-title">💾 Atualizar Arquivos de Referência</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-subtitle">Grave os ajustes e correções feitos nesta tela diretamente no arquivo DOCX de referência original.</div>', unsafe_allow_html=True)
        
        for ref_path, aulas_ref in referencias_para_atualizar.items():
            nome_ref_simpl = os.path.basename(ref_path)
            # Remove duplicatas de aulas_ref
            aulas_ref_unicas = {}
            for a in aulas_ref:
                num = a.get("numero_aula") or a.get("numero") or 0
                aulas_ref_unicas[num] = a
                
            btn_key = f"save_ref_{hashlib.md5(ref_path.encode('utf-8', errors='ignore')).hexdigest()[:8]}"
            confirmar_ref_key = f"confirm_ref_{hashlib.md5(ref_path.encode('utf-8', errors='ignore')).hexdigest()[:8]}"
            confirmar_ref = st.checkbox(
                f"Confirmo que desejo sobrescrever o DOCX de referência '{nome_ref_simpl}'.",
                key=confirmar_ref_key,
            )
            if st.button(
                f"Atualizar '{nome_ref_simpl}' com os ajustes desta tela",
                key=btn_key,
                type="secondary",
                disabled=not confirmar_ref,
            ):
                try:
                    from docx import Document
                    doc = Document()
                    aulas_ordenadas = sorted(aulas_ref_unicas.values(), key=lambda x: int(x.get("numero_aula") or x.get("numero") or 0))
                    for aula in aulas_ordenadas:
                        num = aula.get("numero_aula") or aula.get("numero") or 0
                        tit = aula.get("tema") or ""
                        # Aula Heading
                        doc.add_paragraph(f"AULA {num} - {tit}")
                        doc.add_paragraph()
                        # Metodologia
                        doc.add_paragraph("METODOLOGIA")
                        for etapa in (aula.get("metodologia") or []):
                            if isinstance(etapa, dict):
                                doc.add_paragraph(f"{etapa.get('titulo', '')}: {etapa.get('texto', '')}")
                            else:
                                doc.add_paragraph(str(etapa))
                        doc.add_paragraph()
                        # Acompanhamento
                        doc.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
                        for item in (aula.get("acompanhamento") or []):
                            item_limpo = str(item).replace("☑", "").strip()
                            if item_limpo:
                                doc.add_paragraph(f"☑ {item_limpo}")
                        doc.add_paragraph()
                        # Acessibilidade
                        doc.add_paragraph("ACESSIBILIDADE")
                        for item in (aula.get("acessibilidade") or []):
                            item_limpo = str(item).replace("☑", "").strip()
                            if item_limpo:
                                doc.add_paragraph(f"☑ {item_limpo}")
                        doc.add_paragraph()
                    
                    doc.save(ref_path)
                    st.success(f"✓ O arquivo '{nome_ref_simpl}' foi atualizado e agora contém as versões corrigidas dos planos!")
                except Exception as err:
                    st.error(f"Erro ao salvar arquivo de referência: {err}")
        
    st.markdown(
        """
        <div class="download-panel">
            <div class="panel-title">Última conferência antes do arquivo final</div>
            <div class="panel-text">Se estiver tudo certo na revisão, gere o documento final para liberar os botões de download logo abaixo.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.checkbox(
        "Salvar este plano no histórico",
        key="salvar_historico_geracao",
        value=bool(st.session_state.get("salvar_historico_geracao", False)),
        help="Marque apenas quando o plano estiver realmente ok. Isso guarda o DOCX para consulta futura, sem alterar a aula inicial dos próximos planos.",
    )
    if st.button("GERAR DOCX", type="primary"):
        planos_gerados = []
        for tr in turmas_revisadas:
            planos_gerados.append(_gerar_docx_final(modelo_bytes, tr["aulas"], escola, professor, disciplina, componente_curricular, tr["turma"], mes, bimestre, semana, observacao, aulas_previstas_manual))
        st.session_state["planos_gerados"] = planos_gerados
        
        # Salva fisicamente na pasta local de finalizados
        _salvar_planos_na_pasta_finalizados(planos_gerados, disciplina)
        
        salvou_historico = _salvar_planos_gerados_se_configurado(
            planos_gerados,
            professor,
            disciplina,
            bimestre,
        )
        _registrar_mensagem_memoria_plano(salvou_historico)
        st.success("Planos gerados!")

    # Checa se houve alterações na tela após a geração do DOCX
    alteracoes_detectadas = False
    if st.session_state.get("planos_gerados"):
        for tr in turmas_revisadas:
            plano_gerado = next((p for p in st.session_state["planos_gerados"] if p["turma"] == tr["turma"]), None)
            if plano_gerado:
                if len(plano_gerado.get("aulas", [])) != len(tr["aulas"]):
                    alteracoes_detectadas = True
                    break
                for a1, a2 in zip(plano_gerado.get("aulas", []), tr["aulas"]):
                    if (
                        a1.get("tema") != a2.get("tema")
                        or a1.get("aprendizagem") != a2.get("aprendizagem")
                        or a1.get("acompanhamento") != a2.get("acompanhamento")
                        or a1.get("acessibilidade") != a2.get("acessibilidade")
                        or a1.get("metodologia") != a2.get("metodologia")
                    ):
                        alteracoes_detectadas = True
                        break
            else:
                alteracoes_detectadas = True
                break

if st.session_state.get("planos_gerados"):
    if alteracoes_detectadas:
        st.warning("⚠️ **Alterações detectadas nos campos da tela!** Os arquivos de download abaixo ainda contêm a versão anterior. Clique no botão abaixo para atualizar os arquivos finais com as suas correções.")
        if st.button("🔄 ATUALIZAR ARQUIVOS DOCX COM AS CORREÇÕES DA TELA", type="primary"):
            planos_gerados = []
            for tr in turmas_revisadas:
                planos_gerados.append(_gerar_docx_final(modelo_bytes, tr["aulas"], escola, professor, disciplina, componente_curricular, tr["turma"], mes, bimestre, semana, observacao, aulas_previstas_manual))
            st.session_state["planos_gerados"] = planos_gerados
            
            # Salva localmente e no histórico
            _salvar_planos_na_pasta_finalizados(planos_gerados, disciplina)
            if st.session_state.get("salvar_historico_geracao", False):
                _salvar_planos_gerados_se_configurado(planos_gerados, professor, disciplina, bimestre)
            st.success("✓ Arquivos finais atualizados e salvos com as novas correções da tela!")
            st.rerun()

    planos_gerados = st.session_state["planos_gerados"]
    st.info(f"📂 Os arquivos `.docx` estão salvos e atualizados na pasta: `{PLANOS_FINALIZADOS_DIR}`")
    
    mensagem_historico = str(st.session_state.pop("mensagem_historico_planos", "") or "").strip()
    mensagem_historico_tipo = str(st.session_state.pop("mensagem_historico_planos_tipo", "") or "").strip()
    if mensagem_historico:
        if mensagem_historico_tipo == "success":
            st.success(mensagem_historico)
        else:
            st.info(mensagem_historico)
    st.markdown('<div class="section-title">📥 Passo 3: Download</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-subtitle">Baixe o arquivo final já pronto para envio ou salve um pacote com todas as turmas processadas.</div>', unsafe_allow_html=True)
    st.markdown(
        f"""
        <div class="download-panel">
            <div class="panel-title">Arquivos finais disponíveis</div>
            <div class="panel-text">O sistema preparou <strong>{len(planos_gerados)}</strong> arquivo(s) final(is). Você pode baixar um documento único ou um pacote completo, conforme a quantidade de turmas processadas.</div>
            <div class="panel-pills">
                <span class="panel-pill">DOCX pronto</span>
                <span class="panel-pill">Compatível com modelo</span>
                <span class="panel-pill">Entrega organizada</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if len(planos_gerados) == 1:
        st.download_button("Baixar DOCX", data=planos_gerados[0]["docx_bytes"].getvalue(), file_name=nome_arquivo_plano(planos_gerados[0]["turma"], disciplina), mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.download_button("📦 Baixar ZIP", data=_montar_zip_planos(planos_gerados, disciplina), file_name="planos.zip", mime="application/zip")
        for p in planos_gerados:
            st.download_button(f"Baixar DOCX - {p['turma']}", data=p["docx_bytes"].getvalue(), file_name=nome_arquivo_plano(p["turma"], disciplina), key=f"dl_{p['turma']}")
