from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BASE_DIR = Path(r"D:\PDF novos\MATEMATICA\AF\3_BIMESTRE\6_ANO")
DOCX_PATH = BASE_DIR / "Metodologias_Matematica_6_Ano_Ensino_Fundamental_NOVO.docx"
BACKUP_DIR = Path(r"D:\PLANOS_LUAN\tmp\docx_backups")

COR_TITULO = RGBColor(0x1F, 0x49, 0x7D)
COR_AULA = RGBColor(0x00, 0x47, 0x70)
COR_ETAPA = RGBColor(0x2E, 0x74, 0xB5)


@dataclass
class AulaSecao:
    numero: int
    inicio: int
    fim: int
    metodologia_idx: int
    acompanhamento_idx: int


NOVAS_METODOLOGIAS: dict[int, list[tuple[str, str]]] = {
    4: [
        (
            "Na prática",
            "Atividade 1: Os alunos analisam as alturas de jogadores de vôlei, escrevem os valores em números decimais, representam essas medidas na reta numérica e organizam os dados da menor para a maior altura. O professor acompanha a leitura das informações e retoma com a turma a relação entre medida e posição na reta.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma representa números decimais em uma mesma reta numérica, registra as posições corretas e compara os valores com base na distância entre as marcações. A correção é feita coletivamente para consolidar a leitura e a ordenação dos números.",
        ),
        (
            "Na prática",
            "Atividade 3: Os estudantes observam comprimentos indicados em fita métrica, localizam cada medida na reta e identificam qual objeto apresenta a menor e a maior medida. O professor orienta a leitura das escalas e reforça o uso do decimal em situações de medida.",
        ),
        (
            "Na prática",
            "Atividade 4: Os alunos leem os pontos indicados em uma reta numérica, escrevem os valores decimais correspondentes e justificam qual ponto representa a maior e a menor medida. Ao final, o professor retoma os critérios usados pela turma na interpretação da reta.",
        ),
        (
            "Encerramento",
            'Encerrar a aula retomando como a reta numérica ajuda a localizar, comparar e ordenar medidas em números decimais. Com "COM SUAS PALAVRAS", a turma registra o que precisa observar para identificar corretamente maior, menor e posição intermediária em problemas com medidas.',
        ),
    ],
    6: [
        (
            "Para começar",
            'Iniciar a aula com a situação dos preços do mesmo produto em diferentes estabelecimentos, incentivando a turma a antecipar qual valor é menor e como essa comparação pode ser feita. Com "VIREM E CONVERSEM", os estudantes levantam hipóteses e justificam oralmente as primeiras escolhas.',
        ),
        (
            "Foco no conteúdo",
            'O professor explica estratégias de comparação e ordenação de números decimais, retomando o valor posicional dos algarismos e mostrando por que a leitura da parte inteira e da parte decimal precisa ser feita com atenção. Os exemplos do material são usados para explicitar critérios de comparação.',
        ),
        (
            "Foco no conteúdo",
            'Na sequência, o professor sistematiza o uso de estimativas e arredondamentos para verificar a razoabilidade das respostas, relacionando esse procedimento a situações reais de compra, medida e organização de dados.',
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos analisam três estratégias diferentes para comparar 0,17 e 0,3, identificando semelhanças, diferenças e o erro presente em uma das resoluções. O professor conduz a discussão para destacar a importância do valor posicional.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma compara quantidades de suco representadas por números decimais e por imagens de garrafinhas, registrando a ordem correta dos valores. O professor retoma com os estudantes como observar parte inteira e parte decimal em contextos visuais.",
        ),
        (
            "Na prática",
            "Atividade 3: Os estudantes analisam uma pesquisa de preços para um lanche coletivo, identificam menor e maior valor, fazem arredondamentos e organizam os preços em ordem crescente. As correções são socializadas para validar as estratégias utilizadas.",
        ),
        (
            "Encerramento",
            'Fechar a aula retomando os critérios usados para comparar e ordenar números decimais, destacando a relação entre valor posicional, estimativa e arredondamento. Com "COM SUAS PALAVRAS", a turma registra o que precisa observar antes de decidir qual número decimal é maior ou menor.',
        ),
    ],
    7: [
        (
            "Para começar",
            'Abrir a aula com a situação de compra na lanchonete, incentivando os estudantes a estimar o total gasto e a discutir se o valor fica mais próximo de R$ 20,00 ou de R$ 25,00. Com "VIREM E CONVERSEM", a turma compartilha estratégias de estimativa e compara diferentes raciocínios.',
        ),
        (
            "Foco no conteúdo",
            "O professor explica como realizar adições e subtrações com números decimais por meio da decomposição em parte inteira e parte decimal, retomando o valor posicional dos algarismos e a composição do resultado.",
        ),
        (
            "Foco no conteúdo",
            "Em seguida, o professor apresenta o alinhamento da vírgula e o uso de estimativas como apoio para conferir a razoabilidade dos cálculos, mostrando como essas estratégias ajudam antes e depois do cálculo exato.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos resolvem a compra feita por João na padaria, primeiro por estimativa e depois pelo cálculo exato, comparando os dois resultados. O professor acompanha os registros e retoma com a turma o sentido da aproximação e do resultado final.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma calcula o total gasto em uma compra coletiva no quiosque da escola, estimando o valor da pizza e dos refrigerantes antes de efetuar a soma exata. As soluções são discutidas coletivamente para consolidar estratégias de adição e comparação entre estimativa e valor real.",
        ),
        (
            "Encerramento",
            'Encerrar a aula retomando as estratégias de adição, subtração e estimativa com números decimais, verificando o que a turma compreendeu sobre cálculo aproximado e cálculo exato. Com "COM SUAS PALAVRAS", os estudantes registram uma forma simples de resolver mentalmente uma soma decimal.',
        ),
    ],
    9: [
        (
            "Na prática",
            "Atividade 1: Os estudantes leem e escrevem preços em números decimais a partir de produtos de papelaria, identificando valores, comparando itens e registrando o produto mais barato. O professor acompanha a leitura correta das partes inteira e decimal.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma analisa a conta de uma pizzaria, faz uma estimativa do total e, em seguida, calcula o troco exato a partir de R$ 50,00. O professor retoma com os estudantes a utilidade da aproximação antes do cálculo final.",
        ),
        (
            "Na prática",
            "Atividade 3: Os alunos conferem valores de uma compra no mercado, comparam estimativa e total exato e explicam o raciocínio usado para chegar ao resultado. As respostas são socializadas para validar as estratégias adotadas.",
        ),
        (
            "Na prática",
            "Atividade 4: A turma resolve uma situação envolvendo comparação de valores em reais, estimativa e decisão de compra, registrando o cálculo e justificando a resposta. O professor conduz a correção destacando leitura de preços, adição e subtração com números decimais.",
        ),
        (
            "Encerramento",
            'Fechar a aula retomando as estratégias de leitura, estimativa e cálculo usadas nas questões de verificação. Com "COM SUAS PALAVRAS", os estudantes registram o que mais ajudou a conferir preços, somas e trocos nas situações propostas.',
        ),
    ],
    10: [
        (
            "Abertura",
            'O professor apresenta à turma os conteúdos que serão retomados na plataforma, com foco em adição e subtração de números decimais, estimativas e conferência dos resultados. Antes do acesso, relembra com os estudantes como essas operações aparecem em situações de compra, troco e comparação de valores.',
        ),
        (
            "Prática na Matific",
            'Relembrar com os alunos o passo a passo de acesso à plataforma e orientar a realização das atividades de revisão sobre adição e subtração de números decimais. O professor acompanha os percursos, incentiva novas tentativas quando necessário e pede que a turma consulte anotações e estratégias já registradas no caderno.',
        ),
        (
            "Fechamento",
            "Se necessário, consolidar a aprendizagem retomando uma ou mais atividades no caderno do aluno, com correção coletiva das estratégias mais recorrentes. O professor encerra destacando os erros mais comuns observados na plataforma e os procedimentos que ajudam a evitá-los.",
        ),
    ],
    11: [
        (
            "Para começar",
            'Retomar com a turma situações de adição e subtração com números decimais, incentivando estimativas rápidas e comparação entre diferentes modos de calcular. Com "VIREM E CONVERSEM", os estudantes socializam hipóteses antes do cálculo escrito.',
        ),
        (
            "Foco no conteúdo",
            "O professor sistematiza o alinhamento da vírgula, o valor posicional dos algarismos e as estratégias de cálculo mental e escrito usadas em somas e subtrações com números decimais. Os exemplos do material ajudam a mostrar quando completar casas decimais favorece a organização do cálculo.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos resolvem operações de adição e subtração com números decimais, observando alinhamento, composição das casas decimais e escolha da estratégia mais adequada em cada item. O professor acompanha a resolução e retoma coletivamente os procedimentos usados pela turma.",
        ),
        (
            "Encerramento",
            'Encerrar a aula retomando os procedimentos que ajudam a organizar somas e subtrações com números decimais, verificando se a turma identifica onde posicionar a vírgula e como conferir o resultado. Com "COM SUAS PALAVRAS", os estudantes registram a estratégia que consideram mais segura para esse tipo de cálculo.',
        ),
    ],
    12: [
        (
            "Para começar",
            "Iniciar a aula retomando com a turma o que acontece com um número decimal quando ele é multiplicado por 10, 100 ou 1 000, incentivando previsões antes do cálculo. O professor acolhe as hipóteses e organiza as respostas para preparar a formalização do conteúdo.",
        ),
        (
            "Foco no conteúdo",
            "O professor explica a multiplicação de número natural por número decimal e a multiplicação de números decimais por 10, 100 e 1 000, destacando a relação com o valor posicional dos algarismos e a posição da vírgula.",
        ),
        (
            "Foco no conteúdo",
            "Na sequência, são retomadas estratégias de estimativa e leitura do resultado para que a turma compreenda por que o produto cresce, diminui ou se reorganiza conforme a potência de 10 utilizada na operação.",
        ),
        (
            "Na prática",
            "Atividade 1: Os estudantes resolvem multiplicações mentalmente, observando o deslocamento da vírgula e a alteração do valor posicional. O professor acompanha as justificativas e retoma as regularidades percebidas pela turma.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma analisa a ilustração do material e calcula quanto cada personagem gastou, aplicando a multiplicação em um contexto de consumo. As respostas são conferidas coletivamente para validar a estratégia utilizada.",
        ),
        (
            "Na prática",
            "Atividade 3: Os alunos resolvem uma situação envolvendo libra e quilograma, aplicando multiplicação com números decimais em contexto de medida. O professor destaca a importância de interpretar a unidade e organizar corretamente o cálculo.",
        ),
        (
            "Encerramento",
            'Fechar a aula retomando o efeito da multiplicação sobre a posição da vírgula e sobre o valor do número, verificando se a turma consegue prever resultados antes de calcular. Com "COM SUAS PALAVRAS", os estudantes registram uma regra prática que os ajude a resolver novas situações.',
        ),
    ],
    13: [
        (
            "Para começar",
            "Retomar com a turma a multiplicação de números decimais, incentivando previsões sobre o tamanho do resultado e a posição da vírgula antes do cálculo formal. O professor organiza as hipóteses para conectar o raciocínio inicial ao trabalho da aula.",
        ),
        (
            "Foco no conteúdo",
            "O professor explica a multiplicação de um número decimal por outro número decimal, utilizando decomposição, estimativa, cálculo direto e leitura do valor posicional para mostrar como organizar o produto.",
        ),
        (
            "Foco no conteúdo",
            "Em seguida, retoma com a turma o sentido da vírgula no resultado final e como a estimativa pode ajudar a verificar se a resposta obtida é razoável no contexto apresentado.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos calculam o valor pago por uma quantidade de tecido comprada em promoção, relacionando preço por metro e quantidade. O professor acompanha o cálculo e retoma a posição da vírgula no produto final.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma determina a distância percorrida por um carro em 3,5 horas, aplicando a multiplicação de números decimais em contexto de velocidade e tempo. As resoluções são comparadas para consolidar a estratégia escolhida.",
        ),
        (
            "Na prática",
            "Atividade 3: Os estudantes resolvem uma situação de conversão de polegadas para centímetros, observando decomposição e equivalência de medidas. O professor conduz a correção destacando a organização dos cálculos e a leitura do resultado.",
        ),
        (
            "Encerramento",
            'Encerrar a aula retomando as estratégias usadas para multiplicar números decimais e verificar a coerência do resultado. Com "COM SUAS PALAVRAS", a turma registra o que deve observar para posicionar corretamente a vírgula no produto.',
        ),
    ],
    15: [
        (
            "Abertura",
            'Apresentar à turma os conteúdos que serão retomados na plataforma, com foco na relação entre forma fracionária e forma decimal. O professor relembra exemplos já estudados para que os estudantes iniciem a revisão reconhecendo equivalências entre as duas representações.',
        ),
        (
            "Prática na Matific",
            'Orientar o acesso à plataforma e a realização das atividades de revisão sobre representação fracionária e decimal. Durante a prática, o professor acompanha as tentativas, sugere retomada das anotações do caderno e incentiva os estudantes a persistirem até compreender cada desafio.',
        ),
        (
            "Fechamento",
            "Se necessário, consolidar a aprendizagem retomando no caderno uma ou mais atividades relacionadas à equivalência entre frações e decimais. O professor encerra com correção coletiva, destacando as relações que apareceram com mais frequência na plataforma.",
        ),
    ],
    16: [
        (
            "Para começar",
            "Iniciar a aula retomando situações em que um mesmo fator aparece multiplicado várias vezes, conduzindo a turma à ideia de potenciação. O professor acolhe hipóteses iniciais e relaciona esse raciocínio a números decimais e frações.",
        ),
        (
            "Foco no conteúdo",
            "O professor apresenta a potenciação como multiplicação repetida de fatores iguais, mostrando sua aplicação com números decimais e frações. Os exemplos do material são usados para explicitar base, expoente e resultado.",
        ),
        (
            "Foco no conteúdo",
            "Na sequência, são discutidas regularidades nas potências de números como 0,1, destacando a relação entre o expoente e a quantidade de casas decimais no resultado.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos investigam uma sequência de blocos em que cada letra representa o produto dos números imediatamente abaixo, registrando os cálculos e identificando a potência presente na situação. O professor acompanha as justificativas e retoma a ideia de multiplicação sucessiva.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma calcula potências de 0,1 e observa o comportamento das casas decimais, comparando resultados e formulando conclusões. As respostas são socializadas para consolidar as regularidades encontradas.",
        ),
        (
            "Encerramento",
            'Encerrar a aula retomando o sentido da potenciação e as regularidades percebidas nos resultados com números racionais. Com "COM SUAS PALAVRAS", os estudantes registram como o expoente interfere na escrita decimal das potências estudadas.',
        ),
    ],
    17: [
        (
            "Para começar",
            "Abrir a aula retomando estratégias já conhecidas para dividir números decimais, como representar a divisão por fração ou eliminar a vírgula com potências de 10. O professor organiza as hipóteses da turma antes da sistematização.",
        ),
        (
            "Foco no conteúdo",
            "O professor explica diferentes estratégias para realizar divisão de número decimal por outro número decimal, de número natural por número decimal e de número decimal por número natural, destacando quando vale a pena transformar a escrita ou simplificar a operação.",
        ),
        (
            "Foco no conteúdo",
            "Em seguida, são retomados exemplos do material para mostrar como multiplicar dividendo e divisor por uma potência de 10 ajuda a reorganizar a conta sem alterar o quociente.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos resolvem divisões utilizando uma das estratégias estudadas, escolhendo entre fração, simplificação ou eliminação da vírgula. O professor acompanha a escolha do procedimento e retoma os casos em que cada estratégia favorece a resolução.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma resolve o problema do engarrafamento de vinho em recipientes de 0,75 litro, determinando a quantidade de garrafas necessárias. A correção coletiva destaca interpretação do contexto e organização do cálculo.",
        ),
        (
            "Na prática",
            "Atividade 3: Os estudantes analisam dados de abastecimento em um posto de combustível para identificar, por meio da divisão, qual combustível foi utilizado. O professor conduz a verificação das respostas e reforça a leitura correta das informações numéricas.",
        ),
        (
            "Encerramento",
            'Fechar a aula retomando as estratégias de divisão com números decimais e verificando se a turma consegue justificar por que a multiplicação por potências de 10 preserva o quociente. Com "COM SUAS PALAVRAS", os estudantes registram o procedimento que consideram mais eficiente.',
        ),
    ],
    18: [
        (
            "Na prática",
            "Atividade 1: Os alunos resolvem divisões diretas com números naturais e decimais, retomando procedimentos já estudados e verificando o resultado obtido. O professor acompanha os registros para identificar quais estratégias a turma mobiliza com mais segurança.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma compara opções de cabo elétrico e determina qual marca oferece o menor custo, usando divisão para analisar preço e quantidade. As respostas são discutidas coletivamente para validar a estratégia escolhida.",
        ),
        (
            "Na prática",
            "Atividade 3: Os estudantes resolvem a situação da sorveteria, dividindo o valor pago igualmente entre os participantes. O professor retoma a interpretação do problema e a organização da divisão no contexto apresentado.",
        ),
        (
            "Na prática",
            "Atividade 4: A turma calcula o consumo médio de água por pessoa em uma residência, relacionando divisão com números decimais a uma situação do cotidiano. O professor acompanha as justificativas e retoma com a turma o significado do resultado encontrado.",
        ),
        (
            "Na prática",
            "Atividade 5: Os alunos analisam afirmações incorretas sobre quocientes e corrigem cada uma delas, justificando os ajustes realizados. Esse momento é utilizado para consolidar critérios de verificação de resultados em divisões com decimais.",
        ),
        (
            "Encerramento",
            'Encerrar a aula retomando como a divisão com números decimais aparece em situações de preço, medida e repartição. Com "COM SUAS PALAVRAS", os estudantes registram um cuidado importante para verificar se o quociente encontrado faz sentido no problema.',
        ),
    ],
    19: [
        (
            "Relembre",
            'Retomar com a turma duas situações iniciais envolvendo multiplicação e divisão com decimais: a quantidade total de suco formada por 5 copos de 0,2 litro e o rendimento de uma jarra de 1,2 litro em copos de 0,3 litro. Com "VIREM E CONVERSEM", os estudantes justificam oralmente seus raciocínios antes de iniciar a verificação.',
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos resolvem a questão de múltipla escolha sobre 1,25 ÷ 0,5, registrando a estratégia utilizada para encontrar o quociente correto. O professor retoma a transformação da divisão e a justificativa da alternativa escolhida.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma determina a capacidade de uma jarra a partir da quantidade de copos de 0,2 litro que ela consegue servir. A correção coletiva destaca a relação entre multiplicação e divisão em contextos de medida.",
        ),
        (
            "Na prática",
            "Atividade 3: Os estudantes calculam o valor total de uma máquina de lavar a partir do parcelamento indicado no anúncio, organizando a multiplicação com números decimais. O professor acompanha a montagem da conta e a leitura do resultado final.",
        ),
        (
            "Na prática",
            "Atividade 4: A turma calcula quanto Mariana pagará ao abastecer 5,5 litros de combustível, relacionando preço por litro e quantidade. O professor retoma a importância de organizar corretamente as casas decimais durante a multiplicação.",
        ),
        (
            "Na prática",
            "Atividade 5: Os alunos comparam a compra à vista e a prazo de um micro-ondas, calculando quanto foi pago a mais no parcelamento. As respostas são discutidas para consolidar estratégias de multiplicação e subtração com números decimais.",
        ),
        (
            "Encerramento",
            'Fechar a aula retomando o que acontece com um número decimal quando é multiplicado ou dividido por 10, relacionando esse raciocínio aos exercícios resolvidos na verificação. Com "COM SUAS PALAVRAS", a turma registra a regularidade observada.',
        ),
    ],
    20: [
        (
            "Abertura",
            "O professor apresenta à turma que a revisão na plataforma terá como foco a potenciação com números decimais e frações, retomando a ideia de multiplicação repetida e os padrões observados nas potências estudadas anteriormente.",
        ),
        (
            "Prática na Matific",
            "Relembrar com os estudantes o acesso à plataforma e orientar a realização das atividades de revisão sobre potenciação. Durante a prática, o professor acompanha os percursos, sugere novas tentativas quando necessário e incentiva a consulta ao caderno para relembrar os procedimentos.",
        ),
        (
            "Fechamento",
            "Se necessário, consolidar a aprendizagem retomando no caderno uma ou mais atividades de potenciação, com correção coletiva dos procedimentos usados pela turma. O professor encerra destacando os pontos que ainda precisam de reforço após a experiência na plataforma.",
        ),
    ],
    21: [
        (
            "Para começar",
            "Abrir a aula observando representações em quadros de 10 × 10 e discutindo como partes do todo podem ser lidas como fração, número decimal e porcentagem. O professor recolhe hipóteses iniciais para encaminhar a formalização do conteúdo.",
        ),
        (
            "Foco no conteúdo",
            "O professor explica que frações com denominador 100 representam porcentagens e retoma a equivalência entre frações, números decimais e porcentagens comuns, como 10%, 25%, 50%, 75% e 100%.",
        ),
        (
            "Foco no conteúdo",
            "Na sequência, mostra como essas porcentagens aparecem em contextos reais e como a leitura da parte do todo ajuda a interpretar embalagens, tabelas e descontos simples.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos completam relações entre quadros de 100 partes, frações e escritas decimais, observando quantos quadradinhos coloridos correspondem a cada representação. O professor acompanha os registros e retoma a equivalência entre as diferentes formas de escrita.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma compara descontos de 25% e 10% em um par de tênis, calculando o valor do desconto e o preço final em cada caso. As respostas são socializadas para consolidar a ideia de porcentagem como parte do todo em situações de compra.",
        ),
        (
            "Encerramento",
            'Encerrar a aula retomando a relação entre fração de denominador 100, número decimal e porcentagem, verificando se a turma consegue passar de uma representação para outra. Com "COM SUAS PALAVRAS", os estudantes registram um exemplo dessa equivalência.',
        ),
    ],
    22: [
        (
            "Para começar",
            "Iniciar a aula com a leitura da informação de que 50% dos brasileiros não sabem localizar o Brasil no mapa, conduzindo a turma a discutir como uma porcentagem pode ser escrita de outras maneiras. O professor acolhe as respostas e retoma o significado de metade, quarto e outras partes conhecidas.",
        ),
        (
            "Foco no conteúdo",
            "O professor explica como representar porcentagens por frações equivalentes e utiliza essa relação para calcular percentuais de valores numéricos. Os exemplos do material ajudam a mostrar estratégias de decomposição, como 10% + 5% ou 10% + 10%.",
        ),
        (
            "Foco no conteúdo",
            "Na sequência, sistematiza com a turma que calcular uma porcentagem pode envolver multiplicação de frações equivalentes por valores decimais, sempre relacionando a porcentagem à parte do todo.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos resolvem a questão sobre 50%, reconhecendo essa porcentagem como metade e relacionando a escrita percentual à fração correspondente. O professor retoma as justificativas para consolidar o conceito.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma analisa diferentes estratégias para calcular 15% de R$ 120,00 e 20% de R$ 200,00, comparando os procedimentos usados pelos alunos do material. Esse momento é usado para valorizar estratégias pessoais e conferir a coerência dos resultados.",
        ),
        (
            "Encerramento",
            'Fechar a aula retomando que a porcentagem pode ser pensada como fração e também como estratégia de cálculo. Com "COM SUAS PALAVRAS", a turma registra uma maneira prática de encontrar porcentagens simples em situações do cotidiano.',
        ),
    ],
    23: [
        (
            "Relembre",
            "Retomar com a turma a relação entre porcentagem, fração e número decimal, observando exemplos como 25% e 50% em situações simples do cotidiano. O professor recolhe hipóteses para preparar a resolução das atividades seguintes.",
        ),
        (
            "Foco no conteúdo",
            "O professor sistematiza a representação decimal de porcentagens, como 25% = 0,25 e 50% = 0,5, mostrando que calcular uma porcentagem de um valor equivale a multiplicar esse valor por um número decimal.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos interpretam a situação do bolo e identificam que 25% corresponde a um quarto do todo. O professor retoma a justificativa da turma para consolidar a equivalência entre porcentagem e fração.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma calcula 50% de R$ 50,00 em um problema de compra de brinquedo, relacionando a porcentagem à ideia de metade. As estratégias são comparadas para reforçar o cálculo mental e a leitura do contexto.",
        ),
        (
            "Na prática",
            "Atividade 3: Os estudantes resolvem situações em que 25% e 5% precisam ser calculados a partir de um valor dado, observando decomposição de porcentagens e equivalências com frações e decimais. O professor acompanha as justificativas e retoma os procedimentos mais eficientes.",
        ),
        (
            "Encerramento",
            'Encerrar a aula retomando que a porcentagem pode ser escrita como fração e como número decimal, e que essas representações ajudam a resolver problemas com mais segurança. Com "COM SUAS PALAVRAS", os estudantes registram qual estratégia consideram mais fácil para calcular porcentagens conhecidas.',
        ),
    ],
    24: [
        (
            "Relembre",
            "Retomar com a turma as relações entre porcentagem, fração e número decimal, preparando os estudantes para resolver problemas contextualizados. O professor destaca que o foco da aula será aplicar essas relações em situações reais.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos resolvem a situação da papelaria, escrevendo 25% como fração e como decimal e determinando quantos cadernos de capa dura foram vendidos. O professor acompanha a montagem do cálculo e a interpretação do contexto.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma calcula quantos alunos canhotos há em uma sala com 30% de estudantes nessa condição, registrando a estratégia utilizada. As respostas são comparadas para consolidar o raciocínio percentual.",
        ),
        (
            "Na prática",
            "Atividade 3: Os estudantes resolvem o problema do restaurante, calculando 12% de desconto sobre o valor total da conta e determinando o valor final a pagar. O professor conduz a correção destacando a relação entre porcentagem, decimal e multiplicação.",
        ),
        (
            "Encerramento",
            'Fechar a aula retomando que problemas de porcentagem podem ser resolvidos por diferentes estratégias, desde a escrita fracionária até a multiplicação decimal. Com "COM SUAS PALAVRAS", os estudantes registram o passo que mais os ajuda a organizar esse tipo de cálculo.',
        ),
    ],
    25: [
        (
            "Abertura",
            "O professor apresenta que a revisão na plataforma terá como foco porcentagens comuns e sua relação com frações equivalentes e números decimais, retomando exemplos presentes em contextos reais antes do acesso à atividade.",
        ),
        (
            "Prática na Matific",
            "Orientar a turma no acesso à plataforma e no desenvolvimento das atividades de revisão sobre porcentagem. O professor acompanha as tentativas, incentiva o uso das anotações do caderno e pede que os estudantes refaçam os desafios até compreenderem o raciocínio envolvido.",
        ),
        (
            "Fechamento",
            "Se necessário, consolidar a aprendizagem resolvendo no caderno uma ou mais situações ligadas a porcentagens comuns, com correção coletiva das estratégias mais frequentes observadas na plataforma. O professor encerra retomando os pontos que precisam de maior atenção.",
        ),
    ],
    26: [
        (
            "Para começar",
            'Iniciar a aula com a conversa sobre consumo e escolhas financeiras, analisando a situação do celular em duas lojas e o aumento de 12% em uma delas. Com "VIREM E CONVERSEM", a turma discute qual preço é menor antes e depois do acréscimo e como descobrir o aumento em reais.',
        ),
        (
            "Foco no conteúdo",
            'O professor explica o acréscimo percentual como uma adição proporcional ao valor original, retomando a escrita da porcentagem na forma decimal e o cálculo da parte acrescida.',
        ),
        (
            "Foco no conteúdo",
            'Na sequência, sistematiza com a turma o cálculo do valor final após o aumento, destacando que o acréscimo depende do valor inicial e, por isso, produtos diferentes não aumentam a mesma quantia em reais.',
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos calculam quanto Alfredo pagou de uma dívida ao quitar 20% do valor total, registrando o cálculo em decimal e em reais. O professor acompanha a transformação da porcentagem e a interpretação do resultado.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma determina o valor final de uma poltrona comprada a prazo com 8% de acréscimo, distinguindo valor do aumento e valor final do produto. As soluções são socializadas para consolidar o procedimento.",
        ),
        (
            "Na prática",
            "Atividade 3: Os estudantes calculam a comissão de 8% recebida por um vendedor em uma venda de R$ 750,00 e, depois, o valor total obtido por ele. O professor retoma com a turma a ideia de parte proporcional aplicada a situações de trabalho e rendimento.",
        ),
        (
            "Encerramento",
            'Encerrar a aula retomando que o acréscimo percentual depende do valor inicial e exige calcular primeiro a parte proporcional para depois obter o valor final. Com "COM SUAS PALAVRAS", os estudantes registram como diferenciar acréscimo em reais e preço final.',
        ),
    ],
    27: [
        (
            "Para começar",
            "Abrir a aula retomando situações em que o preço de um produto diminui em porcentagem, conduzindo a turma a perceber que o decréscimo é uma subtração proporcional do valor original. O professor organiza as hipóteses iniciais antes da sistematização.",
        ),
        (
            "Foco no conteúdo",
            "O professor explica o decréscimo percentual como uma retirada proporcional do valor inicial, retomando a escrita decimal da porcentagem e o cálculo do desconto em reais.",
        ),
        (
            "Foco no conteúdo",
            "Na sequência, mostra que aumento e desconto sucessivos não anulam automaticamente um ao outro, pois cada porcentagem é calculada sobre uma base que pode ter mudado.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos analisam a situação da máquina fotográfica que recebe desconto de 10% e depois aumento de 10%, calculando o preço final e justificando por que ele não volta ao valor inicial. O professor conduz a discussão para consolidar a ideia de base de cálculo.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma resolve uma situação ligada à perda percentual de cobertura vegetal da Mata Atlântica, aplicando o cálculo de decréscimo em um contexto ambiental. As respostas são comparadas para reforçar a leitura do problema e o uso do percentual como redução proporcional.",
        ),
        (
            "Encerramento",
            'Fechar a aula retomando o sentido do decréscimo percentual e a diferença entre retirar uma parte do valor original e recalcular sobre um novo valor. Com "COM SUAS PALAVRAS", a turma registra um exemplo de desconto ou redução percentual do cotidiano.',
        ),
    ],
    28: [
        (
            "Relembre",
            "Retomar com a turma a diferença entre acréscimo e decréscimo percentual, preparando os estudantes para resolver problemas em que esses dois movimentos aparecem em contextos de compra e reajuste. O professor relembra que a porcentagem deve ser calculada sempre em relação ao valor inicial de cada etapa.",
        ),
        (
            "Na prática",
            "Atividade 1: Os alunos calculam o acréscimo de 10% no preço de um caderno e determinam o novo valor do produto. O professor acompanha a transformação do percentual em decimal e a adição ao preço original.",
        ),
        (
            "Na prática",
            "Atividade 2: A turma calcula o desconto de 20% no valor de uma camiseta, registrando quanto foi abatido e quanto o cliente pagará ao final. As resoluções são comparadas para consolidar a ideia de decréscimo.",
        ),
        (
            "Na prática",
            "Atividade 3: Os estudantes analisam um serviço de entrega com reajuste de preço, identificando se houve aumento ou diminuição e calculando o novo valor. O professor retoma com a turma a leitura correta do enunciado e a escolha da operação adequada.",
        ),
        (
            "Na prática",
            "Atividade 4: A turma resolve o problema do produto que sofre aumento de 10% e depois desconto de 10%, verificando por que o valor final não retorna a R$ 100,00. Esse momento é usado para consolidar a ideia de que a porcentagem incide sobre bases diferentes.",
        ),
        (
            "Encerramento",
            'Encerrar a aula retomando a diferença entre resolver acréscimos e decréscimos simples, destacando o papel do valor inicial em cada cálculo. Com "COM SUAS PALAVRAS", os estudantes registram o que precisa ser observado para não confundir aumento com desconto.',
        ),
    ],
    30: [
        (
            "Abertura",
            "O professor apresenta que a revisão na plataforma retoma o cálculo de porcentagens por meio da multiplicação entre números decimais, reforçando a leitura da porcentagem como parte do todo e sua aplicação em situações do cotidiano.",
        ),
        (
            "Prática na Matific",
            "Relembrar com a turma o acesso à plataforma e orientar a realização das atividades de revisão sobre cálculo de porcentagem. Durante a prática, o professor acompanha os percursos, incentiva novas tentativas e pede que os estudantes expliquem oralmente como pensaram antes de finalizar cada desafio.",
        ),
        (
            "Fechamento",
            "Se necessário, consolidar a aprendizagem retomando no caderno uma ou mais situações de porcentagem com correção coletiva. O professor encerra destacando os procedimentos que ajudaram a calcular percentuais com mais segurança durante a prática na plataforma.",
        ),
    ],
}


def backup_arquivo(caminho: Path) -> Path:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    destino = BACKUP_DIR / f"{caminho.stem}_{stamp}{caminho.suffix}"
    shutil.copy2(caminho, destino)
    return destino


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p


def extrair_numero_aula(texto: str) -> int | None:
    match = re.match(r"AULA\s+0*(\d+)", texto.strip(), flags=re.I)
    if not match:
        return None
    return int(match.group(1))


def localizar_secoes(doc: Document) -> list[AulaSecao]:
    paragrafos = list(iter_paragraphs(doc))
    titulos = []
    for idx, p in enumerate(paragrafos):
        numero = extrair_numero_aula(p.text)
        if numero is not None:
            titulos.append((numero, idx))

    secoes: list[AulaSecao] = []
    for pos, (numero, inicio) in enumerate(titulos):
        fim = titulos[pos + 1][1] - 1 if pos + 1 < len(titulos) else len(paragrafos) - 1
        metodologia_idx = -1
        acompanhamento_idx = -1
        for idx in range(inicio, fim + 1):
            texto = paragrafos[idx].text.strip().lower()
            if texto == "metodologia":
                metodologia_idx = idx
            elif texto == "acompanhamento da aprendizagem":
                acompanhamento_idx = idx
                break
        if metodologia_idx == -1 or acompanhamento_idx == -1:
            continue
        secoes.append(
            AulaSecao(
                numero=numero,
                inicio=inicio,
                fim=fim,
                metodologia_idx=metodologia_idx,
                acompanhamento_idx=acompanhamento_idx,
            )
        )
    return secoes


def excluir_paragrafo(paragrafo) -> None:
    elemento = paragrafo._element
    parent = elemento.getparent()
    if parent is not None:
        parent.remove(elemento)


def novo_paragrafo_apos(paragrafo):
    novo = OxmlElement("w:p")
    paragrafo._p.addnext(novo)
    from docx.text.paragraph import Paragraph

    return Paragraph(novo, paragrafo._parent)


def aplicar_estilo_etapa(paragrafo, rotulo: str, texto: str, estilo_fonte: str | None) -> None:
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragrafo.paragraph_format.space_after = Pt(0)
    paragrafo.paragraph_format.space_before = Pt(0)
    run_rotulo = paragrafo.add_run(f"{rotulo}: ")
    if estilo_fonte:
        run_rotulo.font.name = estilo_fonte
        run_rotulo._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), estilo_fonte)
    run_rotulo.font.bold = True
    run_rotulo.font.color.rgb = COR_ETAPA
    run_texto = paragrafo.add_run(texto)
    if estilo_fonte:
        run_texto.font.name = estilo_fonte
        run_texto._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), estilo_fonte)


def aplicar_estilo_cabecalho(paragrafo, cor: RGBColor, tamanho: int, estilo_fonte: str | None) -> None:
    if not paragrafo.runs:
        paragrafo.add_run(paragrafo.text)
    for run in paragrafo.runs:
        if estilo_fonte:
            run.font.name = estilo_fonte
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), estilo_fonte)
        run.font.bold = True
        run.font.color.rgb = cor
        run.font.size = Pt(tamanho)


def atualizar_docx() -> tuple[Path, Path]:
    backup = backup_arquivo(DOCX_PATH)
    doc = Document(str(DOCX_PATH))
    paragrafos = list(doc.paragraphs)
    fonte_modelo = None
    if paragrafos and paragrafos[0].runs:
        fonte_modelo = paragrafos[0].runs[0].font.name

    secoes = localizar_secoes(doc)
    for secao in reversed(secoes):
        if secao.numero not in NOVAS_METODOLOGIAS:
            continue

        paragrafos = list(doc.paragraphs)
        metodologia_p = paragrafos[secao.metodologia_idx]
        acompanhamento_p = paragrafos[secao.acompanhamento_idx]

        for idx in range(secao.acompanhamento_idx - 1, secao.metodologia_idx, -1):
            excluir_paragrafo(paragrafos[idx])

        referencia = metodologia_p
        for rotulo, texto in NOVAS_METODOLOGIAS[secao.numero]:
            novo_p = novo_paragrafo_apos(referencia)
            aplicar_estilo_etapa(novo_p, rotulo, texto, fonte_modelo)
            referencia = novo_p

    # reaplica estilos dos cabecalhos principais com o mesmo padrão do arquivo
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("AULA "):
            aplicar_estilo_cabecalho(p, COR_AULA, 14, fonte_modelo)
        elif txt == "Metodologia" or txt == "Acompanhamento da aprendizagem" or txt == "Acessibilidade":
            aplicar_estilo_cabecalho(p, COR_TITULO, 12, fonte_modelo)

    doc.save(str(DOCX_PATH))
    return DOCX_PATH, backup


if __name__ == "__main__":
    destino, backup = atualizar_docx()
    print(f"Arquivo atualizado: {destino}")
    print(f"Backup criado em: {backup}")
