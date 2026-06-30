from __future__ import annotations

import re
import shutil
import sys
import unicodedata
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BASE = Path(r"D:\PDF novos\LINGUA_PORTUGUESA\EM\3_BIMESTRE")
BACKUP_DIR = Path(r"D:\PLANOS_LUAN\tmp\docx_backups")

COR_TITULO = RGBColor(0x1F, 0x49, 0x7D)
COR_SUBTITULO = RGBColor(0x2E, 0x74, 0xB5)
COR_AULA = RGBColor(0x00, 0x47, 0x70)
COR_ETAPA = RGBColor(0x2E, 0x74, 0xB5)

TECNICAS = [
    "VIREM E CONVERSEM",
    "COM SUAS PALAVRAS",
    "HORA DA LEITURA",
    "TODO MUNDO ESCREVE",
    "UM PASSO DE CADA VEZ",
]

SERIES = [
    ("1_ANO", "1º Ano", "Metodologias_Lingua_Portuguesa_1_Ano_Ensino_Medio_3_Bimestre.docx"),
    ("2_ANO", "2º Ano", "Metodologias_Lingua_Portuguesa_2_Ano_Ensino_Medio_3_Bimestre.docx"),
    ("3_ANO", "3º Ano", "Metodologias_Lingua_Portuguesa_3_Ano_Ensino_Medio_3_Bimestre.docx"),
]

PADRAO_ETAPAS = re.compile(
    r"(?im)^(Para começar|Para comecar|Relembre|Foco no conteúdo|Foco no conteudo|Na prática|Na pratica|Socialização|Socializacao|Encerramento|Pause e responda)\b"
)

CANON_ETAPA = {
    "para comecar": "Para começar",
    "relembre": "Relembre",
    "foco no conteudo": "Foco no conteúdo",
    "na pratica": "Na prática",
    "socializacao": "Socialização",
    "encerramento": "Encerramento",
    "pause e responda": "Pause e responda",
}

ABERTURAS = [
    "Abrir a aula",
    "Começar a aula",
    "Dar início ao encontro",
    "Iniciar o trabalho da aula",
    "Abrir o percurso da aula",
]

RETOMADAS = [
    "Retomar a aula",
    "Reativar os pontos da aula anterior",
    "Revisitar o tema já estudado",
    "Recuperar o percurso anterior",
]

FOCOS = [
    "Explorar",
    "Apresentar",
    "Sistematizar",
    "Discutir",
    "Aprofundar",
    "Retomar",
]

ENCERRAMENTOS = [
    "Fechar a aula",
    "Concluir a aula",
    "Encerrar o percurso da aula",
    "Finalizar o trabalho do encontro",
]

VERBOS_OBSERVAVEIS = ["Verificar", "Observar", "Acompanhar"]


@dataclass
class PdfAula:
    numero: int
    titulo: str
    caminho: Path
    texto: str
    blocos: list["Bloco"]


@dataclass
class Bloco:
    etapa: str
    texto: str


def normalizar(texto: str) -> str:
    texto = unicodedata.normalize("NFD", str(texto or "").lower())
    texto = "".join(ch for ch in texto if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", texto).strip()


def escolher(lista: list[str], chave: int) -> str:
    return lista[chave % len(lista)]


def limpar_linha(linha: str) -> str:
    linha = str(linha or "").replace("\xa0", " ")
    linha = re.sub(r"\s+", " ", linha).strip()
    return linha.strip(" -;")


def limpar_bloco_texto(texto: str) -> str:
    linhas = []
    for linha in texto.splitlines():
        limpa = limpar_linha(linha)
        if not limpa:
            continue
        baixa = normalizar(limpa)
        if baixa in {
            "veja no livro!",
            "link para video",
            "correcao",
            "pause e responda",
            "na pratica",
            "foco no conteudo",
            "para comecar",
            "relembre",
            "encerramento",
            "socializacao",
        }:
            continue
        if baixa.startswith("disponivel em:") or "acesso em:" in baixa:
            continue
        if re.fullmatch(r"\d+\s*minutos?", baixa):
            continue
        if re.fullmatch(r"ate\s*\d+\s*min", baixa):
            continue
        if "freepik" in baixa or "getty images" in baixa or "reproducao" in baixa:
            continue
        if re.fullmatch(r"\d+\.", baixa):
            continue
        linhas.append(limpa)
    texto_limpo = " ".join(linhas)
    texto_limpo = re.sub(r"https?://\S+", "", texto_limpo)
    texto_limpo = re.sub(r"\s+", " ", texto_limpo).strip()
    return texto_limpo


def primeiras_frases(texto: str, limite: int = 2) -> list[str]:
    texto = limpar_bloco_texto(texto)
    partes = re.split(r"(?<=[\.\?\!])\s+", texto)
    saida = []
    for parte in partes:
        parte = parte.strip(" .")
        if len(parte) < 12:
            continue
        if parte.lower() in {"atividade 1", "atividade 2", "atividade 3"}:
            continue
        saida.append(parte)
        if len(saida) >= limite:
            break
    return saida


def extrair_titulo_primeira_pagina(texto_pagina: str, fallback: str) -> str:
    linhas = [limpar_linha(l) for l in texto_pagina.splitlines() if limpar_linha(l)]
    inicio = None
    for i, linha in enumerate(linhas):
        if normalizar(linha) in {"lingua portuguesa", "lingua portuguesa", "língua portuguesa"}:
            inicio = i + 1
            break
    if inicio is None:
        return fallback

    partes = []
    for linha in linhas[inicio:]:
        baixa = normalizar(linha)
        if "3o bimestre" in baixa or re.fullmatch(r"aula \d+", baixa) or baixa == "medio":
            break
        partes.append(linha)
    if not partes:
        return fallback
    titulo = " ".join(partes)
    titulo = titulo.replace(" – ", " - ").replace("—", "-").replace("–", "-")
    titulo = re.sub(r"\s+", " ", titulo).strip(" -")
    return titulo or fallback


def extrair_texto_pdf(caminho: Path, max_paginas: int = 16) -> list[str]:
    paginas = []
    with pdfplumber.open(str(caminho)) as pdf:
        for page in pdf.pages[:max_paginas]:
            paginas.append(page.extract_text() or "")
    return paginas


def canonizar_etapa(texto: str) -> str | None:
    chave = normalizar(texto)
    return CANON_ETAPA.get(chave)


def extrair_blocos_por_paginas(paginas: list[str]) -> list[Bloco]:
    blocos: list[Bloco] = []
    for texto_pagina in paginas:
        matches = list(PADRAO_ETAPAS.finditer(texto_pagina))
        if not matches:
            continue
        for idx, match in enumerate(matches):
            etapa = canonizar_etapa(match.group(1))
            if not etapa:
                continue
            inicio = match.end()
            fim = matches[idx + 1].start() if idx + 1 < len(matches) else len(texto_pagina)
            trecho = texto_pagina[inicio:fim].strip()
            if not trecho:
                continue
            if blocos and blocos[-1].etapa == etapa:
                blocos[-1].texto = f"{blocos[-1].texto}\n{trecho}".strip()
            else:
                blocos.append(Bloco(etapa=etapa, texto=trecho))
    return blocos


def carregar_aula(pdf_path: Path) -> PdfAula:
    numero_match = re.search(r"AULA_(\d+)", pdf_path.name, re.I)
    numero = int(numero_match.group(1)) if numero_match else 0
    paginas = extrair_texto_pdf(pdf_path)
    titulo_fallback = pdf_path.stem
    titulo_fallback = re.sub(r"^AULA_\d+\s*-\s*", "", titulo_fallback, flags=re.I)
    titulo_fallback = titulo_fallback.replace("_", " ").strip()
    titulo = extrair_titulo_primeira_pagina(paginas[0] if paginas else "", titulo_fallback)
    texto = "\n".join(paginas)
    blocos = extrair_blocos_por_paginas(paginas)
    return PdfAula(numero=numero, titulo=titulo, caminho=pdf_path, texto=texto, blocos=blocos)


def detectar_tecnica(texto: str) -> str | None:
    texto_norm = normalizar(texto).upper()
    for tecnica in TECNICAS:
        if normalizar(tecnica).upper() in texto_norm:
            return tecnica
    return None


def resumo_tema(titulo: str) -> str:
    tema = re.sub(r"\s*-\s*Parte\s*\d+\s*$", "", titulo, flags=re.I)
    return tema.strip()


def tema_frase(titulo: str) -> str:
    tema = resumo_tema(titulo).lower()
    return re.sub(r"^(o|a|os|as)\s+", "", tema)


def detectar_recurso(texto: str, titulo: str) -> str:
    base = normalizar(titulo)
    mapeamento = [
        ("meme", ["meme"]),
        ("cartaz", ["cartaz", "campanha"]),
        ("infografico", ["infografico"]),
        ("tirinha", ["tirinha"]),
        ("hq", ["historia em quadrinhos", "hq", "quadrinhos"]),
        ("poema", ["poema", "soneto"]),
        ("cancao", ["cancao", "musica", "letra"]),
        ("artigo de opiniao", ["artigo de opiniao"]),
        ("resenha critica", ["resenha critica"]),
        ("manifesto", ["manifesto"]),
        ("debate", ["debate regrado", "debate"]),
        ("texto dissertativo-argumentativo", ["dissertativo", "enem", "redacao"]),
        ("variacao linguistica", ["variacao linguistica", "norma-padrao", "norma padrao"]),
        ("postagem", ["post", "redes sociais", "instagram"]),
        ("video", ["video", "youtube"]),
        ("miniconto", ["miniconto", "microconto"]),
        ("relato de viagem", ["relato de viagem"]),
        ("diario", ["diario pessoal"]),
        ("texto dramatico", ["texto dramatico", "dramatico"]),
        ("literatura", ["literatura", "camoes", "gil vicente", "cortazar", "borges", "machado", "clarice", "guimaraes", "joao cabral", "fernando pessoa"]),
    ]
    for nome, chaves in mapeamento:
        if any(chave in base for chave in chaves):
            return nome
    return "texto"


def contexto_por_titulo(titulo: str) -> str:
    base = normalizar(titulo)
    if "literaturas africanas" in base:
        return "literaturas_africanas"
    if "anuncios publicitarios" in base:
        return "anuncios_digitais"
    if "imparcialidade em textos" in base or "um fato, duas versoes" in base:
        return "imparcialidade"
    if "carta de caminha" in base:
        return "carta_caminha"
    if "relato de viagem" in base:
        return "relato_viagem"
    if "diario pessoal" in base:
        return "diario_pessoal"
    if "texto dramatico" in base:
        return "texto_dramatico"
    if "moldando imagens" in base:
        return "moldando_imagens"
    if "resenha critica" in base:
        return "resenha_critica"
    if "desafios do mundo real" in base:
        return "desafios_mundo_real"
    if "literatura latino" in base:
        return "literatura_latino"
    if "realismo magico" in base:
        return "realismo_magico"
    if "eca de queiros" in base:
        return "realismo_portugal"
    if "machado de assis" in base and "cronica" not in base:
        return "realismo_brasil"
    if "artigo de opiniao" in base:
        return "artigo_opiniao"
    if "romantismo" in base or "cancao do exilio" in base:
        return "identidade_brasileira"
    if "cronica" in base:
        return "cronica"
    if "aluisio azevedo" in base or "o cortico" in base or "quarto de despejo" in base:
        return "realismo_naturalismo"
    if "intervencao urbana" in base:
        return "intervencao_urbana"
    if "parnasianismo" in base or "francisca julia" in base:
        return "parnasianismo"
    if "percurso" in base or "meu caminho" in base or "jornada" in base:
        return "percurso_autoral"
    if "variacao e norma" in base:
        return "variacao_norma"
    if "dissertativo" in base:
        return "dissertativo_argumentativo"
    if "clarice lispector" in base or "guimaraes rosa" in base or "joao cabral" in base:
        return "modernismo_3geracao"
    if "fernando pessoa" in base:
        return "fernando_pessoa"
    if "manifesto" in base:
        return "manifesto"
    if "debate regrado" in base:
        return "debate_regrado"
    if "miniconto" in base or "microconto" in base:
        return "miniconto_microconto"
    return "texto"


def descritores_foco(contexto: str) -> list[str]:
    mapa = {
        "literaturas_africanas": [
            "as literaturas africanas em língua portuguesa, destacando a diversidade de países, autores, contextos históricos e projetos de identidade que atravessam essa produção literária",
            "a relação entre literatura, memória, independência e afirmação cultural, retomando como a linguagem literária dialoga com diferentes experiências africanas",
        ],
        "anuncios_digitais": [
            "a construção de anúncios publicitários em mídias digitais, observando público-alvo, linguagem multissemiótica e estratégias de persuasão",
            "como imagens, palavras e escolhas de design produzem efeitos de sentido e orientam o consumo no ambiente digital",
        ],
        "imparcialidade": [
            "como diferentes textos noticiosos apresentam um mesmo fato, observando seleção lexical, ponto de vista e efeitos de credibilidade",
            "as marcas de parcialidade e imparcialidade na notícia, relacionando linguagem, fontes e interpretação crítica",
        ],
        "carta_caminha": [
            "a Carta de Caminha como relato de viagem e documento histórico, observando descrição, ponto de vista e projeto de colonização",
            "como o texto constrói uma visão sobre a terra e seus habitantes, articulando leitura histórica e análise da linguagem",
        ],
        "relato_viagem": [
            "o relato de viagem contemporâneo, destacando experiências do narrador, impressões pessoais e recursos descritivos",
            "como o percurso vivido se transforma em texto, articulando observação do espaço, marcas de subjetividade e escolhas linguísticas",
        ],
        "diario_pessoal": [
            "o diário pessoal como gênero de registro do cotidiano, observando subjetividade, temporalidade e relação com a experiência narrada",
            "como a escrita de si organiza memórias, emoções e escolhas de linguagem no cotidiano",
        ],
        "texto_dramatico": [
            "o texto dramático, observando falas, rubricas, construção de cena e efeitos de oralidade",
            "como a organização das falas e das ações contribui para a leitura e a interpretação do conflito dramático",
        ],
        "moldando_imagens": [
            "os efeitos das imagens e das redes sociais na construção de sentidos, observando filtros, curadoria e autoimagem",
            "como textos e imagens circulam nas redes, produzindo posicionamentos, comparação social e reflexão crítica",
        ],
        "resenha_critica": [
            "a resenha crítica como gênero de apreciação, observando resumo, avaliação e linguagem persuasiva",
            "como o texto organiza opinião, critérios de análise e convite ao consumo ou à rejeição da obra analisada",
        ],
        "desafios_mundo_real": [
            "como textos, propostas e situações-problema articulam leitura, argumentação e posicionamento diante de desafios do mundo real",
            "estratégias de análise e resposta a questões contemporâneas, relacionando repertório, linguagem e participação social",
        ],
        "literatura_latino": [
            "a literatura latino-americana, destacando autores, contextos culturais e temas que marcam essa produção",
            "como diferentes obras latino-americanas articulam identidade, memória e experimentação estética",
        ],
        "realismo_magico": [
            "o realismo mágico na literatura latino-americana, observando a convivência entre cotidiano, insólito e identidade cultural",
            "como autores e obras do movimento articulam fantástico, crítica social e invenção literária",
        ],
        "realismo_portugal": [
            "o Realismo em Portugal, observando crítica social, construção narrativa e projeto literário em Eça de Queirós",
            "como a leitura do texto evidencia ironia, observação da realidade e análise do comportamento social",
        ],
        "realismo_brasil": [
            "o Realismo no Brasil, observando narrador, crítica social e análise psicológica nas obras de Machado de Assis",
            "como a linguagem e a construção do texto revelam ironia, conflito e leitura crítica da sociedade",
        ],
        "artigo_opiniao": [
            "o artigo de opinião como gênero argumentativo, observando questão polêmica, fato, opinião e posicionamento",
            "como os modos verbais, os dados e as estratégias argumentativas sustentam a defesa de um ponto de vista",
        ],
        "identidade_brasileira": [
            "a construção da identidade brasileira na literatura, observando romantismo, nacionalismo e releituras críticas do tema",
            "como textos e canções articulam memória, pertencimento e representação do país",
        ],
        "cronica": [
            "a crônica como gênero de observação do cotidiano, articulando linguagem, ponto de vista e crítica do dia a dia",
            "como o texto cronístico transforma cenas comuns em reflexão, humor ou comentário social",
        ],
        "realismo_naturalismo": [
            "o Realismo e o Naturalismo, observando contexto histórico, crítica social e representação das relações humanas",
            "como a leitura das obras evidencia determinismo, desigualdade e formas de narrar a realidade",
        ],
        "intervencao_urbana": [
            "a intervenção urbana como linguagem de ocupação crítica do espaço, articulando texto, imagem e posicionamento social",
            "como arte e cidade se relacionam na produção de mensagens públicas e debates sobre convivência urbana",
        ],
        "parnasianismo": [
            "o Parnasianismo e seus projetos estéticos, observando forma, linguagem e construção do poema",
            "como a leitura poética articula rigor formal, escolhas lexicais e efeitos de sentido",
        ],
        "percurso_autoral": [
            "a síntese do percurso autoral da turma, articulando leitura, produção e reflexão sobre os aprendizados do bimestre",
            "como experiências, escolhas e registros do percurso podem ser organizados em uma apresentação autoral",
        ],
        "variacao_norma": [
            "a relação entre variação linguística, norma-padrão e adequação de linguagem em diferentes contextos de uso",
            "como escolhas linguísticas revelam contexto, intenção comunicativa e posicionamento diante da língua",
        ],
        "dissertativo_argumentativo": [
            "o texto dissertativo-argumentativo, observando tese, argumentos, repertório e organização do projeto de texto",
            "como diferentes estratégias argumentativas contribuem para sustentar um posicionamento em contextos como o Enem",
        ],
        "modernismo_3geracao": [
            "a terceira geração modernista, observando temas, linguagem e projetos estéticos dos autores trabalhados",
            "como a obra lida articula subjetividade, invenção formal e leitura crítica do mundo",
        ],
        "fernando_pessoa": [
            "a obra de Fernando Pessoa e a questão dos heterônimos, observando vozes poéticas, identidade e construção literária",
            "como diferentes perspectivas de eu lírico ampliam a leitura da poesia e do projeto estético do autor",
        ],
        "manifesto": [
            "o manifesto como gênero de posicionamento público, observando crítica social, intenção comunicativa e mobilização de leitores",
            "como a linguagem do manifesto constrói denúncia, proposta e visão de mundo em contextos coletivos",
        ],
        "debate_regrado": [
            "o debate regrado como gênero argumentativo da oralidade, observando regras, papéis e formas de participação",
            "como escuta, respeito ao turno de fala e sustentação de argumentos organizam o debate público",
        ],
        "miniconto_microconto": [
            "o miniconto e o microconto, observando concisão, sugestão e construção de sentidos em textos breves",
            "como poucos recursos linguísticos podem produzir surpresa, ambiguidade e densidade interpretativa",
        ],
        "texto": [
            "o texto e o gênero trabalhados na aula, observando leitura, interpretação e construção de sentidos",
            "como a linguagem do material organiza informações, posicionamentos e efeitos de leitura",
        ],
    }
    return mapa.get(contexto, mapa["texto"])


def localizar_linha_interrogativa(texto: str) -> str:
    for linha in texto.splitlines():
        limpa = limpar_linha(linha)
        if "?" in limpa and len(limpa) > 15:
            return limpa
    return ""


def localizar_comando_atividade(texto: str) -> str:
    for linha in texto.splitlines():
        limpa = limpar_linha(linha)
        if re.match(r"^(Leia|Analise|Relacione|Observe|Compare|Releia|Produza|Responda|Assista|Debata|Escreva)\b", limpa, re.I):
            return limpa
    return ""


def compor_para_comecar(aula: PdfAula, bloco: Bloco) -> str:
    contexto = contexto_por_titulo(aula.titulo)
    verbo = escolher(RETOMADAS if bloco.etapa == "Relembre" else ABERTURAS, aula.numero)
    tecnica = detectar_tecnica(bloco.texto)
    tema = resumo_tema(aula.titulo).lower()
    corpo = f"{verbo} com \"{tecnica}\"" if tecnica else verbo
    if contexto in {"debate_regrado", "manifesto", "artigo_opiniao", "variacao_norma", "dissertativo_argumentativo"}:
        return f"{corpo}, retomando situações comunicativas ligadas a {tema} e preparando a discussão inicial da aula."
    if contexto in {"moldando_imagens", "anuncios_digitais", "imparcialidade"}:
        return f"{corpo}, mobilizando impressões iniciais da turma sobre {tema} e preparando a leitura do material."
    return f"{corpo}, retomando referências iniciais da turma sobre {tema} e preparando a leitura do material."


def compor_foco(aula: PdfAula, bloco: Bloco, indice: int) -> str:
    contexto = contexto_por_titulo(aula.titulo)
    verbo = escolher(FOCOS, aula.numero + indice)
    ideias = descritores_foco(contexto)
    base = ideias[min(indice, len(ideias) - 1)]
    return f"{verbo} {base}."


def compor_na_pratica(aula: PdfAula, bloco: Bloco, indice: int) -> str:
    atividade_match = re.search(r"(Atividade\s*\d+|Desafio|Produção)", bloco.texto, re.I)
    rotulo = atividade_match.group(1).title() if atividade_match else f"Atividade {indice + 1}"
    comando = localizar_comando_atividade(bloco.texto)
    tecnica = detectar_tecnica(bloco.texto)
    recurso = detectar_recurso(bloco.texto, aula.titulo)
    partes = []
    if tecnica:
        partes.append(f'Com "{tecnica}",')
    else:
        partes.append("Orientar a turma a")

    comando_invalido = False
    if comando:
        comando_limpo = comando.rstrip(" .:")
        comando_norm = normalizar(comando_limpo)
        if (
            len(comando_limpo) < 18
            or comando.endswith(":")
            or "[...]" in comando
            or ("“" in comando and "”" not in comando)
            or re.search(r"\b(ao|do|da|de|em|no|na|nesse|nessa|nesse texto|ao texto|de um)$", comando_norm)
        ):
            comando_invalido = True
    if comando and not comando_invalido:
        comando = comando.rstrip(".")
        if partes[0].endswith(","):
            partes.append(f" {comando[0].lower() + comando[1:] if len(comando) > 1 else comando.lower()}")
        else:
            partes.append(f" {comando.lower()}")
    else:
        tema = tema_frase(aula.titulo)
        partes.append(f" desenvolver a atividade proposta sobre {tema}")

    texto = f"{rotulo}. {''.join(partes).strip()}."
    if "correcao" in normalizar(bloco.texto):
        texto += " Em seguida, conduzir a correção coletiva para comparar respostas, explicitar critérios e consolidar a aprendizagem."
    else:
        if recurso in {"debate", "artigo de opiniao", "resenha critica", "manifesto", "texto dissertativo-argumentativo"}:
            texto += " Ao final, retomar os argumentos levantados, organizar os registros da turma e conferir coletivamente as conclusões."
        elif recurso in {"meme", "cartaz", "infografico", "postagem", "video"}:
            texto += " Depois, socializar as observações da turma e relacionar as escolhas do material aos efeitos de sentido discutidos na aula."
        else:
            texto += " Depois, acompanhar os registros da turma e retomar coletivamente as respostas para consolidar os pontos centrais da aula."
    return texto


def compor_encerramento(aula: PdfAula, bloco: Bloco | None) -> str:
    verbo = escolher(ENCERRAMENTOS, aula.numero)
    tema = resumo_tema(aula.titulo).lower()
    return f"{verbo}, retomando as ideias centrais sobre {tema} e verificando o que a turma compreendeu ao longo da aula."


def carregar_titulos_existentes(caminho_docx: Path) -> dict[int, str]:
    if not caminho_docx.exists():
        return {}
    doc = Document(caminho_docx)
    titulos: dict[int, str] = {}
    for p in doc.paragraphs:
        texto = p.text.strip()
        match = re.match(r"AULA\s+(\d+)\s*-\s*(.+)$", texto, re.I)
        if match:
            titulos[int(match.group(1))] = match.group(2).strip()
    return titulos


def montar_metodologia(aula: PdfAula) -> list[dict[str, str]]:
    metodologia: list[dict[str, str]] = []
    foco_idx = 0
    pratica_idx = 0
    encerramento_bloco: Bloco | None = None
    for bloco in aula.blocos:
        if bloco.etapa == "Pause e responda":
            continue
        if bloco.etapa == "Para começar" or bloco.etapa == "Relembre":
            if any(item["titulo"] in {"Para começar", "Relembre"} for item in metodologia):
                continue
            metodologia.append({"titulo": bloco.etapa, "texto": compor_para_comecar(aula, bloco)})
        elif bloco.etapa == "Foco no conteúdo":
            if foco_idx >= 2:
                continue
            metodologia.append({"titulo": "Foco no conteúdo", "texto": compor_foco(aula, bloco, foco_idx)})
            foco_idx += 1
        elif bloco.etapa == "Na prática":
            if pratica_idx >= 2:
                continue
            metodologia.append({"titulo": "Na prática", "texto": compor_na_pratica(aula, bloco, pratica_idx)})
            pratica_idx += 1
        elif bloco.etapa == "Socialização":
            metodologia.append({
                "titulo": "Na prática",
                "texto": 'Socialização. Com "COM SUAS PALAVRAS", organizar a partilha das respostas, comparar estratégias de leitura e consolidar coletivamente o que foi aprendido.',
            })
            pratica_idx += 1
        elif bloco.etapa == "Encerramento":
            encerramento_bloco = bloco

    if not any(item["titulo"] in {"Para começar", "Relembre"} for item in metodologia):
        metodologia.insert(0, {"titulo": "Para começar", "texto": compor_para_comecar(aula, Bloco("Para começar", aula.texto[:800]))})
    if not any(item["titulo"] == "Foco no conteúdo" for item in metodologia):
        metodologia.insert(1, {"titulo": "Foco no conteúdo", "texto": compor_foco(aula, Bloco("Foco no conteúdo", aula.texto[:1800]), 0)})
    if not any(item["titulo"] == "Na prática" for item in metodologia):
        metodologia.append({"titulo": "Na prática", "texto": compor_na_pratica(aula, Bloco("Na prática", aula.texto[-1500:]), 0)})

    metodologia.append({"titulo": "Encerramento", "texto": compor_encerramento(aula, encerramento_bloco)})
    return metodologia


def checklist_contexto(aula: PdfAula) -> str:
    contexto = contexto_por_titulo(aula.titulo)
    mapa = {
        "debate_regrado": "debate",
        "artigo_opiniao": "artigo",
        "resenha_critica": "resenha",
        "manifesto": "manifesto",
        "dissertativo_argumentativo": "dissertacao",
        "variacao_norma": "variacao",
        "miniconto_microconto": "miniconto",
        "diario_pessoal": "diario",
        "relato_viagem": "relato",
        "texto_dramatico": "dramatico",
        "moldando_imagens": "postagem",
        "anuncios_digitais": "postagem",
        "imparcialidade": "artigo",
        "literaturas_africanas": "literatura",
        "literatura_latino": "literatura",
        "realismo_magico": "literatura",
        "realismo_portugal": "literatura",
        "realismo_brasil": "literatura",
        "identidade_brasileira": "literatura",
        "cronica": "texto",
        "realismo_naturalismo": "literatura",
        "intervencao_urbana": "texto",
        "parnasianismo": "literatura",
        "fernando_pessoa": "literatura",
        "modernismo_3geracao": "literatura",
    }
    return mapa.get(contexto, "texto")


def gerar_acompanhamento(aula: PdfAula) -> list[str]:
    contexto = checklist_contexto(aula)
    tema = tema_frase(aula.titulo)
    itens_por_contexto = {
        "debate": [
            "☑ Verificar se os estudantes reconhecem a estrutura do debate regrado e o papel de cada participante.",
            "☑ Observar se a turma sustenta opiniões com respeito aos turnos de fala e aos pontos de vista divergentes.",
            f"☑ Acompanhar se os registros e intervenções da turma retomam com clareza as ideias discutidas sobre {tema}.",
        ],
        "artigo": [
            "☑ Verificar se os estudantes identificam tese, argumentos e estratégias de persuasão no artigo lido.",
            "☑ Observar se a turma justifica posicionamentos com base nos textos e discussões da aula.",
            f"☑ Acompanhar se os registros produzidos retomam com clareza os pontos centrais de {tema}.",
        ],
        "resenha": [
            "☑ Verificar se os estudantes distinguem resumo, comentário e avaliação no gênero resenha crítica.",
            "☑ Observar se a turma mobiliza critérios de análise ao comentar a obra ou o texto trabalhado.",
            f"☑ Acompanhar se os registros produzidos articulam leitura, interpretação e posicionamento sobre {tema}.",
        ],
        "manifesto": [
            "☑ Verificar se os estudantes reconhecem a finalidade social e argumentativa do manifesto.",
            "☑ Observar se a turma identifica reivindicações, posicionamentos e recursos persuasivos no texto.",
            f"☑ Acompanhar se os registros retomam com clareza as ideias centrais discutidas em {tema}.",
        ],
        "dissertacao": [
            "☑ Verificar se os estudantes reconhecem a organização do texto dissertativo-argumentativo e a função da tese.",
            "☑ Observar se a turma relaciona argumentos, repertório e projeto de texto às discussões propostas na aula.",
            f"☑ Acompanhar se os registros escritos retomam com clareza os aspectos centrais de {tema}.",
        ],
        "variacao": [
            "☑ Verificar se os estudantes distinguem norma-padrão, variação linguística e adequação de linguagem.",
            "☑ Observar se a turma relaciona exemplos do material aos contextos formais e informais de uso da língua.",
            f"☑ Acompanhar se os registros produzidos retomam com clareza as discussões sobre {tema}.",
        ],
        "miniconto": [
            "☑ Verificar se os estudantes reconhecem concisão, sugestão e construção de sentidos nos minicontos lidos.",
            "☑ Observar se a turma identifica como poucos recursos linguísticos produzem efeitos de interpretação.",
            f"☑ Acompanhar se os registros da turma retomam com clareza os elementos centrais de {tema}.",
        ],
        "diario": [
            "☑ Verificar se os estudantes reconhecem marcas de subjetividade e organização do diário pessoal.",
            "☑ Observar se a turma relaciona linguagem, contexto de produção e sentidos construídos no texto.",
            f"☑ Acompanhar se os registros escritos retomam com clareza os aspectos centrais de {tema}.",
        ],
        "relato": [
            "☑ Verificar se os estudantes identificam experiências, impressões e recursos descritivos no relato de viagem.",
            "☑ Observar se a turma relaciona leitura, contexto e escolhas linguísticas ao efeito produzido no gênero.",
            f"☑ Acompanhar se os registros da turma retomam com clareza os elementos centrais de {tema}.",
        ],
        "dramatico": [
            "☑ Verificar se os estudantes reconhecem falas, rubricas e organização cênica do texto dramático.",
            "☑ Observar se a turma relaciona leitura, interpretação e efeitos de oralidade às cenas trabalhadas.",
            f"☑ Acompanhar se os registros da turma retomam com clareza os aspectos centrais de {tema}.",
        ],
        "postagem": [
            "☑ Verificar se os estudantes analisam criticamente postagens e comentários em seus efeitos de sentido.",
            "☑ Observar se a turma relaciona linguagem, interação em rede e posicionamento crítico às discussões da aula.",
            f"☑ Acompanhar se os registros produzidos retomam com clareza as reflexões sobre {tema}.",
        ],
        "literatura": [
            "☑ Verificar se os estudantes compreendem os sentidos produzidos pelo texto literário trabalhado na aula.",
            "☑ Observar se a turma relaciona leitura, contexto e escolhas de linguagem às discussões propostas.",
            f"☑ Acompanhar se os registros da turma retomam com clareza os aspectos centrais de {tema}.",
        ],
        "texto": [
            "☑ Verificar se os estudantes compreendem o texto ou gênero trabalhado e os sentidos construídos na aula.",
            "☑ Observar se a turma participa das leituras, discussões e análises retomando evidências do material.",
            f"☑ Acompanhar se os registros produzidos retomam com clareza os pontos centrais de {tema}.",
        ],
    }
    return itens_por_contexto[contexto]


def gerar_acessibilidade(aula: PdfAula) -> list[str]:
    contexto = checklist_contexto(aula)
    recurso = detectar_recurso(aula.texto, aula.titulo)

    leitura = {
        "debate": "☑ Disponibilizar roteiro simples com as regras do debate, os turnos de fala e perguntas orientadoras para cada participante.",
        "artigo": "☑ Realizar leitura guiada do artigo de opinião com destaque para tese, argumentos e conectivos importantes.",
        "resenha": "☑ Oferecer quadro de apoio com itens como obra analisada, opinião do autor, argumentos e avaliação final.",
        "manifesto": "☑ Disponibilizar roteiro de leitura com foco em reivindicações, posicionamentos e recursos de convencimento.",
        "dissertacao": "☑ Oferecer esquema de apoio com tese, argumentos, repertório e organização dos parágrafos do texto argumentativo.",
        "variacao": "☑ Disponibilizar exemplos comparativos de linguagem formal e informal para apoiar a discussão sobre adequação linguística.",
        "miniconto": "☑ Realizar leitura mediada do miniconto com pausas para explicitar pistas de sentido e inferências.",
        "diario": "☑ Oferecer roteiro de leitura com marcas de tempo, subjetividade e contexto do diário pessoal.",
        "relato": "☑ Disponibilizar perguntas orientadoras para apoiar a leitura das experiências e impressões presentes no relato.",
        "dramatico": "☑ Realizar leitura dramatizada ou mediada das falas e rubricas para apoiar a compreensão das cenas.",
        "postagem": "☑ Disponibilizar leitura guiada das postagens e comentários, destacando vocabulário, imagens e efeitos de sentido.",
        "literatura": "☑ Realizar leitura mediada do texto literário com pausas para explicação de referências, vocabulário e efeitos de linguagem.",
        "texto": "☑ Disponibilizar leitura guiada do material com pausas para destacar informações principais e apoiar a compreensão da turma.",
    }[contexto]

    apoio_conteudo = {
        "debate": "☑ Permitir anotações em tópicos, cartões de fala ou respostas orais mediadas antes da participação no debate.",
        "artigo": "☑ Oferecer modelo curto para organizar tese, argumentos e conclusão antes do registro escrito.",
        "resenha": "☑ Disponibilizar quadro comparativo ou modelo curto para organizar comentários e avaliação da leitura.",
        "manifesto": "☑ Permitir planejamento em tópicos ou frases curtas para organizar reivindicações e argumentos antes do registro final.",
        "dissertacao": "☑ Oferecer estrutura visual simples para organizar introdução, desenvolvimento e conclusão antes da escrita.",
        "variacao": "☑ Permitir respostas por oralidade, tópicos ou exemplos comparativos ao discutir usos da língua em diferentes contextos.",
        "miniconto": "☑ Permitir registros por tópicos, mapas de ideias ou resposta oral mediada para reconstruir os sentidos do texto.",
        "diario": "☑ Permitir registros em tópicos, quadros ou resposta oral mediada para organizar impressões e marcas do gênero.",
        "relato": "☑ Permitir registros por tópicos, esquemas ou resposta oral mediada para organizar a leitura e a interpretação.",
        "dramatico": "☑ Oferecer apoio por leitura em duplas, marcação de falas e registro em tópicos para organizar a cena trabalhada.",
        "postagem": "☑ Permitir registros por tópicos, leitura em duplas ou resposta oral mediada para discutir criticamente as interações em rede.",
        "literatura": "☑ Disponibilizar palavras-chave e perguntas orientadoras para apoiar a interpretação do texto literário.",
        "texto": "☑ Permitir registros em tópicos, esquemas ou resposta oral mediada para apoiar a participação da turma.",
    }[contexto]

    if recurso in {"meme", "cartaz", "infografico", "hq", "postagem", "video"}:
        terceiro = "☑ Destacar visualmente imagens, legendas, balões ou frames importantes para apoiar a leitura multissemiótica do material."
    else:
        terceiro = "☑ Garantir tempo adicional para leitura, retomada dos comandos e revisão das respostas antes da socialização."

    return [leitura, apoio_conteudo, terceiro]


def aplicar_estilo_run(run, *, bold: bool = False, size: float | None = None, color: RGBColor | None = None) -> None:
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:cs"), "Arial")


def aplicar_estilo_paragrafo(paragrafo, before: int = 0, after: int = 0, line: float = 1.15) -> None:
    pf = paragrafo.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def escrever_docx(destino: Path, ano_label: str, aulas: list[PdfAula]) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Pt(72)
    sec.right_margin = Pt(72)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    aplicar_estilo_paragrafo(p, after=3)
    run = p.add_run("Metodologias - Língua Portuguesa")
    aplicar_estilo_run(run, bold=True, size=18, color=COR_TITULO)

    p = doc.add_paragraph()
    aplicar_estilo_paragrafo(p, after=6)
    run = p.add_run(f"{ano_label} - 3º Bimestre")
    aplicar_estilo_run(run, size=13, color=COR_SUBTITULO)

    p = doc.add_paragraph()
    aplicar_estilo_paragrafo(p, after=8)
    run = p.add_run(
        "Material organizado com a metodologia de cada aula, seguida de acompanhamento da aprendizagem e acessibilidade. "
        "Os textos foram ajustados diretamente a partir dos PDFs reais da série."
    )
    aplicar_estilo_run(run, size=10.5)

    doc.add_paragraph()

    for aula in aulas:
        metodologia = montar_metodologia(aula)
        acompanhamento = gerar_acompanhamento(aula)
        acessibilidade = gerar_acessibilidade(aula)

        p = doc.add_paragraph()
        aplicar_estilo_paragrafo(p, before=8, after=3)
        run = p.add_run(f"AULA {aula.numero} - {aula.titulo}")
        aplicar_estilo_run(run, bold=True, size=14, color=COR_AULA)

        p = doc.add_paragraph()
        aplicar_estilo_paragrafo(p, after=2)
        run = p.add_run("Metodologia")
        aplicar_estilo_run(run, bold=True, size=12, color=COR_TITULO)

        for etapa in metodologia:
            p = doc.add_paragraph()
            aplicar_estilo_paragrafo(p, after=0)
            r1 = p.add_run(f"{etapa['titulo']}: ")
            aplicar_estilo_run(r1, bold=True, color=COR_ETAPA)
            r2 = p.add_run(etapa["texto"])
            aplicar_estilo_run(r2)

        p = doc.add_paragraph()
        aplicar_estilo_paragrafo(p, before=4, after=2)
        run = p.add_run("Acompanhamento da aprendizagem")
        aplicar_estilo_run(run, bold=True, size=12, color=COR_TITULO)

        for item in acompanhamento:
            p = doc.add_paragraph()
            aplicar_estilo_paragrafo(p, after=0)
            run = p.add_run(item)
            aplicar_estilo_run(run)

        p = doc.add_paragraph()
        aplicar_estilo_paragrafo(p, before=4, after=2)
        run = p.add_run("Acessibilidade")
        aplicar_estilo_run(run, bold=True, size=12, color=COR_TITULO)

        for item in acessibilidade:
            p = doc.add_paragraph()
            aplicar_estilo_paragrafo(p, after=0)
            run = p.add_run(item)
            aplicar_estilo_run(run)

        doc.add_paragraph()

    doc.save(str(destino))


def backup_arquivo(caminho: Path) -> Path | None:
    if not caminho.exists():
        return None
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    destino = BACKUP_DIR / f"{caminho.stem}_{stamp}{caminho.suffix}"
    shutil.copy2(caminho, destino)
    return destino


def listar_pdfs(pasta: Path) -> list[Path]:
    def chave(path: Path) -> int:
        match = re.search(r"AULA_(\d+)", path.name, re.I)
        return int(match.group(1)) if match else 999

    return sorted(pasta.glob("AULA_*.pdf"), key=chave)


def gerar_serie(pasta_nome: str, ano_label: str, nome_docx: str, preview: bool = False) -> None:
    pasta = BASE / pasta_nome
    pdfs = listar_pdfs(pasta)
    destino = pasta / nome_docx
    titulos_existentes = carregar_titulos_existentes(destino)
    aulas = []
    for pdf in pdfs:
        aula = carregar_aula(pdf)
        if aula.numero in titulos_existentes:
            aula.titulo = titulos_existentes[aula.numero]
        aulas.append(aula)
    if preview:
        print(f"\n### {ano_label}")
        for aula in aulas[:4]:
            print(f"\nAULA {aula.numero} - {aula.titulo}")
            for etapa in montar_metodologia(aula):
                print(f"{etapa['titulo']}: {etapa['texto']}")
            print("ACOMPANHAMENTO:", *gerar_acompanhamento(aula), sep="\n")
            print("ACESSIBILIDADE:", *gerar_acessibilidade(aula), sep="\n")
        return
    backup_arquivo(destino)
    escrever_docx(destino, ano_label, aulas)
    print(destino)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    preview = "--preview" in sys.argv
    for pasta_nome, ano_label, nome_docx in SERIES:
        gerar_serie(pasta_nome, ano_label, nome_docx, preview=preview)


if __name__ == "__main__":
    main()
