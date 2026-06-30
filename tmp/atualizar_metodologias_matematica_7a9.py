from __future__ import annotations

import re
import shutil
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook


BASE_DIR = Path(r"D:\PDF novos\MATEMATICA\AF\3_BIMESTRE")
BACKUP_DIR = Path(r"D:\PLANOS_LUAN\tmp\docx_backups")

COR_TITULO = RGBColor(0x1F, 0x49, 0x7D)
COR_AULA = RGBColor(0x00, 0x47, 0x70)
COR_ETAPA = RGBColor(0x2E, 0x74, 0xB5)

SERIES = [
    ("7_ANO", "Metodologias_Matematica_7_Ano_Ensino_Fundamental.docx", "GUIA_7_ANO_3_BIMESTRE.xlsx"),
    ("8_ANO", "Metodologias_Matematica_8_Ano_Ensino_Fundamental.docx", "GUIA_8_ANO_3_BIMESTRE.xlsx"),
    ("9_ANO", "Metodologias_Matematica_9_Ano_Ensino_Fundamental.docx", "GUIA_9_ANO_3_BIMESTRE.xlsx"),
]

TECNICAS = [
    "VIREM E CONVERSEM",
    "TODO MUNDO ESCREVE",
    "DE OLHO NO MODELO",
    "UM PASSO DE CADA VEZ",
    "COM SUAS PALAVRAS",
    "HORA DA LEITURA",
]

ETAPAS_PDF = [
    "Relembre",
    "Para começar",
    "Foco no conteúdo",
    "Pause e responda",
    "Na prática",
    "Encerramento",
    "Abertura",
    "Fechamento",
]

FILTROS_LINHA = {
    "na prática",
    "veja no livro!",
    "todo mundo escreve",
    "virem e conversem",
    "de olho no modelo",
    "um passo de cada vez",
    "com suas palavras",
    "relembre",
    "para começar",
    "foco no conteúdo",
    "foco no conteudo",
    "encerramento",
    "abertura",
    "fechamento",
    "resolução",
    "resolucao",
}


@dataclass
class AulaSecao:
    numero: int
    titulo: str
    metodologia_idx: int
    acompanhamento_idx: int


@dataclass
class AulaGuia:
    numero: int
    titulo: str
    conteudos: list[str]
    objetivos: list[str]


@dataclass
class AulaPdf:
    numero: int
    caminho: Path | None = None
    texto: str = ""
    contagens: dict[str, int] = field(default_factory=dict)
    atividades: list[str] = field(default_factory=list)
    tecnicas: list[str] = field(default_factory=list)


def normalizar(texto: str) -> str:
    texto = unicodedata.normalize("NFD", str(texto or "").lower())
    texto = "".join(ch for ch in texto if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", texto).strip()


def limpar_bullets(texto: str | None) -> list[str]:
    if not texto:
        return []
    bruto = str(texto).replace("\r", " ").replace("\n", " ").replace("\u200b", " ")
    bruto = re.sub(r"\s+", " ", bruto).strip()
    partes = bruto.split("•") if "•" in bruto else [bruto]
    saida = []
    for parte in partes:
        linha = re.sub(r"\s+", " ", parte).strip(" -")
        if linha:
            saida.append(linha.rstrip("."))
    return saida


def primeira_minuscula(texto: str) -> str:
    texto = texto.strip()
    if not texto:
        return texto
    return texto[:1].lower() + texto[1:]


def juntar_partes(partes: list[str], max_itens: int = 2) -> str:
    itens = [primeira_minuscula(p.strip().rstrip(".")) for p in partes if p.strip()]
    itens = itens[:max_itens]
    if not itens:
        return ""
    if len(itens) == 1:
        return itens[0]
    return ", ".join(itens[:-1]) + " e " + itens[-1]


def tema_limpo(titulo: str) -> str:
    return re.sub(r"^\s*AULA\s+\d+\s+[—-]\s*", "", titulo, flags=re.I).strip()


def tema_em_estudo(tema: str) -> str:
    tema = tema.strip()
    baixa = normalizar(tema)
    for art, contr in [("a ", "da "), ("o ", "do "), ("as ", "das "), ("os ", "dos ")]:
        if baixa.startswith(art):
            return contr + tema[len(art):]
    return "de " + tema


def extrair_numero_aula(texto: str) -> int | None:
    match = re.match(r"AULA\s+0*(\d+)", texto.strip(), flags=re.I)
    return int(match.group(1)) if match else None


def localizar_secoes(doc: Document) -> list[AulaSecao]:
    paragrafos = list(doc.paragraphs)
    titulos = []
    for idx, p in enumerate(paragrafos):
        numero = extrair_numero_aula(p.text)
        if numero is not None:
            titulos.append((numero, p.text.strip(), idx))

    secoes: list[AulaSecao] = []
    for pos, (numero, titulo, inicio) in enumerate(titulos):
        fim = titulos[pos + 1][2] - 1 if pos + 1 < len(titulos) else len(paragrafos) - 1
        metodologia_idx = -1
        acompanhamento_idx = -1
        for idx in range(inicio, fim + 1):
            atual = paragrafos[idx].text.strip().lower()
            if atual == "metodologia":
                metodologia_idx = idx
            elif atual == "acompanhamento da aprendizagem":
                acompanhamento_idx = idx
                break
        if metodologia_idx != -1 and acompanhamento_idx != -1:
            secoes.append(AulaSecao(numero, titulo, metodologia_idx, acompanhamento_idx))
    return secoes


def carregar_guia(caminho: Path) -> dict[int, AulaGuia]:
    wb = load_workbook(caminho, data_only=True)
    ws = wb[wb.sheetnames[0]]
    aulas: dict[int, AulaGuia] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        aula = row[0]
        if aula is None:
            continue
        try:
            numero = int(str(aula).strip())
        except ValueError:
            continue
        titulo = str(row[1] or "").replace("\n", " ").strip()
        conteudos = limpar_bullets(row[2])
        objetivos = limpar_bullets(row[3])
        aulas[numero] = AulaGuia(numero, titulo, conteudos, objetivos)
    return aulas


def extrair_atividades(caminho: Path) -> list[str]:
    vistas: set[str] = set()
    atividades: list[str] = []
    with pdfplumber.open(str(caminho)) as pdf:
        for page in pdf.pages:
            texto = page.extract_text() or ""
            if "Atividade" not in texto:
                continue
            linhas = [
                re.sub(r"\s+", " ", ln.replace("\u200b", " ").replace("\xa0", " ")).strip()
                for ln in texto.splitlines()
                if ln.strip()
            ]
            numero_atividade = None
            coletadas: list[str] = []
            iniciado = False
            for linha in linhas:
                match = re.search(r"Atividade\s*(\d+)", linha, flags=re.I)
                if match and not iniciado:
                    numero_atividade = match.group(1)
                    if numero_atividade in vistas:
                        break
                    iniciado = True
                    resto = linha[match.end():].strip(" :.-")
                    if resto:
                        coletadas.append(resto)
                    continue
                if not iniciado:
                    continue
                baixa = normalizar(linha)
                if baixa in FILTROS_LINHA:
                    continue
                if baixa.startswith("resolucao") or "disponivel em" in baixa or "acesso em" in baixa:
                    continue
                if re.fullmatch(r"\d+\s*minutos?", baixa):
                    continue
                if len(linha) < 8:
                    continue
                coletadas.append(linha)
                if len(coletadas) >= 4:
                    break
            if numero_atividade and coletadas:
                vistas.add(numero_atividade)
                texto_atividade = " ".join(coletadas)
                texto_atividade = re.sub(r"\s+", " ", texto_atividade).strip(" .")
                atividades.append(f"Atividade {numero_atividade}: {texto_atividade}")
    return atividades


def carregar_pdf(caminho: Path | None, numero: int) -> AulaPdf:
    if caminho is None or not caminho.exists():
        return AulaPdf(numero=numero)
    with pdfplumber.open(str(caminho)) as pdf:
        texto = "\n".join((page.extract_text() or "") for page in pdf.pages)
    contagens = {etapa: texto.lower().count(etapa.lower()) for etapa in ETAPAS_PDF}
    tecnicas = [t for t in TECNICAS if t.lower() in texto.lower()]
    atividades = extrair_atividades(caminho)
    return AulaPdf(
        numero=numero,
        caminho=caminho,
        texto=texto,
        contagens=contagens,
        atividades=atividades,
        tecnicas=tecnicas,
    )


def tecnico(texto: str, aula_pdf: AulaPdf, alvo: str) -> str:
    if alvo in aula_pdf.tecnicas:
        return f' com "{alvo}"'
    return ""


def frase_abertura(tema: str, aula_pdf: AulaPdf, guia: AulaGuia, numero: int) -> str:
    verbos = [
        "Iniciar a aula",
        "Abrir a aula",
        "Começar a aula",
        "Dar início à aula",
    ]
    verbo = verbos[numero % len(verbos)]
    base = f"{verbo} retomando com a turma o estudo {tema_em_estudo(primeira_minuscula(tema))} e convidando os estudantes a observar a situação apresentada no material."
    if guia.objetivos:
        base += f" O professor direciona a conversa para que a turma avance em {primeira_minuscula(guia.objetivos[0])}."
    if "VIREM E CONVERSEM" in aula_pdf.tecnicas:
        base += ' Com "VIREM E CONVERSEM", os estudantes levantam hipóteses iniciais e compartilham estratégias.'
    else:
        base += " Nesse primeiro momento, a turma levanta hipóteses e compartilha estratégias iniciais."
    if "TODO MUNDO ESCREVE" in aula_pdf.tecnicas:
        base += ' Em seguida, os alunos registram no caderno uma ideia inicial com "TODO MUNDO ESCREVE".'
    return base


def frase_relembre(tema: str, guia: AulaGuia, aula_pdf: AulaPdf) -> str:
    base = f"Retomar com a turma os conhecimentos já construídos sobre {primeira_minuscula(tema)}, preparando os estudantes para enfrentar novas situações do material."
    if guia.objetivos:
        base += f" O professor relembra os critérios principais para {primeira_minuscula(guia.objetivos[0])}."
    if "VIREM E CONVERSEM" in aula_pdf.tecnicas:
        base += ' Com "VIREM E CONVERSEM", a turma compartilha rapidamente o raciocínio antes de iniciar as atividades.'
    return base


def foco_conteudo_paragrafos(tema: str, guia: AulaGuia, aula_pdf: AulaPdf, numero: int) -> list[str]:
    saida: list[str] = []
    conteudos = guia.conteudos
    objetivos = guia.objetivos
    if conteudos:
        p1 = f"O professor explica {primeira_minuscula(juntar_partes(conteudos, 2))}, retomando exemplos do material e destacando as ideias centrais da aula."
        if "DE OLHO NO MODELO" in aula_pdf.tecnicas:
            p1 += ' Sempre que o slide indicar, essa explicação é apoiada por "DE OLHO NO MODELO".'
        saida.append(p1)
    if len(conteudos) > 1 or objetivos:
        partes = objetivos[:2] if objetivos else conteudos[1:3]
        p2 = f"Na sequência, o professor sistematiza com a turma os procedimentos necessários para {primeira_minuscula(juntar_partes(partes, 2))}, relacionando o conteúdo às estratégias de resolução presentes na aula."
        if "UM PASSO DE CADA VEZ" in aula_pdf.tecnicas:
            p2 += ' Quando necessário, a condução é organizada "UM PASSO DE CADA VEZ" para tornar o raciocínio mais visível.'
        saida.append(p2)
    if not saida:
        saida.append(f"O professor desenvolve {primeira_minuscula(tema)} de forma dialogada, retomando exemplos do material e sistematizando os pontos principais com a turma.")
    return saida[:2]


def resumir_atividade(atividade: str) -> str:
    texto = re.sub(r"^Atividade\s*\d+:\s*", "", atividade, flags=re.I).strip()
    texto = texto.replace("\u200b", " ").replace("\xa0", " ")
    for lixo in ["Veja no livro!", "TODO MUNDO ESCREVE", "VIREM E CONVERSEM", "DE OLHO NO MODELO", "COM SUAS PALAVRAS"]:
        texto = re.sub(re.escape(lixo), "", texto, flags=re.I)
    texto = re.sub(r"^\s*\d+\s*minutos?\s*", "", texto, flags=re.I)
    texto = re.sub(r"\b\d+\s*minutos?\b", "", texto, flags=re.I)
    texto = texto.replace("  ", " ")
    texto = re.sub(r"\s+", " ", texto)
    return texto[:260].rstrip(" ,.")


def frase_atividade(atividade: str, tema: str, aula_pdf: AulaPdf, numero: int, modo: str) -> str:
    match = re.match(r"Atividade\s*(\d+):", atividade, flags=re.I)
    num = match.group(1) if match else "1"
    resumo = resumir_atividade(atividade)
    if not resumo:
        resumo = f"os estudantes resolvem a proposta principal sobre {primeira_minuscula(tema)}"
    if modo == "problemas":
        texto = f"Atividade {num}: Os estudantes resolvem a situação proposta no material a partir da seguinte situação: {resumo}. O professor acompanha os registros e a justificativa de cada etapa da resolução."
    elif modo == "verificacao":
        texto = f"Atividade {num}: Os alunos resolvem a questão de verificação a partir da seguinte proposta: {resumo}. O professor retoma a estratégia utilizada para chegar à resposta."
    else:
        texto = f"Atividade {num}: A turma desenvolve a proposta do material a partir da seguinte situação: {resumo}. O professor acompanha os registros e compara estratégias durante a correção coletiva."
    if "TODO MUNDO ESCREVE" in aula_pdf.tecnicas and numero % 2 == 0:
        texto += ' O professor pode solicitar registros individuais com "TODO MUNDO ESCREVE" antes da socialização.'
    return texto


def frase_pratica_generica(tema: str, guia: AulaGuia, aula_pdf: AulaPdf, modo: str) -> str:
    if modo == "problemas":
        return f"Os estudantes resolvem problemas relacionados a {primeira_minuscula(tema)}, explicando os procedimentos utilizados e comparando estratégias ao final da correção."
    if modo == "verificacao":
        return f"A turma resolve as questões de verificação sobre {primeira_minuscula(tema)}, registrando os cálculos e justificando as respostas escolhidas."
    objetivo = guia.objetivos[0] if guia.objetivos else f"desenvolver situações sobre {primeira_minuscula(tema)}"
    return f"Os estudantes realizam a atividade principal da aula para {primeira_minuscula(objetivo)}, com acompanhamento do professor durante os registros e a correção."


def frase_encerramento(tema: str, guia: AulaGuia, aula_pdf: AulaPdf, numero: int) -> str:
    verbos = ["Encerrar a aula", "Fechar a aula", "Concluir a aula", "Finalizar a aula"]
    verbo = verbos[numero % len(verbos)]
    base = f'{verbo} retomando os pontos principais sobre {primeira_minuscula(tema)} e verificando o que a turma compreendeu ao longo da aula.'
    if "COM SUAS PALAVRAS" in aula_pdf.tecnicas:
        base += ' Com "COM SUAS PALAVRAS", os estudantes registram a ideia central ou o procedimento que mais os ajudou a resolver a proposta.'
    else:
        if guia.objetivos:
            base += f" Ao final, os estudantes sintetizam o que foi necessário para {primeira_minuscula(guia.objetivos[0])}."
    return base


def gerar_metodologia(secao: AulaSecao, guia: AulaGuia | None, aula_pdf: AulaPdf) -> list[tuple[str, str]]:
    tema = tema_limpo(secao.titulo)
    guia = guia or AulaGuia(secao.numero, tema, [], [])
    titulo_norm = normalizar(tema)
    is_matific = "matific" in titulo_norm
    is_verificacao = "verificacao" in titulo_norm
    is_problemas = "resolucao de problemas" in titulo_norm
    tem_relembre = aula_pdf.contagens.get("Relembre", 0) > 0
    tem_para = aula_pdf.contagens.get("Para começar", 0) > 0

    blocos: list[tuple[str, str]] = []
    if is_matific:
        conteudos = guia.conteudos[:3]
        lista = "; ".join(c.lower() for c in conteudos) if conteudos else primeira_minuscula(tema)
        abertura = f"O professor apresenta à turma os conteúdos que serão retomados na plataforma, com foco em {lista}. Antes do acesso, relembra com os estudantes o que já foi estudado e os procedimentos que serão revisados."
        pratica = "Orientar o acesso à plataforma e acompanhar a realização das atividades de revisão, incentivando novas tentativas quando necessário e pedindo que a turma consulte o caderno para relembrar estratégias já trabalhadas."
        fechamento = "Se necessário, consolidar a aprendizagem retomando no caderno uma ou mais atividades ligadas ao tema da revisão, com correção coletiva das estratégias mais recorrentes observadas durante a prática."
        return [("Abertura", abertura), ("Prática na Matific", pratica), ("Fechamento", fechamento)]

    if tem_relembre:
        blocos.append(("Relembre", frase_relembre(tema, guia, aula_pdf)))
    elif tem_para:
        blocos.append(("Para começar", frase_abertura(tema, aula_pdf, guia, secao.numero)))

    if not is_problemas and not is_verificacao:
        for texto in foco_conteudo_paragrafos(tema, guia, aula_pdf, secao.numero):
            blocos.append(("Foco no conteúdo", texto))

    atividades = aula_pdf.atividades
    modo_atividade = "regular"
    if is_verificacao:
        modo_atividade = "verificacao"
    elif is_problemas:
        modo_atividade = "problemas"

    if atividades:
        for atividade in atividades:
            blocos.append(("Na prática", frase_atividade(atividade, tema, aula_pdf, secao.numero, modo_atividade)))
    elif is_problemas or is_verificacao or aula_pdf.contagens.get("Na prática", 0) > 0:
        blocos.append(("Na prática", frase_pratica_generica(tema, guia, aula_pdf, modo_atividade)))

    blocos.append(("Encerramento", frase_encerramento(tema, guia, aula_pdf, secao.numero)))
    return blocos


def excluir_paragrafo(paragrafo) -> None:
    elemento = paragrafo._element
    parent = elemento.getparent()
    if parent is not None:
        parent.remove(elemento)


def novo_paragrafo_apos(paragrafo):
    novo = OxmlElement("w:p")
    paragrafo._p.addnext(novo)
    from docx.text.paragraph import Paragraph

    return Paragraph(novo, paragrafo._parent)


def aplicar_estilo_etapa(paragrafo, rotulo: str, texto: str, fonte: str | None) -> None:
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragrafo.paragraph_format.space_after = Pt(0)
    paragrafo.paragraph_format.space_before = Pt(0)
    run_rotulo = paragrafo.add_run(f"{rotulo}: ")
    if fonte:
        run_rotulo.font.name = fonte
        run_rotulo._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), fonte)
    run_rotulo.font.bold = True
    run_rotulo.font.color.rgb = COR_ETAPA
    run_texto = paragrafo.add_run(texto)
    if fonte:
        run_texto.font.name = fonte
        run_texto._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), fonte)


def aplicar_estilo_cabecalho(paragrafo, cor: RGBColor, tamanho: int, fonte: str | None) -> None:
    if not paragrafo.runs:
        paragrafo.add_run(paragrafo.text)
    for run in paragrafo.runs:
        if fonte:
            run.font.name = fonte
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), fonte)
        run.font.bold = True
        run.font.color.rgb = cor
        run.font.size = Pt(tamanho)


def backup_arquivo(caminho: Path) -> Path:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    destino = BACKUP_DIR / f"{caminho.stem}_{stamp}{caminho.suffix}"
    shutil.copy2(caminho, destino)
    return destino


def montar_pdf_map(pasta: Path) -> dict[int, Path]:
    saida = {}
    for pdf in pasta.glob("AULA_*.pdf"):
        match = re.match(r"AULA_(\d+)", pdf.name, flags=re.I)
        if match:
            saida[int(match.group(1))] = pdf
    return saida


def atualizar_serie(pasta: Path, docx_nome: str, guia_nome: str, preview: bool = False) -> tuple[Path, Path | None]:
    docx_path = pasta / docx_nome
    guia_path = pasta / guia_nome
    doc = Document(str(docx_path))
    fonte_modelo = doc.paragraphs[0].runs[0].font.name if doc.paragraphs and doc.paragraphs[0].runs else None
    secoes = localizar_secoes(doc)
    guia = carregar_guia(guia_path)
    pdf_map = montar_pdf_map(pasta)

    metodologias: dict[int, list[tuple[str, str]]] = {}
    for secao in secoes:
        aula_pdf = carregar_pdf(pdf_map.get(secao.numero), secao.numero)
        metodologias[secao.numero] = gerar_metodologia(secao, guia.get(secao.numero), aula_pdf)

    if preview:
        for numero in sorted(metodologias)[:4]:
            print(f"\n### AULA {numero}")
            for rotulo, texto in metodologias[numero]:
                print(f"{rotulo}: {texto}")
        return docx_path, None

    backup = backup_arquivo(docx_path)
    for secao in reversed(secoes):
        blocos = metodologias.get(secao.numero)
        if not blocos:
            continue
        paragrafos = list(doc.paragraphs)
        metodologia_p = paragrafos[secao.metodologia_idx]
        acompanhamento_p = paragrafos[secao.acompanhamento_idx]
        for idx in range(secao.acompanhamento_idx - 1, secao.metodologia_idx, -1):
            excluir_paragrafo(paragrafos[idx])
        referencia = metodologia_p
        for rotulo, texto in blocos:
            novo = novo_paragrafo_apos(referencia)
            aplicar_estilo_etapa(novo, rotulo, texto, fonte_modelo)
            referencia = novo

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("AULA "):
            aplicar_estilo_cabecalho(p, COR_AULA, 14, fonte_modelo)
        elif txt in {"Metodologia", "Acompanhamento da aprendizagem", "Acessibilidade"}:
            aplicar_estilo_cabecalho(p, COR_TITULO, 12, fonte_modelo)

    doc.save(str(docx_path))
    return docx_path, backup


def main(preview: bool = False) -> None:
    for pasta_nome, docx_nome, guia_nome in SERIES:
        pasta = BASE_DIR / pasta_nome
        docx_path, backup = atualizar_serie(pasta, docx_nome, guia_nome, preview=preview)
        if preview:
            print(f"\nPreview gerada para {docx_path.name}")
        else:
            print(f"Arquivo atualizado: {docx_path}")
            print(f"Backup criado em: {backup}")


if __name__ == "__main__":
    import sys

    main(preview="--preview" in sys.argv)
