from docx import Document
from pathlib import Path

ARQUIVOS = {
    Path(r'D:\PDF novos\CIENCIAS\AF\3_BIMESTRE\8_ANO\Metodologias_Ciencias_8_Ano_atualizado.docx'): {
        'Foco no conteúdo: Explicar o conceito de clima e suas características, destacando que é a média das condições meteorológicas ao longo do tempo. Usar a descrição de regiões climáticas como tropicais e áridas, abordando a variação climática dentro do estado de São Paulo. Explicar também a diferença entre tempo e clima e a ciência e previsão do tempo, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Explicar a diferença entre clima e tempo, mostrando que o clima corresponde aos padrões observados ao longo de muitos anos, enquanto o tempo descreve condições momentâneas da atmosfera. Destacar rapidamente fatores como latitude, altitude, relevo e maritimidade ou continentalidade, além do papel da meteorologia, do IPCC e das conferências climáticas citadas no material.',
        'Foco no conteúdo: Apresentar a definição e a estrutura de uma estação meteorológica. Explicar as variáveis que são monitoradas, como temperatura e pressão, e a importância desses dados para previsões climáticas. Discutir a instalação das estações meteorológicas. Explicar também a previsão do tempo, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Apresentar a estação meteorológica como espaço organizado para medir variáveis atmosféricas e produzir dados confiáveis para a previsão do tempo. Destacar rapidamente sua estrutura padronizada, o abrigo meteorológico, os critérios de instalação e o funcionamento de instrumentos como termômetro, barômetro e higrômetro.',
        'Foco no conteúdo: Apresentar imagens e descrições dos instrumentos meteorológicos: termômetro, barômetro, pluviômetro, anemômetro e biruta. Discutir suas funções e a importância na previsão do tempo. Explicar também as variáveis envolvidas na previsão do tempo, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Apresentar os instrumentos que serão construídos na aula - termômetro, barômetro, pluviômetro, anemômetro e biruta -, explicando o que cada um mede e como essas informações entram na previsão do tempo. Antecipar que os dados coletados pela miniestação serão registrados e comparados pela turma.',
        'Foco no conteúdo: Observar um mapa de temperatura média e discutir como diferentes elementos climáticos (temperatura, umidade, pressão atmosférica, ventos) influenciam o clima de uma região. Analisar os fatores climáticos e suas influências no clima local, relacionando o relevo e a vegetação como aspectos relevantes. Explicar também a circulação atmosférica, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Analisar como elementos climáticos e fatores geográficos interferem no clima de uma região, relacionando temperatura, umidade, pressão, ventos, relevo e vegetação. Explicar rapidamente a circulação atmosférica global, as células de circulação e os principais ventos mostrados nos esquemas do material.',
        'Foco no conteúdo: Conduzir uma análise das zonas climáticas utilizando mapas e tabelas. Orientar a leitura dos elementos, como legendas e informações climáticas, promovendo a compreensão dos conceitos de zonas tropicais, temperadas, subtropicais, subpolares e polares. Explicar também os climas regionais do planeta, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Conduzir a leitura de mapas e tabelas para identificar as zonas climáticas do planeta e compreender as características dos climas tropical, temperado, subtropical, subpolar e polar. Destacar rapidamente como latitude, circulação atmosférica, massas de ar e correntes oceânicas ajudam a explicar as diferenças entre as regiões analisadas.',
        'Foco no conteúdo: Introduzir a explicação sobre os fenômenos El Niño e La Niña, destacando como ocorrem e a relação com a circulação oceânica, utilizando um esquema da Célula Tropical. Explicar o funcionamento da Célula de Walker e suas implicações, mostrando como a distribuição de temperatura das águas influencia o ciclo climático. Explicar também as alterações climáticas regionais, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Introduzir os fenômenos El Niño e La Niña como fases opostas do ENOS, relacionando-os à circulação atmosférica e oceânica no Pacífico. Explicar rapidamente o papel dos ventos alísios, da Célula de Walker, do aquecimento ou resfriamento das águas superficiais e dos impactos regionais dessas anomalias climáticas.',
    },
    Path(r'D:\PDF novos\CIENCIAS\AF\3_BIMESTRE\9_ANO\Metodologias_Ciencias_9_Ano_atualizado.docx'): {
        'Foco no conteúdo: Apresentar o conceito de Unidades de Conservação (UCs) e discutir seu papel na preservação da biodiversidade e equilíbrio dos ecossistemas. Explicar as características das UCs, detalhando a diferenciação entre Proteção Integral e Uso Sustentável. Explicar também o papel do Instituto Florestal e os critérios para estabelecimento de Unidades de Conservação, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Retomar o conceito de Unidade de Conservação, explicando sua função na proteção da biodiversidade e do patrimônio ambiental. Destacar rapidamente o SNUC, o papel do Instituto Florestal, a diferença entre proteção integral e uso sustentável e exemplos de áreas citadas no material.',
        'Foco no conteúdo: Explicar a importância da água para a vida e desenvolvimento das sociedades, destacando a distribuição desigual dos recursos hídricos no planeta. Utilizar gráficos ou mapas que evidenciem a localização das águas doces. Realizar a leitura guiada de um infográfico sobre a distribuição da água no Brasil. Explicar também a hidrografia do Estado de São Paulo, a poluição da água, as fontes de poluição da água, os impactos da poluição da água, o tratamento da água e a gestão de recursos hídricos, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Retomar a importância da água para a vida e a distribuição desigual dos recursos hídricos, articulando hidrografia, regiões hidrográficas paulistas e conservação das bacias. Destacar rapidamente fontes de poluição, impactos sobre os ecossistemas, tratamento da água, aquíferos e gestão dos recursos hídricos.',
        'Foco no conteúdo: Explicar os conceitos de desenvolvimento sustentável, apresentando os três pilares: econômico, social e ambiental, e sua interconexão. Aprofundar nos Objetivos de Desenvolvimento Sustentável (ODS), discutindo sua importância global e metas específicas. Explicar também os objetivos do Desenvolvimento Sustentável, a pegada ecológica, o papel das ações individuais e coletivas na preservação ambiental, a economia circular, os impactos ambientais ocasionados pela destinação inadequada do “lixo”, a gestão dos resíduos sólidos, a diferenciação de aterro sanitário e “lixão” e o ciclo de vida de um aterro sanitário, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Retomar os conceitos de desenvolvimento sustentável e consumo consciente, explicando os três pilares da sustentabilidade e os Objetivos de Desenvolvimento Sustentável. Destacar rapidamente pegada ecológica, Dia da Sobrecarga da Terra, economia circular, gestão de resíduos, diferença entre aterro sanitário e lixão e impactos do descarte inadequado do lixo.',
    }
}

SUBSTITUICOES_RUN = {
    'coerencia': 'coerência',
    'conclusoes': 'conclusões',
    'cientificos': 'científicos',
    'questoes': 'questões',
    'solucoes': 'soluções',
    'elaboracao': 'elaboração',
}

PREFIXOS = ('Para começar:', 'Relembre:', 'Foco no conteúdo:', 'Pause e responda:', 'Na prática:', 'Socialização:', 'Encerramento:')


def limpar_runs(paragraph):
    p = paragraph._element
    for run in list(paragraph.runs):
        p.remove(run._element)


def aplicar_texto(paragraph, texto):
    limpar_runs(paragraph)
    prefixo = next((item for item in PREFIXOS if texto.startswith(item)), None)
    if prefixo is None:
        paragraph.add_run(texto)
        return
    corpo = texto[len(prefixo):].lstrip()
    run_prefixo = paragraph.add_run(prefixo)
    run_prefixo.bold = True
    if corpo:
        paragraph.add_run(' ' + corpo)


def substituir_em_runs(paragraph):
    for run in paragraph.runs:
        texto = run.text
        novo = texto
        for antigo, novo_txt in SUBSTITUICOES_RUN.items():
            novo = novo.replace(antigo, novo_txt)
        if novo != texto:
            run.text = novo

for caminho, substituicoes in ARQUIVOS.items():
    doc = Document(str(caminho))
    alterados = 0
    for paragraph in doc.paragraphs:
        texto_norm = ' '.join(paragraph.text.split())
        novo = substituicoes.get(texto_norm)
        if novo:
            aplicar_texto(paragraph, novo)
            alterados += 1
        substituir_em_runs(paragraph)
    doc.save(str(caminho))
    print(f'{caminho} -> {alterados} focos ajustados + correções ortográficas aplicadas')
