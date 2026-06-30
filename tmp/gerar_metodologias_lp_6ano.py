#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera o arquivo Word com Metodologias, Acompanhamento da Aprendizagem e Acessibilidade
para Língua Portuguesa - 6º Ano - 3º Bimestre

Baseado no modelo: Metodologias_Educacao_Financeira_8_Ano_CORRIGIDO.docx
"""

import json
import os
import re
import sys

# Garantir path correto
sys.path.insert(0, 'd:/PLANOS_LUAN')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── CONFIGURAÇÕES ─────────────────────────────────────────────────────────────
PASTA_AULAS   = r"D:\PDF novos\LINGUA_PORTUGUESA\AF\3_BIMESTRE\6_ANO"
ARQUIVO_SAIDA = r"D:\PDF novos\LINGUA_PORTUGUESA\AF\3_BIMESTRE\6_ANO\Metodologias_Lingua_Portuguesa_6_Ano_3_Bimestre.docx"

# ─── CORES ─────────────────────────────────────────────────────────────────────
COR_TITULO_DOC   = RGBColor(0x1F, 0x49, 0x7D)   # Azul escuro
COR_SUBTITULO    = RGBColor(0x2E, 0x74, 0xB5)   # Azul médio
COR_HEADING_AULA = RGBColor(0x00, 0x47, 0x70)   # Azul petróleo
COR_HEADING_SEC  = RGBColor(0x1F, 0x49, 0x7D)   # Azul escuro
COR_ETAPA        = RGBColor(0x2E, 0x74, 0xB5)   # Azul médio (etapa em negrito)
COR_TEXTO        = RGBColor(0x00, 0x00, 0x00)   # Preto


def set_cell_background(cell, color_hex):
    """Define cor de fundo de uma célula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Adiciona uma linha horizontal separadora."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_titulo_documento(doc, titulo, subtitulo, descricao):
    """Adiciona cabeçalho do documento."""
    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COR_TITULO_DOC

    # Subtítulo
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(subtitulo)
    run2.bold = False
    run2.font.size = Pt(13)
    run2.font.color.rgb = COR_SUBTITULO

    # Descrição/nota
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(10)
    run3 = p3.add_run(descricao)
    run3.italic = True
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    add_horizontal_rule(doc)


def add_heading_aula(doc, texto_aula):
    """Heading 1: título da aula."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto_aula)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COR_HEADING_AULA

    # borda inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '004770')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_heading_secao(doc, texto):
    """Heading 2: seção (Metodologia / Acompanhamento / Acessibilidade)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COR_HEADING_SEC


def add_objetivo(doc, texto_objetivo):
    """Adiciona objetivo da aula em itálico."""
    if not texto_objetivo:
        return
    # Extrair apenas a parte de objetivos (após 'Objetivos:')
    texto_exibir = texto_objetivo
    if 'Objetivos:' in texto_objetivo:
        partes = texto_objetivo.split('Objetivos:', 1)
        habilidade = partes[0].strip()
        objs = partes[1].strip()
        # Habilidade BNCC/Currículo
        if habilidade:
            ph = doc.add_paragraph()
            ph.paragraph_format.space_before = Pt(2)
            ph.paragraph_format.space_after = Pt(1)
            rh = ph.add_run('Habilidade: ')
            rh.bold = True
            rh.font.size = Pt(9)
            rh.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            rh2 = ph.add_run(habilidade.split(')')[-1].strip() if ')' in habilidade else habilidade)
            rh2.font.size = Pt(9)
            rh2.italic = True
            rh2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        # Objetivos
        texto_exibir = 'Objetivos: ' + objs
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto_exibir)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)


def add_metodologia(doc, etapas):
    """Adiciona as etapas de metodologia."""
    add_heading_secao(doc, 'Metodologia')
    for etapa in etapas:
        titulo_etapa = etapa.get('titulo', '').strip()
        texto_etapa  = etapa.get('texto', '').strip()
        if not titulo_etapa and not texto_etapa:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        # Etapa em negrito colorido
        if titulo_etapa:
            r_titulo = p.add_run(f'{titulo_etapa}: ')
            r_titulo.bold = True
            r_titulo.font.size = Pt(11)
            r_titulo.font.color.rgb = COR_ETAPA
        # Texto da etapa
        if texto_etapa:
            r_texto = p.add_run(texto_etapa)
            r_texto.font.size = Pt(11)
            r_texto.font.color.rgb = COR_TEXTO


def add_lista_itens(doc, itens):
    """Adiciona uma lista de itens de acompanhamento/acessibilidade."""
    for item in itens:
        texto = item.strip()
        if not texto:
            continue
        # Remove checkbox unicode se houver, manter texto limpo
        texto_limpo = texto.replace('☑', '').replace('☐', '').replace('✔', '').strip()
        if texto_limpo.startswith('- '):
            texto_limpo = texto_limpo[2:]
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.75)
        run = p.add_run(texto_limpo)
        run.font.size = Pt(10.5)


def listar_jsons(pasta):
    """Lista todos os JSONs de aula em ordem numérica."""
    arquivos = []
    for nome in os.listdir(pasta):
        if nome.lower().endswith('.json') and nome.upper().startswith('AULA_'):
            # Extrai número da aula
            m = re.search(r'AULA_(\d+)', nome, re.IGNORECASE)
            num = int(m.group(1)) if m else 999
            arquivos.append((num, nome))
    arquivos.sort(key=lambda x: x[0])
    return arquivos


def carregar_json(caminho):
    """Carrega um arquivo JSON com detecção de encoding."""
    for enc in ('utf-8-sig', 'utf-8', 'cp1252', 'latin-1'):
        try:
            with open(caminho, encoding=enc) as f:
                return json.load(f)
        except (UnicodeDecodeError, json.JSONDecodeError):
            continue
    raise RuntimeError(f"Não foi possível carregar: {caminho}")


def main():
    print("=" * 60)
    print("Gerando: Metodologias Língua Portuguesa 6º Ano 3º Bimestre")
    print("=" * 60)

    # Listar JSONs
    jsons = listar_jsons(PASTA_AULAS)
    if not jsons:
        print("ERRO: Nenhum JSON de aula encontrado.")
        sys.exit(1)
    print(f"Encontrados {len(jsons)} arquivos JSON de aula.")

    # Criar documento Word
    doc = Document()

    # Configurar margens da página
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_width  = Cm(21)   # A4
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Estilos padrão
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)

    # ── Cabeçalho do documento ──────────────────────────────────────────────────
    add_titulo_documento(
        doc,
        "Metodologias — Língua Portuguesa",
        "6º Ano · 3º Bimestre",
        "Material organizado com a metodologia de cada aula, seguida de "
        "acompanhamento da aprendizagem e acessibilidade. As técnicas citadas "
        "(VIREM E CONVERSEM, TODO MUNDO ESCREVE, PAUSE E RESPONDA, COM SUAS PALAVRAS "
        "e outras) são baseadas em Lemov (Aula Nota 10) e nos Princípios de Instrução "
        "de Rosenshine."
    )

    # ── Processar cada aula ─────────────────────────────────────────────────────
    for num_aula, nome_json in jsons:
        caminho_json = os.path.join(PASTA_AULAS, nome_json)
        print(f"  Processando Aula {num_aula:02d}: {nome_json}")

        try:
            dados = carregar_json(caminho_json)
        except Exception as e:
            print(f"    AVISO: Erro ao carregar {nome_json}: {e}")
            continue

        tema       = dados.get('tema', '')
        material   = dados.get('material', nome_json.replace('.json', ''))
        aprend     = dados.get('aprendizagem', '')
        metodologia    = dados.get('metodologia', [])
        acompanhamento = dados.get('acompanhamento', [])
        acessibilidade = dados.get('acessibilidade', [])

        # Título da aula
        titulo_aula = material if material else f"Aula {num_aula}"
        if not titulo_aula.upper().startswith('AULA'):
            titulo_aula = f"AULA {num_aula} — {titulo_aula}"
        add_heading_aula(doc, titulo_aula)

        # Objetivo/Habilidade
        if aprend:
            add_objetivo(doc, aprend)

        # Metodologia
        if metodologia:
            add_metodologia(doc, metodologia)
        else:
            add_heading_secao(doc, 'Metodologia')
            p = doc.add_paragraph()
            p.add_run('(Metodologia não disponível para esta aula.)').italic = True

        # Acompanhamento da aprendizagem
        add_heading_secao(doc, 'Acompanhamento da aprendizagem')
        if acompanhamento:
            add_lista_itens(doc, acompanhamento)
        else:
            p = doc.add_paragraph()
            p.add_run('(Sem itens de acompanhamento registrados.)').italic = True

        # Acessibilidade
        add_heading_secao(doc, 'Acessibilidade')
        if acessibilidade:
            add_lista_itens(doc, acessibilidade)
        else:
            p = doc.add_paragraph()
            p.add_run('(Sem itens de acessibilidade registrados.)').italic = True

        # Separador entre aulas
        add_horizontal_rule(doc)

    # ── Salvar ──────────────────────────────────────────────────────────────────
    doc.save(ARQUIVO_SAIDA)
    print()
    print(f"✔ Arquivo gerado com sucesso:")
    print(f"  {ARQUIVO_SAIDA}")
    print(f"  Total de aulas: {len(jsons)}")


if __name__ == '__main__':
    main()
