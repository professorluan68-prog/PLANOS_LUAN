from docx import Document
from pathlib import Path

caminho = Path(r'D:\PDF novos\CIENCIAS\AF\3_BIMESTRE\9_ANO\Metodologias_Ciencias_9_Ano_atualizado.docx')
substituicoes = {
    'Foco no conteúdo: Guiar a interpretação de cladogramas, explicando que eles representam relações de parentesco, destacando a diferença entre grupos monofiléticos e outros tipos. Os alunos devem observar que as sinapomorfias são características herdadas. Explicar também a análise de cladogramas e a evolução convergente e divergente, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Guiar a interpretação de cladogramas, explicando como eles representam relações de parentesco e como as sinapomorfias ajudam a identificar grupos monofiléticos. Destacar rapidamente a diferença entre evolução convergente e divergente e a leitura dos ramos, nós e características herdadas apresentados no material.',
    'Foco no conteúdo: Explicar o conceito de biodiversidade, destacando a variabilidade de organismos vivos, diversidade genética e diversidade ecológica, utilizando imagens como amostras de fungos e recifes de corais para exemplificar. Explicar também a evolução e biodiversidade, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Explicar o conceito de biodiversidade, destacando a variedade de organismos, a diversidade genética e a diversidade ecológica com exemplos como fungos e recifes de corais. Relacionar rapidamente essa diversidade aos processos evolutivos e à importância de sua conservação.',
    'Foco no conteúdo: Orientar a análise das relações sustentáveis com a natureza a partir de exemplos de culturas indígenas, destacando práticas de uso equilibrado dos recursos naturais. Explicar também as ações humanas e impactos na biodiversidade, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Orientar a análise das relações sustentáveis com a natureza a partir de exemplos de culturas indígenas, destacando práticas de uso equilibrado dos recursos naturais. Relacionar rapidamente essas práticas aos impactos das ações humanas sobre a biodiversidade e às formas de preservação discutidas no material.',
    'Foco no conteúdo: Explicar o conceito de bacia hidrográfica e seus componentes. Utilizar um mapa hidrográfico do estado de São Paulo para ilustrar como os rios se interconectam. Discutir a importância das matas ciliares na proteção das bacias hidrográficas. Reforçar como elas ajudam a manter a qualidade da água e a biodiversidade. Explicar também a hidrografia do estado de São Paulo e os mapas hidrográficos, retomando os exemplos, imagens, esquemas e dados apresentados no material.': 'Foco no conteúdo: Explicar o conceito de bacia hidrográfica e seus componentes, utilizando o mapa do estado de São Paulo para mostrar como os rios se interconectam. Destacar rapidamente as regiões hidrográficas paulistas, a leitura de mapas hidrográficos e a importância das matas ciliares para a qualidade da água e a biodiversidade.'
}
PREFIXOS = ('Para começar:', 'Relembre:', 'Foco no conteúdo:', 'Pause e responda:', 'Na prática:', 'Socialização:', 'Encerramento:')

def limpar_runs(paragraph):
    p = paragraph._element
    for run in list(paragraph.runs):
        p.remove(run._element)

def aplicar_texto(paragraph, texto):
    limpar_runs(paragraph)
    prefixo = next((item for item in PREFIXOS if texto.startswith(item)), None)
    if prefixo is None:
        paragraph.add_run(texto)
        return
    corpo = texto[len(prefixo):].lstrip()
    run_prefixo = paragraph.add_run(prefixo)
    run_prefixo.bold = True
    if corpo:
        paragraph.add_run(' ' + corpo)

doc = Document(str(caminho))
alterados = 0
for paragraph in doc.paragraphs:
    texto_norm = ' '.join(paragraph.text.split())
    novo = substituicoes.get(texto_norm)
    if novo:
        aplicar_texto(paragraph, novo)
        alterados += 1

doc.save(str(caminho))
print(f'{caminho} -> {alterados} focos finais ajustados')
