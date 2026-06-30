from __future__ import annotations

import re
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook


BASE = Path(r"D:\PDF novos\LINGUA_PORTUGUESA\AF\3_BIMESTRE")
BACKUP_DIR = Path(r"D:\PLANOS_LUAN\tmp\docx_backups")
WORKSPACE = Path(r"D:\PLANOS_LUAN")
if str(WORKSPACE) not in sys.path:
    sys.path.insert(0, str(WORKSPACE))

COR_TITULO = RGBColor(0x1F, 0x49, 0x7D)
COR_SUBTITULO = RGBColor(0x2E, 0x74, 0xB5)
COR_AULA = RGBColor(0x00, 0x47, 0x70)
COR_ETAPA = RGBColor(0x2E, 0x74, 0xB5)

TECNICAS = [
    "VIREM E CONVERSEM",
    "COM SUAS PALAVRAS",
    "HORA DA LEITURA",
    "TODO MUNDO ESCREVE",
    "UM PASSO DE CADA VEZ",
]

ETAPAS_VALIDAS = {
    "para comecar": "Para começar",
    "relembre": "Relembre",
    "foco no conteudo": "Foco no conteúdo",
    "na pratica": "Na prática",
    "socializacao": "Socialização",
    "encerramento": "Encerramento",
}

ABERTURAS = [
    "Abrir a aula",
    "Começar a aula",
    "Dar início à aula",
    "Iniciar o encontro",
    "Abrir o trabalho da aula",
]
RETOMADAS = [
    "Retomar a aula",
    "Reativar o percurso da aula",
    "Recuperar aprendizagens da aula anterior",
    "Retomar os pontos principais da aula anterior",
]
FOCOS = [
    "Explorar",
    "Apresentar",
    "Sistematizar",
    "Retomar",
    "Esclarecer",
]
ENCERRAMENTOS = [
    "Fechar a aula",
    "Concluir a aula",
    "Encerrar a aula",
    "Finalizar o percurso da aula",
]


@dataclass
class GuiaAula:
    numero: int
    titulo: str
    conteudos: list[str]
    objetivos: list[str]
    habilidade: str
    ae: str
    pratica: str
    campo: str
    objeto: str
    topico: str


@dataclass
class Bloco:
    etapa: str
    texto: str


def normalizar(texto: str) -> str:
    texto = unicodedata.normalize("NFD", str(texto or "").lower())
    texto = "".join(ch for ch in texto if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", texto).strip()


def limpar_item(texto: str) -> str:
    texto = str(texto or "").replace("\xa0", " ").strip()
    texto = re.sub(r"^[•\-\u2022]+\s*", "", texto)
    return re.sub(r"\s+", " ", texto).strip(" .;")


def quebrar_itens(texto: str) -> list[str]:
    bruto = str(texto or "").replace("\r", "\n")
    partes = []
    for linha in bruto.splitlines():
        item = limpar_item(linha)
        if item:
            if partes and item[:1].islower():
                partes[-1] = f"{partes[-1].rstrip(' ;:.')} {item}"
            else:
                partes.append(item)
    return partes


def escolher(lista: list[str], chave: int) -> str:
    return lista[chave % len(lista)]


def sanitizar_topico(texto: str) -> str:
    texto = limpar_item(texto)
    if not texto or texto in {"-", "–", "None"}:
        return ""
    return texto


def encontrar_pdf_por_aula(pasta: Path, numero: int) -> Path:
    candidatos = sorted(
        pasta.glob("AULA_*.pdf"),
        key=lambda p: int(re.search(r"AULA_(\d+)", p.name).group(1)),
    )
    for pdf in candidatos:
        match = re.search(r"AULA_(\d+)", pdf.name)
        if match and int(match.group(1)) == numero:
            return pdf
    raise FileNotFoundError(f"PDF da aula {numero} não encontrado em {pasta}")


def carregar_guia(caminho_xlsx: Path) -> dict[int, GuiaAula]:
    wb = load_workbook(caminho_xlsx, data_only=True)
    ws = wb.active
    linhas: dict[int, GuiaAula] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or not row[0]:
            continue
        numero = int(row[0])
        linhas[numero] = GuiaAula(
            numero=numero,
            titulo=str(row[1] or "").strip(),
            conteudos=quebrar_itens(row[2]),
            objetivos=quebrar_itens(row[3]),
            habilidade=str(row[4] or "").strip(),
            ae=str(row[5] or "").strip(),
            pratica=str(row[6] or "").strip(),
            campo=str(row[7] or "").strip(),
            objeto=str(row[8] or "").strip(),
            topico=str(row[9] or "").strip(),
        )
    return linhas


def extrair_texto_pdf(caminho_pdf: Path, max_paginas: int = 8) -> str:
    partes: list[str] = []
    with pdfplumber.open(str(caminho_pdf)) as pdf:
        for page in pdf.pages[:max_paginas]:
            partes.append(page.extract_text() or "")
    return "\n".join(partes)


def canonizar_etapa(texto: str) -> str | None:
    chave = normalizar(texto)
    chave = chave.replace("ç", "c")
    return ETAPAS_VALIDAS.get(chave)


def extrair_blocos(texto_pdf: str) -> list[Bloco]:
    padrao = re.compile(
        r"(?im)^(Para começar|Para comecar|Relembre|Foco no conteúdo|Foco no conteudo|Na prática|Na pratica|Socialização|Socializacao|Encerramento)\b"
    )
    matches = list(padrao.finditer(texto_pdf))
    blocos: list[Bloco] = []
    for idx, match in enumerate(matches):
        etapa = canonizar_etapa(match.group(1))
        if not etapa or etapa == "Pause e responda":
            continue
        inicio = match.end()
        fim = matches[idx + 1].start() if idx + 1 < len(matches) else len(texto_pdf)
        bloco = re.sub(r"\s+", " ", texto_pdf[inicio:fim]).strip()
        if not bloco:
            continue
        if blocos:
            anterior = blocos[-1]
            if anterior.etapa == etapa:
                a = normalizar(anterior.texto)[:180]
                b = normalizar(bloco)[:180]
                if a and b and (a == b or a[:120] == b[:120]):
                    continue
                if etapa == "Foco no conteúdo" and ("veja no livro" in b or "minutos" in b) and len(b) < 80:
                    continue
                if etapa == "Na prática" and ("veja no livro" in b or "minutos" in b) and len(b) < 80:
                    continue
        blocos.append(Bloco(etapa=etapa, texto=bloco))
    return blocos


def detectar_tecnica(texto: str) -> str | None:
    texto_norm = normalizar(texto).upper()
    for tecnica in TECNICAS:
        if normalizar(tecnica).upper() in texto_norm:
            return tecnica
    return None


def detectar_recurso(texto: str, guia: GuiaAula) -> str:
    base = f"{guia.titulo} {' '.join(guia.conteudos)} {texto}"
    texto_norm = normalizar(base)
    mapeamento = [
        ("texto jornalístico", ["jornal", "editorias", "site noticioso", "jornalistico"]),
        ("meme", ["meme"]),
        ("cartaz", ["cartaz", "campanha"]),
        ("infográfico", ["infografico"]),
        ("tirinha", ["tirinha"]),
        ("história em quadrinhos", ["hq", "quadrinhos"]),
        ("prefácio", ["prefacio"]),
        ("texto de orelha", ["orelha"]),
        ("sinopse", ["sinopse"]),
        ("verbete", ["verbete", "dicionario"]),
        ("poema", ["poema", "soneto"]),
        ("canção", ["cancao", "musica"]),
        ("miniconto", ["miniconto"]),
        ("novela", ["novela"]),
        ("reportagem", ["reportagem"]),
        ("notícia", ["noticia"]),
        ("vídeo", ["video", "youtube"]),
        ("obra visual", ["mona lisa", "obra", "imagem"]),
        ("postagem", ["blog", "post"]),
    ]
    for nome, chaves in mapeamento:
        if any(chave in texto_norm for chave in chaves):
            return nome
    return "texto do material"


def resumir_tema(guia: GuiaAula) -> str:
    primeiro = guia.conteudos[0] if guia.conteudos else guia.titulo
    return limpar_item(primeiro).lower()


def resumir_topico(guia: GuiaAula) -> str:
    topico = sanitizar_topico(guia.topico)
    return topico[0].lower() + topico[1:] if topico else ""


def frase_recurso(recurso: str, caso: str = "objeto") -> str:
    mapa = {
        "meme": {"objeto": "o meme", "de": "do meme"},
        "cartaz": {"objeto": "o cartaz", "de": "do cartaz"},
        "infográfico": {"objeto": "o infográfico", "de": "do infográfico"},
        "tirinha": {"objeto": "a tirinha", "de": "da tirinha"},
        "história em quadrinhos": {"objeto": "a HQ", "de": "da HQ"},
        "prefácio": {"objeto": "o prefácio", "de": "do prefácio"},
        "texto de orelha": {"objeto": "o texto de orelha", "de": "do texto de orelha"},
        "sinopse": {"objeto": "a sinopse", "de": "da sinopse"},
        "verbete": {"objeto": "o verbete", "de": "do verbete"},
        "poema": {"objeto": "o poema", "de": "do poema"},
        "canção": {"objeto": "a canção", "de": "da canção"},
        "miniconto": {"objeto": "o miniconto", "de": "do miniconto"},
        "novela": {"objeto": "o trecho de novela", "de": "do trecho de novela"},
        "reportagem": {"objeto": "a reportagem", "de": "da reportagem"},
        "notícia": {"objeto": "a notícia", "de": "da notícia"},
        "texto jornalístico": {"objeto": "o texto jornalístico", "de": "do texto jornalístico"},
        "vídeo": {"objeto": "o vídeo", "de": "do vídeo"},
        "obra visual": {"objeto": "as imagens", "de": "das imagens"},
        "postagem": {"objeto": "a postagem", "de": "da postagem"},
        "texto do material": {"objeto": "o texto do material", "de": "do texto do material"},
    }
    return mapa.get(recurso, mapa["texto do material"]).get(caso, mapa["texto do material"]["objeto"])


def foco_principal(guia: GuiaAula, indice: int = 0) -> str:
    base = guia.conteudos or guia.objetivos or [guia.titulo]
    return limpar_item(base[min(indice, len(base) - 1)]).lower()


def objetivo_relevante(guia: GuiaAula, indice: int = 0) -> str:
    objetivos = [limpar_item(item).lower() for item in (guia.objetivos or []) if limpar_item(item)]
    invalidos_finais = (" de", " da", " do", " dos", " das", " e", " em", " com")
    candidatos = objetivos[indice:] + objetivos[:indice]
    for candidato in candidatos:
        if len(candidato) < 20:
            continue
        if any(candidato.endswith(sufixo) for sufixo in invalidos_finais):
            continue
        return candidato
    return objetivos[0] if objetivos else limpar_item(guia.titulo).lower()


def frase_para_comecar(guia: GuiaAula, bloco: Bloco) -> str:
    verbo = escolher(RETOMADAS if bloco.etapa == "Relembre" else ABERTURAS, guia.numero)
    tecnica = detectar_tecnica(bloco.texto)
    recurso = detectar_recurso(bloco.texto, guia)
    tema = resumir_tema(guia)
    tecnica_txt = f' com "{tecnica}"' if tecnica else ""
    if recurso in {"meme", "cartaz", "tirinha", "obra visual"}:
        corpo = f"{verbo}{tecnica_txt}, observando {frase_recurso(recurso)} do material e mobilizando hipóteses sobre {tema}."
    elif recurso == "vídeo":
        corpo = f"{verbo}{tecnica_txt}, retomando o vídeo indicado e convidando a turma a comentar como ele se relaciona com {tema}."
    else:
        corpo = f"{verbo}{tecnica_txt}, mobilizando conhecimentos prévios da turma sobre {tema} e preparando a leitura do material."
    return corpo


def frase_foco(guia: GuiaAula, bloco: Bloco, indice: int) -> str:
    verbo = escolher(FOCOS, guia.numero + indice)
    conteudos = guia.conteudos or [guia.titulo]
    topico = resumir_topico(guia)
    base = limpar_item(conteudos[min(indice, len(conteudos) - 1)]).lower()
    tecnica = detectar_tecnica(bloco.texto)
    tecnica_txt = f' com "{tecnica}"' if tecnica else ""
    if indice == 0:
        frase = f"{verbo} {base}{tecnica_txt}, destacando como esse estudo ajuda a compreender o gênero e os sentidos construídos no texto."
    else:
        frase = f"{verbo} {base}{tecnica_txt}, retomando exemplos do material e relacionando a leitura ao desenvolvimento da aula."
    if topico:
        frase += f" Sempre que fizer sentido, articular também {topico} ao texto trabalhado."
    return frase


def frase_pratica(guia: GuiaAula, bloco: Bloco, indice: int) -> str:
    tecnica = detectar_tecnica(bloco.texto)
    tecnica_leitura = "HORA DA LEITURA" if "HORA DA LEITURA" in bloco.texto else tecnica
    recurso = detectar_recurso(bloco.texto, guia)
    objetivo = objetivo_relevante(guia, indice)
    enfoque = foco_principal(guia, min(indice, max(len(guia.conteudos) - 1, 0)))
    atividades = [
        "Atividade 1",
        "Atividade 2",
        "Atividade 3",
        "Desafio",
        "Produção",
    ]
    rotulo = atividades[min(indice, len(atividades) - 1)]
    acao_inicio = {
        "infográfico": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem o infográfico do material',
        "cartaz": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos analisam o cartaz proposto',
        "meme": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos observam o meme apresentado',
        "tirinha": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem a tirinha do material',
        "história em quadrinhos": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem a HQ proposta',
        "poema": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem o poema do material',
        "canção": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem a letra da canção indicada',
        "prefácio": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem o trecho de prefácio apresentado',
        "texto de orelha": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem o texto de orelha proposto',
        "sinopse": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem a sinopse indicada',
        "verbete": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem o verbete do material',
        "miniconto": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem o miniconto proposto',
        "novela": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem o trecho de novela do material',
        "reportagem": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem a reportagem selecionada',
        "notícia": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem a notícia do material',
        "texto jornalístico": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem o texto jornalístico indicado',
        "vídeo": f'Com "{tecnica or "COM SUAS PALAVRAS"}", os alunos retomam o vídeo indicado',
        "obra visual": f'Com "{tecnica or "COM SUAS PALAVRAS"}", os alunos observam as imagens apresentadas',
        "postagem": f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos leem a postagem destacada',
    }.get(recurso, f'Com "{tecnica_leitura or "HORA DA LEITURA"}", os alunos analisam o texto do material')

    if "grupo" in normalizar(bloco.texto) or "trio" in normalizar(bloco.texto) or "dupla" in normalizar(bloco.texto):
        fechamento = ' Depois, em colaboração com os colegas, organizam as respostas e socializam as conclusões com a mediação do professor.'
    elif "estacao" in normalizar(bloco.texto):
        fechamento = " Depois, circulam pelas estações propostas, registrando observações e retomando o que foi lido."
    elif "produ" in normalizar(bloco.texto):
        fechamento = " Depois, produzem o que o material solicita, revisam o que foi elaborado e conferem coletivamente os resultados."
    else:
        fechamento = " Em seguida, respondem às questões do material, registram as ideias principais e acompanham a correção feita pelo professor."

    return f"{rotulo}. {acao_inicio}, com foco em {enfoque}. O trabalho orienta a turma a {objetivo}.{fechamento}"


def frase_encerramento(guia: GuiaAula, bloco: Bloco) -> str:
    verbo = escolher(ENCERRAMENTOS, guia.numero)
    tecnica = detectar_tecnica(bloco.texto) or "COM SUAS PALAVRAS"
    tema = resumir_tema(guia)
    topico = resumir_topico(guia)
    frase = f'{verbo} com "{tecnica}", retomando as ideias centrais sobre {tema} e verificando o que a turma compreendeu ao longo da aula.'
    if topico:
        frase += f" Ao final, reforçar como {topico} aparece articulado ao texto estudado."
    return frase


def gerar_metodologia(guia: GuiaAula, blocos: list[Bloco]) -> list[dict[str, str]]:
    if not blocos:
        blocos = [Bloco("Para começar", ""), Bloco("Foco no conteúdo", ""), Bloco("Na prática", ""), Bloco("Encerramento", "")]
    blocos_filtrados: list[Bloco] = []
    contagem = {"Para começar": 0, "Relembre": 0, "Foco no conteúdo": 0, "Na prática": 0, "Socialização": 0, "Encerramento": 0}
    limites = {"Para começar": 1, "Relembre": 1, "Foco no conteúdo": 2, "Na prática": 2, "Socialização": 1, "Encerramento": 1}
    for bloco in blocos:
        if contagem[bloco.etapa] >= limites[bloco.etapa]:
            continue
        blocos_filtrados.append(bloco)
        contagem[bloco.etapa] += 1
    blocos = blocos_filtrados
    metodologia: list[dict[str, str]] = []
    focos = 0
    praticas = 0
    for bloco in blocos:
        if bloco.etapa == "Para começar" or bloco.etapa == "Relembre":
            texto = frase_para_comecar(guia, bloco)
        elif bloco.etapa == "Foco no conteúdo":
            texto = frase_foco(guia, bloco, focos)
            focos += 1
        elif bloco.etapa == "Na prática":
            texto = frase_pratica(guia, bloco, praticas)
            praticas += 1
        elif bloco.etapa == "Socialização":
            texto = 'Com "COM SUAS PALAVRAS", os alunos socializam as respostas, comparam estratégias de leitura e retomam o que foi aprendido no material.'
        elif bloco.etapa == "Encerramento":
            texto = frase_encerramento(guia, bloco)
        else:
            continue
        metodologia.append({"titulo": bloco.etapa, "texto": texto})

    if not any(item["titulo"] == "Encerramento" for item in metodologia):
        metodologia.append({"titulo": "Encerramento", "texto": frase_encerramento(guia, Bloco("Encerramento", ""))})
    if not any(item["titulo"] == "Foco no conteúdo" for item in metodologia):
        pos = 1 if metodologia else 0
        metodologia.insert(pos, {"titulo": "Foco no conteúdo", "texto": frase_foco(guia, Bloco("Foco no conteúdo", ""), 0)})
    if not any(item["titulo"] == "Na prática" for item in metodologia):
        pos = next((idx for idx, item in enumerate(metodologia) if item["titulo"] == "Encerramento"), len(metodologia))
        metodologia.insert(pos, {"titulo": "Na prática", "texto": frase_pratica(guia, Bloco("Na prática", ""), 0)})
    return metodologia


def gerar_acompanhamento(guia: GuiaAula) -> list[str]:
    objetivo1 = objetivo_relevante(guia, 0)
    objetivo2 = objetivo_relevante(guia, 1)
    topico = resumir_topico(guia)
    itens = [
        f"☑ Verificar se os estudantes conseguem {objetivo1}.",
        f"☑ Observar se a turma relaciona a leitura e as discussões da aula a {objetivo2}.",
    ]
    if topico:
        itens.append(f"☑ Acompanhar se os registros e respostas evidenciam o uso de {topico} no texto estudado.")
    else:
        itens.append("☑ Acompanhar se os registros produzidos durante a aula retomam com clareza as ideias centrais do material.")
    return itens[:3]


def gerar_acessibilidade(guia: GuiaAula, texto_pdf: str) -> list[str]:
    recurso = detectar_recurso(texto_pdf, guia)
    itens = []
    if recurso in {"poema", "canção", "prefácio", "texto de orelha", "sinopse", "miniconto", "novela", "reportagem", "notícia", "texto jornalístico", "verbete", "postagem"}:
        itens.append(f"☑ Realizar leitura guiada {frase_recurso(recurso, 'de')} com pausas para destacar vocabulário, informações principais e pistas de interpretação.")
    elif recurso in {"tirinha", "história em quadrinhos", "obra visual", "cartaz", "infográfico"}:
        itens.append(f"☑ Conduzir leitura mediada {frase_recurso(recurso, 'de')}, chamando atenção para elementos visuais, legendas, balões e relações com o texto verbal.")
    else:
        itens.append("☑ Oferecer leitura mediada do material com pausas curtas para retomada das ideias principais.")

    topico = resumir_topico(guia)
    if topico:
        itens.append(f"☑ Disponibilizar roteiro simples com perguntas orientadoras e exemplos do material para apoiar a compreensão de {topico}.")
    else:
        itens.append("☑ Disponibilizar roteiro simples com perguntas orientadoras e palavras-chave do conteúdo trabalhado na aula.")

    itens.append("☑ Permitir respostas por fala, registro em tópicos, marcações no texto ou organização de ideias com apoio do professor.")
    return itens[:3]


def aplicar_estilo_run(run, *, bold: bool = False, size: float | None = None, color: RGBColor | None = None) -> None:
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:cs"), "Arial")


def aplicar_estilo_paragrafo(paragrafo, before: int = 0, after: int = 0, line: float = 1.15) -> None:
    pf = paragrafo.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def escrever_docx(
    caminho_saida: Path,
    ano_label: str,
    aulas: Iterable[tuple[int, str, list[dict[str, str]], list[str], list[str]]],
) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(72)
    sec.bottom_margin = Pt(72)
    sec.left_margin = Pt(72)
    sec.right_margin = Pt(72)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    aplicar_estilo_paragrafo(p, after=3)
    run = p.add_run("Metodologias — Língua Portuguesa")
    aplicar_estilo_run(run, bold=True, size=18, color=COR_TITULO)

    p = doc.add_paragraph()
    aplicar_estilo_paragrafo(p, after=6)
    run = p.add_run(f"{ano_label} · 3º Bimestre")
    aplicar_estilo_run(run, size=13, color=COR_SUBTITULO)

    p = doc.add_paragraph()
    aplicar_estilo_paragrafo(p, after=8)
    run = p.add_run(
        "Material organizado com a metodologia de cada aula, seguida de acompanhamento da aprendizagem e acessibilidade. "
        'As técnicas citadas (VIREM E CONVERSEM, TODO MUNDO ESCREVE, HORA DA LEITURA, COM SUAS PALAVRAS e outras) '
        "são incorporadas ao texto da aula de forma direta e objetiva."
    )
    aplicar_estilo_run(run, size=10.5)

    doc.add_paragraph()

    for numero, titulo, metodologia, acompanhamento, acessibilidade in aulas:
        p = doc.add_paragraph()
        aplicar_estilo_paragrafo(p, before=8, after=3)
        run = p.add_run(f"AULA {numero} - {titulo}")
        aplicar_estilo_run(run, bold=True, size=14, color=COR_AULA)

        p = doc.add_paragraph()
        aplicar_estilo_paragrafo(p, after=2)
        run = p.add_run("Metodologia")
        aplicar_estilo_run(run, bold=True, size=12, color=COR_TITULO)

        for etapa in metodologia:
            p = doc.add_paragraph()
            aplicar_estilo_paragrafo(p, after=0)
            r1 = p.add_run(f"{etapa['titulo']}: ")
            aplicar_estilo_run(r1, bold=True, color=COR_ETAPA)
            r2 = p.add_run(etapa["texto"])
            aplicar_estilo_run(r2)

        p = doc.add_paragraph()
        aplicar_estilo_paragrafo(p, before=4, after=2)
        run = p.add_run("Acompanhamento da aprendizagem")
        aplicar_estilo_run(run, bold=True, size=12, color=COR_TITULO)

        for item in acompanhamento[:3]:
            p = doc.add_paragraph()
            aplicar_estilo_paragrafo(p, after=0)
            run = p.add_run(item)
            aplicar_estilo_run(run)

        p = doc.add_paragraph()
        aplicar_estilo_paragrafo(p, before=4, after=2)
        run = p.add_run("Acessibilidade")
        aplicar_estilo_run(run, bold=True, size=12, color=COR_TITULO)

        for item in acessibilidade[:3]:
            p = doc.add_paragraph()
            aplicar_estilo_paragrafo(p, after=0)
            run = p.add_run(item)
            aplicar_estilo_run(run)

        doc.add_paragraph()

    doc.save(str(caminho_saida))


def carregar_metodologias_existentes_6(caminho_docx: Path) -> dict[int, dict[str, list]]:
    from core.referencias_portugues import _carregar_referencias_docx, _extrair_metodologia_texto

    bruto = _carregar_referencias_docx(str(caminho_docx))
    return {
        numero: {
            "metodologia": _extrair_metodologia_texto(
                "\n".join(
                    f"{item.get('titulo', '')}: {item.get('texto', '')}"
                    for item in (ref.get("metodologia") or [])
                )
            )
            or list(ref.get("metodologia") or []),
            "acompanhamento": list(ref.get("acompanhamento") or []),
            "acessibilidade": list(ref.get("acessibilidade") or []),
            "titulo": str(ref.get("titulo") or ""),
        }
        for numero, ref in bruto.items()
    }


def backup_arquivo(caminho: Path, sufixo: str) -> None:
    if not caminho.exists():
        return
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    destino = BACKUP_DIR / f"{caminho.stem}_{sufixo}{caminho.suffix}"
    destino.write_bytes(caminho.read_bytes())


def gerar_ano(
    pasta: Path,
    ano_label: str,
    guia_path: Path,
    saida_path: Path,
    preservar_6_ate12: dict[int, dict[str, list]] | None = None,
) -> None:
    guia = carregar_guia(guia_path)
    aulas_saida = []
    for numero in sorted(guia):
        registro = guia[numero]
        pdf = encontrar_pdf_por_aula(pasta, numero)
        texto_pdf = extrair_texto_pdf(pdf)
        blocos = extrair_blocos(texto_pdf)

        if preservar_6_ate12 and numero in preservar_6_ate12 and numero <= 12:
            metodologia = preservar_6_ate12[numero]["metodologia"]
            acompanhamento = [
                item if str(item).startswith("☑") else f"☑ {str(item).strip()}"
                for item in preservar_6_ate12[numero]["acompanhamento"][:3]
            ]
            acessibilidade = [
                item if str(item).startswith("☑") else f"☑ {str(item).strip()}"
                for item in preservar_6_ate12[numero]["acessibilidade"][:3]
            ]
            titulo = preservar_6_ate12[numero]["titulo"] or registro.titulo
        else:
            metodologia = gerar_metodologia(registro, blocos)
            acompanhamento = gerar_acompanhamento(registro)
            acessibilidade = gerar_acessibilidade(registro, texto_pdf)
            titulo = registro.titulo

        aulas_saida.append((numero, titulo, metodologia, acompanhamento, acessibilidade))

    escrever_docx(saida_path, ano_label, aulas_saida)


def imprimir_previas() -> None:
    config = [
        (
            BASE / "6_ANO",
            "6º Ano",
            BASE / "6_ANO" / "GUIA_6_ANO_3_BIMESTRE.xlsx",
        ),
        (
            BASE / "7_ANO",
            "7º Ano",
            BASE / "7_ANO" / "GUIA_7_ANO_3_BIMESTRE.xlsx",
        ),
        (
            BASE / "8_ANO",
            "8º Ano",
            BASE / "8_ANO" / "GUIA_8_ANO_3_BIMESTRE.xlsx",
        ),
        (
            BASE / "9_ANO",
            "9º Ano",
            BASE / "9_ANO" / "RELATORIOS_CONFERENCIA_PLANOS" / "GUIA_9_ANO_3_BIMESTRE.xlsx",
        ),
    ]
    for pasta, _, guia_path in config:
        guia = carregar_guia(guia_path)
        for numero in [1, 13, 24]:
            if numero not in guia:
                continue
            registro = guia[numero]
            pdf = encontrar_pdf_por_aula(pasta, numero)
            texto_pdf = extrair_texto_pdf(pdf)
            blocos = extrair_blocos(texto_pdf)
            metodologia = gerar_metodologia(registro, blocos)
            print(f"\n## {pasta.name} AULA {numero} - {registro.titulo}")
            for etapa in metodologia[:5]:
                print(f"{etapa['titulo']}: {etapa['texto']}")
            print("ACOMPANHAMENTO:", gerar_acompanhamento(registro))
            print("ACESSIBILIDADE:", gerar_acessibilidade(registro, texto_pdf))


def gerar() -> None:
    seis_docx = BASE / "6_ANO" / "Metodologias_Lingua_Portuguesa_6_Ano_3_Bimestre.docx"
    preservar = carregar_metodologias_existentes_6(seis_docx)
    backup_arquivo(seis_docx, "antes_fechamento_total")

    gerar_ano(
        BASE / "6_ANO",
        "6º Ano",
        BASE / "6_ANO" / "GUIA_6_ANO_3_BIMESTRE.xlsx",
        seis_docx,
        preservar_6_ate12=preservar,
    )
    gerar_ano(
        BASE / "7_ANO",
        "7º Ano",
        BASE / "7_ANO" / "GUIA_7_ANO_3_BIMESTRE.xlsx",
        BASE / "7_ANO" / "Metodologias_Lingua_Portuguesa_7_Ano_3_Bimestre.docx",
    )
    gerar_ano(
        BASE / "8_ANO",
        "8º Ano",
        BASE / "8_ANO" / "GUIA_8_ANO_3_BIMESTRE.xlsx",
        BASE / "8_ANO" / "Metodologias_Lingua_Portuguesa_8_Ano_3_Bimestre.docx",
    )
    gerar_ano(
        BASE / "9_ANO",
        "9º Ano",
        BASE / "9_ANO" / "RELATORIOS_CONFERENCIA_PLANOS" / "GUIA_9_ANO_3_BIMESTRE.xlsx",
        BASE / "9_ANO" / "Metodologias_Lingua_Portuguesa_9_Ano_3_Bimestre_REVISADO.docx",
    )


if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding="utf-8")

    if "--preview" in sys.argv:
        imprimir_previas()
    else:
        gerar()
