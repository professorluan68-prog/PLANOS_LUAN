import docx
import sys

def read_docx(file_path):
    doc = docx.Document(file_path)
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n\n".join(lines)

text = read_docx(r"D:\PLANOS_LUAN\Planos feitos\LINGUA_PORTUGUESA_AF\Planos_LP_6_ANO.docx")
with open(r"d:\PLANOS_LUAN\extract_lp_6ano.md", "w", encoding="utf-8") as f:
    f.write(text)
print("Saved to extract_lp_6ano.md")
