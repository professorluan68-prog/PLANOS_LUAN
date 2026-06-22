from __future__ import annotations

import re
from pathlib import Path

import pdfplumber
from docx import Document


BASE_AF_3B = Path(r"D:\PDF novos\LINGUA_INGLESA\AF\3_BIMESTRE")


VARIACOES_TRILHA = [
    {
        "abertura": "Retomar a aula anterior sobre Trilha de aprendizagem individual e conectar os registros ja produzidos ao novo foco do dia.",
        "foco": "Desenvolver a trilha de forma progressiva, articulando explicacao, exemplo e atividade guiada com comandos simples em ingles.",
        "pratica": "Solicitar que os estudantes comparem respostas de atividades anteriores, identifiquem avancos e registrem duvidas persistentes.",
        "fechamento": "Registrar uma sintese parcial e uma pergunta para orientar a proxima aula da sequencia.",
        "acompanhamento": [
            "Verificar se os estudantes compreendem o tema central da aula e reconhecem as ideias principais trabalhadas.",
            "Conferir se as producoes finais apresentam clareza, coerencia e retomada dos conceitos estudados.",
            "Observar a participacao, os registros e a forma como justificam respostas ao longo das atividades propostas.",
        ],
        "acessibilidade": [
            "Organizar momentos de apoio em duplas para favorecer compreensao e participacao.",
            "Permitir diferentes formas de registro, como topicos, frases curtas, esquema, desenho ou resposta oral mediada.",
            "Realizar retomadas coletivas dos comandos e oferecer mediacao individual conforme as necessidades observadas.",
        ],
    },
    {
        "abertura": "Iniciar retomando palavras-chave da trilha anterior e pedindo que a turma indique quais comandos em ingles ja consegue reconhecer.",
        "foco": "Apresentar a atividade da plataforma como continuidade do estudo, destacando vocabulario funcional, leitura de enunciados e organizacao das respostas.",
        "pratica": "Orientar a resolucao por etapas, com leitura guiada, consulta a exemplos e verificacao coletiva de trechos que geram duvida.",
        "fechamento": "Encerrar com um registro rapido do que cada estudante conseguiu compreender melhor e do que ainda precisa praticar.",
        "acompanhamento": [
            "Verificar se os estudantes identificam comandos, palavras-chave e pistas visuais presentes na trilha.",
            "Observar se utilizam exemplos e registros anteriores para resolver as atividades com mais autonomia.",
            "Conferir se conseguem explicar oralmente ou por escrito uma estrategia usada durante a atividade.",
        ],
        "acessibilidade": [
            "Disponibilizar banco de palavras em ingles e portugues para consulta durante a trilha.",
            "Oferecer leitura compartilhada dos enunciados e exemplos resolvidos antes da pratica individual.",
            "Permitir tempo ampliado e resposta oral mediada para estudantes com maior dificuldade de registro escrito.",
        ],
    },
    {
        "abertura": "Retomar o objetivo da trilha individual e relacionar a atividade ao uso real da lingua inglesa em situacoes simples do cotidiano.",
        "foco": "Explorar a trilha com apoio de explicacoes curtas, exemplos no quadro e comparacao entre respostas possiveis.",
        "pratica": "Propor que os estudantes resolvam a atividade, revisem suas escolhas e ajustem respostas a partir da mediacao do professor.",
        "fechamento": "Socializar estrategias de estudo que ajudaram a turma e combinar um ponto de atencao para a continuidade da sequencia.",
        "acompanhamento": [
            "Acompanhar se os estudantes retomam conhecimentos anteriores para compreender a proposta da trilha.",
            "Observar se revisam respostas e fazem ajustes a partir das orientacoes recebidas.",
            "Conferir se registram pelo menos uma estrategia de estudo utilizada na aula.",
        ],
        "acessibilidade": [
            "Apresentar exemplos curtos com destaque visual para palavras-chave e estruturas em ingles.",
            "Organizar pares de apoio para leitura de comandos e conferencia das respostas.",
            "Oferecer roteiro em topicos para orientar inicio, desenvolvimento e fechamento da atividade.",
        ],
    },
    {
        "abertura": "Comecar com uma breve conversa sobre dificuldades encontradas na trilha anterior e selecionar uma delas para resolver coletivamente.",
        "foco": "Conduzir a nova etapa da trilha valorizando leitura atenta, reconhecimento de vocabulario e uso de estrategias de eliminacao quando houver alternativas.",
        "pratica": "Circular pela sala, mediar duvidas individuais e pedir que os estudantes justifiquem algumas escolhas com base no enunciado.",
        "fechamento": "Retomar os erros mais frequentes sem expor estudantes e registrar no quadro orientacoes para evitar esses erros na proxima aula.",
        "acompanhamento": [
            "Verificar se os estudantes leem os enunciados antes de responder e localizam informacoes importantes.",
            "Observar se justificam escolhas usando pistas do texto, imagens ou exemplos da atividade.",
            "Acompanhar se corrigem respostas apos a mediacao e compreendem o motivo do ajuste.",
        ],
        "acessibilidade": [
            "Dividir comandos longos em frases curtas e destacar verbos de acao em ingles e portugues.",
            "Permitir que estudantes verbalizem o raciocinio antes de registrar a resposta.",
            "Retomar coletivamente os erros frequentes com exemplos simples e sem exposicao individual.",
        ],
    },
    {
        "abertura": "Apresentar a trilha do dia como momento de consolidacao e pedir que os estudantes identifiquem o que ja conseguem fazer com mais seguranca.",
        "foco": "Relacionar a atividade a vocabulario, leitura e producao em lingua inglesa, reforcando que a plataforma deve ser acompanhada por registro e reflexao.",
        "pratica": "Orientar a realizacao da atividade com pausas para conferencia, comparacao de respostas e retomada de comandos essenciais.",
        "fechamento": "Finalizar com autoavaliacao breve, em que os estudantes indicam um avanco e uma necessidade de estudo.",
        "acompanhamento": [
            "Conferir se os estudantes registram avancos e dificuldades observados durante a trilha.",
            "Observar se utilizam o vocabulario de apoio para interpretar comandos e alternativas.",
            "Verificar se participam das pausas de conferencia e ajustam as respostas quando necessario.",
        ],
        "acessibilidade": [
            "Disponibilizar quadro de autoavaliacao com opcoes simples de avanco e dificuldade.",
            "Oferecer apoio visual com palavras-chave, setas e exemplos de resposta.",
            "Permitir registros em lista curta, esquema ou fala mediada conforme a necessidade do estudante.",
        ],
    },
]

FOCOS_TRILHA = [
    ("comandos em inglês", "comandos simples em inglês", "leitura dos comandos da plataforma"),
    ("vocabulário essencial", "palavras-chave em inglês e português", "uso de banco de palavras"),
    ("leitura de enunciados", "pistas do enunciado e das imagens", "marcação de informações importantes"),
    ("autonomia de estudo", "organização da rotina individual de estudo", "registro de avanços e dúvidas"),
    ("revisão de respostas", "conferência e ajuste das respostas", "comparação entre alternativas"),
    ("compreensão oral e escrita", "relação entre escuta, leitura e registro", "identificação de informações explícitas"),
    ("produção curta em inglês", "frases curtas com apoio de modelos", "uso de estruturas já estudadas"),
    ("estratégias de plataforma", "navegação, atenção aos feedbacks e retomada", "uso pedagógico das devolutivas"),
    ("participação em duplas", "apoio entre pares e explicação do raciocínio", "troca orientada de estratégias"),
    ("recuperação de dificuldades", "retomada de pontos fragilizados", "mediação individual e coletiva"),
    ("consolidação da sequência", "ligação com aulas anteriores do bimestre", "síntese dos aprendizados"),
    ("leitura visual", "relação entre imagem, ícone e palavra", "uso de pistas visuais para compreender"),
    ("fluência gradual", "repetição funcional e segurança na resposta", "prática progressiva sem pressa"),
    ("autoavaliação", "percepção do que já foi aprendido", "planejamento do próximo passo de estudo"),
]


def _variacao_trilha(indice: int, ano: int) -> dict[str, object]:
    foco, detalhe, pratica = FOCOS_TRILHA[(indice + max(0, ano - 7)) % len(FOCOS_TRILHA)]
    return {
        "abertura": (
            f"Iniciar retomando a continuidade da Trilha de aprendizagem individual e relacionando o trabalho do dia a {foco}. "
            "Convidar os estudantes a localizar registros anteriores e prever que tipo de estratégia pode ajudar na atividade."
        ),
        "foco": (
            f"Desenvolver a trilha de forma progressiva, com explicação dialogada sobre {detalhe}. "
            "Apresentar um exemplo rápido, destacar palavras-chave e orientar a turma a acompanhar a tarefa com registro no caderno."
        ),
        "pratica": (
            f"Propor a realização da atividade com foco em {pratica}. "
            "Circular pela sala, mediar dúvidas, pedir justificativas curtas e incentivar que os estudantes revisem respostas antes de finalizar."
        ),
        "fechamento": (
            f"Encerrar retomando o que a turma conseguiu avançar em {foco}. "
            "Registrar uma síntese breve e um ponto de atenção para a próxima atividade da sequência."
        ),
        "acompanhamento": [
            f"Verificar se os estudantes compreendem o foco da trilha relacionado a {foco}.",
            f"Conferir se os registros apresentam clareza e retomam {detalhe} trabalhado na aula.",
            "Observar a participação, as justificativas e os ajustes realizados durante a resolução das atividades.",
        ],
        "acessibilidade": [
            "Organizar momentos de apoio em duplas para leitura dos comandos e conferencia das respostas.",
            "Permitir diferentes formas de registro, como topicos, frases curtas, esquema, desenho ou resposta oral mediada.",
            f"Disponibilizar palavras-chave e exemplo visual para apoiar estudantes com dificuldade em {foco}.",
        ],
    }


def _texto_pdf(caminho_pdf: Path) -> str:
    try:
        with pdfplumber.open(caminho_pdf) as pdf:
            return "\n".join((pagina.extract_text() or "") for pagina in pdf.pages[:2])
    except Exception:
        return ""


def _eh_pdf_trilha(caminho_pdf: Path) -> bool:
    texto = re.sub(r"\s+", " ", _texto_pdf(caminho_pdf))
    return "Trilha de aprendizagem individual" in texto


def _numeros_aulas_docx(doc: Document) -> list[int]:
    numeros = []
    for paragrafo in doc.paragraphs:
        match = re.match(r"^\s*AULA\s+(\d{1,2})\s*[-–—]", paragrafo.text, flags=re.I)
        if match:
            numeros.append(int(match.group(1)))
    return numeros


def _contar_trilhas_docx(doc: Document) -> int:
    total = 0
    for paragrafo in doc.paragraphs:
        texto = re.sub(r"\s+", " ", paragrafo.text).strip().lower()
        if re.match(r"^aula\s+\d{1,2}\s*[-–—]\s*trilha de aprendizagem individual$", texto):
            total += 1
    return total


def _remover_blocos_trilha(doc: Document) -> int:
    paragrafos = list(doc.paragraphs)
    indices_remover: set[int] = set()
    removendo = False
    removidos = 0

    for indice, paragrafo in enumerate(paragrafos):
        texto = re.sub(r"\s+", " ", paragrafo.text).strip().lower()
        if re.match(r"^aula\s+\d{1,2}\s*[-–—]\s*trilha de aprendizagem individual$", texto):
            removendo = True
            removidos += 1
        elif removendo and re.match(r"^aula\s+\d{1,2}\s*[-–—]", texto):
            removendo = False

        if removendo:
            indices_remover.add(indice)

    for indice in sorted(indices_remover, reverse=True):
        elemento = paragrafos[indice]._element
        elemento.getparent().remove(elemento)
    return removidos


def _adicionar_paragrafo(doc: Document, texto: str) -> None:
    doc.add_paragraph(texto)


def _adicionar_aula_trilha(doc: Document, numero: int, variacao: dict[str, object]) -> None:
    _adicionar_paragrafo(doc, "")
    _adicionar_paragrafo(doc, f"AULA {numero} — Trilha de aprendizagem individual")
    _adicionar_paragrafo(doc, "Metodologia")
    _adicionar_paragrafo(
        doc,
        f"Para começar: {variacao['abertura']}",
    )
    _adicionar_paragrafo(
        doc,
        f"Foco no conteúdo: {variacao['foco']}",
    )
    _adicionar_paragrafo(
        doc,
        f"Na prática: {variacao['pratica']}",
    )
    _adicionar_paragrafo(
        doc,
        f"Encerramento: {variacao['fechamento']}",
    )
    _adicionar_paragrafo(doc, "Acompanhamento da aprendizagem")
    for item in variacao["acompanhamento"]:
        _adicionar_paragrafo(doc, f"☑ {item}")
    _adicionar_paragrafo(doc, "Acessibilidade")
    for item in variacao["acessibilidade"]:
        _adicionar_paragrafo(doc, f"☑ {item}")


def atualizar_docx_trilhas(pasta: Path) -> dict[str, int]:
    docx = next((c for c in sorted(pasta.glob("Metodologias_Lingua_Inglesa*.docx")) if not c.name.startswith("~$")), None)
    if not docx:
        raise FileNotFoundError(f"Sem DOCX de metodologia em: {pasta}")

    pdfs_trilha = [pdf for pdf in sorted(pasta.glob("*.pdf")) if _eh_pdf_trilha(pdf)]
    if not pdfs_trilha:
        return {"trilhas_pdf": 0, "trilhas_ja_docx": 0, "trilhas_atualizadas": 0}

    doc = Document(docx)
    trilhas_existentes = _contar_trilhas_docx(doc)
    _remover_blocos_trilha(doc)
    faltantes = len(pdfs_trilha)
    proximo_numero = (max(_numeros_aulas_docx(doc)) if _numeros_aulas_docx(doc) else 0) + 1
    match_ano = re.search(r"(\d+)", pasta.name)
    ano = int(match_ano.group(1)) if match_ano else 0

    for indice in range(faltantes):
        variacao = _variacao_trilha(indice, ano)
        _adicionar_aula_trilha(doc, proximo_numero + indice, variacao)

    doc.save(docx)

    return {
        "trilhas_pdf": len(pdfs_trilha),
        "trilhas_ja_docx": trilhas_existentes,
        "trilhas_atualizadas": faltantes,
    }


def main() -> None:
    print("Atualizacao dos DOCX de Lingua Inglesa - trilhas de aprendizagem")
    for pasta in sorted(p for p in BASE_AF_3B.iterdir() if p.is_dir()):
        resultado = atualizar_docx_trilhas(pasta)
        print(f"\nPasta: {pasta}")
        for chave, valor in resultado.items():
            print(f"  {chave}: {valor}")


if __name__ == "__main__":
    main()
