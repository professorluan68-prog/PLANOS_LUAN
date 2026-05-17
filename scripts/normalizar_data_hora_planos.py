from __future__ import annotations

import re
import shutil
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(r"D:\PLANOS DE JUNHO")
REPORT_PATH = ROOT / "RELATORIO_PLANOS_SEM_DATA_HORA.docx"
BACKUP_ROOT = ROOT / f"_BACKUP_DATA_HORA_{datetime.now():%Y%m%d_%H%M%S}"

FONT_NAME = "Arial"
FONT_SIZE = Pt(10)
RED = RGBColor(0xEE, 0x00, 0x00)

DATE_RE = re.compile(r"\b(\d{1,2}/\d{1,2})\b")
TIME_RE = re.compile(r"\b(\d{1,2})h(\d{2})?\b", re.I)
TIME_LABEL_RE = re.compile(
    r"(\d{1,2}h(?:\d{2})?)\s*\(?\s*"
    r"((?:\d+|[1-6])[ªº°o]?\s*(?:e\s*(?:\d+|[1-6])[ªº°o]?)?\s*aulas?)\s*\)?",
    re.I,
)
LABEL_TIME_RE = re.compile(
    r"((?:\d+|[1-6])[ªº°o]?\s*(?:e\s*(?:\d+|[1-6])[ªº°o]?)?\s*aulas?)"
    r"\s*[|/-]?\s*(\d{1,2}h(?:\d{2})?(?:\s*[–-]\s*\d{1,2}h(?:\d{2})?)?)",
    re.I,
)

INTERVALOS_AULAS_DUPLAS = {
    ("7h", "1ª e 2ª aula"): "7h – 8h40",
    ("7h50", "2ª e 3ª aula"): "7h50 – 9h50",
    ("8h40", "3ª e 4ª aula"): "8h40 – 10h40",
    ("9h50", "4ª e 5ª aula"): "9h50 – 11h30",
    ("10h40", "5ª e 6ª aula"): "10h40 – 12h20",
    ("13h", "1ª e 2ª aula"): "13h – 14h40",
    ("13h50", "2ª e 3ª aula"): "13h50 – 15h50",
    ("14h40", "3ª e 4ª aula"): "14h40 – 16h40",
    ("15h50", "4ª e 5ª aula"): "15h50 – 17h30",
    ("16h40", "5ª e 6ª aula"): "16h40 – 18h20",
    ("19h", "1ª e 2ª aula"): "19h – 20h30",
    ("19h45", "2ª e 3ª aula"): "19h45 – 21h30",
    ("20h30", "3ª e 4ª aula"): "20h30 – 22h15",
    ("21h30", "4ª e 5ª aula"): "21h30 – 23h",
}


def _norm_space(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _looks_like_schedule(table) -> bool:
    if not table.rows:
        return False
    text = " ".join(cell.text for cell in table.rows[0].cells).upper()
    return "AULA" in text and "APRENDIZAGEM" in text and "DESENVOLVIMENTO" in text


def _header_info(table) -> dict[str, str]:
    info = {"professor": "", "disciplina": "", "turma": "", "semana": "", "aulas_semana": ""}
    if len(table.rows) < 4:
        return info
    row_data = table.rows[2].cells
    if len(row_data) >= 9:
        info["professor"] = _norm_space(row_data[2].text)
        info["disciplina"] = _norm_space(row_data[3].text)
        info["turma"] = _norm_space(row_data[6].text)
    row_week = table.rows[3].cells
    if len(row_week) >= 4:
        info["semana"] = _norm_space(row_week[1].text)
        info["aulas_semana"] = _norm_space(row_week[3].text)
    return info


def _format_time(text: str) -> str:
    text = _norm_space(str(text or "").replace("-", "–"))

    def repl(match):
        hour = str(int(match.group(1)))
        minute = match.group(2) or ""
        return f"{hour}h{minute}"

    return TIME_RE.sub(repl, text)


def _format_label(text: str) -> str:
    text = _norm_space(text)
    text = re.sub(r"\b([1-6])\s*[º°o]\b", r"\1ª", text, flags=re.I)
    text = re.sub(r"\baulas\b", "aula", text, flags=re.I)
    return text


def _complete_time(label: str, time_text: str) -> str:
    label = _format_label(label)
    time_text = _format_time(time_text)
    if "–" in time_text or " e " not in label.lower():
        return time_text
    return INTERVALOS_AULAS_DUPLAS.get((time_text, label), time_text)


def _split_date_rest(text: str) -> tuple[str, str]:
    normalized = text.replace("\n", " | ")
    normalized = re.sub(r"\s*\|\s*", " | ", normalized)
    normalized = _norm_space(normalized)
    match = DATE_RE.search(normalized)
    if not match:
        return "", normalized
    date = match.group(1)
    rest = (normalized[: match.start()] + " " + normalized[match.end() :]).strip(" |")
    return date, _norm_space(rest)


def _normalize_date_time_text(text: str) -> tuple[list[str], list[str]]:
    raw = _norm_space(str(text or "").replace("\n", " | "))
    if not raw:
        return [], ["sem data/hora"]

    date, rest = _split_date_rest(raw)
    issues = []
    if not date:
        issues.append("sem data")

    parts = [p.strip() for p in re.split(r"\s*\|\s*", rest) if p.strip()]
    lines: list[str] = [date] if date else []

    if len(parts) >= 2 and "aula" in parts[0].lower() and TIME_RE.search(parts[1]):
        label = _format_label(parts[0])
        lines.append(label)
        lines.append(_complete_time(label, parts[1]))
    elif len(parts) >= 2 and TIME_RE.search(parts[0]) and "aula" in parts[1].lower():
        label = _format_label(parts[1])
        lines.append(label)
        lines.append(_complete_time(label, parts[0]))
    else:
        remaining = " ".join(parts) if parts else rest
        consumed_spans: list[tuple[int, int]] = []
        pairs: list[tuple[str, str]] = []

        for match in TIME_LABEL_RE.finditer(remaining):
            label = _format_label(match.group(2))
            pairs.append((label, _complete_time(label, match.group(1))))
            consumed_spans.append(match.span())

        if not pairs:
            for match in LABEL_TIME_RE.finditer(remaining):
                label = _format_label(match.group(1))
                pairs.append((label, _complete_time(label, match.group(2))))
                consumed_spans.append(match.span())

        if pairs:
            for label, time in pairs:
                lines.extend([label, time])
        else:
            if remaining:
                lines.extend([_format_time(_format_label(p)) for p in parts or [remaining]])

    has_time = any(TIME_RE.search(line) for line in lines)
    has_label = any("aula" in line.lower() for line in lines)
    if not has_time:
        issues.append("sem horario")
    if date and not has_label and len(lines) == 1:
        issues.append("sem horario/aula")

    # Remove duplicadas preservando ordem.
    cleaned: list[str] = []
    for line in lines:
        line = _norm_space(line)
        if line and (not cleaned or cleaned[-1] != line):
            cleaned.append(line)
    return cleaned, issues


def _clear_cell(cell) -> None:
    cell.text = ""
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._p
        p.getparent().remove(p)


def _write_date_time_cell(cell, lines: list[str]) -> None:
    _clear_cell(cell)
    if not lines:
        return
    first = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for idx, line in enumerate(lines):
        paragraph = first if idx == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.font.color.rgb = RED
        run.bold = False


def _add_report_table(doc, title: str, rows: list[dict[str, str]]) -> None:
    doc.add_heading(title, level=1)
    if not rows:
        doc.add_paragraph("Nenhum item encontrado.")
        return
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Professor", "Arquivo", "Disciplina", "Turma", "Semana/Bloco", "Problema"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row.get("professor", "")
        cells[1].text = row.get("arquivo", "")
        cells[2].text = row.get("disciplina", "")
        cells[3].text = row.get("turma", "")
        cells[4].text = row.get("semana", "")
        cells[5].text = row.get("problema", "")


def _make_report(missing_rows: list[dict[str, str]], summary: dict[str, int]) -> None:
    doc = Document()
    doc.add_heading("Relatorio de planos sem data/hora", level=0)
    doc.add_paragraph(
        "Este relatorio lista os planos ou blocos em que nao foi possivel identificar data e horario "
        "na primeira coluna da tabela de aulas. Linhas reservas vazias do modelo foram ignoradas."
    )
    doc.add_paragraph(f"Planos analisados: {summary['plans']}")
    doc.add_paragraph(f"Planos atualizados: {summary['updated']}")
    doc.add_paragraph(f"Celulas de data/hora formatadas: {summary['cells']}")
    doc.add_paragraph(f"Ocorrencias com pendencia: {len(missing_rows)}")
    _add_report_table(doc, "Pendencias encontradas", missing_rows)
    doc.save(REPORT_PATH)


def normalize_all() -> None:
    missing_rows: list[dict[str, str]] = []
    summary = defaultdict(int)
    BACKUP_ROOT.mkdir(parents=True, exist_ok=True)

    folders = [p for p in ROOT.iterdir() if p.is_dir() and not p.name.startswith("_")]
    for folder in sorted(folders):
        for path in sorted(folder.rglob("*.docx")):
            summary["plans"] += 1
            rel_path = path.relative_to(ROOT)
            backup_path = BACKUP_ROOT / rel_path
            backup_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(path, backup_path)

            doc = Document(path)
            changed = False
            doc_issues: list[dict[str, str]] = []
            tables = list(doc.tables)
            for idx, table in enumerate(tables):
                if not _looks_like_schedule(table):
                    continue

                header = _header_info(tables[idx - 1]) if idx > 0 else {}
                if not header.get("professor"):
                    header = {"professor": folder.name, "disciplina": "", "turma": "", "semana": "", "aulas_semana": ""}

                nonempty_cells = []
                for row_index, row in enumerate(table.rows[1:], start=2):
                    cell = row.cells[0]
                    original = cell.text.strip()
                    if not original:
                        continue
                    nonempty_cells.append(original)
                    lines, issues = _normalize_date_time_text(original)
                    if lines:
                        _write_date_time_cell(cell, lines)
                        summary["cells"] += 1
                        changed = True
                    relevant_issues = [issue for issue in issues if issue != "sem data/hora"]
                    if relevant_issues:
                        doc_issues.append(
                            {
                                "professor": header.get("professor") or folder.name,
                                "arquivo": path.name,
                                "disciplina": header.get("disciplina", ""),
                                "turma": header.get("turma", ""),
                                "semana": header.get("semana", f"Tabela {idx + 1}"),
                                "problema": f"Linha {row_index}: {', '.join(relevant_issues)} ({original})",
                            }
                        )

                if not nonempty_cells:
                    doc_issues.append(
                        {
                            "professor": header.get("professor") or folder.name,
                            "arquivo": path.name,
                            "disciplina": header.get("disciplina", ""),
                            "turma": header.get("turma", ""),
                            "semana": header.get("semana", f"Tabela {idx + 1}"),
                            "problema": "Bloco sem data/hora preenchida.",
                        }
                    )

            if changed:
                doc.save(path)
                summary["updated"] += 1
            missing_rows.extend(doc_issues)

    _make_report(missing_rows, summary)
    print(f"Backup: {BACKUP_ROOT}")
    print(f"Relatorio: {REPORT_PATH}")
    print(f"Planos analisados: {summary['plans']}")
    print(f"Planos atualizados: {summary['updated']}")
    print(f"Celulas formatadas: {summary['cells']}")
    print(f"Pendencias: {len(missing_rows)}")


if __name__ == "__main__":
    normalize_all()
