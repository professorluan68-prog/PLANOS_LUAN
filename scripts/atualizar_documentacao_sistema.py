from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
CORE_DIR = BASE_DIR / "core"
LIB_DIR = CORE_DIR / "lib"
TESTS_DIR = BASE_DIR / "tests"
AE_DIR = BASE_DIR / "assets" / "ae_priorizado"

DOC_ESTRUTURA = BASE_DIR / "DOCUMENTACAO_ESTRUTURA_CORE_PLANOS_LUAN.docx"
TXT_ESTRUTURA = BASE_DIR / "DOCUMENTACAO_ESTRUTURA_CORE_PLANOS_LUAN_extracted.txt"
DOC_SISTEMA = BASE_DIR / "DOCUMENTACAO_SISTEMA_PLANOS_LUAN.docx"
TXT_SISTEMA = BASE_DIR / "DOCUMENTACAO_SISTEMA_PLANOS_LUAN_extracted.txt"


ORDEM_CORE = [
    "__init__.py",
    "ae_priorizado.py",
    "avaliacao.py",
    "base_conhecimento.py",
    "calendario.py",
    "cdp",
    "cdp_em_docx.py",
    "cdp_legacy.py",
    "constantes.py",
    "database.py",
    "disciplinas.py",
    "eja",
    "helpers.py",
    "ia.py",
    "inteligencia_local.py",
    "lib",
    "lote.py",
    "metodologia_texto.py",
    "modelos_docx.py",
    "normalizacao.py",
    "orientacao_estudos_metodologia.py",
    "orientacao_estudos_objetivos.py",
    "planos_luan.db",
    "professores_planos.py",
    "projeto_vida_escopo.py",
    "prompts_por_disciplina.py",
    "qualidade_metodologica.py",
    "redacao_leitura_metodologia.py",
    "referencias_metodologia.py",
    "revisao_final.py",
    "validador_plano.py",
]

ORDEM_LIB = [
    "acessibilidade.py",
    "acessibilidade_perfis.py",
    "acompanhamento.py",
    "acompanhamento_perfis.py",
    "aprofundamento.py",
    "classificador.py",
    "extrator_blocos_pedagogicos.py",
    "extrator_pdf.py",
    "extrator_titulo.py",
    "gerador_colunas_pedagogicas.py",
    "higienizador_pedagogico.py",
    "metodologia.py",
    "modalidades.py",
    "planejamento_colunas.py",
    "progressao.py",
    "tecnicas.py",
]

DESCRICOES_CORE = {
    "__init__.py": ("Arquivo", "Marca a pasta core como pacote Python e habilita os imports internos.", "Tecnico"),
    "ae_priorizado.py": (
        "Arquivo",
        "Aplica AE priorizado quando o contexto da disciplina, turma e bimestre possui base estruturada em assets/ae_priorizado.",
        "Integracao curricular",
    ),
    "avaliacao.py": ("Arquivo", "Gera acompanhamento e acessibilidade dinamicos de forma compatível com o restante do sistema.", "Pedagogico"),
    "base_conhecimento.py": ("Arquivo", "Guarda repertorios e padroes disciplinares usados pelos motores locais.", "Referencia"),
    "calendario.py": ("Arquivo", "Calcula feriados, dias sem aula e apoio de agenda escolar.", "Estrutural"),
    "cdp": ("Pacote", "Concentra a camada atual do fluxo CDP, hoje com destaque para gerador_cdp.py.", "Nucleo CDP"),
    "cdp_em_docx.py": ("Arquivo", "Ajusta documentos Word no fluxo CDP e faz acabamento final em cenarios especificos.", "Ajuste final"),
    "cdp_legacy.py": ("Arquivo", "Mantem a base legada ainda reutilizada por imports e compatibilidade do CDP.", "Compatibilidade"),
    "constantes.py": ("Arquivo", "Centraliza constantes de apoio compartilhadas entre modulos.", "Infraestrutura"),
    "database.py": ("Arquivo", "Controla SQLite, cadastro de professores, historico e memoria de geracao.", "Banco de dados"),
    "disciplinas.py": ("Arquivo", "Define disciplinas, modalidades e regras de comportamento por contexto.", "Estrutural"),
    "eja": ("Pacote", "Agrupa adaptacoes especificas da modalidade EJA.", "Especializacao"),
    "helpers.py": ("Arquivo", "Reune utilitarios de horarios, ordenacao de uploads, relatorios e apoio geral ao app.", "Apoio"),
    "ia.py": (
        "Arquivo",
        "Integra OpenAI e Gemini em modo hibrido, usando rascunho local do sistema como base para refinamento por IA.",
        "Motor IA",
    ),
    "inteligencia_local.py": ("Arquivo", "Mantem heuristicas e geracao local sem dependencia de IA externa.", "Motor local"),
    "lib": ("Pacote", "Agrupa extratores, classificadores, motores metodologicos, higienizacao e perfis pedagogicos.", "Base modular"),
    "lote.py": (
        "Arquivo",
        "Motor central do fluxo PDF, com extracao, cache sidecar JSON, auditoria final, IA opcional e fallback local.",
        "Nucleo central",
    ),
    "metodologia_texto.py": ("Arquivo", "Faz polimento linguistico fino de metodologias ja montadas.", "Apoio textual"),
    "modelos_docx.py": ("Arquivo", "Escolhe template central por contexto, template_id, escola e disciplina.", "Modelos centrais"),
    "normalizacao.py": ("Arquivo", "Concentra normalizacoes textuais reutilizadas por modulos diferentes.", "Apoio"),
    "orientacao_estudos_metodologia.py": ("Arquivo", "Gera metodologia especializada para Orientacao de Estudos.", "Especializado"),
    "orientacao_estudos_objetivos.py": ("Arquivo", "Centraliza objetivos e textos de apoio da Orientacao de Estudos.", "Especializado"),
    "planos_luan.db": ("Artefato tecnico", "Arquivo interno presente em core, sem ser o banco principal utilizado em producao.", "Tecnico"),
    "professores_planos.py": ("Arquivo", "Le e padroniza modelos DOCX dos professores para reaproveitamento no sistema.", "Modelos locais"),
    "projeto_vida_escopo.py": ("Arquivo", "Interpreta o escopo-sequencia de Projeto de Vida por turma, bimestre e aula.", "Especializado"),
    "prompts_por_disciplina.py": ("Arquivo", "Mantem instrucoes e reforcos de prompt por disciplina para a camada de IA.", "Referencia IA"),
    "qualidade_metodologica.py": ("Arquivo", "Revisa, sanitiza e naturaliza metodologia, inclusive correcoes por disciplina e combate a mojibake.", "Controle de qualidade"),
    "redacao_leitura_metodologia.py": ("Arquivo", "Gera estruturas metodologicas autorais para Redacao e Leitura.", "Especializado"),
    "referencias_metodologia.py": ("Arquivo", "Carrega referencias em texto usadas como memoria pedagogica do projeto.", "Memoria pedagogica"),
    "revisao_final.py": ("Arquivo", "Aplica auditoria final, confidence_score, avisos_validacao e grava cache JSON por PDF.", "Auditoria final"),
    "validador_plano.py": ("Arquivo", "Valida tema, aprendizagem, metodologia, acompanhamento e acessibilidade antes da entrega.", "Controle de qualidade"),
}

DESCRICOES_LIB = {
    "acessibilidade.py": "Compositor principal de acessibilidade por camadas.",
    "acessibilidade_perfis.py": "Regras especificas de acessibilidade por disciplina, tipo de aula e tema.",
    "acompanhamento.py": "Compositor principal do acompanhamento da aprendizagem.",
    "acompanhamento_perfis.py": "Regras especificas de acompanhamento por disciplina, perfil e tema.",
    "aprofundamento.py": "Apoia leitura de pistas de aprofundamento, continuidade ou extensao.",
    "classificador.py": "Classifica perfil da disciplina, tipo de aula, nivel de ensino e contexto pedagogico.",
    "extrator_blocos_pedagogicos.py": "Extrai blocos pedagogicos menores diretamente do texto do PDF.",
    "extrator_pdf.py": "Faz leitura, limpeza e limite de texto dos PDFs.",
    "extrator_titulo.py": "Normaliza e limpa titulos das aulas e materiais digitais.",
    "gerador_colunas_pedagogicas.py": "Motor modular para aprendizagem, acompanhamento e acessibilidade.",
    "higienizador_pedagogico.py": "Remove contaminacao tematica e corrige recursos incoerentes com a aula.",
    "metodologia.py": "Grande motor das etapas metodologicas por disciplina e tipo de aula.",
    "modalidades.py": "Apoia variacoes por modalidade, incluindo EJA e cenarios especiais.",
    "planejamento_colunas.py": "Mantem compatibilidade e geracao auxiliar das colunas pedagogicas.",
    "progressao.py": "Varia verbos, progressao de linguagem e nao repeticao entre aulas.",
    "tecnicas.py": "Concentra tecnicas pedagogicas e pistas de uso contextual.",
}


def entrada_ordenada(nomes: Iterable[str], ordem: list[str]) -> list[str]:
    mapa = {nome: idx for idx, nome in enumerate(ordem)}
    return sorted(nomes, key=lambda nome: (mapa.get(nome, 9999), nome.lower()))


def listar_core() -> list[str]:
    nomes = []
    for item in CORE_DIR.iterdir():
        if item.name == "__pycache__":
            continue
        nomes.append(item.name)
    return entrada_ordenada(nomes, ORDEM_CORE)


def listar_lib() -> list[str]:
    nomes = []
    for item in LIB_DIR.iterdir():
        if item.name in {"__init__.py", "__pycache__"}:
            continue
        nomes.append(item.name)
    return entrada_ordenada(nomes, ORDEM_LIB)


def contar_testes() -> int:
    return len(list(TESTS_DIR.glob("test_*.py")))


def listar_bases_ae() -> list[str]:
    return sorted(p.name for p in AE_DIR.glob("*.json"))


def snapshot() -> dict:
    agora = datetime.now()
    core_atual = listar_core()
    lib_atual = listar_lib()
    bases_ae = listar_bases_ae()
    return {
        "agora": agora,
        "data_extenso": agora.strftime("%d/%m/%Y às %H:%M"),
        "data_curta": agora.strftime("%d/%m/%Y"),
        "ano": agora.strftime("%Y"),
        "core_itens": core_atual,
        "lib_itens": lib_atual,
        "qtd_core": len([nome for nome in core_atual if nome.endswith(".py")]),
        "qtd_lib": len(lib_atual),
        "qtd_testes": contar_testes(),
        "bases_ae": bases_ae,
        "qtd_bases_ae": len(bases_ae),
    }


def bloco(tipo: str, **dados) -> dict:
    return {"type": tipo, **dados}


def gerar_blocos_estrutura(info: dict) -> list[dict]:
    tabela_core = []
    for nome in info["core_itens"]:
        tipo, papel, peso = DESCRICOES_CORE.get(nome, ("Arquivo", "Modulo ainda nao classificado nesta documentacao automatica.", "A revisar"))
        etiqueta = f"{nome}/" if (CORE_DIR / nome).is_dir() and not nome.endswith("/") else nome
        tabela_core.append([etiqueta, tipo, papel, peso])

    tabela_lib = []
    for nome in info["lib_itens"]:
        tabela_lib.append([nome, DESCRICOES_LIB.get(nome, "Modulo de apoio ainda sem resumo especifico."), "core/lib"])

    return [
        bloco("title", text="Estrutura Atual da Pasta core do PLANOS_LUAN"),
        bloco(
            "paragraph",
            text=(
                f"Documento atualizado automaticamente em {info['data_extenso']}, em linguagem simples, "
                "com foco na estrutura real observada no codigo."
            ),
        ),
        bloco(
            "paragraph",
            text=(
                "Objetivo: registrar como a pasta core esta organizada hoje, quais modulos ficaram mais importantes "
                "depois das evolucoes recentes e onde entram as camadas de IA, auditoria, higienizacao e CDP."
            ),
        ),
        bloco("heading", level=1, text="Visao geral"),
        bloco(
            "paragraph",
            text=(
                "A pasta core continua sendo o centro da regra de negocio do PLANOS_LUAN, mas hoje ela esta mais distribuida "
                "e mais protegida por camadas de revisao do que estava no inicio do projeto. O fluxo PDF ainda nasce em core/lote.py, "
                "mas agora passa por classificacao modular em core/lib, higienizacao pedagogica, auditoria final e cache sidecar por PDF."
            ),
        ),
        bloco(
            "paragraph",
            text=(
                "A camada de IA tambem mudou de papel: em vez de agir sozinha, ela pode trabalhar em modo hibrido, recebendo um rascunho local "
                "do proprio sistema como base para refinamento. Em paralelo, a pasta core ganhou modulos novos como constantes.py, normalizacao.py "
                "e revisao_final.py, enquanto core/lib ficou mais forte com higienizador_pedagogico.py e regras especializadas por disciplina."
            ),
        ),
        bloco("heading", level=1, text="Snapshot atual"),
        bloco(
            "bullets",
            items=[
                f"Arquivos Python diretos em core: {info['qtd_core']}",
                f"Modulos mapeados em core/lib: {info['qtd_lib']}",
                f"Arquivos de teste em tests/: {info['qtd_testes']}",
                f"Bases de AE priorizado em assets/ae_priorizado: {info['qtd_bases_ae']}",
                "Subpacotes principais vivos hoje: core/lib, core/cdp e core/eja.",
            ],
        ),
        bloco(
            "table",
            title="Mapa rapido dos itens da pasta core",
            headers=["Item", "Tipo", "Papel principal", "Peso no sistema"],
            rows=tabela_core,
        ),
        bloco("heading", level=1, text="Novidades relevantes em relacao a 05/06/2026"),
        bloco(
            "bullets",
            items=[
                "revisao_final.py virou a ultima auditoria da aula, adicionando confidence_score, avisos_validacao e versao do gerador.",
                "lote.py passou a salvar e reler cache sidecar .json por PDF, validado por hash do arquivo e versao do gerador.",
                "ia.py hoje trabalha com rascunho_base, isto e, o sistema primeiro monta uma base local e depois pede refinamento a OpenAI ou Gemini.",
                "core/lib ganhou mais peso com higienizador_pedagogico.py e com reforcos por perfil em acompanhamento e acessibilidade.",
                "helpers.py e planos_luan_app.py passaram a sustentar melhor a ordem real de envio dos PDFs e a memoria da ultima aula gerada.",
                "classificador.py e metodologia.py receberam classificadores e geradores especializados para Ciencias, Biologia e Historia.",
            ],
        ),
        bloco(
            "table",
            title="Mapa rapido do subpacote core/lib",
            headers=["Modulo", "Papel hoje", "Camada"],
            rows=tabela_lib,
        ),
        bloco("heading", level=1, text="Leitura final da arquitetura atual"),
        bloco(
            "paragraph",
            text=(
                "Se a pergunta for onde nascem as aulas comuns, a resposta continua sendo core/lote.py, apoiado por extracao e classificacao em core/lib. "
                "Se a pergunta for onde o texto fica mais confiavel, hoje os pontos mais importantes sao qualidade_metodologica.py, higienizador_pedagogico.py, "
                "validador_plano.py e revisao_final.py. Se a pergunta for onde a IA entra, ela ja nao substitui o motor local: ela entra para refinar um rascunho "
                "base e pode cair de volta no motor heuristico se houver erro, timeout ou indisponibilidade externa."
            ),
        ),
        bloco(
            "paragraph",
            text=(
                "No CDP, a leitura correta continua sendo em duas camadas: core/cdp como nucleo atual e cdp_legacy.py como compatibilidade ainda ativa. "
                "No restante do sistema, a principal mudanca de maturidade foi a presenca de uma trilha completa de saneamento: extracao, classificacao, "
                "geracao, higienizacao, validacao, auditoria final e cache reaproveitavel por PDF."
            ),
        ),
        bloco("heading", level=1, text="Observacao importante"),
        bloco(
            "paragraph",
            text=(
                f"Esta documentacao descreve o estado atual observado em {info['data_curta']}. Ela passou a ser regenerada por script interno "
                "para reduzir envelhecimento entre uma rodada de melhorias e outra."
            ),
        ),
    ]


def gerar_blocos_sistema(info: dict) -> list[dict]:
    tabela_entrada = [
        [
            "Leitura proposta",
            (
                "Documento escrito para quem precisa entender o sistema como ferramenta de trabalho: "
                "o que ele resolve hoje, como a informacao circula e o que mudou de forma concreta."
            ),
        ]
    ]

    tabela_componentes = [
        ["planos_luan_app.py", "Tela principal, memoria de uso, revisao centralizada e orquestracao do fluxo.", "Interface"],
        ["core/lote.py", "Motor principal do fluxo PDF, com cache por PDF, auditoria final e fallback local.", "Nucleo PDF"],
        ["core/lib/", "Base modular de extracao, classificacao, metodologia, higienizacao e perfis pedagogicos.", "Base pedagogica"],
        ["core/ia.py + core/prompts_por_disciplina.py", "Camada de IA com OpenAI/Gemini em modo hibrido a partir de rascunho local.", "IA opcional"],
        ["core/revisao_final.py", "Auditoria final da aula com confidence_score e sidecar JSON.", "Controle final"],
        ["core/database.py + planos_luan.db", "Cadastro persistente, historico e memoria da ultima aula.", "Banco de dados"],
        ["core/professores_planos.py", "Leitura e reaproveitamento de modelos DOCX dos professores.", "Modelos locais"],
        ["core/modelos_docx.py + templates/", "Escolha de templates centrais por contexto e template_id.", "Modelos centrais"],
        ["core/cdp/ + cdp_legacy.py + cdp_em_docx.py", "Fluxo CDP atual com camada de compatibilidade legada.", "Nucleo CDP"],
        ["core/ae_priorizado.py + assets/ae_priorizado/", "Aplicacao de AE priorizado quando houver base compativel.", "Integracao curricular"],
        ["docx_generator/", "Preenchimento final dos documentos Word para fluxo comum e CDP.", "Saida DOCX"],
        ["tests/", "Cobertura automatica das regras mais sensiveis do sistema.", "Qualidade"],
        ["scripts/atualizar_documentacao_sistema.py", "Regeneracao automatica das documentacoes docx e txt.", "Manutencao"],
    ]

    return [
        bloco("title", text="Documentacao Atual do Sistema PLANOS_LUAN"),
        bloco(
            "paragraph",
            text=(
                f"Documento atualizado automaticamente em {info['data_extenso']}, em linguagem simples, descrevendo o comportamento real do sistema hoje."
            ),
        ),
        bloco("table", title="Leitura proposta", headers=["Campo", "Descricao"], rows=tabela_entrada),
        bloco("heading", level=1, text="1. O que faz o sistema hoje"),
        bloco(
            "paragraph",
            text=(
                "O PLANOS_LUAN continua sendo um sistema de montagem de planos de aula em Word, mas hoje ele atua como uma linha de producao pedagogica mais completa. "
                "Ele combina banco de dados local, leitura de modelos DOCX, agenda escolar, entrada de PDFs em modos diferentes, extracao local, IA opcional, revisao humana e geracao final do documento."
            ),
        ),
        bloco(
            "paragraph",
            text=(
                "O sistema atende fluxos comuns por PDF, cenarios CDP, adaptacoes EJA em contextos especificos e AE priorizado quando existe base preparada para a combinacao de disciplina, turma e bimestre."
            ),
        ),
        bloco("heading", level=1, text="2. Como o sistema trabalha por dentro hoje"),
        bloco(
            "bullets",
            items=[
                "Contexto: o usuario escolhe professor, disciplina, turma, mes, bimestre, escola, componente curricular e, quando necessario, modalidades especiais.",
                "Modelo e agenda: o sistema tenta localizar modelo do professor, cadastro salvo ou template central; depois monta datas reais, feriados, dias sem aula e extensao do mes.",
                "Entrada de PDFs: hoje existem tres modos principais de envio — Automatico, Todos de uma vez e Um por aula.",
                "Memoria do sistema: no modo automatico, o app pode pre-selecionar PDFs a partir da ultima aula salva no historico.",
                "Processamento: o motor local extrai tema, aprendizagem e metodologia; se a IA estiver ativa, a camada externa refina um rascunho base do proprio sistema.",
                "Revisao: antes do DOCX final, o usuario revisa tema, aprendizagem, metodologia, acompanhamento e acessibilidade numa tela centralizada.",
                "Geracao final: o DOCX e montado, pode ser salvo no historico ou nao, e o sistema atualiza a memoria da ultima aula apenas quando isso for autorizado.",
            ],
        ),
        bloco(
            "table",
            title="Estruturas principais do sistema",
            headers=["Parte", "Funcao principal", "Tipo de conteudo"],
            rows=tabela_componentes,
        ),
        bloco("heading", level=1, text="3. Mudancas mais importantes desde a versao de 05/06/2026"),
        bloco(
            "bullets",
            items=[
                "A tela principal consolidou o fluxo em contexto, extracao, revisao e download, com mais destaque para revisao pedagógica antes do DOCX.",
                "O envio de PDFs passou a trabalhar claramente com os modos Automatico, Todos de uma vez e Um por aula.",
                "A memoria da ultima aula deixou de ser obrigatoria: hoje o usuario pode gerar o DOCX sem salvar no historico.",
                "A camada de IA virou refinamento sobre rascunho local, em vez de trabalhar isoladamente.",
                "Foi criada uma auditoria final por aula, com confidence_score, avisos_validacao, hash do PDF e cache JSON reaproveitavel.",
                "A higienizacao pedagogica ficou mais forte, reduzindo contaminacao entre disciplinas e trocas indevidas de recurso.",
                "A base de testes cresceu e hoje cobre fluxo principal sem lote, cache JSON, revisao final, timeout de IA e blocos pedagogicos mais recentes.",
                "Padronizacao fisica de PDFs e atualizacao automatica de caminhos no arquivo central mapa_arquivos.csv para Matematica, Ciencias, Biologia e Historia.",
                "Adicionados classificadores e metodologias autorais complexas (como linha do tempo, analise critica de fontes e debate de narrativas) para Historia, Ciencias e Biologia.",
            ],
        ),
        bloco("heading", level=1, text="4. Partes que merecem leitura especial hoje"),
        bloco(
            "paragraph",
            text=(
                "planos_luan_app.py segue sendo o grande orquestrador do uso diario. Ele concentra a experiencia de tela, a memoria de historico, "
                "as escolhas de modo de upload e a revisao centralizada."
            ),
        ),
        bloco(
            "paragraph",
            text=(
                "core/lote.py continua sendo o coracao do fluxo PDF, mas agora conversa mais fortemente com classificadores e higienizadores de core/lib, "
                "alem de revisar e gravar cache sidecar por PDF."
            ),
        ),
        bloco(
            "paragraph",
            text=(
                "core/ia.py, prompts_por_disciplina.py e referencias_metodologia.py formam a trilha de IA opcional. O desenho atual privilegia robustez: "
                "se a IA falhar, o motor local continua sustentando o plano."
            ),
        ),
        bloco(
            "paragraph",
            text=(
                "core/revisao_final.py e tests/ viraram pecas mais relevantes do que eram nos documentos antigos, porque agora o sistema nao apenas gera: "
                "ele pontua, valida, reaproveita e tenta evitar que um PDF ja processado precise ser refeito sem necessidade."
            ),
        ),
        bloco("heading", level=1, text="5. Operacao e artefatos do dia a dia"),
        bloco(
            "bullets",
            items=[
                "O banco principal usado pelo app continua na raiz: planos_luan.db.",
                "Cada PDF processado pode ganhar um sidecar .json com hash, tema, metodologia, confidence_score e avisos.",
                f"Hoje existem {info['qtd_bases_ae']} bases JSON de AE priorizado em assets/ae_priorizado, ainda concentradas no 2o bimestre.",
                f"O projeto possui {info['qtd_testes']} arquivos de teste automatizado em tests/.",
                "As versoes .txt extraidas da documentacao continuam existindo para leitura rapida fora do Word.",
            ],
        ),
        bloco("heading", level=1, text="6. Como a documentacao fica atualizada agora"),
        bloco(
            "paragraph",
            text=(
                "A documentacao passou a ser regenerada por um script proprio do projeto: scripts/atualizar_documentacao_sistema.py. "
                "Esse script atualiza os dois DOCX e tambem as versoes TXT extraidas, reduzindo o risco de a descricao envelhecer depois de novas melhorias."
            ),
        ),
        bloco("heading", level=1, text="7. Pontos fortes atuais"),
        bloco(
            "bullets",
            items=[
                "Fluxo mais guiado e menos manual do que nas versoes mais antigas.",
                "Capacidade de trabalhar com PDF comum, CDP, EJA e AE priorizado sem trocar de sistema.",
                "IA opcional com fallback local e pos-processamento forte.",
                "Revisao humana antes do documento final.",
                "Historico persistente, memoria de ultima aula e opcao de nao avancar essa memoria quando o plano ainda nao estiver ok.",
                "Cobertura de testes mais madura e mais alinhada ao comportamento real do sistema.",
            ],
        ),
        bloco("heading", level=1, text="8. Prioridades naturais de evolucao"),
        bloco(
            "bullets",
            items=[
                "Continuar modularizando planos_luan_app.py, que ainda concentra muitas responsabilidades de interface e orquestracao.",
                "Expandir a base estruturada de AE priorizado para mais contextos quando os dados estiverem prontos.",
                "Reduzir sobreposicao entre camadas novas e legadas do CDP.",
                "Fortalecer diagnosticos automaticos de modelos Word, bases externas e configuracoes locais.",
                "Manter a documentacao e a automacao diaria como artefatos oficiais de manutencao do projeto.",
            ],
        ),
        bloco("heading", level=1, text="9. Conclusao"),
        bloco(
            "paragraph",
            text=(
                "Hoje o PLANOS_LUAN deve ser entendido como uma linha de producao pedagogica com varias camadas de seguranca: "
                "cadastro, agenda, leitura de modelos, extracao, IA opcional, higienizacao, revisao humana, auditoria final e geracao Word. "
                "Ele ja nao funciona apenas como leitor de PDF com Word na saida; ele organiza memoria, contexto e qualidade antes da entrega."
            ),
        ),
    ]


def configurar_documento(doc: Document) -> None:
    secao = doc.sections[0]
    secao.top_margin = Inches(1)
    secao.bottom_margin = Inches(1)
    secao.left_margin = Inches(1)
    secao.right_margin = Inches(1)

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)

    for nome, tamanho in [("Title", 18), ("Heading 1", 14), ("Heading 2", 12)]:
        estilo = doc.styles[nome]
        estilo.font.name = "Calibri"
        estilo.font.size = Pt(tamanho)


def aplicar_bloco_docx(doc: Document, bloco_atual: dict) -> None:
    tipo = bloco_atual["type"]
    if tipo == "title":
        p = doc.add_paragraph(style="Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(bloco_atual["text"])
        run.bold = True
        return

    if tipo == "heading":
        doc.add_heading(bloco_atual["text"], level=bloco_atual.get("level", 1))
        return

    if tipo == "paragraph":
        doc.add_paragraph(bloco_atual["text"])
        return

    if tipo == "bullets":
        for item in bloco_atual["items"]:
            doc.add_paragraph(item, style="List Bullet")
        return

    if tipo == "table":
        titulo = bloco_atual.get("title", "").strip()
        if titulo:
            doc.add_paragraph(titulo)
        headers = bloco_atual["headers"]
        rows = bloco_atual["rows"]
        tabela = doc.add_table(rows=1, cols=len(headers))
        tabela.style = "Table Grid"
        for idx, header in enumerate(headers):
            tabela.rows[0].cells[idx].text = str(header)
        for row in rows:
            cells = tabela.add_row().cells
            for idx, valor in enumerate(row):
                cells[idx].text = str(valor)
        return

    raise ValueError(f"Tipo de bloco nao suportado: {tipo}")


def aplicar_bloco_txt(linhas: list[str], bloco_atual: dict) -> None:
    tipo = bloco_atual["type"]
    if tipo == "title":
        linhas.extend([bloco_atual["text"], ""])
        return

    if tipo == "heading":
        linhas.extend([bloco_atual["text"], ""])
        return

    if tipo == "paragraph":
        linhas.extend([bloco_atual["text"], ""])
        return

    if tipo == "bullets":
        for item in bloco_atual["items"]:
            linhas.append(f"- {item}")
        linhas.append("")
        return

    if tipo == "table":
        titulo = bloco_atual.get("title", "").strip()
        if titulo:
            linhas.append(f"--- Table: {titulo} ---")
        linhas.append(" | ".join(bloco_atual["headers"]))
        for row in bloco_atual["rows"]:
            linhas.append(" | ".join(str(valor) for valor in row))
        linhas.append("")
        return

    raise ValueError(f"Tipo de bloco nao suportado: {tipo}")


def salvar_docx(caminho: Path, blocos: list[dict]) -> None:
    doc = Document()
    configurar_documento(doc)
    for item in blocos:
        aplicar_bloco_docx(doc, item)
    doc.save(caminho)


def salvar_txt(caminho: Path, blocos: list[dict]) -> None:
    linhas: list[str] = []
    for item in blocos:
        aplicar_bloco_txt(linhas, item)
    caminho.write_text("\n".join(linhas).strip() + "\n", encoding="utf-8")


def main() -> None:
    info = snapshot()
    blocos_estrutura = gerar_blocos_estrutura(info)
    blocos_sistema = gerar_blocos_sistema(info)

    salvar_docx(DOC_ESTRUTURA, blocos_estrutura)
    salvar_txt(TXT_ESTRUTURA, blocos_estrutura)
    salvar_docx(DOC_SISTEMA, blocos_sistema)
    salvar_txt(TXT_SISTEMA, blocos_sistema)

    print("Documentacao atualizada com sucesso:")
    print(f"- {DOC_ESTRUTURA}")
    print(f"- {TXT_ESTRUTURA}")
    print(f"- {DOC_SISTEMA}")
    print(f"- {TXT_SISTEMA}")


if __name__ == "__main__":
    main()
