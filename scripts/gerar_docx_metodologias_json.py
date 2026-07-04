from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


COR_TITULO_PRINCIPAL = RGBColor(0x1F, 0x49, 0x7D)
COR_SUBTITULO = RGBColor(0x2E, 0x74, 0xB5)
COR_TITULO_AULA = RGBColor(0x00, 0x47, 0x70)
COR_ROTULO_ETAPA = RGBColor(0x2E, 0x74, 0xB5)
FONTE_PADRAO = "Arial"


def extrair_numero_aula(valor: str) -> int:
    match = re.search(r"\d{1,3}", str(valor or ""))
    return int(match.group(0)) if match else 0


def normalizar_checklist(item: str) -> str:
    texto = re.sub(r"\s+", " ", str(item or "")).strip()
    if not texto:
        return ""
    if not texto.startswith("☑"):
        texto = f"☑ {texto.lstrip('☑ ').strip()}"
    return texto


def carregar_aulas_json(pasta: Path) -> list[dict]:
    aulas: list[dict] = []
    for caminho_json in sorted(
        pasta.glob("AULA_*.json"),
        key=lambda caminho: extrair_numero_aula(caminho.name),
    ):
        dados = json.loads(caminho_json.read_text(encoding="utf-8"))
        dados["_arquivo_json"] = str(caminho_json)
        aulas.append(dados)
    return aulas


def detectar_subtitulo(pasta: Path) -> str:
    nome = pasta.name.upper()
    match = re.search(r"(\d+)_ANO", nome)
    serie = match.group(1) if match else ""
    if serie:
        return f"{serie}º Ano · 3º Bimestre"
    return "3º Bimestre"


def configurar_documento(doc: Document) -> None:
    secao = doc.sections[0]
    secao.top_margin = Inches(1)
    secao.bottom_margin = Inches(1)
    secao.left_margin = Inches(1)
    secao.right_margin = Inches(1)
    secao.header_distance = Inches(0.5)
    secao.footer_distance = Inches(0.5)


def configurar_paragrafo(
    paragrafo,
    *,
    alinhamento=None,
    espaco_antes: float | None = None,
    espaco_depois: float | None = None,
    espacamento_linha: float = 1.15,
) -> None:
    if alinhamento is not None:
        paragrafo.alignment = alinhamento
    formato = paragrafo.paragraph_format
    formato.line_spacing = espacamento_linha
    if espaco_antes is not None:
        formato.space_before = Pt(espaco_antes)
    if espaco_depois is not None:
        formato.space_after = Pt(espaco_depois)


def adicionar_run(
    paragrafo,
    texto: str,
    *,
    negrito: bool = False,
    tamanho: float = 10.5,
    cor: RGBColor | None = None,
) -> None:
    run = paragrafo.add_run(texto)
    run.bold = negrito
    run.font.name = FONTE_PADRAO
    run.font.size = Pt(tamanho)
    if cor is not None:
        run.font.color.rgb = cor


def adicionar_titulo_principal(doc: Document, subtitulo: str) -> None:
    paragrafo = doc.add_paragraph()
    configurar_paragrafo(paragrafo, alinhamento=WD_ALIGN_PARAGRAPH.LEFT, espaco_depois=3)
    adicionar_run(
        paragrafo,
        "Metodologias — Língua Portuguesa",
        negrito=True,
        tamanho=18,
        cor=COR_TITULO_PRINCIPAL,
    )

    paragrafo = doc.add_paragraph()
    configurar_paragrafo(paragrafo, espaco_depois=6)
    adicionar_run(paragrafo, subtitulo, tamanho=13, cor=COR_SUBTITULO)

    paragrafo = doc.add_paragraph()
    configurar_paragrafo(paragrafo, espaco_depois=8)
    adicionar_run(
        paragrafo,
        (
            "Material organizado com a metodologia de cada aula, seguida de "
            "acompanhamento da aprendizagem e acessibilidade. As técnicas "
            "citadas no material são incorporadas ao texto da aula de forma "
            "direta e objetiva."
        ),
        tamanho=10.5,
    )

    doc.add_paragraph()


def adicionar_titulo_aula(doc: Document, titulo: str) -> None:
    paragrafo = doc.add_paragraph()
    configurar_paragrafo(paragrafo, espaco_antes=8, espaco_depois=3)
    adicionar_run(
        paragrafo,
        titulo,
        negrito=True,
        tamanho=14,
        cor=COR_TITULO_AULA,
    )


def adicionar_secao(doc: Document, titulo: str, espaco_antes: float = 0) -> None:
    paragrafo = doc.add_paragraph()
    configurar_paragrafo(paragrafo, espaco_antes=espaco_antes, espaco_depois=2)
    adicionar_run(
        paragrafo,
        titulo,
        negrito=True,
        tamanho=12,
        cor=COR_TITULO_PRINCIPAL,
    )


def adicionar_etapa(doc: Document, titulo: str, texto: str) -> None:
    paragrafo = doc.add_paragraph()
    configurar_paragrafo(paragrafo, espaco_depois=0)
    adicionar_run(
        paragrafo,
        f"{titulo}: ",
        negrito=True,
        tamanho=10.5,
        cor=COR_ROTULO_ETAPA,
    )
    adicionar_run(paragrafo, texto, tamanho=10.5)


def adicionar_checklist(doc: Document, itens: list[str]) -> None:
    for item in itens[:3]:
        texto = normalizar_checklist(item)
        if not texto:
            continue
        paragrafo = doc.add_paragraph()
        configurar_paragrafo(paragrafo, espaco_depois=0)
        adicionar_run(paragrafo, texto, tamanho=10.5)


def titulo_aula(dados_aula: dict) -> str:
    material = re.sub(r"\s+", " ", str(dados_aula.get("material") or "")).strip()
    if material:
        return material
    numero = extrair_numero_aula(dados_aula.get("numero_aula") or "")
    tema = re.sub(r"\s+", " ", str(dados_aula.get("tema") or "")).strip()
    return f"AULA {numero} - {tema}".strip()


def filtrar_metodologia(metodologia: list[dict]) -> list[dict[str, str]]:
    etapas_filtradas: list[dict[str, str]] = []
    for item in metodologia or []:
        if not isinstance(item, dict):
            continue
        titulo = re.sub(r"\s+", " ", str(item.get("titulo") or "")).strip()
        texto = re.sub(r"\s+", " ", str(item.get("texto") or "")).strip()
        if not titulo or not texto:
            continue
        if re.search(r"pause\s*e\s*responda", titulo, flags=re.I):
            continue
        etapas_filtradas.append({"titulo": titulo, "texto": texto})
    return etapas_filtradas


def validar_aula(dados_aula: dict) -> None:
    metodologia = filtrar_metodologia(dados_aula.get("metodologia") or [])
    acompanhamento = [normalizar_checklist(item) for item in dados_aula.get("acompanhamento") or [] if normalizar_checklist(item)]
    acessibilidade = [normalizar_checklist(item) for item in dados_aula.get("acessibilidade") or [] if normalizar_checklist(item)]
    if len(acompanhamento) != 3:
        raise ValueError(
            f"{dados_aula.get('_arquivo_json')}: acompanhamento com {len(acompanhamento)} item(ns)."
        )
    if len(acessibilidade) != 3:
        raise ValueError(
            f"{dados_aula.get('_arquivo_json')}: acessibilidade com {len(acessibilidade)} item(ns)."
        )
    titulos = [item["titulo"] for item in metodologia]
    if "Para começar" in titulos and "Relembre" in titulos:
        raise ValueError(
            f"{dados_aula.get('_arquivo_json')}: aula com 'Para começar' e 'Relembre'."
        )


def adicionar_aula(doc: Document, dados_aula: dict) -> None:
    metodologia = filtrar_metodologia(dados_aula.get("metodologia") or [])
    acompanhamento = [normalizar_checklist(item) for item in dados_aula.get("acompanhamento") or [] if normalizar_checklist(item)]
    acessibilidade = [normalizar_checklist(item) for item in dados_aula.get("acessibilidade") or [] if normalizar_checklist(item)]

    adicionar_titulo_aula(doc, titulo_aula(dados_aula))
    adicionar_secao(doc, "Metodologia")
    for etapa in metodologia:
        adicionar_etapa(doc, etapa["titulo"], etapa["texto"])

    adicionar_secao(doc, "Acompanhamento da aprendizagem", espaco_antes=4)
    adicionar_checklist(doc, acompanhamento)

    adicionar_secao(doc, "Acessibilidade", espaco_antes=4)
    adicionar_checklist(doc, acessibilidade)
    doc.add_paragraph()


def gerar_documento(pasta: Path, saida: Path) -> None:
    aulas = carregar_aulas_json(pasta)
    if not aulas:
        raise FileNotFoundError(f"Nenhum arquivo JSON encontrado em {pasta}")

    for aula in aulas:
        validar_aula(aula)

    doc = Document()
    configurar_documento(doc)
    adicionar_titulo_principal(doc, detectar_subtitulo(pasta))
    for aula in aulas:
        adicionar_aula(doc, aula)

    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Gera um DOCX de metodologias a partir dos JSONs de uma pasta."
    )
    parser.add_argument("--pasta", required=True, help="Pasta com os arquivos AULA_*.json")
    parser.add_argument("--saida", required=True, help="Caminho completo do DOCX de saida")
    args = parser.parse_args()

    gerar_documento(Path(args.pasta), Path(args.saida))


if __name__ == "__main__":
    main()
