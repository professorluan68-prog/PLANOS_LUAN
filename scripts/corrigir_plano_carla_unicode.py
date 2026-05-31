from __future__ import annotations

from pathlib import Path

from docx import Document


PATH = Path(r"D:\PLANOS-FINALIZADOS\CARLA GANDOLF\Plano-Junho.docx")


def set_cell(cell, text: str) -> None:
    cell.text = text


NOVAS_LINHAS = [
    {
        "data_horario": "29/06\n1ª, 2ª e 3ª aula",
        "material": "PORTUGUÊS\nTema:\nFolclore brasileiro",
        "aprendizagem": "HABILIDADE:\n(EF15LP03) Localizar informações explícitas em textos de diferentes gêneros textuais.",
        "desenvolvimento": "1. Abertura (acolhimento e ativação de saberes prévios): iniciar com uma conversa simples e acolhedora sobre o tema Folclore brasileiro, valorizando os conhecimentos de vida dos alunos e registrando no quadro palavras, ideias ou exemplos citados pela turma.\n\n2. Desenvolvimento (leitura mediada e exploração do texto): apresentar leitura em voz alta, pausada e expressiva, explicando vocabulário, informações principais, sentidos e exemplos do material. Retomar oralmente trechos importantes e fazer perguntas curtas para verificar a compreensão.\n\n3. Atividade (prática orientada e registro): orientar os estudantes na realização das atividades de leitura, escrita ou interpretação no caderno, com apoio do professor, retomada coletiva dos comandos e acompanhamento mais próximo de quem apresentar dificuldade.\n\n4. Fechamento (socialização e síntese): corrigir coletivamente as respostas, convidar alguns alunos a compartilharem seus registros e finalizar com uma síntese simples do que foi estudado.",
        "acomp": "☑ Participação durante a leitura e a conversa inicial.\n☑ Compreensão do tema Folclore brasileiro nas respostas orais e escritas.\n☑ Organização das atividades de leitura, escrita e interpretação no caderno.",
        "acess": "☑ Leitura pausada, com explicação oral de palavras e trechos mais difíceis.\n☑ Registro no quadro para organizar as informações principais.\n☑ Apoio individual nas atividades de leitura, escrita e interpretação.",
    },
    {
        "data_horario": "29/06\n4ª e 5ª aula",
        "material": "MATEMÁTICA\nTema:\nReta numérica",
        "aprendizagem": "HABILIDADE:\n(EF03MA04) Estabelecer a relação entre números naturais e pontos da reta numérica para utilizá-la na ordenação dos números naturais.",
        "desenvolvimento": "1. Abertura (contextualização e conexão com a realidade): iniciar com uma situação do cotidiano relacionada ao tema Reta numérica, como contagens, medidas, compras, horários e organização de quantidades, mobilizando exemplos concretos conhecidos pelos alunos.\n\n2. Desenvolvimento (exploração guiada do conteúdo): explicar o conteúdo no quadro com exemplos simples e resolução passo a passo, mostrando como localizar números, avançar, voltar e comparar posições na reta.\n\n3. Atividade (prática com mediação docente): propor exercícios no caderno, permitir que os alunos resolvam com apoio do professor e comparar estratégias durante a correção, retomando oralmente o procedimento quando houver dúvidas.\n\n4. Fechamento (correção e sistematização): conferir os resultados coletivamente, destacar os passos mais importantes e registrar uma síntese do procedimento principal da aula.",
        "acomp": "☑ Participação na resolução das atividades propostas.\n☑ Compreensão do tema Reta numérica por meio dos cálculos, registros e explicações.\n☑ Organização dos procedimentos e participação na correção coletiva.",
        "acess": "☑ Explicação passo a passo com exemplos concretos e registros no quadro.\n☑ Retomada oral dos procedimentos antes da resolução das atividades.\n☑ Apoio individual na organização dos cálculos e respostas.",
    },
    {
        "data_horario": "30/06\n1ª e 2ª aula",
        "material": "HISTÓRIA\nTema:\nPassado e presente",
        "aprendizagem": "HABILIDADE:\n(EF04HI06) Identificar transformações ocorridas no modo de viver, trabalhar e circular das pessoas ao longo do tempo.",
        "desenvolvimento": "1. Abertura (memórias e vivências): iniciar com conversa sobre brincadeiras, objetos, meios de transporte e costumes de antigamente e de hoje, valorizando relatos da turma.\n\n2. Desenvolvimento (explicação dialogada): apresentar o conteúdo de forma simples, mostrando diferenças e permanências entre passado e presente com exemplos do cotidiano, imagens e comparações orais.\n\n3. Atividade (registro e compreensão): orientar pequenos registros, respostas orais e escritas, acompanhando a turma durante a atividade e retomando os pontos principais quando necessário.\n\n4. Fechamento (socialização e síntese): fazer correção coletiva, ouvir alguns relatos e registrar no quadro as ideias centrais estudadas.",
        "acomp": "☑ Participação nas conversas e relatos sobre experiências de vida.\n☑ Compreensão das diferenças entre passado e presente nas respostas e registros.\n☑ Relação entre o conteúdo estudado e situações conhecidas pelos educandos.",
        "acess": "☑ Uso de linguagem simples, pausada e contextualizada.\n☑ Apoio com exemplos visuais e comparações próximas da realidade dos educandos.\n☑ Possibilidade de participação oral, com mediação individual quando necessário.",
    },
    {
        "data_horario": "30/06\n3ª e 4ª aula",
        "material": "GEOGRAFIA\nTema:\nQuem administra o município?",
        "aprendizagem": "HABILIDADE:\n(EF04GE14) Identificar elementos da organização político-administrativa do Brasil, com foco no município onde se vive.",
        "desenvolvimento": "1. Abertura (leitura do espaço vivido): iniciar com conversa sobre o bairro, a cidade, serviços públicos e situações do dia a dia que dependem da organização do município.\n\n2. Desenvolvimento (explicação com exemplos concretos): apresentar o conteúdo com linguagem clara, exemplos próximos da realidade dos alunos e registros no quadro para organizar conceitos e funções da administração municipal.\n\n3. Atividade (observação e registro): propor perguntas simples, identificação de espaços e registro no caderno, acompanhando a turma durante a realização das tarefas.\n\n4. Fechamento (retomada do conceito central): socializar as respostas, corrigir coletivamente e retomar quem administra o município e como isso aparece na vida cotidiana.",
        "acomp": "☑ Participação na conversa inicial e nos exemplos apresentados.\n☑ Compreensão do tema nas atividades de observação e registro.\n☑ Relação entre o conteúdo e experiências do espaço vivido pelos educandos.",
        "acess": "☑ Explicação oral pausada com exemplos da cidade e da comunidade.\n☑ Registro no quadro para organizar conceitos e exemplos.\n☑ Apoio individual na leitura, interpretação e registro das atividades.",
    },
    {
        "data_horario": "30/06\n5ª aula",
        "material": "ARTE\nTema:\nDançando pelo mundo",
        "aprendizagem": "HABILIDADE:\n(EF04AR08) Experimentar e apreciar manifestações da dança presentes em diferentes culturas.",
        "desenvolvimento": "1. Abertura (sensibilização e repertório): iniciar com conversa sobre músicas, festas e movimentos corporais conhecidos pelos educandos, relacionando-os ao tema Dançando pelo mundo.\n\n2. Desenvolvimento (apresentação da proposta): apresentar o conteúdo com exemplos simples, explicação oral, apreciação guiada e demonstração de movimentos ou ritmos quando necessário.\n\n3. Atividade (expressão e acompanhamento): orientar produção, apreciação ou movimento, respeitando diferentes formas de participação e expressão.\n\n4. Fechamento (socialização e valorização): socializar percepções da turma e concluir com uma síntese simples sobre cultura, corpo e expressão artística.",
        "acomp": "☑ Participação na conversa, apreciação ou atividade artística.\n☑ Compreensão do tema por meio de comentários, registros ou movimentos realizados.\n☑ Envolvimento e expressão durante a atividade proposta.",
        "acess": "☑ Orientações orais claras, com demonstração simples da atividade.\n☑ Possibilidade de participação por meio de fala, desenho, movimento ou apreciação.\n☑ Apoio individual na organização da produção ou expressão das ideias.",
    },
    {
        "data_horario": "01/07\n1ª, 2ª e 3ª aula",
        "material": "MATEMÁTICA\nTema:\nAdição e subtração no dia a dia",
        "aprendizagem": "HABILIDADE:\n(EF04MA03) Resolver situações-problema simples envolvendo adição e subtração com apoio de registros e estratégias pessoais.",
        "desenvolvimento": "1. Abertura (contextualização e conexão com a realidade): iniciar com exemplos simples de compra, troco, contagem e organização de quantidades presentes no cotidiano dos educandos.\n\n2. Desenvolvimento (exploração guiada do conteúdo): explicar no quadro como organizar contas e resolver pequenas situações-problema passo a passo, com números acessíveis para a turma multisseriada.\n\n3. Atividade (prática com mediação docente): propor exercícios curtos no caderno, acompanhar a resolução e comparar estratégias com a turma durante a correção.\n\n4. Fechamento (correção e sistematização): conferir os resultados coletivamente e registrar uma síntese simples dos procedimentos usados.",
        "acomp": "☑ Participação na resolução das atividades propostas.\n☑ Compreensão do tema por meio dos cálculos, registros e explicações.\n☑ Organização dos procedimentos e participação na correção coletiva.",
        "acess": "☑ Explicação passo a passo com exemplos concretos e registros no quadro.\n☑ Retomada oral dos procedimentos antes da resolução das atividades.\n☑ Apoio individual na organização dos cálculos e respostas.",
    },
    {
        "data_horario": "01/07\n4ª e 5ª aula",
        "material": "PORTUGUÊS\nTema:\nBilhete e recado",
        "aprendizagem": "HABILIDADE:\n(EF15LP01) Compreender a função social de textos curtos do cotidiano, como bilhetes e recados.",
        "desenvolvimento": "1. Abertura (acolhimento e ativação de saberes prévios): iniciar com conversa sobre recados recebidos em casa, na escola ou no trabalho, valorizando exemplos trazidos pelos educandos.\n\n2. Desenvolvimento (leitura mediada e exploração do texto): apresentar pequenos bilhetes e recados, explicando sua finalidade, estrutura simples e informações principais.\n\n3. Atividade (prática orientada e registro): orientar leitura, interpretação e produção de recados curtos no caderno, com apoio do professor na organização das ideias.\n\n4. Fechamento (socialização e síntese): ler alguns exemplos com a turma, corrigir coletivamente e finalizar com uma síntese simples do que foi estudado.",
        "acomp": "☑ Participação durante a leitura e a conversa inicial.\n☑ Compreensão do tema Bilhete e recado nas respostas orais e escritas.\n☑ Organização das atividades de leitura, escrita e interpretação no caderno.",
        "acess": "☑ Leitura pausada, com explicação oral de palavras e trechos mais difíceis.\n☑ Registro no quadro para organizar as informações principais.\n☑ Apoio individual nas atividades de leitura, escrita e interpretação.",
    },
    {
        "data_horario": "02/07\n1ª, 2ª e 3ª aula",
        "material": "PORTUGUÊS\nTema:\nParlendas e cantigas",
        "aprendizagem": "HABILIDADE:\n(EF15LP18) Relacionar texto verbal a ritmo, oralidade e memórias culturais presentes em parlendas e cantigas.",
        "desenvolvimento": "1. Abertura (acolhimento e ativação de repertório): iniciar com conversa sobre cantigas, parlendas e brincadeiras de infância conhecidas pela turma.\n\n2. Desenvolvimento (leitura mediada e oralidade): apresentar leitura e repetição oral de parlendas e cantigas, explicando palavras, ritmo e sentido do texto.\n\n3. Atividade (prática orientada e registro): orientar leitura, cópia de trechos curtos, identificação de rimas e produção de pequenas respostas no caderno, com apoio do professor.\n\n4. Fechamento (socialização e síntese): retomar oralmente o que foi lido, ouvir a turma e registrar no quadro os pontos mais importantes.",
        "acomp": "☑ Participação durante a leitura, repetição oral e conversa inicial.\n☑ Compreensão do tema nas respostas, leituras e registros realizados.\n☑ Envolvimento nas atividades de oralidade, leitura e escrita.",
        "acess": "☑ Leitura coletiva pausada com repetição oral das parlendas e cantigas.\n☑ Apoio com registro no quadro e destaque para palavras-chave.\n☑ Possibilidade de participação oral antes do registro escrito.",
    },
    {
        "data_horario": "02/07\n4ª e 5ª aula",
        "material": "MATEMÁTICA\nTema:\nDobro e metade",
        "aprendizagem": "HABILIDADE:\n(EF04MA06A) Compreender relações simples de dobro e metade em situações do cotidiano.",
        "desenvolvimento": "1. Abertura (contextualização e conexão com a realidade): iniciar com exemplos simples de dividir e juntar quantidades em situações do dia a dia, como alimentos, objetos e dinheiro.\n\n2. Desenvolvimento (exploração guiada do conteúdo): explicar no quadro o significado de dobro e metade com desenhos, agrupamentos e exemplos concretos.\n\n3. Atividade (prática com mediação docente): propor exercícios curtos com números pequenos, acompanhar a resolução e retomar oralmente o procedimento sempre que necessário.\n\n4. Fechamento (correção e sistematização): corrigir coletivamente e registrar uma síntese simples com exemplos de dobro e metade.",
        "acomp": "☑ Participação na resolução das atividades propostas.\n☑ Compreensão do tema por meio dos registros, cálculos e explicações dadas pela turma.\n☑ Organização dos procedimentos e participação na correção coletiva.",
        "acess": "☑ Explicação passo a passo com exemplos concretos e desenhos no quadro.\n☑ Retomada oral dos procedimentos antes e durante a atividade.\n☑ Apoio individual na organização dos cálculos e respostas.",
    },
    {
        "data_horario": "03/07\n1ª e 2ª aula",
        "material": "CIÊNCIAS\nTema:\nCuidados com a água e a higiene",
        "aprendizagem": "HABILIDADE:\n(EF05CI04) Identificar usos da água e discutir cuidados simples com higiene e saúde no cotidiano.",
        "desenvolvimento": "1. Abertura (observação do cotidiano): iniciar com conversa sobre banho, limpeza da casa, preparo de alimentos e uso da água no dia a dia dos educandos.\n\n2. Desenvolvimento (explicação com exemplos e esquemas): apresentar o conteúdo com linguagem simples, exemplos concretos e registros no quadro sobre economia de água e cuidados com a higiene.\n\n3. Atividade (identificação, classificação e registro): orientar atividades curtas de identificação de atitudes corretas e incorretas, com acompanhamento próximo do professor.\n\n4. Fechamento (retomada e síntese): corrigir coletivamente as respostas e finalizar com uma síntese simples sobre saúde, higiene e uso consciente da água.",
        "acomp": "☑ Participação nas conversas e exemplos apresentados.\n☑ Compreensão do tema nas atividades de identificação, registro ou classificação.\n☑ Respostas apresentadas durante a correção coletiva e a retomada do conteúdo.",
        "acess": "☑ Explicação oral pausada com exemplos concretos e próximos da realidade dos alunos.\n☑ Esquemas simples no quadro para organizar as informações.\n☑ Apoio individual na leitura, escrita e compreensão das atividades.",
    },
    {
        "data_horario": "03/07\n3ª aula",
        "material": "ARTE\nTema:\nBrincadeiras cantadas",
        "aprendizagem": "HABILIDADE:\n(EF04AR11) Explorar formas simples de expressão corporal, ritmo e cultura presentes em brincadeiras cantadas.",
        "desenvolvimento": "1. Abertura (sensibilização e repertório): iniciar com conversa sobre músicas e brincadeiras conhecidas pela turma, valorizando memórias e experiências do grupo.\n\n2. Desenvolvimento (apresentação da proposta): apresentar a atividade com exemplos simples, ritmo marcado e explicação oral clara, permitindo observação e participação gradual.\n\n3. Atividade (expressão e acompanhamento): orientar movimento, canto, palmas ou registro simples, respeitando diferentes formas de participação.\n\n4. Fechamento (socialização e valorização): socializar como a turma participou e finalizar com uma síntese sobre ritmo, corpo e cultura.",
        "acomp": "☑ Participação na conversa, apreciação ou atividade artística.\n☑ Compreensão do tema por meio de comentários, registros ou movimentos realizados.\n☑ Envolvimento e expressão durante a atividade proposta.",
        "acess": "☑ Orientações orais claras, com demonstração simples da atividade antes da realização.\n☑ Possibilidade de participação por meio de fala, desenho, movimento ou apreciação.\n☑ Apoio individual na organização da produção ou expressão das ideias.",
    },
    {
        "data_horario": "03/07\n4ª e 5ª aula",
        "material": "MATEMÁTICA\nTema:\nSequência numérica",
        "aprendizagem": "HABILIDADE:\n(EF04MA01A) Ler, escrever, ordenar e completar sequências numéricas simples.",
        "desenvolvimento": "1. Abertura (contextualização e conexão com a realidade): iniciar com contagens simples, números presentes em calendários, casas, telefones e situações do cotidiano da turma.\n\n2. Desenvolvimento (exploração guiada do conteúdo): explicar no quadro como observar, continuar e organizar sequências numéricas com apoio de exemplos concretos e poucos números por vez.\n\n3. Atividade (prática com mediação docente): propor exercícios curtos de completar sequências e comparar números, acompanhando a turma durante a realização.\n\n4. Fechamento (correção e sistematização): corrigir coletivamente, retomar oralmente o procedimento e registrar uma síntese simples no quadro.",
        "acomp": "☑ Participação na resolução das atividades propostas.\n☑ Compreensão do tema por meio dos registros, contagens e explicações realizadas.\n☑ Organização dos procedimentos e participação na correção coletiva.",
        "acess": "☑ Explicação passo a passo com exemplos concretos e registros no quadro.\n☑ Retomada oral dos procedimentos antes e durante as atividades.\n☑ Apoio individual na organização das respostas e da sequência numérica.",
    },
]


def main() -> None:
    doc = Document(PATH)
    header = doc.tables[-2]
    aulas = doc.tables[-1]

    set_cell(header.rows[3].cells[1], "29/06 a 03/07")
    set_cell(header.rows[3].cells[3], "25")

    for idx, conteudo in enumerate(NOVAS_LINHAS, start=1):
        row = aulas.rows[idx]
        set_cell(row.cells[0], conteudo["data_horario"])
        set_cell(row.cells[1], conteudo["material"])
        set_cell(row.cells[2], conteudo["aprendizagem"])
        set_cell(row.cells[3], conteudo["desenvolvimento"])
        set_cell(row.cells[4], conteudo["acomp"])
        set_cell(row.cells[5], conteudo["acomp"])
        set_cell(row.cells[6], conteudo["acess"])

    doc.save(PATH)
    print(PATH)


if __name__ == "__main__":
    main()
