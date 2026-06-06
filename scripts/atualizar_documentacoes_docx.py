from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_ATUAL = date(2026, 6, 5).strftime("%d/%m/%Y")

CORE_DOCX = ROOT / "DOCUMENTACAO_ESTRUTURA_CORE_PLANOS_LUAN.docx"
CORE_TXT = ROOT / "DOCUMENTACAO_ESTRUTURA_CORE_PLANOS_LUAN_extracted.txt"
SISTEMA_DOCX = ROOT / "DOCUMENTACAO_SISTEMA_PLANOS_LUAN.docx"
SISTEMA_TXT = ROOT / "DOCUMENTACAO_SISTEMA_PLANOS_LUAN_extracted.txt"

CORE_ITEMS = [
    {
        "item": "__init__.py",
        "tipo": "Arquivo",
        "peso": "Tecnico",
        "papel": "Marca a pasta core como pacote Python e habilita os imports internos do sistema.",
        "fluxo": "Aparece em todas as execucoes, porque sem ele o aplicativo, os testes e os scripts perderiam a referencia do pacote core.",
        "obs": "Nao concentra regra de negocio; seu papel hoje e organizacional.",
    },
    {
        "item": "ae_priorizado.py",
        "tipo": "Arquivo",
        "peso": "Integracao piloto",
        "papel": "Aplica o modo AE priorizado para Portugues, lendo a base em assets/ae_priorizado e trocando a aprendizagem pelo AE correspondente.",
        "fluxo": "Entra depois da extracao das aulas dos PDFs. Quando a opcao de AE esta ativa, ele identifica a aula, procura a chave do guia, substitui o texto e ainda reordena as aulas conforme a sequencia do guia.",
        "obs": "Hoje esta em modo de teste para Portugues do Ensino Medio no 2o bimestre, mas ja funciona como uma ponte real entre o sistema e o guia priorizado.",
    },
    {
        "item": "avaliacao.py",
        "tipo": "Arquivo",
        "peso": "Pedagogico",
        "papel": "Gera frases de acompanhamento da aprendizagem e acessibilidade de forma dinamica.",
        "fluxo": "E usado quando o sistema precisa transformar tema, aprendizagem e desenvolvimento em criterios de observacao e apoio ao estudante.",
        "obs": "Hoje convive com a camada mais modular de core/lib, funcionando como apoio e compatibilidade pedagogica.",
    },
    {
        "item": "base_conhecimento.py",
        "tipo": "Arquivo",
        "peso": "Referencia",
        "papel": "Guarda padroes disciplinares e repertorios que ajudam a inteligencia local a montar aulas coerentes.",
        "fluxo": "E consultado por modulos que precisam saber quais tecnicas, focos e estruturas costumam combinar com cada disciplina.",
        "obs": "Serve como memoria pedagogica reutilizavel do projeto.",
    },
    {
        "item": "calendario.py",
        "tipo": "Arquivo",
        "peso": "Estrutural",
        "papel": "Calcula feriados nacionais, datas sem aula, extensao do mes e filtros de agenda.",
        "fluxo": "E chamado na montagem das datas do plano antes do upload dos PDFs, ajudando o app a excluir feriados e pontos facultativos e a completar a ultima semana do mes quando necessario.",
        "obs": "Hoje e uma das pecas mais importantes para evitar erro de contagem de aulas e preenchimento indevido em dias sem aula.",
    },
    {
        "item": "cdp/",
        "tipo": "Pacote",
        "peso": "Nucleo CDP atual",
        "papel": "Concentra a implementacao mais nova do fluxo CDP, especialmente em gerador_cdp.py.",
        "fluxo": "Atende os planos CDP contextual, Ensino Medio, Fundamental e multisseriada, montando tema, aprendizagem, metodologia, acompanhamento e acessibilidade a partir das planilhas e documentos locais.",
        "obs": "O __init__.py do pacote reexporta tambem o conteudo legado, mantendo compatibilidade com imports antigos.",
    },
    {
        "item": "cdp_legacy.py",
        "tipo": "Arquivo",
        "peso": "Compatibilidade",
        "papel": "Mantem a logica legada do CDP que ainda serve de base para funcoes importadas pelo pacote core.cdp.",
        "fluxo": "Ainda participa do modo CDP por meio da camada de compatibilidade do pacote core/cdp.",
        "obs": "Mostra que o sistema esta em transicao: ja existe uma arquitetura mais nova, mas parte da base antiga continua viva.",
    },
    {
        "item": "cdp_em_docx.py",
        "tipo": "Arquivo",
        "peso": "Ajuste final",
        "papel": "Reescreve e ajusta documentos DOCX do fluxo CDP, com foco em cenarios como Matematica contextual e Ensino Medio.",
        "fluxo": "Entra quando o plano CDP precisa ser corrigido ou reorganizado diretamente no Word, sem passar por um PDF comum.",
        "obs": "E a ponte entre o motor CDP e o acabamento final no documento.",
    },
    {
        "item": "database.py",
        "tipo": "Arquivo",
        "peso": "Estrutural",
        "papel": "Controla o banco SQLite do sistema, incluindo professores, turmas, vinculos, template_id, componente curricular e historico dos planos.",
        "fluxo": "E usado desde a abertura do app para carregar cadastros e volta a ser usado no final para gravar historico e recuperar arquivos gerados.",
        "obs": "Hoje o banco principal usado pelo sistema fica na raiz do projeto e e manipulado por este modulo.",
    },
    {
        "item": "disciplinas.py",
        "tipo": "Arquivo",
        "peso": "Estrutural",
        "papel": "Define disciplinas, bimestres e modos de funcionamento como PDF, CDP e CDP Fundamental.",
        "fluxo": "Ajuda a interface e os validadores a entenderem se a disciplina exige PDF, se usa modo contextual ou se ativa algum comportamento especial.",
        "obs": "Funciona como catalogo central das modalidades oficiais do app.",
    },
    {
        "item": "eja/",
        "tipo": "Pacote",
        "peso": "Especializacao",
        "papel": "Agrupa adaptacoes para EJA, como consolidacao de blocos e tecnicas com linguagem adequada a jovens e adultos.",
        "fluxo": "Entra em disciplinas que suportam modalidade EJA, hoje com destaque para Biologia e Ingles.",
        "obs": "Evita que a aula de EJA seja tratada como copia simplificada de uma aula regular.",
    },
    {
        "item": "helpers.py",
        "tipo": "Arquivo",
        "peso": "Apoio",
        "papel": "Reune funcoes pequenas de formatacao e apoio, como horario_para_plano, texto_lista, relatorio simples e preservacao da ordem de envio dos arquivos.",
        "fluxo": "Aparece em varias partes do sistema, principalmente na montagem de horarios, relatorios e tratamento dos PDFs enviados pelo usuario.",
        "obs": "Embora pequeno, foi importante nas correcoes recentes da ordem dos PDFs e no polimento do fluxo.",
    },
    {
        "item": "ia.py",
        "tipo": "Arquivo",
        "peso": "Motor externo",
        "papel": "Integra o sistema com OpenAI e Gemini para extrair e montar tema, aprendizagem e metodologia.",
        "fluxo": "E acionado quando o usuario escolhe usar IA. Ele prepara prompt, aplica orientacoes por disciplina, traz referencias metodologicas e normaliza a resposta estruturada.",
        "obs": "Hoje trabalha junto com validacoes e revisoes textuais para reduzir respostas fracas ou artificiais.",
    },
    {
        "item": "inteligencia_local.py",
        "tipo": "Arquivo",
        "peso": "Motor local",
        "papel": "Executa heuristicas e geracao local sem depender de IA externa.",
        "fluxo": "Serve como alternativa ou apoio ao fluxo com IA, especialmente quando se quer trabalhar so com regras internas do sistema.",
        "obs": "Continua importante como reserva de autonomia do projeto.",
    },
    {
        "item": "lib/",
        "tipo": "Pacote",
        "peso": "Base modular",
        "papel": "Agrupa extratores, classificadores, perfis de acompanhamento e acessibilidade, geradores de colunas pedagogicas, progressao, tecnicas e modalidades.",
        "fluxo": "E a camada modular mais reaproveitada pelo lote.py, pela IA e pelos adaptadores pedagogicos.",
        "obs": "Hoje a inteligencia pedagogica do sistema esta muito mais espalhada aqui do que estava na versao mais antiga da documentacao.",
    },
    {
        "item": "lote.py",
        "tipo": "Arquivo",
        "peso": "Nucleo central",
        "papel": "Transforma PDFs comuns em aulas estruturadas com tema, material, aprendizagem, metodologia, acompanhamento e acessibilidade.",
        "fluxo": "E o motor principal do fluxo PDF: extrai texto, identifica titulo e numero da aula, classifica o perfil e monta as colunas pedagogicas.",
        "obs": "Continua sendo o coracao da geracao das aulas comuns, mesmo com o crescimento dos subpacotes auxiliares.",
    },
    {
        "item": "metodologia_texto.py",
        "tipo": "Arquivo",
        "peso": "Apoio textual",
        "papel": "Faz ajustes linguisticos finos na metodologia, como regularizacao de verbos e polimento de frase.",
        "fluxo": "Entra no pos-processamento do texto antes de a aula seguir para validacao e DOCX.",
        "obs": "Ajuda a evitar metodologia truncada ou com tom mecanico.",
    },
    {
        "item": "modelos_docx.py",
        "tipo": "Arquivo",
        "peso": "Estrutural",
        "papel": "Resolve o template central correto por contexto, template_id, escola, disciplina e componente curricular.",
        "fluxo": "E usado quando o sistema precisa escolher entre MODELOEGLE, MODELOPADRE e MODELOCDP sem depender apenas do arquivo legado do professor.",
        "obs": "Foi uma evolucao importante para deixar o fluxo mais previsivel.",
    },
    {
        "item": "orientacao_estudos_metodologia.py",
        "tipo": "Arquivo",
        "peso": "Especializado",
        "papel": "Gera metodologia altamente especifica para Orientacao de Estudos, com perfis por missao, trilha ou jornada.",
        "fluxo": "E usado quando a disciplina exige uma abordagem propria, baseada em objetos, focos e referencias textuais dessa area.",
        "obs": "Hoje e um dos modulos especializados mais extensos do projeto.",
    },
    {
        "item": "orientacao_estudos_objetivos.py",
        "tipo": "Arquivo",
        "peso": "Especializado",
        "papel": "Centraliza objetivos e apoio textual para o fluxo de Orientacao de Estudos.",
        "fluxo": "Complementa a montagem da aula nessa disciplina, ajudando na coerencia entre titulo, foco e expectativa de aprendizagem.",
        "obs": "Mantem a disciplina menos dependente de regras genericas.",
    },
    {
        "item": "planos_luan.db",
        "tipo": "Arquivo tecnico",
        "peso": "Artefato local",
        "papel": "Arquivo presente dentro da pasta core, mas que nao representa hoje o banco principal de operacao do sistema.",
        "fluxo": "Nao deve ser lido como modulo de negocio. O fluxo principal usa o banco gerenciado por core/database.py na raiz do projeto.",
        "obs": "Vale documentar sua existencia para evitar confusao futuras.",
    },
    {
        "item": "professores_planos.py",
        "tipo": "Arquivo",
        "peso": "Estrutural",
        "papel": "Le, diagnostica, cria e atualiza modelos DOCX encontrados nas pastas dos professores.",
        "fluxo": "Extrai cabecalho, datas, horarios e grade semanal dos planos-modelo e ajuda a converter isso em cadastro reutilizavel no app.",
        "obs": "E a ponte entre os arquivos reais do trabalho escolar e a base interna do sistema.",
    },
    {
        "item": "projeto_vida_escopo.py",
        "tipo": "Arquivo",
        "peso": "Especializado",
        "papel": "Le e interpreta o escopo-sequencia de Projeto de Vida para localizar o item certo por turma, bimestre e aula.",
        "fluxo": "Entra quando Projeto de Vida precisa de um tratamento proprio, sem cair no comportamento generico das outras disciplinas.",
        "obs": "Ajuda o sistema a respeitar a logica pedagogica especifica desse componente.",
    },
    {
        "item": "prompts_por_disciplina.py",
        "tipo": "Arquivo",
        "peso": "Referencia IA",
        "papel": "Guarda orientacoes por disciplina usadas pelo modulo de IA.",
        "fluxo": "Ajuda a montar instrucoes diferentes para Portugues, Matematica, Projeto de Vida, Redacao e outras areas.",
        "obs": "Evita que o comportamento da IA fique uniforme demais entre disciplinas muito diferentes.",
    },
    {
        "item": "qualidade_metodologica.py",
        "tipo": "Arquivo",
        "peso": "Controle de qualidade",
        "papel": "Revisa metodologia, corrige ortografia, detecta contexto pedagogico, naturaliza texto e combate mojibake.",
        "fluxo": "Entra entre a geracao da aula e a validacao final, funcionando como camada de saneamento textual e pedagogico.",
        "obs": "E um dos modulos que mais contribuem para o plano final ficar profissional.",
    },
    {
        "item": "redacao_leitura_metodologia.py",
        "tipo": "Arquivo",
        "peso": "Especializado",
        "papel": "Gera a estrutura metodologica de Redacao e Leitura com etapas fixas e criterios mais autorais.",
        "fluxo": "Atua quando o sistema identifica perfis como leitura literaria, planejamento de producao, devolutiva e producao final.",
        "obs": "Foi uma ampliacao importante do repertorio especializado do projeto.",
    },
    {
        "item": "referencias_metodologia.py",
        "tipo": "Arquivo",
        "peso": "Referencia",
        "papel": "Carrega textos de referencia metodologica salvos no projeto.",
        "fluxo": "Entrega material de apoio por disciplina e turma para a IA e para modulos que precisam reforcar o tom pedagogico esperado.",
        "obs": "Funciona como uma biblioteca interna de memoria metodologica.",
    },
    {
        "item": "validador_plano.py",
        "tipo": "Arquivo",
        "peso": "Controle de qualidade",
        "papel": "Valida aulas geradas e detecta problemas como tema repetido, metodologia pobre, aprendizagem curta e campos vazios.",
        "fluxo": "E uma das ultimas barreiras antes da revisao do usuario e da geracao do DOCX.",
        "obs": "Hoje tambem protege o sistema de contaminacao semantica entre disciplinas.",
    },
]

SISTEMA_PARTES = [
    ("planos_luan_app.py", "Tela principal e orquestrador do fluxo de uso", "Interface"),
    ("core/lote.py", "Motor principal que transforma PDFs comuns em aulas", "Nucleo PDF"),
    ("core/lib/", "Camada modular de extracao, classificacao, colunas pedagogicas e perfis", "Base pedagogica"),
    ("core/cdp/", "Motor atual dos fluxos CDP contextual, fundamental e multisseriada", "Nucleo CDP"),
    ("core/ae_priorizado.py + assets/ae_priorizado/", "Integracao piloto entre o guia priorizado e o plano de aula", "AE priorizado"),
    ("core/calendario.py", "Controle de feriados, dias sem aula e extensao do mes", "Agenda"),
    ("core/database.py + planos_luan.db", "Cadastro persistente e historico dos planos", "Banco de dados"),
    ("core/professores_planos.py", "Leitura e padronizacao dos modelos DOCX dos professores", "Modelos locais"),
    ("core/modelos_docx.py + templates/", "Escolha de templates centrais por contexto", "Modelos centrais"),
    ("docx_generator/", "Preenchimento final dos documentos Word", "Saida DOCX"),
    ("Planos feitos/ + REFERENCIAS_METODOLOGIA/", "Base documental e pedagogica do modo CDP e das referencias metodologicas", "Dados de apoio"),
    ("scripts/", "Ferramentas de manutencao, migracao, extracao e apoio tecnico", "Operacao"),
    ("tests/", "Cobertura automatica das regras mais sensiveis do sistema", "Qualidade"),
]

ESTRUTURAS_SISTEMA = [
    {
        "titulo": "planos_luan_app.py - a tela principal e o grande orquestrador",
        "oque": "E o arquivo principal do aplicativo Streamlit. Hoje ele nao e apenas uma tela: organiza selecao de professor, disciplina, turma, bimestre, mes, agenda, PDF, revisao, geracao e download.",
        "como": "O fluxo atual ficou mais guiado do que na documentacao antiga. O usuario escolhe contexto, o sistema tenta localizar modelo automatico, sincroniza datas e horarios do mes, filtra feriados e dias sem aula, define modo de upload dos PDFs, processa as aulas, abre revisao centralizada e so depois gera o DOCX.",
        "destaques": "Hoje o app tambem oferece: modo de envio todos de uma vez ou um por aula; geracao para 2a turma; divisao de metodologia em dois dias; modalidade EJA para disciplinas suportadas; extensao apos o mes; e a opcao piloto de usar AE no lugar da habilidade em Portugues EM 2o bimestre.",
    },
    {
        "titulo": "config.py - configuracao central de caminhos e padroes",
        "oque": "Centraliza caminhos externos e internos, como pasta principal de trabalho, pasta de backups, templates, banco, modelos padrao de IA e limites de leitura.",
        "como": "Ele liga a instalacao local do PLANOS_LUAN a caminhos como D:\\PLANOS DE JUNHO e D:\\PLANOS-FINALIZADOS, alem de criar automaticamente algumas pastas quando necessario.",
        "destaques": "A documentacao antiga ainda citava varios caminhos espalhados. Hoje o projeto ja concentra melhor essa parte, embora ainda haja espaco para evoluir.",
    },
    {
        "titulo": "core/database.py e planos_luan.db - o cadastro persistente",
        "oque": "Guardam professores, turmas, disciplinas, grade semanal, arquivo_modelo, template_id, componente curricular e historico dos planos em SQLite.",
        "como": "O modulo inicializa tabelas, migra um JSON legado quando existir e oferece operacoes de salvar, atualizar, duplicar, excluir vinculos e recuperar historico de geracao.",
        "destaques": "Em comparacao com a versao antiga da documentacao, o cadastro agora esta mais rico: nao guarda so professor/turma, mas tambem template central e componente curricular associado.",
    },
    {
        "titulo": "core/professores_planos.py - a ponte com os modelos reais dos professores",
        "oque": "Le planos DOCX das pastas dos professores, extrai cabecalhos, semanas, datas e horarios e ajuda a criar ou atualizar modelos padronizados.",
        "como": "Ele identifica tabelas de cabecalho e tabelas de aulas, resume a grade semanal e devolve ao sistema informacoes que depois viram cadastro reaproveitavel.",
        "destaques": "Hoje essa leitura e ainda mais importante porque o sistema usa a agenda detectada no modelo para sincronizar o mes, em vez de depender so da quantidade fixa de aulas por semana.",
    },
    {
        "titulo": "core/modelos_docx.py e templates/ - camada nova de templates centrais",
        "oque": "Escolhem o template central correto entre MODELOEGLE, MODELOPADRE e MODELOCDP.",
        "como": "Quando o cadastro traz template_id ou quando o contexto sugere escola, disciplina ou modo CDP, o sistema consegue cair num modelo central coerente mesmo sem depender exclusivamente do arquivo legado da pasta do professor.",
        "destaques": "Essa e uma das mudancas mais importantes desde a documentacao antiga, porque tornou a selecao de modelo mais previsivel.",
    },
    {
        "titulo": "core/calendario.py - agenda escolar, feriados e dias sem aula",
        "oque": "Calcula feriados nacionais, sugere datas bloqueadas, permite filtragem de dias sem aula e controla a extensao do mes para fechar a ultima semana.",
        "como": "O app monta a agenda do mes a partir do cadastro ou do modelo detectado, depois passa por esse modulo para remover feriados e datas marcadas manualmente.",
        "destaques": "Ajustes recentes fizeram a contagem semanal do plano passar a refletir a carga real daquela semana, inclusive quando a primeira semana tem feriado.",
    },
    {
        "titulo": "core/lote.py, core/lib/ e core/eja/ - o motor pedagogico do fluxo PDF",
        "oque": "Conjunto que extrai texto dos PDFs, identifica titulo e numero da aula, classifica o perfil pedagogico e monta tema, aprendizagem, metodologia, acompanhamento e acessibilidade.",
        "como": "lote.py continua como nucleo, mas hoje trabalha apoiado numa base modular maior: extrator_pdf, extrator_titulo, classificador, gerador_colunas_pedagogicas, modulos de progressao, tecnicas, acompanhamento, acessibilidade e adaptacoes EJA.",
        "destaques": "Na pratica, o sistema ficou menos monolitico no miolo pedagogico, mesmo que planos_luan_app.py ainda seja grande.",
    },
    {
        "titulo": "core/ia.py, prompts_por_disciplina.py e referencias_metodologia.py - a camada opcional de IA",
        "oque": "Permitem usar OpenAI ou Gemini para montar aulas com apoio externo, mantendo orientacoes por disciplina e referencias pedagogicas do proprio projeto.",
        "como": "Quando o usuario ativa IA, o sistema monta um prompt contextualizado, injeta referencias metodologicas e normaliza a resposta para o mesmo formato usado no fluxo local.",
        "destaques": "O sistema continua podendo funcionar sem IA, o que e importante para custo, autonomia e robustez em ambiente escolar.",
    },
    {
        "titulo": "core/ae_priorizado.py, assets/ae_priorizado/ e scripts de extracao - a nova integracao com o guia priorizado",
        "oque": "Representam a nova camada criada para permitir que o plano use AE no lugar da habilidade normal.",
        "como": "Primeiro o guia priorizado foi recortado e transformado em base estruturada. Depois o sistema ganhou uma opcao de tela que, quando ativa no contexto suportado, procura a aula no mapa, troca a aprendizagem pelo AE e reorganiza a sequencia segundo a ordem do guia.",
        "destaques": "Hoje esse recurso esta em fase piloto para Portugues do Ensino Medio, 2o bimestre, mas ja mudou o comportamento do sistema de forma relevante.",
    },
    {
        "titulo": "core/cdp/, cdp_legacy.py, cdp_em_docx.py e Planos feitos/ - o universo CDP",
        "oque": "Controlam o fluxo que gera planos sem depender de PDFs comuns, usando planilhas e documentos locais de habilidades.",
        "como": "Esse bloco trabalha com CDP contextual, Ensino Medio, Fundamental e multisseriada, selecionando habilidades, objetos, titulos e colunas pedagogicas a partir dos arquivos locais.",
        "destaques": "A arquitetura atual mistura uma camada nova em pacote com uma camada de compatibilidade legado. Isso precisa aparecer na documentacao para nao dar a impressao de que existe um unico modulo CDP simples.",
    },
    {
        "titulo": "docx_generator/preencher.py e preencher_cdp.py - montagem final do Word",
        "oque": "Transformam as aulas revisadas em documento DOCX final, respeitando o modelo escolhido.",
        "como": "Esses modulos substituem cabecalho, distribuem aulas por semana, preenchem data, horario, titulo, aprendizagem, metodologia, acompanhamento e acessibilidade. Tambem removem semanas vazias e ajustam o layout das tabelas.",
        "destaques": "Uma correcao recente importante foi a contagem real de aulas previstas na semana com base no horario das aulas, e nao apenas no valor fixo do cadastro.",
    },
    {
        "titulo": "scripts/ - manutencao, migracao e extracao",
        "oque": "Agrupam tarefas de apoio tecnico, como normalizacao de data/hora, padronizacao de nomes dos modelos, migracao de modelos legados e extracao da base de AE priorizado.",
        "como": "Nao fazem parte do uso normal do professor, mas ajudam a manter o sistema limpo, coerente e com base documental atualizada.",
        "destaques": "Os scripts de AE priorizado mostram que o sistema ganhou uma etapa nova de preparacao de dados externos antes da geracao dos planos.",
    },
    {
        "titulo": "assets/, tests/ e arquivos de abertura - acabamento, seguranca e operacao",
        "oque": "assets guarda estilo visual e base de AE priorizado; tests protege regras sensiveis; arquivos .bat/.ps1/.vbs ajudam a abrir, fechar e reiniciar o sistema.",
        "como": "A interface atual usa style.css e hero-planejamento.svg, enquanto os testes cobrem pontos como feriados, AE priorizado, limpeza de DOCX, CDP e perfis pedagogicos.",
        "destaques": "Essa camada de apoio ficou mais rica desde a documentacao anterior e contribui diretamente para estabilidade e usabilidade.",
    },
]

MUDANCAS_RELEVANTES = [
    "A tela principal ficou mais guiada e ganhou uma etapa de revisao centralizada antes do DOCX.",
    "O sistema passou a sincronizar agenda mensal com feriados, extensao do mes e dias sem aula selecionados pelo usuario.",
    "A selecao de template deixou de depender apenas do arquivo da pasta do professor e passou a usar template_id e contexto.",
    "O fluxo PDF ganhou modos de upload diferentes, divisao de metodologia em dois dias e geracao para 2a turma.",
    "Foi criada a integracao piloto com AE priorizado para Portugues EM 2o bimestre.",
    "A contagem semanal de aulas no documento final passou a refletir a carga real da semana, inclusive em semanas com feriado.",
    "A camada core se expandiu com subpacotes especializados como core/lib, core/cdp e core/eja.",
]

FORCAS_ATUAIS = [
    "Fluxo mais guiado e menos manual na tela principal.",
    "Capacidade de trabalhar com PDF comum, CDP, EJA e AE priorizado piloto sem trocar de sistema.",
    "Uso combinado de modelo do professor, template central e banco de dados.",
    "Revisao das aulas antes da geracao do DOCX, o que reduz retrabalho.",
    "Historico persistente dos planos gerados.",
    "Base de testes melhor do que a que existia quando a documentacao antiga foi escrita.",
]

PRIORIDADES = [
    "Continuar modularizando planos_luan_app.py, que ainda concentra muitas responsabilidades.",
    "Expandir o modo AE priorizado para outras disciplinas e para 3o e 4o bimestres quando a base estiver pronta.",
    "Reduzir a sobreposicao entre camada nova e camada legada do CDP.",
    "Fortalecer diagnosticos de modelo Word e de base de dados externos.",
    "Manter scripts e documentacao como artefatos oficiais de manutencao, para evitar novo envelhecimento da descricao do sistema.",
]


def set_page_margins(section) -> None:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size, color, bold in [
        ("Title", 22, RGBColor(25, 67, 115), True),
        ("Heading 1", 15, RGBColor(31, 78, 121), True),
        ("Heading 2", 11.5, RGBColor(31, 78, 121), True),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_label_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    head = table.rows[0].cells
    for idx, header in enumerate(headers):
        head[idx].text = header
        shade(head[idx], "D9EAF7")
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value


def write_text(path: Path, lines: list[str]) -> None:
    path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def build_core_doc() -> None:
    lines = [
        "=== File: DOCUMENTACAO_ESTRUTURA_CORE_PLANOS_LUAN.docx ===",
        "",
        "Estrutura Atual da Pasta core do PLANOS_LUAN",
        "",
        f"Documento atualizado em {DATA_ATUAL}, em linguagem simples, com foco na estrutura real do codigo.",
        "",
        "Objetivo: registrar como a pasta core esta organizada hoje, o papel de cada modulo ou subpacote e em que parte do fluxo ele entra.",
        "",
        "Visao Geral",
        "",
        "A pasta core continua sendo o centro da inteligencia de negocio do PLANOS_LUAN, mas hoje ela esta mais distribuida do que estava na versao anterior da documentacao.",
        "Ela combina arquivos de infraestrutura, motores pedagogicos, camadas especializadas por disciplina ou modalidade e subpacotes como core/lib, core/cdp e core/eja.",
        "Para aulas comuns vindas de PDF, lote.py continua sendo o nucleo principal. Para CDP, o projeto passou a conviver com uma camada nova em pacote e uma camada legada de compatibilidade.",
        "Tambem surgiu uma integracao nova com o guia priorizado por meio de ae_priorizado.py.",
        "",
        "Mapa Rapido dos Itens da Pasta core",
        "",
        "--- Table ---",
        "Item | Tipo | Papel principal | Peso no sistema",
    ]

    doc = Document()
    set_page_margins(doc.sections[0])
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Estrutura Atual da Pasta core do PLANOS_LUAN")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.add_run(
        f"Documento atualizado em {DATA_ATUAL}, em linguagem simples, com foco na estrutura real do codigo."
    ).italic = True

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    intro.add_run("Objetivo: registrar como a pasta ")
    intro.add_run("core").bold = True
    intro.add_run(
        " esta organizada hoje, o papel de cada modulo ou subpacote e em que parte do fluxo ele entra."
    )

    doc.add_heading("Visao Geral", level=1)
    add_bullets(
        doc,
        [
            "A pasta core continua sendo o centro da inteligencia de negocio do PLANOS_LUAN, mas hoje ela esta mais distribuida do que estava na versao anterior da documentacao.",
            "Ela combina arquivos de infraestrutura, motores pedagogicos, camadas especializadas por disciplina ou modalidade e subpacotes como core/lib, core/cdp e core/eja.",
            "Para aulas comuns vindas de PDF, lote.py continua sendo o nucleo principal.",
            "Para CDP, o projeto passou a conviver com uma camada nova em pacote e uma camada legada de compatibilidade.",
            "Tambem surgiu uma integracao nova com o guia priorizado por meio de ae_priorizado.py.",
        ],
    )

    doc.add_heading("Mapa Rapido dos Itens da Pasta core", level=1)
    table_rows = []
    for item in CORE_ITEMS:
        table_rows.append([item["item"], item["tipo"], item["papel"], item["peso"]])
        lines.append(
            f"{item['item']} | {item['tipo']} | {item['papel']} | {item['peso']}"
        )
    lines.extend(["-------------", "", "Explicacao Detalhada dos Itens", ""])
    add_table(doc, ["Item", "Tipo", "Papel principal", "Peso no sistema"], table_rows)

    doc.add_paragraph()
    doc.add_heading("Explicacao Detalhada dos Itens", level=1)
    for item in CORE_ITEMS:
        doc.add_heading(item["item"], level=2)
        add_label_paragraph(doc, "Tipo", item["tipo"])
        add_label_paragraph(doc, "Funcao principal", item["papel"])
        add_label_paragraph(doc, "Quando entra no fluxo", item["fluxo"])
        add_label_paragraph(doc, "Observacoes", item["obs"])

        lines.append(item["item"])
        lines.append("")
        lines.append(f"Tipo: {item['tipo']}")
        lines.append(f"Funcao principal: {item['papel']}")
        lines.append(f"Quando entra no fluxo: {item['fluxo']}")
        lines.append(f"Observacoes: {item['obs']}")
        lines.append("")

    leituras = [
        'Se a pergunta for "onde nascem as aulas comuns a partir dos PDFs?", a resposta principal continua sendo lote.py, apoiado hoje por uma base modular mais forte em core/lib.',
        'Se a pergunta for "onde esta o cadastro e o historico?", a resposta continua sendo core/database.py, mas com mais responsabilidade do que antes por causa de template_id e componente curricular.',
        'Se a pergunta for "onde o sistema entende o CDP?", a resposta agora precisa citar dois niveis: o pacote core/cdp como camada atual e cdp_legacy.py como camada de compatibilidade ainda em uso.',
        'Se a pergunta for "onde entram as novidades do guia priorizado?", a resposta e ae_priorizado.py junto da base em assets/ae_priorizado.',
        'Se a pergunta for "onde o texto final e saneado?", os destaques continuam sendo qualidade_metodologica.py, metodologia_texto.py e validador_plano.py.',
    ]
    doc.add_heading("Leitura Final da Arquitetura Atual", level=1)
    add_bullets(doc, leituras)
    lines.extend(["Leitura Final da Arquitetura Atual", ""])
    lines.extend(leituras)
    lines.append("")

    observacao = (
        "Esta documentacao descreve o estado atual observado em 05/06/2026. Como o projeto ainda tem "
        "camadas novas e legadas convivendo, o texto deve ser revisado sempre que a estrutura de core mudar "
        "de forma relevante."
    )
    doc.add_paragraph()
    add_label_paragraph(doc, "Observacao importante", observacao)
    lines.extend(["Observacao importante", observacao])

    doc.save(CORE_DOCX)
    write_text(CORE_TXT, lines)


def build_system_doc() -> None:
    lines = [
        "=== File: DOCUMENTACAO_SISTEMA_PLANOS_LUAN.docx ===",
        "",
        "Documentacao Atual do Sistema PLANOS_LUAN",
        "",
        f"Documento atualizado em {DATA_ATUAL}, em linguagem simples, descrevendo o comportamento real do sistema no momento atual.",
        "",
        "--- Table ---",
        "Leitura proposta | Este documento foi escrito para quem precisa entender o sistema como ferramenta de trabalho: o que ele resolve hoje, como as informacoes andam por dentro dele e quais partes do projeto sao mais importantes no uso diario.",
        "-------------",
        "",
        "1. O que faz o sistema hoje?",
        "",
        "O PLANOS_LUAN continua sendo um sistema de montagem de planos de aula em Word, mas hoje ele esta mais amplo do que na versao descrita pela documentacao antiga.",
        "Ele gera planos a partir de PDFs comuns, reaproveita modelos DOCX ja existentes dos professores, consulta um banco de dados local, escolhe templates centrais por contexto, sincroniza agenda semanal com feriados e dias sem aula, salva historico e oferece revisao das aulas antes do documento final.",
        "Tambem suporta fluxos especiais como CDP contextual, CDP Fundamental, CDP multisseriada, modalidade EJA em disciplinas especificas e um piloto de AE priorizado para Portugues do Ensino Medio no 2o bimestre.",
        "",
        "2. Como o sistema trabalha por dentro hoje",
        "",
        "O uso atual pode ser lido como um fluxo guiado em seis momentos: contexto, agenda, materiais, processamento, revisao e geracao final.",
        "Primeiro o usuario escolhe professor, disciplina, turma, mes, bimestre, escola e, quando necessario, componente curricular, modalidade EJA ou turma espelho.",
        "Depois o sistema tenta localizar um modelo valido: pode vir da pasta do professor, do cadastro salvo no banco ou de um template central como MODELOEGLE, MODELOPADRE ou MODELOCDP.",
        "Na sequencia, a agenda do mes e montada a partir do cadastro ou do proprio modelo detectado. Nessa etapa entram extensao apos o mes, feriados nacionais e o campo Dias sem aula, que permite bloquear datas especificas.",
        "So depois disso o usuario organiza os PDFs, que podem ser enviados todos de uma vez ou um por aula. O sistema tambem consegue dividir metodologia em dois dias e gerar o mesmo conjunto para uma 2a turma.",
        "No processamento, o motor local ou a IA extraem tema, aprendizagem, metodologia, acompanhamento e acessibilidade. Se a opcao AE estiver ativa e o contexto for suportado, o sistema troca a aprendizagem pelo AE e segue a ordem do guia priorizado.",
        "Por fim, o usuario revisa as aulas na tela e gera o DOCX. O resultado e salvo no historico com professor, disciplina, turma e arquivo final.",
        "",
        "--- Table ---",
        "Parte | Funcao principal | Tipo de conteudo",
    ]

    doc = Document()
    set_page_margins(doc.sections[0])
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Documentacao Atual do Sistema PLANOS_LUAN")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.add_run(
        f"Documento atualizado em {DATA_ATUAL}, em linguagem simples, descrevendo o comportamento real do sistema no momento atual."
    ).italic = True

    leitura = doc.add_table(rows=1, cols=2)
    leitura.style = "Table Grid"
    leitura.rows[0].cells[0].text = "Leitura proposta"
    leitura.rows[0].cells[1].text = (
        "Este documento foi escrito para quem precisa entender o sistema como ferramenta de trabalho: "
        "o que ele resolve hoje, como as informacoes andam por dentro dele e quais partes do projeto sao "
        "mais importantes no uso diario."
    )
    shade(leitura.rows[0].cells[0], "D9EAF7")
    shade(leitura.rows[0].cells[1], "F4F8FB")

    doc.add_paragraph()
    doc.add_heading("1. O que faz o sistema hoje?", level=1)
    for texto in [
        "O PLANOS_LUAN continua sendo um sistema de montagem de planos de aula em Word, mas hoje ele esta mais amplo do que na versao descrita pela documentacao antiga.",
        "Ele gera planos a partir de PDFs comuns, reaproveita modelos DOCX ja existentes dos professores, consulta um banco de dados local, escolhe templates centrais por contexto, sincroniza agenda semanal com feriados e dias sem aula, salva historico e oferece revisao das aulas antes do documento final.",
        "Tambem suporta fluxos especiais como CDP contextual, CDP Fundamental, CDP multisseriada, modalidade EJA em disciplinas especificas e um piloto de AE priorizado para Portugues do Ensino Medio no 2o bimestre.",
    ]:
        doc.add_paragraph(texto)

    doc.add_heading("2. Como o sistema trabalha por dentro hoje", level=1)
    add_bullets(
        doc,
        [
            "Contexto: professor, disciplina, turma, mes, bimestre, escola, componente curricular e opcoes especiais.",
            "Agenda: deteccao de modelo, datas e horarios, com extensao apos o mes e filtro de feriados/dias sem aula.",
            "Materiais: upload em lote ou por aula, com opcao de dividir metodologia em dois dias.",
            "Processamento: motor local ou IA gera as colunas pedagogicas e, quando ativado, aplica AE priorizado.",
            "Revisao: o usuario confere e ajusta tema, aprendizagem, metodologia, acompanhamento e acessibilidade.",
            "Saida: geracao do DOCX final e gravacao no historico.",
        ],
    )

    comp_rows = []
    for parte, funcao, tipo in SISTEMA_PARTES:
        comp_rows.append([parte, funcao, tipo])
        lines.append(f"{parte} | {funcao} | {tipo}")
    lines.extend(["-------------", "", "3. Estruturas principais do sistema", ""])
    add_table(doc, ["Parte", "Funcao principal", "Tipo de conteudo"], comp_rows)

    doc.add_heading("3. Estruturas principais do sistema", level=1)
    for bloco in ESTRUTURAS_SISTEMA:
        doc.add_heading(bloco["titulo"], level=2)
        add_label_paragraph(doc, "O que e e para que serve", bloco["oque"])
        add_label_paragraph(doc, "Como funciona hoje", bloco["como"])
        add_label_paragraph(
            doc, "Destaques em relacao a versao antiga", bloco["destaques"]
        )

        lines.append(bloco["titulo"])
        lines.append("")
        lines.append(f"O que e e para que serve: {bloco['oque']}")
        lines.append(f"Como funciona hoje: {bloco['como']}")
        lines.append(
            f"Destaques em relacao a versao antiga: {bloco['destaques']}"
        )
        lines.append("")

    doc.add_heading("4. Mudancas relevantes desde a documentacao anterior", level=1)
    add_bullets(doc, MUDANCAS_RELEVANTES)
    lines.extend(["4. Mudancas relevantes desde a documentacao anterior", ""])
    lines.extend(MUDANCAS_RELEVANTES)
    lines.append("")

    doc.add_heading("5. Pontos fortes atuais", level=1)
    add_bullets(doc, FORCAS_ATUAIS)
    lines.extend(["5. Pontos fortes atuais", ""])
    lines.extend(FORCAS_ATUAIS)
    lines.append("")

    doc.add_heading("6. Prioridades de evolucao", level=1)
    add_bullets(doc, PRIORIDADES)
    lines.extend(["6. Prioridades de evolucao", ""])
    lines.extend(PRIORIDADES)
    lines.append("")

    conclusao = [
        "Hoje o PLANOS_LUAN deve ser entendido como uma linha de producao pedagogica mais madura do que aquela registrada na documentacao antiga.",
        "A tela principal organiza o uso, o cadastro e os modelos alimentam a agenda, o motor PDF ou CDP monta as aulas, a revisao humana refina o conteudo e o gerador DOCX entrega o documento final no formato esperado pela escola.",
        "O sistema ja nao e apenas um leitor de PDF com Word na saida: ele combina agenda escolar, banco de dados, camadas pedagogicas especializadas, templates centrais, revisao, historico e integracoes novas como o AE priorizado.",
    ]
    doc.add_heading("7. Conclusao", level=1)
    for texto in conclusao:
        doc.add_paragraph(texto)
    lines.extend(["7. Conclusao", ""])
    lines.extend(conclusao)

    doc.save(SISTEMA_DOCX)
    write_text(SISTEMA_TXT, lines)


def main() -> None:
    build_core_doc()
    build_system_doc()
    print(CORE_DOCX)
    print(SISTEMA_DOCX)


if __name__ == "__main__":
    main()
