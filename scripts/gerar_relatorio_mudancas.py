import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def create_report():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("RELATÓRIO DE ATUALIZAÇÕES E AUDITORIA DO SISTEMA\nPLANOS_LUAN")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Documentação técnica das últimas implementações e refatorações para sincronização entre agentes")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    
    # Spacer
    doc.add_paragraph()
    
    # Section 1: Overview
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Visão Geral do Estado Atual")
    h1_run.bold = True
    h1_run.font.size = Pt(14)
    h1_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p = doc.add_paragraph(
        "Este documento registra detalhadamente as melhorias visuais, de arquitetura, pedagógicas e estruturais implementadas recentemente no sistema PLANOS_LUAN. "
        "A base de testes automatizados está com 586 testes passando e 1 falha esperada (xfail) mapeada, o que garante a estabilidade das modificações."
    )
    
    # Section 2: prioritarios
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Guias Priorizados (Aprendizagem Essencial - AE)")
    h2_run.bold = True
    h2_run.font.size = Pt(14)
    h2_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p2 = doc.add_paragraph(
        "Anteriormente restrita ao Ensino Fundamental, a funcionalidade de leitura inteligente do Guia Priorizado "
        "foi estendida com sucesso para o Ensino Médio. Além disso, foram processados os PDFs dos guias prioritários e criadas as respectivas planilhas estruturadas (.xlsx) em \"C:\\Users\\Luan Dias\\PDF novos\" para as seguintes disciplinas no 3º Bimestre:"
    )
    
    disciplinas = [
        "Língua Portuguesa (Ensino Fundamental e Médio)",
        "Matemática (Ensino Fundamental e Médio)",
        "Ciências (Ensino Fundamental)",
        "Biologia (Ensino Médio)",
        "Geografia (Ensino Fundamental e Médio)",
        "Arte (Ensino Fundamental e Médio)",
        "Língua Inglesa (Ensino Fundamental e Médio)",
        "Química (Ensino Médio)",
        "Sociologia (Ensino Médio)",
        "História (Ensino Fundamental e Médio)"
    ]
    for disc in disciplinas:
        doc.add_paragraph(disc, style='List Bullet')
        
    doc.add_paragraph(
        "Esses dados permitem que o sistema aplique e preencha as habilidades e conteúdos essenciais no template .docx "
        "automaticamente ao selecionar a opção correspondente na tela inicial."
    )

    # Section 3: UI
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Modernização Visual da Interface")
    h3_run.bold = True
    h3_run.font.size = Pt(14)
    h3_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p3 = doc.add_paragraph(
        "A tela inicial (Streamlit) passou por uma modernização cosmética completa seguindo um tema escuro profissional. "
        "As principais melhorias incluem:"
    )
    doc.add_paragraph("Fundo com degradê azul-marinho escuro (#0D1B2A → #112236 → #0A1628).", style='List Bullet')
    doc.add_paragraph("Hero section com efeito de neon azul/verde e badge animado de 'SISTEMA ATIVO'.", style='List Bullet')
    doc.add_paragraph("Barra de métricas estilizada exibindo modos de uso, disciplinas e turmas.", style='List Bullet')
    doc.add_paragraph("Menu lateral e menu de navegação escuros com brilho sutil no item selecionado.", style='List Bullet')
    doc.add_paragraph("Correções de contraste nos textos de inputs de rádio, checkboxes e labels para garantir acessibilidade visual.", style='List Bullet')
    
    # Section 4: Duplicidade
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Alerta de Duplicidade de Geração")
    h4_run.bold = True
    h4_run.font.size = Pt(14)
    h4_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p4 = doc.add_paragraph(
        "Para evitar a regeração de planos idênticos por diferentes docentes, foi integrada uma regra de controle no histórico:\n"
        "• Lógica no DB: A função verificar_plano_gerado_por_outro_professor() no core/database.py busca se a combinação de disciplina, série/turma e bimestre já foi gerada por outro professor.\n"
        "• Alerta visual: Caso exista, o Streamlit renderiza um aviso destacado em amarelo (st.warning) antes que o usuário gere o plano, indicando os nomes de quem já gerou, o bimestre e a data da última geração."
    )
    
    # Section 5: Cleanup
    h5 = doc.add_paragraph()
    h5_run = h5.add_run("5. Limpeza de Código Morto e Reorganização")
    h5_run.bold = True
    h5_run.font.size = Pt(14)
    h5_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p5 = doc.add_paragraph(
        "Com o objetivo de manter a integridade arquitetônica do projeto, foi feita uma auditoria geral onde:"
    )
    doc.add_paragraph("O arquivo divisor_metodologia.py foi movido da raiz do projeto para a pasta core/.", style='List Bullet')
    doc.add_paragraph("Os arquivos tela_inicial_moderna.py e ui_components.py foram movidos para a pasta ui/.", style='List Bullet')
    doc.add_paragraph("Todos os rascunhos temporários, backups manuais (ex: planos_luan_app_backup.py) e saídas de testes (.txt) foram movidos de forma limpa para a pasta externa D:\\PLANOS_LUAN_RASCUNHOS\\.", style='List Bullet')
    doc.add_paragraph("A pasta de ambiente virtual venv/ foi deletada para evitar conflitos com a .venv/ oficial (configurada no PyCharm).", style='List Bullet')
    doc.add_paragraph("Remoção do banco planos_luan.db duplicado dentro de core/ (o sistema lê unicamente o banco da raiz).", style='List Bullet')
    doc.add_paragraph("Compactação do banco de dados principal utilizando a instrução SQL VACUUM.", style='List Bullet')
    
    # Section 6: Sincronizacao
    h6 = doc.add_paragraph()
    h6_run = h6.add_run("6. Orientações para o Próximo Agente")
    h6_run.bold = True
    h6_run.font.size = Pt(14)
    h6_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p6 = doc.add_paragraph(
        "Ao realizar novas modificações, siga rigorosamente as seguintes diretrizes:\n"
        "1. Utilize sempre o interpretador da pasta .venv (Python 3.12).\n"
        "2. Sempre execute 'pytest' na raiz após qualquer alteração no código para validar a suite de testes.\n"
        "3. Não adicione novos scripts de rascunhos na raiz do projeto; utilize a pasta tmp/ ou subpastas de rascunhos.\n"
        "4. Mantenha os comentários do código em português.\n"
        "5. O arquivo de banco oficial é unicamente o planos_luan.db na raiz do projeto."
    )
    
    output_dir = Path("D:/Documentos")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "Relatorio_Ultimas_Mudancas.docx"
    doc.save(output_path)
    print(f"Relatório gerado em: {output_path}")

if __name__ == "__main__":
    create_report()
