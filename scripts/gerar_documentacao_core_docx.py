from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "DOCUMENTACAO_ESTRUTURA_CORE_PLANOS_LUAN.docx"


ITENS = [
    {
        "arquivo": "__init__.py",
        "funcao": "Marca a pasta core como pacote Python.",
        "fluxo": "Permite que os outros modulos sejam importados pelo aplicativo principal e pelos testes.",
        "observacoes": "Hoje ele nao concentra regra de negocio; funciona mais como ponto tecnico de organizacao do pacote.",
    },
    {
        "arquivo": "avaliacao.py",
        "funcao": "Gera os textos de acompanhamento da aprendizagem e acessibilidade para cada aula.",
        "fluxo": "Entra quando o sistema precisa completar o plano com orientacoes pedagogicas coerentes com a disciplina, tema e tipo de aula.",
        "observacoes": "Hoje atua como uma camada de roteamento para modulos mais especializados em core/lib, ajudando a manter o texto final mais contextualizado.",
    },
    {
        "arquivo": "base_conhecimento.py",
        "funcao": "Guarda padroes disciplinares, tecnicas esperadas e estruturas base de aula.",
        "fluxo": "Serve de base para a inteligencia local interpretar a disciplina e escolher tom, comandos e organizacao metodologica.",
        "observacoes": "E um arquivo de referencia do sistema: concentra conhecimento pedagogico reutilizado em varios pontos.",
    },
    {
        "arquivo": "calendario.py",
        "funcao": "Calcula datas importantes do calendario escolar, como Pascoa, feriados nacionais e exclusao de dias sem aula.",
        "fluxo": "Apoia a montagem das datas do mes antes da geracao do plano, inclusive filtrando feriados e pontos facultativos marcados pelo usuario.",
        "observacoes": "Foi criado para reduzir erros de agenda e evitar que o plano preencha automaticamente dias em que a escola nao vai funcionar.",
    },
    {
        "arquivo": "cdp.py",
        "funcao": "Lê e organiza os dados do modo CDP a partir de planilhas e documentos locais.",
        "fluxo": "Seleciona habilidades, titulos, objetos de conhecimento e textos-base para montar planos CDP sem depender dos PDFs comuns.",
        "observacoes": "E um dos nucleos do modo CDP. Tambem monta metodologia, acompanhamento e acessibilidade especificos para esse fluxo.",
    },
    {
        "arquivo": "cdp_em_docx.py",
        "funcao": "Reescreve planos CDP prontos em DOCX, especialmente no fluxo contextual de Matematica e Ensino Medio.",
        "fluxo": "Entra quando o sistema precisa corrigir ou reformatar um documento final ja existente, em vez de gerar tudo do zero.",
        "observacoes": "Aproveita partes do preenchimento DOCX e da logica de lote para ajustar tema, material e celulas do plano.",
    },
    {
        "arquivo": "database.py",
        "funcao": "Controla o banco SQLite do sistema.",
        "fluxo": "Salva e consulta professores, turmas, disciplinas, modelos vinculados e historico dos planos gerados.",
        "observacoes": "E a base administrativa do PLANOS_LUAN. Sem ele, o sistema perderia persistencia de cadastro e historico.",
    },
    {
        "arquivo": "disciplinas.py",
        "funcao": "Define a lista oficial de disciplinas, bimestres e configuracoes de modo.",
        "fluxo": "Ajuda a interface e a logica interna a saber se uma disciplina usa PDF, CDP, CDP Fundamental ou outro comportamento especial.",
        "observacoes": "Tambem centraliza regras simples como 'esta disciplina exige PDF?' e 'esta disciplina pertence ao universo CDP?'.",
    },
    {
        "arquivo": "helpers.py",
        "funcao": "Reune utilitarios pequenos usados em varias partes do sistema.",
        "fluxo": "Converte horarios para o formato do plano, transforma listas em texto e gera relatorios simples de conferencia.",
        "observacoes": "Nao e um modulo de regra central, mas evita repeticao de pequenas rotinas espalhadas pelo projeto.",
    },
    {
        "arquivo": "ia.py",
        "funcao": "Integra o sistema com modelos de IA, como OpenAI e Gemini.",
        "fluxo": "Monta o prompt, envia o texto do PDF para o provedor escolhido e normaliza a resposta em tema, aprendizagem e metodologia.",
        "observacoes": "Tambem aplica filtros de qualidade na saida da IA para evitar tema truncado, metodologia quebrada ou texto generico demais.",
    },
    {
        "arquivo": "inteligencia_local.py",
        "funcao": "Executa analise e geracao local sem depender de IA externa.",
        "fluxo": "Classifica conteudo, extrai partes relevantes do PDF e sugere metodologia com base em heuristicas e na base de conhecimento interna.",
        "observacoes": "Funciona como um motor alternativo ou complementar a IA online, ajudando quando se quer reduzir dependencia externa.",
    },
    {
        "arquivo": "lote.py",
        "funcao": "E o coracao da geracao de aulas a partir dos PDFs.",
        "fluxo": "Extrai texto dos materiais, limpa titulos, detecta perfil da disciplina, monta tema, metodologia, aprendizagem, acompanhamento e acessibilidade para cada aula.",
        "observacoes": "E o maior e mais importante modulo pedagogico do projeto hoje. Reune boa parte das regras finas de interpretacao dos materiais digitais.",
    },
    {
        "arquivo": "metodologia_texto.py",
        "funcao": "Faz ajustes linguísticos em textos metodologicos.",
        "fluxo": "Aparece no pos-processamento para deixar a redacao mais natural e coerente, como ajuste de verbos para infinitivo.",
        "observacoes": "Embora pequeno, ajuda no polimento do texto final que vai para o Word.",
    },
    {
        "arquivo": "modelos_docx.py",
        "funcao": "Escolhe qual template DOCX central deve ser usado.",
        "fluxo": "Decide entre modelos como Egle, Padre ou CDP com base na disciplina, escola, componente curricular e contexto.",
        "observacoes": "Ajuda o sistema a manter padrao visual correto sem obrigar o usuario a escolher manualmente toda vez.",
    },
    {
        "arquivo": "professores_planos.py",
        "funcao": "Lê, diagnostica, cria e atualiza planos-modelo dentro das pastas dos professores.",
        "fluxo": "Extrai cabecalho, datas e horarios dos DOCX existentes e cruza isso com o cadastro salvo no banco.",
        "observacoes": "E a ponte entre os arquivos reais dos professores e a estrutura interna do sistema. Tambem ajuda na padronizacao dos modelos.",
    },
    {
        "arquivo": "projeto_vida_escopo.py",
        "funcao": "Lê e interpreta o escopo de Projeto de Vida.",
        "fluxo": "Busca o item correto por turma, bimestre e numero da aula para enriquecer tema e aprendizagem desse componente.",
        "observacoes": "Existe porque Projeto de Vida tem uma logica propria e nao pode ser tratado como uma disciplina comum em todos os casos.",
    },
    {
        "arquivo": "prompts_por_disciplina.py",
        "funcao": "Guarda orientacoes de prompt e instrucoes especificas por disciplina.",
        "fluxo": "Alimenta o modulo de IA com direcionamentos diferentes para Portugues, Matematica, Projeto de Vida e outras areas.",
        "observacoes": "Serve para personalizar o comportamento da IA sem misturar essas regras diretamente no restante do codigo.",
    },
    {
        "arquivo": "qualidade_metodologica.py",
        "funcao": "Revisa, sanitiza e melhora a qualidade pedagogica e textual da metodologia.",
        "fluxo": "Corrige problemas de linguagem, identifica contexto metodologico, extrai conceito central e ajuda a naturalizar o texto final.",
        "observacoes": "Funciona como uma camada de controle de qualidade antes da aula seguir para o documento final.",
    },
    {
        "arquivo": "referencias_metodologia.py",
        "funcao": "Carrega referencias metodologicas salvas em arquivos do projeto.",
        "fluxo": "Seleciona materiais de apoio por disciplina e turma para reforcar a geracao das metodologias, especialmente no uso da IA.",
        "observacoes": "E uma fonte de memoria pedagogica do sistema, trazendo consistencia entre planos da mesma area.",
    },
    {
        "arquivo": "validador_plano.py",
        "funcao": "Valida a qualidade minima das aulas geradas.",
        "fluxo": "Confere se ha tema, metodologia suficiente, aprendizagem, acompanhamento e acessibilidade antes de considerar o plano aceitavel.",
        "observacoes": "Ajuda a bloquear saidas fracas ou incompletas, evitando que um plano ruim avance sem revisao.",
    },
]


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_page_margins(section) -> None:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def _style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size, color, bold in [
        ("Title", 22, RGBColor(25, 67, 115), True),
        ("Heading 1", 15, RGBColor(31, 78, 121), True),
        ("Heading 2", 11.5, RGBColor(31, 78, 121), True),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold


def _add_label_paragraph(doc: Document, label: str, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(texto)


def build_docx() -> Path:
    doc = Document()
    _set_page_margins(doc.sections[0])
    _style_document(doc)

    titulo = doc.add_paragraph(style="Title")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.add_run("Estrutura Atual da Pasta core do PLANOS_LUAN")

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.paragraph_format.space_after = Pt(14)
    subtitulo.add_run(
        "Documento de referencia em linguagem simples, baseado no comportamento atual do codigo."
    ).italic = True

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    intro.add_run(
        "Objetivo: explicar o papel de cada arquivo mostrado na pasta "
    )
    intro.add_run("core").bold = True
    intro.add_run(
        ", indicando para que ele serve no sistema, em que momento entra no fluxo e qual sua importancia pratica."
    )

    doc.add_heading("Visao Geral", level=1)
    for texto in [
        "A pasta core concentra a maior parte da inteligencia de negocio do PLANOS_LUAN.",
        "Alguns arquivos cuidam do cadastro e dos modelos em Word; outros tratam da leitura pedagogica dos PDFs, do modo CDP, do uso de IA e da validacao da qualidade final.",
        "Hoje o modulo mais central para geracao de aulas comuns a partir dos PDFs e o lote.py.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(texto)

    doc.add_heading("Mapa Rapido dos Arquivos", level=1)
    tabela = doc.add_table(rows=1, cols=3)
    tabela.style = "Table Grid"
    tabela.autofit = True
    cab = tabela.rows[0].cells
    cab[0].text = "Arquivo"
    cab[1].text = "Papel principal"
    cab[2].text = "Peso no sistema"
    for cell in cab:
        _set_cell_shading(cell, "D9EAF7")

    pesos = {
        "__init__.py": "Tecnico",
        "helpers.py": "Apoio",
        "metodologia_texto.py": "Apoio",
        "base_conhecimento.py": "Referencia",
        "prompts_por_disciplina.py": "Referencia",
        "referencias_metodologia.py": "Referencia",
        "disciplinas.py": "Estrutural",
        "modelos_docx.py": "Estrutural",
        "calendario.py": "Estrutural",
        "database.py": "Estrutural",
        "professores_planos.py": "Estrutural",
        "avaliacao.py": "Pedagogico",
        "qualidade_metodologica.py": "Pedagogico",
        "validador_plano.py": "Pedagogico",
        "inteligencia_local.py": "Motor",
        "ia.py": "Motor",
        "lote.py": "Nucleo central",
        "cdp.py": "Nucleo CDP",
        "cdp_em_docx.py": "Ajuste final",
        "projeto_vida_escopo.py": "Especializado",
    }
    for item in ITENS:
        row = tabela.add_row().cells
        row[0].text = item["arquivo"]
        row[1].text = item["funcao"]
        row[2].text = pesos.get(item["arquivo"], "Importante")

    doc.add_paragraph()
    doc.add_heading("Explicacao Detalhada Arquivo por Arquivo", level=1)

    for item in ITENS:
        doc.add_heading(item["arquivo"], level=2)
        _add_label_paragraph(doc, "Funcao principal", item["funcao"])
        _add_label_paragraph(doc, "Quando entra no fluxo", item["fluxo"])
        _add_label_paragraph(doc, "Observacoes", item["observacoes"])

    doc.add_paragraph()
    doc.add_heading("Leitura Final da Arquitetura", level=1)
    conclusoes = [
        "Se a pergunta for 'onde o plano realmente nasce?', a resposta principal hoje e lote.py.",
        "Se a pergunta for 'onde ficam cadastros, vinculos e historico?', a resposta principal e database.py.",
        "Se a pergunta for 'onde o sistema entende o modo CDP?', a resposta principal e cdp.py.",
        "Se a pergunta for 'onde a IA e preparada e revisada?', os destaques sao ia.py, prompts_por_disciplina.py e qualidade_metodologica.py.",
        "Se a pergunta for 'onde os modelos DOCX dos professores conversam com o sistema?', o ponto-chave e professores_planos.py.",
    ]
    for texto in conclusoes:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(texto)

    doc.add_section(WD_SECTION.NEW_PAGE)
    _set_page_margins(doc.sections[-1])
    doc.add_heading("Observacao Importante", level=1)
    doc.add_paragraph(
        "Esta documentacao descreve a estrutura atual observada no codigo no momento da analise. "
        "Se algum arquivo ganhar novas responsabilidades no futuro, o texto deve ser revisado para continuar fiel ao sistema."
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    caminho = build_docx()
    print(str(caminho))
