from docx import Document

from core import lote
from core.referencias_lingua_inglesa import (
    localizar_docx_referencia_lingua_inglesa,
    referencia_lingua_inglesa_por_pdf,
    titulos_referencia_lingua_inglesa_por_docx,
)


def _criar_docx_referencia_ingles(caminho):
    doc = Document()
    doc.add_paragraph("METODOLOGIAS - LINGUA INGLESA - 7 ANO")

    doc.add_paragraph("AULA 1 — Leisure-time stories - Part 1")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Retomar simple past em narrativas pessoais.")
    doc.add_paragraph("Foco no conteúdo: Explorar marcadores temporais em inglês.")
    doc.add_paragraph("Na prática: Ler uma narrativa curta e localizar verbos no passado.")
    doc.add_paragraph("Encerramento: Socializar uma frase em inglês sobre leisure time.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se identificam marcadores de tempo em inglês.")
    doc.add_paragraph("☑ Observar se reconhecem verbos regulares e irregulares.")
    doc.add_paragraph("☑ Conferir se registram uma frase adequada ao contexto.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Oferecer glossário bilíngue com palavras-chave.")
    doc.add_paragraph("☑ Disponibilizar modelo de frase em inglês e português.")
    doc.add_paragraph("☑ Permitir resposta oral mediada antes do registro escrito.")

    doc.add_paragraph("AULA 14 — Trilha de aprendizagem individual")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Retomar registros anteriores da trilha e conectar ao novo foco do dia.")
    doc.add_paragraph("Foco no conteúdo: Desenvolver a trilha de forma progressiva com comandos simples em inglês.")
    doc.add_paragraph("Na prática: Resolver atividades, comparar respostas e registrar dúvidas persistentes.")
    doc.add_paragraph("Encerramento: Registrar uma síntese parcial para orientar a continuidade da sequência.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se compreendem o tema central da aula.")
    doc.add_paragraph("☑ Conferir se as produções apresentam clareza e retomada dos conceitos.")
    doc.add_paragraph("☑ Observar participação, registros e justificativas nas atividades.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Organizar apoio em duplas para favorecer compreensão.")
    doc.add_paragraph("☑ Permitir registros por tópicos, frases curtas ou resposta oral.")
    doc.add_paragraph("☑ Realizar retomadas coletivas dos comandos em inglês.")
    doc.save(caminho)


def test_referencia_lingua_inglesa_le_docx_da_pasta_do_pdf(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Lingua_Inglesa_7_Ano.docx"
    caminho_pdf = tmp_path / "AULA_14 - Trilha de aprendizagem individual.pdf"
    _criar_docx_referencia_ingles(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_lingua_inglesa_por_pdf(caminho_pdf, "14")

    assert localizar_docx_referencia_lingua_inglesa(caminho_pdf) == caminho_docx
    assert referencia["numero"] == 14
    assert referencia["titulo"] == "Trilha de aprendizagem individual"
    assert [etapa["titulo"] for etapa in referencia["metodologia"]][:3] == [
        "Para começar",
        "Foco no conteúdo",
        "Na prática",
    ]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_titulos_lingua_inglesa_por_docx(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Lingua_Inglesa_7_Ano.docx"
    _criar_docx_referencia_ingles(caminho_docx)

    titulos = titulos_referencia_lingua_inglesa_por_docx(caminho_docx)

    assert titulos == {
        1: "Leisure-time stories - Part 1",
        14: "Trilha de aprendizagem individual",
    }


def test_lingua_inglesa_resultado_usa_docx_para_trilha(tmp_path, monkeypatch):
    caminho_docx = tmp_path / "Metodologias_Lingua_Inglesa_7_Ano.docx"
    caminho_pdf = tmp_path / "AULA_14 - Trilha de aprendizagem individual.pdf"
    _criar_docx_referencia_ingles(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")
    monkeypatch.setattr(
        lote,
        "_extrair_texto_pdf",
        lambda caminho: (
            "Língua Inglesa\nTrilha de aprendizagem individual\n"
            "Plataforma EF: produção e compreensão oral e escrita."
        ),
    )

    aula = lote._aula_por_pdf(
        str(caminho_pdf),
        "Língua Inglesa",
        "7º ANO A",
        "3º Bimestre",
        usar_ia=False,
        provedor_ia="",
    )

    assert aula["tema"] == "Trilha de aprendizagem individual"
    assert aula["material"] == "AULA 14 - Trilha de aprendizagem individual"
    assert aula["origem_metodologia"] == "docx_referencia_lingua_inglesa"
    assert "trilha" in aula["metodologia"][1]["texto"].lower()
    assert len(aula["acompanhamento"]) == 3
    assert len(aula["acessibilidade"]) == 3
