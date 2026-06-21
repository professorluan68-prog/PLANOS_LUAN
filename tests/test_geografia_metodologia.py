from docx import Document

from core.lote import _montar_resultado_aula_local
from core.referencias_geografia import (
    localizar_docx_referencia_geografia,
    referencia_geografia_por_pdf,
    titulos_referencia_geografia_por_docx,
)


def _criar_docx_referencia_geografia(caminho):
    doc = Document()
    doc.add_paragraph("AULA 1 — Impacto ambiental")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph(
        "Para comecar: Retomar imagens de transformacoes na paisagem e levantar hipoteses da turma sobre impactos ambientais."
    )
    doc.add_paragraph(
        "Foco no conteudo: Orientar a leitura de mapa, dados e exemplos de alteracoes no territorio, relacionando sociedade e natureza."
    )
    doc.add_paragraph(
        "Na pratica: Solicitar que os estudantes comparem situacoes e registrem causas, consequencias e possiveis acoes de mitigacao."
    )
    doc.add_paragraph(
        "Encerramento: Socializar as analises e organizar no quadro uma sintese sobre impacto ambiental."
    )
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se os estudantes identificam causas e consequencias dos impactos ambientais.")
    doc.add_paragraph("☑ Observar a leitura de mapas, dados e exemplos territoriais durante a atividade.")
    doc.add_paragraph("☑ Conferir os registros com relacoes entre sociedade, natureza e mitigacao.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Disponibilizar palavras-chave sobre paisagem, impacto ambiental e mitigacao.")
    doc.add_paragraph("☑ Oferecer leitura guiada de mapas e imagens com perguntas orientadoras.")
    doc.add_paragraph("☑ Permitir resposta oral mediada antes do registro escrito.")
    doc.save(caminho)


def test_referencia_geografia_le_docx_da_pasta_do_pdf(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Geografia_1_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA_01 - Impacto ambiental.pdf"
    _criar_docx_referencia_geografia(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_geografia_por_pdf(caminho_pdf, "1")

    assert referencia["titulo"] == "Impacto ambiental"
    assert [etapa["titulo"] for etapa in referencia["metodologia"]] == [
        "Para comecar",
        "Foco no conteudo",
        "Na pratica",
        "Encerramento",
    ]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_titulos_referencia_geografia_por_docx_aceita_travessao(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Geografia_2_Ano_Ensino_Medio.docx"
    _criar_docx_referencia_geografia(caminho_docx)

    titulos = titulos_referencia_geografia_por_docx(caminho_docx)

    assert titulos == {1: "Impacto ambiental"}


def test_referencia_geografia_localiza_docx_da_serie(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Geografia_1_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA_01 - Impacto ambiental.pdf"
    _criar_docx_referencia_geografia(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    escolhido = localizar_docx_referencia_geografia(caminho_pdf)

    assert escolhido == caminho_docx


def test_geografia_resultado_local_usa_docx_sem_trocar_titulo_oficial(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Geografia_1_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA_01 - Impacto ambiental.pdf"
    _criar_docx_referencia_geografia(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    resultado = _montar_resultado_aula_local(
        texto="Texto qualquer do PDF que nao deve prevalecer sobre o DOCX de referencia.",
        tema="Titulo vindo da planilha",
        material_digital="AULA 1 - Titulo vindo da planilha",
        numero_aula="1",
        disciplina_base="Geografia",
        turma="1a serie A",
        provedor_ia="",
        perfil="geografia",
        contexto_metodologico="",
        indice_aula=0,
        total_aulas=1,
        modalidade_eja_ativa=False,
        metodologia_fixa_pdf=[],
        aprendizagem_pv="",
        objetivos_orientacao=[],
        aprendizagem_orientacao="",
        usar_ia=False,
        ia_erro="",
        caminho_pdf=str(caminho_pdf),
    )

    assert resultado["tema"] == "Titulo vindo da planilha"
    assert resultado["material"] == "AULA 1 - Titulo vindo da planilha"
    assert resultado["origem_metodologia"] == "docx_referencia_geografia"
    assert "impactos ambientais" in resultado["metodologia"][0]["texto"].lower()
    assert len(resultado["acompanhamento"]) == 3
    assert len(resultado["acessibilidade"]) == 3
