from datetime import date
from io import BytesIO

import core.lote as lote
from docx import Document

from core.lib.classificador import detectar_tipo_aula, perfil_disciplina
from core.qualidade_metodologica import normalizar_texto
from docx_generator.preencher import preencher_documento


TEXTO_PORTUGUES_EM_DIARIO = (
    "Lingua Portuguesa\n"
    "O genero diario pessoal:\n"
    "reflexoes do cotidiano - Parte 1\n"
    "3o bimestre Ensino\n"
    "Aula 10 Medio\n"
    "Praticas de leitura e analises de fragmentos literarios sobre o genero diario pessoal.\n"
    "Identificar as caracteristicas do genero diario pessoal.\n"
    "Analisar as tematicas e a estrutura dos diarios pessoais, relacionando-as ao contexto historico e social dos autores.\n"
    "Coordenacao e subordinacao (revisao).\n"
    "Para comecar\n"
    "COM SUAS PALAVRAS\n"
    "Voce ja escreveu um diario ou ja leu o diario escrito por outra pessoa?\n"
    "De que modo se escreve um diario?\n"
    "Foco no conteudo\n"
    "Retomar marcas do genero diario pessoal, o uso da primeira pessoa e a organizacao de reflexoes do cotidiano.\n"
    "Na pratica\n"
    "Ler um fragmento de diario pessoal e registrar as caracteristicas observadas.\n"
)

TEXTO_MATEMATICA_PARABOLA = (
    "Matematica\n"
    "Parabola: grafico de uma\n"
    "funcao quadratica - Parte 2\n"
    "3o bimestre Ensino\n"
    "Aula 10 Medio\n"
    "Construcao do esboco grafico de uma funcao quadratica.\n"
    "Resolver situacoes sobre a construcao e a interpretacao de graficos de uma funcao quadratica.\n"
    "Na pratica\n"
    "Atividade 1 Veja no livro! 10 minutos\n"
    "TODO MUNDO ESCREVE\n"
    "Um engenheiro esta projetando a trajetoria em forma de arco parabolico.\n"
)


def _modelo_minimo():
    doc = Document()
    cabecalho = doc.add_table(rows=4, cols=9)
    cabecalho.rows[0].cells[0].text = "PLANO DE AULAS"
    cabecalho.rows[1].cells[2].text = "PROFESSOR"
    cabecalho.rows[1].cells[3].text = "COMPONENTE CURRICULAR"
    cabecalho.rows[1].cells[6].text = "TURMA"
    cabecalho.rows[1].cells[7].text = "MES"
    cabecalho.rows[1].cells[8].text = "BIMESTRE"
    cabecalho.rows[3].cells[0].text = "SEMANA"
    cabecalho.rows[3].cells[1].text = "01/09 a 05/09"
    cabecalho.rows[3].cells[2].text = "AULAS PREVISTAS NA SEMANA"

    aulas = doc.add_table(rows=2, cols=6)
    aulas.rows[0].cells[0].text = "AULA SEMANAL (Data e Horario)"
    aulas.rows[0].cells[1].text = "NUMERO E TITULO DO MATERIAL DIGITAL"
    aulas.rows[0].cells[2].text = "APRENDIZAGEM ESSENCIAL*"
    aulas.rows[0].cells[3].text = "DESENVOLVIMENTO"
    aulas.rows[0].cells[4].text = "ACOMPANHAMENTO DA APRENDIZAGEM"
    aulas.rows[0].cells[5].text = "ACESSIBILIDADE"

    saida = BytesIO()
    doc.save(saida)
    saida.seek(0)
    return saida


def test_classificacao_portugues_generico_usa_turma_para_identificar_ensino_medio():
    assert perfil_disciplina("Lingua Portuguesa", turma="1 ANO") == "lingua_portuguesa_em"
    assert (
        detectar_tipo_aula(
            TEXTO_PORTUGUES_EM_DIARIO,
            "O genero diario pessoal reflexoes do cotidiano - Parte 1",
            "Lingua Portuguesa",
            turma="1 ANO",
        )
        == "genero_textual"
    )


def test_fluxo_principal_portugues_em_nao_contamina_diario_com_biografia(monkeypatch):
    monkeypatch.setattr(lote, "_extrair_texto_pdf", lambda caminho: TEXTO_PORTUGUES_EM_DIARIO)

    aula = lote._aula_por_pdf(
        "AULA 10.pdf",
        disciplina="Lingua Portuguesa",
        turma="1 ANO",
        bimestre="3o Bimestre",
        usar_ia=False,
        provedor_ia="",
    )

    texto_metodologia = normalizar_texto(" ".join(item["texto"] for item in aula["metodologia"]))

    assert normalizar_texto(aula["tema"]).startswith("o genero diario pessoal")
    assert "biografia" not in texto_metodologia
    assert "trajetoria da pessoa biografada" not in texto_metodologia
    assert "diario" in texto_metodologia

    doc = Document(
        preencher_documento(
            _modelo_minimo(),
            [
                {
                    **aula,
                    "data": date(2026, 9, 2),
                    "horario": "14h40 - 15h30",
                }
            ],
            professor="Teste",
            disciplina="Lingua Portuguesa",
            turma="1 ANO",
            mes="SETEMBRO",
            bimestre="3o Bimestre",
            aulas_previstas_manual="1",
        )
    )
    texto_doc = normalizar_texto(
        "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    )
    assert "diario pessoal" in texto_doc
    assert "trajetoria da pessoa biografada" not in texto_doc
    assert "leitura orientada da biografia" not in texto_doc


def test_fluxo_principal_matematica_real_mantem_tema_e_sem_fragmentos(monkeypatch):
    monkeypatch.setattr(lote, "_extrair_texto_pdf", lambda caminho: TEXTO_MATEMATICA_PARABOLA)

    aula = lote._aula_por_pdf(
        "AULA_10.pdf",
        disciplina="Matematica",
        turma="1 ANO",
        bimestre="3o Bimestre",
        usar_ia=False,
        provedor_ia="",
    )

    textos = [normalizar_texto(item["texto"]) for item in aula["metodologia"] if isinstance(item, dict)]

    assert "parabola" in normalizar_texto(aula["tema"])
    assert "parte 2" in normalizar_texto(aula["material"])
    assert any("funcao quadratica" in texto or "parabola" in texto for texto in textos)
    assert all(not texto.startswith(("que os", "que as", "que cada", "para que")) for texto in textos)
