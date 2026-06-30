from pathlib import Path

from docx import Document

from core import lote
from core.lote import _montar_resultado_aula_local
from core.referencias_matematica import (
    localizar_docx_referencia_matematica,
    referencia_matematica_por_pdf,
)


def _criar_docx_matematica(caminho: Path, titulo_aula: str, etapa_inicial: str) -> None:
    doc = Document()
    doc.add_paragraph(titulo_aula)
    doc.add_paragraph("METODOLOGIA")
    doc.add_paragraph(f"Para começar: {etapa_inicial}")
    doc.add_paragraph("Foco no conteúdo: Retomar os conceitos matemáticos centrais da aula.")
    doc.add_paragraph("Na prática: Resolver e discutir as atividades propostas no material.")
    doc.add_paragraph("Encerramento: Sistematizar as estratégias utilizadas pela turma.")
    doc.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    doc.add_paragraph("☑ Verificar se os estudantes reconhecem os dados do problema.")
    doc.add_paragraph("☑ Observar se organizam estratégias coerentes de resolução.")
    doc.add_paragraph("☑ Acompanhar se justificam os procedimentos utilizados.")
    doc.add_paragraph("ACESSIBILIDADE")
    doc.add_paragraph("☑ Disponibilizar exemplos comentados e palavras-chave da aula.")
    doc.add_paragraph("☑ Oferecer apoio com registro passo a passo no quadro.")
    doc.add_paragraph("☑ Permitir explicação oral mediada durante a resolução.")
    doc.save(caminho)


def test_localizar_docx_referencia_matematica_prioriza_arquivo_novo(tmp_path):
    pdf = tmp_path / "AULA_01 - Retomada de função afim.pdf"
    pdf.write_bytes(b"pdf")

    docx_antigo = tmp_path / "Metodologias_Matematica_1_Ano_Ensino_Medio.docx"
    docx_novo = tmp_path / "Metodologias_Matematica_1_Ano_Ensino_Medio_NOVO.docx"
    _criar_docx_matematica(
        docx_antigo,
        "AULA 1 — Retomada de função afim",
        "Texto antigo.",
    )
    _criar_docx_matematica(
        docx_novo,
        "AULA 1 — Retomada de função afim",
        "Texto novo.",
    )

    localizado = localizar_docx_referencia_matematica(pdf)

    assert localizado == docx_novo.resolve()


def test_referencia_matematica_por_pdf_usa_conteudo_do_docx(tmp_path):
    pdf = tmp_path / "AULA_05 - Estratégias de modelagem algébrica.pdf"
    pdf.write_bytes(b"pdf")

    docx = tmp_path / "Metodologias_Matematica_1_Ano_Ensino_Medio.docx"
    _criar_docx_matematica(
        docx,
        "AULA 5 — Estratégias de modelagem algébrica com a função afim",
        "Retomar o contexto da função afim por meio de uma situação-problema inicial.",
    )

    referencia = referencia_matematica_por_pdf(
        pdf,
        numero_aula=5,
        tema="Estratégias de modelagem algébrica com a função afim",
    )

    assert referencia is not None
    assert referencia["titulo"] == "Estratégias de modelagem algébrica com a função afim"
    assert referencia["metodologia"][0]["titulo"] == "Para começar"
    assert "situação-problema inicial" in referencia["metodologia"][0]["texto"]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_matematica_resultado_local_usa_docx_sem_trocar_tema_pdf(tmp_path, monkeypatch):
    pdf = tmp_path / "AULA_05 - Estratégias de modelagem algébrica.pdf"
    pdf.write_bytes(b"pdf")

    docx = tmp_path / "Metodologias_Matematica_1_Ano_Ensino_Medio.docx"
    _criar_docx_matematica(
        docx,
        "AULA 5 — Estratégias de modelagem algébrica com a função afim",
        "Apresentar uma situação inicial para retomar a relação entre grandezas.",
    )

    monkeypatch.setattr(
        lote._extrator_lib,
        "extrair",
        lambda *args, **kwargs: {
            "habilidade": "EM13MAT101",
            "conceito_extraido": "Função afim",
            "recursos_detectados": ["Quadro", "Caderno"],
        },
    )

    resultado = _montar_resultado_aula_local(
        texto="Texto da aula",
        tema="Tema do PDF preservado",
        material_digital="AULA 5 - Tema do PDF preservado",
        numero_aula="5",
        disciplina_base="Matemática",
        turma="1º ANO A",
        provedor_ia="",
        perfil="matematica",
        contexto_metodologico="regular",
        indice_aula=0,
        total_aulas=1,
        modalidade_eja_ativa=False,
        metodologia_fixa_pdf=[],
        aprendizagem_pv="",
        objetivos_orientacao=[],
        aprendizagem_orientacao="",
        usar_ia=False,
        ia_erro="",
        caminho_pdf=str(pdf),
    )

    assert resultado["tema"] == "Tema do PDF preservado"
    assert resultado["origem_metodologia"] == "docx_referencia_matematica"
    assert "relação entre grandezas" in resultado["metodologia"][0]["texto"]
    assert len(resultado["acompanhamento"]) == 3
    assert len(resultado["acessibilidade"]) == 3
