import pytest
import core.referencias_metodologia as referencias
from core.referencias_metodologia import (
    carregar_referencia_metodologica,
    diagnosticar_referencia_metodologica,
    listar_referencias_disponiveis,
)


def _limpar_cache_referencias():
    diagnosticar_referencia_metodologica.cache_clear()
    carregar_referencia_metodologica.cache_clear()


def test_referencia_aceita_somente_raizes_oficiais(tmp_path, monkeypatch):
    raiz_oficial = tmp_path / "PDF_AULAS"
    raiz_oficial.mkdir()
    referencia_oficial = raiz_oficial / "guia.md"
    referencia_oficial.write_text("Conteúdo oficial", encoding="utf-8")
    referencia_externa = tmp_path / "referencia_antiga.md"
    referencia_externa.write_text("Conteúdo antigo", encoding="utf-8")

    monkeypatch.setattr(referencias, "PASTAS_BUSCA", (raiz_oficial,))

    assert referencias.resolver_caminho_referencia(referencia_oficial) == referencia_oficial
    assert referencias.resolver_caminho_referencia(referencia_externa) is None


def test_diagnostico_informa_ausencia_de_referencia_oficial(tmp_path, monkeypatch):
    raiz_oficial = tmp_path / "REFERENCIAS_METODOLOGICAS"
    raiz_oficial.mkdir()
    monkeypatch.setattr(referencias, "PASTAS_BUSCA", (raiz_oficial,))
    monkeypatch.setattr(referencias, "_buscar_metodologia_automatica", lambda *args: ())
    monkeypatch.setattr(referencias, "_arquivos_novos_para_disciplina", lambda *args: ())
    monkeypatch.setattr(referencias, "_arquivos_para_disciplina", lambda *args: ("ausente.md",))
    _limpar_cache_referencias()

    diagnostico = diagnosticar_referencia_metodologica("História", "8º ano A")

    assert diagnostico.texto == ""
    assert diagnostico.arquivos_ausentes == ("ausente.md",)
    assert "Referência metodológica oficial não encontrada" in diagnostico.aviso


def test_diagnostico_carrega_referencia_oficial(tmp_path, monkeypatch):
    raiz_oficial = tmp_path / "REFERENCIAS_METODOLOGICAS"
    raiz_oficial.mkdir()
    (raiz_oficial / "guia.md").write_text("# Guia\nConteúdo oficial", encoding="utf-8")
    monkeypatch.setattr(referencias, "PASTAS_BUSCA", (raiz_oficial,))
    monkeypatch.setattr(referencias, "_buscar_metodologia_automatica", lambda *args: ())
    monkeypatch.setattr(referencias, "_arquivos_novos_para_disciplina", lambda *args: ())
    monkeypatch.setattr(referencias, "_arquivos_para_disciplina", lambda *args: ("guia.md",))
    _limpar_cache_referencias()

    diagnostico = diagnosticar_referencia_metodologica("História", "8º ano A")

    assert "Conteúdo oficial" in diagnostico.texto
    assert diagnostico.aviso == ""
    assert diagnostico.arquivos_encontrados == (str(raiz_oficial / "guia.md"),)


@pytest.mark.skip(reason="Referencias removidas por solicitacao do usuario")
def test_carrega_referencia_por_disciplina():
    referencia = carregar_referencia_metodologica("Língua Portuguesa", "7º ano A")

    assert "REGRAS FIXAS DO SISTEMA" in referencia
    assert "Língua Portuguesa" in referencia
    assert "Hora da leitura" in referencia
    assert "Não invente técnicas" in referencia


@pytest.mark.skip(reason="Referencias removidas")
def test_referencias_disponiveis_incluem_disciplinas_implantadas():
    referencias = listar_referencias_disponiveis()

    assert "historia" in referencias
    assert "arte" in referencias
    assert "projeto de vida" in referencias
    assert "ciencias" in referencias


@pytest.mark.skip(reason="Referencias removidas")
def test_referencia_interdisciplinar_entra_como_complemento_seguro():
    referencia = carregar_referencia_metodologica("História", "8º ano A")

    assert "REFERÊNCIA INTERDISCIPLINAR COMPLEMENTAR" in referencia
    assert "não presentes nos slides" in referencia
    assert "Riscos de Confusão no Código Python" not in referencia


@pytest.mark.skip(reason="Referencias removidas")
def test_referencia_ciencias_prioriza_nova_analise_dos_anos_finais():
    referencia = carregar_referencia_metodologica("Ciencias", "8 ano A").lower()

    assert "não alucinar materiais" in referencia
    assert "modelagem científica" in referencia
    assert "limitação de modelos" in referencia


def test_ler_docx_com_tabelas_e_paragrafos(tmp_path):
    import docx
    from core.referencias_metodologia import _ler_docx

    doc_file = tmp_path / "temp_ref.docx"
    doc = docx.Document()
    doc.add_paragraph("Este é um parágrafo de teste.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Célula 1A"
    table.cell(0, 1).text = "Célula 1B"
    table.cell(1, 0).text = "Célula 2A"
    table.cell(1, 1).text = "Célula 2B"

    doc.save(doc_file)

    texto = _ler_docx(doc_file)

    assert "Este é um parágrafo de teste." in texto
    assert "Célula 1A | Célula 1B" in texto
    assert "Célula 2A | Célula 2B" in texto
