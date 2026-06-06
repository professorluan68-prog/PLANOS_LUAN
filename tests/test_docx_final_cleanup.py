from datetime import date
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from docx_generator.preencher import preencher_documento


def _modelo_com_semanas(semanas, linhas_aulas=1):
    doc = Document()
    for semana in semanas:
        cabecalho = doc.add_table(rows=4, cols=9)
        cabecalho.rows[0].cells[0].text = "PLANO DE AULAS"
        cabecalho.rows[1].cells[2].text = "PROFESSOR"
        cabecalho.rows[1].cells[3].text = "COMPONENTE CURRICULAR"
        cabecalho.rows[1].cells[6].text = "TURMA"
        cabecalho.rows[1].cells[7].text = "MES"
        cabecalho.rows[1].cells[8].text = "BIMESTRE"
        cabecalho.rows[3].cells[0].text = "SEMANA"
        cabecalho.rows[3].cells[1].text = semana
        cabecalho.rows[3].cells[2].text = "AULAS PREVISTAS NA SEMANA"

        aulas = doc.add_table(rows=1 + linhas_aulas, cols=6)
        aulas.rows[0].cells[0].text = "AULA SEMANAL (Data e Horario)"
        aulas.rows[0].cells[1].text = "NUMERO E TITULO DO MATERIAL DIGITAL"
        aulas.rows[0].cells[2].text = "APRENDIZAGEM ESSENCIAL*"
        aulas.rows[0].cells[3].text = "DESENVOLVIMENTO"
        aulas.rows[0].cells[4].text = "ACOMPANHAMENTO DA APRENDIZAGEM"
        aulas.rows[0].cells[5].text = "ACESSIBILIDADE"

    saida = BytesIO()
    doc.save(saida)
    saida.seek(0)
    return saida


def _aula(data, tema="Tema da aula"):
    return {
        "data": data,
        "horario": "14h40 - 16h40",
        "material": f"AULA 5 - {tema}",
        "aprendizagem": "Habilidade: manter BNCC como veio.",
        "metodologia": [
            {
                "titulo": "Para comecar",
                "texto": (
                    "Retomar registros ja construidos, analisando consequencias possiveis, "
                    "sistematizacao, servico e preco."
                ),
            }
        ],
        "acompanhamento": ["Observar registros da aula."],
        "acessibilidade": ["Oferecer apoio visual."],
    }


def _gerar(aulas):
    modelo = _modelo_com_semanas(
        ["01/06 a 05/06", "08/06 a 12/06", "15/06 a 19/06", "22/06 a 26/06", "29/06 a 03/07"]
    )
    return Document(
        preencher_documento(
            modelo,
            aulas,
            professor="ADRIANA",
            disciplina="Educacao Financeira",
            turma="7o ANO A",
            mes="JUNHO",
            bimestre="2o Bimestre",
            aulas_previstas_manual="2",
        )
    )


def test_remove_apenas_semana_final_vazia():
    doc = _gerar([_aula("05/06"), _aula("12/06"), _aula("19/06"), _aula("26/06")])

    texto = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert len(doc.tables) == 8
    assert "22/06 a 26/06" in texto
    assert "29/06 a 03/07" not in texto


def test_preserva_semana_vazia_no_meio_quando_ha_aula_depois():
    doc = _gerar([_aula("05/06"), _aula("19/06"), _aula("03/07")])

    texto = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert len(doc.tables) == 10
    assert "08/06 a 12/06" in texto
    assert "29/06 a 03/07" in texto
    assert "03/07" in texto


def test_distribui_aulas_na_semana_correta_quando_data_veio_do_app():
    doc = _gerar(
        [
            _aula(date(2026, 6, 3), "Aula semana 1"),
            _aula(date(2026, 6, 10), "Aula semana 2"),
            _aula(date(2026, 6, 17), "Aula semana 3"),
        ]
    )

    tabela_semana_1 = doc.tables[1]
    tabela_semana_2 = doc.tables[3]
    tabela_semana_3 = doc.tables[5]
    cabecalho_semana_1 = doc.tables[0]
    cabecalho_semana_2 = doc.tables[2]
    cabecalho_semana_3 = doc.tables[4]

    texto_semana_1 = "\n".join(cell.text for row in tabela_semana_1.rows for cell in row.cells)
    texto_semana_2 = "\n".join(cell.text for row in tabela_semana_2.rows for cell in row.cells)
    texto_semana_3 = "\n".join(cell.text for row in tabela_semana_3.rows for cell in row.cells)

    assert "03/06" in texto_semana_1
    assert "10/06" not in texto_semana_1
    assert "10/06" in texto_semana_2
    assert "17/06" not in texto_semana_2
    assert "17/06" in texto_semana_3
    assert cabecalho_semana_1.rows[3].cells[1].text == "01/06 a 05/06"
    assert cabecalho_semana_2.rows[3].cells[1].text == "08/06 a 12/06"
    assert cabecalho_semana_3.rows[3].cells[1].text == "15/06 a 19/06"


def test_cabecalho_conta_aulas_reais_da_semana_quando_feriado_remove_um_dia():
    modelo = _modelo_com_semanas(
        ["01/06 a 05/06", "08/06 a 12/06"],
        linhas_aulas=2,
    )
    doc = Document(
        preencher_documento(
            modelo,
            [
                _aula(date(2026, 6, 3), "Semana 1"),
                _aula(date(2026, 6, 10), "Semana 2 - A"),
                _aula(date(2026, 6, 12), "Semana 2 - B"),
            ],
            professor="BEATRIZ",
            disciplina="Lingua Portuguesa",
            turma="1o ANO C",
            mes="JUNHO",
            bimestre="2o Bimestre",
            aulas_previstas_manual="4",
        )
    )

    assert doc.tables[0].rows[3].cells[3].text == "2"
    assert doc.tables[2].rows[3].cells[3].text == "4"


def test_reutiliza_blocos_vazios_quando_cabecalho_nao_casa_com_datas_das_aulas():
    modelo = _modelo_com_semanas(
        ["04/05 a 08/05", "11/05 a 15/05", "18/05 a 22/05"],
        linhas_aulas=4,
    )
    doc = Document(
        preencher_documento(
            modelo,
            [
                _aula(date(2026, 6, 1), "Aula 10"),
                _aula(date(2026, 6, 2), "Aula 11"),
                _aula(date(2026, 6, 8), "Aula 12"),
                _aula(date(2026, 6, 9), "Aula 13"),
            ],
            professor="SILVANA",
            disciplina="Ciencias",
            turma="9o ANO C",
            mes="JUNHO",
            bimestre="2o Bimestre",
            aulas_previstas_manual="4",
        )
    )

    texto_semana_1 = "\n".join(cell.text for row in doc.tables[1].rows for cell in row.cells)
    texto_semana_2 = "\n".join(cell.text for row in doc.tables[3].rows for cell in row.cells)

    assert len(doc.tables) == 4
    assert len(doc.tables[1].rows) == 3
    assert len(doc.tables[3].rows) == 3
    assert doc.tables[0].rows[3].cells[1].text == "01/06 a 05/06"
    assert doc.tables[2].rows[3].cells[1].text == "08/06 a 12/06"
    assert doc.tables[0].rows[3].cells[3].text == "4"
    assert "01/06" in texto_semana_1
    assert "02/06" in texto_semana_1
    assert "08/06" not in texto_semana_1
    assert "08/06" in texto_semana_2
    assert "09/06" in texto_semana_2


def test_reutiliza_blocos_vazios_quando_datas_vem_como_texto():
    modelo = _modelo_com_semanas(
        ["04/05 a 08/05", "11/05 a 15/05"],
        linhas_aulas=4,
    )
    doc = Document(
        preencher_documento(
            modelo,
            [
                _aula("01/06", "Aula 1"),
                _aula("08/06", "Aula 2"),
                _aula("15/06", "Aula 3"),
                _aula("22/06", "Aula 4"),
                _aula("29/06", "Aula 5"),
            ],
            professor="SILVANA",
            disciplina="Biologia",
            turma="1o ANO A",
            mes="JUNHO",
            bimestre="2o Bimestre",
            aulas_previstas_manual="1",
        )
    )

    semanas = [doc.tables[i].rows[3].cells[1].text for i in range(0, len(doc.tables), 2)]
    assert semanas == [
        "01/06 a 05/06",
        "08/06 a 12/06",
        "15/06 a 19/06",
        "22/06 a 26/06",
        "29/06 a 03/07",
    ]
    assert all(len(doc.tables[i].rows) == 2 for i in range(1, len(doc.tables), 2))


def test_remove_linhas_vazias_sobrando_e_quebra_pagina_entre_semanas():
    modelo = _modelo_com_semanas(["01/06 a 05/06", "08/06 a 12/06"], linhas_aulas=3)
    doc = Document(
        preencher_documento(
            modelo,
            [_aula("03/06"), _aula("10/06")],
            professor="SILVANA",
            disciplina="Aprof. em Biologia",
            turma="2o ANO A",
            mes="JUNHO",
            bimestre="2o Bimestre",
            aulas_previstas_manual="2",
        )
    )

    assert len(doc.tables[1].rows) == 2
    assert len(doc.tables[3].rows) == 2
    quebras = doc._element.xpath('.//w:br[@w:type="page"]')
    assert len(quebras) == 0


def test_separa_etapas_do_desenvolvimento_mesmo_quando_vem_tudo_na_mesma_linha():
    aula = _aula("05/06")
    aula["metodologia"] = [
        {
            "titulo": "Desenvolvimento",
            "texto": (
                "Abertura e contextualizacao: Iniciar a aula com observacao da imagem. "
                "Foco no conteudo: Explicar o conceito principal. "
                "Desenvolvimento guiado: Organizar a analise de dados. "
                "Atividade principal: Produzir um texto-sintese. "
                "Fechamento: Retomar os pontos centrais."
            ),
        }
    ]

    doc = _gerar([aula])
    desenvolvimento = doc.tables[1].rows[1].cells[3].text
    linhas = [linha.strip() for linha in desenvolvimento.splitlines() if linha.strip()]

    assert len(linhas) >= 5
    assert linhas[0].startswith("Abertura e contextualização:")
    assert any(linha.startswith("Atividade principal:") for linha in linhas)
    assert any(linha.startswith("Fechamento:") for linha in linhas)
    assert all(linha != "Desenvolvimento:" for linha in linhas)


def test_cria_quinto_bloco_quando_junho_tem_aula_extra_fora_do_modelo():
    modelo = _modelo_com_semanas(
        ["01/06 a 05/06", "08/06 a 12/06", "15/06 a 19/06", "22/06 a 26/06"],
        linhas_aulas=4,
    )
    aulas = [
        _aula(date(2026, 6, 1), "Aula 10"),
        _aula(date(2026, 6, 4), "Aula 11"),
        _aula(date(2026, 6, 8), "Aula 12"),
        _aula(date(2026, 6, 11), "Aula 13"),
        _aula(date(2026, 6, 15), "Aula 14"),
        _aula(date(2026, 6, 18), "Aula 15"),
        _aula(date(2026, 6, 22), "Aula 16"),
        _aula(date(2026, 6, 25), "Aula 17"),
        _aula(date(2026, 6, 29), "Aula 18"),
    ]
    doc = Document(
        preencher_documento(
            modelo,
            aulas,
            professor="ANDREA ROQUE DUARTE",
            disciplina="Lingua Portuguesa",
            turma="9o ANO C",
            mes="JUNHO",
            bimestre="2o Bimestre",
            aulas_previstas_manual="2",
        )
    )

    texto_primeira_semana = "\n".join(cell.text for row in doc.tables[1].rows for cell in row.cells)
    texto_quinta_semana = "\n".join(cell.text for row in doc.tables[9].rows for cell in row.cells)

    assert len(doc.tables) == 10
    assert doc.tables[8].rows[3].cells[1].text == "29/06 a 03/07"
    assert "29/06" not in texto_primeira_semana
    assert "29/06" in texto_quinta_semana
    assert "AULA 5 - Aula 18" in texto_quinta_semana


def test_polimento_corrige_acentos_observados():
    doc = _gerar([_aula("05/06")])

    texto = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert "j" in texto
    assert "consequ" in texto
    assert "sistematiza" in texto
    assert "servi" in texto
    assert "pre" in texto


def test_remove_caracteres_invalidos_para_xml_no_docx():
    aula = _aula("05/06")
    aula["metodologia"] = [
        {
            "titulo": "Para\x00 comecar",
            "texto": "Explicar o conte\x08udo com registro no caderno e correcao coletiva.",
        }
    ]

    doc = _gerar([aula])
    texto = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert "\x00" not in texto
    assert "\x08" not in texto
    assert "Para" in texto
    assert "conte" in texto


def test_tabela_aulas_recebe_geometria_estavel_para_evitar_corte_visual():
    aula = _aula("05/06", "Virologia: mutacoes virais e a necessidade de manutencao da cobertura vacinal.")
    aula["aprendizagem"] = (
        "Utilizar os conhecimentos das Ciencias da Natureza na explicacao da fisiologia humana "
        "e sua relacao com habitos e condicoes de vida, agindo individual e coletivamente para "
        "promocao da saude e bem-estar."
    )

    doc = _gerar([aula])
    tabela = doc.tables[1]
    tbl_w = tabela._tbl.tblPr.find(qn("w:tblW"))
    grid = tabela._tbl.tblGrid

    assert tabela.autofit is False
    assert tbl_w.get(qn("w:type")) == "dxa"
    assert [col.get(qn("w:w")) for col in grid.gridCol_lst] == ["900", "2100", "2350", "6100", "1900", "2050"]
    assert not tabela._element.xpath(".//w:trHeight")

    material_tc_w = tabela.rows[1].cells[1]._tc.tcPr.tcW
    aprendizagem_tc_w = tabela.rows[1].cells[2]._tc.tcPr.tcW
    assert material_tc_w.get(qn("w:w")) == "2100"
    assert aprendizagem_tc_w.get(qn("w:w")) == "2350"
