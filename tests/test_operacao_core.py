from io import BytesIO
import zipfile

from docx import Document

from core.operacao import (
    detectar_alteracoes_planos_revisados,
    gerar_docx_final,
    montar_zip_planos,
)


def _modelo_minimo():
    doc = Document()
    cabecalho = doc.add_table(rows=4, cols=9)
    cabecalho.rows[0].cells[0].text = "PLANO DE AULAS"
    cabecalho.rows[1].cells[2].text = "PROFESSOR"
    cabecalho.rows[1].cells[3].text = "COMPONENTE CURRICULAR"
    cabecalho.rows[1].cells[6].text = "TURMA"
    cabecalho.rows[1].cells[7].text = "MES"
    cabecalho.rows[1].cells[8].text = "BIMESTRE"
    cabecalho.rows[3].cells[0].text = "SEMANA"
    cabecalho.rows[3].cells[1].text = "01/09 a 05/09"
    cabecalho.rows[3].cells[2].text = "AULAS PREVISTAS NA SEMANA"

    aulas = doc.add_table(rows=2, cols=6)
    aulas.rows[0].cells[0].text = "AULA SEMANAL (Data e Horario)"
    aulas.rows[0].cells[1].text = "NUMERO E TITULO DO MATERIAL DIGITAL"
    aulas.rows[0].cells[2].text = "APRENDIZAGEM ESSENCIAL*"
    aulas.rows[0].cells[3].text = "DESENVOLVIMENTO"
    aulas.rows[0].cells[4].text = "ACOMPANHAMENTO DA APRENDIZAGEM"
    aulas.rows[0].cells[5].text = "ACESSIBILIDADE"

    saida = BytesIO()
    doc.save(saida)
    saida.seek(0)
    return saida.getvalue()


def _aula_teste():
    return {
        "tema": "Tema de teste",
        "material": "AULA 1",
        "aprendizagem": "Aprendizagem de teste",
        "metodologia": [
            {"titulo": "Para começar", "texto": "Retomar conhecimentos prévios."},
            {"titulo": "Na prática", "texto": "Resolver atividade orientada."},
            {"titulo": "Encerramento", "texto": "Socializar respostas."},
        ],
        "acompanhamento": [
            "☑ Observar a participação dos estudantes.",
            "☑ Verificar os registros da atividade.",
            "☑ Conferir as respostas socializadas.",
        ],
        "acessibilidade": [
            "☑ Oferecer leitura guiada.",
            "☑ Destacar palavras-chave.",
            "☑ Permitir resposta oral mediada.",
        ],
        "ia_usada": False,
        "data": "02/09",
        "horario": "14h40 - 15h30",
    }


def test_gerar_docx_final_e_montar_zip_planos():
    plano = gerar_docx_final(
        _modelo_minimo(),
        [_aula_teste()],
        escola="Escola Teste",
        professor="Professor Teste",
        disciplina="História",
        componente_curricular="História",
        turma_atual="8º ANO A",
        mes="SETEMBRO",
        bimestre="3º Bimestre",
        semana="01/09 a 05/09",
        observacao="",
        aulas_previstas_manual="1",
    )

    zip_bytes = montar_zip_planos([plano], "História")
    with zipfile.ZipFile(BytesIO(zip_bytes)) as zf:
        nomes = set(zf.namelist())

    assert plano["turma"] == "8º ANO A"
    assert any(nome.endswith(".docx") for nome in nomes)
    assert any(nome.endswith("_relatorio.txt") for nome in nomes)


def test_detectar_alteracoes_planos_revisados():
    plano = {"turma": "8º ANO A", "aulas": [_aula_teste()]}
    assert not detectar_alteracoes_planos_revisados(
        [plano],
        [{"turma": "8º ANO A", "aulas": [_aula_teste()]}],
    )

    aula_alterada = _aula_teste()
    aula_alterada["tema"] = "Tema alterado"
    assert detectar_alteracoes_planos_revisados(
        [plano],
        [{"turma": "8º ANO A", "aulas": [aula_alterada]}],
    )
