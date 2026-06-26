from docx import Document

from core import lote
from core.lote import _montar_resultado_aula_local
from core.referencias_arte import (
    localizar_docx_referencia_arte,
    referencia_arte_por_pdf,
    titulos_referencia_arte_por_docx,
)


def _criar_docx_referencia_arte(caminho, incluir_aula_3: bool = True):
    doc = Document()
    doc.add_paragraph("AULA 01 — Arte nas dobras de papel")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph(
        "Para comecar: Apresentar um exemplo visual ou modelo fisico de dobradura."
    )
    doc.add_paragraph(
        "Foco no conteudo: Explicar o conceito de Arte nas dobras de papel."
    )
    doc.add_paragraph(
        "Na pratica: Propor a experimentacao e confeccao pratica de dobraduras."
    )
    doc.add_paragraph("Encerramento: Conduzir apreciacao compartilhada.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Identificar se os estudantes participam das praticas.")
    doc.add_paragraph("☑ Verificar se os estudantes reconhecem elementos.")
    doc.add_paragraph("☑ Monitorar se os registros revelam ampliacao.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Permitir leitura em dupla ou com apoio.")
    doc.add_paragraph("☑ Descrever oralmente os elementos visuais.")
    doc.add_paragraph("☑ Garantir acessibilidade fisica dos materiais.")

    doc.add_paragraph("AULA 02 — Explorando dobraduras")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para comecar: Retomar o assunto de dobras e origamis.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar a participacao.")
    doc.add_paragraph("☑ Observar a criacao.")
    doc.add_paragraph("☑ Conferir os registros.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Apoiar individualmente.")
    doc.add_paragraph("☑ Oferecer adaptacoes.")
    doc.add_paragraph("☑ Permitir registro oral.")

    if incluir_aula_3:
        doc.add_paragraph("AULA 03 — Aula disponivel futuramente")
        doc.add_paragraph("Metodologia")
        doc.add_paragraph("Para comecar: Texto da aula 3.")
        doc.add_paragraph("Acompanhamento da aprendizagem")
        doc.add_paragraph("☑ Verificar um registro.")
        doc.add_paragraph("☑ Observar uma fala.")
        doc.add_paragraph("☑ Conferir uma resposta.")
        doc.add_paragraph("Acessibilidade")
        doc.add_paragraph("☑ Apoiar com palavras-chave.")
        doc.add_paragraph("☑ Apoiar com perguntas.")
        doc.add_paragraph("☑ Apoiar com resposta oral.")

    doc.add_paragraph("AULA 04 — Gravura")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para comecar: Apresentar conceitos basicos de xilogravura.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se reconhece gravuras.")
    doc.add_paragraph("☑ Observar a criatividade.")
    doc.add_paragraph("☑ Conferir o registro pratico.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Oferecer materiais alternativos.")
    doc.add_paragraph("☑ Registrar passos no quadro.")
    doc.add_paragraph("☑ Permitir resposta oral.")
    doc.save(caminho)


def test_referencia_arte_le_docx_da_pasta_do_pdf(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Arte_6_Ano.docx"
    caminho_pdf = tmp_path / "AULA_01 - Arte nas dobras de papel.pdf"
    _criar_docx_referencia_arte(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_arte_por_pdf(caminho_pdf, "1")

    assert referencia["titulo"] == "Arte nas dobras de papel"
    assert [etapa["titulo"] for etapa in referencia["metodologia"]][:3] == [
        "Para comecar",
        "Foco no conteudo",
        "Na pratica",
    ]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_titulos_arte_6ano_permitem_salto_da_aula_3(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Arte_6_Ano.docx"
    _criar_docx_referencia_arte(caminho_docx, incluir_aula_3=False)

    titulos = titulos_referencia_arte_por_docx(caminho_docx)

    assert titulos == {
        1: "Arte nas dobras de papel",
        2: "Explorando dobraduras",
        4: "Gravura",
    }


def test_localizar_docx_referencia_arte(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Arte_7_Ano.docx"
    caminho_pdf = tmp_path / "AULA_01.pdf"
    _criar_docx_referencia_arte(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    localizado = localizar_docx_referencia_arte(caminho_pdf)
    assert localizado == caminho_docx.resolve()


def test_arte_resultado_local_usa_docx(tmp_path, monkeypatch):
    caminho_docx = tmp_path / "Metodologias_Arte_6_Ano.docx"
    caminho_pdf = tmp_path / "AULA_01.pdf"
    _criar_docx_referencia_arte(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    # Mock _extrator_lib.extrair
    monkeypatch.setattr(lote._extrator_lib, "extrair", lambda *args, **kwargs: {
        "habilidade": "HAB123",
        "conceito_extraido": "Conceito de teste",
        "recursos_detectados": ["Quadro", "Caderno"]
    })

    resultado = _montar_resultado_aula_local(
        texto="Texto da aula",
        tema="Arte nas dobras de papel",
        material_digital="AULA 1",
        numero_aula="1",
        disciplina_base="Arte",
        turma="6º ANO A",
        provedor_ia="",
        perfil="arte",
        contexto_metodologico="regular",
        indice_aula=0,
        total_aulas=1,
        modalidade_eja_ativa=False,
        metodologia_fixa_pdf=[],
        aprendizagem_pv="",
        objetivos_orientacao=[],
        aprendizagem_orientacao="",
        usar_ia=False,
        ia_erro="",
        caminho_pdf=str(caminho_pdf)
    )

    assert resultado["tema"] == "Arte nas dobras de papel"
    assert resultado["origem_metodologia"] == "docx_referencia_arte"
    assert "Apresentar um exemplo visual ou modelo fisico de dobradura." in resultado["metodologia"][0]["texto"]


def test_arte_e_midias_digitais_localiza_docx(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Arte_e_Midias_Digitais_2_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA_01.pdf"
    _criar_docx_referencia_arte(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    localizado = localizar_docx_referencia_arte(caminho_pdf)
    assert localizado == caminho_docx.resolve()


def test_arte_e_midias_digitais_resultado_local_usa_docx(tmp_path, monkeypatch):
    caminho_docx = tmp_path / "Metodologias_Arte_e_Midias_Digitais_2_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA_01.pdf"
    _criar_docx_referencia_arte(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    # Mock _extrator_lib.extrair
    monkeypatch.setattr(lote._extrator_lib, "extrair", lambda *args, **kwargs: {
        "habilidade": "EM13ART01",
        "conceito_extraido": "Conceito de mídias",
        "recursos_detectados": ["Microfone", "Gravador"]
    })

    resultado = _montar_resultado_aula_local(
        texto="Texto da aula",
        tema="Arte nas dobras de papel",
        material_digital="AULA 1",
        numero_aula="1",
        disciplina_base="Arte e Mídias Digitais",
        turma="2º ANO A",
        provedor_ia="",
        perfil="arte",
        contexto_metodologico="regular",
        indice_aula=0,
        total_aulas=1,
        modalidade_eja_ativa=False,
        metodologia_fixa_pdf=[],
        aprendizagem_pv="",
        objetivos_orientacao=[],
        aprendizagem_orientacao="",
        usar_ia=False,
        ia_erro="",
        caminho_pdf=str(caminho_pdf)
    )

    assert resultado["tema"] == "Arte nas dobras de papel"
    assert resultado["origem_metodologia"] == "docx_referencia_arte"
    assert "Apresentar um exemplo visual ou modelo fisico de dobradura." in resultado["metodologia"][0]["texto"]

