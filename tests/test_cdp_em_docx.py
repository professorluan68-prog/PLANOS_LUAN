from io import BytesIO

from docx import Document

from core.cdp_em_docx import reescrever_docx_cdp_ensino_medio


def _docx_base_material_matematica() -> bytes:
    doc = Document()
    tabela = doc.add_table(rows=3, cols=6)
    tabela.rows[0].cells[0].text = "AULA SEMANAL (Data e Horário)"
    tabela.rows[0].cells[1].text = "NÚMERO E TÍTULO DO MATERIAL DIGITAL"
    tabela.rows[0].cells[2].text = "APRENDIZAGEM ESSENCIAL"
    tabela.rows[0].cells[3].text = "DESENVOLVIMENTO"
    tabela.rows[0].cells[4].text = "ACOMPANHAMENTO DA APRENDIZAGEM"
    tabela.rows[0].cells[5].text = "ACESSIBILIDADE"

    tabela.rows[1].cells[0].text = "01/06\n1, 2 e 3 aula\n13h às 15h50"
    tabela.rows[1].cells[1].text = "AULA 13 - Matemática\nRelação entre grandezas: representação algébrica"
    tabela.rows[1].cells[2].text = (
        "Habilidade: (EM13MAT501) Investigar relações entre números expressos em tabelas "
        "para representá-los no plano cartesiano."
    )
    tabela.rows[1].cells[3].text = "Texto antigo"
    tabela.rows[1].cells[4].text = "Item antigo"
    tabela.rows[1].cells[5].text = "Item antigo"

    tabela.rows[2].cells[0].text = "04/06\n4 e 5 aula\n15h50 às 17h30"
    tabela.rows[2].cells[1].text = "AULA 14 - Matemática\nFunção logarítmica"
    tabela.rows[2].cells[2].text = "Habilidade: interpretar variações em escala e relacionar potência e logaritmo."
    tabela.rows[2].cells[3].text = "Texto antigo"
    tabela.rows[2].cells[4].text = "Item antigo"
    tabela.rows[2].cells[5].text = "Item antigo"

    saida = BytesIO()
    doc.save(saida)
    return saida.getvalue()


def test_reescrita_docx_cdp_em_normaliza_tema_e_metodologia():
    corrigido_bytes, relatorio = reescrever_docx_cdp_ensino_medio(_docx_base_material_matematica())
    doc = Document(BytesIO(corrigido_bytes))
    tabela = doc.tables[0]

    material_1 = tabela.rows[1].cells[1].text
    desenvolvimento_1 = tabela.rows[1].cells[3].text.lower()
    acompanhamento_1 = tabela.rows[1].cells[4].text.lower()
    acessibilidade_1 = tabela.rows[1].cells[5].text.lower()

    assert relatorio["linhas_reescritas"] == 2
    assert material_1.startswith("TEMA:")
    assert "AULA 13" not in material_1
    assert "representação algébrica" in material_1
    assert "tabela" in desenvolvimento_1
    assert "expressão algébrica" in desenvolvimento_1
    assert "grandezas" in acompanhamento_1
    assert "tabelas simples" in acessibilidade_1


def test_reescrita_docx_cdp_em_reconhece_funcao_logaritmica():
    corrigido_bytes, _ = reescrever_docx_cdp_ensino_medio(_docx_base_material_matematica())
    doc = Document(BytesIO(corrigido_bytes))
    tabela = doc.tables[0]

    material_2 = tabela.rows[2].cells[1].text
    desenvolvimento_2 = tabela.rows[2].cells[3].text.lower()
    acompanhamento_2 = tabela.rows[2].cells[4].text.lower()
    acessibilidade_2 = tabela.rows[2].cells[5].text.lower()

    assert material_2.startswith("TEMA:")
    assert "função logarítmica" in material_2.lower()
    assert "potência" in desenvolvimento_2
    assert "logaritmo" in acompanhamento_2
    assert "potências simples" in acessibilidade_2
