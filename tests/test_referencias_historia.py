from pathlib import Path

import pytest
from docx import Document

from core.referencias_historia import (
    _carregar_referencias_historia_docx,
    _obter_grade_e_aula_do_pdf,
    _selecionar_referencia,
    localizar_docx_referencia_historia,
    localizar_docx_referencia_historia_cdp,
    referencia_historia_cdp_por_pdf,
    referencia_historia_por_pdf,
)


def test_obter_grade_e_aula_do_pdf():
    path1 = Path("D:/PDF novos/HISTORIA/AF/3_BIMESTRE/6_ANO/AULA_01 - As polis gregas.pdf")
    grade, numero = _obter_grade_e_aula_do_pdf(path1, "1")
    assert grade == 6
    assert numero == 1

    path2 = "D:/PDF novos/HISTORIA/AF/3_BIMESTRE/8_ANO/AULA_07 - O Periodo Regencial.pdf"
    grade, numero = _obter_grade_e_aula_do_pdf(path2, None)
    assert grade == 8
    assert numero == 7

    path3 = Path("D:/SomeOtherFolder/9_ANO/1652639.pdf")
    grade, numero = _obter_grade_e_aula_do_pdf(path3, 17)
    assert grade == 9
    assert numero == 17


def test_selecionar_referencia():
    referencias = {
        (6, 1): {
            "grade": 6,
            "numero": 1,
            "titulo": "As polis gregas",
            "metodologia": [{"titulo": "Para começar", "texto": "Inicie a aula..."}],
            "acompanhamento": ["☑ Item 1"],
            "acessibilidade": ["☑ Adapt 1"],
        },
        (6, 2): {
            "grade": 6,
            "numero": 2,
            "titulo": "Guerras Médicas",
            "metodologia": [{"titulo": "Para começar", "texto": "Conduza..."}],
            "acompanhamento": ["☑ Item 2"],
            "acessibilidade": ["☑ Adapt 2"],
        },
    }

    ref = _selecionar_referencia(referencias, 6, 1, "")
    assert ref is not None
    assert ref["titulo"] == "As polis gregas"

    ref_fallback = _selecionar_referencia(referencias, 6, 99, "As polis gregas e Atenas")
    assert ref_fallback is not None
    assert ref_fallback["numero"] == 1

    ref_none = _selecionar_referencia(referencias, 7, 1, "Guerras Médicas")
    assert ref_none is None


def test_carregar_referencias_historia_docx(tmp_path):
    doc_path = tmp_path / "Metodologias_Historia_Ensino_Regular.docx"
    doc = Document()

    doc.add_paragraph("6º ANO - AULA 1 - As pólis gregas")
    doc.add_paragraph("metodologia")
    doc.add_paragraph("Para começar: Inicie a aula apresentando a pólis grega.")
    doc.add_paragraph("Foco no conteúdo: Explique o conceito de cidade-estado.")
    doc.add_paragraph("acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Item A")
    doc.add_paragraph("☑ Item B")
    doc.add_paragraph("☑ Item C")
    doc.add_paragraph("acessibilidade")
    doc.add_paragraph("☑ Adapt A")
    doc.add_paragraph("☑ Adapt B")
    doc.add_paragraph("☑ Adapt C")
    doc.save(doc_path)

    aulas = _carregar_referencias_historia_docx(str(doc_path))
    assert (6, 1) in aulas
    aula = aulas[(6, 1)]
    assert aula["titulo"] == "As pólis gregas"
    assert len(aula["metodologia"]) == 2
    assert aula["metodologia"][0]["titulo"] == "Para começar"
    assert aula["metodologia"][0]["texto"] == "Inicie a aula apresentando a pólis grega."
    assert len(aula["acompanhamento"]) == 3
    assert aula["acompanhamento"][0] == "☑ Item A"
    assert len(aula["acessibilidade"]) == 3
    assert aula["acessibilidade"][0] == "☑ Adapt A"


def test_localizar_docx_referencia_historia_prioriza_regular(tmp_path):
    base = tmp_path / "3_BIMESTRE"
    ano = base / "6_ANO"
    ano.mkdir(parents=True)
    pdf = ano / "AULA_01 - Tema.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    doc = base / "Metodologias_Historia_Ensino_Regular.docx"
    doc.write_bytes(b"docx")
    (base / "Metodologias_Historia_CDP.docx").write_bytes(b"docx")

    assert localizar_docx_referencia_historia(pdf) == doc


def test_referencia_historia_por_pdf_retorna_aula_do_docx(tmp_path):
    base = tmp_path / "3_BIMESTRE"
    ano = base / "6_ANO"
    ano.mkdir(parents=True)
    pdf = ano / "AULA_01 - As polis gregas.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    doc_path = base / "Metodologias_Historia_Ensino_Regular.docx"

    doc = Document()
    doc.add_paragraph("6º ANO - AULA 1 - As pólis gregas")
    doc.add_paragraph("metodologia")
    doc.add_paragraph("Para começar: Inicie a aula retomando as anotações.")
    doc.add_paragraph("acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar registros")
    doc.add_paragraph("☑ Observar participação")
    doc.add_paragraph("☑ Retomar síntese final")
    doc.add_paragraph("acessibilidade")
    doc.add_paragraph("☑ Garantir leitura mediada")
    doc.add_paragraph("☑ Oferecer apoio visual")
    doc.add_paragraph("☑ Organizar tempo ampliado")
    doc.save(doc_path)

    referencia = referencia_historia_por_pdf(pdf, "1", tema="As pólis gregas")

    assert referencia is not None
    assert referencia["titulo"] == "As pólis gregas"
    assert referencia["metodologia"][0]["texto"] == "Inicie a aula retomando as anotações."
    assert len(referencia["acompanhamento"]) == 3
    assert referencia["fonte"] == str(doc_path)


@pytest.mark.xfail(
    reason=(
        "Arquivo Metodologias_Historia_Ensino_Regular.docx existe mas foi gerado "
        "sem o cabecalho 'Xo ANO - AULA N' que o parser espera. Teste sera corrigido "
        "quando o arquivo for regenerado no formato padrao."
    ),
    strict=False,
)
def test_arquivos_reais_gerados():
    pdf_base = Path("D:/PDF novos/HISTORIA/AF/3_BIMESTRE")
    if not pdf_base.exists():
        pytest.skip("Base directory D:/PDF novos/HISTORIA/AF/3_BIMESTRE not found.")

    doc_reg = pdf_base / "Metodologias_Historia_Ensino_Regular.docx"
    doc_cdp = pdf_base / "Metodologias_Historia_CDP.docx"

    if not doc_reg.exists() or not doc_cdp.exists():
        pytest.skip("Arquivos de referencia gerados nao existem ainda.")

    aulas_reg = _carregar_referencias_historia_docx(str(doc_reg))
    aulas_cdp = _carregar_referencias_historia_docx(str(doc_cdp))

    assert len(aulas_reg) > 0
    assert len(aulas_cdp) > 0

    grades_reg = {g for g, _ in aulas_reg.keys()}
    grades_cdp = {g for g, _ in aulas_cdp.keys()}

    assert {6, 7, 8, 9}.issubset(grades_reg)
    assert {6, 7, 8, 9}.issubset(grades_cdp)
