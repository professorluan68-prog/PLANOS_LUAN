from pathlib import Path

from docx import Document

from core.educacao_financeira_validacao import validar_requisitos_educacao_financeira
from core.helpers import resolver_pasta_pdfs
from core.referencias_educacao_financeira import localizar_docx_referencia


def _criar_referencia(caminho: Path, titulo: str = "Orçamento doméstico") -> None:
    doc = Document()
    doc.add_paragraph(f"AULA 1 - {titulo}")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Retomar receitas e despesas do orçamento doméstico.")
    doc.add_paragraph("Foco no conteúdo: Orientar a leitura da planilha de gastos.")
    doc.add_paragraph("Na prática: Organizar prioridades e reserva financeira.")
    doc.add_paragraph("Encerramento: Socializar critérios de economia doméstica.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Observar se os estudantes identificam receitas e despesas.")
    doc.add_paragraph("☑ Verificar os registros feitos na planilha de gastos.")
    doc.add_paragraph("☑ Conferir as justificativas para a reserva financeira.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Disponibilizar tabela de apoio para receitas e despesas.")
    doc.add_paragraph("☑ Oferecer roteiro com perguntas orientadoras em frases curtas.")
    doc.add_paragraph("☑ Permitir resposta oral antes do registro guiado.")
    doc.save(caminho)


def test_8_ano_corrigido_tem_prioridade_exata(tmp_path):
    esperado = tmp_path / "Metodologias_Educacao_Financeira_8_Ano_CORRIGIDO.docx"
    concorrente = tmp_path / "Metodologias_Educacao_Financeira_8_Ano_ATUALIZADO.docx"
    pdf = tmp_path / "AULA 1.pdf"
    _criar_referencia(esperado)
    _criar_referencia(concorrente, titulo="Título concorrente")
    pdf.write_bytes(b"%PDF-1.4\n")

    assert localizar_docx_referencia(pdf) == esperado


def test_validacao_financeira_aceita_tres_itens_concretos():
    aula = {
        "disciplina": "Educação Financeira",
        "acompanhamento": [
            "☑ Observar se os estudantes identificam receitas e despesas.",
            "☑ Verificar os registros na planilha de gastos.",
            "☑ Conferir as justificativas para a reserva financeira.",
        ],
        "acessibilidade": [
            "☑ Disponibilizar tabela de apoio para receitas e despesas.",
            "☑ Oferecer roteiro com perguntas orientadoras.",
            "☑ Permitir resposta oral antes do registro guiado.",
        ],
    }

    assert validar_requisitos_educacao_financeira(aula) == []


def test_validacao_financeira_aceita_apoios_concretos_do_8_ano_corrigido():
    aula = {
        "disciplina": "Educação Financeira",
        "acompanhamento": [
            "☑ Verificar se os estudantes diferenciam receitas, despesas, orçamento e saldo.",
            "☑ Observar se calculam o saldo familiar e interpretam superávit ou déficit.",
            "☑ Acompanhar se propõem ajustes coerentes para melhorar o orçamento analisado.",
        ],
        "acessibilidade": [
            "☑ Construir no quadro um glossário visual com receitas, despesas, orçamento, saldo, superávit e déficit.",
            "☑ Disponibilizar tabela-modelo para estudantes com dificuldade de organização dos dados.",
            "☑ Permitir cálculo com calculadora e conferência em dupla, mantendo o foco na interpretação financeira.",
        ],
    }

    assert validar_requisitos_educacao_financeira(aula) == []


def test_validacao_financeira_aceita_funcoes_variadas_no_grupo():
    aula = {
        "disciplina": "Educação Financeira",
        "acompanhamento": [
            "☑ Verificar se os estudantes distinguem orçamento pontual e orçamento contínuo.",
            "☑ Observar se os grupos analisam modelos de orçamento antes de tomar decisões.",
            "☑ Acompanhar se justificam escolhas considerando receitas, despesas, prioridades e objetivos.",
        ],
        "acessibilidade": [
            "☑ Apresentar quadro comparativo entre orçamento pontual e contínuo com exemplos simples.",
            "☑ Disponibilizar roteiro de tomada de decisão com perguntas sobre entradas, saídas e ajustes.",
            "☑ Permitir que estudantes participem do grupo com funções variadas, como leitura, cálculo, registro ou apresentação.",
        ],
    }

    assert validar_requisitos_educacao_financeira(aula) == []


def test_validacao_financeira_aceita_apoios_concretos_dos_docx_do_em():
    aula = {
        "disciplina": "Educação Financeira",
        "acompanhamento": [
            "☑ Verificar se os estudantes identificam receitas, despesas e metas presentes na proposta.",
            "☑ Observar se analisam os dados do planejamento financeiro com coerência.",
            "☑ Conferir se justificam as escolhas feitas na organização do orçamento.",
        ],
        "acessibilidade": [
            "☑ Disponibilizar perguntas-guia para apoiar a análise dos dados e a elaboração das estratégias em dupla.",
            "☑ Permitir que a resposta final seja registrada em tópicos, frases curtas ou explicada oralmente com mediação.",
            "☑ Organizar duplas produtivas para leitura dos comandos, pesquisa de preços e conferência dos registros.",
        ],
    }

    assert validar_requisitos_educacao_financeira(aula) == []


def test_validacao_financeira_aceita_leitura_coletiva_como_apoio_concreto():
    aula = {
        "disciplina": "Educação Financeira",
        "acompanhamento": [
            "☑ Verificar se os estudantes identificam modalidades de crédito e seus riscos.",
            "☑ Observar se analisam prazos, valores e custo total antes da decisão.",
            "☑ Conferir se justificam as escolhas com base no orçamento e nos riscos.",
        ],
        "acessibilidade": [
            "☑ Disponibilizar quadro comparativo com definição, exemplo e risco de cada modalidade de crédito.",
            "☑ Ler coletivamente as situações-problema, destacando números, prazos e termos financeiros importantes.",
            "☑ Oferecer perguntas-guia para apoiar a análise: preciso? cabe no orçamento? qual o custo total? quais os riscos?",
        ],
    }

    assert validar_requisitos_educacao_financeira(aula) == []


def test_validacao_financeira_rejeita_item_sem_check_verbo_ou_apoio():
    aula = {
        "disciplina": "Educação Financeira",
        "acompanhamento": [
            "☑ Observar a turma.",
            "☑ Ideias dos estudantes.",
            "Sem check e sem verbo.",
        ],
        "acessibilidade": [
            "☑ Apoiar a turma sem apoio concreto.",
            "☑ Material adaptado genérico.",
            "☑ Orientação individual sem recurso.",
        ],
    }

    problemas = validar_requisitos_educacao_financeira(aula)

    assert any("iniciar com ☑" in problema for problema in problemas)


def test_resolvedor_automatico_monta_pasta_8_ano_3_bimestre(tmp_path):
    caminho = resolver_pasta_pdfs(
        str(tmp_path / "PDF novos"),
        "Educação Financeira",
        "8º ANO",
        "3º BIMESTRE",
    )

    assert caminho == tmp_path / "PDF novos" / "EDUCACAO_FINANCEIRA" / "AF" / "3_BIMESTRE" / "8_ANO"
