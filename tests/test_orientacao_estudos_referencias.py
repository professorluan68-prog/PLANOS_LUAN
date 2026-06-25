from docx import Document

from core import lote
from core.helpers import resolver_pasta_pdfs
from core.referencias_orientacao_estudos import (
    localizar_docx_referencia_orientacao_estudos,
    referencia_orientacao_estudos_por_pdf,
    titulos_referencia_orientacao_estudos_por_docx,
)


def _criar_docx_referencia_orientacao(caminho):
    doc = Document()
    doc.add_paragraph("METODOLOGIAS - ORIENTACAO DE ESTUDOS - 3ª SERIE DO ENSINO MEDIO")

    doc.add_paragraph(
        "AULA 1 — Informações em infográficos, gráficos, tabelas e esquemas\n\n"
        "HABILIDADE: Eixo cognitivo. Reconhecer. LP5LERE02 – Localizar informação explícita."
    )
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Apresentar infográficos e discutir como dados visuais podem informar ou manipular.")
    doc.add_paragraph("Foco no conteúdo: Explicar leitura crítica de gráficos, tabelas e esquemas com atenção às fontes.")
    doc.add_paragraph("Na prática: Analisar textos multissemióticos e identificar tema, fonte e organização visual.")
    doc.add_paragraph("Encerramento: Socializar conclusões sobre verificação de fontes e cidadania digital.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se os estudantes diferenciam a finalidade dos infográficos apresentados.")
    doc.add_paragraph("☑ Observar se identificam a origem e a confiabilidade das fontes de dados.")
    doc.add_paragraph("☑ Acompanhar se relacionam linguagem verbal e não-verbal na análise.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Disponibilizar glossário visual simples para apoiar a leitura dos infográficos.")
    doc.add_paragraph("☑ Oferecer roteiro simplificado com perguntas norteadoras.")
    doc.add_paragraph("☑ Permitir respostas por palavras-chave, esquemas ou registro oral mediado.")

    doc.add_paragraph("AULA 2 — Desenhando para entender melhor")
    doc.add_paragraph("HABILIDADE: Eixo cognitivo. Analisar. LP5LEAN06 – Inferir informações implícitas em textos.")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Retomar formas de organizar e apresentar conhecimentos.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se reconhecem características dos gêneros visuais.")
    doc.add_paragraph("☑ Observar se identificam estratégia visual e público-alvo.")
    doc.add_paragraph("☑ Acompanhar se localizam dados específicos no material.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Disponibilizar banco de termos no quadro.")
    doc.add_paragraph("☑ Oferecer versão ampliada dos infográficos.")
    doc.add_paragraph("☑ Permitir registro por tópicos, símbolos ou resposta oral.")

    doc.add_paragraph("AULA 3 — Seleção de informações e argumentação em textos multissemióticos")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Apresentar questão de exame com leitura de texto verbal e visual.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se relacionam recursos verbais e não-verbais.")
    doc.add_paragraph("☑ Observar se identificam dados usados para sustentar argumentos.")
    doc.add_paragraph("☑ Acompanhar se analisam alternativas com base em evidências.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Disponibilizar diagramas ampliados com marcações.")
    doc.add_paragraph("☑ Permitir resolução em duplas com tempo ampliado.")
    doc.add_paragraph("☑ Oferecer perguntas mediadoras para leitura dos gráficos.")
    doc.save(caminho)


def test_referencia_orientacao_estudos_le_docx_da_pasta_do_pdf(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Orientacao_de_Estudos_3_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "PDF1_com_habilidade_essencial.pdf"
    _criar_docx_referencia_orientacao(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_orientacao_estudos_por_pdf(caminho_pdf, "")

    assert localizar_docx_referencia_orientacao_estudos(caminho_pdf) == caminho_docx
    assert referencia["numero"] == 1
    assert referencia["titulo"] == "Informações em infográficos, gráficos, tabelas e esquemas"
    assert referencia["habilidade"] == "Eixo cognitivo. Reconhecer. LP5LERE02 – Localizar informação explícita."
    assert [etapa["titulo"] for etapa in referencia["metodologia"]][:3] == [
        "Para começar",
        "Foco no conteúdo",
        "Na prática",
    ]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_titulos_orientacao_estudos_por_docx(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Orientacao_de_Estudos_3_Ano_Ensino_Medio.docx"
    _criar_docx_referencia_orientacao(caminho_docx)

    titulos = titulos_referencia_orientacao_estudos_por_docx(caminho_docx)

    assert titulos == {
        1: "Informações em infográficos, gráficos, tabelas e esquemas",
        2: "Desenhando para entender melhor",
        3: "Seleção de informações e argumentação em textos multissemióticos",
    }


def test_orientacao_estudos_resultado_usa_docx_e_titulo_oficial(tmp_path, monkeypatch):
    caminho_docx = tmp_path / "Metodologias_Orientacao_de_Estudos_3_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "PDF1_com_habilidade_essencial.pdf"
    _criar_docx_referencia_orientacao(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")
    monkeypatch.setattr(
        lote,
        "_extrair_texto_pdf",
        lambda caminho: (
            "LINGUA PORTUGUESA\nMATERIAL DE APOIO PEDAGOGICO\n"
            "Habilidade essencial: Compreender criticamente textos de divulgacao cientifica."
        ),
    )

    aula = lote._aula_por_pdf(
        str(caminho_pdf),
        "Orientação de Estudos",
        "3º ANO A",
        "3º Bimestre",
        usar_ia=False,
        provedor_ia="",
    )

    assert aula["tema"] == "Informações em infográficos, gráficos, tabelas e esquemas"
    assert aula["material"] == "AULA 1 - Informações em infográficos, gráficos, tabelas e esquemas"
    assert aula["aprendizagem"] == "Eixo cognitivo. Reconhecer. LP5LERE02 – Localizar informação explícita."
    assert aula["origem_metodologia"] == "docx_referencia_orientacao_estudos"
    assert "fontes" in aula["metodologia"][1]["texto"].lower()
    assert len(aula["acompanhamento"]) == 3
    assert len(aula["acessibilidade"]) == 3


def test_orientacao_estudos_ef_usa_ordem_da_pasta_quando_pdf_tem_intervalo(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Orientacao_de_Estudos_9_Ano_Ensino_Fundamental.docx"
    _criar_docx_referencia_orientacao(caminho_docx)
    matriz = tmp_path / "Matriz_de_Referencia_de_Lingua_Portuguesa.pdf"
    pdf1 = tmp_path / "Missao_11_Um_mergulho_no_cordel-1-3.pdf"
    pdf2 = tmp_path / "Missao_11_Um_mergulho_no_cordel-4-5.pdf"
    for caminho in [matriz, pdf1, pdf2]:
        caminho.write_bytes(b"%PDF-1.4\n")

    referencia1 = referencia_orientacao_estudos_por_pdf(pdf1, "3")
    referencia2 = referencia_orientacao_estudos_por_pdf(pdf2, "5")

    assert referencia1["numero"] == 1
    assert referencia1["titulo"] == "Informações em infográficos, gráficos, tabelas e esquemas"
    assert referencia2["numero"] == 2
    assert referencia2["titulo"] == "Desenhando para entender melhor"
    assert "Inferir informações implícitas" in referencia2["habilidade"]


def test_orientacao_estudos_junta_habilidade_quebrada_em_duas_linhas(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Orientacao_de_Estudos_6_Ano_Ensino_Fundamental.docx"
    caminho_pdf = tmp_path / "Missao_07_A_trama_do_texto-4-5.pdf"
    doc = Document()
    doc.add_paragraph("AULA 1 — MISSAO 7 - A trama do texto - ETAPA 1")
    doc.add_paragraph("HABILIDADE: Eixo cognitivo. Reconhecer. LP5LSRE05 – Identificar os mecanismos de referenciação lexical e")
    doc.add_paragraph("pronominal.")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Retomar exemplos de retomadas no texto.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se localizam retomadas no texto.")
    doc.add_paragraph("☑ Observar se explicam os referentes identificados.")
    doc.add_paragraph("☑ Acompanhar se justificam as respostas com trechos lidos.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Oferecer leitura guiada em trechos curtos.")
    doc.add_paragraph("☑ Disponibilizar palavras-chave no quadro.")
    doc.add_paragraph("☑ Permitir resposta oral mediada com registro em tópicos.")
    doc.save(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_orientacao_estudos_por_pdf(caminho_pdf, "1")

    assert referencia["habilidade"].endswith("lexical e pronominal.")


def test_resolver_pasta_pdfs_orientacao_estudos_em_sem_pasta_bimestre(tmp_path):
    pasta = tmp_path / "ORIENTACAO_DE_ESTUDOS" / "EM" / "3_ANO"
    pasta.mkdir(parents=True)
    (pasta / "PDF1_com_habilidade_essencial.pdf").write_bytes(b"%PDF-1.4\n")

    resolvida = resolver_pasta_pdfs(
        str(tmp_path),
        "Orientação de Estudos",
        "3º ANO A",
        "3º Bimestre",
    )

    assert resolvida == pasta
