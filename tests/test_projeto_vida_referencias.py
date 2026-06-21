from docx import Document

from core import lote
from core.lote import _montar_resultado_aula_local
from core.referencias_projeto_vida import (
    localizar_docx_referencia_projeto_vida,
    referencia_projeto_vida_por_pdf,
    titulos_referencia_projeto_vida_por_docx,
)


def _criar_docx_referencia_projeto_vida(caminho, incluir_aula_3: bool = True):
    doc = Document()
    doc.add_paragraph("AULA 01 — O que me move?")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph(
        "Para comecar: Retomar interesses, habitos e motivacoes pessoais a partir de perguntas orientadoras."
    )
    doc.add_paragraph(
        "Foco no conteudo: Organizar uma conversa guiada sobre metas, escolhas e atitudes que ajudam o estudante a se conhecer melhor."
    )
    doc.add_paragraph(
        "Na pratica: Solicitar que os estudantes registrem uma meta possivel e relacionem essa meta a uma atitude concreta."
    )
    doc.add_paragraph("Encerramento: Socializar algumas metas e combinar um acompanhamento respeitoso dos proximos passos.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se os estudantes identificam interesses e motivacoes pessoais.")
    doc.add_paragraph("☑ Observar a relacao entre meta, atitude concreta e planejamento.")
    doc.add_paragraph("☑ Conferir os registros individuais produzidos durante a aula.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Disponibilizar palavras-chave sobre metas, escolhas e motivacao.")
    doc.add_paragraph("☑ Oferecer perguntas orientadoras em frases curtas.")
    doc.add_paragraph("☑ Permitir resposta oral mediada antes do registro escrito.")

    doc.add_paragraph("AULA 02 — Metas que saem do papel")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para comecar: Retomar metas pessoais e discutir formas de transformar desejo em acao planejada.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se a meta possui acao concreta.")
    doc.add_paragraph("☑ Observar a participacao no planejamento.")
    doc.add_paragraph("☑ Conferir o registro da etapa seguinte.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Usar quadro com exemplo de meta.")
    doc.add_paragraph("☑ Disponibilizar roteiro em topicos.")
    doc.add_paragraph("☑ Apoiar resposta oral mediada.")

    if incluir_aula_3:
        doc.add_paragraph("AULA 03 — Aula disponivel futuramente")
        doc.add_paragraph("Metodologia")
        doc.add_paragraph("Para comecar: Texto da aula 3.")
        doc.add_paragraph("Acompanhamento da aprendizagem")
        doc.add_paragraph("☑ Verificar um registro.")
        doc.add_paragraph("☑ Observar uma fala.")
        doc.add_paragraph("☑ Conferir uma resposta.")
        doc.add_paragraph("Acessibilidade")
        doc.add_paragraph("☑ Apoiar com palavras-chave.")
        doc.add_paragraph("☑ Apoiar com perguntas.")
        doc.add_paragraph("☑ Apoiar com resposta oral.")

    doc.add_paragraph("AULA 04 — Quem me influencia?")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para comecar: Conversar sobre influencias positivas e negativas nas escolhas cotidianas.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se reconhece influencias nas decisoes.")
    doc.add_paragraph("☑ Observar justificativas apresentadas na conversa.")
    doc.add_paragraph("☑ Conferir o registro das escolhas analisadas.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Oferecer exemplos concretos de influencia.")
    doc.add_paragraph("☑ Registrar palavras-chave no quadro.")
    doc.add_paragraph("☑ Permitir resposta oral mediada.")
    doc.save(caminho)


def test_referencia_projeto_vida_le_docx_da_pasta_do_pdf(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Projeto_de_Vida_6_Ano.docx"
    caminho_pdf = tmp_path / "AULA_01 - O que me move.pdf"
    _criar_docx_referencia_projeto_vida(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_projeto_vida_por_pdf(caminho_pdf, "1")

    assert referencia["titulo"] == "O que me move?"
    assert [etapa["titulo"] for etapa in referencia["metodologia"]][:3] == [
        "Para comecar",
        "Foco no conteudo",
        "Na pratica",
    ]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_titulos_projeto_vida_6ano_permitem_salto_da_aula_3(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Projeto_de_Vida_6_Ano.docx"
    _criar_docx_referencia_projeto_vida(caminho_docx, incluir_aula_3=False)

    titulos = titulos_referencia_projeto_vida_por_docx(caminho_docx)

    assert titulos == {
        1: "O que me move?",
        2: "Metas que saem do papel",
        4: "Quem me influencia?",
    }


def test_referencia_projeto_vida_localiza_docx_do_ano(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Projeto_de_Vida_7_Ano.docx"
    caminho_pdf = tmp_path / "AULA_01 - O que me move.pdf"
    _criar_docx_referencia_projeto_vida(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    escolhido = localizar_docx_referencia_projeto_vida(caminho_pdf)

    assert escolhido == caminho_docx


def test_referencia_projeto_vida_localiza_docx_ensino_medio(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Projeto_de_Vida_1_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA_01 - Estresse e ansiedade.pdf"
    _criar_docx_referencia_projeto_vida(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    escolhido = localizar_docx_referencia_projeto_vida(caminho_pdf)

    assert escolhido == caminho_docx


def test_projeto_vida_resultado_local_usa_docx_sem_trocar_titulo_oficial(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Projeto_de_Vida_6_Ano.docx"
    caminho_pdf = tmp_path / "AULA_01 - O que me move.pdf"
    _criar_docx_referencia_projeto_vida(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    resultado = _montar_resultado_aula_local(
        texto="Texto qualquer do PDF que nao deve prevalecer sobre o DOCX de referencia.",
        tema="Titulo vindo da planilha",
        material_digital="AULA 1 - Titulo vindo da planilha",
        numero_aula="1",
        disciplina_base="Projeto de Vida",
        turma="6 ano A",
        provedor_ia="",
        perfil="projeto_de_vida",
        contexto_metodologico="",
        indice_aula=0,
        total_aulas=1,
        modalidade_eja_ativa=False,
        metodologia_fixa_pdf=[],
        aprendizagem_pv="",
        objetivos_orientacao=[],
        aprendizagem_orientacao="",
        usar_ia=False,
        ia_erro="",
        caminho_pdf=str(caminho_pdf),
    )

    assert resultado["tema"] == "Titulo vindo da planilha"
    assert resultado["material"] == "AULA 1 - Titulo vindo da planilha"
    assert resultado["origem_metodologia"] == "docx_referencia_projeto_de_vida"
    assert "motivacoes pessoais" in resultado["metodologia"][0]["texto"].lower()
    assert len(resultado["acompanhamento"]) == 3
    assert len(resultado["acessibilidade"]) == 3


def test_projeto_vida_prefere_titulo_do_docx_quando_escopo_vem_com_texto_do_bimestre(tmp_path, monkeypatch):
    caminho_docx = tmp_path / "Metodologias_Projeto_de_Vida_6_Ano.docx"
    caminho_pdf = tmp_path / "AULA_01 - O que me move.pdf"
    _criar_docx_referencia_projeto_vida(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    monkeypatch.setattr(
        lote,
        "_extrair_texto_pdf",
        lambda caminho: "Projeto de Vida\nAula 01\nO que me move?\n",
    )
    monkeypatch.setattr(
        lote,
        "buscar_item_projeto_vida",
        lambda turma, bimestre, numero: {
            "titulo": (
                "Vozes em Jogo: debater, influenciar e conquistar metas. "
                "Este bimestre se organiza em torno do Roadmap das entregas."
            ),
            "habilidade": "Conhecer habitos pessoais e analisar escolhas.",
            "objeto": "habitos, metas e escolhas pessoais",
            "objetivos": "Identificar habitos pessoais e relaciona-los a metas possiveis.",
            "conteudo": "Projeto de vida e metas",
        },
    )

    aula = lote._aula_por_pdf(
        str(caminho_pdf),
        "Projeto de Vida",
        "6º ANO A",
        "3º Bimestre",
        usar_ia=False,
        provedor_ia="",
    )

    assert aula["tema"] == "O que me move?"
    assert aula["material"] == "AULA 1 - O que me move?"
    assert "Este bimestre" not in aula["material"]
    assert "Roadmap" not in aula["material"]
    assert "metas possiveis" in aula["aprendizagem"]
