from docx import Document

from core.referencias_biologia import (
    referencia_biologia_por_pdf,
    titulos_referencia_biologia_por_docx,
)


def test_referencia_biologia_reconhece_aula_com_underscore(tmp_path):
    caminho_docx = tmp_path / "Metodologia_Biologia_1_Ano_3_Bimestre.docx"
    doc = Document()
    doc.add_paragraph("AULA 5 - Acoes mitigatorias da interferencia humana nos ciclos biogeoquimicos")
    doc.add_paragraph("METODOLOGIA")
    doc.add_paragraph("Para comecar: Retomar estrategias de mitigacao.")
    doc.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    doc.add_paragraph("OK Verificar se identificam as praticas mitigatorias.")
    doc.add_paragraph("OK Analisar relacoes com os ciclos biogeoquimicos.")
    doc.add_paragraph("OK Acompanhar os registros feitos no material.")
    doc.add_paragraph("ACESSIBILIDADE")
    doc.add_paragraph("OK Disponibilizar tabela de sistematizacao.")
    doc.add_paragraph("OK Permitir apresentacao com apoio visual.")
    doc.add_paragraph("OK Retomar oralmente os pontos principais.")
    doc.add_paragraph("AULA_06 - Construcao de hipoteses e teorias cientificas")
    doc.add_paragraph("METODOLOGIA")
    doc.add_paragraph("Para comecar: Apresentar a situacao do cientista em Marte.")
    doc.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    doc.add_paragraph("OK Verificar se diferenciam fato e hipotese.")
    doc.add_paragraph("OK Analisar o uso do metodo cientifico.")
    doc.add_paragraph("OK Acompanhar os registros dos grupos.")
    doc.add_paragraph("ACESSIBILIDADE")
    doc.add_paragraph("OK Realizar leitura guiada da situacao-problema.")
    doc.add_paragraph("OK Disponibilizar quadro-resumo dos conceitos.")
    doc.add_paragraph("OK Permitir respostas por topicos.")
    doc.save(caminho_docx)

    titulos = titulos_referencia_biologia_por_docx(caminho_docx)
    referencia_aula_5 = referencia_biologia_por_pdf(tmp_path / "AULA 5.pdf", "5")
    referencia_aula_6 = referencia_biologia_por_pdf(tmp_path / "AULA 6.pdf", "6")

    assert titulos == {
        5: "Acoes mitigatorias da interferencia humana nos ciclos biogeoquimicos",
        6: "Construcao de hipoteses e teorias cientificas",
    }
    assert len(referencia_aula_5["metodologia"]) == 1
    assert referencia_aula_6["titulo"] == "Construcao de hipoteses e teorias cientificas"
