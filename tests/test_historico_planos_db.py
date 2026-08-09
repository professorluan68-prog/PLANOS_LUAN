from contextlib import contextmanager
from io import BytesIO

import pytest
from docx import Document

from core import database


def _preparar_banco(monkeypatch, tmp_path):
    monkeypatch.setattr(database, "DB_PATH", tmp_path / "planos_teste.db")
    monkeypatch.setattr(database, "HISTORICO_DOCX_DIR", tmp_path / "historico_docx")
    database.init_db()


def _docx_com_aulas(*numeros_aula: int) -> bytes:
    documento = Document()
    for numero in numeros_aula:
        documento.add_paragraph(f"AULA {numero}")
    buffer = BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def _docx_com_cabecalho_plano(mes: str, bimestre: str, *numeros_aula: int) -> bytes:
    documento = Document()
    tabela = documento.add_table(rows=3, cols=9)
    rotulos = ["ESCOLA", "ESCOLA", "PROFESSOR", "COMPONENTE CURRICULAR", "COMPONENTE CURRICULAR", "COMPONENTE CURRICULAR", "TURMA", "MÊS", "BIMESTRE"]
    valores = ["ESCOLA", "ESCOLA", "ANA", "História", "História", "História", "6º ANO A", mes, bimestre]
    for indice, texto in enumerate(rotulos):
        tabela.cell(1, indice).text = texto
    for indice, texto in enumerate(valores):
        tabela.cell(2, indice).text = texto
    for numero in numeros_aula:
        documento.add_paragraph(f"AULA {numero}")
    buffer = BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def test_init_db_cria_indices_e_remove_historico_incompleto(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    with database.get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            INSERT INTO historico_planos
            (professor_nome, disciplina, turma, data_geracao, arquivo_nome, arquivo_path)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            ("ANA", "Matematica", "6 ANO A", "2026-06-05 10:00:00", "valido.docx", "valido.docx"),
        )
        cursor.execute(
            """
            INSERT INTO historico_planos
            (professor_nome, disciplina, turma, data_geracao, arquivo_nome, arquivo_path)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            ("ANA", "Matematica", "6 ANO A", "2026-06-05 10:00:01", "", "sem_nome.docx"),
        )
        cursor.execute(
            """
            INSERT INTO historico_planos
            (professor_nome, disciplina, turma, data_geracao, arquivo_nome, arquivo_path)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            ("ANA", "Matematica", "6 ANO A", "2026-06-05 10:00:02", "sem_path.docx", ""),
        )
        conn.commit()

    database.init_db()

    with database.get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("PRAGMA index_list(historico_planos)")
        indices = {row[1] for row in cursor.fetchall()}
        cursor.execute("PRAGMA table_info(historico_planos)")
        colunas = {row[1] for row in cursor.fetchall()}
        cursor.execute("SELECT arquivo_nome FROM historico_planos ORDER BY id")
        arquivos = [row[0] for row in cursor.fetchall()]
        cursor.execute(
            """
            SELECT
                professor_chave,
                disciplina_chave,
                turma_chave,
                mes_geracao,
                origem
            FROM historico_planos
            WHERE arquivo_nome = ?
            """,
            ("valido.docx",),
        )
        metadados = cursor.fetchone()

    assert "idx_historico_planos_data_id" in indices
    assert "idx_historico_planos_contexto_data" in indices
    assert "idx_historico_planos_chaves_data" in indices
    assert "idx_historico_planos_contexto_chaves_data" in indices
    assert "idx_historico_planos_prof_data" in indices
    assert {"professor_chave", "disciplina_chave", "turma_chave", "mes_geracao", "mes_plano"} <= colunas
    assert arquivos == ["valido.docx"]
    assert metadados == ("ANA", "MATEMATICA", "6 ANO A", "2026-06", "historico_docx")


def test_listar_historico_planos_tem_ordem_estavel_por_id(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    with database.get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            INSERT INTO historico_planos
            (professor_nome, disciplina, turma, data_geracao, arquivo_nome, arquivo_path)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            ("ANA", "Matematica", "6 ANO A", "2026-06-05 10:00:00", "primeiro.docx", "primeiro.docx"),
        )
        cursor.execute(
            """
            INSERT INTO historico_planos
            (professor_nome, disciplina, turma, data_geracao, arquivo_nome, arquivo_path)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            ("ANA", "Matematica", "6 ANO A", "2026-06-05 10:00:00", "segundo.docx", "segundo.docx"),
        )
        conn.commit()

    historico = database.listar_historico_planos(limite=2)

    assert [row[5] for row in historico] == ["segundo.docx", "primeiro.docx"]


def test_salvar_historico_plano_normaliza_metadados(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    database.salvar_historico_plano(
        " ANA ",
        " Matematica ",
        " 6 ANO A ",
        " plano.docx ",
        b"docx",
        bimestre=" 3o BIMESTRE ",
    )

    with database.get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT
                professor_nome,
                disciplina,
                turma,
                bimestre,
                arquivo_nome,
                arquivo_path,
                professor_chave,
                disciplina_chave,
                turma_chave,
                bimestre_chave,
                mes_geracao,
                mes_plano,
                arquivo_hash,
                arquivo_tamanho,
                origem
            FROM historico_planos
            """
        )
        row = cursor.fetchone()

    assert row[0:5] == ("ANA", "Matematica", "6 ANO A", "3o BIMESTRE", "plano.docx")
    assert row[5] != ""
    assert (tmp_path / "historico_docx" / row[5]).exists()
    assert (tmp_path / "historico_docx" / row[5]).read_bytes() == b"docx"
    assert row[6:10] == ("ANA", "MATEMATICA", "6 ANO A", "3 BIMESTRE")
    assert len(row[10]) == 7 and row[10][4] == "-"
    assert row[11] == ""
    assert len(row[12]) == 64
    assert row[13] == 4
    assert row[14] == "historico_docx"


def test_salvar_historico_plano_grava_mes_do_plano(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)
    mes_esperado = f"{database.datetime.now():%Y}-09"

    database.salvar_historico_plano(
        "ANA",
        "Historia",
        "6 ANO A",
        "plano.docx",
        _docx_com_cabecalho_plano("SETEMBRO", "3º Bimestre", 1, 2),
        bimestre="",
        mes_plano="",
    )

    resultado = database.buscar_historico_planos_avancado(
        professor_nome="ANA",
        mes=mes_esperado,
        bimestre="3 Bimestre",
    )

    assert [item["arquivo_nome"] for item in resultado] == ["plano.docx"]
    assert resultado[0]["bimestre"] == "3º Bimestre"
    assert resultado[0]["mes_plano"] == mes_esperado


def test_salvar_historico_plano_grava_resumo_de_aulas(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    database.salvar_historico_plano(
        "ANA",
        "Matematica",
        "6 ANO A",
        "plano.docx",
        _docx_com_aulas(1, 2, 2, 8),
    )

    resultado = database.buscar_historico_planos_avancado(professor_nome="ANA")

    assert resultado[0]["ultima_aula"] == 8
    assert resultado[0]["total_aulas"] == 3


def test_salvar_historico_plano_retencao_limite(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    for i in range(1, 8):
        database.salvar_historico_plano(
            "ANA",
            "Matematica",
            "6 ANO A",
            f"plano_{i}.docx",
            f"conteudo_{i}".encode("utf-8"),
            limite_retencao=5,
            bimestre="3o BIMESTRE",
        )

    with database.get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT arquivo_nome, arquivo_path FROM historico_planos ORDER BY id")
        rows = cursor.fetchall()

    arquivos = [row[0] for row in rows]

    assert len(arquivos) == 5
    assert arquivos == [
        "plano_3.docx",
        "plano_4.docx",
        "plano_5.docx",
        "plano_6.docx",
        "plano_7.docx",
    ]

    for i in range(1, 3):
        exists = any(f"plano_{i}.docx" in f.name for f in (tmp_path / "historico_docx").glob("*"))
        assert not exists

    for i in range(3, 8):
        exists = any(f"plano_{i}.docx" in f.name for f in (tmp_path / "historico_docx").glob("*"))
        assert exists


def test_salvar_historico_plano_retencao_respeita_bimestre(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    for i in range(1, 4):
        database.salvar_historico_plano(
            "ANA",
            "Matematica",
            "6 ANO A",
            f"plano_3b_{i}.docx",
            f"conteudo_3b_{i}".encode("utf-8"),
            limite_retencao=2,
            bimestre="3o BIMESTRE",
        )
    for i in range(1, 3):
        database.salvar_historico_plano(
            "ANA",
            "Matematica",
            "6 ANO A",
            f"plano_4b_{i}.docx",
            f"conteudo_4b_{i}".encode("utf-8"),
            limite_retencao=2,
            bimestre="4o BIMESTRE",
        )

    with database.get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT bimestre, arquivo_nome FROM historico_planos ORDER BY id")
        registros = cursor.fetchall()

    assert registros == [
        ("3o BIMESTRE", "plano_3b_2.docx"),
        ("3o BIMESTRE", "plano_3b_3.docx"),
        ("4o BIMESTRE", "plano_4b_1.docx"),
        ("4o BIMESTRE", "plano_4b_2.docx"),
    ]


def test_salvar_historico_remove_arquivo_novo_quando_transacao_falha(
    monkeypatch,
    tmp_path,
):
    _preparar_banco(monkeypatch, tmp_path)
    connection_scope_original = database.connection_scope

    @contextmanager
    def _connection_scope_com_falha():
        with connection_scope_original() as conn:
            yield conn
            raise RuntimeError("falha de commit simulada")

    monkeypatch.setattr(database, "connection_scope", _connection_scope_com_falha)

    with pytest.raises(RuntimeError, match="falha de commit simulada"):
        database.salvar_historico_plano(
            "ANA",
            "Matematica",
            "6 ANO A",
            "plano.docx",
            b"docx",
        )

    assert not list((tmp_path / "historico_docx").glob("*"))
    with database.get_connection() as conn:
        quantidade = conn.execute("SELECT COUNT(*) FROM historico_planos").fetchone()[0]
    assert quantidade == 0


def test_retencao_nao_apaga_arquivo_externo_ao_historico(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)
    arquivo_externo = tmp_path / "plano_sincronizado.docx"
    arquivo_externo.write_bytes(b"externo")

    with database.get_connection() as conn:
        conn.execute(
            """
            INSERT INTO historico_planos
                (professor_nome, disciplina, turma, bimestre, data_geracao, arquivo_nome, arquivo_path)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                "ANA",
                "Matematica",
                "6 ANO A",
                "3o BIMESTRE",
                "2026-01-01 10:00:00",
                arquivo_externo.name,
                str(arquivo_externo),
            ),
        )
        conn.commit()

    database.salvar_historico_plano(
        "ANA",
        "Matematica",
        "6 ANO A",
        "novo.docx",
        b"novo",
        limite_retencao=1,
        bimestre="3o BIMESTRE",
    )

    assert arquivo_externo.exists()
    with database.get_connection() as conn:
        arquivos = [
            row[0]
            for row in conn.execute(
                "SELECT arquivo_nome FROM historico_planos ORDER BY id"
            ).fetchall()
        ]
    assert arquivos == ["novo.docx"]


def test_retencao_remove_arquivo_somente_depois_do_commit(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)
    database.salvar_historico_plano(
        "ANA",
        "Matematica",
        "6 ANO A",
        "antigo.docx",
        b"antigo",
        limite_retencao=5,
        bimestre="3o BIMESTRE",
    )
    remover_original = database._remover_arquivo_historico_gerenciado
    quantidades_observadas = []

    def _remover_depois_de_consultar(arquivo_path):
        with database.get_connection() as conn:
            quantidade = conn.execute(
                "SELECT COUNT(*) FROM historico_planos WHERE arquivo_nome = ?",
                ("antigo.docx",),
            ).fetchone()[0]
        quantidades_observadas.append(quantidade)
        return remover_original(arquivo_path)

    monkeypatch.setattr(
        database,
        "_remover_arquivo_historico_gerenciado",
        _remover_depois_de_consultar,
    )

    database.salvar_historico_plano(
        "ANA",
        "Matematica",
        "6 ANO A",
        "novo.docx",
        b"novo",
        limite_retencao=1,
        bimestre="3o BIMESTRE",
    )

    assert quantidades_observadas == [0]


def test_verificar_plano_gerado_por_outro_professor_filtra_bimestre(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    database.salvar_historico_plano(
        "CARLA",
        "Historia",
        "8 ANO A",
        "hist_3b.docx",
        b"docx",
        bimestre="3o BIMESTRE",
    )
    database.salvar_historico_plano(
        "CARLA",
        "Historia",
        "8 ANO A",
        "hist_4b.docx",
        b"docx",
        bimestre="4o BIMESTRE",
    )

    outros = database.verificar_plano_gerado_por_outro_professor(
        "BIA",
        "Historia",
        "8 ANO A",
        "3o BIMESTRE",
    )

    assert len(outros) == 1
    assert outros[0]["professor_nome"] == "CARLA"
    assert outros[0]["arquivo_nome"] == "hist_3b.docx"
    assert outros[0]["bimestre"] == "3o BIMESTRE"


def test_listar_ultimos_planos_por_contexto_filtra_bimestre_e_pega_mais_recente(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    database.salvar_historico_plano(
        "ANA",
        "Matematica",
        "6 ANO A",
        "mat_3b_antigo.docx",
        b"docx-1",
        bimestre="3o BIMESTRE",
    )
    database.salvar_historico_plano(
        "ANA",
        "Matematica",
        "6 ANO A",
        "mat_3b_novo.docx",
        b"docx-2",
        bimestre="3o BIMESTRE",
    )
    database.salvar_historico_plano(
        "ANA",
        "Matematica",
        "6 ANO A",
        "mat_4b.docx",
        b"docx-3",
        bimestre="4o BIMESTRE",
    )
    database.salvar_historico_plano(
        "BIA",
        "Historia",
        "7 ANO B",
        "hist_3b.docx",
        b"docx-4",
        bimestre="3o BIMESTRE",
    )

    registros_3b = database.listar_ultimos_planos_por_contexto("3o BIMESTRE")

    assert [
        (item["professor_nome"], item["disciplina"], item["turma"], item["arquivo_nome"], item["bimestre"])
        for item in registros_3b
    ] == [
        ("ANA", "Matematica", "6 ANO A", "mat_3b_novo.docx", "3o BIMESTRE"),
        ("BIA", "Historia", "7 ANO B", "hist_3b.docx", "3o BIMESTRE"),
    ]

    registros_todos = database.listar_ultimos_planos_por_contexto()

    assert [
        (item["professor_nome"], item["disciplina"], item["turma"], item["arquivo_nome"], item["bimestre"])
        for item in registros_todos
    ] == [
        ("ANA", "Matematica", "6 ANO A", "mat_3b_novo.docx", "3o BIMESTRE"),
        ("ANA", "Matematica", "6 ANO A", "mat_4b.docx", "4o BIMESTRE"),
        ("BIA", "Historia", "7 ANO B", "hist_3b.docx", "3o BIMESTRE"),
    ]


def test_sincronizar_historico_planos_com_planos_feitos_indexa_arquivos(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)
    monkeypatch.setattr(database, "PLANOS_FEITOS_DIR", tmp_path / "PLANOS_FEITOS")
    mes_esperado = f"{database.datetime.now():%Y}-09"

    pasta = tmp_path / "PLANOS_FEITOS" / "HELOISA_MORAES_DELFINO" / "LINGUA_PORTUGUESA"
    pasta.mkdir(parents=True)
    arquivo = pasta / "Plano_6o_ANO_A_Lingua_Portuguesa.docx"
    arquivo.write_bytes(_docx_com_cabecalho_plano("SETEMBRO", "3º Bimestre", 1, 2))

    with database.get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("INSERT INTO professores (nome) VALUES (?)", ("HELOISA MORAES DELFINO",))
        conn.commit()

    inseridos = database.sincronizar_historico_planos_com_planos_feitos()
    resultados = database.buscar_historico_planos("HELOISA MORAES DELFINO")

    assert inseridos == 1
    assert len(resultados) == 1
    assert resultados[0]["disciplina"] == "LINGUA PORTUGUESA"
    assert resultados[0]["turma"] == "6O ANO A"
    assert resultados[0]["arquivo_nome"] == "Plano_6o_ANO_A_Lingua_Portuguesa.docx"
    assert resultados[0]["bimestre"] == "3º Bimestre"
    assert resultados[0]["mes_plano"] == mes_esperado

    por_mes = database.buscar_historico_planos_avancado(mes=mes_esperado)
    assert [item["arquivo_nome"] for item in por_mes] == ["Plano_6o_ANO_A_Lingua_Portuguesa.docx"]


def test_buscar_historico_planos_avancado_combina_filtros(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    database.salvar_historico_plano(
        "ANA",
        "Matematica",
        "6 ANO A",
        "mat_3b.docx",
        b"matematica",
        bimestre="3o BIMESTRE",
    )
    database.salvar_historico_plano(
        "ANA",
        "Historia",
        "6 ANO A",
        "hist_4b.docx",
        b"historia",
        bimestre="4o BIMESTRE",
    )
    database.salvar_historico_plano(
        "BIA",
        "Matematica",
        "6 ANO A",
        "bia_mat_3b.docx",
        b"bia",
        bimestre="3o BIMESTRE",
    )

    resultados = database.buscar_historico_planos_avancado(
        professor_nome=" ana ",
        disciplina="matematica",
        turma="6o ano a",
        bimestre="3 bimestre",
        termo_busca="mat",
    )

    assert [item["arquivo_nome"] for item in resultados] == ["mat_3b.docx"]
    assert resultados[0]["arquivo_disponivel"] is True
    assert resultados[0]["origem"] == "historico_docx"
    assert resultados[0]["arquivo_tamanho"] == len(b"matematica")


def test_buscar_historico_planos_avancado_inclui_legado_sem_bimestre(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    database.salvar_historico_plano(
        "HELOÍSA MORAES DELFINO",
        "Língua Portuguesa",
        "8O ANO A",
        "legado_sem_bimestre.docx",
        b"docx",
        bimestre="",
    )

    resultados = database.buscar_historico_planos_avancado(
        professor_nome="HELOÍSA MORAES DELFINO",
        bimestre="3 Bimestre",
        somente_disponiveis=True,
    )

    assert [item["arquivo_nome"] for item in resultados] == ["legado_sem_bimestre.docx"]
    assert resultados[0]["bimestre"] == ""


def test_buscar_historico_planos_avancado_filtra_disponiveis(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    database.salvar_historico_plano(
        "ANA",
        "Matematica",
        "6 ANO A",
        "disponivel.docx",
        b"ok",
    )
    with database.get_connection() as conn:
        conn.execute(
            """
            INSERT INTO historico_planos
                (
                    professor_nome,
                    disciplina,
                    turma,
                    data_geracao,
                    arquivo_nome,
                    arquivo_path,
                    professor_chave,
                    disciplina_chave,
                    turma_chave,
                    mes_geracao,
                    origem
                )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                "ANA",
                "Matematica",
                "6 ANO A",
                "2026-01-01 10:00:00",
                "ausente.docx",
                "ausente.docx",
                "ANA",
                "MATEMATICA",
                "6 ANO A",
                "2026-01",
                "historico_docx",
            ),
        )
        conn.commit()

    todos = database.buscar_historico_planos_avancado(professor_nome="ANA")
    disponiveis = database.buscar_historico_planos_avancado(
        professor_nome="ANA",
        somente_disponiveis=True,
    )

    assert {item["arquivo_nome"] for item in todos} == {"disponivel.docx", "ausente.docx"}
    assert [item["arquivo_nome"] for item in disponiveis] == ["disponivel.docx"]


def test_obter_bimestres_historico_planos_retorna_distintos(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    database.salvar_historico_plano("ANA", "Matematica", "6 ANO A", "a.docx", b"a", bimestre="3o BIMESTRE")
    database.salvar_historico_plano("ANA", "Historia", "7 ANO A", "b.docx", b"b", bimestre="3 Bimestre")
    database.salvar_historico_plano("BIA", "Historia", "8 ANO A", "c.docx", b"c", bimestre="4o BIMESTRE")

    assert database.obter_bimestres_historico_planos() == ["3 Bimestre", "4o BIMESTRE"]


def test_obter_arquivo_historico_aceita_caminho_absoluto(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)

    arquivo = tmp_path / "externo.docx"
    arquivo.write_bytes(b"conteudo-externo")

    with database.get_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            INSERT INTO historico_planos
            (professor_nome, disciplina, turma, data_geracao, arquivo_nome, arquivo_path)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                "ANA",
                "Matematica",
                "6 ANO A",
                "2026-06-05 10:00:00",
                "externo.docx",
                str(arquivo),
            ),
        )
        plano_id = int(cursor.lastrowid)
        conn.commit()

    nome, conteudo = database.obter_arquivo_historico(plano_id)

    assert nome == "externo.docx"
    assert conteudo == b"conteudo-externo"
