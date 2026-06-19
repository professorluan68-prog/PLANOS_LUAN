from docx import Document

from core.lote import _detectar_tipo_aula, _montar_etapas_metodologia, _montar_resultado_aula_local
from core.referencias_educacao_financeira import localizar_docx_referencia, referencia_por_pdf
from core.lib.classificador import perfil_disciplina
from core.lib.acompanhamento import gerar_acompanhamento_aprimorado
from core.lib.acessibilidade import gerar_acessibilidade_aprimorada
from core.lib.higienizador_pedagogico import higienizar_plano


def _criar_docx_referencia_financeira(caminho):
    doc = Document()
    doc.add_paragraph("AULA 1 - Reserva de emergencia")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para comecar: Retomar uma situacao de imprevisto financeiro e levantar hipoteses da turma.")
    doc.add_paragraph("Na pratica: Orientar o registro de receitas, gastos e valor destinado a reserva.")
    doc.add_paragraph("Encerramento: Socializar criterios para manter a reserva sem expor dados pessoais.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("\u2611 Verificar se os estudantes identificam imprevistos e custos envolvidos.")
    doc.add_paragraph("\u2611 Observar os registros de receitas, gastos e reserva planejada.")
    doc.add_paragraph("\u2611 Conferir as justificativas usadas para definir prioridades.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("\u2611 Disponibilizar quadro com receitas, gastos e reserva.")
    doc.add_paragraph("\u2611 Oferecer perguntas orientadoras em frases curtas.")
    doc.add_paragraph("\u2611 Permitir resposta oral mediada antes do registro escrito.")
    doc.save(caminho)


def test_referencia_educacao_financeira_le_docx_da_pasta_do_pdf(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Educacao_Financeira_7_Ano.docx"
    caminho_pdf = tmp_path / "AULA 1.pdf"
    _criar_docx_referencia_financeira(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_por_pdf(caminho_pdf, "1")

    assert referencia["titulo"] == "Reserva de emergencia"
    assert [etapa["titulo"] for etapa in referencia["metodologia"]] == ["Para comecar", "Na pratica", "Encerramento"]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_referencia_educacao_financeira_prefere_docx_corrigido(tmp_path):
    antigo = tmp_path / "Metodologias_Educacao_Financeira_7_Ano.docx"
    corrigido = tmp_path / "Metodologias_Educacao_Financeira_7_Ano_CORRIGIDO.docx"
    caminho_pdf = tmp_path / "AULA 1.pdf"
    _criar_docx_referencia_financeira(antigo)
    _criar_docx_referencia_financeira(corrigido)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    escolhido = localizar_docx_referencia(caminho_pdf)

    assert escolhido == corrigido


def test_referencia_educacao_financeira_pode_casar_por_titulo(tmp_path):
    doc = Document()
    doc.add_paragraph("AULA 1 - Tema antigo")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para comecar: Texto antigo.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("\u2611 Verificar um registro.")
    doc.add_paragraph("\u2611 Observar uma resposta.")
    doc.add_paragraph("\u2611 Conferir uma justificativa.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("\u2611 Apoiar com quadro.")
    doc.add_paragraph("\u2611 Apoiar com perguntas.")
    doc.add_paragraph("\u2611 Apoiar com resposta oral.")
    doc.add_paragraph("AULA 2 - Simulando rendimentos - Parte 2")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para comecar: Retomar os calculos de rendimento da parte 2.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("\u2611 Verificar os calculos de rendimento.")
    doc.add_paragraph("\u2611 Observar as comparacoes realizadas.")
    doc.add_paragraph("\u2611 Conferir a justificativa do resultado.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("\u2611 Disponibilizar tabela de apoio.")
    doc.add_paragraph("\u2611 Oferecer roteiro de calculo.")
    doc.add_paragraph("\u2611 Permitir resposta oral mediada.")
    caminho_docx = tmp_path / "Metodologias_Educacao_Financeira_7_Ano.docx"
    caminho_pdf = tmp_path / "Simulando_rendimentos_Parte_2_01.pdf"
    doc.save(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_por_pdf(caminho_pdf, "1", tema="Simulando rendimentos - Parte 2")

    assert referencia["titulo"] == "Simulando rendimentos - Parte 2"
    assert "parte 2" in referencia["metodologia"][0]["texto"]


def test_educacao_financeira_resultado_local_usa_docx_sem_trocar_titulo_da_planilha(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Educacao_Financeira_7_Ano.docx"
    caminho_pdf = tmp_path / "AULA 1.pdf"
    _criar_docx_referencia_financeira(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    resultado = _montar_resultado_aula_local(
        texto="Aula sobre juros e compras parceladas que nao deve prevalecer sobre o DOCX.",
        tema="Titulo vindo da planilha",
        material_digital="AULA 1 - Titulo vindo da planilha",
        numero_aula="1",
        disciplina_base="Educacao Financeira",
        turma="7 ano A",
        provedor_ia="",
        perfil="educacao_financeira",
        contexto_metodologico="",
        indice_aula=0,
        total_aulas=1,
        modalidade_eja_ativa=False,
        metodologia_fixa_pdf=[],
        aprendizagem_pv="",
        objetivos_orientacao=[],
        aprendizagem_orientacao="",
        usar_ia=False,
        ia_erro="",
        caminho_pdf=str(caminho_pdf),
    )

    assert resultado["tema"] == "Titulo vindo da planilha"
    assert resultado["material"] == "AULA 1 - Titulo vindo da planilha"
    assert resultado["metodologia"][0]["titulo"] == "Para comecar"
    assert "imprevisto financeiro" in resultado["metodologia"][0]["texto"]
    assert len(resultado["acompanhamento"]) == 3
    assert len(resultado["acessibilidade"]) == 3


def test_educacao_financeira_tolera_disciplina_com_caracter_quebrado():
    assert perfil_disciplina("Educa??o Financeira") == "educacao_financeira"

    etapas = _montar_etapas_metodologia(
        texto=(
            "Organize receitas, despesas, prioridades e metas em um planejamento simples, "
            "comparando escolhas e registrando criterios para definir objetivos."
        ),
        disciplina="Educa??o Financeira",
        turma="7 ano A",
        tema="Definicao de objetivos - Parte 1",
    )

    titulos = [etapa["titulo"] for etapa in etapas]
    texto = " ".join(etapa["texto"] for etapa in etapas).lower()

    assert "Analise de caso" in titulos
    assert "planejamento" in texto or "objetivos" in texto
    assert "texto literario" not in texto


def test_educacao_financeira_classifica_tipos_especificos():
    assert (
        _detectar_tipo_aula(
            "Organize receita, despesa, renda e gastos em um orcamento mensal.",
            "Orcamento pessoal",
            "Educacao Financeira",
        )
        == "orcamento_planejamento"
    )
    assert (
        _detectar_tipo_aula(
            "Compare compra parcelada, credito, juros, parcelas e custo total.",
            "Credito e juros",
            "Educacao Financeira",
        )
        == "credito_endividamento"
    )
    assert (
        _detectar_tipo_aula(
            "Analise poupanca, reserva de emergencia, rendimento e metas.",
            "Por que poupamos?",
            "Educacao Financeira",
        )
        == "investimento_poupanca"
    )
    assert (
        _detectar_tipo_aula(
            "Texto contaminado com produto, servico e viabilidade, mas o tema da aula e sobre noticias e percentuais.",
            "Percentuais na midia analisando noticias - Parte 1",
            "Educacao Financeira",
        )
        == "analise_percentuais_noticias"
    )
    assert (
        _detectar_tipo_aula(
            "Texto contaminado com juros, parcelas e custo total.",
            "Onde guardamos o dinheiro?",
            "Educacao Financeira",
        )
        == "instituicoes_financeiras"
    )
    assert (
        _detectar_tipo_aula(
            "Texto contaminado com negocio, produto e lucro.",
            "O papel do governo na economia",
            "Educacao Financeira",
        )
        == "governo_economia"
    )
    assert (
        _detectar_tipo_aula(
            "Texto contaminado com produto, servico, mercado, lucro e viabilidade.",
            "Entendendo a Economia Domestica",
            "Educacao Financeira",
        )
        == "orcamento_planejamento"
    )


def test_educacao_financeira_higieniza_residuos_de_empreendedorismo_e_frases_quebradas():
    metodologia = [
        {
            "titulo": "Para começar",
            "texto": (
                "Utilizar a técnica para que os alunos discuta impactos no orçamento. "
                "Aplicar o conversa inicial em duplas, solicitando que os alunos discutam "
                "itens essenciais para uma feira cultural, relacionando essa atividade com o orçamento familiar."
            ),
        },
        {
            "titulo": "Na prática",
            "texto": (
                "Na prática: em que os estudantes organizam receitas e despesas. "
                "Realizar uma parada estratégica para verificação para conferir a compreensão."
            ),
        },
    ]
    acompanhamento = [
        "☑ Observar se os estudantes identificam custos, preço, público, recursos necessários e viabilidade em propostas empreendedoras simples.",
        "☑ Conferir se articulam ideia, necessidade, produto ou serviço e organização financeira.",
        "☑ Acompanhar os registros.",
    ]
    acessibilidade = [
        "☑ Organizar o projeto em etapas curtas: ideia, público, recursos, custos, preço, viabilidade e revisão.",
        "☑ Utilizar quadro de apoio.",
        "☑ Permitir resposta oral.",
    ]

    metodologia_h, acompanhamento_h, acessibilidade_h = higienizar_plano(
        metodologia,
        acompanhamento,
        acessibilidade,
        perfil="educacao_financeira",
        disciplina="Educacao Financeira",
        tema="Classificando e analisando as despesas de uma familia - Parte 1",
        recursos_reais={"tabela": True},
    )
    texto = " ".join(
        [etapa["texto"] for etapa in metodologia_h]
        + acompanhamento_h
        + acessibilidade_h
    ).lower()

    assert "os alunos discuta" not in texto
    assert "para que os estudantes discutam" in texto
    assert "aplicar o conversa" not in texto
    assert "verificação para conferir" not in texto
    assert "feira cultural" not in texto
    assert "propostas empreendedoras" not in texto
    assert "produto ou serviço" not in texto
    assert "viabilidade" not in texto
    assert "orçamento familiar" in texto or "organizacao financeira" in texto or "organização financeira" in texto


def test_educacao_financeira_metodologia_usa_regras_da_analise():
    etapas = _montar_etapas_metodologia(
        texto=(
            "Compra parcelada com juros. Compare valor a vista, parcelas, credito e custo total. "
            "Na pratica, resolva as situacoes-problema e justifique a decisao."
        ),
        disciplina="Educacao Financeira",
        turma="7 ano A",
        tema="Juros e credito",
    )

    titulos = [etapa["titulo"] for etapa in etapas]
    texto = " ".join(etapa["texto"] for etapa in etapas)

    assert "Analise de caso" in titulos
    assert "Calculos financeiros" in titulos
    assert "sem exigir relatos pessoais" in texto
    assert "valor a vista" in texto or "custo total" in texto
    assert "REGISTREM" not in texto


def test_educacao_financeira_percentuais_nao_vira_empreendedorismo():
    etapas = _montar_etapas_metodologia(
        texto=(
            "Texto antigo contaminado por produto, servico, custos e viabilidade. "
            "Agora a aula precisa analisar noticias, manchetes e porcentagens com leitura de dados."
        ),
        disciplina="Educacao Financeira",
        turma="8 ano A",
        tema="Percentuais na midia analisando noticias - Parte 1",
    )

    titulos = [etapa["titulo"] for etapa in etapas]
    texto = " ".join(etapa["texto"] for etapa in etapas).lower()

    assert "Calculos financeiros" in titulos
    assert "noticias" in texto
    assert "manchetes" in texto or "graficos" in texto
    assert "projeto empreendedor" not in texto
    assert "viabilidade" not in texto


def test_educacao_financeira_acompanhamento_e_acessibilidade_por_tipo():
    acompanhamento = gerar_acompanhamento_aprimorado(
        tema="Credito e juros",
        disciplina="Educacao Financeira",
        tipo="credito_endividamento",
    )
    acessibilidade = gerar_acessibilidade_aprimorada(
        tema="Credito e juros",
        disciplina="Educacao Financeira",
        tipo="credito_endividamento",
    )

    assert any("custo total" in item for item in acompanhamento)
    assert any("valor \u00e0 vista" in item for item in acessibilidade)


def test_educacao_financeira_acompanhamento_e_acessibilidade_para_percentuais():
    acompanhamento = gerar_acompanhamento_aprimorado(
        tema="Percentuais na midia analisando noticias",
        disciplina="Educacao Financeira",
        tipo="analise_percentuais_noticias",
    )
    acessibilidade = gerar_acessibilidade_aprimorada(
        tema="Percentuais na midia analisando noticias",
        disciplina="Educacao Financeira",
        tipo="analise_percentuais_noticias",
    )

    assert any("percentuais" in item.lower() for item in acompanhamento)
    assert any("noticia" in item.lower() or "grafico" in item.lower() for item in acessibilidade)


def test_matematica_parte_2_ganha_tom_de_continuidade_sem_perder_estrutura():
    etapas = _montar_etapas_metodologia(
        texto=(
            "Princípios de contagem - Parte 2. "
            "Retome os registros da aula anterior e resolva novas situações com árvore de possibilidades."
        ),
        disciplina="Matematica",
        turma="2 ano A",
        tema="Princípios de contagem - Parte 2",
    )

    titulos = [etapa["titulo"] for etapa in etapas]
    texto = " ".join(etapa["texto"] for etapa in etapas).lower()

    assert "Para comecar" in titulos or "Para começar" in titulos
    assert "retomar brevemente o conceito central da aula anterior" in texto
    assert "com suas palavras" in texto


def test_matematica_principios_de_contagem_ativa_regras_de_combinatoria():
    etapas = _montar_etapas_metodologia(
        texto=(
            "Princípios de contagem. Resolver situações com diagrama de árvore, "
            "evento favorável e espaço amostral antes da fórmula."
        ),
        disciplina="Matematica",
        turma="2 ano A",
        tema="Princípios de contagem",
    )

    texto = " ".join(etapa["texto"] for etapa in etapas).lower()

    assert "diagrama de árvore" in texto or "diagrama de arvore" in texto
    assert "espaço amostral" in texto or "espaco amostral" in texto
