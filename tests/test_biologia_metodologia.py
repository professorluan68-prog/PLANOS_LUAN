# -*- coding: utf-8 -*-
from docx import Document

from core.lib.classificador import detectar_tipo_aula, perfil_disciplina
from core.lib.metodologia import MotorMetodologico, _etapas_por_perfil
from core.lib.acompanhamento import gerar_acompanhamento_aprimorado
from core.lib.acessibilidade import gerar_acessibilidade_aprimorada
from core.lote import _montar_resultado_aula_local
from core.referencias_biologia import (
    localizar_docx_referencia_biologia,
    referencia_biologia_por_pdf,
    titulos_referencia_biologia_por_docx,
)


def _criar_docx_referencia_biologia(caminho):
    doc = Document()
    doc.add_paragraph("AULA 1 - Ciclos biogeoquimicos do carbono e do oxigenio")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para comecar: Retomar situacoes do cotidiano relacionadas a combustao, respiracao e fotossintese para ativar conhecimentos previos.")
    doc.add_paragraph("Foco no conteudo: Orientar a leitura do esquema dos ciclos e comparar a circulacao do carbono e do oxigenio nos ambientes terrestre e aquatico.")
    doc.add_paragraph("Na pratica: Solicitar que a turma organize um quadro com processos, reservatorios e interferencias humanas identificadas no material.")
    doc.add_paragraph("Encerramento: Socializar as conclusoes e registrar uma sintese sobre equilibrio ambiental.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("\u2611 Verificar se os estudantes identificam processos e reservatorios dos ciclos estudados.")
    doc.add_paragraph("\u2611 Observar as relacoes estabelecidas entre fotossintese, respiracao e combustao.")
    doc.add_paragraph("\u2611 Conferir os registros feitos sobre impactos humanos no equilibrio ambiental.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("\u2611 Disponibilizar quadro comparativo com palavras-chave dos ciclos e seus processos.")
    doc.add_paragraph("\u2611 Oferecer roteiro de leitura guiada com perguntas curtas sobre cada esquema.")
    doc.add_paragraph("\u2611 Permitir resposta oral mediada antes do registro escrito da sintese.")
    doc.save(caminho)


def test_perfil_biologia_detectado_corretamente():
    assert perfil_disciplina("Biologia") == "biologia"


def test_tipo_aula_biologia_detectado_corretamente():
    # 1. etico_biotecnologico
    texto_etico = "Celulas HeLa: a importancia da bioetica em biotecnologia. CEP, CONEP, consentimento livre e esclarecido, dignidade e sigilo."
    assert detectar_tipo_aula(texto_etico, "Bioetica em pesquisa", "Biologia") == "etico_biotecnologico"

    # 2. debate_critico
    texto_debate = "Estudo do darwinismo social e eugenia. Racismo cientifico, pseudociencia e determinismo biologico na historia."
    assert detectar_tipo_aula(texto_debate, "Eugenia e racismo cientifico", "Biologia") == "debate_critico"

    # 3. molecular_genetico
    texto_molecular = "Bases nitrogenadas: adenina, timina, citosina e guanina. Replicacao semiconservativa do DNA e transcricao do RNA."
    assert detectar_tipo_aula(texto_molecular, "DNA e RNA", "Biologia") == "molecular_genetico"

    # 4. aplicacao_biotecnologica
    texto_biotec = "Vacinas e soros. Imunidade adquirida e resposta imunológica no Instituto Butantan e Fiocruz."
    assert detectar_tipo_aula(texto_biotec, "Imunidade e vacinacao", "Biologia") == "aplicacao_biotecnologica"

    # 5. revisao_aprofundamento
    texto_revisao = "Retomada dos conceitos de genetica mendeliana. Relembre o cruzamento de ervilhas e a segregacao."
    assert detectar_tipo_aula(texto_revisao, "Revisao de Genetica", "Biologia") == "revisao_aprofundamento"


def test_etapas_por_perfil_biologia():
    etapas_etico = _etapas_por_perfil("biologia", "etico_biotecnologico")
    chaves_etico = [e[1] for e in etapas_etico]
    assert "para_comecar" in chaves_etico
    assert "foco_1" in chaves_etico
    assert "foco_2" in chaves_etico
    assert "pause" in chaves_etico
    assert "pratica" in chaves_etico
    assert "encerramento" in chaves_etico

    etapas_molecular = _etapas_por_perfil("biologia", "molecular_genetico")
    chaves_molecular = [e[1] for e in etapas_molecular]
    assert "relembre" in chaves_molecular
    assert "foco_1" in chaves_molecular
    assert "foco_2" in chaves_molecular
    assert "pause" in chaves_molecular
    assert "pratica" in chaves_molecular
    assert "encerramento" in chaves_molecular


def test_motor_metodologico_biologia_etico():
    motor = MotorMetodologico()
    etapas = motor.gerar(
        texto_pdf='Assista ao video "A mulher que mudou a medicina" no canal Nerdologia com duracao de 7 minutos. Discutir bioetica e consentimento.',
        disciplina="Biologia",
        turma="1º ANO A",
        tema="Células HeLa e Bioética",
    )
    titulos = [e["titulo"] for e in etapas]
    assert "Para comecar" in titulos
    assert "Foco no conteudo" in titulos
    assert "Encerramento" in titulos

    textos = " ".join(e["texto"] for e in etapas)
    # Deve conter menções a vídeos e canais extraídos
    assert "Nerdologia" in textos
    assert "A mulher que mudou a medicina" in textos
    
    # Nenhuma etapa deve começar com definição direta de conceitos (ex: "X é...")
    for etapa in etapas:
        texto_etapa = etapa["texto"].strip()
        assert not texto_etapa.startswith(("Definir", "Apresentar a definição", "Explicar a definição", "O conceito de", "Conceito:"))


def test_acompanhamento_biologia_etico():
    acompanhamento = gerar_acompanhamento_aprimorado(
        tema="Células HeLa e Bioética",
        desenvolvimento='Para comecar: video sobre Henrietta Lacks. Foco no conteudo: bioetica e consentimento. Na pratica: analise do caso.',
        disciplina="Biologia",
    )
    assert len(acompanhamento) == 3
    # Todos os itens de biologia devem conter o checkmark ☑
    for item in acompanhamento:
        assert item.startswith("☑")
    
    assert any("bioética" in item.lower() or "bioetica" in item.lower() for item in acompanhamento)
    assert any("dignidade" in item.lower() or "autonomia" in item.lower() for item in acompanhamento)


def test_acessibilidade_biologia_molecular():
    acessibilidade = gerar_acessibilidade_aprimorada(
        tema="Cruzamento Genético",
        desenvolvimento='Relembre: genotipo e fenotipo. Foco no conteudo: Primeira Lei de Mendel. Na pratica: quadro de Punnett.',
        disciplina="Biologia",
    )
    assert len(acessibilidade) == 3
    # Todos os itens de biologia devem conter o checkmark ☑
    for item in acessibilidade:
        assert item.startswith("☑")

    # Deve conter templates de ferramentas práticas ou glossários
    assert any("glossário" in item.lower() or "glossario" in item.lower() for item in acessibilidade)
    assert any("punnett" in item.lower() or "heredograma" in item.lower() for item in acessibilidade)


def test_referencia_biologia_le_docx_da_pasta_do_pdf(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Biologia_1_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA 1.pdf"
    _criar_docx_referencia_biologia(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_biologia_por_pdf(caminho_pdf, "1")

    assert referencia["titulo"] == "Ciclos biogeoquimicos do carbono e do oxigenio"
    assert [etapa["titulo"] for etapa in referencia["metodologia"]] == [
        "Para comecar",
        "Foco no conteudo",
        "Na pratica",
        "Encerramento",
    ]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_titulos_referencia_biologia_por_docx_expoe_mapa_de_aulas(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Biologia_2_Ano_Ensino_Medio.docx"
    _criar_docx_referencia_biologia(caminho_docx)

    titulos = titulos_referencia_biologia_por_docx(caminho_docx)

    assert titulos == {1: "Ciclos biogeoquimicos do carbono e do oxigenio"}


def test_titulos_referencia_biologia_aceita_travessao_no_docx(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Aprofundamento_Biologia_3_Ano_Ensino_Medio.docx"
    doc = Document()
    doc.add_paragraph("AULA 01 — Organização filogenética dos grandes grupos vegetais")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Apresentar o tema central da aula.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se compreendem a organização filogenética.")
    doc.add_paragraph("☑ Observar se usam critérios de comparação.")
    doc.add_paragraph("☑ Acompanhar os registros no caderno.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Oferecer esquema visual dos grupos vegetais.")
    doc.add_paragraph("☑ Disponibilizar palavras-chave no quadro.")
    doc.add_paragraph("☑ Permitir resposta oral mediada.")
    doc.save(caminho_docx)

    titulos = titulos_referencia_biologia_por_docx(caminho_docx)

    assert titulos == {1: "Organização filogenética dos grandes grupos vegetais"}


def test_referencia_biologia_localiza_docx_da_serie(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Biologia_2_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA 1.pdf"
    _criar_docx_referencia_biologia(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    escolhido = localizar_docx_referencia_biologia(caminho_pdf)

    assert escolhido == caminho_docx


def test_biologia_resultado_local_usa_docx_sem_trocar_titulo_oficial(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Biologia_1_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA 1.pdf"
    _criar_docx_referencia_biologia(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    resultado = _montar_resultado_aula_local(
        texto="Texto qualquer do PDF que nao deve prevalecer sobre o DOCX de referencia.",
        tema="Titulo vindo da planilha",
        material_digital="AULA 1 - Titulo vindo da planilha",
        numero_aula="1",
        disciplina_base="Biologia",
        turma="1a serie A",
        provedor_ia="",
        perfil="biologia",
        contexto_metodologico="",
        indice_aula=0,
        total_aulas=1,
        modalidade_eja_ativa=False,
        metodologia_fixa_pdf=[],
        aprendizagem_pv="",
        objetivos_orientacao=[],
        aprendizagem_orientacao="",
        usar_ia=False,
        ia_erro="",
        caminho_pdf=str(caminho_pdf),
    )

    assert resultado["tema"] == "Titulo vindo da planilha"
    assert resultado["material"] == "AULA 1 - Titulo vindo da planilha"
    assert resultado["origem_metodologia"] == "docx_referencia_biologia"
    assert "combustao" in resultado["metodologia"][0]["texto"].lower()
    assert len(resultado["acompanhamento"]) == 3
    assert len(resultado["acessibilidade"]) == 3
