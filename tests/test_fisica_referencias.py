from docx import Document

from core.resultados_aula import montar_resultado_aula_local
from core.referencias_fisica import (
    localizar_docx_referencia_fisica,
    referencia_fisica_por_pdf,
    titulos_referencia_fisica_por_docx,
)
from tests.test_resultados_aula_core import _deps_resultados_base


def _criar_docx_referencia_fisica(caminho):
    doc = Document()
    doc.add_paragraph("AULA 1 - Movimento uniforme")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para comecar: Retomar situacoes cotidianas de deslocamento e velocidade para ativar conhecimentos previos.")
    doc.add_paragraph("Foco no conteudo: Explicar a relacao entre espaco, tempo e velocidade, usando exemplos simples e registros no quadro.")
    doc.add_paragraph("Na pratica: Solicitar que a turma analise dados de deslocamento e resolva exercicios guiados.")
    doc.add_paragraph("Encerramento: Sistematizar os conceitos principais e registrar uma sintese final da aula.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("OK Verificar se os estudantes identificam as grandezas fisicas envolvidas nas situacoes propostas.")
    doc.add_paragraph("OK Observar se aplicam corretamente a relacao entre espaco, tempo e velocidade.")
    doc.add_paragraph("OK Conferir os registros feitos ao longo da resolucao dos exercicios.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("OK Disponibilizar quadro de formulas, palavras-chave e exemplos resolvidos.")
    doc.add_paragraph("OK Oferecer roteiro passo a passo para leitura e interpretacao das situacoes-problema.")
    doc.add_paragraph("OK Permitir resposta oral mediada antes do registro matematico final.")
    doc.save(caminho)


def test_referencia_fisica_le_docx_da_pasta_do_pdf(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Fisica_1_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA 1.pdf"
    _criar_docx_referencia_fisica(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_fisica_por_pdf(caminho_pdf, "1")

    assert referencia["titulo"] == "Movimento uniforme"
    assert [etapa["titulo"] for etapa in referencia["metodologia"]] == [
        "Para comecar",
        "Foco no conteudo",
        "Na pratica",
        "Encerramento",
    ]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_titulos_referencia_fisica_por_docx_expoe_mapa_de_aulas(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Fisica_2_Ano_Ensino_Medio.docx"
    _criar_docx_referencia_fisica(caminho_docx)

    titulos = titulos_referencia_fisica_por_docx(caminho_docx)

    assert titulos == {1: "Movimento uniforme"}


def test_referencia_fisica_localiza_docx_da_serie(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Fisica_2_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA 1.pdf"
    _criar_docx_referencia_fisica(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    escolhido = localizar_docx_referencia_fisica(caminho_pdf)

    assert escolhido == caminho_docx


def test_fisica_resultado_local_usa_docx_sem_motor_interno(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Fisica_1_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA 1.pdf"
    _criar_docx_referencia_fisica(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    deps = _deps_resultados_base()
    deps.referencia_docx_por_perfil_fn = lambda *args, **kwargs: referencia_fisica_por_pdf(caminho_pdf, "1")
    deps.tentar_gerador_colunas_pedagogicas_fn = (
        lambda **kwargs: (_ for _ in ()).throw(AssertionError("Nao deveria montar colunas"))
    )
    deps.montar_etapas_metodologia_fn = (
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("Nao deveria montar motor local"))
    )

    resultado = montar_resultado_aula_local(
        texto="Texto qualquer do PDF.",
        tema="Titulo vindo da planilha",
        material_digital="AULA 1 - Titulo vindo da planilha",
        numero_aula="1",
        disciplina_base="Fisica",
        turma="1a serie A",
        provedor_ia="",
        perfil="fisica",
        contexto_metodologico="",
        indice_aula=0,
        total_aulas=1,
        modalidade_eja_ativa=False,
        metodologia_fixa_pdf=[],
        aprendizagem_pv="",
        objetivos_orientacao=[],
        aprendizagem_orientacao="",
        usar_ia=False,
        ia_erro="",
        dependencias=deps,
        caminho_pdf=str(caminho_pdf),
    )

    assert resultado["origem_metodologia"] == "docx_referencia_fisica"
    assert "deslocamento" in resultado["metodologia"][0]["texto"].lower()
    assert len(resultado["acompanhamento"]) == 3
    assert len(resultado["acessibilidade"]) == 3
    assert any(
        "copiados exatamente do arquivo .docx" in aviso.lower()
        for aviso in resultado["avisos_validacao"]
    )


def test_fisica_resultado_local_sem_docx_retorna_aviso_sem_motor():
    deps = _deps_resultados_base()
    deps.referencia_docx_por_perfil_fn = lambda *args, **kwargs: None
    deps.tentar_gerador_colunas_pedagogicas_fn = (
        lambda **kwargs: (_ for _ in ()).throw(AssertionError("Nao deveria montar colunas"))
    )
    deps.montar_etapas_metodologia_fn = (
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("Nao deveria montar motor local"))
    )

    resultado = montar_resultado_aula_local(
        texto="Texto qualquer do PDF.",
        tema="Titulo vindo da planilha",
        material_digital="AULA 1 - Titulo vindo da planilha",
        numero_aula="1",
        disciplina_base="Fisica",
        turma="1a serie A",
        provedor_ia="",
        perfil="fisica",
        contexto_metodologico="",
        indice_aula=0,
        total_aulas=1,
        modalidade_eja_ativa=False,
        metodologia_fixa_pdf=[],
        aprendizagem_pv="",
        objetivos_orientacao=[],
        aprendizagem_orientacao="",
        usar_ia=False,
        ia_erro="",
        dependencias=deps,
        caminho_pdf="AULA_1.pdf",
    )

    assert resultado["origem_metodologia"] == "referencia_docx_fisica_ausente"
    assert resultado["metodologia"] == []
    assert any("fisica" in aviso.lower() for aviso in resultado["avisos_validacao"])
