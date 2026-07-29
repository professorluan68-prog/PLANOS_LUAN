from io import BytesIO

from docx import Document

from core import database, gestao_aulas


def _docx_com_aulas(*numeros_aula: int) -> bytes:
    documento = Document()
    tabela = documento.add_table(rows=1, cols=2)
    tabela.rows[0].cells[0].text = "AULA"
    tabela.rows[0].cells[1].text = "APRENDIZAGEM"
    for numero_aula in numeros_aula:
        linha = tabela.add_row()
        linha.cells[0].text = "Conteúdo"
        linha.cells[1].text = f"Aula {numero_aula}"

    buffer = BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def _docx_financeiro_com_aulas(bimestre: str, *numeros_aula: int) -> bytes:
    documento = Document()
    cabecalho = documento.add_table(rows=2, cols=3)
    cabecalho.rows[0].cells[0].text = "PROFESSOR"
    cabecalho.rows[0].cells[1].text = "TURMA"
    cabecalho.rows[0].cells[2].text = "BIMESTRE"
    cabecalho.rows[1].cells[0].text = "MARTA DE ARAÚJO"
    cabecalho.rows[1].cells[1].text = "1º ANO C"
    cabecalho.rows[1].cells[2].text = bimestre

    aulas = documento.add_table(rows=len(numeros_aula), cols=2)
    for linha, numero_aula in zip(aulas.rows, numeros_aula):
        linha.cells[0].text = "Data e horário"
        linha.cells[1].text = (
            f"AULA {numero_aula} - Educação Financeira\n"
            "RECUPERAÇÃO DA APRENDIZAGEM\n1º BIMESTRE"
        )

    buffer = BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def _preparar_banco(monkeypatch, tmp_path):
    monkeypatch.setattr(database, "DB_PATH", tmp_path / "planos_teste.db")
    monkeypatch.setattr(database, "HISTORICO_DOCX_DIR", tmp_path / "historico_docx")
    database.init_db()


def test_obter_ultima_aula_gerada_sistema_impl_reinicia_da_primeira_aula(monkeypatch):
    monkeypatch.setattr(database, "obter_ultimo_plano_docx", lambda *args, **kwargs: b"docx_antigo")
    monkeypatch.setattr(gestao_aulas, "detectar_ultima_aula_de_docx_bytes", lambda *args, **kwargs: 9)
    monkeypatch.setattr(gestao_aulas, "obter_aula_parada_do_json", lambda *args, **kwargs: 7)

    ultima_aula = gestao_aulas.obter_ultima_aula_gerada_sistema_impl(
        "ANA",
        "Ciências",
        "6 ANO A",
        "3º BIMESTRE",
    )

    assert ultima_aula == 0


def test_obter_referencia_ultima_aula_historico_consulta_ultimo_plano_do_bimestre(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)
    database.salvar_historico_plano(
        "SILVANA",
        "Biologia",
        "1 ANO A",
        "plano_2_bimestre.docx",
        _docx_com_aulas(1, 2, 3),
        bimestre="2º BIMESTRE",
    )
    database.salvar_historico_plano(
        "SILVANA",
        "Biologia",
        "1 ANO A",
        "plano_3_bimestre.docx",
        _docx_com_aulas(1, 2, 9),
        bimestre="3º BIMESTRE",
    )

    referencia = gestao_aulas.obter_referencia_ultima_aula_historico(
        "Silvana",
        "Biologia",
        "1 ano a",
        "3º BIMESTRE",
    )

    assert referencia is not None
    assert referencia["ultima_aula"] == 9
    assert referencia["arquivo_nome"] == "plano_3_bimestre.docx"


def test_obter_referencia_ultima_aula_historico_reconhece_modelo_financeiro_antigo(monkeypatch, tmp_path):
    _preparar_banco(monkeypatch, tmp_path)
    database.salvar_historico_plano(
        "MARTA DE ARAÚJO",
        "EDUCAÇÃO FINANCEIRA",
        "1O ANO C",
        "plano_financeiro.docx",
        _docx_financeiro_com_aulas("3º Bimestre", 3, 4, 5, 6),
        bimestre="",
    )

    referencia = gestao_aulas.obter_referencia_ultima_aula_historico(
        "Marta de Araújo",
        "Educação Financeira",
        "1º ANO C",
        "3º Bimestre",
    )

    assert referencia is not None
    assert referencia["ultima_aula"] == 6
    assert gestao_aulas.detectar_ultima_aula_de_docx_bytes(
        _docx_financeiro_com_aulas("2º Bimestre", 9),
        "3º Bimestre",
    ) == 0
