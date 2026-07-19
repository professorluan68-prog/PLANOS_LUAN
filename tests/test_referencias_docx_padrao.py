from docx import Document

from core.referencias_docx_padrao import (
    carregar_referencias_docx_padrao,
    localizar_docx_referencia_padrao,
    referencia_docx_padrao_por_pdf,
)


def _adicionar_aula_valida(documento: Document, numero: int = 1) -> None:
    documento.add_paragraph(f"AULA {numero:02d} — Energia: formas e transformações")
    documento.add_paragraph("METODOLOGIA")
    documento.add_paragraph(
        "Para começar: Retomar a pergunta: “De onde vem a energia?”"
    )
    documento.add_paragraph(
        "Foco no conteúdo: Explicar conservação e transformação sem trocar as palavras."
    )
    documento.add_paragraph("Este parágrafo continua exatamente a etapa anterior.")
    documento.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    documento.add_paragraph("☑ Verificar o registro: conceito e exemplo.")
    documento.add_paragraph("☑ Observar a explicação oral dos estudantes.")
    documento.add_paragraph("☑ Conferir a síntese escrita ao final.")
    documento.add_paragraph("☑ Este quarto item não deve ser retornado.")
    documento.add_paragraph("ACESSIBILIDADE")
    documento.add_paragraph("☑ Oferecer apoio visual com setas e legendas.")
    documento.add_paragraph("☑ Permitir resposta oral, sem alterar o enunciado.")
    documento.add_paragraph("☑ Disponibilizar tempo ampliado quando necessário.")
    documento.add_paragraph("☑ Este quarto apoio não deve ser retornado.")


def test_localizador_prioriza_nome_singular_e_ignora_backup_e_temporario(
    tmp_path,
):
    pdf = tmp_path / "AULA_01 - Energia.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    oficial = tmp_path / "METODOLOGIA_CIENCIAS_6_ANO.docx"
    oficial.write_bytes(b"arquivo oficial")
    (tmp_path / "METODOLOGIA_CIENCIAS_6_ANO_backup.docx").write_bytes(
        b"backup"
    )
    (tmp_path / "~$METODOLOGIA_CIENCIAS_6_ANO.docx").write_bytes(b"temporario")
    (tmp_path / "Metodologias_Ciencias_6_Ano.docx").write_bytes(b"plural")

    assert localizar_docx_referencia_padrao(pdf) == oficial


def test_localizador_aceita_metodologias_quando_nao_ha_nome_singular(tmp_path):
    pdf = tmp_path / "AULA_01 - Energia.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    referencia = tmp_path / "Guia_de_Metodologias_Ciencias.docx"
    referencia.write_bytes(b"referencia")
    (tmp_path / "Guia_de_Metodologias_Ciencias_BACKUP.docx").write_bytes(
        b"backup"
    )

    assert localizar_docx_referencia_padrao(pdf) == referencia


def test_parser_preserva_textos_e_retorna_somente_tres_itens(tmp_path):
    caminho_docx = tmp_path / "METODOLOGIA_CIENCIAS_6_ANO.docx"
    documento = Document()
    _adicionar_aula_valida(documento)
    documento.save(caminho_docx)

    referencias = carregar_referencias_docx_padrao(caminho_docx)

    assert list(referencias) == [1]
    aula = referencias[1]
    assert aula["titulo"] == "Energia: formas e transformações"
    assert aula["metodologia"] == [
        {
            "titulo": "Para começar",
            "texto": "Retomar a pergunta: “De onde vem a energia?”",
        },
        {
            "titulo": "Foco no conteúdo",
            "texto": (
                "Explicar conservação e transformação sem trocar as palavras. "
                "Este parágrafo continua exatamente a etapa anterior."
            ),
        },
    ]
    assert aula["acompanhamento"] == [
        "☑ Verificar o registro: conceito e exemplo.",
        "☑ Observar a explicação oral dos estudantes.",
        "☑ Conferir a síntese escrita ao final.",
    ]
    assert aula["acessibilidade"] == [
        "☑ Oferecer apoio visual com setas e legendas.",
        "☑ Permitir resposta oral, sem alterar o enunciado.",
        "☑ Disponibilizar tempo ampliado quando necessário.",
    ]


def test_referencia_resolve_somente_numero_e_nao_substitui_aula_ausente(
    tmp_path,
):
    pdf = tmp_path / "AULA_02 - Energia formas e transformações.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    caminho_docx = tmp_path / "METODOLOGIA_CIENCIAS_6_ANO.docx"
    documento = Document()
    _adicionar_aula_valida(documento, numero=1)
    documento.save(caminho_docx)

    referencia_existente = referencia_docx_padrao_por_pdf(
        pdf,
        1,
        tema="Tema completamente diferente",
    )
    referencia_ausente = referencia_docx_padrao_por_pdf(
        pdf,
        2,
        tema="Energia: formas e transformações",
    )

    assert referencia_existente is not None
    assert referencia_existente["fonte"] == str(caminho_docx)
    assert referencia_existente["referencia_pedagogica_aplicada"] is True
    assert referencia_ausente is None


def test_parser_descarta_aula_incompleta(tmp_path):
    caminho_docx = tmp_path / "METODOLOGIA_INCOMPLETA.docx"
    documento = Document()
    documento.add_paragraph("AULA 3 - Aula incompleta")
    documento.add_paragraph("METODOLOGIA")
    documento.add_paragraph("Abertura: Texto literal da abertura.")
    documento.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    documento.add_paragraph("Item 1")
    documento.add_paragraph("Item 2")
    documento.add_paragraph("ACESSIBILIDADE")
    documento.add_paragraph("Apoio 1")
    documento.add_paragraph("Apoio 2")
    documento.add_paragraph("Apoio 3")
    documento.save(caminho_docx)

    assert carregar_referencias_docx_padrao(caminho_docx) == {}


def test_leitor_padrao_separa_habilidade_escrita_no_cabecalho(tmp_path):
    pdf = tmp_path / "AULA_001__ORIENTACAO_DE_ESTUDOS.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    caminho_docx = tmp_path / "METODOLOGIA_ORIENTACAO_DE_ESTUDOS.docx"
    documento = Document()
    documento.add_paragraph(
        "AULA 1 — Leitura de gráficos\nHABILIDADE: Localizar informação explícita."
    )
    documento.add_paragraph("METODOLOGIA")
    documento.add_paragraph("Abertura: Texto literal.")
    documento.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    documento.add_paragraph("Item 1")
    documento.add_paragraph("Item 2")
    documento.add_paragraph("Item 3")
    documento.add_paragraph("ACESSIBILIDADE")
    documento.add_paragraph("Apoio 1")
    documento.add_paragraph("Apoio 2")
    documento.add_paragraph("Apoio 3")
    documento.save(caminho_docx)

    referencia = referencia_docx_padrao_por_pdf(pdf, "1")

    assert referencia["titulo"] == "Leitura de gráficos"
    assert referencia["habilidade"] == "Localizar informação explícita."
