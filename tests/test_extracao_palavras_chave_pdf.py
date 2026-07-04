from core.extracao_palavras_chave_pdf import (
    BlocoSecao,
    PaginaExtraida,
    agrupar_paginas_por_secao,
    classificar_secao,
    detectar_elemento_principal,
    extrair_descricao_atividade,
    extrair_titulo_pagina,
    gerar_esboco,
)


def _pagina(
    numero: int,
    texto: str,
    *,
    imagem: bool = False,
    tabela: bool = False,
    tecnicas: list[str] | None = None,
    correcao: bool = False,
    numero_atividade: int | None = None,
) -> PaginaExtraida:
    return PaginaExtraida(
        numero=numero,
        texto=texto,
        linhas=[linha.strip() for linha in texto.splitlines() if linha.strip()],
        tem_imagem=imagem,
        tem_tabela=tabela,
        tecnicas=list(tecnicas or []),
        secao_detectada=classificar_secao(texto),
        eh_correcao=correcao,
        numero_atividade_detectada=numero_atividade,
    )


def test_classificar_secao_detecta_labels_principais():
    assert classificar_secao("Para começar\nPerguntas iniciais") == "Para começar"
    assert classificar_secao("Foco no conteúdo\nTexto base") == "Foco no conteúdo"
    assert classificar_secao("Na prática\nAtividade 1") == "Na prática"
    assert classificar_secao("Encerramento\nSíntese final") == "Encerramento"
    assert classificar_secao("Pause e responda\nQuestão") == "IGNORAR"


def test_classificar_secao_nao_confunde_texto_de_orientacao():
    texto = (
        "Slide 21\n"
        "Dinâmica de condução: o Encerramento foi desenvolvido para levar os estudantes\n"
        "a retomarem a aula oralmente."
    )
    assert classificar_secao(texto) is None


def test_agrupar_paginas_por_secao_agrupa_foco_e_separa_atividades():
    paginas = [
        _pagina(1, "História\nCapa"),
        _pagina(2, "Para começar\nPergunta inicial"),
        _pagina(3, "Foco no conteúdo\nTema 1"),
        _pagina(4, "Foco no conteúdo\nTema 2"),
        _pagina(5, "Na prática\nAtividade 1"),
        _pagina(6, "Correção\nComentário"),
        _pagina(7, "Na prática\nAtividade 2"),
        _pagina(8, "Encerramento\nFechamento"),
    ]

    blocos = agrupar_paginas_por_secao(paginas)

    assert [bloco.secao for bloco in blocos] == [
        "Para começar",
        "Foco no conteúdo",
        "Na prática",
        "Na prática",
        "Encerramento",
    ]
    assert len(blocos[1].paginas) == 2
    assert blocos[2].numero_atividade == 1
    assert len(blocos[2].paginas) == 2
    assert blocos[3].numero_atividade == 2


def test_agrupar_paginas_por_secao_une_para_comecar_e_mesma_atividade():
    paginas = [
        _pagina(1, "Para começar\nPergunta 1"),
        _pagina(2, "Para começar\nComentário de continuidade"),
        _pagina(3, "Na prática\nAtividade 1\nLeia o texto", numero_atividade=1),
        _pagina(4, "Na prática\nAtividade 1\nContinuação da leitura", numero_atividade=1),
        _pagina(5, "Na prática\nGabarito", correcao=True),
        _pagina(6, "Encerramento\nCom suas palavras"),
    ]

    blocos = agrupar_paginas_por_secao(paginas)

    assert [bloco.secao for bloco in blocos] == ["Para começar", "Na prática", "Encerramento"]
    assert len(blocos[0].paginas) == 2
    assert blocos[1].numero_atividade == 1
    assert len(blocos[1].paginas) == 3


def test_detectar_elemento_principal_prioriza_mapa_mental_lista_e_comparacao():
    pagina_mapa_mental = _pagina(1, "Foco no conteúdo\nMapa mental\nAtenas", imagem=True)
    pagina_lista = _pagina(2, "Foco no conteúdo\n1\nPrimeiro item\n2\nSegundo item\n3\nTerceiro item", imagem=True)
    pagina_comparacao = _pagina(3, "Foco no conteúdo\nVamos comparar?\nQuadro comparativo", tabela=True)

    assert detectar_elemento_principal(pagina_mapa_mental) == "MAPA MENTAL"
    assert detectar_elemento_principal(pagina_lista) == "Lista numerada"
    assert detectar_elemento_principal(pagina_comparacao) == "QUADRO DE COMPARAÇÃO"


def test_extrair_titulo_pagina_junta_linhas_quebradas_do_titulo():
    pagina = _pagina(
        1,
        "Foco no conteúdo\nContexto anterior às Guerras\ngreco-pérsicas\nAntes do século V a.C.",
        imagem=True,
    )

    assert extrair_titulo_pagina(pagina) == "Contexto anterior às Guerras greco-pérsicas"


def test_extrair_titulo_pagina_reconstroi_pergunta_quebrada():
    pagina = _pagina(
        1,
        "Para começar\nVocê sabe onde surgiram termos como “política” e\n“democracia”?",
        imagem=True,
    )

    assert extrair_titulo_pagina(pagina) == "Você sabe onde surgiram termos como “política” e “democracia”?"


def test_extrair_titulo_pagina_pega_apenas_primeira_pergunta_numerada():
    pagina = _pagina(
        1,
        (
            "Para começar\nCompare as imagens e suas legendas. Na sequência, debata:\n"
            "1. Como a República é representada nas imagens?\n"
            "2. Você acha que a República brasileira é parecida com o que foi a romana?"
        ),
        imagem=True,
    )

    assert extrair_titulo_pagina(pagina) == "Como a República é representada nas imagens?"


def test_extrair_titulo_pagina_ignora_rotulos_genericos():
    pagina = _pagina(
        1,
        "Foco no conteúdo\nAtenas e Esparta\nProduzido pela SEDUC-SP\nDisponível em: exemplo",
        imagem=True,
    )

    assert extrair_titulo_pagina(pagina) == "Atenas e Esparta"


def test_extrair_descricao_atividade_prefere_instrucao():
    paginas = [
        _pagina(
            1,
            "Na prática\nAtividade 1\nLeia com atenção o fragmento e responda às questões.\nHORA DA LEITURA",
            tecnicas=["HORA DA LEITURA"],
        )
    ]

    assert extrair_descricao_atividade(paginas) == "Leia com atenção o fragmento e responda às questões"


def test_gerar_esboco_monta_linhas_principais():
    blocos = [
        BlocoSecao(
            secao="Para começar",
            paginas=[
                _pagina(
                    1,
                    "Para começar\nPolítica e democracia\nVIREM E CONVERSEM",
                    imagem=True,
                    tecnicas=["VIREM E CONVERSEM"],
                )
            ],
        ),
        BlocoSecao(
            secao="Foco no conteúdo",
            paginas=[
                _pagina(2, "Foco no conteúdo\nAs cidades-estado gregas\nMapa da Grécia", imagem=True),
                _pagina(3, "Foco no conteúdo\nOs principais espaços das pólis\n1\n2\n3", imagem=True),
            ],
        ),
        BlocoSecao(
            secao="Na prática",
            numero_atividade=1,
            paginas=[
                _pagina(
                    4,
                    "Na prática\nObserve a tirinha e responda ao que se pede.\nTODO MUNDO ESCREVE",
                    tecnicas=["TODO MUNDO ESCREVE"],
                ),
                _pagina(5, "Correção\nComentário final"),
            ],
        ),
        BlocoSecao(secao="Encerramento", paginas=[_pagina(6, "Encerramento\nCom suas palavras")]),
    ]

    linhas = gerar_esboco(blocos)

    assert linhas[0] == 'PARA COMEÇAR: IMAGEM / Política e democracia – "VIREM E CONVERSEM"'
    assert linhas[1] == "FOCO NO CONTEÚDO: (2 PÁGINAS)"
    assert "1ª PÁGINA: MAPA – As cidades-estado gregas" in linhas[2]
    assert 'NA PRÁTICA: ATIVIDADE 1: Observe a tirinha e responda ao que se pede – "TODO MUNDO ESCREVE" + correção' == linhas[4]
    assert linhas[5] == "ENCERRAMENTO: Com suas palavras"


def test_extrair_palavras_chave_docx(tmp_path):
    import docx
    from docx.enum.text import WD_COLOR_INDEX
    from core.extracao_palavras_chave_pdf import extrair_palavras_chave_docx
    
    doc_path = tmp_path / "teste_realce.docx"
    doc = docx.Document()
    
    p1 = doc.add_paragraph("Este é um parágrafo de teste com ")
    run1 = p1.add_run("palavra destacada")
    run1.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p1.add_run(" no meio do texto.")
    
    p2 = doc.add_paragraph("Aqui está outra ")
    run2 = p2.add_run("segunda palavra")
    run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p2.add_run(" que deve ser capturada.")
    
    # Adicionar também em uma tabela
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p_cell = cell.paragraphs[0]
    p_cell.text = "Texto na célula com "
    run_cell = p_cell.add_run("terceira palavra")
    run_cell.font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    doc.save(doc_path)
    
    palavras = extrair_palavras_chave_docx(doc_path)
    
    assert palavras == ["palavra destacada", "segunda palavra", "terceira palavra"]

