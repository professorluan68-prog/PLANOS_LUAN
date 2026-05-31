from pathlib import Path

from docx import Document

from core.professores_planos import extrair_info_plano


def test_extrair_info_plano_preserva_quinta_apos_feriado(tmp_path: Path):
    caminho = tmp_path / "plano_portugues.docx"
    doc = Document()

    cabecalho = doc.add_table(rows=4, cols=9)
    cabecalho.rows[0].cells[0].text = "PLANO DE AULAS"
    cabecalho.rows[1].cells[2].text = "PROFESSOR"
    cabecalho.rows[2].cells[2].text = "LUÍS HENRIQUE GABRIEL"
    cabecalho.rows[2].cells[3].text = "LÍNGUA PORTUGUESA"
    cabecalho.rows[2].cells[6].text = "2º ANO B"
    cabecalho.rows[2].cells[7].text = "JUNHO"
    cabecalho.rows[3].cells[1].text = "01/06 a 05/06"
    cabecalho.rows[3].cells[3].text = "4"

    tabela_aulas = doc.add_table(rows=6, cols=6)
    tabela_aulas.rows[0].cells[0].text = "AULA"
    tabela_aulas.rows[0].cells[2].text = "APRENDIZAGEM"
    tabela_aulas.rows[0].cells[3].text = "DESENVOLVIMENTO"
    tabela_aulas.rows[1].cells[0].text = "01/06\n6ª aula\n11h30"
    tabela_aulas.rows[2].cells[0].text = "03/06\n4ª aula\n9h50"
    tabela_aulas.rows[3].cells[0].text = "08/06\n6ª aula\n11h30"
    tabela_aulas.rows[4].cells[0].text = "10/06\n4ª aula\n9h50"
    tabela_aulas.rows[5].cells[0].text = "11/06\n5ª e 6ª aula\n10h40 - 12h20"

    doc.save(caminho)

    info = extrair_info_plano(caminho, "Professor Teste")

    assert info["dia_semana"] == "Segunda - Quarta - Quinta"
    assert info["horario"] == "6ª aula 11h30, 4ª aula 9h50, 5ª e 6ª aula 10h40 - 12h20"
