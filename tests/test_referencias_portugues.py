from pathlib import Path

from docx import Document

from core.referencias_portugues import (
    localizar_docx_referencia_portugues,
    referencia_portugues_por_pdf,
)


def _criar_docx_portugues(caminho: Path, titulo_aula: str, texto_etapa: str) -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(titulo_aula).bold = True

    p = doc.add_paragraph()
    p.add_run("METODOLOGIA").bold = True
    p = doc.add_paragraph()
    p.add_run("Para começar: ").bold = True
    p.add_run(texto_etapa)
    p = doc.add_paragraph()
    p.add_run("Foco no conteúdo: ").bold = True
    p.add_run("Desenvolver a leitura e a interpretação do texto proposto.")

    p = doc.add_paragraph()
    p.add_run("ACOMPANHAMENTO DA APRENDIZAGEM").bold = True
    doc.add_paragraph("☑ Verificar a leitura.")
    doc.add_paragraph("☑ Observar a interpretação.")
    doc.add_paragraph("☑ Acompanhar os registros.")

    p = doc.add_paragraph()
    p.add_run("ACESSIBILIDADE").bold = True
    doc.add_paragraph("☑ Disponibilizar palavras-chave.")
    doc.add_paragraph("☑ Oferecer leitura guiada.")
    doc.add_paragraph("☑ Permitir resposta oral mediada.")

    doc.save(caminho)


def test_localizar_docx_referencia_portugues_prioriza_revisado_em_relatorios(tmp_path):
    pasta = tmp_path / "9_ANO"
    pasta.mkdir()
    pasta_relatorios = pasta / "RELATORIOS_CONFERENCIA_PLANOS"
    pasta_relatorios.mkdir()

    pdf = pasta / "AULA_03 - Versos que envolvem Parte 1.pdf"
    pdf.write_bytes(b"pdf")

    docx_raiz = pasta / "Plano_9o_ANO_A_Lingua_Portuguesa.docx"
    docx_revisado = pasta_relatorios / "Plano_9o_ANO_A_Lingua_Portuguesa_REVISADO.docx"

    _criar_docx_portugues(docx_raiz, "AULA 3 — Titulo antigo", "Texto antigo da raiz.")
    _criar_docx_portugues(docx_revisado, "AULA 3 — Versos que envolvem – Parte 1", "Texto revisado correto.")

    encontrado = localizar_docx_referencia_portugues(pdf)

    assert encontrado == docx_revisado


def test_referencia_portugues_por_pdf_usa_conteudo_do_revisado(tmp_path):
    pasta = tmp_path / "9_ANO"
    pasta.mkdir()
    pasta_relatorios = pasta / "RELATORIOS_CONFERENCIA_PLANOS"
    pasta_relatorios.mkdir()

    pdf = pasta / "AULA_03 - Versos que envolvem Parte 1.pdf"
    pdf.write_bytes(b"pdf")

    docx_revisado = pasta_relatorios / "Plano_9o_ANO_A_Lingua_Portuguesa_REVISADO.docx"
    _criar_docx_portugues(
        docx_revisado,
        "AULA 3 — Versos que envolvem – Parte 1",
        "Texto revisado correto.",
    )

    referencia = referencia_portugues_por_pdf(
        pdf,
        numero_aula=3,
        tema="Versos que envolvem – Parte 1",
    )

    assert referencia is not None
    assert "REVISADO" in str(referencia["fonte"]).upper()
    assert referencia["titulo"] == "Versos que envolvem – Parte 1"
    assert referencia["metodologia"][0]["texto"] == "Texto revisado correto."


def test_referencia_portugues_por_pdf_ler_docx_em_tabela(tmp_path):
    pasta = tmp_path / "9_ANO"
    pasta.mkdir()
    pasta_relatorios = pasta / "RELATORIOS_CONFERENCIA_PLANOS"
    pasta_relatorios.mkdir()

    pdf = pasta / "AULA_03 - Versos que envolvem Parte 1.pdf"
    pdf.write_bytes(b"pdf")

    docx_revisado = pasta_relatorios / "Plano_9o_ANO_A_Lingua_Portuguesa_REVISADO.docx"
    doc = Document()
    tabela = doc.add_table(rows=2, cols=6)
    tabela.rows[0].cells[1].text = "NÚMERO E TÍTULO DO MATERIAL DIGITAL"
    tabela.rows[0].cells[2].text = "APRENDIZAGEM ESSENCIAL*"
    tabela.rows[0].cells[3].text = "DESENVOLVIMENTO"
    tabela.rows[0].cells[4].text = "ACOMPANHAMENTO DA APRENDIZAGEM"
    tabela.rows[0].cells[5].text = "ACESSIBILIDADE"
    tabela.rows[1].cells[1].text = "AULA 3 - Versos que envolvem – Parte 1"
    tabela.rows[1].cells[2].text = "AE3 - Texto correto da aula 3"
    tabela.rows[1].cells[3].text = (
        "Para começar: Retomar o tema.\n"
        "Foco no conteúdo: Explorar o poema.\n"
        "Encerramento: Sistematizar a leitura."
    )
    tabela.rows[1].cells[4].text = (
        "☑ Verificar a leitura.\n"
        "☑ Observar a interpretação.\n"
        "☑ Acompanhar os registros."
    )
    tabela.rows[1].cells[5].text = (
        "☑ Disponibilizar palavras-chave.\n"
        "☑ Oferecer leitura guiada.\n"
        "☑ Permitir resposta oral mediada."
    )
    doc.save(docx_revisado)

    referencia = referencia_portugues_por_pdf(
        pdf,
        numero_aula=3,
        tema="Versos que envolvem – Parte 1",
    )

    assert referencia is not None
    assert referencia["titulo"] == "Versos que envolvem – Parte 1"
    assert referencia["habilidade"] == "AE3 - Texto correto da aula 3"
    assert referencia["metodologia"][1]["titulo"] == "Foco no conteúdo"
    assert referencia["acompanhamento"][0] == "☑ Verificar a leitura."
