from pathlib import Path

from docx import Document

from core import seletor_referencias


def test_itens_referencia_docx_normaliza_e_limita_tres():
    referencia = {
        "acompanhamento": ["Item 1", "☑ Item 2", "", "Item 4"],
    }

    itens = seletor_referencias.itens_referencia_docx(referencia, "acompanhamento")

    assert itens == ["☑ Item 1", "☑ Item 2"]


def test_sobrescrever_listas_pedagogicas_com_referencia_exige_tres_itens():
    referencia = {
        "acompanhamento": ["Item 1", "Item 2", "Item 3"],
        "acessibilidade": ["Apoio 1", "Apoio 2"],
    }

    acompanhamento, acessibilidade = (
        seletor_referencias.sobrescrever_listas_pedagogicas_com_referencia(
            referencia,
            ["Base 1"],
            ["Base A"],
        )
    )

    assert acompanhamento == ["☑ Item 1", "☑ Item 2", "☑ Item 3"]
    assert acessibilidade == ["Base A"]


def test_assinatura_docx_referencia_usa_docx_localizado(monkeypatch, tmp_path):
    caminho_pdf = tmp_path / "AULA_1.pdf"
    caminho_pdf.write_bytes(b"%PDF-1.4")
    caminho_docx = tmp_path / "referencia.docx"
    caminho_docx.write_bytes(b"docx")

    monkeypatch.setattr(
        seletor_referencias,
        "localizar_docx_referencia_por_perfil",
        lambda *args, **kwargs: caminho_docx,
    )

    assinatura = seletor_referencias.assinatura_docx_referencia(
        str(caminho_pdf),
        "Matematica",
        "1 ANO A",
    )

    assert assinatura.startswith("referencia.docx|")


def test_resolver_caminho_original_reconhece_nome_temporario_reversivel(
    monkeypatch, tmp_path
):
    import config

    raiz_pdf = tmp_path / "PDF_AULAS"
    original = (
        raiz_pdf
        / "ORIENTACAO_DE_ESTUDOS"
        / "EM"
        / "3_ANO"
        / "1_Jornada_07_Etapa1.pdf"
    )
    original.parent.mkdir(parents=True)
    original.write_bytes(b"%PDF-1.4")

    temporario = (
        tmp_path
        / "temp_upload"
        / "planos_luan_upload_AbC1__1_Jornada_07_Etapa1.pdf"
    )
    temporario.parent.mkdir()
    monkeypatch.setattr(config, "PDF_AULAS_DIR", raiz_pdf)

    resolvido = seletor_referencias._resolver_caminho_original(
        str(temporario), "orientacao_estudos", "3º ANO A"
    )

    assert resolvido == original


def test_resolver_caminho_original_reconhece_diretorio_real_do_upload_cdp(
    monkeypatch, tmp_path
):
    import config

    raiz_pdf = tmp_path / "PDF_AULAS"
    pasta_6_7 = (
        raiz_pdf
        / "HISTORIA"
        / "AF"
        / "3_BIMESTRE"
        / "CDP-EF"
        / "6_ANO_7_ANO"
    )
    pasta_8_9 = pasta_6_7.parent / "8_ANO_9_ANO"
    pasta_6_7.mkdir(parents=True)
    pasta_8_9.mkdir(parents=True)
    original = pasta_6_7 / "01 A Historia na sua vida.pdf"
    original.write_bytes(b"%PDF-1.4")
    (pasta_8_9 / original.name).write_bytes(b"%PDF-1.4")

    temporario = (
        tmp_path
        / "planos_luan_upload_diretorio"
        / "planos_luan_upload_AbC1__01 A Historia na sua vida.pdf"
    )
    temporario.parent.mkdir()
    temporario.write_bytes(b"%PDF-1.4")
    monkeypatch.setattr(config, "PDF_AULAS_DIR", raiz_pdf)

    resolvido = seletor_referencias.resolver_caminho_pdf_original(
        str(temporario),
        "Historia",
        "6o/7o E.F",
    )

    assert resolvido == original


def test_upload_cdp_nao_usa_docx_regular_da_pasta_pai(monkeypatch, tmp_path):
    import config

    raiz_pdf = tmp_path / "PDF_AULAS"
    original = (
        raiz_pdf
        / "HISTORIA"
        / "EM"
        / "3_BIMESTRE"
        / "CDP_EM"
        / "01 - ATIVIDADE 1 - Fontes.pdf"
    )
    original.parent.mkdir(parents=True)
    original.write_bytes(b"%PDF-1.4")
    docx_regular = original.parent.parent / "Metodologias_Historia_Ensino_Regular.docx"
    docx_regular.write_bytes(b"not a parsed docx")

    temporario = (
        tmp_path
        / "temp_upload"
        / "planos_luan_upload_AbC1__01 - ATIVIDADE 1 - Fontes.pdf"
    )
    temporario.parent.mkdir()
    temporario.write_bytes(b"%PDF-1.4")
    monkeypatch.setattr(config, "PDF_AULAS_DIR", raiz_pdf)

    assert (
        seletor_referencias.localizar_docx_referencia_por_perfil(
            str(temporario), "historia", ""
        )
        is None
    )
    assert (
        seletor_referencias.referencia_docx_por_perfil(
            str(temporario), "1", "Fontes", "historia"
        )
        is None
    )


def test_deve_aplicar_referencia_docx_no_resultado_ia_respeita_perfil():
    assert not seletor_referencias.deve_aplicar_referencia_docx_no_resultado_ia(
        "lingua_portuguesa_ef",
        {"metodologia": [{"titulo": "Etapa", "texto": "Texto IA"}]},
    )
    assert not seletor_referencias.deve_aplicar_referencia_docx_no_resultado_ia(
        "historia",
        {"metodologia": [{"titulo": "Etapa", "texto": "Texto IA"}]},
    )
    assert seletor_referencias.deve_aplicar_referencia_docx_no_resultado_ia(
        "historia",
        {"metodologia": []},
    )


def test_perfil_prioriza_docx_sobre_cache_json_para_perfis_com_referencia():
    assert seletor_referencias.perfil_prioriza_docx_sobre_cache_json("historia")
    assert seletor_referencias.perfil_prioriza_docx_sobre_cache_json("arte")
    assert seletor_referencias.perfil_prioriza_docx_sobre_cache_json("lingua_portuguesa_ef")
    assert seletor_referencias.perfil_prioriza_docx_sobre_cache_json("filosofia")


def test_material_aula_com_titulo_monta_rotulo_padrao():
    assert (
        seletor_referencias.material_aula_com_titulo("AULA 7", "Equacoes")
        == "AULA 7 - Equacoes"
    )


def test_seletor_usa_docx_padronizado_em_perfil_sem_leitor_especifico(tmp_path):
    pdf = tmp_path / "AULA_001__MATERIA__QUIMICA.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    docx = tmp_path / "METODOLOGIA_QUIMICA_2_ANO_3_B.docx"
    documento = Document()
    documento.add_paragraph("AULA 1 - Materia")
    documento.add_paragraph("METODOLOGIA")
    documento.add_paragraph("Abertura: Texto literal de Quimica.")
    documento.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    documento.add_paragraph("Item 1")
    documento.add_paragraph("Item 2")
    documento.add_paragraph("Item 3")
    documento.add_paragraph("ACESSIBILIDADE")
    documento.add_paragraph("Apoio 1")
    documento.add_paragraph("Apoio 2")
    documento.add_paragraph("Apoio 3")
    documento.save(docx)

    referencia = seletor_referencias.referencia_docx_por_perfil(
        str(pdf),
        "1",
        "Materia",
        "quimica",
    )

    assert referencia is not None
    assert referencia["metodologia"] == [
        {"titulo": "Abertura", "texto": "Texto literal de Quimica."}
    ]
    assert referencia["fonte"] == str(docx)
