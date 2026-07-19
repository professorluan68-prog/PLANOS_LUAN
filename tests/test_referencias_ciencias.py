import json

from docx import Document
from core import lote
from core.referencias_ciencias import (
    localizar_docx_referencia_ciencias,
    referencia_ciencias_por_pdf,
    titulos_referencia_ciencias_por_docx,
)
from core.revisao_final import VERSAO_GERADOR_ATUAL


def _criar_docx_referencia_ciencias(caminho, incluir_aula_3: bool = True):
    doc = Document()
    doc.add_paragraph("AULA 01 — A célula")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Apresentar a imagem inicial da aula e pedir que a turma levante hipóteses sobre a menor unidade que forma os seres vivos.")
    doc.add_paragraph("Foco no conteúdo: Explicar a teoria celular com apoio das imagens observadas e destacar a contribuição da microscopia para essa descoberta.")
    doc.add_paragraph("Na prática: Orientar a comparação entre exemplos de seres unicelulares e pluricelulares, com registro das características principais no caderno.")
    doc.add_paragraph("Encerramento: Retomar os pilares da teoria celular e pedir uma síntese curta sobre o que permite reconhecer uma célula.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se os estudantes identificam os pilares da teoria celular nas explicações da aula.")
    doc.add_paragraph("☑ Observar se utilizam vocabulário científico ao comparar seres unicelulares e pluricelulares.")
    doc.add_paragraph("☑ Conferir se os registros articulam imagem observada, conceito estudado e conclusão construída.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Utilizar imagens ampliadas da célula e nomear oralmente as estruturas principais antes do registro individual.")
    doc.add_paragraph("☑ Fornecer banco de palavras-chave com célula, microscópio, unicelular e pluricelular.")
    doc.add_paragraph("☑ Permitir resposta em tópicos, desenho identificado ou explicação oral mediada antes do texto completo.")

    doc.add_paragraph("AULA 02 — Desenvolvimento da microscopia")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Relacionar as imagens do material à pergunta sobre como a ciência observava estruturas invisíveis antes do microscópio.")
    doc.add_paragraph("Foco no conteúdo: Explicar a evolução da microscopia e sua importância para observar células e microrganismos.")
    doc.add_paragraph("Na prática: Orientar a leitura do caso apresentado no material e pedir que os estudantes registrem o papel do microscópio na investigação científica.")
    doc.add_paragraph("Encerramento: Solicitar que a turma sintetize como a microscopia ampliou o conhecimento sobre os seres vivos.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se reconhecem o papel do microscópio na produção de conhecimento científico.")
    doc.add_paragraph("☑ Observar se relacionam imagens, instrumentos e descobertas apresentados na aula.")
    doc.add_paragraph("☑ Conferir se o registro final explicita a função do microscópio no estudo dos seres vivos.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Usar imagem simplificada dos microscópios e destacar visualmente suas diferenças principais.")
    doc.add_paragraph("☑ Disponibilizar roteiro com perguntas orientadoras para leitura do texto e das imagens.")
    doc.add_paragraph("☑ Apoiar resposta oral mediada antes do registro escrito final.")

    if incluir_aula_3:
        doc.add_paragraph("AULA 03 — Aula fictícia")
        doc.add_paragraph("Metodologia")
        doc.add_paragraph("Para começar: Texto da aula 3.")
        doc.add_paragraph("Foco no conteúdo: Texto da aula 3.")
        doc.add_paragraph("Na prática: Texto da aula 3.")
        doc.add_paragraph("Encerramento: Texto da aula 3.")
        doc.add_paragraph("Acompanhamento da aprendizagem")
        doc.add_paragraph("☑ Verificar um registro.")
        doc.add_paragraph("☑ Observar uma fala.")
        doc.add_paragraph("☑ Conferir uma resposta.")
        doc.add_paragraph("Acessibilidade")
        doc.add_paragraph("☑ Apoiar com palavras-chave.")
        doc.add_paragraph("☑ Apoiar com perguntas.")
        doc.add_paragraph("☑ Apoiar com resposta oral.")

    doc.add_paragraph("AULA 04 — Seres procariontes")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Observar as imagens do material e levantar diferenças iniciais entre células procariontes e eucariontes.")
    doc.add_paragraph("Foco no conteúdo: Explicar a ausência ou presença de núcleo organizado e relacionar essa diferença aos exemplos apresentados.")
    doc.add_paragraph("Na prática: Orientar a comparação entre os esquemas e pedir que a turma registre características de cada grupo.")
    doc.add_paragraph("Encerramento: Socializar as comparações e retomar como identificar cada tipo celular.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se diferenciam procariontes e eucariontes a partir das imagens e explicações da aula.")
    doc.add_paragraph("☑ Observar se utilizam critérios científicos ao justificar as comparações feitas.")
    doc.add_paragraph("☑ Conferir se o registro final destaca presença ou ausência de núcleo e exemplos de cada grupo.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Oferecer imagens ampliadas e destacar visualmente núcleo, membrana e material genético.")
    doc.add_paragraph("☑ Registrar palavras-chave no quadro para apoiar a comparação entre os dois grupos celulares.")
    doc.add_paragraph("☑ Permitir resposta oral, em tópicos ou por esquema comparativo antes do texto completo.")
    doc.save(caminho)


def _criar_docx_referencia_ciencias_legada(caminho):
    doc = Document()
    doc.add_paragraph("AULA 01 — A célula")
    doc.add_paragraph("METODOLOGIA")
    doc.add_paragraph("Para começar: Aplicar o VIREM E CONVERSEM: Os estudantes analisam a imagem de uma célula.")
    doc.add_paragraph("Foco no conteúdo: Explicar a teoria celular.")
    doc.add_paragraph("Na prática: Realizar a atividade de verificar a compreensão dos alunos.")
    doc.add_paragraph("Encerramento: Realizar um breve quiz para verificar o que foi aprendido.")
    doc.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    doc.add_paragraph("☑ Verificar se os estudantes compreendem o conceito cientifico central de A célula.")
    doc.add_paragraph("☑ Observar se utilizam vocabulário cientifico, justificativas e evidências do material.")
    doc.add_paragraph("☑ Acompanhar a atividade escrita, conferindo se os registros apresentam clareza e sintese propria.")
    doc.add_paragraph("ACESSIBILIDADE")
    doc.add_paragraph("☑ Utilizar imagens, esquemas, tabelas e exemplos do material para tornar mais concreto o estudo.")
    doc.add_paragraph("☑ Organizar o registro em etapas curtas ligadas ao tema.")
    doc.add_paragraph("☑ Permitir diferentes formas de resposta, como topicos, desenho, setas, frases curtas ou explicação oral mediada.")
    doc.save(caminho)


def test_referencia_ciencias_le_docx_da_pasta_do_pdf(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Ciencias_6_Ano.docx"
    caminho_pdf = tmp_path / "AULA_01 - A célula.pdf"
    _criar_docx_referencia_ciencias(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_ciencias_por_pdf(caminho_pdf, "1")

    assert referencia is not None
    assert referencia["titulo"] == "A célula"
    assert [etapa["titulo"] for etapa in referencia["metodologia"]][:3] == [
        "Para começar",
        "Foco no conteúdo",
        "Na prática",
    ]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_titulos_ciencias_6ano_permitem_salto_da_aula_3(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Ciencias_6_Ano.docx"
    _criar_docx_referencia_ciencias(caminho_docx, incluir_aula_3=False)

    titulos = titulos_referencia_ciencias_por_docx(caminho_docx)

    assert titulos == {
        1: "A célula",
        2: "Desenvolvimento da microscopia",
        4: "Seres procariontes",
    }


def test_referencia_ciencias_localiza_docx_do_ano(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Ciencias_6_Ano.docx"
    caminho_pdf = tmp_path / "AULA_02 - Microscopia.pdf"
    _criar_docx_referencia_ciencias(caminho_docx)

    caminho_localizado = localizar_docx_referencia_ciencias(caminho_pdf)
    assert caminho_localizado is not None
    assert caminho_localizado.name == "Metodologias_Ciencias_6_Ano.docx"


def test_referencia_ciencias_aceita_texto_valido_sem_filtrar_palavras(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Ciencias_6_Ano.docx"
    caminho_pdf = tmp_path / "AULA_01 - A célula.pdf"
    _criar_docx_referencia_ciencias_legada(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_ciencias_por_pdf(caminho_pdf, "1", tema="A célula")

    assert referencia is not None
    assert referencia["titulo"] == "A célula"
    assert referencia["metodologia"][0]["texto"].startswith("Aplicar o VIREM E CONVERSEM")


def test_ciencias_prioriza_docx_em_vez_do_cache_json(monkeypatch, tmp_path):
    caminho_docx = tmp_path / "Metodologias_Ciencias_6_Ano_atualizado.docx"
    caminho_pdf = tmp_path / "AULA_01 - A célula como unidade básica da vida teoria celular.pdf"
    _criar_docx_referencia_ciencias(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4 dummy contents")

    caminho_json = caminho_pdf.with_suffix(".json")
    caminho_json.write_text(
        json.dumps(
            {
                "disciplina": "Ciências",
                "tema": "Tema antigo em cache",
                "material": "AULA 1 - Tema antigo em cache",
                "numero_aula": "1",
                "aprendizagem": "Aprendizagem antiga do cache.",
                "metodologia": [{"titulo": "Para começar", "texto": "Texto antigo em cache."}],
                "acompanhamento": ["Item antigo 1", "Item antigo 2", "Item antigo 3"],
                "acessibilidade": ["Acesso antigo 1", "Acesso antigo 2", "Acesso antigo 3"],
                "ia_usada": False,
                "ia_provedor": "",
                "ia_erro": "",
                "versao_gerador": VERSAO_GERADOR_ATUAL,
                "fingerprint_contexto": "fingerprint-antigo",
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    monkeypatch.setattr(
        lote,
        "_preparar_contexto_aula_pdf",
        lambda **kwargs: {
            "texto": "Texto da aula de Ciências.",
            "tema": "A célula como unidade básica da vida teoria celular",
            "material_digital": "AULA 1 - A célula como unidade básica da vida teoria celular",
            "numero_aula": "1",
            "cdp_contextual": False,
            "disciplina_base": "Ciências",
            "perfil": "ciencias_ef",
            "objetivos_orientacao": [],
            "aprendizagem_orientacao": "",
            "extracao_pdf": {},
            "tipo": "regular",
            "metodologia_fixa_pdf": [],
            "modalidade_eja_ativa": False,
            "contexto_metodologico": "regular",
            "escopo_pv": {},
            "aprendizagem_pv": "",
            "fonte_extracao": "pdf",
            "arquivo_fonte_extracao": str(caminho_pdf),
        },
    )

    def _montar_resultado_local_fake(**kwargs):
        referencia = referencia_ciencias_por_pdf(
            kwargs["caminho_pdf"],
            kwargs["numero_aula"],
            tema=kwargs["tema"],
        )
        assert referencia is not None
        return {
            "disciplina": "Ciências",
            "tema": referencia["titulo"],
            "material": f"AULA 1 - {referencia['titulo']}",
            "numero_aula": "1",
            "aprendizagem": "Habilidade de Ciências.",
            "metodologia": referencia["metodologia"],
            "acompanhamento": referencia["acompanhamento"],
            "acessibilidade": referencia["acessibilidade"],
            "ia_usada": False,
            "ia_provedor": "",
            "ia_erro": "",
            "origem_metodologia": "docx_referencia_ciencias",
            "fonte_referencia_metodologia": referencia["fonte"],
            "avisos_validacao": [],
        }

    monkeypatch.setattr(lote, "_montar_resultado_aula_local", _montar_resultado_local_fake)

    resultado = lote._aula_por_pdf(
        caminho_pdf=str(caminho_pdf),
        disciplina="Ciências",
        turma="6º ANO",
        bimestre="3º Bimestre",
        usar_ia=False,
        provedor_ia="",
        professor="Luan",
    )

    assert resultado["tema"] == "A célula"
    assert resultado["metodologia"][0]["titulo"] == "Para começar"
    assert resultado["cache_reutilizado"] is False
