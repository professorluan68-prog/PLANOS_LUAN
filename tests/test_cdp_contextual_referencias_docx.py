from docx import Document

from core.helpers import resolver_pasta_pdfs
from core.lote import _montar_resultado_cdp_contextual
from core.referencias_cdp_contextual import (
    localizar_docx_referencia_cdp_contextual,
    referencia_cdp_compativel,
    referencia_cdp_contextual_por_pdf,
    titulos_referencia_cdp_contextual_por_docx,
)


def _criar_docx_referencia_cdp(caminho):
    doc = Document()
    doc.add_paragraph("AULA 8 — Aquecimento global e acordos internacionais")
    doc.add_paragraph("METODOLOGIA")
    doc.add_paragraph("Para começar: Retomar impactos ambientais ja estudados.")
    doc.add_paragraph("Foco no conteúdo: Explicar aquecimento global e mitigacao.")
    doc.add_paragraph("Na prática: Analisar uma situacao-problema sobre acordos internacionais.")
    doc.add_paragraph("Encerramento: Registrar uma sintese coletiva.")
    doc.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    doc.add_paragraph("• Verificar se relacionam impactos e mitigacao. • Observar se interpretam a situacao-problema. • Conferir se registram a sintese.")
    doc.add_paragraph("ACESSIBILIDADE")
    doc.add_paragraph("• Disponibilizar palavras-chave no quadro. • Usar perguntas orientadoras curtas. • Permitir registro em topicos.")

    doc.add_paragraph("AULA 8.1 — Globalização e fluxos culturais")
    doc.add_paragraph("METODOLOGIA")
    doc.add_paragraph("Para começar: Levantar exemplos de trocas culturais presentes no cotidiano.")
    doc.add_paragraph("Foco no conteúdo: Discutir globalizacao, circulacao cultural e identidades.")
    doc.add_paragraph("Verificacao: Orientar resposta breve sobre influencias culturais.")
    doc.add_paragraph("Na prática: Comparar exemplos de fluxos culturais em diferentes escalas.")
    doc.add_paragraph("Encerramento: Socializar conclusoes e registrar uma ideia central.")
    doc.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    doc.add_paragraph("• Verificar se identificam exemplos de fluxos culturais. • Observar se comparam escalas de circulacao. • Acompanhar se justificam a ideia central.")
    doc.add_paragraph("ACESSIBILIDADE")
    doc.add_paragraph("• Apresentar exemplos proximos da turma. • Organizar quadro comparativo simples. • Permitir resposta oral mediada.")
    doc.save(caminho)


def test_referencia_cdp_contextual_le_aula_decimal_e_bullets_no_mesmo_paragrafo(tmp_path):
    caminho_docx = tmp_path / "metodologias.docx"
    caminho_pdf = tmp_path / "2_ano_AULA_08.1 - Globalizacao e fluxos culturais.pdf"
    _criar_docx_referencia_cdp(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_cdp_contextual_por_pdf(caminho_pdf, "8")

    assert localizar_docx_referencia_cdp_contextual(caminho_pdf) == caminho_docx
    assert referencia["numero"] == "8.1"
    assert referencia["titulo"] == "Globalização e fluxos culturais"
    assert [etapa["titulo"] for etapa in referencia["metodologia"]] == [
        "Para começar",
        "Foco no conteúdo",
        "Verificacao",
        "Na prática",
        "Encerramento",
    ]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3
    assert all(item.startswith("☑") for item in referencia["acompanhamento"])


def test_titulos_cdp_contextual_por_docx_mantem_aula_decimal(tmp_path):
    caminho_docx = tmp_path / "metodologias.docx"
    _criar_docx_referencia_cdp(caminho_docx)

    titulos = titulos_referencia_cdp_contextual_por_docx(caminho_docx)

    assert titulos["8"] == "Aquecimento global e acordos internacionais"
    assert titulos["8.1"] == "Globalização e fluxos culturais"


def test_referencia_cdp_contextual_nao_usa_numero_quando_titulo_nao_bate(tmp_path):
    caminho_docx = tmp_path / "metodologias.docx"
    caminho_pdf = tmp_path / "2_ano_AULA_05 - Fluxos de capitais e investimentos internacionais.pdf"
    _criar_docx_referencia_cdp(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_cdp_contextual_por_pdf(
        caminho_pdf,
        "5",
        tema="Fluxos de capitais e investimentos internacionais",
    )

    assert referencia is None


def test_referencia_cdp_contextual_preserva_numeros_repetidos_por_titulo(tmp_path):
    caminho_docx = tmp_path / "metodologias.docx"
    caminho_pdf = tmp_path / "2_ano_AULA_05 - Fluxos de capitais e investimentos internacionais.pdf"
    _criar_docx_referencia_cdp(caminho_docx)
    doc = Document(caminho_docx)
    doc.add_paragraph("AULA 7 — Fluxos de capitais e investimentos internacionais")
    doc.add_paragraph("METODOLOGIA")
    doc.add_paragraph("Para começar: Retomar exemplos de empresas globais na regiao.")
    doc.add_paragraph("Foco no conteúdo: Explicar fluxos de capitais e investimentos internacionais.")
    doc.add_paragraph("Na prática: Analisar impactos em economias emergentes.")
    doc.add_paragraph("Encerramento: Registrar uma conclusao sobre o mercado financeiro global.")
    doc.add_paragraph("ACOMPANHAMENTO DA APRENDIZAGEM")
    doc.add_paragraph("• Verificar se diferenciam fluxos de capitais. • Observar se analisam impactos economicos. • Conferir se registram conclusao coerente.")
    doc.add_paragraph("ACESSIBILIDADE")
    doc.add_paragraph("• Usar glossario de palavras-chave. • Organizar esquema de causa e efeito. • Permitir resposta oral mediada.")
    doc.save(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_cdp_contextual_por_pdf(
        caminho_pdf,
        "5",
        tema="Fluxos de capitais e investimentos internacionais",
    )

    assert referencia["titulo"] == "Fluxos de capitais e investimentos internacionais"
    assert "investimentos internacionais" in referencia["metodologia"][1]["texto"]


def test_referencia_cdp_compativel_rejeita_lemov_e_agrupamentos():
    assert not referencia_cdp_compativel(
        {
            "metodologia": [
                {"titulo": "Etapa", "texto": "Usar a tecnica LEMOV Virem e conversem."}
            ],
            "acompanhamento": [],
            "acessibilidade": [],
        }
    )
    assert not referencia_cdp_compativel(
        {
            "metodologia": [
                {"titulo": "Etapa", "texto": "Realizar atividade em grupos."}
            ],
            "acompanhamento": [],
            "acessibilidade": [],
        }
    )
    assert referencia_cdp_compativel(
        {
            "metodologia": [
                {"titulo": "Etapa", "texto": "Orientar registro individual no caderno."}
            ],
            "acompanhamento": ["Verificar a compreensão."],
            "acessibilidade": ["Usar palavras-chave no quadro."],
        }
    )


def test_resolver_pasta_pdfs_cdp_ensino_medio_multisseriado(tmp_path):
    pasta = tmp_path / "CDP_ENSINO_MEDIO" / "GEOGRAFIA_CDP" / "3_BIMESTRE" / "MULTISSERIADO_J"
    pasta.mkdir(parents=True)
    (pasta / "1_ano_AULA_03 - Consumo sustentavel.pdf").write_bytes(b"%PDF-1.4\n")

    caminho = resolver_pasta_pdfs(
        str(tmp_path),
        "CDP-ENSINO MÉDIO",
        "MULTISSERIADO J",
        "3º Bimestre",
    )

    assert caminho == pasta


def test_resolver_pasta_pdfs_geografia_cdp_ensino_medio_multisseriado(tmp_path):
    pasta = tmp_path / "CDP_ENSINO_MEDIO" / "GEOGRAFIA_CDP" / "3_BIMESTRE" / "MULTISSERIADO_J"
    pasta.mkdir(parents=True)
    (pasta / "1_ano_AULA_03 - Consumo sustentavel.pdf").write_bytes(b"%PDF-1.4\n")

    caminho = resolver_pasta_pdfs(
        str(tmp_path),
        "Geografia CDP Ensino Médio",
        "MULTISSERIADO J",
        "3º Bimestre",
    )

    assert caminho == pasta


def test_resultado_cdp_contextual_usa_docx_referencia(tmp_path):
    caminho_docx = tmp_path / "metodologias.docx"
    caminho_pdf = tmp_path / "2_ano_AULA_08.1 - Globalizacao e fluxos culturais.pdf"
    _criar_docx_referencia_cdp(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    resultado = _montar_resultado_cdp_contextual(
        texto="Ensino Medio Geografia Globalizacao e fluxos culturais",
        tema="Globalizacao e fluxos culturais",
        disciplina_base="Geografia",
        numero_aula="8",
        indice_aula=0,
        perfil="geografia",
        tipo="",
        extracao_pdf={"conceito_extraido": "globalizacao e fluxos culturais"},
        caminho_pdf=str(caminho_pdf),
    )

    assert resultado["numero_aula"] == "8.1"
    assert resultado["tema"] == "Globalização e fluxos culturais"
    assert resultado["origem_metodologia"] == "docx_referencia_cdp_contextual"
    assert "circulacao cultural" in resultado["metodologia"][1]["texto"].lower()
    assert len(resultado["acompanhamento"]) == 3
    assert len(resultado["acessibilidade"]) == 3
    assert any(
        "copiados exatamente do arquivo .docx" in aviso.lower()
        for aviso in resultado["avisos_validacao"]
    )


def test_resultado_cdp_contextual_sem_docx_usa_gerador_local(tmp_path):
    caminho_pdf = tmp_path / "2_ano_AULA_09 - Tema sem referencia.pdf"
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    resultado = _montar_resultado_cdp_contextual(
        texto="Ensino Medio Geografia tema sem referencia",
        tema="Tema sem referencia",
        disciplina_base="Geografia",
        numero_aula="9",
        indice_aula=0,
        perfil="geografia",
        tipo="",
        extracao_pdf={"conceito_extraido": "tema sem referencia"},
        caminho_pdf=str(caminho_pdf),
    )

    assert resultado["origem_metodologia"] == "fallback_cdp_contextual"
    assert resultado["metodologia"]
    assert all(
        isinstance(etapa, dict) and etapa.get("titulo") and etapa.get("texto")
        for etapa in resultado["metodologia"]
    )
    assert len(resultado["acompanhamento"]) == 3
    assert len(resultado["acessibilidade"]) == 3
    assert any(
        "nao encontrei o arquivo .docx de referencia" in aviso.lower()
        for aviso in resultado["avisos_validacao"]
    )
