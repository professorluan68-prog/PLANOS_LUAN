from docx import Document

from core.lote import _montar_resultado_aula_local
from core.referencias_lideranca_oratoria import (
    localizar_docx_referencia_lideranca_oratoria,
    referencia_lideranca_oratoria_por_pdf,
    titulos_referencia_lideranca_oratoria_por_docx,
)


def _criar_docx_referencia_lideranca(caminho):
    doc = Document()
    doc.add_paragraph("AULA 1 — Fundamentos da negociação")
    doc.add_paragraph("Metodologia")
    doc.add_paragraph("Para começar: Apresentar uma situação simples de negociação cotidiana.")
    doc.add_paragraph("Foco no conteúdo: Explicar escuta ativa, interesses e construção de acordos.")
    doc.add_paragraph("Na prática: Propor simulação curta em duplas com papéis combinados.")
    doc.add_paragraph("Encerramento: Socializar estratégias usadas e registrar uma síntese.")
    doc.add_paragraph("Acompanhamento da aprendizagem")
    doc.add_paragraph("☑ Verificar se os estudantes identificam interesses e posições na negociação.")
    doc.add_paragraph("☑ Observar se utilizam escuta ativa durante a simulação.")
    doc.add_paragraph("☑ Acompanhar se registram estratégias para construir acordos respeitosos.")
    doc.add_paragraph("Acessibilidade")
    doc.add_paragraph("☑ Disponibilizar roteiro com etapas da negociação.")
    doc.add_paragraph("☑ Oferecer exemplos concretos antes da simulação.")
    doc.add_paragraph("☑ Permitir participação por fala, escrita ou mediação em dupla.")
    doc.save(caminho)


def test_referencia_lideranca_oratoria_le_docx_da_pasta_do_pdf(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Lideranca_e_Oratoria_2_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA_01 - Fundamentos da negociacao.pdf"
    _criar_docx_referencia_lideranca(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    referencia = referencia_lideranca_oratoria_por_pdf(caminho_pdf, "1")

    assert localizar_docx_referencia_lideranca_oratoria(caminho_pdf) == caminho_docx
    assert referencia["titulo"] == "Fundamentos da negociação"
    assert [etapa["titulo"] for etapa in referencia["metodologia"]] == [
        "Para começar",
        "Foco no conteúdo",
        "Na prática",
        "Encerramento",
    ]
    assert len(referencia["acompanhamento"]) == 3
    assert len(referencia["acessibilidade"]) == 3


def test_titulos_lideranca_oratoria_por_docx(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Lideranca_e_Oratoria_2_Ano_Ensino_Medio.docx"
    _criar_docx_referencia_lideranca(caminho_docx)

    titulos = titulos_referencia_lideranca_oratoria_por_docx(caminho_docx)

    assert titulos == {1: "Fundamentos da negociação"}


def test_lideranca_resultado_local_usa_docx_sem_trocar_titulo_oficial(tmp_path):
    caminho_docx = tmp_path / "Metodologias_Lideranca_e_Oratoria_2_Ano_Ensino_Medio.docx"
    caminho_pdf = tmp_path / "AULA_01 - Fundamentos da negociacao.pdf"
    _criar_docx_referencia_lideranca(caminho_docx)
    caminho_pdf.write_bytes(b"%PDF-1.4\n")

    resultado = _montar_resultado_aula_local(
        texto="Texto do PDF que nao deve prevalecer sobre o DOCX de referencia.",
        tema="Fundamentos da negociação",
        material_digital="AULA 1 - Fundamentos da negociação",
        numero_aula="1",
        disciplina_base="Liderança e Oratória",
        turma="2a serie A",
        provedor_ia="",
        perfil="lideranca_oratoria",
        contexto_metodologico="",
        indice_aula=0,
        total_aulas=1,
        modalidade_eja_ativa=False,
        metodologia_fixa_pdf=[],
        aprendizagem_pv="",
        objetivos_orientacao=[],
        aprendizagem_orientacao="",
        usar_ia=False,
        ia_erro="",
        caminho_pdf=str(caminho_pdf),
    )

    assert resultado["origem_metodologia"] == "docx_referencia_lideranca_oratoria"
    assert "escuta ativa" in resultado["metodologia"][1]["texto"].lower()
    assert len(resultado["acompanhamento"]) == 3
    assert len(resultado["acessibilidade"]) == 3
