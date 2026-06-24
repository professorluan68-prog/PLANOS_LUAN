import os
import glob
import json
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    # Caminho dos JSONs
    json_dir = r"D:\PDF novos\LIDERANCA_E_ORATORIA\EM\3_BIMESTRE\2_ANO"
    json_files = glob.glob(os.path.join(json_dir, "*.json"))
    
    # Criar documento
    doc = docx.Document()
    
    # Título principal
    title = doc.add_heading("REFERÊNCIA METODOLÓGICA - LIDERANÇA E ORATÓRIA CDP", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Processar cada arquivo
    for json_file in sorted(json_files):
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        tema = data.get("tema", "Sem tema")
        aula_num = data.get("numero_aula", "")
        
        doc.add_heading(f"Aula {aula_num} - {tema}", level=2)
        
        # Metodologia
        doc.add_heading("Metodologia / Desenvolvimento:", level=3)
        metodologia = data.get("diagnostico_geracao", {}).get("metodologia_final", [])
        if not metodologia:
            metodologia = data.get("metodologia", [])
            
        for m in metodologia:
            p = doc.add_paragraph()
            p.add_run(f"[{m.get('titulo', '')}]: ").bold = True
            p.add_run(m.get('texto', ''))
            
        # Acompanhamento da aprendizagem
        doc.add_heading("Acompanhamento da Aprendizagem:", level=3)
        acompanhamento = data.get("acompanhamento", [])
        if isinstance(acompanhamento, list):
            for a in acompanhamento:
                doc.add_paragraph(a, style='List Bullet')
        else:
            doc.add_paragraph(str(acompanhamento))
            
        # Acessibilidade
        doc.add_heading("Acessibilidade:", level=3)
        acessibilidade = data.get("acessibilidade", [])
        if isinstance(acessibilidade, list):
            for ac in acessibilidade:
                doc.add_paragraph(ac, style='List Bullet')
        else:
            doc.add_paragraph(str(acessibilidade))
            
        doc.add_paragraph("\n")

    # Salvar docx
    output_path = r"d:\PLANOS_LUAN\REFERENCIAS_METODOLOGIA\LIDERANCA_E_ORATORIA_CDP_metodologia.docx"
    doc.save(output_path)
    print(f"Documento criado com sucesso em: {output_path}")

if __name__ == "__main__":
    create_docx()
