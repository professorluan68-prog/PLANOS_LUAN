from io import BytesIO
from copy import deepcopy
from datetime import date, datetime, timedelta
import re
import unicodedata
import logging

logger = logging.getLogger(__name__)

from docx import Document
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from core.helpers import texto_lista
from core.qualidade_metodologica import (
    CORRECOES_ORTOGRAFIA,
    corrigir_mojibake,
    corrigir_ortografia_basica,
)


DESTAQUES_TEXTO = {
    "VIREM E CONVERSEM": "VIREM E CONVERSEM",
    "TODO MUNDO ESCREVE": "TODO MUNDO ESCREVE",
    "COM SUAS PALAVRAS": "COM SUAS PALAVRAS",
    "HORA DA LEITURA": "HORA DA LEITURA",
    "DE OLHO NO MODELO": "DE OLHO NO MODELO",
    "UM PASSO DE CADA VEZ": "UM PASSO DE CADA VEZ",
    "LISTEN AND REPEAT": "Listen and repeat",
    "WRITE AND SHARE": "Write and share",
    "SAY IT IN ENGLISH": "Say it in English",
}

TITULOS_ETAPAS = {
    "Para comecar": "Para come\u00e7ar",
    "Contextualizacao": "Contextualiza\u00e7\u00e3o",
    "Leitura analitica": "Leitura anal\u00edtica",
    "Leitura e construcao do conteudo": "Leitura e constru\u00e7\u00e3o do conte\u00fado",
    "Exploracao": "Explora\u00e7\u00e3o",
    "Disparo inicial / contextualizacao": "Disparo inicial / contextualização",
    "Leitura ou exploracao inicial": "Leitura ou exploração inicial",
    "Analise guiada": "Análise guiada",
    "Foco no conteudo": "Foco no conte\u00fado",
    "Formalizacao": "Formaliza\u00e7\u00e3o",
    "Sistematizacao": "Sistematiza\u00e7\u00e3o",
    "Sistematiza\u00e7\u00e3o": "Sistematiza\u00e7\u00e3o",
    "Pause e responda": "Pause e responda",
    "Na pratica": "Na pr\u00e1tica",
    "Analise de caso": "An\u00e1lise de caso",
    "Calculos financeiros": "C\u00e1lculos financeiros",
    "Planejamento orcamentario": "Planejamento or\u00e7ament\u00e1rio",
    "Projeto empreendedor": "Projeto empreendedor",
    "Producao textual": "Produção textual",
    "Revisao e reescrita": "Revis\u00e3o e reescrita",
    "Revisao e fechamento": "Revisão e fechamento",
    "Relembre": "Relembre",
    "Encerramento": "Encerramento",
    "Para come\u00e7ar": "Para come\u00e7ar",
    "Contextualiza\u00e7\u00e3o": "Contextualiza\u00e7\u00e3o",
    "Leitura anal\u00edtica": "Leitura anal\u00edtica",
    "Leitura e constru\u00e7\u00e3o do conte\u00fado": "Leitura e constru\u00e7\u00e3o do conte\u00fado",
    "Explora\u00e7\u00e3o": "Explora\u00e7\u00e3o",
    "Na pr\u00e1tica": "Na pr\u00e1tica",
    "An\u00e1lise de caso": "An\u00e1lise de caso",
    "C\u00e1lculos financeiros": "C\u00e1lculos financeiros",
    "Planejamento or\u00e7ament\u00e1rio": "Planejamento or\u00e7ament\u00e1rio",
    "Revis\u00e3o e reescrita": "Revis\u00e3o e reescrita",
}


def _substituir_texto(paragraph, substituicoes: dict[str, str]) -> None:
    """
    Substitui placeholders preservando a formatação do primeiro run.
    Se o parágrafo tem apenas um run, preserva fonte, tamanho, negrito e cor.
    """
    if not paragraph.runs:
        return
    texto_original = paragraph.text
    texto_novo = texto_original
    for chave, valor in substituicoes.items():
        texto_novo = texto_novo.replace(chave, _sanitizar_texto_xml(valor))
    if texto_novo == texto_original:
        return

    # Preservar formatação do primeiro run antes de limpar
    primeiro_run = paragraph.runs[0]
    fonte_nome = primeiro_run.font.name
    fonte_tamanho = primeiro_run.font.size
    fonte_bold = primeiro_run.bold
    fonte_cor = primeiro_run.font.color.rgb if primeiro_run.font.color and primeiro_run.font.color.type else None

    paragraph.clear()
    novo_run = paragraph.add_run(_sanitizar_texto_xml(texto_novo))

    # Restaurar formatação
    if fonte_nome:
        novo_run.font.name = fonte_nome
    if fonte_tamanho:
        novo_run.font.size = fonte_tamanho
    if fonte_bold is not None:
        novo_run.bold = fonte_bold
    if fonte_cor is not None:
        novo_run.font.color.rgb = fonte_cor


def _substituir_em_tabela(tabela, substituicoes: dict[str, str]) -> None:
    for linha in tabela.rows:
        for celula in linha.cells:
            for paragrafo in celula.paragraphs:
                _substituir_texto(paragrafo, substituicoes)
            for tabela_interna in celula.tables:
                _substituir_em_tabela(tabela_interna, substituicoes)


def _texto_metodologia(aula: dict) -> str:
    metodologia = aula.get("metodologia") or []
    blocos = []
    for item in metodologia:
        if isinstance(item, dict):
            titulo = item.get("titulo", "")
            texto = item.get("texto", "")
            blocos.append(f"{titulo}\n{texto}".strip())
        else:
            blocos.append(str(item))
    return "\n\n".join(blocos)


def _texto_metodologia_lista(metodologia) -> str:
    blocos = []
    for item in metodologia or []:
        if isinstance(item, dict):
            blocos.append(f"{item.get('titulo', '')}\n{item.get('texto', '')}".strip())
        else:
            blocos.append(str(item))
    return "\n\n".join(blocos)


# ── Constantes de formatação ────────────────────────────────────────────────
_FONTE_PADRAO = "Arial"
_TAMANHO_PADRAO = Pt(10)
_COR_VERMELHA = RGBColor(0xEE, 0x00, 0x00)
_LARGURAS_TABELA_AULAS = [900, 2100, 2350, 6100, 1900, 2050]
_TURNOS_REFERENCIA_AULAS = (
    ["07h", "07h50", "08h40", "09h50", "10h40", "11h30", "12h20"],
    ["13h", "13h50", "14h40", "15h50", "16h40", "17h30", "18h20"],
    ["19h", "19h45", "20h30", "21h30", "22h15", "23h"],
)
_PADRAO_BNCC = re.compile(r'(\([A-Z]{2}\d{2}[A-Z]{2,4}\d{0,3}[A-Z]?\))')
_PADRAO_TURMA_METODOLOGIA = re.compile(
    r"\b(da turma|com a turma)\s+\d{1,2}\s*[º°oªa?]?\s*(?:ano|s[ée]rie|em|ef)?\s*[A-Z]?\b",
    flags=re.I,
)

_CORRECOES_TEXTO_FINAL = dict(CORRECOES_ORTOGRAFIA)


def _aplicar_fonte(run, nome=_FONTE_PADRAO, tamanho=_TAMANHO_PADRAO, bold=None, color=None):
    """Aplica formatação padrão a um run."""
    run.font.name = nome
    run.font.size = tamanho
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def _elemento_filho(parent, tag: str):
    filho = parent.find(qn(tag))
    if filho is None:
        filho = OxmlElement(tag)
        parent.append(filho)
    return filho


def _remover_filhos(parent, tag: str) -> None:
    for filho in list(parent.findall(qn(tag))):
        parent.remove(filho)


def _definir_largura_celula(celula, largura: int) -> None:
    tc_pr = celula._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(largura))
    tc_w.set(qn("w:type"), "dxa")
    _remover_filhos(tc_pr, "w:noWrap")


def _definir_margens_celula(celula, margem: int = 90) -> None:
    tc_pr = celula._tc.get_or_add_tcPr()
    tc_mar = _elemento_filho(tc_pr, "w:tcMar")
    for lado in ("top", "left", "bottom", "right"):
        item = _elemento_filho(tc_mar, f"w:{lado}")
        item.set(qn("w:w"), str(margem))
        item.set(qn("w:type"), "dxa")


def _remover_alturas_fixas(tabela) -> None:
    for linha in tabela.rows:
        tr_pr = linha._tr.get_or_add_trPr()
        _remover_filhos(tr_pr, "w:trHeight")


def _normalizar_layout_tabela_aulas(tabela) -> None:
    _encurtar_cabecalho_data_horario(tabela)
    if tabela._tbl.tblGrid is not None and len(tabela._tbl.tblGrid.xpath("w:gridCol")) > 6:
        return
    tabela.autofit = False
    tbl_pr = tabela._tbl.tblPr
    tbl_w = _elemento_filho(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), str(sum(_LARGURAS_TABELA_AULAS)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = _elemento_filho(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = tabela._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tabela._tbl.insert(0, grid)
    for col in list(grid):
        grid.remove(col)
    for largura in _LARGURAS_TABELA_AULAS:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(largura))
        grid.append(col)

    _remover_alturas_fixas(tabela)
    for linha in tabela.rows:
        celulas = _celulas_unicas(linha)
        for indice, celula in enumerate(celulas[: len(_LARGURAS_TABELA_AULAS)]):
            _definir_largura_celula(celula, _LARGURAS_TABELA_AULAS[indice])
            _definir_margens_celula(celula)


def _encurtar_cabecalho_data_horario(tabela) -> None:
    if not tabela.rows:
        return
    celulas = _celulas_unicas(tabela.rows[0])
    if not celulas:
        return
    primeira = celulas[0]
    texto = _normalizar_cabecalho_coluna(primeira.text)
    if "AULA SEMANAL" not in texto and "DATA" not in texto:
        return
    _limpar_celula(primeira)
    paragrafo = _paragrafo_base(primeira)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.paragraph_format.space_before = Pt(0)
    paragrafo.paragraph_format.space_after = Pt(0)
    _aplicar_fonte(paragrafo.add_run("DATA\nHOR."), tamanho=Pt(8.5), bold=True)


def _tamanho_por_texto(texto: str, padrao: float = 10.0, medio: float = 9.0, pequeno: float = 8.5) -> Pt:
    tamanho = padrao
    texto = str(texto or "")
    if len(texto) > 520:
        tamanho = pequeno
    elif len(texto) > 160:
        tamanho = medio
    return Pt(tamanho)


def _ajustar_fonte_celula(celula, tamanho: Pt) -> None:
    for paragrafo in celula.paragraphs:
        paragrafo.paragraph_format.space_before = Pt(0)
        paragrafo.paragraph_format.space_after = Pt(0)
        for run in paragrafo.runs:
            run.font.size = tamanho


def _limpar_celula(celula) -> None:
    celula.text = ""


def _paragrafo_base(celula):
    if not celula.paragraphs:
        return celula.add_paragraph()
    return celula.paragraphs[0]


def _normalizar_destaques(texto: str) -> str:
    texto_final = str(texto or "")
    for original, exibicao in DESTAQUES_TEXTO.items():
        texto_final = texto_final.replace(original, exibicao)
    return texto_final


def _capitalizar_como(original: str, corrigido: str) -> str:
    if original.isupper():
        return corrigido.upper()
    if original[:1].isupper():
        return corrigido[:1].upper() + corrigido[1:]
    return corrigido


def _sanitizar_texto_xml(texto: str) -> str:
    texto = str(texto or "")
    return "".join(
        ch
        for ch in texto
        if ch in ("\t", "\n", "\r") or 0x20 <= ord(ch) <= 0xD7FF or 0xE000 <= ord(ch) <= 0xFFFD
    )


def _polir_texto_docx(texto: str) -> str:
    texto_final = _PADRAO_TURMA_METODOLOGIA.sub(lambda m: m.group(1), _sanitizar_texto_xml(texto))
    texto_final = "\n".join(
        re.sub(r"[ \t\r\f\v]+", " ", linha).strip()
        for linha in texto_final.splitlines()
    ).strip()
    for tecnica in DESTAQUES_TEXTO.keys():
        texto_final = re.sub(
            rf"\(\s*{re.escape(tecnica)}\s*\)",
            f"“{tecnica}”",
            texto_final,
            flags=re.I,
        )
        texto_final = re.sub(
            rf"(?<![\"“])\b{re.escape(tecnica)}\b(?![\"”])",
            f"“{tecnica}”",
            texto_final,
            flags=re.I,
        )
    texto_final = re.sub(r"\bde o conceito\b", "do conceito", texto_final, flags=re.I)
    texto_final = re.sub(r"\b1o\b", "1º", texto_final, flags=re.I)
    texto_final = re.sub(r"\b1\s*o\s+grau\b", "1º grau", texto_final, flags=re.I)
    for sem_acento, com_acento in _CORRECOES_TEXTO_FINAL.items():
        texto_final = re.sub(
            rf"\b{re.escape(sem_acento)}\b",
            lambda m, novo=com_acento: _capitalizar_como(m.group(0), novo),
            texto_final,
            flags=re.I,
        )
    texto_final = corrigir_ortografia_basica(texto_final)
    return corrigir_mojibake(texto_final)


def _validar_docx_gerado(buffer: BytesIO) -> BytesIO:
    conteudo = buffer.getvalue()
    Document(BytesIO(conteudo))
    buffer.seek(0)
    return buffer


def _titulo_exibicao(titulo: str) -> str:
    valor = TITULOS_ETAPAS.get(str(titulo or "").strip(), str(titulo or "").strip())
    return corrigir_mojibake(valor)


def _adicionar_texto_com_destaques(paragrafo, texto: str) -> None:
    restante = _normalizar_destaques(texto)
    if not restante:
        return

    padrao = "|".join(re.escape(valor) for valor in DESTAQUES_TEXTO.values())
    if not padrao:
        paragrafo.add_run(restante)
        return

    partes = re.split(f"({padrao})", restante)
    for parte in partes:
        if not parte:
            continue
        run = paragrafo.add_run(parte)
        if parte in DESTAQUES_TEXTO.values():
            run.bold = True


def _preencher_celula_centralizada(celula, texto: str, bold: bool = False, color=None) -> None:
    _limpar_celula(celula)
    paragrafo = _paragrafo_base(celula)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragrafo.add_run(_polir_texto_docx(str(texto or "")))
    _aplicar_fonte(run, tamanho=_tamanho_por_texto(texto), bold=bold, color=color)


def _preencher_celula_tema_material(celula, texto: str) -> None:
    bruto = str(texto or "").strip()
    if not bruto:
        _limpar_celula(celula)
        return
    if not _polir_texto_docx(bruto).upper().startswith("TEMA:"):
        _preencher_celula_centralizada(celula, bruto, bold=True, color=_COR_VERMELHA)
        return

    linhas = [linha.strip() for linha in bruto.splitlines() if linha.strip()]
    rotulo = linhas[0] if linhas else "TEMA:"
    titulo = " ".join(linhas[1:]).strip() if len(linhas) > 1 else ""

    _limpar_celula(celula)
    primeiro = _paragrafo_base(celula)
    primeiro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    primeiro.paragraph_format.space_before = Pt(0)
    primeiro.paragraph_format.space_after = Pt(0)
    _aplicar_fonte(primeiro.add_run(_polir_texto_docx(rotulo)), tamanho=Pt(9), bold=True, color=_COR_VERMELHA)

    if titulo:
        paragrafo_titulo = celula.add_paragraph()
        paragrafo_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafo_titulo.paragraph_format.space_before = Pt(0)
        paragrafo_titulo.paragraph_format.space_after = Pt(0)
        _aplicar_fonte(
            paragrafo_titulo.add_run(_polir_texto_docx(titulo)),
            tamanho=_tamanho_por_texto(titulo, medio=9.0, pequeno=8.5),
            bold=True,
        )


def _preencher_celula_data_horario(celula, texto: str) -> None:
    _limpar_celula(celula)
    linhas = [linha.strip() for linha in str(texto or "").splitlines() if linha.strip()]
    if not linhas:
        return

    primeiro = _paragrafo_base(celula)
    for indice, linha in enumerate(linhas):
        paragrafo = primeiro if indice == 0 else celula.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafo.paragraph_format.space_before = Pt(0)
        paragrafo.paragraph_format.space_after = Pt(0)
        run = paragrafo.add_run(_polir_texto_docx(linha))
        _aplicar_fonte(run, tamanho=Pt(9), bold=False, color=_COR_VERMELHA)


def _preencher_celula_lista(celula, itens) -> None:
    _limpar_celula(celula)
    itens_lista = list(itens or [])
    if not itens_lista:
        return

    tamanho = _tamanho_por_texto(" ".join(str(item or "") for item in itens_lista), medio=8.8, pequeno=8.2)
    primeiro = _paragrafo_base(celula)
    for indice, item in enumerate(itens_lista):
        paragrafo = primeiro if indice == 0 else celula.add_paragraph()
        paragrafo.alignment = None
        texto_item = re.sub(r"^\s*(?:\u2611|\u2713|\u2022|-)\s*", "", str(item or "").strip())
        texto_item = _polir_texto_docx(texto_item)
        run_check = paragrafo.add_run("\u2611 ")
        run_check.font.name = "Segoe UI Symbol"
        run_check.font.size = tamanho
        _adicionar_texto_com_destaques_formatado(paragrafo, texto_item, tamanho=tamanho)


def _adicionar_texto_com_destaques_formatado(paragrafo, texto: str, tamanho=_TAMANHO_PADRAO) -> None:
    """Igual a _adicionar_texto_com_destaques mas aplica Arial 10pt a cada run."""
    restante = _normalizar_destaques(texto)
    if not restante:
        return
    padrao = "|".join(re.escape(valor) for valor in DESTAQUES_TEXTO.values())
    if not padrao:
        _aplicar_fonte(paragrafo.add_run(restante), tamanho=tamanho)
        return
    partes = re.split(f"({padrao})", restante)
    for parte in partes:
        if not parte:
            continue
        run = paragrafo.add_run(parte)
        _aplicar_fonte(run, tamanho=tamanho, bold=True if parte in DESTAQUES_TEXTO.values() else None)


def _remover_acentos(texto: str) -> str:
    texto = unicodedata.normalize("NFKD", str(texto or ""))
    return "".join(ch for ch in texto if not unicodedata.combining(ch))


def _normalizar_para_busca(texto: str) -> str:
    """Normaliza texto apenas para fins de comparação — NUNCA usar no conteúdo final."""
    return _remover_acentos(texto).lower().strip()


def _eh_aula_educacao_financeira(aula: dict) -> bool:
    return "educacao financeira" in _normalizar_para_busca(aula.get("disciplina") or "")


def _limitar_texto_etapa_docx(texto: str, max_frases: int = 2, max_chars: int = 280) -> str:
    texto = _polir_texto_docx(texto)
    frases = re.split(r"(?<=[.!?])\s+", texto)
    if len(frases) <= max_frases and len(texto) <= max_chars:
        return texto

    resumo = " ".join(frase.strip() for frase in frases[:max_frases] if frase.strip()).strip()
    if not resumo:
        resumo = texto[:max_chars].rsplit(" ", 1)[0].strip()
    if len(resumo) > max_chars:
        resumo = resumo[:max_chars].rsplit(" ", 1)[0].strip()
    if resumo and resumo[-1] not in ".!?":
        resumo += "."
    return resumo


def _metodologia_compacta_educacao_financeira_docx(metodologia) -> list:
    itens = [item for item in list(metodologia or []) if item]
    if len(_texto_metodologia_lista(itens)) <= 1700 and len(itens) <= 5:
        return itens

    por_titulo = {}
    for item in itens:
        if not isinstance(item, dict):
            continue
        titulo = str(item.get("titulo") or "").strip()
        titulo_norm = _normalizar_para_busca(titulo)
        por_titulo.setdefault(titulo_norm, item)

    ordem_preferida = [
        ("para comecar", "Para começar"),
        ("foco no conteudo", "Foco no conteúdo"),
        ("pause e responda", "Pause e responda"),
        ("na pratica", "Na prática"),
        ("encerramento", "Encerramento"),
    ]

    compactos = []
    for chave, titulo_padrao in ordem_preferida:
        item = por_titulo.get(chave)
        if not item:
            continue
        compactos.append(
            {
                "titulo": item.get("titulo") or titulo_padrao,
                "texto": _limitar_texto_etapa_docx(item.get("texto", ""), max_frases=3, max_chars=420),
            }
        )

    if compactos:
        return compactos

    resultado = []
    for item in itens[:5]:
        if isinstance(item, dict):
            resultado.append(
                {
                    "titulo": item.get("titulo", ""),
                    "texto": _limitar_texto_etapa_docx(item.get("texto", ""), max_frases=3, max_chars=420),
                }
            )
        else:
            resultado.append(_limitar_texto_etapa_docx(str(item), max_frases=3, max_chars=420))
    return resultado


def _metodologia_para_docx(aula: dict):
    metodologia = aula.get("metodologia")
    if _eh_aula_educacao_financeira(aula):
        return _metodologia_compacta_educacao_financeira_docx(metodologia)
    return metodologia


def _preencher_celula_aprendizagem(celula, texto: str) -> None:
    """Preenche a coluna Aprendizagem: código BNCC em vermelho+bold, resto bold preto, centralizado."""
    _limpar_celula(celula)
    paragrafo = _paragrafo_base(celula)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    texto = _polir_texto_docx(str(texto or "").strip())
    if not texto:
        return
    tamanho = _tamanho_por_texto(texto, medio=9.0, pequeno=8.5)
    # Tenta separar o código BNCC do texto descritivo
    match = _PADRAO_BNCC.search(texto)
    if match:
        codigo = match.group(1)
        pos = match.start()
        antes = texto[:pos].strip()
        depois = texto[pos + len(codigo):].strip()
        if antes:
            _aplicar_fonte(paragrafo.add_run(antes + " "), tamanho=tamanho, bold=True)
        _aplicar_fonte(paragrafo.add_run(codigo + " "), tamanho=tamanho, bold=True, color=_COR_VERMELHA)
        if depois:
            _aplicar_fonte(paragrafo.add_run(depois), tamanho=tamanho, bold=True)
    else:
        _aplicar_fonte(paragrafo.add_run(texto), tamanho=tamanho, bold=True)


_TITULOS_METODOLOGIA_INLINE = [
    "Para comecar",
    "Disparo inicial / contextualizacao",
    "Abertura",
    "Abertura e contextualizacao",
    "Contextualizacao",
    "Leitura ou exploracao inicial",
    "Analise guiada",
    "Foco no conteudo",
    "Conceituacao",
    "Exploracao conceitual",
    "Desenvolvimento guiado",
    "Esclarecimento de conceitos",
    "Atividade",
    "Atividade principal",
    "Atividade em grupo",
    "Atividade pratica",
    "Atividade de legendar",
    "Classificacao",
    "Discussao",
    "Discussao em duplas",
    "Discussao em grupo",
    "Aplicacao",
    "Aprofundamento",
    "Producao textual",
    "Socializacao",
    "Socializacao e correcao",
    "Socializacao/correcao",
    "Socializacao das descobertas",
    "Socializacao de resultados",
    "Correcao",
    "Revisao e fechamento",
    "Conclusao",
    "Sintese e fechamento",
    "Fechamento",
    "Fechamento reflexivo",
    "Fechamento e reflexao",
]


def _quebrar_texto_metodologia_em_linhas(texto: str) -> list[str]:
    texto_base = _polir_texto_docx(texto)
    if not texto_base:
        return []

    partes = [linha.strip() for linha in texto_base.splitlines() if linha.strip()]
    padroes = sorted({_titulo_exibicao(t) for t in _TITULOS_METODOLOGIA_INLINE}, key=len, reverse=True)

    linhas: list[str] = []
    for parte in partes:
        trecho = parte
        for titulo in padroes:
            trecho = re.sub(
                rf"(?<=[.!?])\s+({re.escape(titulo)}:)",
                r"\n\1",
                trecho,
                flags=re.I,
            )
        linhas.extend(linha.strip() for linha in trecho.splitlines() if linha.strip())
    return linhas


def _texto_ja_comeca_com_etapa(texto: str) -> bool:
    primeira_linha = next((linha for linha in _quebrar_texto_metodologia_em_linhas(texto) if linha.strip()), "")
    return bool(re.match(r"^[^:]{2,40}:\s*", primeira_linha))


def _titulo_metodologia_deve_prefixar(titulo: str, texto: str) -> bool:
    titulo_limpo = str(titulo or "").strip()
    if not titulo_limpo:
        return False
    if _texto_ja_comeca_com_etapa(texto):
        return False

    titulo_norm = re.sub(r"\s+", " ", titulo_limpo).strip().lower()
    linhas = _quebrar_texto_metodologia_em_linhas(texto)
    if titulo_norm == "desenvolvimento" and any(re.match(r"^[^:]{2,40}:\s*", linha) for linha in linhas):
        return False

    return True


def _preencher_celula_metodologia(celula, metodologia) -> None:
    _limpar_celula(celula)
    itens = list(metodologia or [])
    if not itens:
        return

    primeiro = _paragrafo_base(celula)
    primeiro.text = ""
    paragrafo_atual = primeiro
    tamanho = _tamanho_por_texto(_texto_metodologia_lista(metodologia), medio=9.2, pequeno=8.8)

    for indice, item in enumerate(itens):
        if isinstance(item, dict):
            titulo = str(item.get("titulo") or "").strip()
            texto = str(item.get("texto") or "").strip()
            if _titulo_metodologia_deve_prefixar(titulo, texto):
                texto_item = f"{_titulo_exibicao(_normalizar_destaques(titulo))}: {texto}"
            else:
                texto_item = texto
        else:
            texto_item = str(item).strip()

        linhas = _quebrar_texto_metodologia_em_linhas(texto_item)
        for linha in linhas:
            if not linha:
                continue

            if paragrafo_atual.text:
                paragrafo_atual = celula.add_paragraph()

            # Procurar por um padrão "Titulo: texto" para colocar em negrito
            match = re.match(r'^([^:]{2,60}):\s*(.*)$', linha)
            if match:
                titulo_bold = match.group(1) + ":"
                resto_texto = " " + match.group(2)
                _aplicar_fonte(paragrafo_atual.add_run(titulo_bold), tamanho=tamanho, bold=True)
                _adicionar_texto_com_destaques_formatado(paragrafo_atual, resto_texto, tamanho=tamanho)
            else:
                _adicionar_texto_com_destaques_formatado(paragrafo_atual, linha, tamanho=tamanho)


def _texto_tabela(tabela) -> str:
    texto = " ".join(celula.text.upper() for linha in tabela.rows for celula in linha.cells)
    return " ".join(texto.split())


def _eh_cabecalho_plano(tabela) -> bool:
    texto = _texto_tabela(tabela)
    return len(tabela.rows) >= 4 and "PLANO DE AULAS" in texto and "PROFESSOR" in texto


def _eh_tabela_aulas(tabela) -> bool:
    texto = _texto_tabela(tabela)
    return len(tabela.rows) >= 2 and "AULA SEMANAL" in texto and "APRENDIZAGEM" in texto


def _celulas_unicas(linha):
    vistas = set()
    celulas = []
    for celula in linha.cells:
        chave = id(celula._tc)
        if chave not in vistas:
            vistas.add(chave)
            celulas.append(celula)
    return celulas


def _definir_texto(celula, texto: str) -> None:
    celula.text = _sanitizar_texto_xml(texto)


def _limpar_linha(linha) -> None:
    for celula in _celulas_unicas(linha):
        _definir_texto(celula, "")


def _remover_tabela(tabela) -> None:
    elemento = tabela._element
    elemento.getparent().remove(elemento)


def _remover_linha(linha) -> None:
    elemento = linha._tr
    elemento.getparent().remove(elemento)


def _paragrafo_ooxml_vazio(elemento) -> bool:
    if elemento.tag != qn("w:p"):
        return False
    if "".join(elemento.xpath(".//w:t/text()")).strip():
        return False
    if elemento.xpath(".//w:drawing|.//w:pict|.//w:object"):
        return False
    return True


def _remover_paragrafos_vazios_finais(documento) -> None:
    body = documento._element.body
    filhos = list(body)
    indice = len(filhos) - 1
    while indice >= 0 and filhos[indice].tag == qn("w:sectPr"):
        indice -= 1
    while indice >= 0 and _paragrafo_ooxml_vazio(filhos[indice]):
        body.remove(filhos[indice])
        indice -= 1


def _clonar_par_semana(pares: list[tuple]) -> tuple:
    cabecalho_ref, tabela_ref = pares[-1]
    novo_cabecalho_xml = deepcopy(cabecalho_ref._element)
    nova_tabela_xml = deepcopy(tabela_ref._element)
    
    p = OxmlElement("w:p")
    tabela_ref._element.addnext(p)
    p.addnext(novo_cabecalho_xml)
    novo_cabecalho_xml.addnext(nova_tabela_xml)
    
    return (
        Table(novo_cabecalho_xml, cabecalho_ref._parent),
        Table(nova_tabela_xml, tabela_ref._parent),
    )


def _formatar_horario_modelo(texto: str) -> str:
    texto = re.sub(r"\s+", " ", str(texto or "")).strip()
    texto = texto.replace(" - ", " – ").replace("-", "–")

    def remover_zero(match):
        hora = str(int(match.group(1)))
        minutos = match.group(2) or ""
        return f"{hora}h{minutos}"

    return re.sub(r"\b0?(\d{1,2})h(\d{2})?\b", remover_zero, texto, flags=re.I)


def _formatar_data_horario(aula: dict) -> str:
    data_bruta = aula.get("data")
    if isinstance(data_bruta, datetime):
        data = data_bruta.strftime("%d/%m")
    elif isinstance(data_bruta, date):
        data = data_bruta.strftime("%d/%m")
    else:
        data = str(data_bruta or "").strip()
    horario = str(aula.get("horario") or "").strip()
    partes = [parte.strip() for parte in horario.splitlines() if parte.strip()]

    if len(partes) >= 2:
        return "\n".join([data, partes[1], _formatar_horario_modelo(partes[0])]).strip()
    if horario:
        return "\n".join([data, _formatar_horario_modelo(horario)]).strip()
    return data


def _extrair_horarios_do_texto(texto: str) -> list[str]:
    horarios = []
    for hora, minuto in re.findall(r"\b0?(\d{1,2})h(\d{2})?\b", str(texto or ""), flags=re.I):
        horarios.append(f"{int(hora):02d}h{minuto or ''}")
    return horarios


def _quantidade_aulas_por_horario(horario) -> int:
    texto = str(horario or "").strip()
    if not texto:
        return 0

    horarios = _extrair_horarios_do_texto(texto)
    if len(horarios) >= 2:
        inicio, fim = horarios[0], horarios[1]
        for slots in _TURNOS_REFERENCIA_AULAS:
            if inicio in slots and fim in slots:
                inicio_idx = slots.index(inicio)
                fim_idx = slots.index(fim)
                if fim_idx > inicio_idx:
                    return fim_idx - inicio_idx

    numeros = [int(valor) for valor in re.findall(r"\b(\d+)\s*(?:ª|º|a|o)\b", texto.lower())]
    if numeros:
        return max(numeros) - min(numeros) + 1

    return 1


def _quantidade_aulas_semana(aulas_da_semana) -> int:
    total = 0
    for _, aula in aulas_da_semana or []:
        total += _quantidade_aulas_por_horario((aula or {}).get("horario"))
    return total


def _data_ddmm(texto: str):
    if isinstance(texto, datetime):
        return date(2000, texto.month, texto.day)
    if isinstance(texto, date):
        return date(2000, texto.month, texto.day)
    partes = str(texto or "").strip().split("/")
    if len(partes) >= 2:
        try:
            return date(2000, int(partes[1]), int(partes[0]))
        except ValueError:
            return None

    partes_iso = str(texto or "").strip().split("-")
    if len(partes_iso) == 3:
        try:
            return date(2000, int(partes_iso[1]), int(partes_iso[2]))
        except ValueError:
            return None
    return None


def _data_para_semana(data_bruta):
    if isinstance(data_bruta, datetime):
        return data_bruta.date()
    if isinstance(data_bruta, date):
        return data_bruta
    data_parseada = _data_ddmm(data_bruta)
    if data_parseada:
        return date(date.today().year, data_parseada.month, data_parseada.day)
    return None


def _intervalo_cabecalho(tabela):
    if len(tabela.rows) < 4 or len(tabela.rows[3].cells) < 2:
        return None
    texto = tabela.rows[3].cells[1].text
    if " a " not in texto:
        return None
    inicio, fim = texto.split(" a ", 1)
    inicio_data = _data_ddmm(inicio)
    fim_data = _data_ddmm(fim)
    if not inicio_data or not fim_data:
        return None
    return inicio_data, fim_data


def _aula_pertence_ao_intervalo(aula: dict, intervalo) -> bool:
    data_aula = _data_ddmm(aula.get("data"))
    if not data_aula or not intervalo:
        return False
    inicio, fim = intervalo
    return inicio <= data_aula <= fim


def _semana_automatica_por_aulas(aulas_da_semana) -> str:
    datas = []
    for _, aula in aulas_da_semana or []:
        data_aula = aula.get("data") if isinstance(aula, dict) else None
        data_semana = _data_para_semana(data_aula)
        if data_semana:
            datas.append(data_semana)

    if not datas:
        return ""

    referencia = min(datas)
    inicio = referencia - timedelta(days=referencia.weekday())
    fim = inicio + timedelta(days=4)
    return f"{inicio.strftime('%d/%m')} a {fim.strftime('%d/%m')}"


def _inicio_semana_aula(aula: dict):
    data_aula = aula.get("data") if isinstance(aula, dict) else None
    data_semana = _data_para_semana(data_aula)
    if data_semana:
        return data_semana - timedelta(days=data_semana.weekday())
    return None


def _agrupar_sobras_por_semana(sobras):
    grupos = []
    sem_data = []
    indices_por_semana = {}
    for numero, aula in sobras:
        inicio = _inicio_semana_aula(aula)
        if inicio is None:
            sem_data.append((numero, aula))
            continue
        if inicio not in indices_por_semana:
            indices_por_semana[inicio] = len(grupos)
            grupos.append([])
        grupos[indices_por_semana[inicio]].append((numero, aula))
    return grupos, sem_data


def _semana_atual_cabecalho(tabela) -> str:
    if len(tabela.rows) < 4 or len(tabela.rows[3].cells) < 2:
        return ""
    return str(tabela.rows[3].cells[1].text or "").strip()


def _titulo_aula(aula: dict, numero: int) -> str:
    if aula.get("aula_vazia"):
        return ""
        
    material = str(aula.get("material") or aula.get("titulo_material") or "").strip()
    if material:
        return material

    titulo = str(aula.get("titulo") or aula.get("tema") or "").strip()
    numero_aula = str(aula.get("numero_aula") or "").strip()
    rotulo = f"AULA {numero_aula}" if numero_aula else f"AULA {numero}"
    if not titulo:
        return rotulo
    if titulo.upper().startswith("AULA"):
        return titulo
    return f"{rotulo} - {titulo}"


def _preencher_cabecalho(
    tabela,
    escola: str,
    professor: str,
    disciplina: str,
    turma: str,
    mes: str,
    bimestre: str,
    semana: str,
    observacao: str,
    aulas_previstas: str,
) -> None:
    if len(tabela.rows) < 4:
        return

    linha_dados = tabela.rows[2].cells
    if len(linha_dados) >= 9:
        if escola:
            _definir_texto(linha_dados[0], escola)
        _definir_texto(linha_dados[2], professor)
        _definir_texto(linha_dados[3], disciplina)
        _definir_texto(linha_dados[6], turma)
        _definir_texto(linha_dados[7], mes)
        _definir_texto(linha_dados[8], bimestre)

    linha_semana = tabela.rows[3].cells
    if len(linha_semana) >= 4:
        if semana:
            _definir_texto(linha_semana[1], semana)
        _definir_texto(linha_semana[3], aulas_previstas)
    if observacao and len(linha_semana) >= 6:
        _definir_texto(linha_semana[5], observacao)


def _normalizar_cabecalho_coluna(texto: str) -> str:
    texto = unicodedata.normalize("NFD", texto or "")
    texto = "".join(ch for ch in texto if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", texto).strip().upper()


def _indice_cabecalho(celulas, termo: str) -> int | None:
    termo_norm = _normalizar_cabecalho_coluna(termo)
    for indice, celula in enumerate(celulas):
        if termo_norm in _normalizar_cabecalho_coluna(celula.text):
            return indice
    return None


def _indice_distinto_apos(celulas, indice_base: int) -> int | None:
    if indice_base < 0 or indice_base >= len(celulas):
        return None
    tc_base = celulas[indice_base]._tc
    for indice in range(indice_base + 1, len(celulas)):
        if celulas[indice]._tc is not tc_base:
            return indice
    return None


def _indices_linha_aula(linha, cabecalho=None) -> dict[str, int] | None:
    celulas = linha.cells
    if len(celulas) < 6:
        return None

    indices = {
        "data": 0,
        "material": 1,
        "aprendizagem": 2,
        "desenvolvimento": 3,
        "acompanhamento": 4,
        "acessibilidade": 5,
    }
    if cabecalho is not None and getattr(cabecalho, "cells", None):
        cab = cabecalho.cells
        indices["data"] = _indice_cabecalho(cab, "AULA") or indices["data"]
        indices["material"] = _indice_cabecalho(cab, "MATERIAL") or indices["material"]
        indices["aprendizagem"] = _indice_cabecalho(cab, "APRENDIZAGEM") or indices["aprendizagem"]
        indices["desenvolvimento"] = _indice_cabecalho(cab, "DESENVOLVIMENTO") or indices["desenvolvimento"]
        indices["acompanhamento"] = _indice_cabecalho(cab, "ACOMPANHAMENTO") or indices["acompanhamento"]
        indices["acessibilidade"] = _indice_cabecalho(cab, "ACESSIBILIDADE") or indices["acessibilidade"]

    acomp_idx = indices["acompanhamento"]
    acess_idx = indices["acessibilidade"]
    if acomp_idx < len(celulas) and acess_idx < len(celulas) and celulas[acess_idx]._tc is celulas[acomp_idx]._tc:
        distinto = _indice_distinto_apos(celulas, acomp_idx)
        if distinto is not None:
            indices["acessibilidade"] = distinto

    if any(indice >= len(celulas) for indice in indices.values()):
        return None
    return indices


def _preencher_linha_aula(linha, aula: dict, numero: int, cabecalho=None) -> None:
    celulas = linha.cells
    indices = _indices_linha_aula(linha, cabecalho)
    if not indices:
        return

    usadas = set()
    for campo, idx in indices.items():
        if idx >= len(celulas):
            continue
        tc_id = id(celulas[idx]._tc)
        if tc_id in usadas:
            raise RuntimeError(f"Colisão de gridSpan detectada: o campo '{campo}' tentou usar a mesma célula de outro campo na tabela.")
        usadas.add(tc_id)


    # Col 0: Data/Horário — vermelho, centralizado, Arial 10
    _preencher_celula_data_horario(celulas[indices["data"]], _formatar_data_horario(aula))
    # Col 1: Título — vermelho + bold, centralizado, Arial 10
    _preencher_celula_tema_material(celulas[indices["material"]], _titulo_aula(aula, numero))
    # Col 2: Aprendizagem — código BNCC vermelho, texto bold preto, centralizado
    _preencher_celula_aprendizagem(celulas[indices["aprendizagem"]], aula.get("aprendizagem", ""))
    # Col 3: Metodologia — título bold, texto normal, Arial 10
    _preencher_celula_metodologia(celulas[indices["desenvolvimento"]], _metodologia_para_docx(aula))
    # Col 4: Acompanhamento — Arial 10
    _preencher_celula_lista(celulas[indices["acompanhamento"]], aula.get("acompanhamento"))
    # Col 5: Acessibilidade — Arial 10
    _preencher_celula_lista(celulas[indices["acessibilidade"]], aula.get("acessibilidade"))


def _preencher_tabelas_modelo(
    documento,
    aulas,
    escola: str,
    professor: str,
    disciplina: str,
    turma: str,
    mes: str,
    bimestre: str,
    semana: str,
    observacao: str,
    aulas_previstas_manual: str,
) -> bool:
    from core.disciplinas import eh_cdp_contextual
    is_cdp_ctx = eh_cdp_contextual(disciplina)

    pares = []
    tabelas = list(documento.tables)
    for indice, tabela in enumerate(tabelas):
        if not _eh_cabecalho_plano(tabela):
            continue
        proxima = tabelas[indice + 1] if indice + 1 < len(tabelas) else None
        if proxima is not None and _eh_tabela_aulas(proxima):
            pares.append((tabela, proxima))

    if not pares:
        return False

    aulas = list(aulas or [])
    aulas_por_par = [[] for _ in pares]
    usadas = set()

    for aula_indice, aula in enumerate(aulas):
        for par_indice, (cabecalho, _) in enumerate(pares):
            if _aula_pertence_ao_intervalo(aula, _intervalo_cabecalho(cabecalho)):
                aulas_por_par[par_indice].append((aula_indice + 1, aula))
                usadas.add(aula_indice)
                break

    sobras = [(indice + 1, aula) for indice, aula in enumerate(aulas) if indice not in usadas]
    grupos_sobra_por_semana, sobras = _agrupar_sobras_por_semana(sobras)
    for grupo in grupos_sobra_por_semana:
        par_livre = next(
            (indice for indice, aulas_do_par in enumerate(aulas_por_par) if not aulas_do_par),
            None,
        )
        if par_livre is None:
            novo_par = _clonar_par_semana(pares)
            pares.append(novo_par)
            aulas_por_par.append(grupo)
        else:
            aulas_por_par[par_livre] = grupo

    for par_indice, (_, tabela_aulas) in enumerate(pares):
        vagas = max(0, len(tabela_aulas.rows) - 1 - len(aulas_por_par[par_indice]))
        if vagas and sobras:
            aulas_por_par[par_indice].extend(sobras[:vagas])
            sobras = sobras[vagas:]

    ultimo_par_com_aula = None
    for par_indice, aulas_do_par in enumerate(aulas_por_par):
        if aulas_do_par:
            ultimo_par_com_aula = par_indice

    if ultimo_par_com_aula is not None and ultimo_par_com_aula < len(pares) - 1:
        for cabecalho, tabela_aulas in pares[ultimo_par_com_aula + 1 :]:
            _remover_tabela(cabecalho)
            _remover_tabela(tabela_aulas)
        pares = pares[: ultimo_par_com_aula + 1]
        aulas_por_par = aulas_por_par[: ultimo_par_com_aula + 1]

    for par_indice, (cabecalho, tabela_aulas) in enumerate(pares):
        if not is_cdp_ctx:
            _normalizar_layout_tabela_aulas(tabela_aulas)
        linhas_conteudo = list(tabela_aulas.rows[1:])
        aulas_da_semana = aulas_por_par[par_indice][: len(linhas_conteudo)]
        quantidade_semana = _quantidade_aulas_semana(aulas_da_semana)
        if quantidade_semana > 0:
            aulas_previstas = str(quantidade_semana)
        elif aulas_da_semana:
            aulas_previstas = str(len([a for a in aulas_da_semana if a])).strip()
        else:
            aulas_previstas = "0"
        semana_cabecalho = (
            _semana_automatica_por_aulas(aulas_da_semana)
            or _semana_atual_cabecalho(cabecalho)
            or semana
        )

        _preencher_cabecalho(
            cabecalho,
            escola,
            professor,
            disciplina,
            turma,
            mes,
            bimestre,
            semana_cabecalho,
            observacao,
            aulas_previstas,
        )

        for linha in linhas_conteudo:
            _limpar_linha(linha)

        cabecalho_aulas = tabela_aulas.rows[0] if tabela_aulas.rows else None
        for linha, (numero, aula) in zip(linhas_conteudo, aulas_da_semana):
            _preencher_linha_aula(linha, aula, numero, cabecalho_aulas)

        for linha in linhas_conteudo[len(aulas_da_semana) :]:
            _remover_linha(linha)
        if not is_cdp_ctx:
            _normalizar_layout_tabela_aulas(tabela_aulas)

    return True


def preencher_documento(
    modelo_stream,
    aulas,
    escola: str = "",
    professor: str = "",
    disciplina: str = "",
    turma: str = "",
    mes: str = "",
    bimestre: str = "",
    semana: str = "",
    observacao: str = "",
    aulas_previstas_manual: str = "",
):
    logger.info("Iniciando preenchimento de documento Word para o professor %s (disciplina: %s, turma: %s, total de aulas: %d)", professor, disciplina, turma, len(aulas or []))
    documento = Document(modelo_stream)
    primeira_aula = (aulas or [{}])[0]
    substituicoes = {
        "{{PROFESSOR}}": professor or "",
        "{{DISCIPLINA}}": disciplina or "",
        "{{TURMA}}": turma or "",
        "{{MES}}": mes or "",
        "{{BIMESTRE}}": bimestre or "",
        "{{OBSERVACAO}}": observacao or "",
        "{{AULAS_PREVISTAS}}": str(aulas_previstas_manual or ""),
        "{{TEMA}}": primeira_aula.get("tema", ""),
        "{{DATA}}": primeira_aula.get("data", ""),
        "{{HORARIO}}": primeira_aula.get("horario", ""),
        "{{APRENDIZAGEM}}": primeira_aula.get("aprendizagem", ""),
        "{{METODOLOGIA}}": _texto_metodologia(primeira_aula),
        "{{ACOMPANHAMENTO}}": texto_lista(primeira_aula.get("acompanhamento")),
        "{{ACESSIBILIDADE}}": texto_lista(primeira_aula.get("acessibilidade")),
    }
    for paragrafo in documento.paragraphs:
        _substituir_texto(paragrafo, substituicoes)
    for tabela in documento.tables:
        _substituir_em_tabela(tabela, substituicoes)

    _preencher_tabelas_modelo(
        documento,
        aulas,
        escola,
        professor,
        disciplina,
        turma,
        mes,
        bimestre,
        semana,
        observacao,
        aulas_previstas_manual,
    )
    _remover_paragrafos_vazios_finais(documento)

    saida = BytesIO()
    documento.save(saida)
    logger.info("Documento Word preenchido com sucesso para %s (%s, %s)", professor, disciplina, turma)
    return _validar_docx_gerado(saida)
