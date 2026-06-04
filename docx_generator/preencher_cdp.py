from io import BytesIO
from datetime import datetime
from typing import Dict

from docx import Document

from core.cdp import (
    componente_da_linha_multisseriada,
    disciplina_da_linha,
    habilidade_item_cdp,
    montar_acessibilidade_cdp,
    montar_acompanhamento_cdp,
    montar_metodologia_cdp,
    selecionar_item,
    titulo_item_cdp,
)
from core.ia import processar_item_cdp_ia
from docx_generator.preencher import (
    _eh_cabecalho_plano,
    _preencher_cabecalho,
    _preencher_celula_aprendizagem,
    _preencher_celula_data_horario,
    _preencher_celula_lista,
    _preencher_celula_metodologia,
    _preencher_celula_tema_material,
)


def _linha_de_aula_cdp(row) -> bool:
    if len(row.cells) < 5:
        return False
    primeira = (row.cells[0].text or "").strip().lower()
    material = (row.cells[1].text or "").strip().lower()
    if not primeira or not material:
        return False
    if "aula" in primeira and "semanal" in primeira:
        return False
    if "número e título" in material or "numero e titulo" in material:
        return False
    return bool(disciplina_da_linha(material))


def _first_distinct_cell_index(row, start_idx: int):
    if start_idx >= len(row.cells):
        return None
    vistos = {row.cells[i]._tc for i in range(start_idx + 1)}
    for idx in range(start_idx + 1, len(row.cells)):
        if row.cells[idx]._tc not in vistos:
            return idx
    return None


def _indices_cdp(row) -> Dict[str, int | None]:
    desenvolvimento_idx = 3
    acompanhamento_idx = _first_distinct_cell_index(row, desenvolvimento_idx)
    acessibilidade_idx = None
    if acompanhamento_idx is not None:
        vistos = {row.cells[i]._tc for i in range(acompanhamento_idx + 1)}
        for idx in range(len(row.cells) - 1, acompanhamento_idx, -1):
            if row.cells[idx]._tc not in vistos:
                acessibilidade_idx = idx
                break
    return {
        "material": 1,
        "aprendizagem": 2,
        "desenvolvimento": desenvolvimento_idx,
        "acompanhamento": acompanhamento_idx,
        "acessibilidade": acessibilidade_idx,
    }


def _disciplina_exibicao(disciplina: str) -> str:
    nomes = {
        "português": "PORTUGUÊS",
        "matematica": "MATEMÁTICA",
        "história": "HISTÓRIA",
        "geografia": "GEOGRAFIA",
        "ciências": "CIÊNCIAS",
        "arte": "ARTE",
    }
    return nomes.get(disciplina, disciplina.upper())


_DIAS_SEMANA_CDP = {
    0: "Segunda",
    1: "Terça",
    2: "Quarta",
    3: "Quinta",
    4: "Sexta",
    5: "Sábado",
    6: "Domingo",
}


def _dia_semana_cdp(valor) -> str:
    if hasattr(valor, "weekday"):
        return _DIAS_SEMANA_CDP.get(int(valor.weekday()), "")
    texto = str(valor or "").strip()
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d/%m"):
        try:
            data = datetime.strptime(texto, fmt)
            return _DIAS_SEMANA_CDP.get(int(data.weekday()), "")
        except ValueError:
            continue
    return ""


def _ajustar_rotulos_data_cdp(doc) -> None:
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    texto = paragrafo.text or ""
                    if "Data e Horário" in texto:
                        paragrafo.text = texto.replace("Data e Horário", "Data")
                    elif "Data e Horario" in texto:
                        paragrafo.text = texto.replace("Data e Horario", "Data")


def _formatar_data_aula_cdp(aula: dict) -> str:
    data_bruta = aula.get("data")
    if hasattr(data_bruta, "strftime"):
        data = data_bruta.strftime("%d/%m")
    else:
        data = str(data_bruta or "").strip()
    dia_semana = str(aula.get("dia_semana") or "").strip() or _dia_semana_cdp(data_bruta)
    partes = [parte for parte in [data, dia_semana] if parte]
    return "\n".join(partes).strip()


def _habilidade(item: Dict[str, str]) -> str:
    habilidade = habilidade_item_cdp(item)
    return f"HABILIDADE:\n{habilidade}" if habilidade else ""


def _material(disciplina: str, item: Dict[str, str]) -> str:
    titulo = titulo_item_cdp(item)
    return f"TEMA:\n{titulo}"


def _validar_docx_gerado(buffer: BytesIO) -> BytesIO:
    conteudo = buffer.getvalue()
    Document(BytesIO(conteudo))
    buffer.seek(0)
    return buffer


def _metodologia_dict(texto: str):
    return [{"titulo": "", "texto": texto}]


def prever_aulas_cdp(
    modelo_docx,
    aula_inicial: int = 1,
    fundamental: bool = False,
    multisseriada: bool = False,
    serie_cdp: str = "",
    componente_cdp: str = "",
    bimestre: str = "",
) -> list[Dict[str, str]]:
    doc = Document(modelo_docx)
    contadores: Dict[str, int] = {}
    preview: list[Dict[str, str]] = []

    for table in doc.tables:
        for row in table.rows:
            if not _linha_de_aula_cdp(row):
                continue

            idxs = _indices_cdp(row)
            material_modelo = (row.cells[idxs["material"]].text or "").strip()
            disciplina = disciplina_da_linha(material_modelo)
            turma_selecao = serie_cdp or ""
            contador = contadores.get(disciplina, 0)
            item = selecionar_item(
                disciplina,
                contador,
                turma=turma_selecao,
                bimestre=bimestre,
                aula_inicial=aula_inicial,
                fundamental=fundamental,
                multisseriada=multisseriada,
                componente_cdp=componente_cdp,
            )
            contadores[disciplina] = contadores.get(disciplina, 0) + 1

            aula_planilha = str(item.get("AULA", "") or "").strip()
            titulo = titulo_item_cdp(item)
            habilidade = habilidade_item_cdp(item)
            preview.append(
                {
                    "ordem": str(len(preview) + 1),
                    "disciplina": _disciplina_exibicao(disciplina),
                    "componente_planilha": componente_da_linha_multisseriada(material_modelo) if multisseriada else _disciplina_exibicao(disciplina),
                    "material_modelo": material_modelo,
                    "aula_planilha": aula_planilha,
                    "titulo": titulo,
                    "habilidade": habilidade,
                }
            )

    return preview


def _preencher_cabecalhos_cdp(
    doc,
    escola: str,
    professor: str,
    componente: str,
    turma: str,
    mes: str,
    bimestre: str,
    semana: str,
    observacao: str,
    aulas_previstas_manual: str,
) -> None:
    for table in doc.tables:
        if _eh_cabecalho_plano(table):
            _preencher_cabecalho(
                table,
                escola,
                professor,
                componente,
                turma,
                mes,
                bimestre,
                semana,
                observacao,
                aulas_previstas_manual,
            )


def preencher_documento_cdp(
    modelo_docx,
    escola: str,
    professor: str,
    turma: str,
    mes: str = "",
    bimestre: str = "",
    aula_inicial: int = 1,
    fundamental: bool = False,
    multisseriada: bool = False,
    serie_cdp: str = "",
    componente_cdp: str = "",
    item_cdp: Dict[str, str] | None = None,
    usar_ia: bool = False,
    provedor_ia: str = "",
    modelo_ia: str = "",
    datas_horarios: list[dict] | None = None,
    semana: str = "",
    observacao: str = "",
    aulas_previstas_manual: str = "",
) -> BytesIO:
    doc = Document(modelo_docx)
    _ajustar_rotulos_data_cdp(doc)
    componente = "CDP - CICLO I" if fundamental else "MULTISSERIADA - EJA FUNDAMENTAL - ANOS INICIAIS"
    _preencher_cabecalhos_cdp(
        doc,
        escola,
        professor,
        componente,
        turma,
        mes,
        bimestre,
        semana,
        observacao,
        aulas_previstas_manual,
    )

    contadores: Dict[str, int] = {}
    indice_data_horario = 0
    for table in doc.tables:
        for row in table.rows:
            if not _linha_de_aula_cdp(row):
                continue

            idxs = _indices_cdp(row)
            disciplina = disciplina_da_linha(row.cells[idxs["material"]].text)
            turma_selecao = serie_cdp or turma

            if item_cdp and contadores.get(disciplina, 0) == 0:
                item = item_cdp
            else:
                contador = contadores.get(disciplina, 0)
                item = selecionar_item(
                    disciplina,
                    contador,
                    turma=turma_selecao,
                    bimestre=bimestre,
                    aula_inicial=aula_inicial,
                    fundamental=fundamental,
                    multisseriada=multisseriada,
                    componente_cdp=componente_cdp,
                )

            contadores[disciplina] = contadores.get(disciplina, 0) + 1
            if not item:
                continue

            if datas_horarios and indice_data_horario < len(datas_horarios):
                _preencher_celula_data_horario(
                    row.cells[0],
                    _formatar_data_aula_cdp(datas_horarios[indice_data_horario]),
                )
            indice_data_horario += 1

            aprendizagem = _habilidade(item)
            metodologia = montar_metodologia_cdp(disciplina, item, fundamental=fundamental)
            acompanhamento = montar_acompanhamento_cdp(disciplina, item, fundamental=fundamental)
            acessibilidade = montar_acessibilidade_cdp(disciplina, item, fundamental=fundamental)

            if usar_ia and provedor_ia:
                try:
                    plano_ia = processar_item_cdp_ia(item, disciplina, turma_selecao, provedor_ia, modelo_ia)
                    if plano_ia.get("aprendizagem"):
                        aprendizagem = f"HABILIDADE:\n{plano_ia['aprendizagem']}"
                    if plano_ia.get("metodologia"):
                        metodologia = "\n\n".join(
                            f"{etapa.get('titulo', '').strip()}: {etapa.get('texto', '').strip()}".strip(": ")
                            for etapa in plano_ia["metodologia"]
                            if etapa.get("texto")
                        ) or metodologia
                except Exception:
                    pass

            _preencher_celula_aprendizagem(row.cells[idxs["aprendizagem"]], aprendizagem)
            _preencher_celula_metodologia(
                row.cells[idxs["desenvolvimento"]],
                _metodologia_dict(metodologia),
            )
            _preencher_celula_tema_material(row.cells[idxs["material"]], _material(disciplina, item))

            if idxs["acompanhamento"] is not None:
                _preencher_celula_lista(
                    row.cells[idxs["acompanhamento"]],
                    acompanhamento,
                )
            if idxs["acessibilidade"] is not None:
                _preencher_celula_lista(
                    row.cells[idxs["acessibilidade"]],
                    acessibilidade,
                )

    out = BytesIO()
    doc.save(out)
    return _validar_docx_gerado(out)
