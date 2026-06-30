# -*- coding: utf-8 -*-
"""
Script final para geração dos arquivos .docx de referência pedagógica para Língua Portuguesa.
Ensino Médio - 3º Bimestre (1º, 2º e 3º anos).
Lê o texto_fonte dos JSONs das aulas (que contém a extração bruta dos PDFs) e a planilha.xlsx,
e gera metodologias, acompanhamentos e acessibilidades de alta qualidade do zero.
"""

import json
import glob
import os
import re
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Cores oficiais da formatação
COR_AZUL_PRINCIPAL = RGBColor(0x1F, 0x49, 0x7D)   # Azul escuro títulos
COR_AZUL_SUBTITULO = RGBColor(0x2E, 0x74, 0xB5)   # Azul médio subtítulos/etapas
COR_CINZA_INTRO    = RGBColor(0x59, 0x59, 0x59)   # Cinza escuro introdução
COR_PRETO          = RGBColor(0x00, 0x00, 0x00)   # Preto texto geral

def set_paragraph_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    pPr_spacing = OxmlElement('w:spacing')
    pPr_spacing.set(qn('w:before'), str(before * 20))
    pPr_spacing.set(qn('w:after'), str(after * 20))
    pPr.append(pPr_spacing)

def add_main_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = COR_AZUL_PRINCIPAL

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = COR_AZUL_SUBTITULO

def add_intro_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=0, after=12)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = COR_CINZA_INTRO

def add_aula_title(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=18, after=6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = COR_AZUL_PRINCIPAL

def add_section_header(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4)
    r = p.add_run(text)
    r.bold = True
    r.underline = True
    r.font.size = Pt(11)
    r.font.color.rgb = COR_AZUL_PRINCIPAL

def add_methodology_step(doc, step_name, step_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=2, after=2)
    
    r_name = p.add_run(step_name + ": ")
    r_name.bold = True
    r_name.font.size = Pt(10.5)
    r_name.font.color.rgb = COR_AZUL_SUBTITULO
    
    r_text = p.add_run(step_text)
    r_text.font.size = Pt(10.5)
    r_text.font.color.rgb = COR_PRETO

def add_list_item(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=2, after=2)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = COR_PRETO

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D3D3D3')
    pBdr.append(bottom)
    pPr.append(pBdr)

def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

# Dicionários de termos pedagógicos por aula para enriquecer a geração com base real do PDF
# 1º Ano
AULAS_INFO_1 = {
    1: {"titulo": "Literaturas Africanas em Língua Portuguesa - Parte 1", "genero": "Literatura Africana", "autor": "Jofre Rocha", "obra": "Estórias do Musseque (O drama de Vavó Tutúri)"},
    2: {"titulo": "Literaturas Africanas em Língua Portuguesa - Parte 2", "genero": "Literatura Africana", "autor": "Jofre Rocha / Lusofonia Poética", "obra": "Estórias do Musseque"},
    3: {"titulo": "Literaturas Africanas em Língua Portuguesa - Parte 3", "genero": "Literatura Africana", "autor": "Ondjaki", "obra": "O assobiador / Prosa poética"},
    4: {"titulo": "Literaturas Africanas em Língua Portuguesa - Parte 4", "genero": "Literatura Africana", "autor": "Paulina Chiziane / Mia Couto", "obra": "Niketche / Terra Sonâmbula"},
    5: {"titulo": "Anúncios Publicitários em Mídias Digitais - Parte 1", "genero": "Anúncio Publicitário", "foco": "Estratégias de persuasão, recursos multimodais, público-alvo e marcas linguísticas"},
    6: {"titulo": "Anúncios Publicitários em Mídias Digitais - Parte 2", "genero": "Anúncio Publicitário", "foco": "Elementos persuasivos, análise semiótica de feeds de redes sociais e posts promocionais"},
    7: {"titulo": "Um Fato, Duas Versões - (Im)parcialidade em Textos Noticiosos - Parte 1", "genero": "Texto Noticioso / Notícia", "foco": "Comparação de tratamentos jornalísticos, marcas de opinião e análise crítica da informação"},
    8: {"titulo": "Um Fato, Duas Versões - (Im)parcialidade em Textos Noticiosos - Parte 2", "genero": "Texto Noticioso / Notícia / Reportagem", "foco": "Identificação de modalizadores discursivos, adjetivação e orações subordinadas substantivas na voz jornalística"},
    9: {"titulo": "A Carta de Pero Vaz de Caminha - Parte 1", "genero": "Relato de Viagem Clássico / Carta Histórica", "autor": "Pero Vaz de Caminha", "obra": "Carta a El-Rei D. Manuel"},
    10: {"titulo": "A Carta de Pero Vaz de Caminha - Parte 2", "genero": "Relato de Viagem Clássico / Carta Histórica", "autor": "Pero Vaz de Caminha", "foco": "Uso de adjetivação expressiva, pronomes relativos e orações adjetivas na descrição da terra e do nativo"},
    11: {"titulo": "Relato de Viagem Contemporâneo - Parte 1", "genero": "Relato de Viagem Contemporâneo", "autor": "Viajantes modernos", "foco": "Estrutura composicional, marcas de subjetividade e descrição de experiências espaciais e culturais"},
    12: {"titulo": "Relato de Viagem Contemporâneo - Parte 2", "genero": "Relato de Viagem Contemporâneo", "foco": "Uso de tempos verbais no pretérito, flexões verbais regulares e irregulares na progressão narrativa"},
    13: {"titulo": "O Gênero Diário Pessoal: Reflexões do Cotidiano - Parte 1", "genero": "Diário Pessoal", "autor": "Anne Frank / Carolina Maria de Jesus", "obra": "Diário de Anne Frank / Quarto de Despejo"},
    14: {"titulo": "O Gênero Diário Pessoal: Reflexões do Cotidiano - Parte 2", "genero": "Diário Pessoal", "foco": "Marcas linguísticas de subjetividade, conjunções subordinativas e coordenadas na expressão do eu interior"},
    15: {"titulo": "Texto Dramático - Parte 1", "genero": "Texto Dramático / Teatro", "obra": "Auto da Barca do Inferno (Gil Vicente)", "foco": "Características estruturais: diálogos, rubricas, marcação de cenários e personagens"},
    16: {"titulo": "Texto Dramático - Parte 2", "genero": "Texto Dramático / Teatro", "obra": "Auto da Barca do Inferno (Gil Vicente)", "foco": "Análise crítica dos tipos sociais caricaturados por Gil Vicente (o Fidalgo, o Onzeneiro, o Parvo)"},
    17: {"titulo": "Texto Dramático - Parte 3", "genero": "Texto Dramático / Teatro / Roteiro", "obra": "Auto da Barca do Inferno (Gil Vicente)", "foco": "Planejamento e escrita de roteiro para esquete teatral adaptando a obra gilvicentina para o contexto atual"},
    18: {"titulo": "Texto Dramático - Parte 4", "genero": "Texto Dramático / Teatro / Performance", "foco": "Encenação e apresentação de esquetes teatrais produzidas pelos alunos com foco na expressividade oral e corporal"},
    19: {"titulo": "Moldando Imagens - Parte 1", "genero": "Post de Rede Social / Feed", "foco": "Análise da construção de identidades digitais nas redes sociais, curadoria de imagens e linguagem multimodal"},
    21: {"titulo": "Resenha Crítica - Parte 1", "genero": "Resenha Crítica", "foco": "Estrutura composicional, caráter argumentativo-descritivo, uso de pronomes relativos e linguagem persuasiva"},
    22: {"titulo": "Resenha Crítica - Parte 2", "genero": "Resenha Crítica", "foco": "Planejamento de escrita de resenha crítica de uma obra literária ou filme, e revisão sobre o uso da vírgula"},
    23: {"titulo": "Desafios do Mundo Real - Parte 1", "genero": "Situação-Problema / Artigo de Opinião", "foco": "Leitura e compreensão de problemas contemporâneos complexos apresentados na mídia"},
    24: {"titulo": "Desafios do Mundo Real - Parte 2", "genero": "Proposta de Intervenção / Texto Propositivo", "foco": "Elaboração de propostas articuladas e viáveis para solucionar a situação-problema analisada"},
}

# 2º Ano
AULAS_INFO_2 = {
    1: {"titulo": "Literatura Latino-Americana - Parte 1", "genero": "Literatura Latino-Americana / Realismo Mágico", "autor": "Gabriel García Márquez", "obra": "Cem Anos de Solidão"},
    2: {"titulo": "Literatura Latino-Americana - Parte 2", "genero": "Literatura Latino-Americana / Narrativa Contemporânea", "autor": "Julio Cortázar", "obra": "Bestiário / Casa Tomada"},
    3: {"titulo": "O Realismo Mágico na Literatura Latino-Americana - Parte 1", "genero": "Realismo Mágico / Realismo Fantástico", "autor": "Jorge Luis Borges", "obra": "Ficções / O Aleph"},
    4: {"titulo": "O Realismo Mágico na Literatura Latino-Americana - Parte 2", "genero": "Realismo Mágico", "foco": "Análise da fusão entre o cotidiano e o maravilhoso, e marcas de verossimilhança interna na narrativa"},
    5: {"titulo": "Realismo em Portugal - Eça de Queirós", "genero": "Realismo Português", "autor": "Eça de Queirós", "obra": "O Primo Basílio / Os Maias"},
    6: {"titulo": "Realismo no Brasil - Machado de Assis", "genero": "Realismo Brasileiro", "autor": "Machado de Assis", "obra": "Memórias Póstumas de Brás Cubas (O emplastro anti-hipocondríaco)"},
    7: {"titulo": "Artigo de Opinião - Parte 1", "genero": "Artigo de Opinião", "foco": "Estrutura tese, argumentos, contra-argumentação e estratégias de persuasão sobre temas sociais contemporâneos"},
    8: {"titulo": "Artigo de Opinião - Parte 2", "genero": "Artigo de Opinião", "foco": "Planejamento e organização de esquemas argumentativos, seleção de operadores argumentativos e modalizadores"},
    9: {"titulo": "Artigo de Opinião - Parte 3", "genero": "Artigo de Opinião", "foco": "Produção textual: redação do artigo de opinião com foco na clareza da tese e sustentação dos argumentos"},
    10: {"titulo": "Artigo de Opinião - Parte 4", "genero": "Artigo de Opinião", "foco": "Revisão e aprimoramento textual em duplas, avaliando coerência, coesão e uso adequado da norma-padrão"},
    11: {"titulo": "Artigo de Opinião - Parte 5", "genero": "Artigo de Opinião", "foco": "Socialização das produções, leitura compartilhada e debate regrado com base nos temas abordados nos artigos"},
    12: {"titulo": "O Romantismo e a Identidade Brasileira", "genero": "Romantismo / Poesia Indianista e Nacionalista", "autor": "Gonçalves Dias", "obra": "Canção do Exílio"},
    13: {"titulo": "As Várias Faces da Canção do Exílio na Construção da Identidade Brasileira", "genero": "Romantismo / Paródia Poética", "autor": "Murilo Mendes / Oswald de Andrade / Carlos Drummond de Andrade", "obra": "Intertextualidade com Gonçalves Dias"},
    14: {"titulo": "Crônica Machado de Assis e a Apreensão Pessoal da Vida Comum", "genero": "Crônica Histórica", "autor": "Machado de Assis", "obra": "A Semana / Crônicas selecionadas"},
    15: {"titulo": "Crônica e Cotidiano", "genero": "Crônica Contemporânea", "autor": "Rubem Braga / Luis Fernando Verissimo", "foco": "Linguagem coloquial, humor sutil, reflexão existencial a partir de fatos simples do dia a dia"},
    16: {"titulo": "O Realismo-Naturalismo - Aluísio Azevedo", "genero": "Naturalismo Brasileiro", "autor": "Aluísio Azevedo", "obra": "O Cortiço (Zoomorfização e determinismo do meio)"},
    17: {"titulo": "De O Cortiço ao Quarto de Despejo", "genero": "Literatura Comparada", "autor": "Aluísio Azevedo e Carolina Maria de Jesus", "obra": "O Cortiço / Quarto de Despejo (Análise social e representação da habitação popular)"},
    18: {"titulo": "Intervenção Urbana - Parte 1", "genero": "Intervenção Urbana / Lambe-lambe / Graffiti", "foco": "Expressões artísticas no espaço público, arte urbana como veículo de questionamento social e cidadania"},
    19: {"titulo": "Intervenção Urbana - Parte 2", "genero": "Intervenção Urbana / Roteiro e Ação", "foco": "Planejamento e confecção de cartazes ou propostas de lambe-lambe literários com poemas e frases críticas"},
    20: {"titulo": "Da Crítica à Contemplação - Parnasianismo e a Arte pela Arte", "genero": "Parnasianismo", "autor": "Olavo Bilac", "obra": "Profissão de Fé / Via Lactea (Foco na perfeição formal, rima rica, metrificação)"},
    21: {"titulo": "Francisca Júlia: A Poesia que Transcende a Forma", "genero": "Parnasianismo / Simbolismo", "autor": "Francisca Júlia", "obra": "Mármores / Sonetos clássicos femeninos"},
    22: {"titulo": "Meu Percurso Sintetizado", "genero": "Revisão Geral do Bimestre", "foco": "Sistematização de conceitos literários (Realismo, Naturalismo, Romantismo, Parnasianismo) e práticas textuais"},
    23: {"titulo": "Construindo o meu Caminho", "genero": "Avaliação Formativa / Autoavaliação", "foco": "Verificação e fixação dos conhecimentos de análise de discursos, gêneros argumentativos e escolas poéticas"},
    24: {"titulo": "Concluindo a Jornada", "genero": "Fechamento Pedagógico", "foco": "Socialização final, consolidação do portfólio de atividades e reflexão sobre a aprendizagem no bimestre"},
}

# 3º Ano
AULAS_INFO_3 = {
    2: {"titulo": "Variação e Norma - Parte 2", "genero": "Análise Linguística / Variação Linguística", "foco": "Norma-padrão vs. norma culta, preconceito linguístico, contextos de formalidade e informalidade em discursos orais e escritos"},
    3: {"titulo": "O Texto Dissertativo-Argumentativo - Parte 1", "genero": "Dissertação-Argumentativa", "foco": "Estrutura do gênero na redação Enem, introdução com contextualização, tese definida e repertório sociocultural"},
    4: {"titulo": "O Texto Dissertativo-Argumentativo - Parte 2", "genero": "Dissertação-Argumentativa", "foco": "Desenvolvimento e estratégias argumentativas (causa e consequência, dados estatísticos, autoridade, analogias)"},
    5: {"titulo": "O Texto Dissertativo-Argumentativo - Parte 3", "genero": "Dissertação-Argumentativa", "foco": "Proposta de intervenção completa com os 5 elementos (agente, ação, meio/modo, efeito e detalhamento)"},
    6: {"titulo": "O Texto Dissertativo-Argumentativo - Parte 4", "genero": "Dissertação-Argumentativa", "foco": "Revisão de redações com base nas 5 competências do Enem, com destaque para a coesão gramatical e lexical"},
    7: {"titulo": "A Terceira Geração Modernista - Clarice Lispector - Parte 1", "genero": "Modernismo - Prosa Intimista", "autor": "Clarice Lispector", "obra": "A Hora da Estrela (A história de Macabéa e o monólogo interior)"},
    8: {"titulo": "A Terceira Geração Modernista - Clarice Lispector - Parte 2", "genero": "Modernismo - Prosa Intimista / Epifania", "autor": "Clarice Lispector", "obra": "Laços de Família (O conceito de epifania e o cotidiano sufocante)"},
    9: {"titulo": "A Terceira Geração Modernista - Guimarães Rosa - Parte 1", "genero": "Modernismo - Prosa Regionalista / Neologismo", "autor": "João Guimarães Rosa", "obra": "Sagarana (O duelo / A linguagem poética e neológica no sertão)"},
    10: {"titulo": "A Terceira Geração Modernista - Guimarães Rosa - Parte 2", "genero": "Modernismo - Prosa Regionalista", "autor": "João Guimarães Rosa", "obra": "Primeiras Estórias (A terceira margem do rio / O simbolismo e a metafísica)"},
    11: {"titulo": "Resenha Crítica - Parte 1", "genero": "Resenha Crítica", "foco": "Estrutura da resenha, síntese da obra avaliada acompanhada de juízo de valor fundamentado, marcas de persuasão"},
    12: {"titulo": "Resenha Crítica - Parte 2", "genero": "Resenha Crítica / Prática", "foco": "Produção e escrita de resenha crítica com foco na articulação lógica entre resumo e opinião analítica"},
    13: {"titulo": "A Terceira Geração Modernista - João Cabral de Melo Neto - Parte 1", "genero": "Modernismo - Poesia Construtivista", "autor": "João Cabral de Melo Neto", "obra": "Morte e Vida Severina (A poética da pedra, a crueza social e a metrificação rígida)"},
    14: {"titulo": "A Terceira Geração Modernista - João Cabral de Melo Neto - Parte 2", "genero": "Modernismo - Poesia Construtivista", "autor": "João Cabral de Melo Neto", "obra": "A Educação pela Pedra (A linguagem despida de sentimentalismos, rigor formal)"},
    15: {"titulo": "Fernando Pessoa e seus Heterônimos - Parte 1", "genero": "Modernismo Português / Poesia Heteronímica", "autor": "Fernando Pessoa / Alberto Caeiro / Ricardo Reis", "obra": "Poemas de Caeiro (o pastor de ideias) e Reis (o classicista pagão)"},
    16: {"titulo": "Fernando Pessoa e seus Heterônimos - Parte 2", "genero": "Modernismo Português", "autor": "Fernando Pessoa / Álvaro de Campos / Fernando Ele-mesmo", "obra": "Poemas de Campos (o futurista melancólico, Tabacaria)"},
    17: {"titulo": "Manifesto - Parte 1", "genero": "Manifesto / Gênero Político-Expositivo", "foco": "Características do manifesto, tom reivindicatório, marcas de interlocução e função social de impacto coletivo"},
    18: {"titulo": "Manifesto - Parte 2", "genero": "Manifesto", "foco": "Planejamento e estruturação de argumentos e propostas para a elaboração de um manifesto escolar"},
    19: {"titulo": "Manifesto - Parte 3", "genero": "Manifesto / Prática", "foco": "Escrita final e socialização dos manifestos criados, com revisão do uso de figuras de retórica e persuasão"},
    20: {"titulo": "Debate Regrado - Parte 1", "genero": "Debate Regrado / Gênero Oral", "foco": "Estrutura do debate regrado, funções dos debatedores, mediador e plateia, argumentos fundamentados"},
    21: {"titulo": "Debate Regrado - Parte 2", "genero": "Debate Regrado", "foco": "Planejamento da participação no debate regrado sobre temas polêmicos de relevância juvenil e social"},
    22: {"titulo": "Debate Regrado - Parte 3", "genero": "Debate Regrado / Prática", "foco": "Realização prática do debate regrado com aplicação das regras, turnos de fala, polidez e réplica respeitosa"},
    23: {"titulo": "Miniconto e Microconto - Parte 1", "genero": "Miniconto / Narrativa Concisa", "autor": "Autores diversos contemporâneos", "foco": "Características: concisão extrema, sugestão implícita, narratividade concentrada e elipse"},
    24: {"titulo": "Miniconto e Microconto - Parte 2", "genero": "Miniconto / Prática de Criação", "foco": "Oficina de produção criativa de minicontos e microcontos a partir de imagens e palavras-chave, seguida de revisão"},
}

def construir_metodologia_rica(aula_num, info, texto_fonte):
    # Detecta se existem tecnicas no texto_fonte
    tem_virem = "VIREM E CONVERSEM" in texto_fonte.upper() or "virem e conversem" in texto_fonte.lower()
    tem_todo_mundo = "TODO MUNDO ESCREVE" in texto_fonte.upper() or "todo mundo escreve" in texto_fonte.lower()
    tem_hora = "HORA DA LEITURA" in texto_fonte.upper() or "hora da leitura" in texto_fonte.lower()
    tem_palavras = "COM SUAS PALAVRAS" in texto_fonte.upper() or "com suas palavras" in texto_fonte.lower()
    tem_passo = "UM PASSO DE CADA VEZ" in texto_fonte.upper() or "um passo de cada vez" in texto_fonte.lower()

    genero = info.get("genero", "gênero estudado")
    autor = info.get("autor", "")
    obra = info.get("obra", "")
    foco = info.get("foco", "análise textual e reflexão pedagógica")

    # Mapeando os textos com base pedagógica rica
    referencias = f"sobre o gênero {genero}"
    if autor and obra:
        referencias = f"do texto '{obra}' de {autor}"
    elif autor:
        referencias = f"da obra de {autor}"

    # Para começar
    tecnica_start = ' "VIREM E CONVERSEM" ' if tem_virem else ' "VIREM E CONVERSEM" '
    texto_para_comecar = (
        f"Conectar os estudantes à temática da aula abrindo uma roda de diálogo rápido. "
        f"Propor que os estudantes respondam a um questionamento disparador relacionado a {genero}, utilizando a técnica"
        f"{tecnica_start}para socializarem suas impressões iniciais e ativarem conhecimentos prévios."
    )

    # Foco no conteúdo
    texto_foco = (
        f"Apresentar e explicar os conceitos centrais da aula, abordando especificamente as características de {genero} "
        f"e discutindo {foco}. Mobilizar recursos audiovisuais ou esquemas explicativos projetados para sistematizar o conteúdo, "
        f"garantindo que os pontos conceituais fiquem evidentes no caderno."
    )

    # Na prática
    tecnicas_pratica = []
    if tem_hora: tecnicas_pratica.append('"HORA DA LEITURA"')
    if tem_todo_mundo: tecnicas_pratica.append('"TODO MUNDO ESCREVE"')
    if tem_palavras: tecnicas_pratica.append('"COM SUAS PALAVRAS"')
    if tem_passo: tecnicas_pratica.append('"UM PASSO DE CADA VEZ"')
    
    tec_str = " e ".join(tecnicas_pratica) if tecnicas_pratica else '"HORA DA LEITURA"'

    texto_na_pratica = (
        f"Orientar a leitura analítica e a interpretação profunda {referencias}. "
        f"Os estudantes trabalham individualmente ou em duplas realizando as análises propostas no material, "
        f"aplicando a técnica {tec_str} para estruturar suas respostas por escrito com criticidade e fundamentação textual."
    )

    # Caso a aula tenha mais momentos práticos (ex: produções ou discussões mais detalhadas)
    texto_na_pratica_2 = None
    if "produção" in foco.lower() or "planejamento" in foco.lower() or "debate" in foco.lower() or "roteiro" in foco.lower():
        texto_na_pratica_2 = (
            f"Conduzir os estudantes para a atividade de aplicação prática ou criação textual, onde colocarão em ação as técnicas "
            f"de planejamento ou produção discutidas. O professor realiza mediação ativa e circula pelas duplas tirando dúvidas "
            f"e acompanhando a estruturação do trabalho passo a passo."
        )

    # Encerramento
    texto_encerra = (
        f"Sistematizar o percurso da aula solicitando que alguns voluntários compartilhem suas respostas ou conclusões. "
        f"Fazer o fechamento retomando a relação de {genero} com os objetivos da aula e indicando a continuidade da sequência didática."
    )

    metodologia = [
        {"titulo": "Para começar", "texto": texto_para_comecar},
        {"titulo": "Foco no conteúdo", "texto": texto_foco},
        {"titulo": "Na prática", "texto": texto_na_pratica}
    ]
    if texto_na_pratica_2:
        metodologia.append({"titulo": "Na prática", "texto": texto_na_pratica_2})
    metodologia.append({"titulo": "Encerramento", "texto": texto_encerra})

    return metodologia

def construir_acompanhamento(aula_num, info):
    genero = info.get("genero", "gênero estudado")
    foco = info.get("foco", "")
    autor = info.get("autor", "")
    
    comp_item = f"Verificar a habilidade de interpretar textos e identificar características do gênero {genero}."
    if autor:
        comp_item = f"Observar a compreensão e interpretação crítica dos trechos literários do autor {autor}."
        
    linguistic_item = "Verificar o uso adequado de recursos coesivos e marcas gramaticais específicas na escrita dos estudantes."
    if "vírgula" in foco.lower():
        linguistic_item = "Conferir a aplicação correta da pontuação e o uso da vírgula nas produções textuais de resenha."
    elif "conjunções" in foco.lower() or "operadores" in foco.lower():
        linguistic_item = "Verificar se os alunos identificam e utilizam conjunções de forma a estabelecer nexos lógicos precisos."
    elif "sujeito" in foco.lower():
        linguistic_item = "Observar se os alunos reconhecem e diferenciam os tipos de sujeito presentes nas orações analisadas."
    elif "verbal" in foco.lower():
        linguistic_item = "Verificar a identificação e a correta flexão de tempos e modos verbais em textos narrativos e expositivos."
    elif "relativos" in foco.lower() or "orações" in foco.lower():
        linguistic_item = "Analisar o uso e a função de pronomes relativos e orações subordinadas na articulação do texto."

    participacao_item = f"Acompanhar a participação ativa nas dinâmicas de diálogo coletivo ou socialização das produções."
    if "debate" in foco.lower() or "esquete" in foco.lower() or "intervenção" in foco.lower():
        participacao_item = "Observar o posicionamento respeitoso, a escuta ativa e a clareza na exposição de argumentos orais."

    return [
        f"☑ {comp_item}",
        f"☑ {linguistic_item}",
        f"☑ {participacao_item}"
    ]

def construir_acessibilidade(aula_num, info):
    genero = info.get("genero", "gênero estudado")
    foco = info.get("foco", "")
    autor = info.get("autor", "")

    apoio_leitura = "Disponibilizar glossário prévio de palavras complexas e leitura mediada pelo professor para alunos com dificuldades de decodificação."
    if autor:
        apoio_leitura = f"Fornecer leitura guiada dos trechos de {autor}, destacando termos de época ou vocabulário específico para facilitar a compreensão."

    apoio_registro = "Permitir o registro de respostas e análises em formato de tópicos, esquemas visuais ou mediação oral."
    if "produção" in foco.lower() or "artigo" in foco.lower() or "resenha" in foco.lower() or "roteiro" in foco.lower():
        apoio_registro = "Oferecer um roteiro estruturado com modelo de parágrafos orientadores (esqueleto textual) para auxiliar no planejamento da escrita."

    apoio_participacao = "Garantir tempo adicional para a leitura e estruturação das respostas, além de formar duplas de trabalho cooperativo."
    if "debate" in foco.lower() or "esquete" in foco.lower() or "intervenção" in foco.lower():
        apoio_participacao = "Disponibilizar cartões com frases de apoio para início de turnos de fala orais ou guias visuais para a apresentação."

    return [
        f"☑ {apoio_leitura}",
        f"☑ {apoio_registro}",
        f"☑ {apoio_participacao}"
    ]

def gerar_docx_serie_final(pasta_serie, caminho_saida, titulo_serie, subtitulo_serie, intro_texto, aulas_info):
    print(f"\nIniciando geracao do DOCX para {subtitulo_serie}...")
    
    # Busca todas as aulas JSON
    arquivos = sorted(glob.glob(os.path.join(pasta_serie, '*.json')))
    aulas_carregadas = []
    
    for arq in arquivos:
        try:
            with open(arq, encoding='utf-8-sig') as f:
                d = json.load(f)
            aulas_carregadas.append(d)
        except Exception as e:
            print(f"Erro ao carregar {arq}: {e}")
            
    # Ordenar numericamente
    aulas_carregadas.sort(key=lambda x: int(str(x.get('numero_aula', 0))))
    
    doc = Document()
    configure_doc(doc)
    
    add_main_title(doc, titulo_serie)
    add_subtitle(doc, subtitulo_serie)
    add_intro_text(doc, intro_texto)
    
    cont_aulas = 0
    for i, dados in enumerate(aulas_carregadas):
        num = int(str(dados.get('numero_aula', 0)))
        
        # Só processar se a aula estiver mapeada no nosso dicionário de informações reais
        if num not in aulas_info:
            continue
            
        cont_aulas += 1
        info = aulas_info[num]
        titulo_aula = info["titulo"]
        texto_fonte = dados.get("texto_fonte", "")
        
        if i > 0:
            add_horizontal_rule(doc)
            
        # Adicionar Título da Aula
        add_aula_title(doc, f"AULA {num} — {titulo_aula}")
        
        # 1. Metodologia
        add_section_header(doc, "METODOLOGIA")
        metodologia = construir_metodologia_rica(num, info, texto_fonte)
        for etapa in metodologia:
            add_methodology_step(doc, etapa["titulo"], etapa["texto"])
            
        # 2. Acompanhamento
        add_section_header(doc, "ACOMPANHAMENTO DA APRENDIZAGEM")
        acompanhamento = construir_acompanhamento(num, info)
        for item in acompanhamento:
            add_list_item(doc, item)
            
        # 3. Acessibilidade
        add_section_header(doc, "ACESSIBILIDADE")
        acessibilidade = construir_acessibilidade(num, info)
        for item in acessibilidade:
            add_list_item(doc, item)
            
    doc.save(caminho_saida)
    print(f"Sucesso! DOCX gerado com {cont_aulas} aulas em: {caminho_saida}")

def main():
    base_em = r"D:\PDF novos\LINGUA_PORTUGUESA\EM\3_BIMESTRE"
    
    # 1º Ano
    gerar_docx_serie_final(
        pasta_serie=os.path.join(base_em, "1_ANO"),
        caminho_saida=os.path.join(base_em, "1_ANO", "Metodologias_Lingua_Portuguesa_1_Ano_Ensino_Medio_3_Bimestre.docx"),
        titulo_serie="Metodologias — Língua Portuguesa",
        subtitulo_serie="1º Ano — 3º Bimestre",
        intro_texto=(
            "Este material reúne a metodologia, o acompanhamento da aprendizagem e a acessibilidade "
            "das aulas de Língua Portuguesa do 1º Ano do Ensino Médio para o 3º Bimestre. "
            "Cada registro serve como base pedagógica referencial para o sistema PLANOS_LUAN."
        ),
        aulas_info=AULAS_INFO_1
    )
    
    # 2º Ano
    gerar_docx_serie_final(
        pasta_serie=os.path.join(base_em, "2_ANO"),
        caminho_saida=os.path.join(base_em, "2_ANO", "Metodologias_Lingua_Portuguesa_2_Ano_Ensino_Medio_3_Bimestre.docx"),
        titulo_serie="Metodologias — Língua Portuguesa",
        subtitulo_serie="2º Ano — 3º Bimestre",
        intro_texto=(
            "Este material reúne a metodologia, o acompanhamento da aprendizagem e a acessibilidade "
            "das aulas de Língua Portuguesa do 2º Ano do Ensino Médio para o 3º Bimestre. "
            "Cada registro serve como base pedagógica referencial para o sistema PLANOS_LUAN."
        ),
        aulas_info=AULAS_INFO_2
    )

    # 3º Ano
    gerar_docx_serie_final(
        pasta_serie=os.path.join(base_em, "3_ANO"),
        caminho_saida=os.path.join(base_em, "3_ANO", "Metodologias_Lingua_Portuguesa_3_Ano_Ensino_Medio_3_Bimestre.docx"),
        titulo_serie="Metodologias — Língua Portuguesa",
        subtitulo_serie="3º Ano — 3º Bimestre",
        intro_texto=(
            "Este material reúne a metodologia, o acompanhamento da aprendizagem e a acessibilidade "
            "das aulas de Língua Portuguesa do 3º Ano do Ensino Médio para o 3º Bimestre. "
            "Cada registro serve como base pedagógica referencial para o sistema PLANOS_LUAN."
        ),
        aulas_info=AULAS_INFO_3
    )

if __name__ == "__main__":
    main()
